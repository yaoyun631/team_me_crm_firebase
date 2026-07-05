# -*- coding: utf-8 -*-
# ✅ FULL_READY_20260621：包含 /line-card-preview、/line_card_preview、/debug/routes
"""
完整主程式版本（對應最新需求）
- buyer 卡片式 + 顯示最後一筆追蹤
- seller 卡片式（寬度較窄，由 templates/sellers.html 控制）
- development 使用「目前狀況 / 下一步 / 下次時間」
- LINE Bot 回覆：已註記客需 / 已註記委託 / 已註記開發

注意：
1. 這支檔案請直接覆蓋你目前的主程式，不要再疊 patch。
2. 請搭配最新 templates/buyers.html、templates/sellers.html、templates/developments.html 使用。
"""

from flask import (
    Flask, render_template, render_template_string, request, redirect,
    url_for, flash, session, Response, Blueprint, send_file
)
import os
import sys
import json
from datetime import datetime
from zoneinfo import ZoneInfo
import csv
from io import StringIO, BytesIO
import hmac
import hashlib
import base64
import re
from uuid import uuid4
from PIL import Image
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Mm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
except Exception:
    Document = None
    Mm = Pt = None
    WD_ALIGN_PARAGRAPH = None
    WD_CELL_VERTICAL_ALIGNMENT = None

import firebase_admin
from firebase_admin import credentials, firestore, storage  
from werkzeug.security import generate_password_hash, check_password_hash
from urllib.parse import urlparse, unquote
import time
from werkzeug.utils import secure_filename

print("Working directory:", os.getcwd())

# ========= 專案路徑 + 台北時區設定 =========
# 打包成 exe 時，BASE_DIR 會指向 exe 所在資料夾；一般執行時，指向 app.py 所在資料夾。
if getattr(sys, "frozen", False):
    BASE_DIR = os.path.dirname(sys.executable)
else:
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))

TAIPEI_TZ = ZoneInfo("Asia/Taipei")

def now_taipei():
    """統一取得台北時區時間，避免 Render / 伺服器使用 UTC 導致時間差 8 小時。"""
    return datetime.now(TAIPEI_TZ)


# ========= Firestore + Storage 初始化（Render + 本機皆可用） =========
def init_firebase():
    """
    初始化 Firebase：
    - Render / 雲端：優先從環境變數 FIREBASE_CREDENTIALS 讀取
    - 本機 / 打包 exe：自動尋找 serviceAccountKey.json
    - 備援：使用你指定的完整憑證路徑
    """
    # 如果已初始化過，就直接回傳 Firestore client
    if firebase_admin._apps:
        return firestore.client()

    cred = None

    # 1️⃣ Render / 伺服器：從環境變數讀 FIREBASE_CREDENTIALS
    cred_json = os.environ.get("FIREBASE_CREDENTIALS")
    if cred_json:
        try:
            cred_dict = json.loads(cred_json)
            cred = credentials.Certificate(cred_dict)
            print("✅ 使用 FIREBASE_CREDENTIALS 初始化 Firebase")
        except Exception as e:
            print("⚠️ 解析 FIREBASE_CREDENTIALS 失敗：", e)

    # 2️⃣ 本機 / exe：依序尋找憑證
    credential_paths = [
        os.path.join(BASE_DIR, "serviceAccountKey.json"),
        os.path.join(os.getcwd(), "serviceAccountKey.json"),
        r"C:\Users\ellen\Desktop\00_Workspace\08_程式設計\team_me_crm_firebase\serviceAccountKey.json",
    ]

    if not cred:
        for key_path in credential_paths:
            if key_path and os.path.exists(key_path):
                cred = credentials.Certificate(key_path)
                print(f"✅ 使用 Firebase 憑證初始化：{key_path}")
                break

    if not cred:
        checked = "\n".join(f"- {p}" for p in credential_paths)
        raise RuntimeError(
            "找不到 Firebase 憑證：請確認 serviceAccountKey.json 是否存在。\n"
            f"目前工作目錄：{os.getcwd()}\n"
            f"程式資料夾：{BASE_DIR}\n"
            f"已檢查路徑：\n{checked}\n"
            "或請在 Render 設定 FIREBASE_CREDENTIALS 環境變數。"
        )

    # ⭐ 這裡同時指定 Storage bucket
    firebase_admin.initialize_app(cred, {
        "storageBucket": "team-me-98acf.firebasestorage.app"
    })

    return firestore.client()


# 全域 Firestore client
db = init_firebase()

# ========= 圖片上傳相關設定 =========
ALLOWED_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "gif", "webp"}

def allowed_image(filename: str) -> bool:
    if not filename:
        return False
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


# 後端統一壓縮設定
MAX_WIDTH = 1600
MAX_HEIGHT = 1600
MAX_BYTES = 800 * 1024   # 800 KB 以內（約 5 張 ~ 4MB 左右）
QUALITY_STEPS = [80, 70, 60, 50]   # 逐步降品質直到符合大小


def upload_image_to_storage(file, folder: str, object_id: str):
    """
    前端已壓縮一次，後端再保險壓縮一次，避免超肥圖炸 Render / Storage。
    回傳公開網址，失敗回傳 None。
    """
    if not file or not file.filename:
        return None

    try:
        # 讀入圖片
        img = Image.open(file.stream)

        # 統一轉成 RGB，避免 PNG / HEIC 等模式出問題
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        # 等比縮圖：不超過 1600x1600
        img.thumbnail((MAX_WIDTH, MAX_HEIGHT))

        best_buf = None
        best_size = None

        # 依序嘗試不同品質，挑一個 <= MAX_BYTES 或最後一個
        for q in QUALITY_STEPS:
            buf = BytesIO()
            img.save(buf, format="JPEG", quality=q, optimize=True)
            size = buf.tell()

            # 先記住（以防全部都 > MAX_BYTES，也至少有一個）
            if best_buf is None or size < best_size:
                best_buf = buf
                best_size = size

            if size <= MAX_BYTES:
                break

        # 使用選到的 buffer
        best_buf.seek(0)

        # 產生檔名
        base_name = secure_filename(file.filename.rsplit(".", 1)[0]) or "image"
        filename = f"{base_name}_{int(time.time())}.jpg"
        blob_path = f"{folder}/{object_id}_{filename}"

        bucket = storage.bucket()
        blob = bucket.blob(blob_path)

        # 上傳
        blob.upload_from_file(best_buf, content_type="image/jpeg")
        # 如果你原本有用 make_public，就保留；沒有就用 token 方式
        blob.make_public()
        return blob.public_url

    except Exception as e:
        print("❌ 上傳圖片至 Storage 時發生錯誤：", e)
        return None


def delete_image_from_storage(folder: str, object_id: str):
    """
    刪除 Firebase Storage 裡這個 buyer/seller 的圖片
    - folder: "buyers" / "sellers"
    - object_id: Firestore 文件 id
    這裡會嘗試所有常見副檔名，有存在就刪掉。
    """
    try:
        bucket = storage.bucket()
        for ext in ALLOWED_IMAGE_EXTENSIONS:
            blob_path = f"{folder}/{object_id}.{ext}"
            blob = bucket.blob(blob_path)
            # 避免亂砍，先檢查有沒有存在
            if blob.exists():
                blob.delete()
                print("🗑️ 已刪除圖片：", blob_path)
    except Exception as e:
        print("⚠️ 刪除圖片時發生錯誤：", e)
        



# ========= Flask 基本設定 =========
app = Flask(__name__)
app.secret_key = os.environ.get("APP_SECRET_KEY", "team_me_super_secret")

# BASE_DIR 已在檔案前段依本機 / exe 模式設定

# blog Blueprint 改成可選；如果專案裡沒有 blog.py，也能正常啟動
try:
    from blog import blog_bp
    app.register_blueprint(blog_bp)
    print("✅ blog blueprint 已載入")
except ModuleNotFoundError:
    print("ℹ️ 找不到 blog.py，略過 blog blueprint 載入")
except Exception as e:
    print(f"⚠️ blog blueprint 載入失敗：{e}")

# 限制單一請求最大 5MB（可依需求調整）
app.config["MAX_CONTENT_LENGTH"] = 5 * 1024 * 1024


# ========= LINE Webhook / 分類設定 =========
LINE_CHANNEL_SECRET = os.environ.get("LINE_CHANNEL_SECRET", "").strip()
LINE_CHANNEL_ACCESS_TOKEN = os.environ.get("LINE_CHANNEL_ACCESS_TOKEN", "").strip()

DEFAULT_LABEL_OPTIONS = [
    "開發紀錄",
    "插街",
    "待盤點客戶",
    "售-客戶需求",
    "租-客戶需求",
    "委託",
    "買賣委託",
    "出租委託",
    "廣告",
    "影片待剪/排程",
    "影片上架",
    "LINE紀錄",
    "群組回覆註記",
]


# ========= 小工具 =========
def doc_to_dict(doc):
    """把 Firestore Document 轉成 dict 並加上 id 欄位"""
    data = doc.to_dict()
    data["id"] = doc.id
    return data


def delete_by_field(collection_name, field_name, field_value):
    """
    把 collection_name 中 field_name == field_value 的文件全部刪掉
    用來刪掉某個客戶底下所有追蹤紀錄
    """
    ref = db.collection(collection_name).where(field_name, "==", field_value)
    docs = list(ref.stream())
    for d in docs:
        d.reference.delete()



def ensure_list(value):
    if not value:
        return []
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    return [str(value).strip()] if str(value).strip() else []


def parse_label_csv(text_value: str):
    if not text_value:
        return []
    parts = re.split(r"[，,、\n]+", text_value)
    return [p.strip() for p in parts if p.strip()]


def dedupe_keep_order(items):
    seen = set()
    result = []
    for item in items:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result


def normalize_phone(phone: str) -> str:
    return re.sub(r"\D", "", phone or "")


def get_request_labels(form):
    labels = []
    labels.extend(form.getlist("labels"))
    labels.extend(parse_label_csv(form.get("labels_csv", "").strip()))
    return dedupe_keep_order([x.strip() for x in labels if x and str(x).strip()])


def get_label_options_from_docs(docs):
    label_set = set(DEFAULT_LABEL_OPTIONS)
    for item in docs:
        for label in ensure_list(item.get("labels")):
            if label:
                label_set.add(label)
    return sorted(label_set)


def append_note_block(old_note: str, content: str, source_label: str = "LINE"):
    stamp = now_taipei().strftime("%Y-%m-%d %H:%M")
    block = f"[{stamp}][{source_label}] {content}".strip()
    if old_note and old_note.strip():
        return old_note.rstrip() + "\n\n" + block
    return block


def verify_line_signature(raw_body: bytes, signature: str) -> bool:
    if not LINE_CHANNEL_SECRET or not signature:
        return False
    digest = hmac.new(
        LINE_CHANNEL_SECRET.encode("utf-8"),
        raw_body,
        hashlib.sha256
    ).digest()
    expected = base64.b64encode(digest).decode("utf-8")
    return hmac.compare_digest(expected, signature)


def line_api_headers():
    return {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {LINE_CHANNEL_ACCESS_TOKEN}",
    }


def reply_line_text(reply_token: str, text_message: str):
    if not LINE_CHANNEL_ACCESS_TOKEN or not reply_token:
        return

    payload = {
        "replyToken": reply_token,
        "messages": [
            {"type": "text", "text": text_message[:5000]}
        ],
    }

    try:
        import requests
        res = requests.post(
            "https://api.line.me/v2/bot/message/reply",
            headers=line_api_headers(),
            json=payload,
            timeout=8,
        )
        print("LINE reply status:", res.status_code, res.text[:300])
    except Exception as e:
        print("⚠️ LINE reply 發生錯誤：", e)



def normalize_line_key(key: str):
    k = (key or "").strip().replace(" ", "")
    mapping = {
        "對象": "target_type",
        "類型": "target_type",
        "目標": "target_type",
        "客戶類型": "target_type",
        "電話": "phone",
        "手機": "phone",
        "電話號碼": "phone",
        "姓名": "name",
        "客戶": "name",
        "買方": "name",
        "賣方": "name",
        "客戶來源": "source",
        "來源": "source",
        "需求類型": "intent_type_raw",
        "委託類型": "deal_type_raw",
        "ID": "record_id",
        "客戶ID": "record_id",
        "buyer_id": "record_id",
        "seller_id": "record_id",
        "內容": "content",
        "紀錄": "content",
        "備註": "content",
        "說明": "content",
        "進度內容": "content",
        "進程": "stage",
        "階段": "stage",
        "狀態": "stage",
        "標籤": "labels",
        "分類": "labels",
        "labels": "labels",
        "下一步": "next_action",
        "下次行動": "next_action",
        "next_action": "next_action",
        "下次聯絡日": "next_contact_date",
        "下次聯絡日期": "next_contact_date",
        "next_contact_date": "next_contact_date",
        "地址": "address",
        "物件": "address",
        "案名": "address",
        "總價": "price",
        "價格": "price",
        "預算": "budget",
        "區域": "preferred_areas",
        "產品類型": "property_type",
        "房數": "room_range",
        "車位": "car_need",
        "開價": "expected_price",
        "底價": "min_price",
        "委託到期日": "contract_end_date",
        "筆數": "limit",
    }
    return mapping.get(k, k)


def normalize_target_type(value: str) -> str:
    v = (value or "").strip()
    if v in ("買方", "客需", "buyer", "buyers"):
        return "buyer"
    if v in ("賣方", "委託", "seller", "sellers", "屋主"):
        return "seller"
    return ""


def normalize_intent_type(value: str, fields=None) -> str:
    fields = fields or {}
    raw = (value or "").strip().lower()
    if raw in ("租", "租屋", "出租", "承租", "rent"):
        return "rent"
    if raw in ("買", "買房", "買賣", "購屋", "buy", "sale", "售屋"):
        return "buy"

    budget = str(fields.get("budget", "") or "")
    price = str(fields.get("price", "") or "")
    if "萬" in budget or "萬" in price:
        return "buy"
    digits = re.sub(r"\D", "", budget or price)
    if digits:
        try:
            num = int(digits)
            if num >= 1000000:
                return "buy"
            if num <= 100000:
                return "rent"
        except Exception:
            pass
    return ""


def normalize_deal_type(value: str) -> str:
    raw = (value or "").strip().lower()
    if raw in ("租", "出租", "委租", "rent"):
        return "rent"
    if raw in ("賣", "出售", "賣屋", "買賣", "委售", "sale", "sell"):
        return "sale"
    return ""


def parse_int_limit(value, default=10, max_value=30):
    try:
        num = int(str(value).strip())
        return max(1, min(max_value, num))
    except Exception:
        return default


def get_line_sender_display_name(event):
    source = event.get("source") or {}
    user_id = source.get("userId", "")
    if not user_id or not LINE_CHANNEL_ACCESS_TOKEN:
        return ""

    cache_key = f"{source.get('groupId', '')}:{source.get('roomId', '')}:{user_id}"
    cache = app.config.setdefault("LINE_PROFILE_CACHE", {})
    if cache_key in cache:
        return cache[cache_key]

    headers = line_api_headers()
    url = ""
    if source.get("groupId"):
        url = f"https://api.line.me/v2/bot/group/{source['groupId']}/member/{user_id}"
    elif source.get("roomId"):
        url = f"https://api.line.me/v2/bot/room/{source['roomId']}/member/{user_id}"
    else:
        url = f"https://api.line.me/v2/bot/profile/{user_id}"

    try:
        import requests
        res = requests.get(url, headers=headers, timeout=8)
        if res.status_code == 200:
            data = res.json() or {}
            display_name = data.get("displayName", "") or ""
            cache[cache_key] = display_name
            return display_name
    except Exception as e:
        print("⚠️ 取得 LINE 使用者名稱失敗：", e)

    return ""


def reply_line_text(reply_token: str, text_message: str):
    if not LINE_CHANNEL_ACCESS_TOKEN or not reply_token:
        return None

    payload = {
        "replyToken": reply_token,
        "messages": [
            {"type": "text", "text": (text_message or "")[:5000]}
        ],
    }

    try:
        import requests
        res = requests.post(
            "https://api.line.me/v2/bot/message/reply",
            headers=line_api_headers(),
            json=payload,
            timeout=8,
        )
        print("LINE reply status:", res.status_code, res.text[:300])
        data = {}
        try:
            data = res.json() if res.text else {}
        except Exception:
            data = {}
        return {
            "status_code": res.status_code,
            "data": data,
            "sent_messages": (data or {}).get("sentMessages", []),
        }
    except Exception as e:
        print("⚠️ LINE reply 發生錯誤：", e)
        return None


def save_line_message_link(message_id: str, target_type: str, target_id: str, tag="", action="", customer_name="", phone="", source_event=None):
    if not message_id or not target_type or not target_id:
        return
    data = {
        "message_id": message_id,
        "target_type": target_type,
        "target_id": target_id,
        "tag": tag,
        "action": action,
        "customer_name": customer_name,
        "phone": phone,
        "created_at": now_taipei().isoformat(),
    }
    if source_event:
        source = source_event.get("source", {})
        data["line_group_id"] = source.get("groupId", "")
        data["line_room_id"] = source.get("roomId", "")
        data["line_user_id"] = source.get("userId", "")
        data["sender_display_name"] = get_line_sender_display_name(source_event)
    db.collection("line_message_links").document(str(message_id)).set(data, merge=True)


def get_line_message_link(message_id: str):
    if not message_id:
        return None
    doc = db.collection("line_message_links").document(str(message_id)).get()
    if doc.exists:
        return doc.to_dict() or {}
    return None


def build_line_operator_label(event):
    sender = get_line_sender_display_name(event) or "未知成員"
    return f"LINE/{sender}"


def build_line_summary(content: str, event):
    sender = get_line_sender_display_name(event) or "未知成員"
    text = (content or "").strip()
    return f"{sender}：{text}" if text else f"{sender}：LINE 更新"


def parse_line_formatted_message(text: str):
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return None

    first = lines[0]
    if not first.startswith("#"):
        return None

    tag = first.lstrip("#").strip()
    tag_map = {
        "新增客需": "create_buyer_need",
        "新增委託": "create_seller_listing",
        "買方追蹤": "buyer_followup",
        "賣方追蹤": "seller_followup",
        "客戶分類": "classify",
        "查詢紀錄": "query_records",
        "查詢委託到期": "query_contract_end",
        "帶看": "buyer_followup",
        "成交": "buyer_followup",
        "委託": "seller_followup",
        "紀錄": "generic_note",
    }
    action = tag_map.get(tag)
    if not action:
        return None

    fields = {}
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = m.group(2).strip()
        if key == "labels":
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value

    if tag in ("買方追蹤", "帶看", "成交") and not fields.get("target_type"):
        fields["target_type"] = "buyer"
    if tag in ("賣方追蹤", "委託") and not fields.get("target_type"):
        fields["target_type"] = "seller"

    fields["target_type"] = normalize_target_type(fields.get("target_type", "")) or fields.get("target_type", "")
    fields["intent_type"] = normalize_intent_type(fields.get("intent_type_raw", ""), fields)
    fields["deal_type"] = normalize_deal_type(fields.get("deal_type_raw", ""))
    fields["limit"] = parse_int_limit(fields.get("limit", 10), default=10, max_value=30)

    if action == "create_buyer_need":
        if not (fields.get("name") and fields.get("phone") and fields.get("source")):
            return None
    elif action == "create_seller_listing":
        if not (fields.get("name") and fields.get("phone") and fields.get("source")):
            return None
    elif action in ("buyer_followup", "seller_followup", "classify", "query_records", "query_contract_end"):
        if not (fields.get("record_id") or fields.get("phone") or fields.get("name")):
            return None

    return {
        "tag": tag,
        "action": action,
        "fields": fields,
        "raw_text": text,
    }


def find_records_by_phone(collection_name: str, phone: str):
    normalized_phone = normalize_phone(phone)
    if not normalized_phone:
        return []
    matches = []
    for doc in db.collection(collection_name).stream():
        data = doc.to_dict() or {}
        if normalize_phone(data.get("phone", "")) == normalized_phone:
            matches.append(doc)
    return matches


def find_customer_record(target_type: str, record_id: str = "", phone: str = "", name: str = ""):
    collection_name = "buyers" if target_type == "buyer" else "sellers"

    if record_id:
        doc = db.collection(collection_name).document(record_id).get()
        if doc.exists:
            return doc

    if phone:
        matches = find_records_by_phone(collection_name, phone)
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            return None

    if name:
        docs = list(db.collection(collection_name).where("name", "==", name.strip()).limit(2).stream())
        if len(docs) == 1:
            return docs[0]

    return None


def resolve_customer_record(fields, preferred_target_type=""):
    record_id = fields.get("record_id", "")
    phone = fields.get("phone", "")
    name = fields.get("name", "")
    target_type = preferred_target_type or normalize_target_type(fields.get("target_type", ""))

    if target_type in ("buyer", "seller"):
        doc = find_customer_record(target_type, record_id, phone, name)
        return target_type, doc

    buyer_doc = find_customer_record("buyer", record_id, phone, name)
    seller_doc = find_customer_record("seller", record_id, phone, name)

    if buyer_doc and not seller_doc:
        return "buyer", buyer_doc
    if seller_doc and not buyer_doc:
        return "seller", seller_doc
    return "", None


def update_customer_note_and_labels(target_type: str, doc_ref, content: str, labels=None, stage="", source="LINE", event=None):
    labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(labels))
    snapshot = doc_ref.get()
    current = snapshot.to_dict() or {}
    old_note = current.get("note", "")

    source_label = build_line_operator_label(event) if event else "LINE"
    updates = {
        "note": append_note_block(old_note, content, source_label),
        "labels": firestore.ArrayUnion(labels),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }
    if stage:
        updates["stage"] = stage
    if source:
        updates["source"] = source

    doc_ref.update(updates)


def add_customer_followup(target_type: str, customer_id: str, content: str, next_action="", next_contact_date="", labels=None, line_event=None):
    collection_name = "buyer_followups" if target_type == "buyer" else "seller_followups"
    key_name = "buyer_id" if target_type == "buyer" else "seller_id"

    sender_display_name = get_line_sender_display_name(line_event) if line_event else ""
    data = {
        key_name: customer_id,
        "contact_time": now_taipei().strftime("%Y-%m-%d %H:%M"),
        "channel": "LINE",
        "content": content,
        "next_action": next_action,
        "next_contact_date": next_contact_date,
        "labels": dedupe_keep_order(["LINE紀錄"] + ensure_list(labels)),
        "created_at": now_taipei().isoformat(),
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "sender_display_name": sender_display_name,
    }

    if line_event:
        source = line_event.get("source", {})
        data["line_group_id"] = source.get("groupId", "")
        data["line_room_id"] = source.get("roomId", "")
        data["line_user_id"] = source.get("userId", "")
        data["line_message_id"] = (line_event.get("message") or {}).get("id", "")
        data["quoted_message_id"] = (line_event.get("message") or {}).get("quotedMessageId", "")

    db.collection(collection_name).add(data)


def save_line_log(parsed, event, status, target_type="", target_id="", note="", sender_display_name=""):
    source = event.get("source", {})
    msg = event.get("message") or {}
    db.collection("line_logs").add({
        "tag": parsed.get("tag", ""),
        "action": parsed.get("action", ""),
        "fields": parsed.get("fields", {}),
        "raw_text": parsed.get("raw_text", ""),
        "status": status,
        "target_type": target_type,
        "target_id": target_id,
        "note": note,
        "sender_display_name": sender_display_name,
        "line_group_id": source.get("groupId", ""),
        "line_room_id": source.get("roomId", ""),
        "line_user_id": source.get("userId", ""),
        "message_id": msg.get("id", ""),
        "quoted_message_id": msg.get("quotedMessageId", ""),
        "webhook_event_id": event.get("webhookEventId", ""),
        "created_at": now_taipei().isoformat(),
    })


def build_buyer_labels(intent_type: str, extra_labels=None):
    base = ["LINE紀錄", "租-客戶需求" if intent_type == "rent" else "售-客戶需求"]
    return dedupe_keep_order(base + ensure_list(extra_labels))


def build_seller_labels(deal_type: str, extra_labels=None):
    base = ["LINE紀錄", "委託", "出租委託" if deal_type == "rent" else "買賣委託"]
    return dedupe_keep_order(base + ensure_list(extra_labels))


def create_buyer_need(fields, event):
    sender_name = get_line_sender_display_name(event) or "未知成員"
    phone = fields.get("phone", "").strip()
    matches = find_records_by_phone("buyers", phone)

    intent_type = normalize_intent_type(fields.get("intent_type_raw", "") or fields.get("intent_type", ""), fields)
    if not intent_type:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：#新增客需 請填 需求類型: 租 或 買賣",
        }

    labels = build_buyer_labels(intent_type, fields.get("labels"))
    budget = fields.get("budget", "").strip()

    payload = {
        "name": fields.get("name", "").strip(),
        "phone": phone,
        "source": fields.get("source", "").strip(),
        "intent_type": intent_type,
        "stage": fields.get("stage", "").strip() or "接觸",
        "preferred_areas": fields.get("preferred_areas", "").strip(),
        "property_type": fields.get("property_type", "").strip(),
        "room_range": fields.get("room_range", "").strip(),
        "car_need": fields.get("car_need", "").strip(),
        "labels": labels,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }

    if intent_type == "rent":
        payload["rent_max"] = fields.get("rent_max", "").strip() or budget
    else:
        payload["budget_max"] = fields.get("budget_max", "").strip() or budget

    extra_content = fields.get("content", "").strip()
    note_content = extra_content or "LINE 新增客需"
    note_content = build_line_summary(note_content, event)

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection("buyers").document(doc.id)
        update_customer_note_and_labels(
            target_type="buyer",
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=payload["stage"],
            source=payload["source"],
            event=event,
        )
        doc_ref.update(payload)
        add_customer_followup(
            target_type="buyer",
            customer_id=doc.id,
            content=note_content,
            labels=labels,
            line_event=event,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已註記客需：{updated_doc.get('name', '')}",
            "target_type": "buyer",
            "target_id": doc.id,
            "customer_name": updated_doc.get("name", ""),
            "phone": updated_doc.get("phone", ""),
            "parsed_tag": "新增客需",
        }

    if len(matches) > 1:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：同電話有多位買方，請先到後台整理或改用客戶ID",
        }

    now = now_taipei().isoformat()
    payload.update({
        "created_at": now,
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "note": append_note_block("", note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection("buyers").document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type="buyer",
        customer_id=doc_ref.id,
        content=note_content,
        labels=labels,
        line_event=event,
    )
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記客需：{payload['name']}（{'租' if intent_type == 'rent' else '買賣'}）",
        "target_type": "buyer",
        "target_id": doc_ref.id,
        "customer_name": payload["name"],
        "phone": payload["phone"],
        "parsed_tag": "新增客需",
    }


def create_seller_listing(fields, event):
    phone = fields.get("phone", "").strip()
    matches = find_records_by_phone("sellers", phone)

    deal_type = normalize_deal_type(fields.get("deal_type_raw", "") or fields.get("deal_type", ""))
    if not deal_type:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：#新增委託 請填 委託類型: 賣 或 出租",
        }

    labels = build_seller_labels(deal_type, fields.get("labels"))
    note_content = build_line_summary(fields.get("content", "").strip() or fields.get("address", "").strip() or "LINE 新增委託", event)

    payload = {
        "name": fields.get("name", "").strip(),
        "phone": phone,
        "source": fields.get("source", "").strip(),
        "address": fields.get("address", "").strip(),
        "property_type": fields.get("property_type", "").strip(),
        "stage": fields.get("stage", "").strip() or "委託中",
        "deal_type": deal_type,
        "expected_price": fields.get("expected_price", "").strip() or fields.get("price", "").strip(),
        "min_price": fields.get("min_price", "").strip(),
        "contract_end_date": fields.get("contract_end_date", "").strip(),
        "labels": labels,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection("sellers").document(doc.id)
        update_customer_note_and_labels(
            target_type="seller",
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=payload["stage"],
            source=payload["source"],
            event=event,
        )
        doc_ref.update(payload)
        add_customer_followup(
            target_type="seller",
            customer_id=doc.id,
            content=note_content,
            labels=labels,
            line_event=event,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已註記委託：{updated_doc.get('name', '')}",
            "target_type": "seller",
            "target_id": doc.id,
            "customer_name": updated_doc.get("name", ""),
            "phone": updated_doc.get("phone", ""),
            "parsed_tag": "新增委託",
        }

    if len(matches) > 1:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：同電話有多位委託，請先到後台整理或改用客戶ID",
        }

    now = now_taipei().isoformat()
    payload.update({
        "created_at": now,
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "note": append_note_block("", note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection("sellers").document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type="seller",
        customer_id=doc_ref.id,
        content=note_content,
        labels=labels,
        line_event=event,
    )
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記委託：{payload['name']}（{'出租' if deal_type == 'rent' else '買賣'}）",
        "target_type": "seller",
        "target_id": doc_ref.id,
        "customer_name": payload["name"],
        "phone": payload["phone"],
        "parsed_tag": "新增委託",
    }


def format_record_timeline(target_type: str, doc_snapshot, limit=10):
    data = doc_snapshot.to_dict() or {}
    record_id = doc_snapshot.id

    followup_collection = "buyer_followups" if target_type == "buyer" else "seller_followups"
    key_name = "buyer_id" if target_type == "buyer" else "seller_id"

    followups = []
    for d in db.collection(followup_collection).where(key_name, "==", record_id).stream():
        item = d.to_dict() or {}
        followups.append({
            "time": item.get("contact_time") or item.get("created_at") or "",
            "channel": item.get("channel", "LINE"),
            "text": (item.get("content") or "").strip(),
            "created_by_name": item.get("created_by_name", "") or "",
            "sender_display_name": item.get("sender_display_name", "") or "",
        })

    followups = [x for x in followups if x.get("text")]
    followups.sort(key=lambda x: x.get("time", ""), reverse=True)

    lines = ["客戶資訊"]

    if target_type == "buyer":
        intent_map = {
            "rent": "租屋",
            "buy": "買賣",
            "both": "租買皆可",
        }
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {data.get('source', '') or '-'}",
            f"需求類型: {intent_map.get(data.get('intent_type', ''), data.get('intent_type', '') or '-')}",
            f"預算: {data.get('budget_min', '') or data.get('rent_min', '') or '-'} ~ {data.get('budget_max', '') or data.get('rent_max', '') or '-'}",
            f"偏好區域: {data.get('preferred_areas', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"房數需求: {data.get('room_range', '') or '-'}",
            f"車位需求: {data.get('car_need', '') or '-'}",
        ])
    else:
        deal_map = {
            "sale": "買賣",
            "rent": "出租",
        }
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {data.get('source', '') or '-'}",
            f"委託類型: {deal_map.get(data.get('deal_type', ''), data.get('deal_type', '') or '-')}",
            f"地址: {data.get('address', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"開價: {data.get('expected_price', '') or '-'}",
            f"底價: {data.get('min_price', '') or '-'}",
            f"委託到期日: {data.get('contract_end_date', '') or '-'}",
        ])

    lines.append("")
    lines.append("追蹤進度")

    if not followups:
        lines.append("目前沒有追蹤紀錄")
    else:
        for item in followups[:limit]:
            header_parts = [item.get('time', ''), item.get('channel', 'LINE')]
            creator = (item.get("created_by_name") or "").strip()
            sender = (item.get("sender_display_name") or "").strip()
            if creator:
                header_parts.append(f"KEYIN: {creator}")
            if sender:
                header_parts.append(f"留言者: {sender}")
            lines.append("｜".join([x for x in header_parts if x]))
            lines.append(item.get("text", ""))
            lines.append("")

    output = "\n".join(lines).strip()
    return output[:4500]


def query_contract_end_text(fields):
    target_type, doc = resolve_customer_record(fields, preferred_target_type="seller")
    if target_type != "seller" or not doc:
        return False, "查無唯一委託資料，請提供正確電話或客戶ID", None

    data = doc.to_dict() or {}
    deal_type = data.get("deal_type", "")
    deal_label = "出租" if deal_type == "rent" else "買賣" if deal_type == "sale" else "未設定"
    text = "\n".join([
        "查詢委託到期",
        f"姓名: {data.get('name', '')}",
        f"電話: {data.get('phone', '')}",
        f"地址: {data.get('address', '')}",
        f"委託類型: {deal_label}",
        f"委託到期日: {data.get('contract_end_date', '') or '未填寫'}",
        f"目前進程: {data.get('stage', '') or '未填寫'}",
    ]).strip()
    return True, text[:4500], {
        "target_type": "seller",
        "target_id": doc.id,
        "customer_name": data.get("name", ""),
        "phone": data.get("phone", ""),
        "parsed_tag": "查詢委託到期",
    }


def process_quote_context_message(event):
    message = event.get("message") or {}
    quoted_message_id = message.get("quotedMessageId", "")
    raw_text = (message.get("text") or "").strip()
    if not quoted_message_id or not raw_text:
        return {"handled": False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {"handled": False}

    target_type = link.get("target_type", "")
    target_id = link.get("target_id", "")
    if target_type not in ("buyer", "seller") or not target_id:
        return {"handled": False}

    collection_name = "buyers" if target_type == "buyer" else "sellers"
    doc_ref = db.collection(collection_name).document(target_id)
    doc = doc_ref.get()
    if not doc.exists:
        return {"handled": True, "ok": False, "reply_text": "未寫入：引用的客戶資料不存在"}

    labels = dedupe_keep_order(["LINE紀錄", "群組回覆註記"])
    # 引用/回覆型註記只記錄「新的回覆內容」；
    # 發話者姓名會由 note 的 source_label 與 followup 的 sender_display_name 另行保留。
    reply_only_text = raw_text

    update_customer_note_and_labels(
        target_type=target_type,
        doc_ref=doc_ref,
        content=reply_only_text,
        labels=labels,
        source="LINE",
        event=event,
    )
    add_customer_followup(
        target_type=target_type,
        customer_id=target_id,
        content=reply_only_text,
        labels=labels,
        line_event=event,
    )

    parsed = {
        "tag": "群組回覆註記",
        "action": "quoted_context_note",
        "fields": {"quoted_message_id": quoted_message_id},
        "raw_text": raw_text,
    }
    save_line_log(parsed, event, "success", target_type=target_type, target_id=target_id, sender_display_name=get_line_sender_display_name(event))
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記到{'客需' if target_type == 'buyer' else '委託'}：{(doc.to_dict() or {}).get('name', '')}",
        "target_type": target_type,
        "target_id": target_id,
        "customer_name": (doc.to_dict() or {}).get("name", ""),
        "phone": (doc.to_dict() or {}).get("phone", ""),
        "parsed_tag": "群組回覆註記",
    }


def process_line_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    sender_display_name = get_line_sender_display_name(event)

    parsed = parse_line_formatted_message(message.get("text", ""))
    if not parsed:
        quoted_result = process_quote_context_message(event)
        if quoted_result.get("handled"):
            return quoted_result
        return {"handled": False}

    fields = parsed["fields"]
    action = parsed["action"]

    if action == "create_buyer_need":
        result = create_buyer_need(fields, event)
    elif action == "create_seller_listing":
        result = create_seller_listing(fields, event)
    elif action == "query_records":
        target_type, doc = resolve_customer_record(fields)
        if not doc:
            result = {"handled": True, "ok": False, "reply_text": "查無唯一客戶，請補電話或客戶ID"}
        else:
            result = {
                "handled": True,
                "ok": True,
                "reply_text": format_record_timeline(target_type, doc, limit=fields.get("limit", 10)),
                "target_type": target_type,
                "target_id": doc.id,
                "customer_name": (doc.to_dict() or {}).get("name", ""),
                "phone": (doc.to_dict() or {}).get("phone", ""),
                "parsed_tag": parsed.get("tag", ""),
            }
    elif action == "query_contract_end":
        ok, text, ctx = query_contract_end_text(fields)
        result = {"handled": True, "ok": ok, "reply_text": text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get("target_type", "")
        if action == "buyer_followup":
            target_type = "buyer"
        elif action == "seller_followup":
            target_type = "seller"

        if target_type not in ("buyer", "seller"):
            result = {"handled": True, "ok": False, "reply_text": "請提供對象：買方 或 賣方"}
        else:
            doc = find_customer_record(
                target_type=target_type,
                record_id=fields.get("record_id", ""),
                phone=fields.get("phone", ""),
                name=fields.get("name", ""),
            )
            if not doc:
                result = {"handled": True, "ok": False, "reply_text": "找不到唯一客戶，請補客戶ID或正確電話"}
            else:
                doc_ref = db.collection("buyers" if target_type == "buyer" else "sellers").document(doc.id)
                labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(fields.get("labels")))
                summary_parts = []
                if fields.get("content"):
                    summary_parts.append(fields["content"])
                if fields.get("address"):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get("price"):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary("；".join(summary_parts).strip() or "LINE 更新", event)

                update_customer_note_and_labels(
                    target_type=target_type,
                    doc_ref=doc_ref,
                    content=summary_text,
                    labels=labels,
                    stage=fields.get("stage", ""),
                    source=fields.get("source", "LINE"),
                    event=event,
                )

                if action in ("buyer_followup", "seller_followup", "classify"):
                    add_customer_followup(
                        target_type=target_type,
                        customer_id=doc.id,
                        content=summary_text,
                        next_action=fields.get("next_action", ""),
                        next_contact_date=fields.get("next_contact_date", ""),
                        labels=labels,
                        line_event=event,
                    )

                current_data = doc_ref.get().to_dict() or {}
                result = {
                    "handled": True,
                    "ok": True,
                    "reply_text": f"已寫入{'買方' if target_type == 'buyer' else '賣方'}：{current_data.get('name', '')}",
                    "target_type": target_type,
                    "target_id": doc.id,
                    "customer_name": current_data.get("name", ""),
                    "phone": current_data.get("phone", ""),
                    "parsed_tag": parsed.get("tag", ""),
                }

    # 寫 line_logs
    save_line_log(
        parsed,
        event,
        "success" if result.get("ok") else "failed",
        target_type=result.get("target_type", ""),
        target_id=result.get("target_id", ""),
        note=result.get("reply_text", ""),
        sender_display_name=sender_display_name,
    )

    # 將使用者原始訊息與客戶關聯起來，方便之後引用回覆
    if result.get("ok") and result.get("target_type") and result.get("target_id"):
        incoming_message_id = message.get("id", "")
        save_line_message_link(
            incoming_message_id,
            result["target_type"],
            result["target_id"],
            tag=result.get("parsed_tag", ""),
            action=action,
            customer_name=result.get("customer_name", ""),
            phone=result.get("phone", ""),
            source_event=event,
        )

    return result


def build_label_options(*doc_lists):
    label_set = set(DEFAULT_LABEL_OPTIONS)
    for docs in doc_lists:
        for item in docs:
            for label in ensure_list(item.get("labels")):
                label_set.add(label)
    return sorted(label_set)


BUYER_STAGE_OPTIONS = [
    "待聯繫",
    "已聯繫",
    "帶看中",
    "持續追蹤",
    "成交",
    "無效",
]

SELLER_STAGE_OPTIONS = [
    "待聯繫",
    "已聯繫",
    "持續追蹤",
    "委託中",
    "成交",
    "無效",
]

def attach_latest_followup(records, followup_collection, key_name):
    followups = [doc_to_dict(d) for d in db.collection(followup_collection).stream()]
    latest_map = {}
    for f in followups:
        rid = f.get(key_name)
        if not rid:
            continue
        ts = f.get("contact_time") or f.get("created_at") or ""
        prev = latest_map.get(rid)
        prev_ts = (prev or {}).get("contact_time") or (prev or {}).get("created_at") or ""
        if (not prev) or ts > prev_ts:
            latest_map[rid] = f
    for item in records:
        item["latest_followup"] = latest_map.get(item.get("id"))
    return records







# ========= 登入保護 =========
def login_required(view_func):
    from functools import wraps

    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if "user_id" not in session:
            flash("請先登入", "warning")
            return redirect(url_for("login"))
        return view_func(*args, **kwargs)

    return wrapper


# ========= 登入 / 登出 =========
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")

        if not email or not password:
            flash("請輸入帳號與密碼", "danger")
            return redirect(url_for("login"))

        users_ref = db.collection("users").where("email", "==", email).limit(1)
        docs = list(users_ref.stream())

        if not docs:
            flash("帳號或密碼錯誤", "danger")
            return redirect(url_for("login"))

        user_doc = docs[0]
        user = user_doc.to_dict()

        if not check_password_hash(user.get("password_hash", ""), password):
            flash("帳號或密碼錯誤", "danger")
            return redirect(url_for("login"))

        session["user_id"] = user_doc.id
        session["user_name"] = user.get("name") or user.get("email")
        session["user_email"] = user.get("email")

        flash(f"歡迎回來，{session['user_name']}！", "success")
        return redirect(url_for("buyers"))

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    flash("已登出", "info")
    return redirect(url_for("login"))


# ========= 首頁 =========
@app.route("/")
def index():
    if "user_id" in session:
        return redirect(url_for("buyers"))
    return redirect(url_for("login"))


# ========= 買方列表 + 篩選 / 排序 =========
# ========= 買方列表 =========
@app.route("/buyers")
@login_required
def buyers():
    q = request.args.get("q", "").strip()
    level = request.args.get("level", "").strip()
    intent_type = request.args.get("intent_type", "").strip()
    stage = request.args.get("stage", "").strip()
    source = request.args.get("source", "").strip()
    label = request.args.get("label", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")

    docs = db.collection("buyers").stream()
    all_buyers = [doc_to_dict(d) for d in docs]

    source_set = set()
    for b in all_buyers:
        s = (b.get("source") or "").strip()
        if s:
            source_set.add(s)
    source_options = sorted(source_set)
    label_options = build_label_options(all_buyers)

    buyers_list = list(all_buyers)

    if q:
        buyers_list = [
            b for b in buyers_list
            if q in (b.get("name") or "") or q in (b.get("phone") or "")
        ]

    if level:
        buyers_list = [b for b in buyers_list if b.get("level") == level]

    if intent_type:
        buyers_list = [b for b in buyers_list if b.get("intent_type") == intent_type]

    if stage:
        buyers_list = [b for b in buyers_list if (b.get("stage") or "") == stage]

    if source:
        buyers_list = [b for b in buyers_list if (b.get("source") or "") == source]

    if label:
        buyers_list = [b for b in buyers_list if label in ensure_list(b.get("labels"))]

    def parse_created_at(b):
        return b.get("created_at") or ""

    if sort_by == "created_at_asc":
        buyers_list.sort(key=parse_created_at)
    elif sort_by == "created_at_desc":
        buyers_list.sort(key=parse_created_at, reverse=True)
    elif sort_by == "name_asc":
        buyers_list.sort(key=lambda b: (b.get("name") or ""))
    elif sort_by == "name_desc":
        buyers_list.sort(key=lambda b: (b.get("name") or ""), reverse=True)

    buyers_list = attach_latest_followup(buyers_list, "buyer_followups", "buyer_id")

    return render_template(
        "buyers.html",
        buyers=buyers_list,
        q=q,
        level=level,
        intent_type=intent_type,
        stage=stage,
        source=source,
        source_options=source_options,
        label=label,
        label_options=label_options,
        sort_by=sort_by,
        buyer_stage_options=BUYER_STAGE_OPTIONS,
        total_count=len(all_buyers),
        filtered_count=len(buyers_list),
    )


# ========= 新增買方 =========
@app.route("/buyers/new", methods=["POST"])
@login_required
def buyers_new():
    form = request.form
    file = request.files.get("photo")   # ⭐ 新增：抓圖片

    name = form.get("name", "").strip()
    phone = form.get("phone", "").strip()
    email = form.get("email", "").strip()
    line_id = form.get("line_id", "").strip()
    source = form.get("source", "").strip()
    level = form.get("level", "").strip()
    intent_type = form.get("intent_type", "").strip()
    rent_min = form.get("rent_min", "").strip()
    rent_max = form.get("rent_max", "").strip()
    budget_min = form.get("budget_min", "").strip()
    budget_max = form.get("budget_max", "").strip()
    preferred_areas = form.get("preferred_areas", "").strip()
    property_type = form.get("property_type", "").strip()
    room_range = form.get("room_range", "").strip()
    car_need = form.get("car_need", "").strip()
    job = form.get("job", "").strip()
    family_info = form.get("family_info", "").strip()
    requirement_must = form.get("requirement_must", "").strip()
    requirement_nice = form.get("requirement_nice", "").strip()
    other_background = form.get("other_background", "").strip()
    note = form.get("note", "").strip()
    labels = get_request_labels(form)

    if not name:
        flash("買方姓名必填", "danger")
        return redirect(url_for("buyers"))

    now = now_taipei().isoformat()

    # ⭐ 先建立一個空的 document，拿到 id
    doc_ref = db.collection("buyers").document()
    buyer_id = doc_ref.id

    # ⭐ 如果有上傳圖片，就丟到 Storage
    photo_url = None
    if file and file.filename:
        photo_url = upload_image_to_storage(file, folder="buyers", object_id=buyer_id)

    crm_vis = crm_record_visibility_payload_from_form(form)

    data = {
        **crm_vis,
        "name": name,
        "phone": phone,
        "email": email,
        "line_id": line_id,
        "source": source,
        "level": level,
        "intent_type": intent_type,
        "rent_min": rent_min,
        "rent_max": rent_max,
        "budget_min": budget_min,
        "budget_max": budget_max,
        "preferred_areas": preferred_areas,
        "property_type": property_type,
        "room_range": room_range,
        "car_need": car_need,
        "job": job,
        "family_info": family_info,
        "requirement_must": requirement_must,
        "requirement_nice": requirement_nice,
        "other_background": other_background,
        "note": note,
        "labels": labels,
        "created_at": now,
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
    }

    if photo_url:
        data["photo_url"] = photo_url   # ⭐ 存圖片網址

    doc_ref.set(data)

    flash("已新增買方", "success")
    return redirect(url_for("buyers"))


# ========= 買方詳細 =========
@app.route("/buyers/<buyer_id>")
@login_required
def buyer_detail(buyer_id):
    doc = db.collection("buyers").document(buyer_id).get()
    if not doc.exists:
        flash("找不到這位買方", "danger")
        return redirect(url_for("buyers"))

    buyer = doc_to_dict(doc)

    # 追蹤紀錄
    followups_ref = db.collection("buyer_followups").where("buyer_id", "==", buyer_id)
    followups = [doc_to_dict(f) for f in followups_ref.stream()]

    # 依 contact_time 排序（新到舊）
    followups.sort(key=lambda x: x.get("contact_time", ""), reverse=True)

    return render_template("buyer_detail.html", buyer=buyer, followups=followups)

@app.route("/buyers/<buyer_id>/quick-stage", methods=["POST"])
@login_required
def buyer_quick_stage(buyer_id):
    stage = (request.form.get("stage", "") or request.form.get("current_stage", "")).strip()
    if not stage:
        flash("請選擇進度", "warning")
        return redirect(request.referrer or url_for("buyers"))
    if stage not in BUYER_STAGE_OPTIONS:
        flash("買方進度不在可選清單內", "danger")
        return redirect(request.referrer or url_for("buyers"))
    updates = {
        "stage": stage,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id"),
        "updated_by_name": session.get("user_name"),
    }
    db.collection("buyers").document(buyer_id).update(updates)
    flash("已更新買方進度", "success")
    return redirect(request.referrer or url_for("buyers"))


# ========= 新增買方追蹤紀錄 =========
@app.route("/buyers/<buyer_id>/followup", methods=["POST"])
@login_required
def add_buyer_followup(buyer_id):
    contact_time = request.form.get("contact_time", "").strip()
    channel = request.form.get("channel", "").strip()
    content = request.form.get("content", "").strip()
    next_action = request.form.get("next_action", "").strip()
    next_contact_date = request.form.get("next_contact_date", "").strip()

    if not contact_time:
        contact_time = now_taipei().strftime("%Y-%m-%d %H:%M")

    now = now_taipei().isoformat()

    db.collection("buyer_followups").add(
        {
            "buyer_id": buyer_id,
            "contact_time": contact_time,
            "channel": channel,
            "content": content,
            "next_action": next_action,
            "next_contact_date": next_contact_date,
            "created_at": now,
            "created_by_id": session.get("user_id"),
            "created_by_name": session.get("user_name"),
        }
    )

    flash("已新增追蹤紀錄", "success")
    return redirect(url_for("buyer_detail", buyer_id=buyer_id))


# ========= 編輯買方 =========
@app.route("/buyers/<buyer_id>/edit", methods=["GET", "POST"])
@login_required
def buyer_edit(buyer_id):
    # 先取得該買方文件
    doc_ref = db.collection("buyers").document(buyer_id)
    doc = doc_ref.get()

    if not doc.exists:
        flash("找不到這位買方", "danger")
        return redirect(url_for("buyers"))

    buyer = doc_to_dict(doc)

    # 處理送出編輯表單（POST）
    if request.method == "POST":
        form = request.form

        # ⭐ 必填姓名檢查
        name = form.get("name", "").strip()
        if not name:
            flash("姓名為必填", "danger")
            # 更新 buyer 物件，讓表單保留剛剛輸入的東西
            buyer.update({
                "name": name,
                "phone": form.get("phone", "").strip(),
                "email": form.get("email", "").strip(),
                "line_id": form.get("line_id", "").strip(),
                "source": form.get("source", "").strip(),
                "level": form.get("level", "").strip(),
                "intent_type": form.get("intent_type", "").strip(),
                "rent_min": form.get("rent_min", "").strip(),
                "rent_max": form.get("rent_max", "").strip(),
                "budget_min": form.get("budget_min", "").strip(),
                "budget_max": form.get("budget_max", "").strip(),
                "preferred_areas": form.get("preferred_areas", "").strip(),
                "property_type": form.get("property_type", "").strip(),
                "room_range": form.get("room_range", "").strip(),
                "car_need": form.get("car_need", "").strip(),
                "job": form.get("job", "").strip(),
                "family_info": form.get("family_info", "").strip(),
                "requirement_must": form.get("requirement_must", "").strip(),
                "requirement_nice": form.get("requirement_nice", "").strip(),
                "other_background": form.get("other_background", "").strip(),
                "note": form.get("note", "").strip(),
                "stage": form.get("stage", "").strip(),
                "labels": get_request_labels(form),
            })
            return render_template("buyer_edit.html", buyer=buyer)

        # ✅ 先處理一般文字欄位
        labels = get_request_labels(form)
        crm_vis = crm_record_visibility_payload_from_form(form, buyer)

        updated = {
            **crm_vis,
            "name": name,
            "phone": form.get("phone", "").strip(),
            "email": form.get("email", "").strip(),
            "line_id": form.get("line_id", "").strip(),
            "source": form.get("source", "").strip(),
            "level": form.get("level", "").strip(),
            "intent_type": form.get("intent_type", "").strip(),
            "rent_min": form.get("rent_min", "").strip(),
            "rent_max": form.get("rent_max", "").strip(),
            "budget_min": form.get("budget_min", "").strip(),
            "budget_max": form.get("budget_max", "").strip(),
            "preferred_areas": form.get("preferred_areas", "").strip(),
            "property_type": form.get("property_type", "").strip(),
            "room_range": form.get("room_range", "").strip(),
            "car_need": form.get("car_need", "").strip(),
            "job": form.get("job", "").strip(),
            "family_info": form.get("family_info", "").strip(),
            "requirement_must": form.get("requirement_must", "").strip(),
            "requirement_nice": form.get("requirement_nice", "").strip(),
            "other_background": form.get("other_background", "").strip(),
            "note": form.get("note", "").strip(),
            "labels": labels,
            "stage": form.get("stage", "").strip() or buyer.get("stage", ""),  # 保留最後狀態
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": session.get("user_id"),
            "updated_by_name": session.get("user_name"),
        }

        # ====== 圖片處理：多張刪除 + 多張新增 ======

        # 現在 Firestore 中的圖片列表：優先用 photo_urls，舊資料就用 photo_url
        current_photos = buyer.get("photo_urls") or []
        if not current_photos and buyer.get("photo_url"):
            current_photos = [buyer["photo_url"]]

        # 1️⃣ 要刪除的 index（來自 checkbox name="delete_photos"）
        delete_indexes_raw = form.getlist("delete_photos")
        delete_indexes = set()
        for idx in delete_indexes_raw:
            try:
                delete_indexes.add(int(idx))
            except ValueError:
                pass

        # 🔥 先記錄「要被刪除的 URL」（拿來刪 Storage）
        deleted_urls = [
            url for i, url in enumerate(current_photos)
            if i in delete_indexes
        ]

        # 保留沒勾選刪除的圖片
        new_photos = [
            url for i, url in enumerate(current_photos)
            if i not in delete_indexes
        ]

        # 2️⃣ 多張上傳：input name="photos" multiple
        files = request.files.getlist("photos")
        for f in files:
            if f and f.filename:
                photo_url = upload_image_to_storage(f, folder="buyers", object_id=buyer_id)
                if photo_url:
                    new_photos.append(photo_url)

        # 3️⃣ 寫回 Firestore：主要用 photo_urls，photo_url 當第一張給舊版用
        updated["photo_urls"] = new_photos
        if new_photos:
            updated["photo_url"] = new_photos[0]
        else:
            updated["photo_url"] = ""

        # ✅ 先更新 Firestore
        doc_ref.update(updated)

        # ✅ 再刪除 Firebase Storage 檔案
        if deleted_urls:
            delete_storage_files(deleted_urls)

        flash("已更新買方資料", "success")
        return redirect(url_for("buyer_detail", buyer_id=buyer_id))

    # GET：第一次進來編輯頁
    return render_template("buyer_edit.html", buyer=buyer)

# ========= 刪除買方（含追蹤） =========
@app.route("/buyers/<buyer_id>/delete", methods=["POST"])
@login_required
def buyer_delete(buyer_id):
    # 先刪除追蹤紀錄
    delete_by_field("buyer_followups", "buyer_id", buyer_id)
    # 再刪除買方本身
    db.collection("buyers").document(buyer_id).delete()

    flash("已刪除買方與相關追蹤紀錄", "info")
    return redirect(url_for("buyers"))


# ========= 買方追蹤紀錄：編輯 =========
@app.route("/buyers/<buyer_id>/followup/<followup_id>/edit", methods=["GET", "POST"])
@login_required
def buyer_followup_edit(buyer_id, followup_id):
    doc_ref = db.collection("buyer_followups").document(followup_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆追蹤紀錄", "danger")
        return redirect(url_for("buyer_detail", buyer_id=buyer_id))

    followup = doc_to_dict(doc)

    if request.method == "POST":
        contact_time = request.form.get("contact_time", "").strip()
        channel = request.form.get("channel", "").strip()
        content = request.form.get("content", "").strip()
        next_action = request.form.get("next_action", "").strip()
        next_contact_date = request.form.get("next_contact_date", "").strip()

        if not contact_time:
            contact_time = now_taipei().strftime("%Y-%m-%d %H:%M")

        doc_ref.update(
            {
                "contact_time": contact_time,
                "channel": channel,
                "content": content,
                "next_action": next_action,
                "next_contact_date": next_contact_date,
            }
        )

        flash("已更新追蹤紀錄", "success")
        return redirect(url_for("buyer_detail", buyer_id=buyer_id))

    return render_template("buyer_followup_edit.html", buyer_id=buyer_id, followup=followup)


# ========= 買方追蹤紀錄：刪除 =========
@app.route("/buyers/<buyer_id>/followup/<followup_id>/delete", methods=["POST"])
@login_required
def buyer_followup_delete(buyer_id, followup_id):
    db.collection("buyer_followups").document(followup_id).delete()
    flash("已刪除追蹤紀錄", "info")
    return redirect(url_for("buyer_detail", buyer_id=buyer_id))


# ========= 賣方列表 + 篩選 / 排序 =========
# ========= 賣方列表 =========
@app.route("/sellers")
@login_required
def sellers():
    q = request.args.get("q", "").strip()
    level = request.args.get("level", "").strip()
    stage = request.args.get("stage", "").strip()
    source = request.args.get("source", "").strip()
    label = request.args.get("label", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")

    docs = db.collection("sellers").stream()
    all_sellers = [doc_to_dict(d) for d in docs]

    source_set = set()
    for s in all_sellers:
        val = (s.get("source") or "").strip()
        if val:
            source_set.add(val)
    source_options = sorted(source_set)
    label_options = build_label_options(all_sellers)

    sellers_list = list(all_sellers)

    if q:
        sellers_list = [
            s for s in sellers_list
            if q in (s.get("name") or "") or q in (s.get("phone") or "")
        ]

    if level:
        sellers_list = [s for s in sellers_list if s.get("level") == level]

    if stage:
        sellers_list = [s for s in sellers_list if (s.get("stage") or "") == stage]

    if source:
        sellers_list = [s for s in sellers_list if (s.get("source") or "") == source]

    if label:
        sellers_list = [s for s in sellers_list if label in ensure_list(s.get("labels"))]

    def parse_created_at(s):
        return s.get("created_at") or ""

    if sort_by == "created_at_asc":
        sellers_list.sort(key=parse_created_at)
    elif sort_by == "created_at_desc":
        sellers_list.sort(key=parse_created_at, reverse=True)
    elif sort_by == "name_asc":
        sellers_list.sort(key=lambda s: (s.get("name") or ""))
    elif sort_by == "name_desc":
        sellers_list.sort(key=lambda s: (s.get("name") or ""), reverse=True)

    sellers_list = attach_latest_followup(sellers_list, "seller_followups", "seller_id")

    return render_template(
        "sellers.html",
        sellers=sellers_list,
        q=q,
        level=level,
        stage=stage,
        source=source,
        source_options=source_options,
        label=label,
        label_options=label_options,
        sort_by=sort_by,
        seller_stage_options=SELLER_STAGE_OPTIONS,
        total_count=len(all_sellers),
        filtered_count=len(sellers_list),
    )


# ========= 新增賣方 =========
@app.route("/sellers/new", methods=["POST"])
@login_required
def sellers_new():
    form = request.form

    name = form.get("name", "").strip()
    phone = form.get("phone", "").strip()
    email = form.get("email", "").strip()
    line_id = form.get("line_id", "").strip()
    address = form.get("address", "").strip()
    property_type = form.get("property_type", "").strip()
    level = form.get("level", "").strip()
    stage = form.get("stage", "").strip()   # 進程
    reason = form.get("reason", "").strip()
    expected_price = form.get("expected_price", "").strip()
    min_price = form.get("min_price", "").strip()
    timeline = form.get("timeline", "").strip()
    occupancy_status = form.get("occupancy_status", "").strip()
    contract_end_date = form.get("contract_end_date", "").strip()
    next_action = form.get("next_action", "").strip()
    next_contact_date = form.get("next_contact_date", "").strip()
    note = form.get("note", "").strip()
    labels = get_request_labels(form)

    # ⭐ 加上“來源 source”
    source = form.get("source", "").strip()

    if not name:
        flash("賣方姓名必填", "danger")
        return redirect(url_for("sellers"))

    now = now_taipei().isoformat()

    # 先產生一個 document id → 用來放圖片
    sellers_collection = db.collection("sellers")
    doc_ref = sellers_collection.document()
    seller_id = doc_ref.id

    # ========== 圖片（多張上傳） ==========
    photo_urls = []
    files = request.files.getlist("photos")   # <input name="photos" multiple>

    for f in files:
        if f and f.filename:
            url = upload_image_to_storage(f, folder="sellers", object_id=seller_id)
            if url:
                photo_urls.append(url)


    # ========== Firestore 要存的資料 ==========
    crm_vis = crm_record_visibility_payload_from_form(form)

    data = {
        **crm_vis,
        "name": name,
        "phone": phone,
        "email": email,
        "line_id": line_id,
        "address": address,
        "property_type": property_type,
        "level": level,
        "stage": stage,
        "reason": reason,
        "expected_price": expected_price,
        "min_price": min_price,
        "timeline": timeline,
        "occupancy_status": occupancy_status,
        "contract_end_date": contract_end_date,
        "next_action": next_action,
        "next_contact_date": next_contact_date,
        "note": note,
        "labels": labels,
        "source": source,               # ⭐ 加進 Firestore
        "created_at": now,
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
    }

    # 如果有圖片就寫入，沒有就給空
    if photo_urls:
        data["photo_urls"] = photo_urls
        data["photo_url"] = photo_urls[0]
    else:
        data["photo_urls"] = []
        data["photo_url"] = ""

    # 寫入 Firestore
    doc_ref.set(data)

    flash("已新增賣方", "success")
    return redirect(url_for("sellers"))





# ========= 賣方詳細 =========
@app.route("/sellers/<seller_id>")
@login_required
def seller_detail(seller_id):
    doc = db.collection("sellers").document(seller_id).get()
    if not doc.exists:
        flash("找不到這位賣方", "danger")
        return redirect(url_for("sellers"))

    seller = doc_to_dict(doc)

    followups_ref = db.collection("seller_followups").where("seller_id", "==", seller_id)
    followups = [doc_to_dict(f) for f in followups_ref.stream()]
    followups.sort(key=lambda x: x.get("contact_time", ""), reverse=True)

    return render_template("seller_detail.html", seller=seller, followups=followups)

@app.route("/sellers/<seller_id>/quick-stage", methods=["POST"])
@login_required
def seller_quick_stage(seller_id):
    stage = (request.form.get("stage", "") or request.form.get("current_stage", "")).strip()
    if not stage:
        flash("請選擇目前狀態", "warning")
        return redirect(request.referrer or url_for("sellers"))
    if stage not in SELLER_STAGE_OPTIONS:
        flash("委託狀態不在可選清單內", "danger")
        return redirect(request.referrer or url_for("sellers"))
    updates = {
        "stage": stage,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id"),
        "updated_by_name": session.get("user_name"),
    }
    db.collection("sellers").document(seller_id).update(updates)
    flash("已更新委託狀態", "success")
    return redirect(request.referrer or url_for("sellers"))


# ========= 新增賣方追蹤紀錄 =========
@app.route("/sellers/<seller_id>/followup", methods=["POST"])
@login_required
def add_seller_followup(seller_id):
    contact_time = request.form.get("contact_time", "").strip()
    channel = request.form.get("channel", "").strip()
    content = request.form.get("content", "").strip()
    next_action = request.form.get("next_action", "").strip()
    next_contact_date = request.form.get("next_contact_date", "").strip()

    if not contact_time:
        contact_time = now_taipei().strftime("%Y-%m-%d %H:%M")

    now = now_taipei().isoformat()

    db.collection("seller_followups").add(
        {
            "seller_id": seller_id,
            "contact_time": contact_time,
            "channel": channel,
            "content": content,
            "next_action": next_action,
            "next_contact_date": next_contact_date,
            "created_at": now,
            "created_by_id": session.get("user_id"),
            "created_by_name": session.get("user_name"),
        }
    )

    flash("已新增追蹤紀錄", "success")
    return redirect(url_for("seller_detail", seller_id=seller_id))


# ========= 編輯賣方 =========
@app.route("/sellers/<seller_id>/edit", methods=["GET", "POST"])
@login_required
def seller_edit(seller_id):
    doc_ref = db.collection("sellers").document(seller_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這位賣方", "danger")
        return redirect(url_for("sellers"))

    seller = doc_to_dict(doc)

    if request.method == "POST":
        form = request.form

        labels = get_request_labels(form)

        updated = {
            "name": form.get("name", "").strip(),
            "phone": form.get("phone", "").strip(),
            "email": form.get("email", "").strip(),
            "line_id": form.get("line_id", "").strip(),
            "address": form.get("address", "").strip(),
            "property_type": form.get("property_type", "").strip(),
            "level": form.get("level", "").strip(),
            "stage": form.get("stage", "").strip() or seller.get("stage", ""),  # 保留最後狀態
            "reason": form.get("reason", "").strip(),
            "expected_price": form.get("expected_price", "").strip(),
            "min_price": form.get("min_price", "").strip(),
            "timeline": form.get("timeline", "").strip(),
            "occupancy_status": form.get("occupancy_status", "").strip(),
            "contract_end_date": form.get("contract_end_date", "").strip(),  # 委託到期日
            "note": form.get("note", "").strip(),
            "labels": labels,
            # ⭐ 新增：客源來源
            "source": form.get("source", "").strip(),
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": session.get("user_id"),
            "updated_by_name": session.get("user_name"),
        }

        # ====== 圖片處理：多張刪除 + 多張新增 ======

        # 目前 Firestore 中的圖片列表（支援舊欄位 photo_url）
        current_photos = seller.get("photo_urls") or []
        if not current_photos and seller.get("photo_url"):
            current_photos = [seller["photo_url"]]

        # 1️⃣ 取得要刪除的 index（來自 checkbox）
        delete_indexes_raw = form.getlist("delete_photos")  # name="delete_photos"
        delete_indexes = set()
        for idx in delete_indexes_raw:
            try:
                delete_indexes.add(int(idx))
            except ValueError:
                pass

        # 把沒勾選的留下來
        new_photos = [
            url for i, url in enumerate(current_photos)
            if i not in delete_indexes
        ]

        # 2️⃣ 多張上傳：input name="photos" multiple
        files = request.files.getlist("photos")
        for f in files:
            if f and f.filename:
                photo_url = upload_image_to_storage(f, folder="sellers", object_id=seller_id)
                if photo_url:
                    new_photos.append(photo_url)

        # 3️⃣ 寫回 Firestore（主要用 photo_urls，photo_url 當第一張方便舊版使用）
        updated["photo_urls"] = new_photos
        if new_photos:
            updated["photo_url"] = new_photos[0]
        else:
            updated["photo_url"] = ""

        doc_ref.update(updated)
        flash("已更新賣方資料", "success")
        return redirect(url_for("seller_detail", seller_id=seller_id))

    # GET：首次載入編輯頁
    return render_template("seller_edit.html", seller=seller)


def delete_storage_file_by_url(url: str):
    """
    傳入 Firebase Storage 的檔案 URL（支援三種常見格式）：
    1) https://firebasestorage.googleapis.com/v0/b/<bucket>/o/<encoded_path>?...
    2) https://storage.googleapis.com/<bucket>/<path>
    3) gs://<bucket>/<path>
    自動解析出 bucket 與 blob path 並刪除。
    """
    if not url:
        return

    try:
        bucket = None
        blob_path = None

        # --- 格式 3：gs://bucket/path/to/file ---
        if url.startswith("gs://"):
            no_scheme = url[len("gs://"):]      # bucket/path/to/file
            parts = no_scheme.split("/", 1)
            bucket_name = parts[0]
            blob_path = parts[1] if len(parts) > 1 else ""
            bucket = storage.bucket(bucket_name)

        else:
            parsed = urlparse(url)
            netloc = parsed.netloc
            path = parsed.path  # e.g. /team-me-98acf.firebassestorage.app/sellers/xxx.jpg

            # --- 格式 1：firebasestorage.googleapis.com/v0/b/<bucket>/o/<encoded_path> ---
            if "firebasestorage.googleapis.com" in netloc and "/o/" in path:
                # 通常用預設 bucket 即可
                bucket = storage.bucket()
                encoded_blob_path = path.split("/o/", 1)[1]   # buyers%2Fabc%2Fxxx.jpg
                blob_path = unquote(encoded_blob_path)        # buyers/abc/xxx.jpg

            # --- 格式 2：storage.googleapis.com/<bucket>/<path> ---
            elif "storage.googleapis.com" in netloc:
                # path: /<bucket>/<blob_path>
                segments = path.lstrip("/").split("/", 1)
                if len(segments) == 2:
                    bucket_name, blob_path = segments
                    bucket = storage.bucket(bucket_name)

        if not bucket or not blob_path:
            print("⚠️ 無法解析 Storage URL：", url)
            return

        blob = bucket.blob(blob_path)
        if blob.exists():
            blob.delete()
            print(f"🔥 已刪除 Storage 檔案：{bucket.name}/{blob_path}")
        else:
            print(f"⚠️ 找不到 Storage 檔案：{bucket.name}/{blob_path}")

    except Exception as e:
        print("⚠️ 刪除 Storage 檔案發生錯誤：", e)


def delete_storage_files(urls: list):
    """一次刪多個 URL 對應的 Storage 檔案"""
    for url in urls:
        delete_storage_file_by_url(url)
# ========= 刪除賣方（含追蹤） =========
@app.route("/sellers/<seller_id>/delete", methods=["POST"])
@login_required
def seller_delete(seller_id):
    delete_by_field("seller_followups", "seller_id", seller_id)
    db.collection("sellers").document(seller_id).delete()

    flash("已刪除賣方與相關追蹤紀錄", "info")
    return redirect(url_for("sellers"))


# ========= 賣方追蹤紀錄：編輯 =========
@app.route("/sellers/<seller_id>/followup/<followup_id>/edit", methods=["GET", "POST"])
@login_required
def seller_followup_edit(seller_id, followup_id):
    doc_ref = db.collection("seller_followups").document(followup_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆追蹤紀錄", "danger")
        return redirect(url_for("seller_detail", seller_id=seller_id))

    followup = doc_to_dict(doc)

    if request.method == "POST":
        contact_time = request.form.get("contact_time", "").strip()
        channel = request.form.get("channel", "").strip()
        content = request.form.get("content", "").strip()
        next_action = request.form.get("next_action", "").strip()
        next_contact_date = request.form.get("next_contact_date", "").strip()

        if not contact_time:
            contact_time = now_taipei().strftime("%Y-%m-%d %H:%M")

        doc_ref.update(
            {
                "contact_time": contact_time,
                "channel": channel,
                "content": content,
                "next_action": next_action,
                "next_contact_date": next_contact_date,
            }
        )

        flash("已更新追蹤紀錄", "success")
        return redirect(url_for("seller_detail", seller_id=seller_id))

    return render_template("seller_followup_edit.html", seller_id=seller_id, followup=followup)


# ========= 賣方追蹤紀錄：刪除 =========
@app.route("/sellers/<seller_id>/followup/<followup_id>/delete", methods=["POST"])
@login_required
def seller_followup_delete(seller_id, followup_id):
    db.collection("seller_followups").document(followup_id).delete()
    flash("已刪除追蹤紀錄", "info")
    return redirect(url_for("seller_detail", seller_id=seller_id))




# ========= CSV：賣方 =========
@app.route("/sellers/download")
@login_required
def download_sellers():
    # 從 Firestore 抓全部賣方資料
    docs = db.collection("sellers").stream()
    sellers_list = [doc_to_dict(d) for d in docs]

    si = StringIO()
    writer = csv.writer(si)

    # 表頭（有進程 + 委託到期日）
    writer.writerow([
        "id",
        "姓名",
        "電話",
        "Email",
        "LINE ID",
        "物件地址",
        "產品類型",
        "客戶等級",
        "進程",              # 開發中 / 委託中 / 成交
        "出售原因",
        "期望售價(萬)",
        "可接受底價(萬)",
        "預計出售時程",
        "目前使用狀態",
        "委託到期日",
        "內部備註",
        "分類標籤",
        "建立時間",
        "建立者",
        "最後編輯時間",
        "最後編輯者",
    ])

    for s in sellers_list:
        writer.writerow([
            s.get("id", ""),
            s.get("name", ""),
            s.get("phone", ""),
            s.get("email", ""),
            s.get("line_id", ""),
            s.get("address", ""),
            s.get("property_type", ""),
            s.get("level", ""),
            s.get("stage", ""),                # 開發中 / 委託中 / 成交
            s.get("reason", ""),
            s.get("expected_price", ""),
            s.get("min_price", ""),
            s.get("timeline", ""),
            s.get("occupancy_status", ""),
            s.get("contract_end_date", ""),    # 委託到期日
            s.get("note", ""),
            "、".join(ensure_list(s.get("labels"))),
            s.get("created_at", ""),
            s.get("created_by_name", ""),
            s.get("updated_at", ""),
            s.get("updated_by_name", ""),
        ])

    csv_data = si.getvalue()
    csv_data = '\ufeff' + csv_data  # UTF-8 BOM

    filename = f"sellers.csv"
    response = Response(csv_data, mimetype="text/csv; charset=utf-8")
    response.headers["Content-Disposition"] = f"attachment; filename={filename}"
    return response

# ========= CSV：買方 =========
@app.route("/buyers/download")
@login_required
def download_buyers():
    # 從 Firestore 抓全部買方資料
    docs = db.collection("buyers").stream()
    buyers_list = [doc_to_dict(d) for d in docs]

    # 用 StringIO 暫存 CSV 文字
    si = StringIO()
    writer = csv.writer(si)

    # 表頭（你可以自行調整順序 / 欄位）
    writer.writerow([
        "id",
        "姓名",
        "電話",
        "Email",
        "LINE ID",
        "客源來源",
        "客戶等級",
        "進程",          # 接觸 / 帶看 / 斡旋 / 成交
        "需求類型",      # 買房 / 租屋 / 租買皆可
        "預算最低(萬)",
        "預算最高(萬)",
        "租金最低",
        "租金最高",
        "偏好區域",
        "產品類型",
        "房數需求",
        "車位需求",
        "職業/收入",
        "家庭成員/生活型態",
        "必備條件(Must Have)",
        "加分條件(Nice to Have)",
        "背景補充",
        "內部備註",
        "分類標籤",
        "建立時間",
        "建立者",
        "最後編輯時間",
        "最後編輯者",
    ])

    for b in buyers_list:
        writer.writerow([
            b.get("id", ""),
            b.get("name", ""),
            b.get("phone", ""),
            b.get("email", ""),
            b.get("line_id", ""),
            b.get("source", ""),
            b.get("level", ""),
            b.get("stage", ""),               # 進程
            b.get("intent_type", ""),         # 原始值（buy/rent/both），你也可以改成中文後再匯出
            b.get("budget_min", ""),
            b.get("budget_max", ""),
            b.get("rent_min", ""),
            b.get("rent_max", ""),
            b.get("preferred_areas", ""),
            b.get("property_type", ""),
            b.get("room_range", ""),
            b.get("car_need", ""),
            b.get("job", ""),
            b.get("family_info", ""),
            b.get("requirement_must", ""),
            b.get("requirement_nice", ""),
            b.get("other_background", ""),
            b.get("note", ""),
            "、".join(ensure_list(b.get("labels"))),
            b.get("created_at", ""),
            b.get("created_by_name", ""),
            b.get("updated_at", ""),
            b.get("updated_by_name", ""),
        ])

    # 取出 CSV 字串，加上 BOM 讓 Excel 顯示中文不亂碼
    csv_data = si.getvalue()
    csv_data = '\ufeff' + csv_data  # UTF-8 BOM

    # 回傳 Response，讓瀏覽器下載
    filename = f"buyers.csv"
    response = Response(csv_data, mimetype="text/csv; charset=utf-8")
    response.headers["Content-Disposition"] = f"attachment; filename={filename}"
    return response


# ========= LINE Webhook =========
@app.route("/line/ping")
def line_ping():
    return {"ok": True, "message": "line webhook ready"}, 200


@app.route("/line/webhook", methods=["POST"])
def line_webhook():
    raw_body = request.get_data(cache=False, as_text=False)
    signature = request.headers.get("x-line-signature", "")

    if not verify_line_signature(raw_body, signature):
        return "Invalid signature", 400

    try:
        payload = json.loads(raw_body.decode("utf-8"))
    except Exception as e:
        print("⚠️ LINE webhook JSON 解析失敗：", e)
        return "Bad Request", 400

    events = payload.get("events", [])
    for event in events:
        try:
            result = process_line_message_event(event)
            if not result or not result.get("handled"):
                continue

            if event.get("replyToken") and result.get("reply_text"):
                reply_result = reply_line_text(
                    event["replyToken"],
                    result["reply_text"] if result.get("ok") else result["reply_text"]
                )
                if result.get("ok") and result.get("target_type") and result.get("target_id") and reply_result:
                    for sent in reply_result.get("sent_messages", []):
                        sent_id = str(sent.get("id", "")).strip()
                        if sent_id:
                            save_line_message_link(
                                sent_id,
                                result["target_type"],
                                result["target_id"],
                                tag=result.get("parsed_tag", ""),
                                action="bot_reply",
                                customer_name=result.get("customer_name", ""),
                                phone=result.get("phone", ""),
                                source_event=event,
                            )
        except Exception as e:
            print("⚠️ 處理 LINE event 發生錯誤：", e)

    return "OK", 200

# ========= CLI：建立後台使用者 =========
@app.cli.command("create-user")
def create_user_cmd():
    """
    在命令列執行：
      flask --app team_me_firebase.py create-user
    然後依照提示輸入
    """
    import getpass

    email = input("Email: ").strip().lower()
    name = input("Name: ").strip()
    password = getpass.getpass("Password: ")

    if not email or not password:
        print("Email / Password 不可空白")
        return

    users_ref = db.collection("users").where("email", "==", email).limit(1)
    docs = list(users_ref.stream())
    if docs:
        print("此 Email 已存在")
        return

    pwd_hash = generate_password_hash(password)

    db.collection("users").add(
        {
            "email": email,
            "name": name or email,
            "password_hash": pwd_hash,
            "created_at": now_taipei().isoformat(),
        }
    )

    print("使用者建立完成")




# ========= 開發資料（developments） =========
def build_development_labels(extra=None):
    return dedupe_keep_order(["開發", "LINE紀錄"] + ensure_list(extra))


def normalize_line_key(key: str):
    k = (key or "").strip().replace(" ", "")
    mapping = {
        "對象": "target_type",
        "類型": "target_type",
        "目標": "target_type",
        "客戶類型": "target_type",
        "電話": "phone",
        "手機": "phone",
        "電話號碼": "phone",
        "姓名": "name",
        "客戶": "name",
        "買方": "name",
        "賣方": "name",
        "屋主": "name",
        "客戶來源": "source",
        "來源": "source",
        "需求類型": "intent_type_raw",
        "委託類型": "deal_type_raw",
        "ID": "record_id",
        "客戶ID": "record_id",
        "buyer_id": "record_id",
        "seller_id": "record_id",
        "development_id": "record_id",
        "內容": "content",
        "紀錄": "content",
        "備註": "content",
        "說明": "content",
        "進度內容": "content",
        "進程": "stage",
        "階段": "stage",
        "狀態": "stage",
        "標籤": "labels",
        "分類": "labels",
        "labels": "labels",
        "下一步": "next_action",
        "下次行動": "next_action",
        "next_action": "next_action",
        "下次聯絡日": "next_contact_date",
        "下次聯絡日期": "next_contact_date",
        "next_contact_date": "next_contact_date",
        "地址": "address",
        "物件": "address",
        "案名": "address",
        "總價": "price",
        "價格": "price",
        "預算": "budget",
        "區域": "preferred_areas",
        "產品類型": "property_type",
        "房數": "room_range",
        "車位": "car_need",
        "開價": "expected_price",
        "底價": "min_price",
        "委託到期日": "contract_end_date",
        "筆數": "limit",
        "網址": "url",
        "連結": "url",
        "網站": "url",
        "日期": "record_date",
    }
    return mapping.get(k, k)


def normalize_target_type(value: str) -> str:
    v = (value or "").strip()
    if v in ("買方", "客需", "buyer", "buyers"):
        return "buyer"
    if v in ("賣方", "委託", "seller", "sellers", "屋主"):
        return "seller"
    if v in ("開發", "development", "developments", "名單", "開發名單"):
        return "development"
    return ""


def parse_line_formatted_message(text: str):
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return None

    first = lines[0]
    if not first.startswith("#"):
        return None

    tag = first.lstrip("#").strip()
    tag_map = {
        "新增客需": "create_buyer_need",
        "新增委託": "create_seller_listing",
        "新增開發": "create_development",
        "買方追蹤": "buyer_followup",
        "賣方追蹤": "seller_followup",
        "開發追蹤": "development_followup",
        "客戶分類": "classify",
        "查詢紀錄": "query_records",
        "查詢委託到期": "query_contract_end",
        "查詢開發": "query_records",
        "帶看": "buyer_followup",
        "成交": "buyer_followup",
        "委託": "seller_followup",
        "紀錄": "generic_note",
    }
    action = tag_map.get(tag)
    if not action:
        return None

    fields = {}
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = m.group(2).strip()
        if key == "labels":
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value

    if tag in ("買方追蹤", "帶看", "成交") and not fields.get("target_type"):
        fields["target_type"] = "buyer"
    if tag in ("賣方追蹤", "委託") and not fields.get("target_type"):
        fields["target_type"] = "seller"
    if tag in ("開發追蹤", "查詢開發") and not fields.get("target_type"):
        fields["target_type"] = "development"

    fields["target_type"] = normalize_target_type(fields.get("target_type", "")) or fields.get("target_type", "")
    fields["intent_type"] = normalize_intent_type(fields.get("intent_type_raw", ""), fields)
    fields["deal_type"] = normalize_deal_type(fields.get("deal_type_raw", ""))
    fields["limit"] = parse_int_limit(fields.get("limit", 10), default=10, max_value=30)

    if action == "create_buyer_need":
        if not (fields.get("name") and fields.get("phone") and fields.get("source")):
            return None
    elif action == "create_seller_listing":
        if not (fields.get("name") and fields.get("phone") and fields.get("source")):
            return None
    elif action == "create_development":
        if not ((fields.get("phone") or fields.get("name")) and (fields.get("url") or fields.get("content"))):
            return None
    elif action in ("buyer_followup", "seller_followup", "development_followup", "classify", "query_records", "query_contract_end"):
        if not (fields.get("record_id") or fields.get("phone") or fields.get("name")):
            return None

    return {
        "tag": tag,
        "action": action,
        "fields": fields,
        "raw_text": text,
    }


def parse_potential_development_freeform(text: str):
    raw = (text or "").strip()
    if not raw or raw.startswith("#"):
        return None
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    if len(lines) < 3:
        return None

    url = ""
    phone = ""
    name = ""
    note_lines = []

    for ln in lines:
        if not url:
            m = re.search(r"https?://\S+", ln)
            if m:
                url = m.group(0)
                continue
        cleaned = re.sub(r"[^0-9]", "", ln)
        if not phone and len(cleaned) >= 8 and len(cleaned) <= 12:
            phone = ln
            continue

    if not url or not phone:
        return None

    # 姓名：取電話上一行或下一行最像名字的那行
    try:
        phone_idx = next(i for i, ln in enumerate(lines) if re.sub(r"[^0-9]", "", ln) == re.sub(r"[^0-9]", "", phone))
    except StopIteration:
        phone_idx = -1

    candidate_indices = []
    if phone_idx > 0:
        candidate_indices.append(phone_idx - 1)
    if phone_idx + 1 < len(lines):
        candidate_indices.append(phone_idx + 1)
    for idx in candidate_indices:
        ln = lines[idx]
        if ln != url and ln != phone and not re.search(r"https?://", ln):
            if len(ln) <= 20:
                name = ln
                break
    if not name:
        name = "未填姓名"

    for ln in lines:
        if ln in (url, phone, name):
            continue
        if re.search(r"https?://", ln):
            continue
        note_lines.append(ln)

    return {
        "tag": "新增開發",
        "action": "create_development",
        "fields": {
            "name": name,
            "phone": phone,
            "source": "LINE",
            "url": url,
            "content": "\n".join(note_lines).strip(),
            "stage": "待追蹤",
        },
        "raw_text": raw,
    }


def find_customer_record(target_type: str, record_id: str = "", phone: str = "", name: str = ""):
    if target_type == "buyer":
        collection_name = "buyers"
    elif target_type == "seller":
        collection_name = "sellers"
    else:
        collection_name = "developments"

    if record_id:
        doc = db.collection(collection_name).document(record_id).get()
        if doc.exists:
            return doc

    if phone:
        matches = find_records_by_phone(collection_name, phone)
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            return None

    if name:
        docs = list(db.collection(collection_name).where("name", "==", name.strip()).limit(2).stream())
        if len(docs) == 1:
            return docs[0]

    return None


def resolve_customer_record(fields, preferred_target_type=""):
    record_id = fields.get("record_id", "")
    phone = fields.get("phone", "")
    name = fields.get("name", "")
    target_type = preferred_target_type or normalize_target_type(fields.get("target_type", ""))

    if target_type in ("buyer", "seller", "development"):
        doc = find_customer_record(target_type, record_id, phone, name)
        return target_type, doc

    candidates = []
    for t in ("buyer", "seller", "development"):
        doc = find_customer_record(t, record_id, phone, name)
        if doc:
            candidates.append((t, doc))
    if len(candidates) == 1:
        return candidates[0]
    return "", None


def update_customer_note_and_labels(target_type: str, doc_ref, content: str, labels=None, stage="", source="LINE", event=None):
    labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(labels))
    snapshot = doc_ref.get()
    current = snapshot.to_dict() or {}
    old_note = current.get("note", "")

    source_label = build_line_operator_label(event) if event else "LINE"
    updates = {
        "note": append_note_block(old_note, content, source_label),
        "labels": firestore.ArrayUnion(labels),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }
    if stage:
        updates["stage"] = stage
    if source:
        updates["source"] = source

    doc_ref.update(updates)


def add_customer_followup(target_type: str, customer_id: str, content: str, next_action="", next_contact_date="", labels=None, line_event=None):
    if target_type == "buyer":
        collection_name = "buyer_followups"
        key_name = "buyer_id"
    elif target_type == "seller":
        collection_name = "seller_followups"
        key_name = "seller_id"
    else:
        collection_name = "development_followups"
        key_name = "development_id"

    sender_display_name = get_line_sender_display_name(line_event) if line_event else ""
    data = {
        key_name: customer_id,
        "contact_time": now_taipei().strftime("%Y-%m-%d %H:%M"),
        "channel": "LINE",
        "content": content,
        "next_action": next_action,
        "next_contact_date": next_contact_date,
        "labels": dedupe_keep_order(["LINE紀錄"] + ensure_list(labels)),
        "created_at": now_taipei().isoformat(),
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "sender_display_name": sender_display_name,
    }

    if line_event:
        source = line_event.get("source", {})
        data["line_group_id"] = source.get("groupId", "")
        data["line_room_id"] = source.get("roomId", "")
        data["line_user_id"] = source.get("userId", "")
        data["line_message_id"] = (line_event.get("message") or {}).get("id", "")
        data["quoted_message_id"] = (line_event.get("message") or {}).get("quotedMessageId", "")

    db.collection(collection_name).add(data)


def create_development(fields, event):
    phone = fields.get("phone", "").strip()
    name = fields.get("name", "").strip() or "未填姓名"
    url = fields.get("url", "").strip()
    matches = find_records_by_phone("developments", phone) if phone else []

    labels = build_development_labels(fields.get("labels"))
    content_text = fields.get("content", "").strip() or fields.get("address", "").strip() or url or "LINE 新增開發"
    note_content = build_line_summary(content_text, event)

    payload = {
        "name": name,
        "phone": phone,
        "source": fields.get("source", "").strip() or "LINE",
        "url": url,
        "current_stage": normalize_development_status(fields.get("current_stage", "").strip() or fields.get("stage", "").strip() or "待聯繫"),
        "stage": normalize_development_status(fields.get("current_stage", "").strip() or fields.get("stage", "").strip() or "待聯繫"),
        "next_action": normalize_development_next_action(fields.get("next_action", "").strip()),
        "next_action_date": fields.get("next_action_date", "").strip() or fields.get("next_contact_date", "").strip(),
        "record_date": fields.get("record_date", "").strip() or now_taipei().strftime("%Y-%m-%d"),
        "note": "",
        "labels": labels,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
        "sender_display_name": get_line_sender_display_name(event) or "",
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection("developments").document(doc.id)
        update_customer_note_and_labels(
            target_type="development",
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=payload["stage"],
            source=payload["source"],
            event=event,
        )
        doc_ref.update({k: v for k, v in payload.items() if v != "" and k != "note"})
        add_customer_followup(
            target_type="development",
            customer_id=doc.id,
            content=note_content,
            next_action=payload.get("next_action", ""),
            next_contact_date=payload.get("next_action_date", ""),
            labels=labels,
            line_event=event,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已註記開發：{updated_doc.get('name', '')}",
            "target_type": "development",
            "target_id": doc.id,
            "customer_name": updated_doc.get("name", ""),
            "phone": updated_doc.get("phone", ""),
            "parsed_tag": "新增開發",
        }

    if len(matches) > 1:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：同電話有多筆開發資料，請先到後台整理或改用客戶ID",
        }

    now = now_taipei().isoformat()
    payload.update({
        "created_at": now,
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "note": append_note_block("", note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection("developments").document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type="development",
        customer_id=doc_ref.id,
        content=note_content,
        next_action=payload.get("next_action", ""),
        next_contact_date=payload.get("next_action_date", ""),
        labels=labels,
        line_event=event,
    )
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記開發：{payload['name']}",
        "target_type": "development",
        "target_id": doc_ref.id,
        "customer_name": payload["name"],
        "phone": payload["phone"],
        "parsed_tag": "新增開發",
    }


def format_record_timeline(target_type: str, doc_snapshot, limit=10):
    data = doc_snapshot.to_dict() or {}
    record_id = doc_snapshot.id

    if target_type == "buyer":
        followup_collection = "buyer_followups"
        key_name = "buyer_id"
    elif target_type == "seller":
        followup_collection = "seller_followups"
        key_name = "seller_id"
    else:
        followup_collection = "development_followups"
        key_name = "development_id"

    followups = []
    for d in db.collection(followup_collection).where(key_name, "==", record_id).stream():
        item = d.to_dict() or {}
        followups.append({
            "time": item.get("contact_time") or item.get("created_at") or "",
            "channel": item.get("channel", "LINE"),
            "text": (item.get("content") or "").strip(),
            "created_by_name": item.get("created_by_name", "") or "",
            "sender_display_name": item.get("sender_display_name", "") or "",
        })

    followups = [x for x in followups if x.get("text")]
    followups.sort(key=lambda x: x.get("time", ""), reverse=True)

    lines = ["客戶資訊"]

    if target_type == "buyer":
        intent_map = {"rent": "租屋", "buy": "買賣", "both": "租買皆可"}
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {data.get('source', '') or '-'}",
            f"需求類型: {intent_map.get(data.get('intent_type', ''), data.get('intent_type', '') or '-')}",
            f"預算: {data.get('budget_min', '') or data.get('rent_min', '') or '-'} ~ {data.get('budget_max', '') or data.get('rent_max', '') or '-'}",
            f"偏好區域: {data.get('preferred_areas', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"房數需求: {data.get('room_range', '') or '-'}",
            f"車位需求: {data.get('car_need', '') or '-'}",
        ])
    elif target_type == "seller":
        deal_map = {"sale": "買賣", "rent": "出租"}
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {data.get('source', '') or '-'}",
            f"委託類型: {deal_map.get(data.get('deal_type', ''), data.get('deal_type', '') or '-')}",
            f"地址: {data.get('address', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"開價: {data.get('expected_price', '') or '-'}",
            f"底價: {data.get('min_price', '') or '-'}",
            f"委託到期日: {data.get('contract_end_date', '') or '-'}",
        ])
    else:
        lines.extend([
            f"日期: {data.get('record_date', '') or '-'}",
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '') or '-'}",
            f"來源: {data.get('source', '') or '-'}",
            f"進度: {data.get('stage', '') or '-'}",
            f"網址: {data.get('url', '') or '-'}",
            f"主紀錄: {data.get('note', '').splitlines()[-1] if data.get('note') else '-'}",
        ])

    lines.append("")
    lines.append("追蹤進度")

    if not followups:
        lines.append("目前沒有追蹤紀錄")
    else:
        for item in followups[:limit]:
            header_parts = [item.get('time', ''), item.get('channel', 'LINE')]
            creator = (item.get("created_by_name") or "").strip()
            sender = (item.get("sender_display_name") or "").strip()
            if creator:
                header_parts.append(f"KEYIN: {creator}")
            if sender:
                header_parts.append(f"留言者: {sender}")
            lines.append("｜".join([p for p in header_parts if p]))
            lines.append(item.get("text", ""))
            lines.append("")

    return "\n".join(lines).strip()[:4500]


def query_contract_end_text(fields):
    target_type, doc = resolve_customer_record(fields, preferred_target_type="seller")
    if target_type != "seller" or not doc:
        return False, "查無唯一委託資料，請提供正確電話或客戶ID", None

    data = doc.to_dict() or {}
    deal_type = data.get("deal_type", "")
    deal_label = "出租" if deal_type == "rent" else "買賣" if deal_type == "sale" else "未設定"
    text = "\n".join([
        "查詢委託到期",
        f"姓名: {data.get('name', '')}",
        f"電話: {data.get('phone', '')}",
        f"地址: {data.get('address', '')}",
        f"委託類型: {deal_label}",
        f"委託到期日: {data.get('contract_end_date', '') or '未填寫'}",
        f"目前進程: {data.get('stage', '') or '未填寫'}",
    ]).strip()
    return True, text[:4500], {
        "target_type": "seller",
        "target_id": doc.id,
        "customer_name": data.get("name", ""),
        "phone": data.get("phone", ""),
        "parsed_tag": "查詢委託到期",
    }


def process_quote_context_message(event):
    message = event.get("message") or {}
    quoted_message_id = message.get("quotedMessageId", "")
    raw_text = (message.get("text") or "").strip()
    if not quoted_message_id or not raw_text:
        return {"handled": False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {"handled": False}

    target_type = link.get("target_type", "")
    target_id = link.get("target_id", "")
    if target_type not in ("buyer", "seller", "development") or not target_id:
        return {"handled": False}

    if target_type == "buyer":
        collection_name = "buyers"
    elif target_type == "seller":
        collection_name = "sellers"
    else:
        collection_name = "developments"

    doc_ref = db.collection(collection_name).document(target_id)
    doc = doc_ref.get()
    if not doc.exists:
        return {"handled": True, "ok": False, "reply_text": "未寫入：引用的客戶資料不存在"}

    labels = dedupe_keep_order(["LINE紀錄", "群組回覆註記"])
    reply_only_text = raw_text

    update_customer_note_and_labels(
        target_type=target_type,
        doc_ref=doc_ref,
        content=reply_only_text,
        labels=labels,
        source="LINE",
        event=event,
    )
    add_customer_followup(
        target_type=target_type,
        customer_id=target_id,
        content=reply_only_text,
        labels=labels,
        line_event=event,
    )

    parsed = {
        "tag": "群組回覆註記",
        "action": "quoted_context_note",
        "fields": {"quoted_message_id": quoted_message_id},
        "raw_text": raw_text,
    }
    save_line_log(parsed, event, "success", target_type=target_type, target_id=target_id, sender_display_name=get_line_sender_display_name(event))
    target_label = "買方" if target_type == "buyer" else "賣方" if target_type == "seller" else "開發"
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記到{target_label}：{(doc.to_dict() or {}).get('name', '')}",
        "target_type": target_type,
        "target_id": target_id,
        "customer_name": (doc.to_dict() or {}).get("name", ""),
        "phone": (doc.to_dict() or {}).get("phone", ""),
        "parsed_tag": "群組回覆註記",
    }


def process_line_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    sender_display_name = get_line_sender_display_name(event)

    parsed = parse_line_formatted_message(message.get("text", ""))
    if not parsed:
        quoted_result = process_quote_context_message(event)
        if quoted_result.get("handled"):
            return quoted_result
        parsed = parse_potential_development_freeform(message.get("text", ""))
        if not parsed:
            return {"handled": False}

    fields = parsed["fields"]
    action = parsed["action"]

    if action == "create_buyer_need":
        result = create_buyer_need(fields, event)
    elif action == "create_seller_listing":
        result = create_seller_listing(fields, event)
    elif action == "create_development":
        result = create_development(fields, event)
    elif action == "query_records":
        target_type, doc = resolve_customer_record(fields)
        if not doc:
            result = {"handled": True, "ok": False, "reply_text": "查無唯一資料，請補電話或客戶ID"}
        else:
            result = {
                "handled": True,
                "ok": True,
                "reply_text": format_record_timeline(target_type, doc, limit=fields.get("limit", 10)),
                "target_type": target_type,
                "target_id": doc.id,
                "customer_name": (doc.to_dict() or {}).get("name", ""),
                "phone": (doc.to_dict() or {}).get("phone", ""),
                "parsed_tag": parsed.get("tag", ""),
            }
    elif action == "query_contract_end":
        ok, text, ctx = query_contract_end_text(fields)
        result = {"handled": True, "ok": ok, "reply_text": text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get("target_type", "")
        if action == "buyer_followup":
            target_type = "buyer"
        elif action == "seller_followup":
            target_type = "seller"
        elif action == "development_followup":
            target_type = "development"

        if target_type not in ("buyer", "seller", "development"):
            result = {"handled": True, "ok": False, "reply_text": "請提供對象：買方、賣方 或 開發"}
        else:
            doc = find_customer_record(
                target_type=target_type,
                record_id=fields.get("record_id", ""),
                phone=fields.get("phone", ""),
                name=fields.get("name", ""),
            )
            if not doc:
                target_label = "買方" if target_type == "buyer" else "賣方" if target_type == "seller" else "開發"
                result = {"handled": True, "ok": False, "reply_text": f"未寫入：找不到唯一一位{target_label}，請補客戶ID或正確電話"}
            else:
                content = fields.get("content", "").strip() or build_line_summary("LINE 更新", event)
                labels = fields.get("labels") or ["LINE紀錄"]
                collection_name = "buyers" if target_type == "buyer" else "sellers" if target_type == "seller" else "developments"
                doc_ref = db.collection(collection_name).document(doc.id)
                update_customer_note_and_labels(
                    target_type=target_type,
                    doc_ref=doc_ref,
                    content=content,
                    labels=labels,
                    stage=fields.get("stage", "").strip(),
                    source="LINE",
                    event=event,
                )
                add_customer_followup(
                    target_type=target_type,
                    customer_id=doc.id,
                    content=content,
                    next_action=fields.get("next_action", "").strip(),
                    next_contact_date=fields.get("next_contact_date", "").strip(),
                    labels=labels,
                    line_event=event,
                )
                data = doc_ref.get().to_dict() or {}
                target_label = "買方" if target_type == "buyer" else "賣方" if target_type == "seller" else "開發"
                result = {
                    "handled": True,
                    "ok": True,
                    "reply_text": f"已新增{target_label}追蹤：{data.get('name', '')}",
                    "target_type": target_type,
                    "target_id": doc.id,
                    "customer_name": data.get("name", ""),
                    "phone": data.get("phone", ""),
                    "parsed_tag": parsed.get("tag", ""),
                }

    if result.get("handled"):
        save_line_log(
            parsed=parsed,
            event=event,
            status="success" if result.get("ok") else "failed",
            target_type=result.get("target_type", ""),
            target_id=result.get("target_id", ""),
            note=result.get("reply_text", ""),
            sender_display_name=sender_display_name,
        )
    return result


def delete_by_field(collection_name, field_name, field_value):
    ref = db.collection(collection_name).where(field_name, "==", field_value)
    docs = list(ref.stream())
    for d in docs:
        d.reference.delete()




# ========= 開發模組加強：目前狀態 / 下一動作 / 下次日期、戶籍地址、批次掃街、LINE 即時更新 =========

DEVELOPMENT_STATUS_OPTIONS = [
    "待調謄本",
    "已調謄本",
    "待寄開發信",
    "已寄開發信待跑開發",
    "已跑開發",
    "待聯繫",
    "已聯繫",
    "持續追蹤",
    "已簽回",
    "已拋轉客需",
    "已拋轉委託",
    "無效",
]

DEVELOPMENT_NEXT_ACTION_OPTIONS = [
    "調謄本",
    "寄開發信",
    "跑開發",
    "電話聯繫",
    "LINE聯繫",
    "再次聯繫",
    "現場拜訪",
    "約面談",
    "簽回委託",
    "拋轉客需",
    "拋轉委託",
    "暫不處理",
    "結案",
]

DEVELOPMENT_HIDDEN_BY_DEFAULT = {"已簽回", "已拋轉客需", "已拋轉委託"}


def normalize_development_status(raw: str) -> str:
    v = (raw or "").strip()
    mapping = {
        "待調謄本": "待調謄本",
        "已調謄本": "已調謄本",
        "待寄開發信": "待寄開發信",
        "已寄開發信": "已寄開發信待跑開發",
        "已寄開發信待跑開發": "已寄開發信待跑開發",
        "寄出開發信": "已寄開發信待跑開發",
        "已跑開發": "已跑開發",
        "待聯絡": "待聯繫",
        "待聯繫": "待聯繫",
        "已聯絡": "已聯繫",
        "已連繫": "已聯繫",
        "已聯繫": "已聯繫",
        "再聯絡": "持續追蹤",
        "再聯繫": "持續追蹤",
        "持續追蹤": "持續追蹤",
        "已簽回": "已簽回",
        "已拋轉客需": "已拋轉客需",
        "已拋轉委託": "已拋轉委託",
        "無效": "無效",
    }
    return mapping.get(v, v)

def normalize_development_next_action(raw: str) -> str:
    v = (raw or "").strip()
    mapping = {
        "調謄本": "調謄本",
        "寄開發信": "寄開發信",
        "跑開發": "跑開發",
        "電話聯絡": "電話聯繫",
        "電話聯繫": "電話聯繫",
        "LINE聯絡": "LINE聯繫",
        "LINE聯繫": "LINE聯繫",
        "再次聯絡": "再次聯繫",
        "再次聯繫": "再次聯繫",
        "現場拜訪": "現場拜訪",
        "約面談": "約面談",
        "簽回委託": "簽回委託",
        "拋轉客需": "拋轉客需",
        "拋轉委託": "拋轉委託",
        "暫不處理": "暫不處理",
        "結案": "結案",
    }
    return mapping.get(v, v)


def normalize_line_key(key: str):
    k = (key or "").strip().replace(" ", "")
    mapping = {
        "對象": "target_type",
        "類型": "target_type",
        "目標": "target_type",
        "客戶類型": "target_type",
        "電話": "phone",
        "手機": "phone",
        "電話號碼": "phone",
        "姓名": "name",
        "客戶": "name",
        "買方": "name",
        "賣方": "name",
        "屋主": "name",
        "客戶來源": "source",
        "來源": "source",
        "需求類型": "intent_type_raw",
        "委託類型": "deal_type_raw",
        "ID": "record_id",
        "客戶ID": "record_id",
        "buyer_id": "record_id",
        "seller_id": "record_id",
        "development_id": "record_id",
        "內容": "content",
        "紀錄": "content",
        "備註": "content",
        "說明": "content",
        "進度內容": "content",
        "進度": "stage",
        "進程": "stage",
        "階段": "stage",
        "狀態": "stage",
        "目前狀態": "current_stage",
        "下一動作": "next_action",
        "下次日期": "next_action_date",
        "下次動作日期": "next_action_date",
        "標籤": "labels",
        "分類": "labels",
        "labels": "labels",
        "下一步": "next_action",
        "下次行動": "next_action",
        "next_action": "next_action",
        "下次聯絡日": "next_contact_date",
        "下次聯絡日期": "next_contact_date",
        "next_contact_date": "next_contact_date",
        "地址": "address",
        "物件": "address",
        "案名": "address",
        "總價": "price",
        "價格": "price",
        "預算": "budget",
        "區域": "preferred_areas",
        "產品類型": "property_type",
        "房數": "room_range",
        "車位": "car_need",
        "開價": "expected_price",
        "底價": "min_price",
        "委託到期日": "contract_end_date",
        "筆數": "limit",
        "網址": "url",
        "連結": "url",
        "網站": "url",
        "日期": "record_date",
        "戶籍地址": "registered_address",
        "戶籍地": "registered_address",
    }
    return mapping.get(k, k)


def normalize_target_type(value: str) -> str:
    v = (value or "").strip()
    if v in ("買方", "客需", "buyer", "buyers"):
        return "buyer"
    if v in ("賣方", "委託", "seller", "sellers", "屋主"):
        return "seller"
    if v in ("開發", "development", "developments", "名單", "開發名單"):
        return "development"
    return ""


def find_customer_record(target_type: str, record_id: str = "", phone: str = "", name: str = "", address: str = ""):
    if target_type == "buyer":
        collection_name = "buyers"
    elif target_type == "seller":
        collection_name = "sellers"
    else:
        collection_name = "developments"

    if record_id:
        doc = db.collection(collection_name).document(record_id).get()
        if doc.exists:
            return doc

    if phone:
        matches = find_records_by_phone(collection_name, phone)
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            return None

    if name:
        docs = list(db.collection(collection_name).where("name", "==", name.strip()).limit(2).stream())
        if len(docs) == 1:
            return docs[0]

    if target_type == "development" and address:
        docs = [d for d in db.collection(collection_name).stream() if (d.to_dict() or {}).get("address", "").strip() == address.strip()]
        if len(docs) == 1:
            return docs[0]

    return None


def resolve_customer_record(fields, preferred_target_type=""):
    record_id = fields.get("record_id", "")
    phone = fields.get("phone", "")
    name = fields.get("name", "")
    address = fields.get("address", "")
    target_type = preferred_target_type or normalize_target_type(fields.get("target_type", ""))

    if target_type in ("buyer", "seller", "development"):
        doc = find_customer_record(target_type, record_id, phone, name, address)
        return target_type, doc

    buyer_doc = find_customer_record("buyer", record_id, phone, name, address)
    seller_doc = find_customer_record("seller", record_id, phone, name, address)
    development_doc = find_customer_record("development", record_id, phone, name, address)

    hits = [(t, d) for t, d in [("buyer", buyer_doc), ("seller", seller_doc), ("development", development_doc)] if d]
    if len(hits) == 1:
        return hits[0][0], hits[0][1]
    return "", None



def update_customer_note_and_labels(target_type: str, doc_ref, content: str, labels=None, stage="", source="LINE", event=None, registered_address="", extra_updates=None):
    labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(labels))
    snapshot = doc_ref.get()
    current = snapshot.to_dict() or {}
    old_note = current.get("note", "")

    updates = {
        "labels": firestore.ArrayUnion(labels),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }
    if content:
        source_label = build_line_operator_label(event) if event else "LINE"
        updates["note"] = append_note_block(old_note, content, source_label)
    if stage:
        updates["stage"] = stage
    if source:
        updates["source"] = source
    if target_type == "development" and registered_address:
        updates["registered_address"] = registered_address
    if extra_updates:
        cleaned = {k: v for k, v in extra_updates.items() if v not in (None, "")}
    
    if target_type == "development":
        updates, note_text = parse_reply_field_updates(raw_text)

        extra_updates = {k: v for k, v in updates.items() if k in ("address", "phone", "name", "url", "current_stage", "next_action", "next_action_date")}
        if updates:
            update_customer_note_and_labels(
                target_type="development",
                doc_ref=doc_ref,
                content=note_text or reply_only_text,
                labels=labels,
                stage=updates.get("current_stage", ""),
                source=updates.get("source", "LINE"),
                event=event,
                registered_address=updates.get("registered_address", ""),
                extra_updates=extra_updates,
            )
        else:
            update_customer_note_and_labels(
                target_type="development",
                doc_ref=doc_ref,
                content=reply_only_text,
                labels=labels,
                source="LINE",
                event=event,
            )

        followup_text = note_text or reply_only_text or "更新開發資料"
        add_customer_followup(
            target_type="development",
            customer_id=target_id,
            content=followup_text,
            labels=labels,
            line_event=event,
            current_stage=updates.get("current_stage", "") if updates else "",
            registered_address=updates.get("registered_address", "") if updates else "",
            next_action=updates.get("next_action", "") if updates else "",
            next_action_date=updates.get("next_action_date", "") if updates else "",
        )
    else:
        collection_name = "development_followups"
        key_name = "development_id"

    sender_display_name = get_line_sender_display_name(line_event) if line_event else ""
    data = {
        key_name: customer_id,
        "contact_time": now_taipei().strftime("%Y-%m-%d %H:%M"),
        "channel": channel,
        "content": content,
        "next_action": next_action,
        "next_contact_date": next_contact_date,
        "labels": dedupe_keep_order(["LINE紀錄"] + ensure_list(labels)),
        "created_at": now_taipei().isoformat(),
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "sender_display_name": sender_display_name,
    }
    if target_type == "development":
        cs = normalize_development_status(current_stage or stage)
        na = normalize_development_next_action(next_action)
        data["current_stage"] = cs
        data["stage"] = cs
        data["registered_address"] = registered_address.strip() if registered_address else ""
        data["next_action"] = na
        data["next_action_date"] = (next_action_date or next_contact_date or "").strip()

    if line_event:
        source = line_event.get("source", {})
        data["line_group_id"] = source.get("groupId", "")
        data["line_room_id"] = source.get("roomId", "")
        data["line_user_id"] = source.get("userId", "")
        data["line_message_id"] = (line_event.get("message") or {}).get("id", "")
        data["quoted_message_id"] = (line_event.get("message") or {}).get("quotedMessageId", "")

    db.collection(collection_name).add(data)


def parse_line_formatted_message(text: str):
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return None

    first = lines[0]
    if not first.startswith("#"):
        return None

    tag = first.lstrip("#").strip()
    if tag == "新增開發批次":
        return {
            "tag": tag,
            "action": "create_development_batch",
            "fields": {},
            "raw_text": text,
        }

    tag_map = {
        "新增客需": "create_buyer_need",
        "新增委託": "create_seller_listing",
        "新增開發": "create_development",
        "買方追蹤": "buyer_followup",
        "賣方追蹤": "seller_followup",
        "開發追蹤": "development_followup",
        "客戶分類": "classify",
        "查詢紀錄": "query_records",
        "查詢開發": "query_records",
        "查詢委託到期": "query_contract_end",
        "帶看": "buyer_followup",
        "成交": "buyer_followup",
        "委託": "seller_followup",
        "紀錄": "generic_note",
    }
    action = tag_map.get(tag)
    if not action:
        return None

    fields = {}
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = m.group(2).strip()
        if key == "labels":
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value

    if tag in ("買方追蹤", "帶看", "成交") and not fields.get("target_type"):
        fields["target_type"] = "buyer"
    if tag in ("賣方追蹤", "委託") and not fields.get("target_type"):
        fields["target_type"] = "seller"
    if tag in ("開發追蹤", "查詢開發") and not fields.get("target_type"):
        fields["target_type"] = "development"

    fields["target_type"] = normalize_target_type(fields.get("target_type", "")) or fields.get("target_type", "")
    fields["intent_type"] = normalize_intent_type(fields.get("intent_type_raw", ""), fields)
    fields["deal_type"] = normalize_deal_type(fields.get("deal_type_raw", ""))
    fields["limit"] = parse_int_limit(fields.get("limit", 10), default=10, max_value=30)
    if fields.get("stage"):
        fields["stage"] = normalize_development_status(fields.get("stage", ""))

    if action == "create_buyer_need":
        if not (fields.get("name") and fields.get("phone") and fields.get("source")):
            return None
    elif action == "create_seller_listing":
        if not (fields.get("name") and fields.get("phone") and fields.get("source")):
            return None
    elif action == "create_development":
        if not (fields.get("name") or fields.get("phone") or fields.get("address") or fields.get("url")):
            return None
    elif action in ("buyer_followup", "seller_followup", "development_followup", "classify", "query_records", "query_contract_end"):
        if not (fields.get("record_id") or fields.get("phone") or fields.get("name") or fields.get("address")):
            return None

    return {
        "tag": tag,
        "action": action,
        "fields": fields,
        "raw_text": text,
    }


def parse_potential_development_freeform(text: str):
    raw = (text or "").strip()
    if not raw or raw.startswith("#"):
        return None
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    if len(lines) < 3:
        return None

    url = ""
    phone = ""
    name = ""
    address = ""
    note_lines = []

    for ln in lines:
        if not url:
            m = re.search(r"https?://\S+", ln)
            if m:
                url = m.group(0)
                continue

    # 找電話
    for ln in lines:
        cleaned = re.sub(r"[^0-9]", "", ln)
        if len(cleaned) >= 8 and len(cleaned) <= 12:
            phone = ln
            break

    # 找地址：含台中/臺中/區/路/街/巷/號者優先
    for ln in lines:
        if re.search(r"(台中|臺中).*(區).*(路|街|巷).*(號)", ln) or re.search(r"(路|街|巷).*(號)", ln):
            address = ln
            break

    # 姓名：優先找先生/小姐/太太/女士或短字串
    for ln in lines:
        if ln in (url, phone, address):
            continue
        if re.search(r"(先生|小姐|太太|女士)$", ln):
            name = ln
            break
    if not name:
        for ln in lines:
            if ln in (url, phone, address):
                continue
            if len(ln) <= 12 and not re.search(r"https?://", ln):
                name = ln
                break

    for ln in lines:
        if ln in (url, phone, name, address):
            continue
        if re.search(r"https?://", ln):
            continue
        note_lines.append(ln)

    if not (phone or url or address):
        return None

    return {
        "tag": "新增開發",
        "action": "create_development",
        "fields": {
            "name": name or "未填姓名",
            "phone": phone,
            "address": address,
            "source": "LINE",
            "url": url,
            "content": "\n".join(note_lines).strip(),
            "stage": "待聯繫",
        },
        "raw_text": raw,
    }



def parse_development_batch_message(raw_text: str):
    text = (raw_text or "").replace("\r\n", "\n").strip()
    lines = text.split("\n")
    if lines and lines[0].strip().startswith("#"):
        lines = lines[1:]
    body = "\n".join(lines).strip()
    if not body:
        return []

    # 支援 --- 與 ---- 當分隔線；優先使用「獨立一行」的分隔，避免誤切備註內容。
    chunks = [c.strip() for c in re.split(r"(?m)^-{3,}\s*$", body) if c.strip()]
    items = []
    for chunk in chunks:
        fields = {}
        notes = []
        for line in chunk.split("\n"):
            ln = line.strip()
            if not ln:
                continue
            m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", ln)
            if m:
                key = normalize_line_key(m.group(1))
                value = (m.group(2) or "").strip()
                if key == "labels":
                    fields[key] = parse_label_csv(value)
                else:
                    fields[key] = value
            else:
                notes.append(ln)
        if notes:
            fields["content"] = (fields.get("content", "") + ("\n" if fields.get("content") else "") + "\n".join(notes)).strip()

        if not fields.get("current_stage"):
            fields["current_stage"] = normalize_development_status(fields.get("stage", "") or "待聯繫")
        else:
            fields["current_stage"] = normalize_development_status(fields["current_stage"])
        fields["stage"] = fields["current_stage"]

        if fields.get("next_action"):
            fields["next_action"] = normalize_development_next_action(fields["next_action"])

        if fields.get("name") or fields.get("phone") or fields.get("address") or fields.get("url") or fields.get("content"):
            items.append(fields)
    return items


def infer_development_source(explicit_source: str, url: str) -> str:
    explicit = (explicit_source or "").strip()
    if explicit in ("掃街", "踩線"):
        return explicit
    if explicit and explicit.upper() != "LINE":
        return explicit
    return "踩線" if (url or "").strip() else "掃街"


def create_development(fields, event):
    sender_name = get_line_sender_display_name(event) or "未知成員"
    phone = fields.get("phone", "").strip()
    name = fields.get("name", "").strip() or "未填姓名"
    url = fields.get("url", "").strip()
    address = fields.get("address", "").strip()
    registered_address = fields.get("registered_address", "").strip()
    current_stage = normalize_development_status(fields.get("current_stage", "").strip() or fields.get("stage", "").strip() or "待聯繫")
    next_action = normalize_development_next_action(fields.get("next_action", "").strip())
    next_action_date = (fields.get("next_action_date", "") or fields.get("next_contact_date", "")).strip()
    source = infer_development_source(fields.get("source", ""), url)

    matches = []
    if phone:
        matches = find_records_by_phone("developments", phone)
    elif address:
        matches = [d for d in db.collection("developments").stream() if (d.to_dict() or {}).get("address", "").strip() == address]

    labels = build_development_labels(fields.get("labels"))
    summary_content = fields.get("content", "").strip() or address or url or "LINE 新增開發"
    note_content = build_line_summary(summary_content, event)

    payload = {
        "name": name,
        "phone": phone,
        "source": source,
        "url": url,
        "address": address,
        "registered_address": registered_address,
        "current_stage": current_stage,
        "stage": current_stage,
        "next_action": next_action,
        "next_action_date": next_action_date,
        "record_date": fields.get("record_date", "").strip() or now_taipei().strftime("%Y-%m-%d"),
        "note": "",
        "labels": labels,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
        "sender_display_name": sender_name,
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection("developments").document(doc.id)
        update_customer_note_and_labels(
            target_type="development",
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=current_stage,
            source=source,
            event=event,
            registered_address=registered_address,
            extra_updates={
                "address": address,
                "url": url,
                "name": name,
                "current_stage": current_stage,
                "next_action": next_action,
                "next_action_date": next_action_date,
            },
        )
        add_customer_followup(
            target_type="development",
            customer_id=doc.id,
            content=note_content,
            labels=labels,
            line_event=event,
            current_stage=current_stage,
            registered_address=registered_address,
            next_action=next_action,
            next_action_date=next_action_date,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已註記開發：{updated_doc.get('name', '')}",
            "target_type": "development",
            "target_id": doc.id,
            "customer_name": updated_doc.get("name", ""),
            "phone": updated_doc.get("phone", ""),
            "parsed_tag": "新增開發",
        }

    if len(matches) > 1:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：同電話或同地址有多筆開發資料，請先到後台整理或改用客戶ID",
        }

    now = now_taipei().isoformat()
    payload.update({
        "created_at": now,
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "note": append_note_block("", note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection("developments").document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type="development",
        customer_id=doc_ref.id,
        content=note_content,
        labels=labels,
        line_event=event,
        current_stage=current_stage,
        registered_address=registered_address,
        next_action=next_action,
        next_action_date=next_action_date,
    )
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記開發：{payload['name']}",
        "target_type": "development",
        "target_id": doc_ref.id,
        "customer_name": payload["name"],
        "phone": payload["phone"],
        "parsed_tag": "新增開發",
    }

def create_development_batch(raw_text, event):

    items = parse_development_batch_message(raw_text)
    if not items:
        return {"handled": True, "ok": False, "reply_text": "未寫入：#新增開發批次 沒有解析到有效資料"}

    ok_count = 0
    fail_count = 0
    last_result = None
    for fields in items:
        result = create_development(fields, event)
        last_result = result
        if result.get("ok"):
            ok_count += 1
        else:
            fail_count += 1

    reply_text = f"批量註記完成：成功 {ok_count} 筆，失敗 {fail_count} 筆"
    if last_result and last_result.get("target_type") and last_result.get("target_id"):
        return {
            "handled": True,
            "ok": ok_count > 0,
            "reply_text": reply_text,
            "target_type": last_result.get("target_type", ""),
            "target_id": last_result.get("target_id", ""),
            "customer_name": last_result.get("customer_name", ""),
            "phone": last_result.get("phone", ""),
            "parsed_tag": "新增開發批次",
        }
    return {"handled": True, "ok": ok_count > 0, "reply_text": reply_text}



def parse_reply_field_updates(text: str):
    updates = {}
    note_lines = []

    for raw in (text or "").splitlines():
        line = raw.strip()
        if not line:
            continue

        if "：" in line:
            key, value = line.split("：", 1)
        elif ":" in line:
            key, value = line.split(":", 1)
        else:
            maybe_status = normalize_development_status(line)
            maybe_action = normalize_development_next_action(line)
            if maybe_status in DEVELOPMENT_STATUS_OPTIONS:
                updates["current_stage"] = maybe_status
            elif maybe_action in DEVELOPMENT_NEXT_ACTION_OPTIONS:
                updates["next_action"] = maybe_action
            else:
                note_lines.append(line)
            continue

        key = normalize_line_key(key.strip())
        value = value.strip()

        if key in ("stage", "current_stage"):
            updates["current_stage"] = normalize_development_status(value)
        elif key == "registered_address":
            updates["registered_address"] = value
        elif key == "next_action":
            updates["next_action"] = normalize_development_next_action(value)
        elif key in ("next_action_date", "next_contact_date"):
            updates["next_action_date"] = value
        elif key in ("address", "phone", "name", "url", "source"):
            updates[key] = value
        elif key == "content":
            note_lines.append(value)

    return updates, "\n".join(note_lines).strip()


def process_quote_context_message(event):
    message = event.get("message") or {}
    quoted_message_id = message.get("quotedMessageId", "")
    raw_text = (message.get("text") or "").strip()
    if not quoted_message_id or not raw_text:
        return {"handled": False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {"handled": False}

    target_type = link.get("target_type", "")
    target_id = link.get("target_id", "")
    if target_type not in ("buyer", "seller", "development") or not target_id:
        return {"handled": False}

    if target_type == "buyer":
        collection_name = "buyers"
        label_text = "買方"
    elif target_type == "seller":
        collection_name = "sellers"
        label_text = "賣方"
    else:
        collection_name = "developments"
        label_text = "開發"

    doc_ref = db.collection(collection_name).document(target_id)
    doc = doc_ref.get()
    if not doc.exists:
        return {"handled": True, "ok": False, "reply_text": "未寫入：引用的資料不存在"}

    labels = dedupe_keep_order(["LINE紀錄", "群組回覆註記"])
    reply_only_text = raw_text
    sender_display_name = get_line_sender_display_name(event)

    if target_type == "development":
        updates = parse_reply_field_updates(raw_text)
        if updates:
            update_customer_note_and_labels(
                target_type="development",
                doc_ref=doc_ref,
                content=reply_only_text,
                labels=labels,
                stage=updates.get("stage", ""),
                source=updates.get("source", "LINE"),
                event=event,
                registered_address=updates.get("registered_address", ""),
                extra_updates={k: v for k, v in updates.items() if k in ("address", "phone", "name", "url")},
            )
        else:
            update_customer_note_and_labels(
                target_type="development",
                doc_ref=doc_ref,
                content=reply_only_text,
                labels=labels,
                source="LINE",
                event=event,
            )

        add_customer_followup(
            target_type="development",
            customer_id=target_id,
            content=reply_only_text,
            labels=labels,
            line_event=event,
            stage=updates.get("stage", "") if updates else "",
            registered_address=updates.get("registered_address", "") if updates else "",
        )
    else:
        update_customer_note_and_labels(
            target_type=target_type,
            doc_ref=doc_ref,
            content=reply_only_text,
            labels=labels,
            source="LINE",
            event=event,
        )
        add_customer_followup(
            target_type=target_type,
            customer_id=target_id,
            content=reply_only_text,
            labels=labels,
            line_event=event,
        )

    parsed = {
        "tag": "群組回覆註記",
        "action": "quoted_context_note",
        "fields": {"quoted_message_id": quoted_message_id},
        "raw_text": raw_text,
    }
    save_line_log(parsed, event, "success", target_type=target_type, target_id=target_id, sender_display_name=sender_display_name)
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記到{label_text}：{(doc.to_dict() or {}).get('name', '')}",
        "target_type": target_type,
        "target_id": target_id,
        "customer_name": (doc.to_dict() or {}).get("name", ""),
        "phone": (doc.to_dict() or {}).get("phone", ""),
        "parsed_tag": "群組回覆註記",
    }


def format_record_timeline(target_type: str, doc_snapshot, limit=10):
    data = doc_snapshot.to_dict() or {}
    record_id = doc_snapshot.id

    if target_type == "buyer":
        followup_collection = "buyer_followups"
        key_name = "buyer_id"
    elif target_type == "seller":
        followup_collection = "seller_followups"
        key_name = "seller_id"
    else:
        followup_collection = "development_followups"
        key_name = "development_id"

    followups = []
    for d in db.collection(followup_collection).where(key_name, "==", record_id).stream():
        item = d.to_dict() or {}
        followups.append({
            "time": item.get("contact_time") or item.get("created_at") or "",
            "channel": item.get("channel", "LINE"),
            "text": (item.get("content") or "").strip(),
            "created_by_name": item.get("created_by_name", "") or "",
            "sender_display_name": item.get("sender_display_name", "") or "",
            "stage": item.get("stage", "") or "",
            "registered_address": item.get("registered_address", "") or "",
        })

    followups = [x for x in followups if x.get("text") or x.get("registered_address")]
    followups.sort(key=lambda x: x.get("time", ""), reverse=True)

    lines = ["客戶資訊"]

    if target_type == "buyer":
        intent_map = {"rent": "租屋", "buy": "買賣", "both": "租買皆可"}
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {data.get('source', '') or '-'}",
            f"需求類型: {intent_map.get(data.get('intent_type', ''), data.get('intent_type', '') or '-')}",
            f"預算: {data.get('budget_min', '') or data.get('rent_min', '') or '-'} ~ {data.get('budget_max', '') or data.get('rent_max', '') or '-'}",
            f"偏好區域: {data.get('preferred_areas', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"房數需求: {data.get('room_range', '') or '-'}",
            f"車位需求: {data.get('car_need', '') or '-'}",
        ])
    elif target_type == "seller":
        deal_map = {"sale": "買賣", "rent": "出租"}
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {data.get('source', '') or '-'}",
            f"委託類型: {deal_map.get(data.get('deal_type', ''), data.get('deal_type', '') or '-')}",
            f"地址: {data.get('address', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"開價: {data.get('expected_price', '') or '-'}",
            f"底價: {data.get('min_price', '') or '-'}",
            f"委託到期日: {data.get('contract_end_date', '') or '-'}",
        ])
    else:
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '') or '-'}",
            f"來源: {data.get('source', '') or '-'}",
            f"地址: {data.get('address', '') or '-'}",
            f"戶籍地址: {data.get('registered_address', '') or '-'}",
            f"網址: {data.get('url', '') or '-'}",
            f"進度: {data.get('stage', '') or '-'}",
        ])

    lines.append("")
    lines.append("追蹤進度")

    if not followups:
        lines.append("目前沒有追蹤紀錄")
    else:
        for item in followups[:limit]:
            header_parts = [item.get('time', ''), item.get('channel', 'LINE')]
            if item.get("stage"):
                header_parts.append(item["stage"])
            creator = (item.get("created_by_name") or "").strip()
            sender = (item.get("sender_display_name") or "").strip()
            if creator:
                header_parts.append(f"KEYIN: {creator}")
            if sender:
                header_parts.append(f"留言者: {sender}")
            lines.append("｜".join([x for x in header_parts if x]))
            if item.get("text"):
                lines.append(item.get("text", ""))
            if item.get("registered_address"):
                lines.append(f"戶籍地址：{item['registered_address']}")
            lines.append("")

    output = "\n".join(lines).strip()
    return output[:4500]


def process_line_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    sender_display_name = get_line_sender_display_name(event)
    raw_text = message.get("text", "")

    parsed = parse_line_formatted_message(raw_text)
    if not parsed:
        quoted_result = process_quote_context_message(event)
        if quoted_result.get("handled"):
            return quoted_result

        freeform = parse_potential_development_freeform(raw_text)
        if freeform:
            parsed = freeform
        else:
            return {"handled": False}

    fields = parsed["fields"]
    action = parsed["action"]

    if action == "create_buyer_need":
        result = create_buyer_need(fields, event)
    elif action == "create_seller_listing":
        result = create_seller_listing(fields, event)
    elif action == "create_development":
        result = create_development(fields, event)
    elif action == "create_development_batch":
        result = create_development_batch(parsed.get("raw_text", raw_text), event)
    elif action == "query_records":
        target_type, doc = resolve_customer_record(fields)
        if not doc:
            result = {"handled": True, "ok": False, "reply_text": "查無唯一客戶，請補電話、地址或客戶ID"}
        else:
            result = {
                "handled": True,
                "ok": True,
                "reply_text": format_record_timeline(target_type, doc, limit=fields.get("limit", 10)),
                "target_type": target_type,
                "target_id": doc.id,
                "customer_name": (doc.to_dict() or {}).get("name", ""),
                "phone": (doc.to_dict() or {}).get("phone", ""),
                "parsed_tag": parsed.get("tag", ""),
            }
    elif action == "query_contract_end":
        ok, text, ctx = query_contract_end_text(fields)
        result = {"handled": True, "ok": ok, "reply_text": text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get("target_type", "")
        if action == "buyer_followup":
            target_type = "buyer"
        elif action == "seller_followup":
            target_type = "seller"
        elif action == "development_followup":
            target_type = "development"

        if target_type not in ("buyer", "seller", "development"):
            result = {"handled": True, "ok": False, "reply_text": "請提供對象：客需、委託 或 開發"}
        else:
            doc = find_customer_record(
                target_type=target_type,
                record_id=fields.get("record_id", ""),
                phone=fields.get("phone", ""),
                name=fields.get("name", ""),
                address=fields.get("address", ""),
            )
            if not doc:
                result = {"handled": True, "ok": False, "reply_text": "找不到唯一資料，請補客戶ID、正確電話或地址"}
            else:
                collection_name = "buyers" if target_type == "buyer" else ("sellers" if target_type == "seller" else "developments")
                doc_ref = db.collection(collection_name).document(doc.id)
                labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(fields.get("labels")))
                summary_parts = []
                if fields.get("content"):
                    summary_parts.append(fields["content"])
                if fields.get("address") and target_type != "development":
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get("price") and target_type != "development":
                    summary_parts.append(f"價格：{fields['price']}")
                if fields.get("registered_address") and target_type == "development":
                    summary_parts.append(f"戶籍地址：{fields['registered_address']}")
                summary_text = build_line_summary("；".join(summary_parts).strip() or "LINE 更新", event)

                update_kwargs = {
                    "target_type": target_type,
                    "doc_ref": doc_ref,
                    "content": summary_text,
                    "labels": labels,
                    "stage": normalize_development_status(fields.get("stage", "")) if target_type == "development" else fields.get("stage", ""),
                    "source": fields.get("source", "LINE"),
                    "event": event,
                }
                if target_type == "development":
                    update_kwargs["registered_address"] = fields.get("registered_address", "")
                    update_kwargs["extra_updates"] = {
                        "address": fields.get("address", ""),
                        "url": fields.get("url", ""),
                    }

                update_customer_note_and_labels(**update_kwargs)

                if action in ("buyer_followup", "seller_followup", "development_followup", "classify"):
                    add_customer_followup(
                        target_type=target_type,
                        customer_id=doc.id,
                        content=summary_text,
                        next_action=fields.get("next_action", ""),
                        next_contact_date=fields.get("next_contact_date", ""),
                        labels=labels,
                        line_event=event,
                        stage=normalize_development_status(fields.get("stage", "")) if target_type == "development" else "",
                        registered_address=fields.get("registered_address", "") if target_type == "development" else "",
                    )

                current_data = doc_ref.get().to_dict() or {}
                label_text = "客需" if target_type == "buyer" else ("委託" if target_type == "seller" else "開發")
                result = {
                    "handled": True,
                    "ok": True,
                    "reply_text": f"已寫入{label_text}：{current_data.get('name', '')}",
                    "target_type": target_type,
                    "target_id": doc.id,
                    "customer_name": current_data.get("name", ""),
                    "phone": current_data.get("phone", ""),
                    "parsed_tag": parsed.get("tag", ""),
                }

    save_line_log(
        parsed,
        event,
        "success" if result.get("ok") else "failed",
        target_type=result.get("target_type", ""),
        target_id=result.get("target_id", ""),
        note=result.get("reply_text", ""),
        sender_display_name=sender_display_name,
    )

    if result.get("ok") and result.get("target_type") and result.get("target_id"):
        incoming_message_id = message.get("id", "")
        save_line_message_link(
            incoming_message_id,
            result.get("target_type", ""),
            result.get("target_id", ""),
            customer_name=result.get("customer_name", ""),
            phone=result.get("phone", ""),
            source_event=event,
        )

    return result



@app.route("/developments")
@login_required
def developments():
    q = request.args.get("q", "").strip()
    current_stage = request.args.get("current_stage", "").strip()
    next_action = request.args.get("next_action", "").strip()
    source = request.args.get("source", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")
    show_done = request.args.get("show_done", "").strip()

    docs = db.collection("developments").stream()
    all_items = [doc_to_dict(d) for d in docs]
    total_count = len(all_items)
    items = list(all_items)
    source_options = sorted({(x.get("source") or "").strip() for x in all_items if (x.get("source") or "").strip()})

    if q:
        items = [
            x for x in items
            if q in (x.get("name") or "")
            or q in (x.get("phone") or "")
            or q in (x.get("address") or "")
            or q in (x.get("registered_address") or "")
        ]

    if current_stage:
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") == current_stage]
    if next_action:
        items = [x for x in items if (x.get("next_action") or "") == next_action]
    if source:
        items = [x for x in items if (x.get("source") or "") == source]
    if show_done != "1":
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") not in DEVELOPMENT_HIDDEN_BY_DEFAULT]

    def parse_created_at(x):
        return x.get("created_at") or ""

    if sort_by == "created_at_asc":
        items.sort(key=parse_created_at)
    elif sort_by == "created_at_desc":
        items.sort(key=parse_created_at, reverse=True)
    elif sort_by == "name_asc":
        items.sort(key=lambda x: (x.get("name") or ""))
    elif sort_by == "name_desc":
        items.sort(key=lambda x: (x.get("name") or ""), reverse=True)
    else:
        items.sort(key=parse_created_at, reverse=True)

    return render_template(
        "developments.html",
        developments=items,
        q=q,
        current_stage=current_stage,
        next_action=next_action,
        source=source,
        source_options=source_options,
        show_done=show_done,
        sort_by=sort_by,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
        total_count=total_count,
        filtered_count=len(items),
        label_docx_enabled=(next_action == "寄開發信"),
        label_docx_count=len([x for x in items if (x.get("registered_address") or "").strip()]),
    )


@app.route("/developments/new", methods=["POST"])
@login_required
def developments_new():
    form = request.form
    current_stage = normalize_development_status(form.get("current_stage", "").strip() or form.get("stage", "").strip() or "待聯繫")
    next_action = normalize_development_next_action(form.get("next_action", "").strip())
    next_action_date = form.get("next_action_date", "").strip()

    _manual_url = form.get("url", "").strip()
    data = {
        "name": form.get("name", "").strip() or "未填姓名",
        "phone": form.get("phone", "").strip(),
        "source": infer_development_source(form.get("source", "").strip(), _manual_url),
        "address": form.get("address", "").strip(),
        "registered_address": form.get("registered_address", "").strip(),
        "url": _manual_url,
        "current_stage": current_stage,
        "current_stage": current_stage,
        "stage": current_stage,
        "next_action": next_action,
        "next_action_date": next_action_date,
        "note": form.get("note", "").strip(),
        "record_date": now_taipei().strftime("%Y-%m-%d"),
        "created_at": now_taipei().isoformat(),
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id"),
        "updated_by_name": session.get("user_name"),
    }
    if not (data["name"] or data["phone"] or data["address"] or data["url"]):
        flash("至少請填姓名、電話、地址、網址其中一項", "danger")
        return redirect(url_for("developments"))

    doc_ref = db.collection("developments").document()
    doc_ref.set(data)
    if data["note"]:
        db.collection("development_followups").add({
            "development_id": doc_ref.id,
            "contact_time": now_taipei().strftime("%Y-%m-%d %H:%M"),
            "channel": "手動新增",
            "current_stage": current_stage,
            "stage": current_stage,
            "next_action": next_action,
            "next_action_date": next_action_date,
            "registered_address": data["registered_address"],
            "content": data["note"],
            "next_contact_date": next_action_date,
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id"),
            "created_by_name": session.get("user_name"),
            "sender_display_name": session.get("user_name"),
        })

    flash("已新增開發", "success")
    return redirect(url_for("developments"))


@app.route("/developments/<development_id>")
@login_required
def development_detail(development_id):
    doc = db.collection("developments").document(development_id).get()
    if not doc.exists:
        flash("找不到這筆開發", "danger")
        return redirect(url_for("developments"))

    development = doc_to_dict(doc)
    followups_ref = db.collection("development_followups").where("development_id", "==", development_id)
    followups = [doc_to_dict(f) for f in followups_ref.stream()]
    followups.sort(key=lambda x: x.get("contact_time", ""), reverse=True)
    return render_template(
        "development_detail.html",
        development=development,
        followups=followups,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )


@app.route("/developments/<development_id>/quick-flow", methods=["POST"])
@app.route("/developments/<development_id>/quick-stage", methods=["POST"])
@login_required
def development_quick_flow(development_id):
    current_stage = normalize_development_status(request.form.get("current_stage", "").strip() or request.form.get("stage", "").strip())
    next_action = normalize_development_next_action(request.form.get("next_action", "").strip())
    next_action_date = request.form.get("next_action_date", "").strip()

    if not current_stage and not next_action and not next_action_date:
        flash("請至少調整一個欄位", "warning")
        return redirect(request.referrer or url_for("developments"))

    updates = {
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id"),
        "updated_by_name": session.get("user_name"),
    }
    if current_stage:
        updates["current_stage"] = current_stage
        updates["stage"] = current_stage
    if next_action:
        updates["next_action"] = next_action
    if next_action_date:
        updates["next_action_date"] = next_action_date

    db.collection("developments").document(development_id).update(updates)
    flash("已更新開發流程", "success")
    return redirect(request.referrer or url_for("developments"))


@app.route("/developments/<development_id>/followup", methods=["POST"])
@login_required
def add_development_followup(development_id):
    contact_time = request.form.get("contact_time", "").strip() or now_taipei().strftime("%Y-%m-%d %H:%M")
    channel = request.form.get("channel", "").strip()
    current_stage = normalize_development_status(request.form.get("current_stage", "").strip() or request.form.get("stage", "").strip())
    next_action = normalize_development_next_action(request.form.get("next_action", "").strip())
    next_action_date = request.form.get("next_action_date", "").strip() or request.form.get("next_contact_date", "").strip()
    registered_address = request.form.get("registered_address", "").strip()
    content = request.form.get("content", "").strip()
    note_extra = request.form.get("note", "").strip()
    if note_extra:
        content = (content + ("\n" if content else "") + note_extra).strip()

    now = now_taipei().isoformat()
    db.collection("development_followups").add({
        "development_id": development_id,
        "contact_time": contact_time,
        "channel": channel,
        "current_stage": current_stage,
        "current_stage": current_stage,
        "stage": current_stage,
        "next_action": next_action,
        "next_action_date": next_action_date,
        "registered_address": registered_address,
        "content": content,
        "next_contact_date": next_action_date,
        "created_at": now,
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
        "sender_display_name": session.get("user_name"),
    })

    updates = {
        "updated_at": now,
        "updated_by_id": session.get("user_id"),
        "updated_by_name": session.get("user_name"),
    }
    if current_stage:
        updates["current_stage"] = current_stage
        updates["stage"] = current_stage
    if next_action:
        updates["next_action"] = next_action
    if next_action_date:
        updates["next_action_date"] = next_action_date
    if registered_address:
        updates["registered_address"] = registered_address

    ref = db.collection("developments").document(development_id)
    snap = ref.get()
    current = snap.to_dict() or {}
    if content:
        updates["note"] = append_note_block(current.get("note", ""), content, session.get("user_name") or "後台")
    ref.update(updates)

    flash("已新增開發追蹤紀錄", "success")
    return redirect(url_for("development_detail", development_id=development_id))


@app.route("/developments/<development_id>/edit", methods=["GET", "POST"])
@login_required
def development_edit(development_id):
    doc_ref = db.collection("developments").document(development_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆開發", "danger")
        return redirect(url_for("developments"))

    development = doc_to_dict(doc)
    if request.method == "POST":
        form = request.form
        current_stage = normalize_development_status(form.get("current_stage", "").strip() or form.get("stage", "").strip())
        next_action = normalize_development_next_action(form.get("next_action", "").strip())
        next_action_date = form.get("next_action_date", "").strip()
        updated = {
            "name": form.get("name", "").strip() or "未填姓名",
            "phone": form.get("phone", "").strip(),
            "source": form.get("source", "").strip(),
            "address": form.get("address", "").strip(),
            "registered_address": form.get("registered_address", "").strip(),
            "url": form.get("url", "").strip(),
            "current_stage": current_stage,
            "stage": current_stage,
            "next_action": next_action,
            "next_action_date": next_action_date,
            "note": form.get("note", "").strip(),
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": session.get("user_id"),
            "updated_by_name": session.get("user_name"),
        }
        doc_ref.update(updated)
        flash("已更新開發資料", "success")
        return redirect(url_for("development_detail", development_id=development_id))

    return render_template(
        "development_edit.html",
        development=development,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )


@app.route("/developments/<development_id>/delete", methods=["POST"])
@login_required
def development_delete(development_id):
    delete_by_field("development_followups", "development_id", development_id)
    db.collection("developments").document(development_id).delete()
    flash("已刪除開發與相關追蹤紀錄", "info")
    return redirect(url_for("developments"))


@app.route("/developments/<development_id>/followup/<followup_id>/edit", methods=["GET", "POST"])
@login_required
def development_followup_edit(development_id, followup_id):
    doc_ref = db.collection("development_followups").document(followup_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆追蹤紀錄", "danger")
        return redirect(url_for("development_detail", development_id=development_id))

    followup = doc_to_dict(doc)
    if request.method == "POST":
        contact_time = request.form.get("contact_time", "").strip() or now_taipei().strftime("%Y-%m-%d %H:%M")
        channel = request.form.get("channel", "").strip()
        current_stage = normalize_development_status(request.form.get("current_stage", "").strip() or request.form.get("stage", "").strip())
        next_action = normalize_development_next_action(request.form.get("next_action", "").strip())
        next_action_date = request.form.get("next_action_date", "").strip() or request.form.get("next_contact_date", "").strip()
        registered_address = request.form.get("registered_address", "").strip()
        content = request.form.get("content", "").strip()

        doc_ref.update({
            "contact_time": contact_time,
            "channel": channel,
            "current_stage": current_stage,
            "stage": current_stage,
            "next_action": next_action,
            "next_action_date": next_action_date,
            "registered_address": registered_address,
            "content": content,
            "next_contact_date": next_action_date,
        })

        updates = {
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": session.get("user_id"),
            "updated_by_name": session.get("user_name"),
        }
        if current_stage:
            updates["current_stage"] = current_stage
            updates["stage"] = current_stage
        if next_action:
            updates["next_action"] = next_action
        if next_action_date:
            updates["next_action_date"] = next_action_date
        if registered_address:
            updates["registered_address"] = registered_address
        db.collection("developments").document(development_id).update(updates)

        flash("已更新追蹤紀錄", "success")
        return redirect(url_for("development_detail", development_id=development_id))

    return render_template(
        "development_followup_edit.html",
        development_id=development_id,
        followup=followup,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )


@app.route("/developments/<development_id>/followup/<followup_id>/delete", methods=["POST"])
@login_required
def development_followup_delete(development_id, followup_id):
    db.collection("development_followups").document(followup_id).delete()
    flash("已刪除追蹤紀錄", "info")
    return redirect(url_for("development_detail", development_id=development_id))




def _filter_development_items_from_args(args):
    q = (args.get("q") or "").strip()
    current_stage = (args.get("current_stage") or "").strip()
    next_action = (args.get("next_action") or "").strip()
    source = (args.get("source") or "").strip()
    sort_by = (args.get("sort_by") or "created_at_desc").strip()
    show_done = (args.get("show_done") or "").strip()

    items = [doc_to_dict(d) for d in db.collection("developments").stream()]
    if q:
        items = [
            x for x in items
            if q in (x.get("name") or "")
            or q in (x.get("phone") or "")
            or q in (x.get("address") or "")
            or q in (x.get("registered_address") or "")
        ]
    if current_stage:
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") == current_stage]
    if next_action:
        items = [x for x in items if (x.get("next_action") or "") == next_action]
    if source:
        items = [x for x in items if (x.get("source") or "") == source]
    if show_done != "1":
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") not in DEVELOPMENT_HIDDEN_BY_DEFAULT]

    def parse_created_at(x):
        return x.get("created_at") or ""

    if sort_by == "created_at_asc":
        items.sort(key=parse_created_at)
    elif sort_by == "created_at_desc":
        items.sort(key=parse_created_at, reverse=True)
    elif sort_by == "name_asc":
        items.sort(key=lambda x: (x.get("name") or ""))
    elif sort_by == "name_desc":
        items.sort(key=lambda x: (x.get("name") or ""), reverse=True)
    else:
        items.sort(key=parse_created_at, reverse=True)
    return items



def _build_label_recipient_text(item) -> str:
    """
    第二行優先顯示使用者輸入的標題，沒有的話才退回姓名。
    不再自動補「負責人收」。
    """
    if isinstance(item, dict):
        return (item.get("label_title") or item.get("name") or "").strip()
    return (item or "").strip()


def _build_development_label_rows(items):
    rows = []
    for item in items:
        # 開發列表目前有顯示的資料都要能印。
        # 優先用戶籍地址；若沒有戶籍地址，就退回一般地址。
        address = (item.get("registered_address") or item.get("address") or "").strip()
        if not address:
            continue
        rows.append({
            "address": address,
            "recipient": _build_label_recipient_text(item),
        })
    return rows


def _find_development_label_template_path():
    """
    以使用者提供的 Word 標籤模板為準，直接套用原始版型，
    讓字體、字級、段落間距、表格尺寸都跟模板完全一致。
    """
    candidates = []
    env_path = (os.environ.get("DEVELOPMENT_LABEL_TEMPLATE_PATH") or "").strip()
    if env_path:
        candidates.append(env_path)

    candidates.extend([
        os.path.join(BASE_DIR, "空白標籤輸入表.docx"),
        os.path.join(BASE_DIR, "development_label_template.docx"),
        "空白標籤輸入表.docx",
        "development_label_template.docx",
    ])

    for path in candidates:
        if path and os.path.exists(path):
            return path
    return ""


def _clear_paragraph_text_preserve_format(paragraph):
    for run in paragraph.runs:
        run.text = ""


def _write_text_into_template_paragraph(paragraph, text: str):
    text = text or ""
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _fill_label_cell_like_template(cell, address: str, recipient: str):
    """
    直接沿用模板 cell 內原有的段落結構：
    第 1 段空白
    第 2 段地址
    第 3 段空白
    第 4 段標題/姓名
    這樣字型、字級、粗細、每段之間的距離都會維持跟模板一致。
    """
    while len(cell.paragraphs) < 4:
        cell.add_paragraph()

    for p in cell.paragraphs:
        _clear_paragraph_text_preserve_format(p)

    _write_text_into_template_paragraph(cell.paragraphs[1], (address or "").strip())
    _write_text_into_template_paragraph(cell.paragraphs[3], (recipient or "").strip())


def _render_development_labels_docx(rows):
    """
    使用使用者提供的空白 Word 標籤模板直接填值，
    讓輸出格式與原模板完全一致（字體 / 字級 / 行距 / 欄距 / 表格尺寸）。
    """
    if Document is None:
        raise RuntimeError("python-docx 未安裝")

    template_path = _find_development_label_template_path()
    if not template_path:
        raise RuntimeError("找不到標籤模板，請把『空白標籤輸入表.docx』放在專案根目錄")

    doc = Document(template_path)
    body = doc._body._element

    labels_per_page = 16
    needed_pages = max(1, (len(rows) + labels_per_page - 1) // labels_per_page)

    if not doc.tables:
        raise RuntimeError("標籤模板內沒有表格，無法產生標籤")

    # 模板結構：每頁 1 個 8x2 表格，表格後接一個分頁段落
    body_children = list(body)
    if len(body_children) < 2:
        raise RuntimeError("標籤模板結構異常，請重新提供模板檔")

    proto_table = deepcopy(body_children[0])
    proto_separator = deepcopy(body_children[1])

    current_pages = len(doc.tables)
    while current_pages < needed_pages:
        insert_at = max(0, len(body) - 2)   # 保留最後的空白段落與 sectPr
        body.insert(insert_at, deepcopy(proto_table))
        body.insert(insert_at + 1, deepcopy(proto_separator))
        current_pages += 1

    # 只保留實際需要的頁數，其餘模板頁刪掉，避免印出空白頁
    keep_prefix_count = needed_pages * 2
    for el in list(body)[keep_prefix_count:-2]:
        body.remove(el)

    all_cells = []
    for table in doc.tables[:needed_pages]:
        for row in table.rows:
            for cell in row.cells:
                all_cells.append(cell)

    for idx, cell in enumerate(all_cells):
        if idx < len(rows):
            row = rows[idx]
            _fill_label_cell_like_template(
                cell,
                row.get("address", ""),
                row.get("recipient", ""),
            )
        else:
            _fill_label_cell_like_template(cell, "", "")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


@app.route("/developments/labels-docx")
@login_required
def development_labels_docx():
    args = request.args.to_dict(flat=True)
    if not (args.get("next_action") or "").strip():
        args["next_action"] = "寄開發信"
    items = _filter_development_items_from_args(args)
    rows = _build_development_label_rows(items)
    if not rows:
        flash("目前沒有可列印的寄開發信標籤資料（需有戶籍地址或地址）", "warning")
        return redirect(request.referrer or url_for("developments", **args))
    try:
        bio = _render_development_labels_docx(rows)
    except Exception as e:
        flash(f"產生 Word 標籤失敗：{e}", "danger")
        return redirect(request.referrer or url_for("developments", **args))
    filename = f"開發信標籤_{now_taipei().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.route("/developments/download")
@login_required
def download_developments():
    docs = db.collection("developments").stream()
    rows = [doc_to_dict(d) for d in docs]
    si = StringIO()
    writer = csv.writer(si)
    writer.writerow(["id","日期","姓名","電話","地址","戶籍地址","網址","目前狀態","下一動作","下次日期","來源","內部備註","建立時間","建立者"])
    for r in rows:
        writer.writerow([
            r.get("id",""),
            r.get("record_date",""),
            r.get("name",""),
            r.get("phone",""),
            r.get("address",""),
            r.get("registered_address",""),
            r.get("url",""),
            r.get("current_stage","") or r.get("stage",""),
            r.get("next_action",""),
            r.get("next_action_date",""),
            r.get("source",""),
            r.get("note",""),
            r.get("created_at",""),
            r.get("created_by_name",""),
        ])
    csv_data = "\ufeff" + si.getvalue()
    response = Response(csv_data, mimetype="text/csv; charset=utf-8")
    response.headers["Content-Disposition"] = "attachment; filename=developments.csv"
    return response

# ========= 開發自由格式解析（新版覆蓋） =========
DEVELOPMENT_SOURCE_SELLER_SELF = "屋主自售/踩線"
DEVELOPMENT_SOURCE_STREET = "掃街"

def infer_development_source(explicit_source: str, url: str) -> str:
    explicit = (explicit_source or "").strip()
    if explicit in (DEVELOPMENT_SOURCE_STREET, DEVELOPMENT_SOURCE_SELLER_SELF):
        return explicit
    if explicit and explicit.upper() != "LINE":
        return explicit
    return DEVELOPMENT_SOURCE_SELLER_SELF if (url or "").strip() else DEVELOPMENT_SOURCE_STREET

def normalize_line_key(key: str):
    k = (key or "").strip().replace(" ", "")
    mapping = {
        "對象": "target_type",
        "類型": "target_type",
        "目標": "target_type",
        "客戶類型": "target_type",
        "電話": "phone",
        "手機": "phone",
        "電話號碼": "phone",
        "姓名": "name",
        "客戶": "name",
        "買方": "name",
        "賣方": "name",
        "屋主": "name",
        "客戶來源": "source",
        "來源": "source",
        "需求類型": "intent_type_raw",
        "委託類型": "deal_type_raw",
        "ID": "record_id",
        "客戶ID": "record_id",
        "buyer_id": "record_id",
        "seller_id": "record_id",
        "development_id": "record_id",
        "內容": "content",
        "紀錄": "content",
        "備註": "content",
        "說明": "content",
        "進度內容": "content",
        "進程": "current_stage",
        "階段": "current_stage",
        "狀態": "current_stage",
        "進度": "current_stage",
        "目前狀況": "current_stage",
        "目前狀態": "current_stage",
        "標籤": "labels",
        "分類": "labels",
        "labels": "labels",
        "下一步": "next_action",
        "下次行動": "next_action",
        "next_action": "next_action",
        "下一次時間": "next_action_date",
        "下次時間": "next_action_date",
        "下次日期": "next_action_date",
        "下次聯絡日": "next_action_date",
        "下次聯絡日期": "next_action_date",
        "next_contact_date": "next_action_date",
        "地址": "address",
        "物件": "address",
        "案名": "address",
        "戶籍地址": "registered_address",
        "戶籍地": "registered_address",
        "網址": "url",
        "連結": "url",
        "網站": "url",
        "日期": "record_date",
        "總價": "price",
        "價格": "price",
        "預算": "budget",
        "區域": "preferred_areas",
        "產品類型": "property_type",
        "房數": "room_range",
        "車位": "car_need",
        "開價": "expected_price",
        "底價": "min_price",
        "委託到期日": "contract_end_date",
        "筆數": "limit",
    }
    return mapping.get(k, k)

def _looks_like_url(line: str) -> bool:
    s = (line or "").strip()
    return bool(re.search(r'(https?://|www\.|591\.com|rakuya|house|chyi\.com)', s, re.I))

def _extract_first_url(line: str) -> str:
    m = re.search(r'(https?://\S+|www\.\S+)', line or '', re.I)
    return m.group(1).strip() if m else ''

def _looks_like_phone(line: str) -> bool:
    s = re.sub(r'\s+', '', line or '')
    return bool(re.search(r'(09\d{2}[- ]?\d{3}[- ]?\d{3}|0\d{1,2}[- ]?\d{6,8})', s))

def _extract_first_phone(line: str) -> str:
    m = re.search(r'(09\d{2}[- ]?\d{3}[- ]?\d{3}|0\d{1,2}[- ]?\d{6,8})', line or '')
    return m.group(1).strip() if m else ''

def _looks_like_address(line: str) -> bool:
    s = (line or '').strip()
    if not s:
        return False
    patterns = [
        r'[台臺]中市?.{0,15}(區|里|路|街|段|巷|弄|號)',
        r'.+(路|街|大道|段).*(巷|弄|號)',
        r'.+(路|街|大道|段).*(號)',
        r'.+(區).*(路|街|段|巷|弄|號)',
    ]
    return any(re.search(p, s) for p in patterns)

def _looks_like_name(line: str) -> bool:
    s = (line or '').strip()
    if not s or len(s) > 12:
        return False
    if _looks_like_url(s) or _looks_like_phone(s) or _looks_like_address(s):
        return False
    bad = ['目前狀況','目前狀態','下一步','下次','網址','地址','戶籍地址','備註','來源']
    if any(b in s for b in bad):
        return False
    return True

def _extract_date_text(line: str) -> str:
    s = (line or '').strip()
    m = re.search(r'(\d{4}-\d{1,2}-\d{1,2})', s)
    if m:
        return m.group(1)
    m = re.search(r'(\d{1,2}/\d{1,2})', s)
    if m:
        mm, dd = m.group(1).split('/')
        year = now_taipei().year
        return f"{year}-{int(mm):02d}-{int(dd):02d}"
    return ''

def _find_status_in_text(text: str) -> str:
    options = sorted(DEVELOPMENT_STATUS_OPTIONS, key=len, reverse=True)
    raw = text or ''
    for opt in options:
        if opt in raw:
            return opt
    aliases = {
        '已寄開發信': '已寄開發信待跑開發',
        '待聯絡': '待聯繫',
        '已聯絡': '已聯繫',
        '已連繫': '已聯繫',
        '再聯絡': '持續追蹤',
        '再聯繫': '持續追蹤',
    }
    for k, v in aliases.items():
        if k in raw:
            return v
    return ''

def _find_next_action_in_text(text: str) -> str:
    options = sorted(DEVELOPMENT_NEXT_ACTION_OPTIONS, key=len, reverse=True)
    raw = text or ''
    for opt in options:
        if opt in raw:
            return opt
    aliases = {
        '電話聯絡': '電話聯繫',
        'LINE聯絡': 'LINE聯繫',
        '再次聯絡': '再次聯繫',
    }
    for k, v in aliases.items():
        if k in raw:
            return v
    return ''

def parse_flexible_development_chunk(raw_text: str) -> dict:
    text = (raw_text or '').replace('\r\n', '\n').strip()
    if not text:
        return {}

    fields = {}
    notes = []
    candidate_names = []

    for raw_line in text.split('\n'):
        line = raw_line.strip()
        if not line:
            continue

        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if m:
            key = normalize_line_key(m.group(1))
            value = (m.group(2) or '').strip()
            if key == 'labels':
                fields[key] = parse_label_csv(value)
            elif value:
                fields[key] = value
            continue

        url = _extract_first_url(line)
        phone = _extract_first_phone(line)
        date_txt = _extract_date_text(line)

        if url and not fields.get('url'):
            fields['url'] = url
            leftover = line.replace(url, '').strip()
            if leftover:
                notes.append(leftover)
            continue

        if phone and not fields.get('phone'):
            fields['phone'] = phone
            leftover = line.replace(phone, '').strip(' ：:')
            if leftover and _looks_like_name(leftover):
                candidate_names.append(leftover)
            elif leftover:
                notes.append(leftover)
            continue

        if _looks_like_address(line):
            if not fields.get('address'):
                fields['address'] = line
            elif not fields.get('registered_address') and ('戶籍' in line or '住址' in line):
                fields['registered_address'] = line
            else:
                notes.append(line)
            continue

        if date_txt and not fields.get('next_action_date') and any(word in line for word in ['下次','再','聯絡','時間','日期']):
            fields['next_action_date'] = date_txt
            leftover = line.replace(date_txt, '').strip(' ：:')
            if leftover and leftover not in ('下次','下次時間','下一次時間','下次日期'):
                notes.append(leftover)
            continue

        if _looks_like_name(line):
            candidate_names.append(line)
            continue

        notes.append(line)

    if not fields.get('name') and candidate_names:
        fields['name'] = candidate_names[0]

    joined_notes = '\n'.join(notes).strip()
    if joined_notes:
        fields['content'] = ((fields.get('content', '') + '\n' + joined_notes).strip() if fields.get('content') else joined_notes)

    # 自動從整段文字抓狀態 / 下一步 / 日期
    big_text = text
    if not fields.get('current_stage'):
        auto_stage = _find_status_in_text(big_text)
        fields['current_stage'] = normalize_development_status(auto_stage or '待聯繫')
    else:
        fields['current_stage'] = normalize_development_status(fields.get('current_stage'))

    if not fields.get('next_action'):
        fields['next_action'] = normalize_development_next_action(_find_next_action_in_text(big_text))
    else:
        fields['next_action'] = normalize_development_next_action(fields.get('next_action'))

    if not fields.get('next_action_date'):
        auto_date = _extract_date_text(big_text)
        if auto_date:
            fields['next_action_date'] = auto_date

    fields['source'] = infer_development_source(fields.get('source', ''), fields.get('url', ''))
    if not fields.get('record_date'):
        fields['record_date'] = now_taipei().strftime('%Y-%m-%d')

    has_any = any(fields.get(k) for k in ['name','phone','address','url','content'])
    return fields if has_any else {}

def parse_development_batch_message(raw_text: str):
    text = (raw_text or '').replace('\r\n', '\n').strip()
    lines = text.split('\n')
    if lines and lines[0].strip().startswith('#'):
        lines = lines[1:]
    body = '\n'.join(lines).strip()
    if not body:
        return []

    if re.search(r'(?m)^/{3,}\s*$', body):
        chunks = [c.strip() for c in re.split(r'(?m)^/{3,}\s*$', body) if c.strip()]
    elif re.search(r'(?m)^-{3,}\s*$', body):
        chunks = [c.strip() for c in re.split(r'(?m)^-{3,}\s*$', body) if c.strip()]
    else:
        # 沒有分隔線時，嘗試依地址行切段
        lines2 = [ln.rstrip() for ln in body.split('\n')]
        chunks, current = [], []
        seen_first = False
        for ln in lines2:
            s = ln.strip()
            if not s:
                continue
            if _looks_like_address(s) and seen_first and current:
                chunks.append('\n'.join(current).strip())
                current = [s]
            else:
                current.append(s)
                if _looks_like_address(s):
                    seen_first = True
        if current:
            chunks.append('\n'.join(current).strip())

    items = []
    for chunk in chunks:
        fields = parse_flexible_development_chunk(chunk)
        if fields:
            items.append(fields)
    return items

def parse_line_formatted_message(text: str):
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    if not lines:
        return None

    first = lines[0]
    if not first.startswith("#"):
        return None

    tag = first.lstrip("#").strip()
    tag_map = {
        "新增客需": "create_buyer_need",
        "新增委託": "create_seller_listing",
        "新增開發": "create_development",
        "新增開發批次": "create_development_batch",
        "開發追蹤": "development_followup",
        "買方追蹤": "buyer_followup",
        "賣方追蹤": "seller_followup",
        "客戶分類": "classify",
        "查詢紀錄": "query_records",
        "查詢委託到期": "query_contract_end",
        "帶看": "buyer_followup",
        "成交": "buyer_followup",
        "委託": "seller_followup",
        "紀錄": "generic_note",
    }
    action = tag_map.get(tag)
    if not action:
        return None

    raw_body = "\n".join(lines[1:]).strip()

    if action == "create_development_batch":
        return {
            "tag": tag,
            "action": action,
            "fields": {},
            "raw_text": text,
            "raw_body": raw_body,
        }

    if action == "create_development":
        fields = parse_flexible_development_chunk(raw_body)
        if not fields:
            return None
        return {
            "tag": tag,
            "action": action,
            "fields": fields,
            "raw_text": text,
            "raw_body": raw_body,
        }

    fields = {}
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = m.group(2).strip()
        if key == "labels":
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value

    if tag in ("買方追蹤", "帶看", "成交") and not fields.get("target_type"):
        fields["target_type"] = "buyer"
    if tag in ("賣方追蹤", "委託") and not fields.get("target_type"):
        fields["target_type"] = "seller"
    if tag == "開發追蹤" and not fields.get("target_type"):
        fields["target_type"] = "development"

    fields["target_type"] = normalize_target_type(fields.get("target_type", "")) or fields.get("target_type", "")
    fields["intent_type"] = normalize_intent_type(fields.get("intent_type_raw", ""), fields)
    fields["deal_type"] = normalize_deal_type(fields.get("deal_type_raw", ""))
    fields["limit"] = parse_int_limit(fields.get("limit", 10), default=10, max_value=30)

    if action == "create_buyer_need":
        if not (fields.get("name") and fields.get("phone")):
            return None
    elif action == "create_seller_listing":
        if not (fields.get("name") and fields.get("phone")):
            return None
    elif action in ("buyer_followup", "seller_followup", "classify", "query_records", "query_contract_end", "development_followup"):
        if not (fields.get("record_id") or fields.get("phone") or fields.get("name") or fields.get("address")):
            return None

    return {
        "tag": tag,
        "action": action,
        "fields": fields,
        "raw_text": text,
        "raw_body": raw_body,
    }

def find_development_record(record_id: str = "", phone: str = "", name: str = "", address: str = ""):
    col = db.collection("developments")
    if record_id:
        doc = col.document(record_id).get()
        if doc.exists:
            return doc
    if phone:
        matches = find_records_by_phone("developments", phone)
        if len(matches) == 1:
            return matches[0]
    if name:
        docs = list(col.where("name", "==", name.strip()).limit(2).stream())
        if len(docs) == 1:
            return docs[0]
    if address:
        docs = list(col.where("address", "==", address.strip()).limit(2).stream())
        if len(docs) == 1:
            return docs[0]
    return None

def create_development(fields, event):
    phone = fields.get("phone", "").strip()
    name = fields.get("name", "").strip() or "未填姓名"
    url = fields.get("url", "").strip()
    source = infer_development_source(fields.get("source", ""), url)
    address = fields.get("address", "").strip()

    matches = find_records_by_phone("developments", phone) if phone else []
    if not matches and address:
        doc = find_development_record(address=address)
        if doc:
            matches = [doc]

    labels = build_development_labels(fields.get("labels"))
    content_text = fields.get("content", "").strip() or address or url or "LINE 新增開發"
    note_content = build_line_summary(content_text, event)

    payload = {
        "name": name,
        "phone": phone,
        "source": source,
        "url": url,
        "address": address,
        "registered_address": fields.get("registered_address", "").strip(),
        "current_stage": normalize_development_status(fields.get("current_stage", "").strip() or fields.get("stage", "").strip() or "待聯繫"),
        "stage": normalize_development_status(fields.get("current_stage", "").strip() or fields.get("stage", "").strip() or "待聯繫"),
        "next_action": normalize_development_next_action(fields.get("next_action", "").strip()),
        "next_action_date": fields.get("next_action_date", "").strip() or fields.get("next_contact_date", "").strip(),
        "record_date": fields.get("record_date", "").strip() or now_taipei().strftime("%Y-%m-%d"),
        "note": "",
        "labels": labels,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
        "sender_display_name": get_line_sender_display_name(event) or "",
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection("developments").document(doc.id)
        update_customer_note_and_labels(
            target_type="development",
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=payload["stage"],
            source=payload["source"],
            event=event,
        )
        doc_ref.update({k: v for k, v in payload.items() if v != "" and k != "note"})
        add_customer_followup(
            target_type="development",
            customer_id=doc.id,
            content=note_content,
            next_action=payload.get("next_action", ""),
            next_contact_date=payload.get("next_action_date", ""),
            labels=labels,
            line_event=event,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已註記開發：{updated_doc.get('name', '')}",
            "target_type": "development",
            "target_id": doc.id,
            "customer_name": updated_doc.get("name", ""),
            "phone": updated_doc.get("phone", ""),
            "parsed_tag": "新增開發",
        }

    if len(matches) > 1:
        return {
            "handled": True,
            "ok": False,
            "reply_text": "未寫入：同電話有多筆開發資料，請補地址或客戶ID",
        }

    now = now_taipei().isoformat()
    payload.update({
        "created_at": now,
        "created_by_id": "line_bot",
        "created_by_name": "LINE Bot",
        "note": append_note_block("", note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection("developments").document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type="development",
        customer_id=doc_ref.id,
        content=note_content,
        next_action=payload.get("next_action", ""),
        next_contact_date=payload.get("next_action_date", ""),
        labels=labels,
        line_event=event,
    )
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記開發：{name}",
        "target_type": "development",
        "target_id": doc_ref.id,
        "customer_name": name,
        "phone": phone,
        "parsed_tag": "新增開發",
    }

def create_development_batch(raw_text, event):
    items = parse_development_batch_message(raw_text)
    if not items:
        return {"handled": True, "ok": False, "reply_text": "未寫入：沒有解析到有效開發資料"}

    ok_count = 0
    fail_count = 0
    last_result = None
    for fields in items:
        result = create_development(fields, event)
        last_result = result
        if result.get("ok"):
            ok_count += 1
        else:
            fail_count += 1

    reply_text = f"批量註記完成：成功 {ok_count} 筆，失敗 {fail_count} 筆"
    if last_result and last_result.get("target_type") and last_result.get("target_id"):
        return {
            "handled": True,
            "ok": ok_count > 0,
            "reply_text": reply_text,
            "target_type": last_result.get("target_type", ""),
            "target_id": last_result.get("target_id", ""),
            "customer_name": last_result.get("customer_name", ""),
            "phone": last_result.get("phone", ""),
            "parsed_tag": "新增開發批次",
        }
    return {"handled": True, "ok": ok_count > 0, "reply_text": reply_text}

def add_development_followup_via_line(fields, event):
    doc = find_development_record(
        record_id=fields.get("record_id", ""),
        phone=fields.get("phone", ""),
        name=fields.get("name", ""),
        address=fields.get("address", ""),
    )
    if not doc:
        return {"handled": True, "ok": False, "reply_text": "找不到唯一開發資料，請補電話、地址或客戶ID"}

    content = (fields.get("content") or "").strip() or "LINE 更新"
    current_stage = normalize_development_status(fields.get("current_stage", "").strip() or fields.get("stage", "").strip())
    next_action = normalize_development_next_action(fields.get("next_action", "").strip())
    next_action_date = fields.get("next_action_date", "").strip() or fields.get("next_contact_date", "").strip()
    registered_address = fields.get("registered_address", "").strip()
    source = infer_development_source(fields.get("source",""), fields.get("url",""))

    doc_ref = db.collection("developments").document(doc.id)
    update_customer_note_and_labels(
        target_type="development",
        doc_ref=doc_ref,
        content=build_line_summary(content, event),
        labels=build_development_labels(fields.get("labels")),
        stage=current_stage,
        source=source or None,
        event=event,
    )
    updates = {
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }
    if current_stage:
        updates["current_stage"] = current_stage
        updates["stage"] = current_stage
    if next_action:
        updates["next_action"] = next_action
    if next_action_date:
        updates["next_action_date"] = next_action_date
    if registered_address:
        updates["registered_address"] = registered_address
    if source:
        updates["source"] = source
    doc_ref.update(updates)

    add_customer_followup(
        target_type="development",
        customer_id=doc.id,
        content=build_line_summary(content, event),
        next_action=next_action,
        next_contact_date=next_action_date,
        labels=build_development_labels(fields.get("labels")),
        line_event=event,
    )

    data = doc_ref.get().to_dict() or {}
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記開發：{data.get('name', '')}",
        "target_type": "development",
        "target_id": doc.id,
        "customer_name": data.get("name", ""),
        "phone": data.get("phone", ""),
        "parsed_tag": "開發追蹤",
    }

def _maybe_freeform_development(text: str):
    raw = (text or '').strip()
    if not raw:
        return None
    body = raw
    if raw.startswith('#'):
        return None
    if '///' in raw or re.search(r'(?m)^/{3,}\s*$', raw):
        items = parse_development_batch_message(raw)
        return {"batch": True, "items": items} if items else None
    fields = parse_flexible_development_chunk(raw)
    return {"batch": False, "fields": fields} if fields else None

def process_line_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    sender_display_name = get_line_sender_display_name(event)
    raw_text = message.get("text", "")

    parsed = parse_line_formatted_message(raw_text)
    if not parsed:
        free = _maybe_freeform_development(raw_text)
        if free:
            if free.get("batch"):
                result = create_development_batch(raw_text, event)
            else:
                result = create_development(free["fields"], event)
            save_line_log(
                {"tag": "自由格式開發", "action": "create_development_batch" if free.get("batch") else "create_development",
                 "fields": free.get("fields", {}), "raw_text": raw_text},
                event,
                "success" if result.get("ok") else "failed",
                target_type=result.get("target_type", ""),
                target_id=result.get("target_id", ""),
                note=result.get("reply_text", ""),
                sender_display_name=sender_display_name,
            )
            if result.get("ok") and result.get("target_type") and result.get("target_id"):
                incoming_message_id = message.get("id", "")
                save_line_message_link(
                    incoming_message_id,
                    result["target_type"],
                    result["target_id"],
                    tag=result.get("parsed_tag", "自由格式開發"),
                    action="create_development",
                    customer_name=result.get("customer_name", ""),
                    phone=result.get("phone", ""),
                    source_event=event,
                )
            return result

        quoted_result = process_quote_context_message(event)
        if quoted_result.get("handled"):
            return quoted_result
        return {"handled": False}

    fields = parsed["fields"]
    action = parsed["action"]

    if action == "create_buyer_need":
        result = create_buyer_need(fields, event)
    elif action == "create_seller_listing":
        result = create_seller_listing(fields, event)
    elif action == "create_development":
        result = create_development(fields, event)
    elif action == "create_development_batch":
        result = create_development_batch(parsed.get("raw_body") or raw_text, event)
    elif action == "development_followup":
        result = add_development_followup_via_line(fields, event)
    elif action == "query_records":
        target_type, doc = resolve_customer_record(fields)
        if not doc and (fields.get("record_id") or fields.get("phone") or fields.get("name") or fields.get("address")):
            doc = find_development_record(fields.get("record_id",""), fields.get("phone",""), fields.get("name",""), fields.get("address",""))
            if doc:
                target_type = "development"
        if not doc:
            result = {"handled": True, "ok": False, "reply_text": "查無唯一客戶，請補電話、地址或客戶ID"}
        else:
            result = {
                "handled": True,
                "ok": True,
                "reply_text": format_record_timeline(target_type, doc, limit=fields.get("limit", 10)),
                "target_type": target_type,
                "target_id": doc.id,
                "customer_name": (doc.to_dict() or {}).get("name", ""),
                "phone": (doc.to_dict() or {}).get("phone", ""),
                "parsed_tag": parsed.get("tag", ""),
            }
    elif action == "query_contract_end":
        ok, text, ctx = query_contract_end_text(fields)
        result = {"handled": True, "ok": ok, "reply_text": text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get("target_type", "")
        if action == "buyer_followup":
            target_type = "buyer"
        elif action == "seller_followup":
            target_type = "seller"

        if target_type not in ("buyer", "seller"):
            result = {"handled": True, "ok": False, "reply_text": "請提供對象：買方 或 賣方"}
        else:
            doc = find_customer_record(
                target_type=target_type,
                record_id=fields.get("record_id", ""),
                phone=fields.get("phone", ""),
                name=fields.get("name", ""),
            )
            if not doc:
                result = {"handled": True, "ok": False, "reply_text": "找不到唯一客戶，請補客戶ID或正確電話"}
            else:
                doc_ref = db.collection("buyers" if target_type == "buyer" else "sellers").document(doc.id)
                labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(fields.get("labels")))
                summary_parts = []
                if fields.get("content"):
                    summary_parts.append(fields["content"])
                if fields.get("address"):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get("price"):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary("；".join(summary_parts).strip() or "LINE 更新", event)

                update_customer_note_and_labels(
                    target_type=target_type,
                    doc_ref=doc_ref,
                    content=summary_text,
                    labels=labels,
                    stage=fields.get("stage", ""),
                    source=fields.get("source", "LINE"),
                    event=event,
                )

                if action in ("buyer_followup", "seller_followup", "classify"):
                    add_customer_followup(
                        target_type=target_type,
                        customer_id=doc.id,
                        content=summary_text,
                        next_action=fields.get("next_action", ""),
                        next_contact_date=fields.get("next_contact_date", ""),
                        labels=labels,
                        line_event=event,
                    )

                label_text = "客需" if target_type == "buyer" else "委託"
                current_data = doc_ref.get().to_dict() or {}
                result = {
                    "handled": True,
                    "ok": True,
                    "reply_text": f"已註記{label_text}：{current_data.get('name', '')}",
                    "target_type": target_type,
                    "target_id": doc.id,
                    "customer_name": current_data.get("name", ""),
                    "phone": current_data.get("phone", ""),
                    "parsed_tag": parsed.get("tag", ""),
                }

    save_line_log(
        parsed,
        event,
        "success" if result.get("ok") else "failed",
        target_type=result.get("target_type", ""),
        target_id=result.get("target_id", ""),
        note=result.get("reply_text", ""),
        sender_display_name=sender_display_name,
    )

    if result.get("ok") and result.get("target_type") and result.get("target_id"):
        incoming_message_id = message.get("id", "")
        save_line_message_link(
            incoming_message_id,
            result["target_type"],
            result["target_id"],
            tag=result.get("parsed_tag", ""),
            action=action,
            customer_name=result.get("customer_name", ""),
            phone=result.get("phone", ""),
            source_event=event,
        )

    return result


# ========= 引用回覆：開發以 phone / target_id 自動對標（新版覆蓋） =========

def find_developments_by_phone(phone: str):
    normalized = normalize_phone(phone)
    if not normalized:
        return []
    matches = []
    for doc in db.collection("developments").stream():
        data = doc.to_dict() or {}
        if normalize_phone(data.get("phone", "")) == normalized:
            matches.append(doc)
    return matches

def _status_alias_from_plain_text(text: str) -> str:
    raw = (text or "").strip()
    mapping = {
        "已連繫": "已聯繫",
        "已聯繫": "已聯繫",
        "已連絡": "已聯繫",
        "已聯絡": "已聯繫",
        "待聯繫": "待聯繫",
        "待聯絡": "待聯繫",
        "已寄開發信": "已寄開發信待跑開發",
        "已寄信": "已寄開發信待跑開發",
        "已開發": "已開發",
        "待調謄本": "待調謄本",
        "已調謄本": "已調謄本",
        "持續追蹤": "持續追蹤",
        "已簽回": "已簽回",
        "無效": "無效",
    }
    return mapping.get(raw, "")

def _parse_development_quote_updates(raw_text: str):
    text = (raw_text or "").strip()
    updates = {}
    note_lines = []
    field_hit = False

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line)
        if m:
            key = normalize_line_key(m.group(1))
            value = (m.group(2) or "").strip()
            field_hit = True

            if key in ("current_stage", "stage"):
                if value:
                    updates["current_stage"] = normalize_development_status(value)
            elif key == "next_action":
                updates["next_action"] = normalize_development_next_action(value)
            elif key == "next_action_date":
                updates["next_action_date"] = value
            elif key == "registered_address":
                updates["registered_address"] = value
            elif key == "source":
                updates["source"] = infer_development_source(value, updates.get("url", ""))
            elif key == "url":
                updates["url"] = value
                if not updates.get("source"):
                    updates["source"] = infer_development_source("", value)
            elif key in ("content", "address", "name", "phone"):
                note_lines.append(f"{m.group(1).strip()}: {value}" if value else m.group(1).strip())
            else:
                note_lines.append(line)
        else:
            note_lines.append(line)

    joined = "\n".join([x for x in note_lines if x]).strip()

    # 沒寫欄位，但直接回「已連繫」這種狀態
    if not updates.get("current_stage"):
        status = _status_alias_from_plain_text(text)
        if status:
            updates["current_stage"] = status

    # 沒寫欄位，但整句像地址，當戶籍地址
    if not updates.get("registered_address") and _looks_like_address(text):
        updates["registered_address"] = text

    followup_text = joined or text
    if field_hit and not followup_text:
        parts = []
        if updates.get("current_stage"):
            parts.append(f"目前狀況：{updates['current_stage']}")
        if updates.get("next_action"):
            parts.append(f"下一步：{updates['next_action']}")
        if updates.get("next_action_date"):
            parts.append(f"下一次時間：{updates['next_action_date']}")
        if updates.get("registered_address"):
            parts.append(f"戶籍地址：{updates['registered_address']}")
        followup_text = "；".join(parts)

    return updates, followup_text.strip()

def _describe_development(doc):
    data = doc.to_dict() or {}
    return f"{data.get('name','未填姓名')} / {data.get('phone','-')} / {data.get('address','-')}"

def _resolve_development_from_link(link):
    # 先用 target_id
    target_id = (link.get("target_id") or "").strip()
    if target_id:
        doc = db.collection("developments").document(target_id).get()
        if doc.exists:
            return doc, []

    # 再用 phone
    phone = (link.get("phone") or "").strip()
    matches = find_developments_by_phone(phone) if phone else []
    if len(matches) == 1:
        return matches[0], []
    if len(matches) > 1:
        return None, matches

    # 再用 name 當最後手段
    name = (link.get("customer_name") or "").strip()
    if name:
        docs = list(db.collection("developments").where("name", "==", name).limit(5).stream())
        if len(docs) == 1:
            return docs[0], []
        if len(docs) > 1:
            return None, docs

    return None, []

def process_quote_context_message(event):
    message = event.get("message") or {}
    quoted_message_id = message.get("quotedMessageId", "")
    raw_text = (message.get("text") or "").strip()
    if not quoted_message_id or not raw_text:
        return {"handled": False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {"handled": False}

    target_type = link.get("target_type", "")

    # ===== 開發：用 target_id / phone 自動對標 =====
    if target_type == "development":
        doc, multi = _resolve_development_from_link(link)
        if multi:
            preview = "\n".join([f"- {_describe_development(d)}" for d in multi[:5]])
            return {
                "handled": True,
                "ok": False,
                "reply_text": f"找到同電話/同姓名多筆開發，請確認要註記哪一筆：\n{preview}",
            }
        if not doc:
            phone = (link.get("phone") or "").strip()
            who = (link.get("customer_name") or "").strip()
            hint = f"（電話：{phone}，姓名：{who}）" if phone or who else ""
            return {
                "handled": True,
                "ok": False,
                "reply_text": f"找不到對應的開發資料{hint}",
            }

        updates, followup_text = _parse_development_quote_updates(raw_text)
        doc_ref = db.collection("developments").document(doc.id)
        current = doc.to_dict() or {}

        payload = {
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": "line_bot",
            "updated_by_name": "LINE Bot",
        }

        if updates.get("current_stage"):
            payload["current_stage"] = updates["current_stage"]
            payload["stage"] = updates["current_stage"]   # 舊欄位相容
        if updates.get("next_action"):
            payload["next_action"] = updates["next_action"]
        if updates.get("next_action_date"):
            payload["next_action_date"] = updates["next_action_date"]
        if updates.get("registered_address"):
            payload["registered_address"] = updates["registered_address"]
        if updates.get("source") or updates.get("url"):
            payload["source"] = infer_development_source(updates.get("source", current.get("source", "")), updates.get("url", current.get("url", "")))
        if updates.get("url"):
            payload["url"] = updates["url"]

        doc_ref.update(payload)

        # note 與 followup 都寫進去
        labels = dedupe_keep_order(["LINE紀錄", "群組回覆註記"])
        update_customer_note_and_labels(
            target_type="development",
            doc_ref=doc_ref,
            content=followup_text,
            labels=labels,
            stage=updates.get("current_stage", ""),
            source=payload.get("source", current.get("source", "")),
            event=event,
        )
        add_customer_followup(
            target_type="development",
            customer_id=doc.id,
            content=followup_text,
            next_action=updates.get("next_action", ""),
            next_contact_date=updates.get("next_action_date", ""),
            labels=labels,
            line_event=event,
        )

        latest = doc_ref.get().to_dict() or {}
        reply_name = latest.get("name", "") or "未填姓名"
        reply_phone = latest.get("phone", "") or "-"
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已註記開發：{reply_name}（{reply_phone}）",
            "target_type": "development",
            "target_id": doc.id,
            "customer_name": reply_name,
            "phone": reply_phone,
            "parsed_tag": "群組回覆註記",
        }

    # ===== 客需 / 委託維持原邏輯 =====
    target_id = link.get("target_id", "")
    if target_type not in ("buyer", "seller") or not target_id:
        return {"handled": False}

    collection_name = "buyers" if target_type == "buyer" else "sellers"
    doc_ref = db.collection(collection_name).document(target_id)
    doc = doc_ref.get()
    if not doc.exists:
        return {"handled": True, "ok": False, "reply_text": "未寫入：引用的客戶資料不存在"}

    labels = dedupe_keep_order(["LINE紀錄", "群組回覆註記"])
    reply_only_text = raw_text

    update_customer_note_and_labels(
        target_type=target_type,
        doc_ref=doc_ref,
        content=reply_only_text,
        labels=labels,
        source="LINE",
        event=event,
    )
    add_customer_followup(
        target_type=target_type,
        customer_id=target_id,
        content=reply_only_text,
        labels=labels,
        line_event=event,
    )

    parsed = {
        "tag": "群組回覆註記",
        "action": "quoted_context_note",
        "fields": {"quoted_message_id": quoted_message_id},
        "raw_text": raw_text,
    }
    save_line_log(parsed, event, "success", target_type=target_type, target_id=target_id, sender_display_name=get_line_sender_display_name(event))
    target_label = "客需" if target_type == "buyer" else "委託"
    data = doc.to_dict() or {}
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記{target_label}：{data.get('name', '')}（{data.get('phone','-')}）",
        "target_type": target_type,
        "target_id": target_id,
        "customer_name": data.get("name", ""),
        "phone": data.get("phone", ""),
        "parsed_tag": "群組回覆註記",
    }

def process_line_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    sender_display_name = get_line_sender_display_name(event)
    raw_text = message.get("text", "")

    # 先處理引用回覆，避免「已連繫」被當成新開發
    quoted_result = process_quote_context_message(event)
    if quoted_result.get("handled"):
        parsed = {
            "tag": "群組回覆註記",
            "action": "quoted_context_note",
            "fields": {"quoted_message_id": (message or {}).get("quotedMessageId", "")},
            "raw_text": raw_text,
        }
        save_line_log(
            parsed,
            event,
            "success" if quoted_result.get("ok") else "failed",
            target_type=quoted_result.get("target_type", ""),
            target_id=quoted_result.get("target_id", ""),
            note=quoted_result.get("reply_text", ""),
            sender_display_name=sender_display_name,
        )
        if quoted_result.get("ok") and quoted_result.get("target_type") and quoted_result.get("target_id"):
            incoming_message_id = message.get("id", "")
            save_line_message_link(
                incoming_message_id,
                quoted_result["target_type"],
                quoted_result["target_id"],
                tag=quoted_result.get("parsed_tag", ""),
                action="quoted_context_note",
                customer_name=quoted_result.get("customer_name", ""),
                phone=quoted_result.get("phone", ""),
                source_event=event,
            )
        return quoted_result

    parsed = parse_line_formatted_message(raw_text)
    if not parsed:
        free = _maybe_freeform_development(raw_text)
        if free:
            if free.get("batch"):
                result = create_development_batch(raw_text, event)
            else:
                result = create_development(free["fields"], event)
            save_line_log(
                {"tag": "自由格式開發", "action": "create_development_batch" if free.get("batch") else "create_development",
                 "fields": free.get("fields", {}), "raw_text": raw_text},
                event,
                "success" if result.get("ok") else "failed",
                target_type=result.get("target_type", ""),
                target_id=result.get("target_id", ""),
                note=result.get("reply_text", ""),
                sender_display_name=sender_display_name,
            )
            if result.get("ok") and result.get("target_type") and result.get("target_id"):
                incoming_message_id = message.get("id", "")
                save_line_message_link(
                    incoming_message_id,
                    result["target_type"],
                    result["target_id"],
                    tag=result.get("parsed_tag", "自由格式開發"),
                    action="create_development",
                    customer_name=result.get("customer_name", ""),
                    phone=result.get("phone", ""),
                    source_event=event,
                )
            return result
        return {"handled": False}

    fields = parsed["fields"]
    action = parsed["action"]

    if action == "create_buyer_need":
        result = create_buyer_need(fields, event)
    elif action == "create_seller_listing":
        result = create_seller_listing(fields, event)
    elif action == "create_development":
        result = create_development(fields, event)
    elif action == "create_development_batch":
        result = create_development_batch(parsed.get("raw_body") or raw_text, event)
    elif action == "development_followup":
        result = add_development_followup_via_line(fields, event)
    elif action == "query_records":
        target_type, doc = resolve_customer_record(fields)
        if not doc and (fields.get("record_id") or fields.get("phone") or fields.get("name") or fields.get("address")):
            doc = find_development_record(fields.get("record_id",""), fields.get("phone",""), fields.get("name",""), fields.get("address",""))
            if doc:
                target_type = "development"
        if not doc:
            result = {"handled": True, "ok": False, "reply_text": "查無唯一客戶，請補電話、地址或客戶ID"}
        else:
            result = {
                "handled": True,
                "ok": True,
                "reply_text": format_record_timeline(target_type, doc, limit=fields.get("limit", 10)),
                "target_type": target_type,
                "target_id": doc.id,
                "customer_name": (doc.to_dict() or {}).get("name", ""),
                "phone": (doc.to_dict() or {}).get("phone", ""),
                "parsed_tag": parsed.get("tag", ""),
            }
    elif action == "query_contract_end":
        ok, text, ctx = query_contract_end_text(fields)
        result = {"handled": True, "ok": ok, "reply_text": text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get("target_type", "")
        if action == "buyer_followup":
            target_type = "buyer"
        elif action == "seller_followup":
            target_type = "seller"

        if target_type not in ("buyer", "seller"):
            result = {"handled": True, "ok": False, "reply_text": "請提供對象：買方 或 賣方"}
        else:
            doc = find_customer_record(
                target_type=target_type,
                record_id=fields.get("record_id", ""),
                phone=fields.get("phone", ""),
                name=fields.get("name", ""),
            )
            if not doc:
                result = {"handled": True, "ok": False, "reply_text": "找不到唯一客戶，請補客戶ID或正確電話"}
            else:
                doc_ref = db.collection("buyers" if target_type == "buyer" else "sellers").document(doc.id)
                labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(fields.get("labels")))
                summary_parts = []
                if fields.get("content"):
                    summary_parts.append(fields["content"])
                if fields.get("address"):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get("price"):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary("；".join(summary_parts).strip() or "LINE 更新", event)

                update_customer_note_and_labels(
                    target_type=target_type,
                    doc_ref=doc_ref,
                    content=summary_text,
                    labels=labels,
                    stage=fields.get("stage", ""),
                    source=fields.get("source", "LINE"),
                    event=event,
                )

                if action in ("buyer_followup", "seller_followup", "classify"):
                    add_customer_followup(
                        target_type=target_type,
                        customer_id=doc.id,
                        content=summary_text,
                        next_action=fields.get("next_action", ""),
                        next_contact_date=fields.get("next_contact_date", ""),
                        labels=labels,
                        line_event=event,
                    )

                label_text = "客需" if target_type == "buyer" else "委託"
                current_data = doc_ref.get().to_dict() or {}
                result = {
                    "handled": True,
                    "ok": True,
                    "reply_text": f"已註記{label_text}：{current_data.get('name', '')}（{current_data.get('phone','-')}）",
                    "target_type": target_type,
                    "target_id": doc.id,
                    "customer_name": current_data.get("name", ""),
                    "phone": current_data.get("phone", ""),
                    "parsed_tag": parsed.get("tag", ""),
                }

    save_line_log(
        parsed,
        event,
        "success" if result.get("ok") else "failed",
        target_type=result.get("target_type", ""),
        target_id=result.get("target_id", ""),
        note=result.get("reply_text", ""),
        sender_display_name=sender_display_name,
    )

    if result.get("ok") and result.get("target_type") and result.get("target_id"):
        incoming_message_id = message.get("id", "")
        save_line_message_link(
            incoming_message_id,
            result["target_type"],
            result["target_id"],
            tag=result.get("parsed_tag", ""),
            action=action,
            customer_name=result.get("customer_name", ""),
            phone=result.get("phone", ""),
            source_event=event,
        )

    return result



# ========= quoted reply followup fix for buyer / seller / development =========

def _collection_key_by_target(target_type: str):
    if target_type == "buyer":
        return "buyers", "buyer_followups", "buyer_id", "客需"
    if target_type == "seller":
        return "sellers", "seller_followups", "seller_id", "委託"
    return "developments", "development_followups", "development_id", "開發"

def _find_docs_by_phone_in_collection(target_type: str, phone: str):
    normalized = normalize_phone(phone)
    if not normalized:
        return []
    collection_name, _, _, _ = _collection_key_by_target(target_type)
    hits = []
    for doc in db.collection(collection_name).stream():
        data = doc.to_dict() or {}
        if normalize_phone(data.get("phone", "")) == normalized:
            hits.append(doc)
    return hits

def _resolve_target_doc_from_link(link):
    target_type = (link.get("target_type") or "").strip()
    if target_type not in ("buyer", "seller", "development"):
        return target_type, None, []
    collection_name, _, _, _ = _collection_key_by_target(target_type)
    target_id = (link.get("target_id") or "").strip()
    if target_id:
        doc = db.collection(collection_name).document(target_id).get()
        if doc.exists:
            return target_type, doc, []
    phone = (link.get("phone") or "").strip()
    if phone:
        docs = _find_docs_by_phone_in_collection(target_type, phone)
        if len(docs) == 1:
            return target_type, docs[0], []
        if len(docs) > 1:
            return target_type, None, docs
    name = (link.get("customer_name") or "").strip()
    if name:
        docs = list(db.collection(collection_name).where("name", "==", name).limit(5).stream())
        if len(docs) == 1:
            return target_type, docs[0], []
        if len(docs) > 1:
            return target_type, None, docs
    return target_type, None, []

def _describe_doc_brief(target_type: str, doc):
    data = doc.to_dict() or {}
    return f"{data.get('name','未填姓名')} / {data.get('phone','-')} / {data.get('address','-')}"

def _parse_quoted_reply_updates(target_type: str, raw_text: str):
    text = (raw_text or "").strip()
    updates = {}
    note_lines = []
    touched_field = False
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line)
        if not m:
            note_lines.append(line)
            continue
        key = normalize_line_key(m.group(1))
        value = (m.group(2) or "").strip()
        touched_field = True
        if key in ("stage", "current_stage"):
            updates["stage"] = normalize_development_status(value) if target_type == "development" else value
        elif key == "next_action":
            updates["next_action"] = normalize_development_next_action(value) if target_type == "development" else value
        elif key in ("next_contact_date", "next_action_date"):
            updates["next_contact_date"] = value
        elif key == "registered_address" and target_type == "development":
            updates["registered_address"] = value
        elif key in ("address", "phone", "name", "url", "source"):
            updates[key] = value
        elif key == "content":
            if value:
                note_lines.append(value)
        else:
            note_lines.append(line)
    joined_note = "\n".join([x for x in note_lines if x]).strip()
    if target_type == "development":
        plain_status_map = {
            "已連繫": "已聯繫",
            "已聯繫": "已聯繫",
            "已連絡": "已聯繫",
            "已聯絡": "已聯繫",
            "待聯繫": "待聯繫",
            "待聯絡": "待聯繫",
            "已寄開發信": "已寄開發信待跑開發",
            "持續追蹤": "持續追蹤",
            "未接": "待聯繫",
            "未接，再聯繫": "待聯繫",
        }
        if not updates.get("stage"):
            s = plain_status_map.get(text, "")
            if s:
                updates["stage"] = s
        if not updates.get("registered_address") and _looks_like_address(text):
            updates["registered_address"] = text
    followup_text = joined_note or text
    return updates, followup_text

def process_quote_context_message(event):
    message = event.get("message") or {}
    quoted_message_id = message.get("quotedMessageId", "")
    raw_text = (message.get("text") or "").strip()
    if not quoted_message_id or not raw_text:
        return {"handled": False}
    link = get_line_message_link(quoted_message_id)
    if not link:
        return {"handled": False}
    target_type, doc, multi = _resolve_target_doc_from_link(link)
    if target_type not in ("buyer", "seller", "development"):
        return {"handled": False}
    if multi:
        preview = "\n".join([f"- {_describe_doc_brief(target_type, d)}" for d in multi[:5]])
        return {"handled": True, "ok": False, "reply_text": f"找到多筆同電話/同姓名資料，請確認要註記哪一筆：\n{preview}"}
    if not doc:
        label = "客需" if target_type == "buyer" else "委託" if target_type == "seller" else "開發"
        return {"handled": True, "ok": False, "reply_text": f"找不到對應的{label}資料"}
    collection_name, _, _, label_text = _collection_key_by_target(target_type)
    updates, followup_text = _parse_quoted_reply_updates(target_type, raw_text)
    doc_ref = db.collection(collection_name).document(doc.id)
    labels = dedupe_keep_order(["LINE紀錄", "群組回覆註記"])
    extra_updates = {}
    for k in ("address", "url", "phone", "name"):
        if updates.get(k):
            extra_updates[k] = updates[k]
    if updates.get("next_action") and target_type == "development":
        extra_updates["next_action"] = updates["next_action"]
    if updates.get("next_contact_date") and target_type == "development":
        extra_updates["next_action_date"] = updates["next_contact_date"]
    update_kwargs = {
        "target_type": target_type,
        "doc_ref": doc_ref,
        "content": build_line_summary(followup_text, event),
        "labels": labels,
        "stage": updates.get("stage", ""),
        "source": link.get("source", "LINE"),
        "event": event,
    }
    if target_type == "development":
        update_kwargs["registered_address"] = updates.get("registered_address", "")
        if extra_updates:
            update_kwargs["extra_updates"] = extra_updates
    elif extra_updates:
        update_kwargs["extra_updates"] = extra_updates
    update_customer_note_and_labels(**update_kwargs)
    add_customer_followup(
        target_type=target_type,
        customer_id=doc.id,
        content=build_line_summary(followup_text, event),
        next_action=updates.get("next_action", ""),
        next_contact_date=updates.get("next_contact_date", ""),
        labels=labels,
        line_event=event,
        stage=updates.get("stage", "") if target_type == "development" else "",
        registered_address=updates.get("registered_address", "") if target_type == "development" else "",
    )
    current = doc_ref.get().to_dict() or {}
    reply_name = current.get("name", "") or "未填姓名"
    reply_phone = current.get("phone", "") or "-"
    return {
        "handled": True,
        "ok": True,
        "reply_text": f"已註記{label_text}：{reply_name}（{reply_phone}）",
        "target_type": target_type,
        "target_id": doc.id,
        "customer_name": reply_name,
        "phone": reply_phone,
        "parsed_tag": "群組回覆註記",
    }

def process_line_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}
    sender_display_name = get_line_sender_display_name(event)
    raw_text = message.get("text", "")
    quoted_result = process_quote_context_message(event)
    if quoted_result.get("handled"):
        parsed = {"tag": "群組回覆註記", "action": "quoted_context_note", "fields": {"quoted_message_id": (message or {}).get("quotedMessageId", "")}, "raw_text": raw_text}
        save_line_log(parsed, event, "success" if quoted_result.get("ok") else "failed", target_type=quoted_result.get("target_type", ""), target_id=quoted_result.get("target_id", ""), note=quoted_result.get("reply_text", ""), sender_display_name=sender_display_name)
        if quoted_result.get("ok") and quoted_result.get("target_type") and quoted_result.get("target_id"):
            incoming_message_id = message.get("id", "")
            save_line_message_link(incoming_message_id, quoted_result["target_type"], quoted_result["target_id"], tag=quoted_result.get("parsed_tag", ""), action="quoted_context_note", customer_name=quoted_result.get("customer_name", ""), phone=quoted_result.get("phone", ""), source_event=event)
        return quoted_result
    parsed = parse_line_formatted_message(raw_text)
    if not parsed:
        free = _maybe_freeform_development(raw_text) if '_maybe_freeform_development' in globals() else None
        if free:
            if free.get("batch"):
                result = create_development_batch(raw_text, event)
            else:
                result = create_development(free["fields"], event)
            save_line_log({"tag": "自由格式開發", "action": "create_development_batch" if free.get("batch") else "create_development", "fields": free.get("fields", {}), "raw_text": raw_text}, event, "success" if result.get("ok") else "failed", target_type=result.get("target_type", ""), target_id=result.get("target_id", ""), note=result.get("reply_text", ""), sender_display_name=sender_display_name)
            if result.get("ok") and result.get("target_type") and result.get("target_id"):
                incoming_message_id = message.get("id", "")
                save_line_message_link(incoming_message_id, result["target_type"], result["target_id"], tag=result.get("parsed_tag", "自由格式開發"), action="create_development", customer_name=result.get("customer_name", ""), phone=result.get("phone", ""), source_event=event)
            return result
        return {"handled": False}
    fields = parsed["fields"]
    action = parsed["action"]
    if action == "create_buyer_need":
        result = create_buyer_need(fields, event)
    elif action == "create_seller_listing":
        result = create_seller_listing(fields, event)
    elif action == "create_development":
        result = create_development(fields, event)
    elif action == "create_development_batch":
        result = create_development_batch(parsed.get("raw_body") or raw_text, event)
    elif action == "development_followup":
        result = add_development_followup_via_line(fields, event)
    elif action == "query_records":
        target_type, doc = resolve_customer_record(fields)
        if not doc and (fields.get("record_id") or fields.get("phone") or fields.get("name") or fields.get("address")):
            doc = find_customer_record("development", fields.get("record_id",""), fields.get("phone",""), fields.get("name",""), fields.get("address",""))
            if doc:
                target_type = "development"
        if not doc:
            result = {"handled": True, "ok": False, "reply_text": "查無唯一客戶，請補電話、地址或客戶ID"}
        else:
            result = {"handled": True, "ok": True, "reply_text": format_record_timeline(target_type, doc, limit=fields.get("limit", 10)), "target_type": target_type, "target_id": doc.id, "customer_name": (doc.to_dict() or {}).get("name", ""), "phone": (doc.to_dict() or {}).get("phone", ""), "parsed_tag": parsed.get("tag", "")}
    elif action == "query_contract_end":
        ok, text, ctx = query_contract_end_text(fields)
        result = {"handled": True, "ok": ok, "reply_text": text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get("target_type", "")
        if action == "buyer_followup":
            target_type = "buyer"
        elif action == "seller_followup":
            target_type = "seller"
        if target_type not in ("buyer", "seller"):
            result = {"handled": True, "ok": False, "reply_text": "請提供對象：買方 或 賣方"}
        else:
            doc = find_customer_record(target_type=target_type, record_id=fields.get("record_id", ""), phone=fields.get("phone", ""), name=fields.get("name", ""))
            if not doc:
                result = {"handled": True, "ok": False, "reply_text": "找不到唯一客戶，請補客戶ID或正確電話"}
            else:
                doc_ref = db.collection("buyers" if target_type == "buyer" else "sellers").document(doc.id)
                labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(fields.get("labels")))
                summary_parts = []
                if fields.get("content"):
                    summary_parts.append(fields["content"])
                if fields.get("address"):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get("price"):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary("；".join(summary_parts).strip() or "LINE 更新", event)
                update_customer_note_and_labels(target_type=target_type, doc_ref=doc_ref, content=summary_text, labels=labels, stage=fields.get("stage", ""), source=fields.get("source", "LINE"), event=event)
                if action in ("buyer_followup", "seller_followup", "classify"):
                    add_customer_followup(target_type=target_type, customer_id=doc.id, content=summary_text, next_action=fields.get("next_action", ""), next_contact_date=fields.get("next_contact_date", ""), labels=labels, line_event=event)
                label_text = "客需" if target_type == "buyer" else "委託"
                current_data = doc_ref.get().to_dict() or {}
                result = {"handled": True, "ok": True, "reply_text": f"已註記{label_text}：{current_data.get('name', '')}（{current_data.get('phone','-')}）", "target_type": target_type, "target_id": doc.id, "customer_name": current_data.get("name", ""), "phone": current_data.get("phone", ""), "parsed_tag": parsed.get("tag", "")}
    save_line_log(parsed, event, "success" if result.get("ok") else "failed", target_type=result.get("target_type", ""), target_id=result.get("target_id", ""), note=result.get("reply_text", ""), sender_display_name=sender_display_name)
    if result.get("ok") and result.get("target_type") and result.get("target_id"):
        incoming_message_id = message.get("id", "")
        save_line_message_link(incoming_message_id, result["target_type"], result["target_id"], tag=result.get("parsed_tag", ""), action=action, customer_name=result.get("customer_name", ""), phone=result.get("phone", ""), source_event=event)
    return result


# ========= final reply / strict development patch =========

def _build_development_input_help(batch: bool = False) -> str:
    if batch:
        return (
            "新增開發批次格式錯誤，請用下面格式：\n"
            "#新增開發批次\n"
            "姓名: 王先生\n"
            "電話: 0911000111\n"
            "地址: 台中市沙鹿區A路1號\n"
            "戶籍地址: 台中市沙鹿區OO路OO號\n"
            "網址: https://...\n"
            "來源: 掃街\n"
            "目前狀況: 待聯繫\n"
            "下一步: 電話聯繫\n"
            "下次時間: 2026-04-10\n"
            "內容: 門口自售\n"
            "///\n"
            "姓名: 陳小姐\n"
            "電話: 0922000222\n"
            "地址: 台中市梧棲區B路2號"
        )
    return (
        "新增開發格式錯誤，請用下面格式：\n"
        "#新增開發\n"
        "姓名: 王先生\n"
        "電話: 0911000111\n"
        "地址: 台中市沙鹿區OO路OO號\n"
        "戶籍地址: 台中市沙鹿區OO路OO號\n"
        "網址: https://...\n"
        "來源: 掃街\n"
        "目前狀況: 待聯繫\n"
        "下一步: 電話聯繫\n"
        "下次時間: 2026-04-10\n"
        "內容: 門口自售"
    )


def _make_google_nav_url(address: str) -> str:
    addr = (address or '').strip()
    if not addr:
        return ''
    from urllib.parse import quote
    return f"https://www.google.com/maps/search/?api=1&query={quote(addr, safe='')}"


def _strict_parse_development_fields(raw_body: str):
    lines = [ln.strip() for ln in (raw_body or '').splitlines() if ln.strip()]
    if not lines:
        return None
    fields = {}
    invalid_lines = []
    for line in lines:
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.+)$', line)
        if not m:
            invalid_lines.append(line)
            continue
        key = normalize_line_key(m.group(1))
        value = (m.group(2) or '').strip()
        if key == 'labels':
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value
    if invalid_lines:
        return None
    if not any(fields.get(k) for k in ('address', 'phone', 'name')):
        return None
    fields['target_type'] = 'development'
    if fields.get('current_stage') and not fields.get('stage'):
        fields['stage'] = fields['current_stage']
    if fields.get('next_action_date') and not fields.get('next_contact_date'):
        fields['next_contact_date'] = fields['next_action_date']
    return fields


def _strict_parse_development_batch(raw_body: str):
    body = (raw_body or '').strip()
    if not body:
        return []
    chunks = [c.strip() for c in re.split(r'(?m)^\s*(?:///+|---+)\s*$', body) if c.strip()]
    items = []
    for chunk in chunks:
        fields = _strict_parse_development_fields(chunk)
        if not fields:
            return []
        items.append(fields)
    return items


def _maybe_freeform_development(text: str):
    return None


def parse_line_formatted_message(text: str):
    lines = [ln.strip() for ln in (text or '').splitlines() if ln.strip()]
    if not lines:
        return None
    first = lines[0]
    if not first.startswith('#'):
        return None

    tag = first.lstrip('#').strip()
    tag_map = {
        '新增客需': 'create_buyer_need',
        '新增委託': 'create_seller_listing',
        '新增開發': 'create_development',
        '新增開發批次': 'create_development_batch',
        '開發追蹤': 'development_followup',
        '買方追蹤': 'buyer_followup',
        '賣方追蹤': 'seller_followup',
        '客戶分類': 'classify',
        '查詢紀錄': 'query_records',
        '查詢委託到期': 'query_contract_end',
        '帶看': 'buyer_followup',
        '成交': 'buyer_followup',
        '委託': 'seller_followup',
        '紀錄': 'generic_note',
    }
    action = tag_map.get(tag)
    if not action:
        return None

    raw_body = '\n'.join(lines[1:]).strip()

    if action == 'create_development':
        fields = _strict_parse_development_fields(raw_body)
        if not fields:
            return None
        return {'tag': tag, 'action': action, 'fields': fields, 'raw_text': text, 'raw_body': raw_body}

    if action == 'create_development_batch':
        items = _strict_parse_development_batch(raw_body)
        if not items:
            return None
        return {'tag': tag, 'action': action, 'fields': {}, 'raw_text': text, 'raw_body': raw_body}

    fields = {}
    for line in lines[1:]:
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.+)$', line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = m.group(2).strip()
        if key == 'labels':
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value

    if tag in ('買方追蹤', '帶看', '成交') and not fields.get('target_type'):
        fields['target_type'] = 'buyer'
    if tag in ('賣方追蹤', '委託') and not fields.get('target_type'):
        fields['target_type'] = 'seller'
    if tag == '開發追蹤' and not fields.get('target_type'):
        fields['target_type'] = 'development'

    fields['target_type'] = normalize_target_type(fields.get('target_type', '')) or fields.get('target_type', '')
    fields['intent_type'] = normalize_intent_type(fields.get('intent_type_raw', ''), fields)
    fields['deal_type'] = normalize_deal_type(fields.get('deal_type_raw', ''))
    fields['limit'] = parse_int_limit(fields.get('limit', 10), default=10, max_value=30)

    if action == 'create_buyer_need':
        if not (fields.get('name') and fields.get('phone')):
            return None
    elif action == 'create_seller_listing':
        if not (fields.get('name') and fields.get('phone')):
            return None
    elif action in ('buyer_followup', 'seller_followup', 'classify', 'query_records', 'query_contract_end', 'development_followup'):
        if not (fields.get('record_id') or fields.get('phone') or fields.get('name') or fields.get('address')):
            return None

    return {'tag': tag, 'action': action, 'fields': fields, 'raw_text': text, 'raw_body': raw_body}


def create_development(fields, event):
    phone = (fields.get('phone') or '').strip()
    name = (fields.get('name') or '').strip() or '未填姓名'
    url = (fields.get('url') or '').strip()
    source = infer_development_source(fields.get('source', ''), url)
    address = (fields.get('address') or '').strip()
    registered_address = (fields.get('registered_address') or '').strip()
    nav_url = _make_google_nav_url(registered_address or address)

    matches = find_records_by_phone('developments', phone) if phone else []
    if not matches and address:
        doc = find_development_record(address=address)
        if doc:
            matches = [doc]

    labels = build_development_labels(fields.get('labels'))
    content_text = (fields.get('content') or '').strip() or address or url or 'LINE 新增開發'
    note_content = build_line_summary(content_text, event)

    payload = {
        'name': name,
        'phone': phone,
        'source': source,
        'url': url,
        'address': address,
        'registered_address': registered_address,
        'registered_address_google_maps_url': nav_url,
        'current_stage': normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip() or '待聯繫'),
        'stage': normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip() or '待聯繫'),
        'next_action': normalize_development_next_action((fields.get('next_action') or '').strip()),
        'next_action_date': (fields.get('next_action_date') or '').strip() or (fields.get('next_contact_date') or '').strip(),
        'record_date': (fields.get('record_date') or '').strip() or now_taipei().strftime('%Y-%m-%d'),
        'labels': labels,
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
        'sender_display_name': get_line_sender_display_name(event) or '',
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection('developments').document(doc.id)
        update_customer_note_and_labels(
            target_type='development',
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=payload['stage'],
            source=payload['source'],
            event=event,
            registered_address=registered_address,
        )
        clean_updates = {k: v for k, v in payload.items() if v not in ('', None)}
        doc_ref.update(clean_updates)
        add_customer_followup(
            target_type='development',
            customer_id=doc.id,
            content=note_content + (f"\nGoogle導航: {nav_url}" if nav_url else ''),
            next_action=payload.get('next_action', ''),
            next_contact_date=payload.get('next_action_date', ''),
            labels=labels,
            line_event=event,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        reply_text = f"已註記開發：{updated_doc.get('name', '')}（{updated_doc.get('phone', '-') or '-'}）"
        if nav_url:
            reply_text += f"\nGoogle導航: {nav_url}"
        return {
            'handled': True,
            'ok': True,
            'reply_text': reply_text[:5000],
            'target_type': 'development',
            'target_id': doc.id,
            'customer_name': updated_doc.get('name', ''),
            'phone': updated_doc.get('phone', ''),
            'parsed_tag': '新增開發',
        }

    if len(matches) > 1:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：同電話有多筆開發資料，請補地址或客戶ID'}

    now = now_taipei().isoformat()
    payload.update({
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': 'LINE Bot',
        'note': append_note_block('', note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection('developments').document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type='development',
        customer_id=doc_ref.id,
        content=note_content + (f"\nGoogle導航: {nav_url}" if nav_url else ''),
        next_action=payload.get('next_action', ''),
        next_contact_date=payload.get('next_action_date', ''),
        labels=labels,
        line_event=event,
    )
    reply_text = f"已註記開發：{name}（{phone or '-'}）"
    if nav_url:
        reply_text += f"\nGoogle導航: {nav_url}"
    return {
        'handled': True,
        'ok': True,
        'reply_text': reply_text[:5000],
        'target_type': 'development',
        'target_id': doc_ref.id,
        'customer_name': name,
        'phone': phone,
        'parsed_tag': '新增開發',
    }


def create_development_batch(raw_text, event):
    body = (raw_text or '').strip()
    if body.startswith('#新增開發批次'):
        body = '\n'.join(body.splitlines()[1:]).strip()
    items = _strict_parse_development_batch(body)
    if not items:
        return {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=True)}
    ok_count = 0
    fail_count = 0
    last_result = None
    for fields in items:
        result = create_development(fields, event)
        last_result = result
        if result.get('ok'):
            ok_count += 1
        else:
            fail_count += 1
    reply_text = f"批量註記完成：成功 {ok_count} 筆，失敗 {fail_count} 筆"
    if last_result and last_result.get('target_type') and last_result.get('target_id'):
        return {
            'handled': True,
            'ok': ok_count > 0,
            'reply_text': reply_text,
            'target_type': last_result.get('target_type', ''),
            'target_id': last_result.get('target_id', ''),
            'customer_name': last_result.get('customer_name', ''),
            'phone': last_result.get('phone', ''),
            'parsed_tag': '新增開發批次',
        }
    return {'handled': True, 'ok': ok_count > 0, 'reply_text': reply_text}


def process_quote_context_message(event):
    message = event.get('message') or {}
    quoted_message_id = (message.get('quotedMessageId') or '').strip()
    raw_text = (message.get('text') or '').strip()
    if not quoted_message_id or not raw_text:
        return {'handled': False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {
            'handled': True,
            'ok': False,
            'reply_text': '這則回覆找不到對標資料，請直接回覆 bot 建立成功的那則訊息或已對標的追蹤訊息。'
        }

    target_type, doc, multi = _resolve_target_doc_from_link(link)
    if target_type not in ('buyer', 'seller', 'development'):
        return {
            'handled': True,
            'ok': False,
            'reply_text': '這則回覆目前無法判定對應客戶，請直接回覆 bot 建立成功的那則訊息。'
        }
    if multi:
        preview = '\n'.join([f"- {_describe_doc_brief(target_type, d)}" for d in multi[:5]])
        return {'handled': True, 'ok': False, 'reply_text': f"找到多筆同電話/同姓名資料，請確認要註記哪一筆：\n{preview}"}
    if not doc:
        label = '客需' if target_type == 'buyer' else '委託' if target_type == 'seller' else '開發'
        return {'handled': True, 'ok': False, 'reply_text': f'找不到對應的{label}資料'}

    collection_name, _, _, label_text = _collection_key_by_target(target_type)
    updates, followup_text = _parse_quoted_reply_updates(target_type, raw_text)
    doc_ref = db.collection(collection_name).document(doc.id)
    labels = dedupe_keep_order(['LINE紀錄', '群組回覆註記'])

    stage_for_note = updates.get('stage', '')
    source_for_note = updates.get('source', link.get('source', 'LINE'))
    registered_address = updates.get('registered_address', '') if target_type == 'development' else ''
    followup_content = build_line_summary(followup_text, event)

    update_customer_note_and_labels(
        target_type=target_type,
        doc_ref=doc_ref,
        content=followup_content,
        labels=labels,
        stage=stage_for_note,
        source=source_for_note,
        event=event,
        registered_address=registered_address,
    )

    manual_updates = {
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
    }
    for k in ('address', 'url', 'phone', 'name', 'source'):
        if updates.get(k):
            manual_updates[k] = updates[k]

    nav_url = ''
    if target_type == 'development':
        if updates.get('stage'):
            manual_updates['current_stage'] = updates['stage']
            manual_updates['stage'] = updates['stage']
        if updates.get('next_action'):
            manual_updates['next_action'] = updates['next_action']
        if updates.get('next_contact_date'):
            manual_updates['next_action_date'] = updates['next_contact_date']
        if updates.get('registered_address'):
            manual_updates['registered_address'] = updates['registered_address']
            nav_url = _make_google_nav_url(updates['registered_address'])
            if nav_url:
                manual_updates['registered_address_google_maps_url'] = nav_url
        elif doc.to_dict().get('registered_address'):
            nav_url = _make_google_nav_url(doc.to_dict().get('registered_address'))
    else:
        if updates.get('stage'):
            manual_updates['stage'] = updates['stage']

    clean_updates = {k: v for k, v in manual_updates.items() if v not in ('', None)}
    if clean_updates:
        doc_ref.update(clean_updates)

    followup_with_nav = followup_content + (f"\nGoogle導航: {nav_url}" if nav_url else '')
    add_customer_followup(
        target_type=target_type,
        customer_id=doc.id,
        content=followup_with_nav,
        next_action=updates.get('next_action', ''),
        next_contact_date=updates.get('next_contact_date', ''),
        labels=labels,
        line_event=event,
    )

    current = doc_ref.get().to_dict() or {}
    reply_name = current.get('name', '') or '未填姓名'
    reply_phone = current.get('phone', '') or '-'
    reply_text = f"已註記{label_text}：{reply_name}（{reply_phone}）"
    if nav_url and target_type == 'development':
        reply_text += f"\nGoogle導航: {nav_url}"
    return {
        'handled': True,
        'ok': True,
        'reply_text': reply_text[:5000],
        'target_type': target_type,
        'target_id': doc.id,
        'customer_name': reply_name,
        'phone': reply_phone,
        'parsed_tag': '群組回覆註記',
    }


def process_line_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    sender_display_name = get_line_sender_display_name(event)
    raw_text = (message.get('text') or '').strip()
    quoted_message_id = (message.get('quotedMessageId') or '').strip()

    if quoted_message_id:
        quoted_result = process_quote_context_message(event)
        parsed = {'tag': '群組回覆註記', 'action': 'quoted_context_note', 'fields': {'quoted_message_id': quoted_message_id}, 'raw_text': raw_text}
        save_line_log(parsed, event, 'success' if quoted_result.get('ok') else 'failed', target_type=quoted_result.get('target_type', ''), target_id=quoted_result.get('target_id', ''), note=quoted_result.get('reply_text', ''), sender_display_name=sender_display_name)
        if quoted_result.get('ok') and quoted_result.get('target_type') and quoted_result.get('target_id'):
            incoming_message_id = message.get('id', '')
            save_line_message_link(incoming_message_id, quoted_result['target_type'], quoted_result['target_id'], tag=quoted_result.get('parsed_tag', ''), action='quoted_context_note', customer_name=quoted_result.get('customer_name', ''), phone=quoted_result.get('phone', ''), source_event=event)
        return quoted_result

    if raw_text.startswith('#新增開發批次'):
        parsed = parse_line_formatted_message(raw_text)
        if not parsed:
            result = {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=True)}
            save_line_log({'tag': '新增開發批次', 'action': 'create_development_batch', 'fields': {}, 'raw_text': raw_text}, event, 'failed', note=result['reply_text'], sender_display_name=sender_display_name)
            return result
    elif raw_text.startswith('#新增開發'):
        parsed = parse_line_formatted_message(raw_text)
        if not parsed:
            result = {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=False)}
            save_line_log({'tag': '新增開發', 'action': 'create_development', 'fields': {}, 'raw_text': raw_text}, event, 'failed', note=result['reply_text'], sender_display_name=sender_display_name)
            return result
    else:
        parsed = parse_line_formatted_message(raw_text)

    if not parsed:
        return {'handled': False}

    fields = parsed['fields']
    action = parsed['action']

    if action == 'create_buyer_need':
        result = create_buyer_need(fields, event)
    elif action == 'create_seller_listing':
        result = create_seller_listing(fields, event)
    elif action == 'create_development':
        result = create_development(fields, event)
    elif action == 'create_development_batch':
        result = create_development_batch(parsed.get('raw_body') or raw_text, event)
    elif action == 'development_followup':
        result = add_development_followup_via_line(fields, event)
    elif action == 'query_records':
        target_type, doc = resolve_customer_record(fields)
        if not doc and (fields.get('record_id') or fields.get('phone') or fields.get('name') or fields.get('address')):
            doc = find_development_record(fields.get('record_id',''), fields.get('phone',''), fields.get('name',''), fields.get('address',''))
            if doc:
                target_type = 'development'
        if not doc:
            result = {'handled': True, 'ok': False, 'reply_text': '查無唯一客戶，請補電話、地址或客戶ID'}
        else:
            result = {'handled': True, 'ok': True, 'reply_text': format_record_timeline(target_type, doc, limit=fields.get('limit', 10)), 'target_type': target_type, 'target_id': doc.id, 'customer_name': (doc.to_dict() or {}).get('name', ''), 'phone': (doc.to_dict() or {}).get('phone', ''), 'parsed_tag': parsed.get('tag', '')}
    elif action == 'query_contract_end':
        ok, text, ctx = query_contract_end_text(fields)
        result = {'handled': True, 'ok': ok, 'reply_text': text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get('target_type', '')
        if action == 'buyer_followup':
            target_type = 'buyer'
        elif action == 'seller_followup':
            target_type = 'seller'
        if target_type not in ('buyer', 'seller'):
            result = {'handled': True, 'ok': False, 'reply_text': '請提供對象：買方 或 賣方'}
        else:
            doc = find_customer_record(target_type=target_type, record_id=fields.get('record_id', ''), phone=fields.get('phone', ''), name=fields.get('name', ''), address=fields.get('address', ''))
            if not doc:
                result = {'handled': True, 'ok': False, 'reply_text': '找不到唯一客戶，請補客戶ID或正確電話'}
            else:
                doc_ref = db.collection('buyers' if target_type == 'buyer' else 'sellers').document(doc.id)
                labels = dedupe_keep_order(['LINE紀錄'] + ensure_list(fields.get('labels')))
                summary_parts = []
                if fields.get('content'):
                    summary_parts.append(fields['content'])
                if fields.get('address'):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get('price'):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary('；'.join(summary_parts).strip() or 'LINE 更新', event)
                update_customer_note_and_labels(target_type=target_type, doc_ref=doc_ref, content=summary_text, labels=labels, stage=fields.get('stage', ''), source=fields.get('source', 'LINE'), event=event)
                if action in ('buyer_followup', 'seller_followup', 'classify'):
                    add_customer_followup(target_type=target_type, customer_id=doc.id, content=summary_text, next_action=fields.get('next_action', ''), next_contact_date=fields.get('next_contact_date', ''), labels=labels, line_event=event)
                label_text = '客需' if target_type == 'buyer' else '委託'
                current_data = doc_ref.get().to_dict() or {}
                result = {'handled': True, 'ok': True, 'reply_text': f"已註記{label_text}：{current_data.get('name', '')}（{current_data.get('phone','-')}）", 'target_type': target_type, 'target_id': doc.id, 'customer_name': current_data.get('name', ''), 'phone': current_data.get('phone', ''), 'parsed_tag': parsed.get('tag', '')}

    save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', target_type=result.get('target_type', ''), target_id=result.get('target_id', ''), note=result.get('reply_text', ''), sender_display_name=sender_display_name)
    if result.get('ok') and result.get('target_type') and result.get('target_id'):
        incoming_message_id = message.get('id', '')
        save_line_message_link(incoming_message_id, result['target_type'], result['target_id'], tag=result.get('parsed_tag', ''), action=action, customer_name=result.get('customer_name', ''), phone=result.get('phone', ''), source_event=event)
    return result


# ========= final v3 patch: relaxed development header mode + reliable quote mapping =========

def update_customer_note_and_labels(target_type: str, doc_ref, content: str, labels=None, stage="", source="LINE", event=None, registered_address="", extra_updates=None):
    labels = dedupe_keep_order(["LINE紀錄"] + ensure_list(labels))
    snapshot = doc_ref.get()
    current = snapshot.to_dict() or {}
    old_note = current.get("note", "")

    updates = {
        "labels": firestore.ArrayUnion(labels),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": "line_bot",
        "updated_by_name": "LINE Bot",
    }
    if content:
        source_label = build_line_operator_label(event) if event else "LINE"
        updates["note"] = append_note_block(old_note, content, source_label)
    if stage:
        if target_type == 'development':
            normalized_stage = normalize_development_status(stage)
            updates['stage'] = normalized_stage
            updates['current_stage'] = normalized_stage
        else:
            updates['stage'] = stage
    if source:
        updates['source'] = source
    if registered_address:
        updates['registered_address'] = registered_address
        nav_url = _make_google_nav_url(registered_address)
        if nav_url:
            updates['registered_address_google_maps_url'] = nav_url
    if extra_updates:
        cleaned = {k: v for k, v in extra_updates.items() if v not in (None, "")}
        if cleaned:
            updates.update(cleaned)
    doc_ref.update(updates)


def add_customer_followup(target_type: str, customer_id: str, content: str, next_action="", next_contact_date="", labels=None, line_event=None, stage="", registered_address=""):
    if target_type == 'buyer':
        collection_name = 'buyer_followups'
        key_name = 'buyer_id'
    elif target_type == 'seller':
        collection_name = 'seller_followups'
        key_name = 'seller_id'
    else:
        collection_name = 'development_followups'
        key_name = 'development_id'

    sender_display_name = get_line_sender_display_name(line_event) if line_event else ''
    data = {
        key_name: customer_id,
        'contact_time': now_taipei().strftime('%Y-%m-%d %H:%M'),
        'channel': 'LINE',
        'content': content,
        'next_action': next_action,
        'next_contact_date': next_contact_date,
        'labels': dedupe_keep_order(['LINE紀錄'] + ensure_list(labels)),
        'created_at': now_taipei().isoformat(),
        'created_by_id': 'line_bot',
        'created_by_name': 'LINE Bot',
        'sender_display_name': sender_display_name,
    }
    if stage:
        data['stage'] = stage
        if target_type == 'development':
            data['current_stage'] = stage
    if registered_address:
        data['registered_address'] = registered_address
        nav_url = _make_google_nav_url(registered_address)
        if nav_url:
            data['registered_address_google_maps_url'] = nav_url
    if line_event:
        source = line_event.get('source', {})
        data['line_group_id'] = source.get('groupId', '')
        data['line_room_id'] = source.get('roomId', '')
        data['line_user_id'] = source.get('userId', '')
        data['line_message_id'] = (line_event.get('message') or {}).get('id', '')
        data['quoted_message_id'] = (line_event.get('message') or {}).get('quotedMessageId', '')
    db.collection(collection_name).add(data)


def _build_development_input_help(batch: bool = False) -> str:
    if batch:
        return (
            "新增開發批次格式範例（欄位可留空）：\n"
            "#新增開發批次\n"
            "姓名: 王先生\n"
            "電話: \n"
            "地址: 台中市沙鹿區A路1號\n"
            "戶籍地址: \n"
            "網址: \n"
            "來源: 掃街\n"
            "目前狀況: 待聯繫\n"
            "下一步: \n"
            "下次時間: \n"
            "內容: \n"
            "///\n"
            "姓名: 陳小姐\n"
            "地址: 台中市梧棲區B路2號"
        )
    return (
        "新增開發格式範例（只要有 #新增開發 就可以，欄位可留空）：\n"
        "#新增開發\n"
        "姓名: 王先生\n"
        "電話: \n"
        "地址: 台中市沙鹿區OO路OO號\n"
        "戶籍地址: \n"
        "網址: \n"
        "來源: 掃街\n"
        "目前狀況: 待聯繫\n"
        "下一步: \n"
        "下次時間: \n"
        "內容: "
    )


def _relaxed_parse_kv_block(raw_body: str):
    fields = {}
    note_lines = []
    lines = (raw_body or '').splitlines()
    saw_any = False
    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        saw_any = True
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if m:
            key = normalize_line_key(m.group(1))
            value = (m.group(2) or '').strip()
            if key == 'labels':
                fields[key] = parse_label_csv(value)
            else:
                fields[key] = value
        else:
            note_lines.append(line)
    if note_lines:
        existing = (fields.get('content') or '').strip()
        combined = '\n'.join(note_lines).strip()
        fields['content'] = (existing + ('\n' if existing and combined else '') + combined).strip()
    return fields, saw_any


def _strict_parse_development_fields(raw_body: str):
    fields, saw_any = _relaxed_parse_kv_block(raw_body)
    if not saw_any and not (raw_body or '').strip():
        fields = {}
    fields['target_type'] = 'development'
    if fields.get('current_stage') and not fields.get('stage'):
        fields['stage'] = fields['current_stage']
    if fields.get('next_action_date') and not fields.get('next_contact_date'):
        fields['next_contact_date'] = fields['next_action_date']
    return fields


def _strict_parse_development_batch(raw_body: str):
    body = (raw_body or '').strip()
    if not body:
        return []
    chunks = [c.strip() for c in re.split(r'(?m)^\s*(?:///+|---+)\s*$', body) if c.strip()]
    items = []
    for chunk in chunks:
        fields = _strict_parse_development_fields(chunk)
        # skip completely empty chunk
        meaningful = any((str(v).strip() if not isinstance(v, list) else any(str(x).strip() for x in v)) for k, v in fields.items() if k != 'target_type')
        if not meaningful and not chunk:
            continue
        items.append(fields)
    return items


def parse_line_formatted_message(text: str):
    lines = [ln.rstrip() for ln in (text or '').splitlines()]
    nonempty = [ln.strip() for ln in lines if ln.strip()]
    if not nonempty:
        return None
    first = nonempty[0]
    if not first.startswith('#'):
        return None

    tag = first.lstrip('#').strip()
    tag_map = {
        '新增客需': 'create_buyer_need',
        '新增委託': 'create_seller_listing',
        '新增開發': 'create_development',
        '新增開發批次': 'create_development_batch',
        '開發追蹤': 'development_followup',
        '買方追蹤': 'buyer_followup',
        '賣方追蹤': 'seller_followup',
        '客戶分類': 'classify',
        '查詢紀錄': 'query_records',
        '查詢委託到期': 'query_contract_end',
        '帶看': 'buyer_followup',
        '成交': 'buyer_followup',
        '委託': 'seller_followup',
        '紀錄': 'generic_note',
    }
    action = tag_map.get(tag)
    if not action:
        return None

    # raw body should preserve blank values after header
    started = False
    body_lines = []
    for ln in lines:
        stripped = ln.strip()
        if not started:
            if stripped == first:
                started = True
            continue
        body_lines.append(ln)
    raw_body = '\n'.join(body_lines).strip('\n')

    if action == 'create_development':
        fields = _strict_parse_development_fields(raw_body)
        return {'tag': tag, 'action': action, 'fields': fields, 'raw_text': text, 'raw_body': raw_body}

    if action == 'create_development_batch':
        items = _strict_parse_development_batch(raw_body)
        if not items:
            return None
        return {'tag': tag, 'action': action, 'fields': {}, 'raw_text': text, 'raw_body': raw_body}

    fields = {}
    for raw in body_lines:
        line = raw.strip()
        if not line:
            continue
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = (m.group(2) or '').strip()
        if key == 'labels':
            fields[key] = parse_label_csv(value)
        else:
            fields[key] = value

    if tag in ('買方追蹤', '帶看', '成交') and not fields.get('target_type'):
        fields['target_type'] = 'buyer'
    if tag in ('賣方追蹤', '委託') and not fields.get('target_type'):
        fields['target_type'] = 'seller'
    if tag == '開發追蹤' and not fields.get('target_type'):
        fields['target_type'] = 'development'

    fields['target_type'] = normalize_target_type(fields.get('target_type', '')) or fields.get('target_type', '')
    fields['intent_type'] = normalize_intent_type(fields.get('intent_type_raw', ''), fields)
    fields['deal_type'] = normalize_deal_type(fields.get('deal_type_raw', ''))
    fields['limit'] = parse_int_limit(fields.get('limit', 10), default=10, max_value=30)

    if action == 'create_buyer_need':
        if not (fields.get('name') and fields.get('phone')):
            return None
    elif action == 'create_seller_listing':
        if not (fields.get('name') and fields.get('phone')):
            return None
    elif action in ('buyer_followup', 'seller_followup', 'classify', 'query_records', 'query_contract_end', 'development_followup'):
        if not (fields.get('record_id') or fields.get('phone') or fields.get('name') or fields.get('address')):
            return None

    return {'tag': tag, 'action': action, 'fields': fields, 'raw_text': text, 'raw_body': raw_body}


def create_development(fields, event):
    phone = (fields.get('phone') or '').strip()
    name = (fields.get('name') or '').strip() or '未填姓名'
    url = (fields.get('url') or '').strip()
    source = infer_development_source(fields.get('source', ''), url)
    address = (fields.get('address') or '').strip()
    registered_address = (fields.get('registered_address') or '').strip()
    content_value = (fields.get('content') or '').strip()
    nav_url = _make_google_nav_url(registered_address or address)

    matches = find_records_by_phone('developments', phone) if phone else []
    if not matches and address:
        doc = find_development_record(address=address)
        if doc:
            matches = [doc]

    labels = build_development_labels(fields.get('labels'))
    content_text = content_value or registered_address or address or url or 'LINE 新增開發'
    note_content = build_line_summary(content_text, event)

    normalized_stage = normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip() or '待聯繫')
    normalized_next = normalize_development_next_action((fields.get('next_action') or '').strip())
    next_date = (fields.get('next_action_date') or '').strip() or (fields.get('next_contact_date') or '').strip()
    record_date = (fields.get('record_date') or '').strip() or now_taipei().strftime('%Y-%m-%d')

    payload = {
        'name': name,
        'phone': phone,
        'source': source,
        'url': url,
        'address': address,
        'registered_address': registered_address,
        'registered_address_google_maps_url': nav_url,
        'current_stage': normalized_stage,
        'stage': normalized_stage,
        'next_action': normalized_next,
        'next_action_date': next_date,
        'record_date': record_date,
        'labels': labels,
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
        'sender_display_name': get_line_sender_display_name(event) or '',
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection('developments').document(doc.id)
        update_customer_note_and_labels(
            target_type='development',
            doc_ref=doc_ref,
            content=note_content,
            labels=labels,
            stage=normalized_stage,
            source=source,
            event=event,
            registered_address=registered_address,
            extra_updates={k: v for k, v in payload.items() if k not in ('labels', 'updated_at', 'updated_by_id', 'updated_by_name') and v not in ('', None)}
        )
        add_customer_followup(
            target_type='development',
            customer_id=doc.id,
            content=note_content + (f"\nGoogle導航: {nav_url}" if nav_url else ''),
            next_action=normalized_next,
            next_contact_date=next_date,
            labels=labels,
            line_event=event,
            stage=normalized_stage,
            registered_address=registered_address,
        )
        updated_doc = doc_ref.get().to_dict() or {}
        reply_text = f"已註記開發：{updated_doc.get('name', '')}（{updated_doc.get('phone', '-') or '-'}）"
        if nav_url:
            reply_text += f"\nGoogle導航: {nav_url}"
        return {
            'handled': True,
            'ok': True,
            'reply_text': reply_text[:5000],
            'target_type': 'development',
            'target_id': doc.id,
            'customer_name': updated_doc.get('name', ''),
            'phone': updated_doc.get('phone', ''),
            'parsed_tag': '新增開發',
        }

    if len(matches) > 1:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：同電話有多筆開發資料，請補地址或客戶ID'}

    now = now_taipei().isoformat()
    payload.update({
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': 'LINE Bot',
        'note': append_note_block('', note_content, build_line_operator_label(event)),
    })
    doc_ref = db.collection('developments').document()
    doc_ref.set(payload)
    add_customer_followup(
        target_type='development',
        customer_id=doc_ref.id,
        content=note_content + (f"\nGoogle導航: {nav_url}" if nav_url else ''),
        next_action=normalized_next,
        next_contact_date=next_date,
        labels=labels,
        line_event=event,
        stage=normalized_stage,
        registered_address=registered_address,
    )
    reply_text = f"已註記開發：{name}（{phone or '-'}）"
    if nav_url:
        reply_text += f"\nGoogle導航: {nav_url}"
    return {
        'handled': True,
        'ok': True,
        'reply_text': reply_text[:5000],
        'target_type': 'development',
        'target_id': doc_ref.id,
        'customer_name': name,
        'phone': phone,
        'parsed_tag': '新增開發',
    }


def create_development_batch(raw_text, event):
    body = (raw_text or '').strip()
    if body.startswith('#新增開發批次'):
        body = '\n'.join(body.splitlines()[1:]).strip()
    items = _strict_parse_development_batch(body)
    if not items:
        return {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=True)}
    ok_count = 0
    fail_count = 0
    last_result = None
    for fields in items:
        result = create_development(fields, event)
        last_result = result
        if result.get('ok'):
            ok_count += 1
        else:
            fail_count += 1
    reply_text = f'批量註記完成：成功 {ok_count} 筆，失敗 {fail_count} 筆'
    if last_result and last_result.get('target_type') and last_result.get('target_id'):
        return {
            'handled': True,
            'ok': ok_count > 0,
            'reply_text': reply_text,
            'target_type': last_result.get('target_type', ''),
            'target_id': last_result.get('target_id', ''),
            'customer_name': last_result.get('customer_name', ''),
            'phone': last_result.get('phone', ''),
            'parsed_tag': '新增開發批次',
        }
    return {'handled': True, 'ok': ok_count > 0, 'reply_text': reply_text}


def _parse_quoted_reply_updates(target_type: str, raw_text: str):
    text = (raw_text or '').strip()
    updates = {}
    note_lines = []
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if not m:
            note_lines.append(line)
            continue
        key = normalize_line_key(m.group(1))
        value = (m.group(2) or '').strip()
        if key in ('stage', 'current_stage'):
            updates['stage'] = normalize_development_status(value) if target_type == 'development' else value
        elif key == 'next_action':
            updates['next_action'] = normalize_development_next_action(value) if target_type == 'development' else value
        elif key in ('next_contact_date', 'next_action_date'):
            updates['next_contact_date'] = value
        elif key == 'registered_address':
            updates['registered_address'] = value
        elif key in ('address', 'phone', 'name', 'url', 'source'):
            updates[key] = value
        elif key == 'content':
            if value:
                note_lines.append(value)
        else:
            note_lines.append(line)

    joined_note = '\n'.join([x for x in note_lines if x]).strip()
    plain_status_map = {
        '已連繫': '已聯繫',
        '已聯繫': '已聯繫',
        '已連絡': '已聯繫',
        '已聯絡': '已聯繫',
        '待聯繫': '待聯繫',
        '待聯絡': '待聯繫',
        '已寄開發信': '已寄開發信待跑開發',
        '持續追蹤': '持續追蹤',
        '未接': '待聯繫',
        '未接，再聯繫': '待聯繫',
    }
    if target_type == 'development' and not updates.get('stage'):
        s = plain_status_map.get(text, '')
        if s:
            updates['stage'] = s
    if not updates.get('registered_address') and _looks_like_address(text):
        updates['registered_address'] = text
    followup_text = joined_note or text
    return updates, followup_text


def process_quote_context_message(event):
    message = event.get('message') or {}
    quoted_message_id = (message.get('quotedMessageId') or '').strip()
    raw_text = (message.get('text') or '').strip()
    if not quoted_message_id or not raw_text:
        return {'handled': False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {
            'handled': True,
            'ok': False,
            'reply_text': '這則回覆找不到對標資料，請直接回覆 bot 建立成功的那則訊息，或確認這則原始訊息是用指令新增成功的。'
        }

    target_type, doc, multi = _resolve_target_doc_from_link(link)
    if target_type not in ('buyer', 'seller', 'development'):
        return {
            'handled': True,
            'ok': False,
            'reply_text': '這則回覆目前無法判定對應客戶，請直接回覆 bot 建立成功的那則訊息。'
        }
    if multi:
        preview = '\n'.join([f"- {_describe_doc_brief(target_type, d)}" for d in multi[:5]])
        return {'handled': True, 'ok': False, 'reply_text': f'找到多筆同電話/同姓名資料，請確認要註記哪一筆：\n{preview}'}
    if not doc:
        label = '客需' if target_type == 'buyer' else '委託' if target_type == 'seller' else '開發'
        return {'handled': True, 'ok': False, 'reply_text': f'找不到對應的{label}資料'}

    collection_name, _, _, label_text = _collection_key_by_target(target_type)
    updates, followup_text = _parse_quoted_reply_updates(target_type, raw_text)
    doc_ref = db.collection(collection_name).document(doc.id)
    labels = dedupe_keep_order(['LINE紀錄', '群組回覆註記'])

    extra_updates = {}
    for k in ('address', 'url', 'phone', 'name', 'source'):
        if updates.get(k):
            extra_updates[k] = updates[k]
    if target_type == 'development':
        if updates.get('stage'):
            extra_updates['stage'] = updates['stage']
            extra_updates['current_stage'] = updates['stage']
        if updates.get('next_action'):
            extra_updates['next_action'] = updates['next_action']
        if updates.get('next_contact_date'):
            extra_updates['next_action_date'] = updates['next_contact_date']
    else:
        if updates.get('stage'):
            extra_updates['stage'] = updates['stage']

    registered_address = updates.get('registered_address', '')
    followup_content = build_line_summary(followup_text, event)
    update_customer_note_and_labels(
        target_type=target_type,
        doc_ref=doc_ref,
        content=followup_content,
        labels=labels,
        stage=updates.get('stage', ''),
        source=updates.get('source', link.get('source', 'LINE')),
        event=event,
        registered_address=registered_address,
        extra_updates=extra_updates,
    )

    current_after_update = doc_ref.get().to_dict() or {}
    nav_url = ''
    if registered_address:
        nav_url = _make_google_nav_url(registered_address)
    elif current_after_update.get('registered_address'):
        nav_url = _make_google_nav_url(current_after_update.get('registered_address'))

    followup_with_nav = followup_content
    if nav_url:
        followup_with_nav += f"\nGoogle導航: {nav_url}"

    add_customer_followup(
        target_type=target_type,
        customer_id=doc.id,
        content=followup_with_nav,
        next_action=updates.get('next_action', ''),
        next_contact_date=updates.get('next_contact_date', ''),
        labels=labels,
        line_event=event,
        stage=updates.get('stage', ''),
        registered_address=registered_address,
    )

    current = doc_ref.get().to_dict() or {}
    reply_name = current.get('name', '') or '未填姓名'
    reply_phone = current.get('phone', '') or '-'
    reply_text = f"已註記{label_text}：{reply_name}（{reply_phone}）"
    if registered_address:
        reply_text += f"\n已更新戶籍地址：{registered_address}"
    if nav_url:
        reply_text += f"\nGoogle導航: {nav_url}"
    return {
        'handled': True,
        'ok': True,
        'reply_text': reply_text[:5000],
        'target_type': target_type,
        'target_id': doc.id,
        'customer_name': reply_name,
        'phone': reply_phone,
        'parsed_tag': '群組回覆註記',
    }


def process_line_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    sender_display_name = get_line_sender_display_name(event)
    raw_text = (message.get('text') or '').strip()
    quoted_message_id = (message.get('quotedMessageId') or '').strip()

    if quoted_message_id:
        quoted_result = process_quote_context_message(event)
        parsed = {'tag': '群組回覆註記', 'action': 'quoted_context_note', 'fields': {'quoted_message_id': quoted_message_id}, 'raw_text': raw_text}
        save_line_log(parsed, event, 'success' if quoted_result.get('ok') else 'failed', target_type=quoted_result.get('target_type', ''), target_id=quoted_result.get('target_id', ''), note=quoted_result.get('reply_text', ''), sender_display_name=sender_display_name)
        if quoted_result.get('ok') and quoted_result.get('target_type') and quoted_result.get('target_id'):
            incoming_message_id = message.get('id', '')
            save_line_message_link(incoming_message_id, quoted_result['target_type'], quoted_result['target_id'], tag=quoted_result.get('parsed_tag', ''), action='quoted_context_note', customer_name=quoted_result.get('customer_name', ''), phone=quoted_result.get('phone', ''), source_event=event)
        return quoted_result

    if raw_text.startswith('#新增開發批次'):
        parsed = parse_line_formatted_message(raw_text)
        if not parsed:
            result = {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=True)}
            save_line_log({'tag': '新增開發批次', 'action': 'create_development_batch', 'fields': {}, 'raw_text': raw_text}, event, 'failed', note=result['reply_text'], sender_display_name=sender_display_name)
            return result
    elif raw_text.startswith('#新增開發'):
        parsed = parse_line_formatted_message(raw_text)
        if not parsed:
            result = {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=False)}
            save_line_log({'tag': '新增開發', 'action': 'create_development', 'fields': {}, 'raw_text': raw_text}, event, 'failed', note=result['reply_text'], sender_display_name=sender_display_name)
            return result
    else:
        parsed = parse_line_formatted_message(raw_text)

    if not parsed:
        return {'handled': False}

    fields = parsed['fields']
    action = parsed['action']

    if action == 'create_buyer_need':
        result = create_buyer_need(fields, event)
    elif action == 'create_seller_listing':
        result = create_seller_listing(fields, event)
    elif action == 'create_development':
        result = create_development(fields, event)
    elif action == 'create_development_batch':
        result = create_development_batch(parsed.get('raw_body') or raw_text, event)
    elif action == 'development_followup':
        result = add_development_followup_via_line(fields, event)
    elif action == 'query_records':
        target_type, doc = resolve_customer_record(fields)
        if not doc and (fields.get('record_id') or fields.get('phone') or fields.get('name') or fields.get('address')):
            doc = find_development_record(fields.get('record_id',''), fields.get('phone',''), fields.get('name',''), fields.get('address',''))
            if doc:
                target_type = 'development'
        if not doc:
            result = {'handled': True, 'ok': False, 'reply_text': '查無唯一客戶，請補電話、地址或客戶ID'}
        else:
            result = {'handled': True, 'ok': True, 'reply_text': format_record_timeline(target_type, doc, limit=fields.get('limit', 10)), 'target_type': target_type, 'target_id': doc.id, 'customer_name': (doc.to_dict() or {}).get('name', ''), 'phone': (doc.to_dict() or {}).get('phone', ''), 'parsed_tag': parsed.get('tag', '')}
    elif action == 'query_contract_end':
        ok, text, ctx = query_contract_end_text(fields)
        result = {'handled': True, 'ok': ok, 'reply_text': text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get('target_type', '')
        if action == 'buyer_followup':
            target_type = 'buyer'
        elif action == 'seller_followup':
            target_type = 'seller'
        if target_type not in ('buyer', 'seller'):
            result = {'handled': True, 'ok': False, 'reply_text': '請提供對象：買方 或 賣方'}
        else:
            doc = find_customer_record(target_type=target_type, record_id=fields.get('record_id', ''), phone=fields.get('phone', ''), name=fields.get('name', ''), address=fields.get('address', ''))
            if not doc:
                result = {'handled': True, 'ok': False, 'reply_text': '找不到唯一客戶，請補客戶ID或正確電話'}
            else:
                doc_ref = db.collection('buyers' if target_type == 'buyer' else 'sellers').document(doc.id)
                labels = dedupe_keep_order(['LINE紀錄'] + ensure_list(fields.get('labels')))
                summary_parts = []
                if fields.get('content'):
                    summary_parts.append(fields['content'])
                if fields.get('address'):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get('price'):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary('；'.join(summary_parts).strip() or 'LINE 更新', event)
                update_customer_note_and_labels(target_type=target_type, doc_ref=doc_ref, content=summary_text, labels=labels, stage=fields.get('stage', ''), source=fields.get('source', 'LINE'), event=event)
                if action in ('buyer_followup', 'seller_followup', 'classify'):
                    add_customer_followup(target_type=target_type, customer_id=doc.id, content=summary_text, next_action=fields.get('next_action', ''), next_contact_date=fields.get('next_contact_date', ''), labels=labels, line_event=event, stage=fields.get('stage', ''))
                label_text = '客需' if target_type == 'buyer' else '委託'
                current_data = doc_ref.get().to_dict() or {}
                result = {'handled': True, 'ok': True, 'reply_text': f"已註記{label_text}：{current_data.get('name', '')}（{current_data.get('phone','-')}）", 'target_type': target_type, 'target_id': doc.id, 'customer_name': current_data.get('name', ''), 'phone': current_data.get('phone', ''), 'parsed_tag': parsed.get('tag', '')}

    save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', target_type=result.get('target_type', ''), target_id=result.get('target_id', ''), note=result.get('reply_text', ''), sender_display_name=sender_display_name)
    if result.get('ok') and result.get('target_type') and result.get('target_id'):
        incoming_message_id = message.get('id', '')
        save_line_message_link(incoming_message_id, result['target_type'], result['target_id'], tag=result.get('parsed_tag', ''), action=action, customer_name=result.get('customer_name', ''), phone=result.get('phone', ''), source_event=event)
    return result



# ========= source display + no-LINE-default patch (v4) =========

def infer_development_source(explicit_source: str, url: str) -> str:
    explicit = (explicit_source or '').strip()
    if explicit and explicit.upper() != 'LINE':
        return explicit
    return '踩線/屋主自售' if (url or '').strip() else '掃街'


def _store_source_value(target_type: str, incoming_source: str = '', current_source: str = '', url: str = '') -> str:
    incoming = (incoming_source or '').strip()
    current = (current_source or '').strip()
    if target_type == 'development':
        return infer_development_source(incoming or current, url)
    if incoming and incoming.upper() != 'LINE':
        return incoming
    if current and current.upper() != 'LINE':
        return current
    return ''


def _display_source_value(target_type: str, data: dict) -> str:
    raw = (data.get('source') or '').strip()
    if target_type == 'development':
        return raw if raw and raw.upper() != 'LINE' else infer_development_source('', data.get('url', ''))
    return '' if raw.upper() == 'LINE' else raw


def update_customer_note_and_labels(target_type: str, doc_ref, content: str, labels=None, stage='', source=None, event=None, registered_address='', extra_updates=None):
    labels = dedupe_keep_order(['LINE紀錄'] + ensure_list(labels))
    snapshot = doc_ref.get()
    current = snapshot.to_dict() or {}
    old_note = current.get('note', '')

    updates = {
        'labels': firestore.ArrayUnion(labels),
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
    }
    if content:
        source_label = build_line_operator_label(event) if event else 'LINE'
        updates['note'] = append_note_block(old_note, content, source_label)
    if stage:
        if target_type == 'development':
            normalized_stage = normalize_development_status(stage)
            updates['stage'] = normalized_stage
            updates['current_stage'] = normalized_stage
        else:
            updates['stage'] = stage

    # source: buyer/seller 空白就不要寫 LINE；development 則自動推斷
    url_for_source = ''
    if extra_updates and extra_updates.get('url'):
        url_for_source = extra_updates.get('url', '')
    else:
        url_for_source = current.get('url', '')
    normalized_source = _store_source_value(target_type, incoming_source=(source or ''), current_source=current.get('source', ''), url=url_for_source)
    if normalized_source:
        updates['source'] = normalized_source
    elif target_type in ('buyer', 'seller') and ((source or '').strip() == '' or str(source).strip().upper() == 'LINE'):
        # buyer/seller 空白時不覆蓋成 LINE，也不強制清空既有來源
        pass

    if registered_address:
        updates['registered_address'] = registered_address
        nav_url = _make_google_nav_url(registered_address)
        if nav_url:
            updates['registered_address_google_maps_url'] = nav_url
    if extra_updates:
        cleaned = {k: v for k, v in extra_updates.items() if v not in (None, '') and k != 'source'}
        if cleaned:
            updates.update(cleaned)
    doc_ref.update(updates)


def create_buyer_need(fields, event):
    phone = (fields.get('phone') or '').strip()
    matches = find_records_by_phone('buyers', phone)

    intent_type = normalize_intent_type(fields.get('intent_type_raw', '') or fields.get('intent_type', ''), fields)
    if not intent_type:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：#新增客需 請填 需求類型: 租 或 買賣'}

    labels = build_buyer_labels(intent_type, fields.get('labels'))
    budget = (fields.get('budget') or '').strip()
    source_value = _store_source_value('buyer', incoming_source=fields.get('source', ''))

    payload = {
        'name': (fields.get('name') or '').strip(),
        'phone': phone,
        'source': source_value,
        'intent_type': intent_type,
        'stage': (fields.get('stage') or '').strip() or '接觸',
        'preferred_areas': (fields.get('preferred_areas') or '').strip(),
        'property_type': (fields.get('property_type') or '').strip(),
        'room_range': (fields.get('room_range') or '').strip(),
        'car_need': (fields.get('car_need') or '').strip(),
        'labels': labels,
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
    }
    if intent_type == 'rent':
        payload['rent_max'] = (fields.get('rent_max') or '').strip() or budget
    else:
        payload['budget_max'] = (fields.get('budget_max') or '').strip() or budget

    extra_content = (fields.get('content') or '').strip()
    note_content = build_line_summary(extra_content or '新增客需', event)

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection('buyers').document(doc.id)
        update_customer_note_and_labels(target_type='buyer', doc_ref=doc_ref, content=note_content, labels=labels, stage=payload['stage'], source=source_value or None, event=event)
        clean_payload = {k: v for k, v in payload.items() if not (k == 'source' and not source_value)}
        doc_ref.update(clean_payload)
        add_customer_followup(target_type='buyer', customer_id=doc.id, content=note_content, labels=labels, line_event=event)
        updated_doc = doc_ref.get().to_dict() or {}
        return {'handled': True, 'ok': True, 'reply_text': f"已註記客需：{updated_doc.get('name', '')}", 'target_type': 'buyer', 'target_id': doc.id, 'customer_name': updated_doc.get('name', ''), 'phone': updated_doc.get('phone', ''), 'parsed_tag': '新增客需'}
    if len(matches) > 1:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：同電話有多位買方，請先到後台整理或改用客戶ID'}

    now = now_taipei().isoformat()
    payload.update({'created_at': now, 'created_by_id': 'line_bot', 'created_by_name': 'LINE Bot', 'note': append_note_block('', note_content, build_line_operator_label(event))})
    doc_ref = db.collection('buyers').document()
    doc_ref.set(payload)
    add_customer_followup(target_type='buyer', customer_id=doc_ref.id, content=note_content, labels=labels, line_event=event)
    return {'handled': True, 'ok': True, 'reply_text': f"已註記客需：{payload['name']}（{'租' if intent_type == 'rent' else '買賣'}）", 'target_type': 'buyer', 'target_id': doc_ref.id, 'customer_name': payload['name'], 'phone': payload['phone'], 'parsed_tag': '新增客需'}


def create_seller_listing(fields, event):
    phone = (fields.get('phone') or '').strip()
    matches = find_records_by_phone('sellers', phone)

    deal_type = normalize_deal_type(fields.get('deal_type_raw', '') or fields.get('deal_type', ''))
    if not deal_type:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：#新增委託 請填 委託類型: 賣 或 出租'}

    labels = build_seller_labels(deal_type, fields.get('labels'))
    source_value = _store_source_value('seller', incoming_source=fields.get('source', ''))
    note_content = build_line_summary((fields.get('content') or '').strip() or (fields.get('address') or '').strip() or '新增委託', event)

    payload = {
        'name': (fields.get('name') or '').strip(),
        'phone': phone,
        'source': source_value,
        'address': (fields.get('address') or '').strip(),
        'property_type': (fields.get('property_type') or '').strip(),
        'stage': (fields.get('stage') or '').strip() or '委託中',
        'deal_type': deal_type,
        'expected_price': (fields.get('expected_price') or '').strip() or (fields.get('price') or '').strip(),
        'min_price': (fields.get('min_price') or '').strip(),
        'contract_end_date': (fields.get('contract_end_date') or '').strip(),
        'labels': labels,
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection('sellers').document(doc.id)
        update_customer_note_and_labels(target_type='seller', doc_ref=doc_ref, content=note_content, labels=labels, stage=payload['stage'], source=source_value or None, event=event)
        clean_payload = {k: v for k, v in payload.items() if not (k == 'source' and not source_value)}
        doc_ref.update(clean_payload)
        add_customer_followup(target_type='seller', customer_id=doc.id, content=note_content, labels=labels, line_event=event)
        updated_doc = doc_ref.get().to_dict() or {}
        return {'handled': True, 'ok': True, 'reply_text': f"已註記委託：{updated_doc.get('name', '')}", 'target_type': 'seller', 'target_id': doc.id, 'customer_name': updated_doc.get('name', ''), 'phone': updated_doc.get('phone', ''), 'parsed_tag': '新增委託'}
    if len(matches) > 1:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：同電話有多位委託，請先到後台整理或改用客戶ID'}

    now = now_taipei().isoformat()
    payload.update({'created_at': now, 'created_by_id': 'line_bot', 'created_by_name': 'LINE Bot', 'note': append_note_block('', note_content, build_line_operator_label(event))})
    doc_ref = db.collection('sellers').document()
    doc_ref.set(payload)
    add_customer_followup(target_type='seller', customer_id=doc_ref.id, content=note_content, labels=labels, line_event=event)
    return {'handled': True, 'ok': True, 'reply_text': f"已註記委託：{payload['name']}（{'出租' if deal_type == 'rent' else '買賣'}）", 'target_type': 'seller', 'target_id': doc_ref.id, 'customer_name': payload['name'], 'phone': payload['phone'], 'parsed_tag': '新增委託'}


def create_development(fields, event):
    phone = (fields.get('phone') or '').strip()
    name = (fields.get('name') or '').strip() or '未填姓名'
    url = (fields.get('url') or '').strip()
    source_value = infer_development_source(fields.get('source', ''), url)
    address = (fields.get('address') or '').strip()
    registered_address = (fields.get('registered_address') or '').strip()
    nav_url = _make_google_nav_url(registered_address or address)

    matches = find_records_by_phone('developments', phone) if phone else []
    if not matches and address:
        doc = find_development_record(address=address)
        if doc:
            matches = [doc]

    labels = build_development_labels(fields.get('labels'))
    content_text = (fields.get('content') or '').strip() or registered_address or address or url or '新增開發'
    note_content = build_line_summary(content_text, event)

    payload = {
        'name': name,
        'phone': phone,
        'source': source_value,
        'url': url,
        'address': address,
        'registered_address': registered_address,
        'registered_address_google_maps_url': nav_url,
        'current_stage': normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip() or '待聯繫'),
        'stage': normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip() or '待聯繫'),
        'next_action': normalize_development_next_action((fields.get('next_action') or '').strip()),
        'next_action_date': (fields.get('next_action_date') or '').strip() or (fields.get('next_contact_date') or '').strip(),
        'record_date': (fields.get('record_date') or '').strip() or now_taipei().strftime('%Y-%m-%d'),
        'labels': labels,
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
        'sender_display_name': get_line_sender_display_name(event) or '',
    }

    if len(matches) == 1:
        doc = matches[0]
        doc_ref = db.collection('developments').document(doc.id)
        update_customer_note_and_labels(target_type='development', doc_ref=doc_ref, content=note_content, labels=labels, stage=payload['stage'], source=source_value, event=event, registered_address=registered_address)
        clean_updates = {k: v for k, v in payload.items() if v not in ('', None)}
        doc_ref.update(clean_updates)
        add_customer_followup(target_type='development', customer_id=doc.id, content=note_content, next_action=payload.get('next_action', ''), next_contact_date=payload.get('next_action_date', ''), labels=labels, line_event=event)
        updated_doc = doc_ref.get().to_dict() or {}
        reply_text = f"已註記開發：{updated_doc.get('name', '')}（{updated_doc.get('phone', '-') or '-'}）"
        if registered_address or nav_url:
            reply_text += '\n已更新戶籍地址與 Google導航連結'
        return {'handled': True, 'ok': True, 'reply_text': reply_text[:5000], 'target_type': 'development', 'target_id': doc.id, 'customer_name': updated_doc.get('name', ''), 'phone': updated_doc.get('phone', ''), 'parsed_tag': '新增開發'}
    if len(matches) > 1:
        return {'handled': True, 'ok': False, 'reply_text': '未寫入：同電話有多筆開發資料，請補地址或客戶ID'}

    now = now_taipei().isoformat()
    payload.update({'created_at': now, 'created_by_id': 'line_bot', 'created_by_name': 'LINE Bot', 'note': append_note_block('', note_content, build_line_operator_label(event))})
    doc_ref = db.collection('developments').document()
    doc_ref.set(payload)
    add_customer_followup(target_type='development', customer_id=doc_ref.id, content=note_content, next_action=payload.get('next_action', ''), next_contact_date=payload.get('next_action_date', ''), labels=labels, line_event=event)
    reply_text = f"已註記開發：{name}（{phone or '-'}）"
    if registered_address or nav_url:
        reply_text += '\n已更新戶籍地址與 Google導航連結'
    return {'handled': True, 'ok': True, 'reply_text': reply_text[:5000], 'target_type': 'development', 'target_id': doc_ref.id, 'customer_name': name, 'phone': phone, 'parsed_tag': '新增開發'}


def add_development_followup_via_line(fields, event):
    doc = find_development_record(record_id=fields.get('record_id', ''), phone=fields.get('phone', ''), name=fields.get('name', ''), address=fields.get('address', ''))
    if not doc:
        return {'handled': True, 'ok': False, 'reply_text': '找不到唯一開發資料，請補電話、地址或客戶ID'}

    current = doc.to_dict() or {}
    content = (fields.get('content') or '').strip() or '開發追蹤更新'
    current_stage = normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip())
    next_action = normalize_development_next_action((fields.get('next_action') or '').strip())
    next_action_date = (fields.get('next_action_date') or '').strip() or (fields.get('next_contact_date') or '').strip()
    registered_address = (fields.get('registered_address') or '').strip()
    source_value = infer_development_source(fields.get('source', '') or current.get('source', ''), (fields.get('url') or '').strip() or current.get('url', ''))

    doc_ref = db.collection('developments').document(doc.id)
    update_customer_note_and_labels(target_type='development', doc_ref=doc_ref, content=build_line_summary(content, event), labels=build_development_labels(fields.get('labels')), stage=current_stage, source=source_value, event=event, registered_address=registered_address)
    updates = {'updated_at': now_taipei().isoformat(), 'updated_by_id': 'line_bot', 'updated_by_name': 'LINE Bot'}
    if current_stage:
        updates['current_stage'] = current_stage
        updates['stage'] = current_stage
    if next_action:
        updates['next_action'] = next_action
    if next_action_date:
        updates['next_action_date'] = next_action_date
    if registered_address:
        updates['registered_address'] = registered_address
        nav_url = _make_google_nav_url(registered_address)
        if nav_url:
            updates['registered_address_google_maps_url'] = nav_url
    if source_value:
        updates['source'] = source_value
    doc_ref.update(updates)
    add_customer_followup(target_type='development', customer_id=doc.id, content=build_line_summary(content, event), next_action=next_action, next_contact_date=next_action_date, labels=build_development_labels(fields.get('labels')), line_event=event)

    data = doc_ref.get().to_dict() or {}
    reply_text = f"已註記開發：{data.get('name', '')}"
    if registered_address:
        reply_text += '\n已更新戶籍地址與 Google導航連結'
    return {'handled': True, 'ok': True, 'reply_text': reply_text, 'target_type': 'development', 'target_id': doc.id, 'customer_name': data.get('name', ''), 'phone': data.get('phone', ''), 'parsed_tag': '開發追蹤'}


def format_record_timeline(target_type: str, doc_snapshot, limit=10):
    data = doc_snapshot.to_dict() or {}
    record_id = doc_snapshot.id
    if target_type == 'buyer':
        followup_collection = 'buyer_followups'
        key_name = 'buyer_id'
    elif target_type == 'seller':
        followup_collection = 'seller_followups'
        key_name = 'seller_id'
    else:
        followup_collection = 'development_followups'
        key_name = 'development_id'

    followups = []
    for d in db.collection(followup_collection).where(key_name, '==', record_id).stream():
        item = d.to_dict() or {}
        followups.append({'time': item.get('contact_time') or item.get('created_at') or '', 'channel': item.get('channel', 'LINE'), 'text': (item.get('content') or '').strip(), 'created_by_name': item.get('created_by_name', '') or '', 'sender_display_name': item.get('sender_display_name', '') or '', 'stage': item.get('stage', '') or '', 'registered_address': item.get('registered_address', '') or ''})
    followups = [x for x in followups if x.get('text') or x.get('registered_address')]
    followups.sort(key=lambda x: x.get('time', ''), reverse=True)

    lines = ['客戶資訊']
    source_text = _display_source_value(target_type, data)
    if target_type == 'buyer':
        intent_map = {'rent': '租屋', 'buy': '買賣', 'both': '租買皆可'}
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {source_text}",
            f"需求類型: {intent_map.get(data.get('intent_type', ''), data.get('intent_type', '') or '-')}",
            f"預算: {data.get('budget_min', '') or data.get('rent_min', '') or '-'} ~ {data.get('budget_max', '') or data.get('rent_max', '') or '-'}",
            f"偏好區域: {data.get('preferred_areas', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"房數需求: {data.get('room_range', '') or '-'}",
            f"車位需求: {data.get('car_need', '') or '-'}",
        ])
    elif target_type == 'seller':
        deal_map = {'sale': '買賣', 'rent': '出租'}
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '')}",
            f"客源來源: {source_text}",
            f"委託類型: {deal_map.get(data.get('deal_type', ''), data.get('deal_type', '') or '-')}",
            f"地址: {data.get('address', '') or '-'}",
            f"產品類型: {data.get('property_type', '') or '-'}",
            f"開價: {data.get('expected_price', '') or '-'}",
            f"底價: {data.get('min_price', '') or '-'}",
            f"委託到期日: {data.get('contract_end_date', '') or '-'}",
        ])
    else:
        lines.extend([
            f"姓名: {data.get('name', '')}",
            f"電話: {data.get('phone', '') or '-'}",
            f"來源: {source_text}",
            f"地址: {data.get('address', '') or '-'}",
            f"戶籍地址: {data.get('registered_address', '') or '-'}",
            f"網址: {data.get('url', '') or '-'}",
            f"進度: {data.get('stage', '') or '-'}",
        ])
    lines.append('')
    lines.append('追蹤進度')
    if not followups:
        lines.append('目前沒有追蹤紀錄')
    else:
        for item in followups[:limit]:
            header_parts = [item.get('time', ''), item.get('channel', 'LINE')]
            if item.get('stage'):
                header_parts.append(item['stage'])
            creator = (item.get('created_by_name') or '').strip()
            sender = (item.get('sender_display_name') or '').strip()
            if creator:
                header_parts.append(f"KEYIN: {creator}")
            if sender:
                header_parts.append(f"留言者: {sender}")
            lines.append('｜'.join([x for x in header_parts if x]))
            if item.get('text'):
                lines.append(item.get('text', ''))
            if item.get('registered_address'):
                lines.append(f"戶籍地址：{item['registered_address']}")
            lines.append('')
    return '\n'.join(lines).strip()[:4500]


def process_quote_context_message(event):
    message = event.get('message') or {}
    quoted_message_id = (message.get('quotedMessageId') or '').strip()
    raw_text = (message.get('text') or '').strip()
    if not quoted_message_id or not raw_text:
        return {'handled': False}

    link = get_line_message_link(quoted_message_id)
    if not link:
        return {'handled': True, 'ok': False, 'reply_text': '這則回覆找不到對標資料，請直接回覆 bot 建立成功的那則訊息，或確認這則原始訊息是用指令新增成功的。'}

    target_type, doc, multi = _resolve_target_doc_from_link(link)
    if target_type not in ('buyer', 'seller', 'development'):
        return {'handled': True, 'ok': False, 'reply_text': '這則回覆目前無法判定對應客戶，請直接回覆 bot 建立成功的那則訊息。'}
    if multi:
        preview = '\n'.join([f"- {_describe_doc_brief(target_type, d)}" for d in multi[:5]])
        return {'handled': True, 'ok': False, 'reply_text': f'找到多筆同電話/同姓名資料，請確認要註記哪一筆：\n{preview}'}
    if not doc:
        label = '客需' if target_type == 'buyer' else '委託' if target_type == 'seller' else '開發'
        return {'handled': True, 'ok': False, 'reply_text': f'找不到對應的{label}資料'}

    collection_name, _, _, label_text = _collection_key_by_target(target_type)
    updates, followup_text = _parse_quoted_reply_updates(target_type, raw_text)
    doc_ref = db.collection(collection_name).document(doc.id)
    labels = dedupe_keep_order(['LINE紀錄', '群組回覆註記'])

    current_before = doc.to_dict() or {}
    url_for_source = (updates.get('url') or '').strip() or current_before.get('url', '')
    source_for_note = _store_source_value(target_type, incoming_source=updates.get('source', ''), current_source=current_before.get('source', ''), url=url_for_source)

    extra_updates = {}
    for k in ('address', 'url', 'phone', 'name'):
        if updates.get(k):
            extra_updates[k] = updates[k]
    if source_for_note:
        extra_updates['source'] = source_for_note
    if target_type == 'development':
        if updates.get('stage'):
            extra_updates['stage'] = updates['stage']
            extra_updates['current_stage'] = updates['stage']
        if updates.get('next_action'):
            extra_updates['next_action'] = updates['next_action']
        if updates.get('next_contact_date'):
            extra_updates['next_action_date'] = updates['next_contact_date']
    else:
        if updates.get('stage'):
            extra_updates['stage'] = updates['stage']

    registered_address = updates.get('registered_address', '')
    followup_content = build_line_summary(followup_text, event)
    update_customer_note_and_labels(target_type=target_type, doc_ref=doc_ref, content=followup_content, labels=labels, stage=updates.get('stage', ''), source=source_for_note or None, event=event, registered_address=registered_address, extra_updates=extra_updates)

    add_customer_followup(target_type=target_type, customer_id=doc.id, content=followup_content, next_action=updates.get('next_action', ''), next_contact_date=updates.get('next_contact_date', ''), labels=labels, line_event=event, stage=updates.get('stage', ''), registered_address=registered_address if target_type == 'development' else '')

    current = doc_ref.get().to_dict() or {}
    reply_name = current.get('name', '') or '未填姓名'
    reply_phone = current.get('phone', '') or '-'
    reply_text = f"已註記{label_text}：{reply_name}（{reply_phone}）"
    if registered_address and target_type == 'development':
        reply_text += '\n已更新戶籍地址與 Google導航連結'
    return {'handled': True, 'ok': True, 'reply_text': reply_text[:5000], 'target_type': target_type, 'target_id': doc.id, 'customer_name': reply_name, 'phone': reply_phone, 'parsed_tag': '群組回覆註記'}


def process_line_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}
    sender_display_name = get_line_sender_display_name(event)
    raw_text = (message.get('text') or '').strip()
    quoted_message_id = (message.get('quotedMessageId') or '').strip()

    if quoted_message_id:
        quoted_result = process_quote_context_message(event)
        parsed = {'tag': '群組回覆註記', 'action': 'quoted_context_note', 'fields': {'quoted_message_id': quoted_message_id}, 'raw_text': raw_text}
        save_line_log(parsed, event, 'success' if quoted_result.get('ok') else 'failed', target_type=quoted_result.get('target_type', ''), target_id=quoted_result.get('target_id', ''), note=quoted_result.get('reply_text', ''), sender_display_name=sender_display_name)
        if quoted_result.get('ok') and quoted_result.get('target_type') and quoted_result.get('target_id'):
            incoming_message_id = message.get('id', '')
            save_line_message_link(incoming_message_id, quoted_result['target_type'], quoted_result['target_id'], tag=quoted_result.get('parsed_tag', ''), action='quoted_context_note', customer_name=quoted_result.get('customer_name', ''), phone=quoted_result.get('phone', ''), source_event=event)
        return quoted_result

    if raw_text.startswith('#新增開發批次'):
        parsed = parse_line_formatted_message(raw_text)
        if not parsed:
            result = {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=True)}
            save_line_log({'tag': '新增開發批次', 'action': 'create_development_batch', 'fields': {}, 'raw_text': raw_text}, event, 'failed', note=result['reply_text'], sender_display_name=sender_display_name)
            return result
    elif raw_text.startswith('#新增開發'):
        parsed = parse_line_formatted_message(raw_text)
        if not parsed:
            result = {'handled': True, 'ok': False, 'reply_text': _build_development_input_help(batch=False)}
            save_line_log({'tag': '新增開發', 'action': 'create_development', 'fields': {}, 'raw_text': raw_text}, event, 'failed', note=result['reply_text'], sender_display_name=sender_display_name)
            return result
    else:
        parsed = parse_line_formatted_message(raw_text)

    if not parsed:
        return {'handled': False}

    fields = parsed['fields']
    action = parsed['action']

    if action == 'create_buyer_need':
        result = create_buyer_need(fields, event)
    elif action == 'create_seller_listing':
        result = create_seller_listing(fields, event)
    elif action == 'create_development':
        result = create_development(fields, event)
    elif action == 'create_development_batch':
        result = create_development_batch(parsed.get('raw_body') or raw_text, event)
    elif action == 'development_followup':
        result = add_development_followup_via_line(fields, event)
    elif action == 'query_records':
        target_type, doc = resolve_customer_record(fields)
        if not doc and (fields.get('record_id') or fields.get('phone') or fields.get('name') or fields.get('address')):
            doc = find_development_record(fields.get('record_id',''), fields.get('phone',''), fields.get('name',''), fields.get('address',''))
            if doc:
                target_type = 'development'
        if not doc:
            result = {'handled': True, 'ok': False, 'reply_text': '查無唯一客戶，請補電話、地址或客戶ID'}
        else:
            result = {'handled': True, 'ok': True, 'reply_text': format_record_timeline(target_type, doc, limit=fields.get('limit', 10)), 'target_type': target_type, 'target_id': doc.id, 'customer_name': (doc.to_dict() or {}).get('name', ''), 'phone': (doc.to_dict() or {}).get('phone', ''), 'parsed_tag': parsed.get('tag', '')}
    elif action == 'query_contract_end':
        ok, text, ctx = query_contract_end_text(fields)
        result = {'handled': True, 'ok': ok, 'reply_text': text}
        if ctx:
            result.update(ctx)
    else:
        target_type = fields.get('target_type', '')
        if action == 'buyer_followup':
            target_type = 'buyer'
        elif action == 'seller_followup':
            target_type = 'seller'
        if target_type not in ('buyer', 'seller'):
            result = {'handled': True, 'ok': False, 'reply_text': '請提供對象：買方 或 賣方'}
        else:
            doc = find_customer_record(target_type=target_type, record_id=fields.get('record_id', ''), phone=fields.get('phone', ''), name=fields.get('name', ''), address=fields.get('address', ''))
            if not doc:
                result = {'handled': True, 'ok': False, 'reply_text': '找不到唯一客戶，請補客戶ID或正確電話'}
            else:
                doc_ref = db.collection('buyers' if target_type == 'buyer' else 'sellers').document(doc.id)
                labels = dedupe_keep_order(['LINE紀錄'] + ensure_list(fields.get('labels')))
                summary_parts = []
                if fields.get('content'):
                    summary_parts.append(fields['content'])
                if fields.get('address'):
                    summary_parts.append(f"地址/物件：{fields['address']}")
                if fields.get('price'):
                    summary_parts.append(f"價格：{fields['price']}")
                summary_text = build_line_summary('；'.join(summary_parts).strip() or 'LINE 更新', event)
                source_value = _store_source_value(target_type, incoming_source=fields.get('source', ''), current_source=(doc.to_dict() or {}).get('source', ''), url=(doc.to_dict() or {}).get('url', ''))
                update_customer_note_and_labels(target_type=target_type, doc_ref=doc_ref, content=summary_text, labels=labels, stage=fields.get('stage', ''), source=source_value or None, event=event)
                if action in ('buyer_followup', 'seller_followup', 'classify'):
                    add_customer_followup(target_type=target_type, customer_id=doc.id, content=summary_text, next_action=fields.get('next_action', ''), next_contact_date=fields.get('next_contact_date', ''), labels=labels, line_event=event, stage=fields.get('stage', ''))
                label_text = '客需' if target_type == 'buyer' else '委託'
                current_data = doc_ref.get().to_dict() or {}
                result = {'handled': True, 'ok': True, 'reply_text': f"已註記{label_text}：{current_data.get('name', '')}（{current_data.get('phone','-')}）", 'target_type': target_type, 'target_id': doc.id, 'customer_name': current_data.get('name', ''), 'phone': current_data.get('phone', ''), 'parsed_tag': parsed.get('tag', '')}

    save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', target_type=result.get('target_type', ''), target_id=result.get('target_id', ''), note=result.get('reply_text', ''), sender_display_name=sender_display_name)
    if result.get('ok') and result.get('target_type') and result.get('target_id'):
        incoming_message_id = message.get('id', '')
        save_line_message_link(incoming_message_id, result['target_type'], result['target_id'], tag=result.get('parsed_tag', ''), action=action, customer_name=result.get('customer_name', ''), phone=result.get('phone', ''), source_event=event)
    return result


# ========= LINE Bot 代辦事項指令 Patch =========
# 使用方式：
# #新增代辦
# 日期: 2026-05-29
# 事項: 打給王小姐確認貸款資料
# 備註: 可空白
#
# #今日代辦
# #查詢代辦
# 日期: 2026-05-29
#
# #完成代辦
# ID: abc123
#
# #清除代辦
# ID: abc123

from datetime import timedelta

LINE_TODO_COLLECTION = os.environ.get('LINE_TODO_COLLECTION', 'line_todos')


def _line_todo_target_from_event(event):
    """取得這則 LINE 訊息要綁定的提醒對象：群組、聊天室或個人。"""
    source = event.get('source') or {}
    if source.get('groupId'):
        return source.get('groupId'), 'group'
    if source.get('roomId'):
        return source.get('roomId'), 'room'
    return source.get('userId', ''), 'user'


def _parse_line_todo_date(value: str):
    """支援 今天/明天/後天、YYYY-MM-DD、YYYY/MM/DD、MM/DD。回傳 YYYY-MM-DD。"""
    raw = (value or '').strip()
    today = now_taipei().date()
    if not raw:
        return today.strftime('%Y-%m-%d')

    aliases = {
        '今天': 0,
        '今日': 0,
        '明天': 1,
        '明日': 1,
        '後天': 2,
    }
    if raw in aliases:
        return (today + timedelta(days=aliases[raw])).strftime('%Y-%m-%d')

    m = re.search(r'(\d{4})[/-](\d{1,2})[/-](\d{1,2})', raw)
    if m:
        y, mo, d = map(int, m.groups())
        try:
            return datetime(y, mo, d, tzinfo=TAIPEI_TZ).date().strftime('%Y-%m-%d')
        except Exception:
            return ''

    m = re.search(r'(?<!\d)(\d{1,2})[/-](\d{1,2})(?!\d)', raw)
    if m:
        mo, d = map(int, m.groups())
        y = today.year
        try:
            return datetime(y, mo, d, tzinfo=TAIPEI_TZ).date().strftime('%Y-%m-%d')
        except Exception:
            return ''

    return ''


def _split_inline_todo_text(text: str):
    """處理：#新增代辦 明天 打給客戶。"""
    s = (text or '').strip()
    if not s:
        return '', ''
    parts = s.split(maxsplit=1)
    possible_date = _parse_line_todo_date(parts[0]) if parts else ''
    if possible_date and len(parts) > 1:
        return possible_date, parts[1].strip()
    return '', s


def _parse_line_todo_fields(raw_text: str, tag: str):
    lines = [ln.rstrip() for ln in (raw_text or '').splitlines()]
    if not lines:
        return {}

    first = lines[0].strip()
    inline = first.replace('#' + tag, '', 1).strip() if first.startswith('#' + tag) else ''
    fields = {}
    note_lines = []

    key_map = {
        'ID': 'todo_id',
        'id': 'todo_id',
        '編號': 'todo_id',
        '代辦ID': 'todo_id',
        '事項': 'title',
        '代辦': 'title',
        '內容': 'title',
        '工作': 'title',
        '日期': 'todo_date_raw',
        '時間': 'todo_date_raw',
        '期限': 'todo_date_raw',
        '提醒日期': 'todo_date_raw',
        '備註': 'note',
        '說明': 'note',
    }

    if inline:
        inline_date, inline_title = _split_inline_todo_text(inline)
        if inline_date:
            fields['todo_date'] = inline_date
        if inline_title:
            if tag in ('完成代辦', '清除代辦', '刪除代辦'):
                fields['todo_id'] = inline_title
            else:
                fields['title'] = inline_title

    for raw in lines[1:]:
        line = raw.strip()
        if not line:
            continue
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if m:
            key = key_map.get(m.group(1).strip(), m.group(1).strip())
            value = (m.group(2) or '').strip()
            fields[key] = value
        else:
            note_lines.append(line)

    if note_lines:
        if tag in ('完成代辦', '清除代辦', '刪除代辦') and not fields.get('todo_id'):
            fields['todo_id'] = note_lines[0]
        elif not fields.get('title'):
            fields['title'] = note_lines[0]
            if len(note_lines) > 1:
                fields['note'] = '\n'.join(note_lines[1:]).strip()
        else:
            fields['note'] = (fields.get('note', '') + '\n' + '\n'.join(note_lines)).strip()

    if fields.get('todo_date_raw') and not fields.get('todo_date'):
        fields['todo_date'] = _parse_line_todo_date(fields.get('todo_date_raw'))
    elif not fields.get('todo_date'):
        fields['todo_date'] = now_taipei().strftime('%Y-%m-%d')

    return fields


def _format_line_todo_list(items, title):
    if not items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    for idx, item in enumerate(items, 1):
        data = item.to_dict() or {}
        short_id = item.id[:6]
        note = (data.get('note') or '').strip()
        line = f"{idx}. [{short_id}] {data.get('title', '')}"
        if note:
            line += f"\n   備註: {note}"
        lines.append(line)
    lines.append('')
    lines.append('完成請回：')
    lines.append('#完成代辦')
    lines.append('ID: 上方編號')
    return '\n'.join(lines)[:5000]


def _get_open_line_todos(todo_date='', target_id=''):
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    docs = list(db.collection(LINE_TODO_COLLECTION).where('todo_date', '==', query_date).stream())
    result = []
    for doc in docs:
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if target_id and data.get('line_target_id') != target_id:
            continue
        result.append(doc)
    result.sort(key=lambda d: ((d.to_dict() or {}).get('created_at', ''), d.id))
    return result


def _find_line_todo(todo_key: str, target_id=''):
    key = (todo_key or '').strip()
    if not key:
        return None, '請提供代辦 ID 或事項關鍵字。'

    # 先用完整文件 ID 找
    direct = db.collection(LINE_TODO_COLLECTION).document(key).get()
    if direct.exists:
        data = direct.to_dict() or {}
        if target_id and data.get('line_target_id') != target_id:
            return None, '這筆代辦不是在目前這個 LINE 對話建立的。'
        return direct, ''

    matches = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if target_id and data.get('line_target_id') != target_id:
            continue
        title = data.get('title', '')
        if doc.id.startswith(key) or key in title:
            matches.append(doc)

    if len(matches) == 1:
        return matches[0], ''
    if len(matches) > 1:
        preview = '\n'.join([f"- [{d.id[:6]}] {(d.to_dict() or {}).get('title','')}" for d in matches[:8]])
        return None, '找到多筆代辦，請用 ID 完成：\n' + preview
    return None, '找不到這筆未完成代辦，請先輸入 #今日代辦 查看 ID。'


def create_line_todo(fields, event):
    title = (fields.get('title') or '').strip()
    todo_date = _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    note = (fields.get('note') or '').strip()

    if not title:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：請填「事項」。\n\n範例：\n#新增代辦\n日期: 明天\n事項: 打給王小姐確認貸款資料'}
    if not todo_date:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：日期格式看不懂，請用 2026-05-29、5/29、今天、明天。'}

    target_id, target_type = _line_todo_target_from_event(event)
    source = event.get('source') or {}
    sender_display_name = get_line_sender_display_name(event)
    now = now_taipei().isoformat()

    doc_ref = db.collection(LINE_TODO_COLLECTION).document()
    doc_ref.set({
        'title': title,
        'todo_date': todo_date,
        'note': note,
        'status': 'open',
        'line_target_id': target_id,
        'line_target_type': target_type,
        'line_group_id': source.get('groupId', ''),
        'line_room_id': source.get('roomId', ''),
        'line_user_id': source.get('userId', ''),
        'sender_display_name': sender_display_name,
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': sender_display_name or 'LINE Bot',
        'reminder_sent_dates': [],
    })

    return {
        'handled': True,
        'ok': True,
        'reply_text': f"已新增代辦：{title}\n日期：{todo_date}\nID：{doc_ref.id[:6]}\n\n當天提醒清單會包含這筆代辦。",
        'parsed_tag': '新增代辦',
    }


def query_line_todos(fields, event, force_today=False):
    target_id, _ = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    items = _get_open_line_todos(todo_date=todo_date, target_id=target_id)
    title = f'{todo_date} 代辦清單'
    return {'handled': True, 'ok': True, 'reply_text': _format_line_todo_list(items, title), 'parsed_tag': '查詢代辦'}


def complete_line_todo(fields, event):
    target_id, _ = _line_todo_target_from_event(event)
    todo_key = fields.get('todo_id') or fields.get('title') or ''
    doc, err = _find_line_todo(todo_key, target_id=target_id)
    if err:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '完成代辦'}

    data = doc.to_dict() or {}
    doc.reference.update({
        'status': 'done',
        'done_at': now_taipei().isoformat(),
        'done_by_id': 'line_bot',
        'done_by_name': get_line_sender_display_name(event) or 'LINE Bot',
        'updated_at': now_taipei().isoformat(),
    })
    return {'handled': True, 'ok': True, 'reply_text': f"已完成並從未完成清單移除：{data.get('title', '')}", 'parsed_tag': '完成代辦'}


def delete_line_todo(fields, event):
    target_id, _ = _line_todo_target_from_event(event)
    todo_key = fields.get('todo_id') or fields.get('title') or ''
    doc, err = _find_line_todo(todo_key, target_id=target_id)
    if err:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '清除代辦'}
    data = doc.to_dict() or {}
    doc.reference.delete()
    return {'handled': True, 'ok': True, 'reply_text': f"已清除代辦：{data.get('title', '')}", 'parsed_tag': '清除代辦'}


def process_line_todo_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if not raw_text.startswith('#'):
        return {'handled': False}

    command_table = [
        ('新增代辦', create_line_todo),
        ('今日代辦', lambda fields, ev: query_line_todos(fields, ev, force_today=True)),
        ('查詢代辦', query_line_todos),
        ('代辦', lambda fields, ev: query_line_todos(fields, ev, force_today=True)),
        ('完成代辦', complete_line_todo),
        ('清除代辦', delete_line_todo),
        ('刪除代辦', delete_line_todo),
    ]

    for tag, handler in command_table:
        if raw_text.startswith('#' + tag):
            fields = _parse_line_todo_fields(raw_text, tag)
            result = handler(fields, event)
            parsed = {'tag': result.get('parsed_tag', tag), 'action': 'line_todo', 'fields': fields, 'raw_text': raw_text}
            save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
            return result

    return {'handled': False}


def push_line_text(to_id: str, text_message: str):
    if not LINE_CHANNEL_ACCESS_TOKEN or not to_id:
        return False, 'LINE_CHANNEL_ACCESS_TOKEN 或 to_id 為空'
    payload = {
        'to': to_id,
        'messages': [{'type': 'text', 'text': (text_message or '')[:5000]}],
    }
    try:
        import requests
        res = requests.post(
            'https://api.line.me/v2/bot/message/push',
            headers=line_api_headers(),
            json=payload,
            timeout=8,
        )
        print('LINE push status:', res.status_code, res.text[:300])
        return res.status_code in (200, 202), res.text[:300]
    except Exception as e:
        print('⚠️ LINE push 發生錯誤：', e)
        return False, str(e)


def send_today_line_todo_reminders():
    today = now_taipei().strftime('%Y-%m-%d')
    docs = list(db.collection(LINE_TODO_COLLECTION).where('todo_date', '==', today).stream())
    grouped = {}

    for doc in docs:
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        sent_dates = data.get('reminder_sent_dates') or []
        if today in sent_dates:
            continue
        target_id = data.get('line_target_id', '')
        if not target_id:
            continue
        grouped.setdefault(target_id, []).append(doc)

    sent_count = 0
    failed = []
    for target_id, items in grouped.items():
        text = _format_line_todo_list(items, f'今天 {today} 要做的事情')
        ok, msg = push_line_text(target_id, text)
        if ok:
            sent_count += 1
            for doc in items:
                doc.reference.update({
                    'reminder_sent_dates': firestore.ArrayUnion([today]),
                    'last_reminded_at': now_taipei().isoformat(),
                })
        else:
            failed.append({'target_id': target_id, 'error': msg})

    return {'date': today, 'target_count': len(grouped), 'sent_count': sent_count, 'failed': failed}


@app.route('/line/todos/remind-today', methods=['GET', 'POST'])
def line_todos_remind_today():
    # 建議在 Render / Cron Job 設定環境變數 TODO_REMINDER_SECRET，並用 ?key=你的密鑰 呼叫。
    secret = os.environ.get('TODO_REMINDER_SECRET', '').strip()
    key = request.args.get('key', '').strip() or request.form.get('key', '').strip()
    if secret and key != secret:
        return {'ok': False, 'message': 'Invalid key'}, 403
    result = send_today_line_todo_reminders()
    return {'ok': True, 'result': result}, 200


# 保留原本 LINE 訊息處理流程；代辦指令先攔截，其它指令照舊走原本客需/委託/開發流程。
_process_line_message_event_before_todo = process_line_message_event


def process_line_message_event(event):
    todo_result = process_line_todo_message_event(event)
    if todo_result.get('handled'):
        return todo_result
    return _process_line_message_event_before_todo(event)

# ========= LINE Bot 代辦事項指令 Patch End =========

# ========= 原本 app.run 已移到檔案最底部，避免後續 Patch / Route 未載入 =========

# ========= LINE Bot 代辦事項「多行批次新增 / 用序號完成」強化 Patch v2 =========
# 貼在原本代辦事項 Patch 的最底部即可。
# 支援格式：
# #新增代辦
# 日期: 明天
# 厝米排版 土地現廣稿
# 拍水哥爸爸土地
# 遠6屋主
# 下午要去找刺青大哥
#
# 也支援：#完成代辦 1 2 3  或  #完成代辦\n1\n3


def _clean_todo_item_line(line: str) -> str:
    """把清單符號、序號拿掉，保留真正的代辦文字。"""
    s = (line or '').strip()
    s = re.sub(r'^[-*•●▪▫□☐✅✔]+\s*', '', s)
    s = re.sub(r'^\d+[\.\)）、．]\s*', '', s)
    return s.strip()


def _looks_like_date_text(text: str) -> bool:
    raw = (text or '').strip()
    if raw in ('今天', '今日', '明天', '明日', '後天'):
        return True
    if re.fullmatch(r'\d{4}[/-]\d{1,2}[/-]\d{1,2}', raw):
        return True
    if re.fullmatch(r'\d{1,2}[/-]\d{1,2}', raw):
        return True
    return False


def _parse_line_todo_bulk_fields(raw_text: str, tag: str = '新增代辦'):
    """
    批次新增解析：
    - 日期只要寫一次，下面每一行都會變成一筆代辦
    - 沒寫日期就預設今天
    - 仍相容舊格式「事項: xxx」
    """
    lines = [ln.strip() for ln in (raw_text or '').splitlines() if ln.strip()]
    if not lines:
        return []

    first = lines[0]
    inline = first.replace('#' + tag, '', 1).strip() if first.startswith('#' + tag) else ''

    default_date = now_taipei().strftime('%Y-%m-%d')
    default_note = ''
    items = []

    # 例如：#新增代辦 明天
    # 或：#新增代辦 明天 打給王小姐
    if inline:
        if _looks_like_date_text(inline):
            parsed_date = _parse_line_todo_date(inline)
            if parsed_date:
                default_date = parsed_date
        else:
            inline_date, inline_title = _split_inline_todo_text(inline)
            if inline_date and inline_title:
                default_date = inline_date
                items.append(inline_title)
            else:
                items.append(inline)

    key_map = {
        '日期': 'todo_date_raw',
        '時間': 'todo_date_raw',
        '期限': 'todo_date_raw',
        '提醒日期': 'todo_date_raw',
        '事項': 'title',
        '代辦': 'title',
        '內容': 'title',
        '工作': 'title',
        '備註': 'note',
        '說明': 'note',
    }

    for raw_line in lines[1:]:
        line = raw_line.strip()
        if not line:
            continue

        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if m:
            key = key_map.get(m.group(1).strip(), m.group(1).strip())
            value = (m.group(2) or '').strip()
            if key == 'todo_date_raw':
                parsed_date = _parse_line_todo_date(value)
                if parsed_date:
                    default_date = parsed_date
            elif key == 'note':
                default_note = value
            elif key == 'title':
                title = _clean_todo_item_line(value)
                if title:
                    items.append(title)
            continue

        cleaned = _clean_todo_item_line(line)
        if not cleaned:
            continue

        # 允許這種格式：
        # #新增代辦
        # 明天
        # 打給王小姐
        if not items and _looks_like_date_text(cleaned):
            parsed_date = _parse_line_todo_date(cleaned)
            if parsed_date:
                default_date = parsed_date
                continue

        items.append(cleaned)

    # 去重但保留順序，避免同一行不小心貼兩次
    items = dedupe_keep_order(items)

    return [
        {
            'title': title,
            'todo_date': default_date,
            'note': default_note,
        }
        for title in items
        if title
    ]


def create_line_todos_bulk(fields_list, event):
    if not fields_list:
        return {
            'handled': True,
            'ok': False,
            'reply_text': (
                '未新增：請在 #新增代辦 下一行開始貼代辦事項。\n\n'
                '範例：\n'
                '#新增代辦\n'
                '日期: 明天\n'
                '厝米排版 土地現廣稿\n'
                '拍水哥爸爸土地\n'
                '遠6屋主\n'
                '下午要去找刺青大哥'
            ),
            'parsed_tag': '新增代辦',
        }

    ok_titles = []
    failed_msgs = []

    for fields in fields_list:
        result = create_line_todo(fields, event)
        if result.get('ok'):
            ok_titles.append((fields.get('todo_date', ''), fields.get('title', '')))
        else:
            failed_msgs.append(result.get('reply_text', '新增失敗'))

    if ok_titles:
        # 多數情況同一天，先拿第一筆日期當摘要
        date_label = ok_titles[0][0]
        lines = [f'已新增 {len(ok_titles)} 筆代辦', f'日期：{date_label}', '']
        for idx, (_, title) in enumerate(ok_titles, 1):
            lines.append(f'{idx}. {title}')
        lines.append('')
        lines.append('查詢請回：#今日代辦')
        lines.append('完成可回：#完成代辦 1')
        reply = '\n'.join(lines)
        if failed_msgs:
            reply += '\n\n有部分未新增：\n' + '\n'.join(failed_msgs[:3])
        return {'handled': True, 'ok': True, 'reply_text': reply[:5000], 'parsed_tag': '新增代辦'}

    return {'handled': True, 'ok': False, 'reply_text': '\n'.join(failed_msgs)[:5000], 'parsed_tag': '新增代辦'}


def _extract_todo_keys_from_command(raw_text: str, tag: str):
    """支援 #完成代辦 1 2 3、ID: xxx、或換行多個序號。"""
    lines = [ln.strip() for ln in (raw_text or '').splitlines() if ln.strip()]
    if not lines:
        return []

    first = lines[0]
    inline = first.replace('#' + tag, '', 1).strip() if first.startswith('#' + tag) else ''
    chunks = []
    if inline:
        chunks.append(inline)

    for line in lines[1:]:
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if m:
            key = m.group(1).strip()
            value = (m.group(2) or '').strip()
            if key in ('ID', 'id', '編號', '代辦ID', '事項', '代辦') and value:
                chunks.append(value)
            # 日期只用來判斷序號查哪一天，不當成 key
        else:
            chunks.append(line)

    keys = []
    for chunk in chunks:
        for part in re.split(r'[\s,，、]+', chunk):
            part = part.strip()
            if part:
                keys.append(part)
    return dedupe_keep_order(keys)


def _find_line_todo_v2(todo_key: str, target_id='', todo_date=''):
    key = (todo_key or '').strip()
    if not key:
        return None, '請提供代辦序號、ID 或事項關鍵字。'

    # 讓使用者可以直接用 #今日代辦 的序號完成，例如：#完成代辦 2
    if re.fullmatch(r'\d+', key):
        items = _get_open_line_todos(
            todo_date=todo_date or now_taipei().strftime('%Y-%m-%d'),
            target_id=target_id,
        )
        idx = int(key) - 1
        if 0 <= idx < len(items):
            return items[idx], ''
        return None, f'找不到第 {key} 筆代辦，請先輸入 #今日代辦 確認序號。'

    return _find_line_todo(key, target_id=target_id)


def complete_line_todos_bulk(fields, event, raw_text='', delete_mode=False):
    target_id, _ = _line_todo_target_from_event(event)
    tag = '清除代辦' if delete_mode else '完成代辦'
    todo_date = fields.get('todo_date') or fields.get('todo_date_raw') or now_taipei().strftime('%Y-%m-%d')
    todo_date = _parse_line_todo_date(todo_date) or now_taipei().strftime('%Y-%m-%d')

    keys = _extract_todo_keys_from_command(raw_text, tag)
    if not keys:
        fallback = fields.get('todo_id') or fields.get('title') or ''
        if fallback:
            keys = [fallback]

    if not keys:
        return {
            'handled': True,
            'ok': False,
            'reply_text': '請提供要完成的序號或 ID。\n例如：#完成代辦 1\n或：#完成代辦 1 3',
            'parsed_tag': tag,
        }

    done_titles = []
    errors = []

    for key in keys:
        doc, err = _find_line_todo_v2(key, target_id=target_id, todo_date=todo_date)
        if err:
            errors.append(f'{key}：{err}')
            continue

        data = doc.to_dict() or {}
        title = data.get('title', '')
        if delete_mode:
            doc.reference.delete()
            done_titles.append(f'已清除：{title}')
        else:
            doc.reference.update({
                'status': 'done',
                'done_at': now_taipei().isoformat(),
                'done_by_id': 'line_bot',
                'done_by_name': get_line_sender_display_name(event) or 'LINE Bot',
                'updated_at': now_taipei().isoformat(),
            })
            done_titles.append(f'已完成：{title}')

    lines = []
    if done_titles:
        lines.extend(done_titles)
    if errors:
        if lines:
            lines.append('')
        lines.append('未處理：')
        lines.extend(errors[:8])

    return {
        'handled': True,
        'ok': bool(done_titles),
        'reply_text': '\n'.join(lines)[:5000] if lines else '沒有完成任何代辦。',
        'parsed_tag': tag,
    }


# 覆寫代辦指令處理：新增代辦改成可批次，多筆完成/清除也可用序號。
def process_line_todo_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if not raw_text.startswith('#'):
        return {'handled': False}

    if raw_text.startswith('#新增代辦'):
        fields_list = _parse_line_todo_bulk_fields(raw_text, '新增代辦')
        result = create_line_todos_bulk(fields_list, event)
        parsed = {'tag': result.get('parsed_tag', '新增代辦'), 'action': 'line_todo_bulk_create', 'fields': {'items': fields_list}, 'raw_text': raw_text}
        save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#完成代辦'):
        fields = _parse_line_todo_fields(raw_text, '完成代辦')
        result = complete_line_todos_bulk(fields, event, raw_text=raw_text, delete_mode=False)
        parsed = {'tag': result.get('parsed_tag', '完成代辦'), 'action': 'line_todo_bulk_done', 'fields': fields, 'raw_text': raw_text}
        save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#清除代辦') or raw_text.startswith('#刪除代辦'):
        tag = '清除代辦' if raw_text.startswith('#清除代辦') else '刪除代辦'
        fields = _parse_line_todo_fields(raw_text, tag)
        result = complete_line_todos_bulk(fields, event, raw_text=raw_text, delete_mode=True)
        parsed = {'tag': result.get('parsed_tag', tag), 'action': 'line_todo_bulk_delete', 'fields': fields, 'raw_text': raw_text}
        save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    command_table = [
        ('今日代辦', lambda fields, ev: query_line_todos(fields, ev, force_today=True)),
        ('查詢代辦', query_line_todos),
        ('代辦', lambda fields, ev: query_line_todos(fields, ev, force_today=True)),
    ]

    for tag, handler in command_table:
        if raw_text.startswith('#' + tag):
            fields = _parse_line_todo_fields(raw_text, tag)
            result = handler(fields, event)
            parsed = {'tag': result.get('parsed_tag', tag), 'action': 'line_todo_query', 'fields': fields, 'raw_text': raw_text}
            save_line_log(parsed, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
            return result

    return {'handled': False}

# ========= LINE Bot 代辦事項強化 Patch v2 End =========

# ========= LINE Bot 代辦事項「逾期尚未完成 / 隔天延續提醒」強化 Patch v3 =========
# 貼在 v2 代辦事項 Patch 的最底部即可。
# 功能：
# 1. 今天以前未完成的代辦，隔天會出現在「尚未完成」區塊。
# 2. #今日代辦 會顯示：尚未完成（以前）＋ 今天要做。
# 3. 每日提醒也會推送：尚未完成（以前）＋ 今天要做。
# 4. #完成代辦 1 會依照畫面上的序號完成，包含尚未完成區塊。


def _todo_date_value(doc):
    data = doc.to_dict() or {}
    return (data.get('todo_date') or '').strip()


def _is_open_todo_doc(doc, target_id=''):
    data = doc.to_dict() or {}
    if data.get('status', 'open') != 'open':
        return False
    if target_id and data.get('line_target_id') != target_id:
        return False
    if not (data.get('todo_date') or '').strip():
        return False
    return True


def _sort_line_todo_docs(items):
    return sorted(
        items,
        key=lambda d: (
            (d.to_dict() or {}).get('todo_date', ''),
            (d.to_dict() or {}).get('created_at', ''),
            d.id,
        )
    )


def _get_open_line_todos(todo_date='', target_id='', include_overdue=False):
    """
    覆寫 v2 的查詢邏輯。
    include_overdue=False：只查指定日期。
    include_overdue=True：查指定日期以前含當天，讓未完成代辦隔天延續顯示。
    """
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    result = []

    # 用 stream + Python 過濾，避免 Firestore 複合索引問題，也能兼容舊資料。
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = _todo_date_value(doc)
        if include_overdue:
            if d <= query_date:
                result.append(doc)
        else:
            if d == query_date:
                result.append(doc)

    return _sort_line_todo_docs(result)


def _get_overdue_line_todos(todo_date='', target_id=''):
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    result = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = _todo_date_value(doc)
        if d < query_date:
            result.append(doc)
    return _sort_line_todo_docs(result)


def _get_display_line_todos(todo_date='', target_id=''):
    """回傳畫面顯示順序：逾期未完成在前，指定日期在後。"""
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    overdue_items = _get_overdue_line_todos(query_date, target_id=target_id)
    today_items = _get_open_line_todos(query_date, target_id=target_id, include_overdue=False)
    return overdue_items + today_items


def _format_line_todo_sections(overdue_items, today_items, title, today_label='今天'):
    if not overdue_items and not today_items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    idx = 1

    if overdue_items:
        lines.append('')
        lines.append('【尚未完成】')
        for doc in overdue_items:
            data = doc.to_dict() or {}
            short_id = doc.id[:6]
            note = (data.get('note') or '').strip()
            old_date = (data.get('todo_date') or '').strip()
            lines.append(f"{idx}. [{short_id}] {old_date}｜{data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if today_items:
        lines.append('')
        lines.append(f'【{today_label}要做】')
        for doc in today_items:
            data = doc.to_dict() or {}
            short_id = doc.id[:6]
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. [{short_id}] {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    lines.append('一次完成多筆：#完成代辦 1 3')
    lines.append('清除請回：#清除代辦 1')
    return '\n'.join(lines)[:5000]


def query_line_todos(fields, event, force_today=False):
    target_id, _ = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    if not todo_date:
        todo_date = now_taipei().strftime('%Y-%m-%d')

    overdue_items = _get_overdue_line_todos(todo_date=todo_date, target_id=target_id)
    today_items = _get_open_line_todos(todo_date=todo_date, target_id=target_id, include_overdue=False)

    if todo_date == now_taipei().strftime('%Y-%m-%d'):
        title = f'{todo_date} 代辦清單'
        today_label = '今天'
    else:
        title = f'{todo_date} 代辦清單'
        today_label = todo_date

    return {
        'handled': True,
        'ok': True,
        'reply_text': _format_line_todo_sections(overdue_items, today_items, title, today_label=today_label),
        'parsed_tag': '查詢代辦',
    }


def _find_line_todo_v2(todo_key: str, target_id='', todo_date=''):
    """
    覆寫 v2 的序號完成邏輯。
    序號會依照 #今日代辦 畫面順序：尚未完成在前，今天要做在後。
    """
    key = (todo_key or '').strip()
    if not key:
        return None, '請提供代辦序號、ID 或事項關鍵字。'

    query_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')

    if re.fullmatch(r'\d+', key):
        items = _get_display_line_todos(todo_date=query_date, target_id=target_id)
        idx = int(key) - 1
        if 0 <= idx < len(items):
            return items[idx], ''
        return None, f'找不到第 {key} 筆代辦，請先輸入 #今日代辦 確認序號。'

    return _find_line_todo(key, target_id=target_id)


def send_today_line_todo_reminders():
    """
    覆寫 v2 的每日提醒。
    今天提醒會包含：
    - 今天以前尚未完成的代辦
    - 今天要做的代辦
    每天只會對同一筆代辦提醒一次；隔天還沒完成，會再提醒一次。
    """
    today = now_taipei().strftime('%Y-%m-%d')
    grouped = {}

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        todo_date = (data.get('todo_date') or '').strip()
        if not todo_date or todo_date > today:
            continue
        sent_dates = data.get('reminder_sent_dates') or []
        if today in sent_dates:
            continue
        target_id = data.get('line_target_id', '')
        if not target_id:
            continue
        grouped.setdefault(target_id, []).append(doc)

    sent_count = 0
    failed = []

    for target_id, items in grouped.items():
        overdue_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) < today])
        today_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) == today])
        text = _format_line_todo_sections(
            overdue_items,
            today_items,
            f'今日代辦提醒 {today}',
            today_label='今天',
        )
        ok, msg = push_line_text(target_id, text)
        if ok:
            sent_count += 1
            for doc in items:
                doc.reference.update({
                    'reminder_sent_dates': firestore.ArrayUnion([today]),
                    'last_reminded_at': now_taipei().isoformat(),
                })
        else:
            failed.append({'target_id': target_id, 'error': msg})

    return {'date': today, 'target_count': len(grouped), 'sent_count': sent_count, 'failed': failed}

# ========= LINE Bot 代辦事項強化 Patch v3 End =========



# ========= LINE Bot 代辦事項顯示日期 Patch v4：只顯示月/日 =========
# 貼在 v3 代辦事項 Patch 的最底部即可。
# 注意：Firestore 裡的 todo_date 仍然保留 YYYY-MM-DD，方便判斷逾期；只有 LINE 顯示文字改成 M/D。


def _todo_display_md(date_text: str) -> str:
    """把 YYYY-MM-DD / YYYY/MM/DD / MM-DD 轉成 M/D 顯示；解析失敗就回傳原文字。"""
    raw = (date_text or '').strip()
    if not raw:
        return ''

    # YYYY-MM-DD 或 YYYY/MM/DD -> M/D
    m = re.search(r'^(\d{4})[/-](\d{1,2})[/-](\d{1,2})$', raw)
    if m:
        _, mo, d = m.groups()
        return f'{int(mo)}/{int(d)}'

    # MM-DD 或 MM/DD -> M/D
    m = re.search(r'^(\d{1,2})[/-](\d{1,2})$', raw)
    if m:
        mo, d = m.groups()
        return f'{int(mo)}/{int(d)}'

    return raw


def create_line_todo(fields, event):
    """覆寫新增回覆：日期顯示改成 M/D，但資料庫仍存 YYYY-MM-DD。"""
    title = (fields.get('title') or '').strip()
    todo_date = _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    note = (fields.get('note') or '').strip()

    if not title:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：請填「事項」。\n\n範例：\n#新增代辦\n日期: 明天\n事項: 打給王小姐確認貸款資料'}
    if not todo_date:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：日期格式看不懂，請用 5/29、今天、明天。'}

    target_id, target_type = _line_todo_target_from_event(event)
    source = event.get('source') or {}
    sender_display_name = get_line_sender_display_name(event)
    now = now_taipei().isoformat()

    doc_ref = db.collection(LINE_TODO_COLLECTION).document()
    doc_ref.set({
        'title': title,
        'todo_date': todo_date,  # 資料庫維持完整日期
        'note': note,
        'status': 'open',
        'line_target_id': target_id,
        'line_target_type': target_type,
        'line_group_id': source.get('groupId', ''),
        'line_room_id': source.get('roomId', ''),
        'line_user_id': source.get('userId', ''),
        'sender_display_name': sender_display_name,
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': sender_display_name or 'LINE Bot',
        'reminder_sent_dates': [],
    })

    return {
        'handled': True,
        'ok': True,
        'reply_text': f"已新增代辦：{title}\n日期：{_todo_display_md(todo_date)}\nID：{doc_ref.id[:6]}\n\n當天提醒清單會包含這筆代辦。",
        'parsed_tag': '新增代辦',
    }


def create_line_todos_bulk(fields_list, event):
    """覆寫批次新增摘要：日期顯示改成 M/D。"""
    if not fields_list:
        return {
            'handled': True,
            'ok': False,
            'reply_text': (
                '未新增：請在 #新增代辦 下一行開始貼代辦事項。\n\n'
                '範例：\n'
                '#新增代辦\n'
                '日期: 明天\n'
                '厝米排版 土地現廣稿\n'
                '拍水哥爸爸土地\n'
                '遠6屋主\n'
                '下午要去找刺青大哥'
            ),
            'parsed_tag': '新增代辦',
        }

    ok_titles = []
    failed_msgs = []

    for fields in fields_list:
        result = create_line_todo(fields, event)
        if result.get('ok'):
            ok_titles.append((fields.get('todo_date', ''), fields.get('title', '')))
        else:
            failed_msgs.append(result.get('reply_text', '新增失敗'))

    if ok_titles:
        date_label = _todo_display_md(ok_titles[0][0])
        lines = [f'已新增 {len(ok_titles)} 筆代辦', f'日期：{date_label}', '']
        for idx, (_, title) in enumerate(ok_titles, 1):
            lines.append(f'{idx}. {title}')
        lines.append('')
        lines.append('查詢請回：#今日代辦')
        lines.append('完成可回：#完成代辦 1')
        reply = '\n'.join(lines)
        if failed_msgs:
            reply += '\n\n有部分未新增：\n' + '\n'.join(failed_msgs[:3])
        return {'handled': True, 'ok': True, 'reply_text': reply[:5000], 'parsed_tag': '新增代辦'}

    return {'handled': True, 'ok': False, 'reply_text': '\n'.join(failed_msgs)[:5000], 'parsed_tag': '新增代辦'}


def _format_line_todo_sections(overdue_items, today_items, title, today_label='今天'):
    """覆寫清單顯示：日期只顯示 M/D。"""
    if not overdue_items and not today_items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    idx = 1

    if overdue_items:
        lines.append('')
        lines.append('【尚未完成】')
        for doc in overdue_items:
            data = doc.to_dict() or {}
            short_id = doc.id[:6]
            note = (data.get('note') or '').strip()
            old_date = _todo_display_md((data.get('todo_date') or '').strip())
            lines.append(f"{idx}. [{short_id}] {old_date}｜{data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if today_items:
        lines.append('')
        lines.append(f'【{today_label}要做】')
        for doc in today_items:
            data = doc.to_dict() or {}
            short_id = doc.id[:6]
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. [{short_id}] {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    lines.append('一次完成多筆：#完成代辦 1 3')
    lines.append('清除請回：#清除代辦 1')
    return '\n'.join(lines)[:5000]


def _format_line_todo_list(items, title):
    """保險覆寫舊版清單格式。"""
    if not items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    for idx, item in enumerate(items, 1):
        data = item.to_dict() or {}
        short_id = item.id[:6]
        note = (data.get('note') or '').strip()
        line = f"{idx}. [{short_id}] {data.get('title', '')}"
        if note:
            line += f"\n   備註: {note}"
        lines.append(line)
    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    return '\n'.join(lines)[:5000]


def query_line_todos(fields, event, force_today=False):
    """覆寫查詢標題：日期顯示改成 M/D。"""
    target_id, _ = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    if not todo_date:
        todo_date = now_taipei().strftime('%Y-%m-%d')

    overdue_items = _get_overdue_line_todos(todo_date=todo_date, target_id=target_id)
    today_items = _get_open_line_todos(todo_date=todo_date, target_id=target_id, include_overdue=False)

    display_date = _todo_display_md(todo_date)
    if todo_date == now_taipei().strftime('%Y-%m-%d'):
        title = f'{display_date} 代辦清單'
        today_label = '今天'
    else:
        title = f'{display_date} 代辦清單'
        today_label = display_date

    return {
        'handled': True,
        'ok': True,
        'reply_text': _format_line_todo_sections(overdue_items, today_items, title, today_label=today_label),
        'parsed_tag': '查詢代辦',
    }


def send_today_line_todo_reminders():
    """覆寫每日提醒標題：日期顯示改成 M/D，但提醒判斷仍使用 YYYY-MM-DD。"""
    today = now_taipei().strftime('%Y-%m-%d')
    grouped = {}

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        todo_date = (data.get('todo_date') or '').strip()
        if not todo_date or todo_date > today:
            continue
        sent_dates = data.get('reminder_sent_dates') or []
        if today in sent_dates:
            continue
        target_id = data.get('line_target_id', '')
        if not target_id:
            continue
        grouped.setdefault(target_id, []).append(doc)

    sent_count = 0
    failed = []

    for target_id, items in grouped.items():
        overdue_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) < today])
        today_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) == today])
        text = _format_line_todo_sections(
            overdue_items,
            today_items,
            f'今日代辦提醒 {_todo_display_md(today)}',
            today_label='今天',
        )
        ok, msg = push_line_text(target_id, text)
        if ok:
            sent_count += 1
            for doc in items:
                doc.reference.update({
                    'reminder_sent_dates': firestore.ArrayUnion([today]),
                    'last_reminded_at': now_taipei().isoformat(),
                })
        else:
            failed.append({'target_id': target_id, 'error': msg})

    return {'date': today, 'target_count': len(grouped), 'sent_count': sent_count, 'failed': failed}

# ========= LINE Bot 代辦事項顯示日期 Patch v4 End =========

# ========= LINE Bot 代辦事項 Patch v5：隱藏 ID / 多日期摘要 / 明天與今天雙提醒 =========
# 貼在 v4 代辦事項 Patch 的最底部即可。
# 功能：
# 1. LINE 清單不顯示 Firestore 文件 ID，完成/清除一律用畫面序號。
# 2. 新增代辦回覆會顯示日期；若同批有不同日期，會依日期分組顯示。
# 3. 支援每一行前面直接寫日期，例如：5/29 厝米排版、明天 拍水哥爸爸土地。
# 4. 新增 /line/todos/remind-tomorrow，用於每天晚上 23:00 推播「明天代辦」。
# 5. 原本 /line/todos/remind-today 用於每天早上 08:00 推播「今日代辦＋尚未完成」。


def _split_todo_line_date_prefix(line: str):
    """
    支援單行指定日期：
    - 5/29 厝米排版
    - 5/29｜厝米排版
    - 2026-05-29 厝米排版
    - 明天 拍水哥爸爸土地
    - 後天：找刺青大哥
    回傳：(YYYY-MM-DD 或 '', 事項文字)
    """
    s = _clean_todo_item_line(line or '')
    if not s:
        return '', ''

    # YYYY-MM-DD / YYYY/MM/DD 開頭
    m = re.match(r'^(\d{4}[/-]\d{1,2}[/-]\d{1,2})\s*[｜|:：\-—、,，]?\s*(.+)$', s)
    if m:
        d = _parse_line_todo_date(m.group(1))
        title = _clean_todo_item_line(m.group(2))
        return d, title

    # M/D 或 M-D 開頭
    m = re.match(r'^(\d{1,2}[/-]\d{1,2})\s*[｜|:：\-—、,，]?\s*(.+)$', s)
    if m:
        d = _parse_line_todo_date(m.group(1))
        title = _clean_todo_item_line(m.group(2))
        return d, title

    # 今天 / 明天 / 後天 開頭
    m = re.match(r'^(今天|今日|明天|明日|後天)\s*[｜|:：\-—、,，]?\s*(.+)$', s)
    if m:
        d = _parse_line_todo_date(m.group(1))
        title = _clean_todo_item_line(m.group(2))
        return d, title

    return '', s


def _parse_line_todo_bulk_fields(raw_text: str, tag: str = '新增代辦'):
    """
    覆寫批次新增解析 v5：
    - 日期: 5/29 會作為下面代辦的預設日期。
    - 每一行也能自己帶日期，例如「5/30 拍水哥爸爸土地」。
    - 沒寫日期仍預設今天。
    """
    lines = [ln.strip() for ln in (raw_text or '').splitlines() if ln.strip()]
    if not lines:
        return []

    first = lines[0]
    inline = first.replace('#' + tag, '', 1).strip() if first.startswith('#' + tag) else ''

    default_date = now_taipei().strftime('%Y-%m-%d')
    default_note = ''
    default_remind_time = ''
    items = []

    # 例如：#新增代辦 明天
    # 或：#新增代辦 明天 打給王小姐
    if inline:
        if _looks_like_date_text(inline):
            parsed_date = _parse_line_todo_date(inline)
            if parsed_date:
                default_date = parsed_date
        else:
            inline_date, inline_title = _split_todo_line_date_prefix(inline)
            if inline_date and inline_title:
                items.append({'title': inline_title, 'todo_date': inline_date, 'note': default_note, 'remind_time': default_remind_time})
            elif inline_title:
                items.append({'title': inline_title, 'todo_date': default_date, 'note': default_note, 'remind_time': default_remind_time})

    key_map = {
        '日期': 'todo_date_raw',
        '時間': 'todo_date_raw',
        '期限': 'todo_date_raw',
        '提醒日期': 'todo_date_raw',
        '提醒時間': 'remind_time',
        '代辦時間': 'remind_time',
        '事項': 'title',
        '代辦': 'title',
        '內容': 'title',
        '工作': 'title',
        '備註': 'note',
        '說明': 'note',
    }

    for raw_line in lines[1:]:
        line = raw_line.strip()
        if not line:
            continue

        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if m:
            key = key_map.get(m.group(1).strip(), m.group(1).strip())
            value = (m.group(2) or '').strip()
            if key == 'todo_date_raw':
                parsed_date = _parse_line_todo_date(value)
                if parsed_date:
                    default_date = parsed_date
            elif key == 'remind_time':
                default_remind_time = value
            elif key == 'note':
                default_note = value
            elif key == 'title':
                line_date, title = _split_todo_line_date_prefix(value)
                if title:
                    items.append({
                        'title': title,
                        'todo_date': line_date or default_date,
                        'note': default_note,
                        'remind_time': default_remind_time,
                    })
            continue

        cleaned = _clean_todo_item_line(line)
        if not cleaned:
            continue

        # 允許中途切換日期：
        # 5/29
        # A事項
        # 5/30
        # B事項
        if _looks_like_date_text(cleaned):
            parsed_date = _parse_line_todo_date(cleaned)
            if parsed_date:
                default_date = parsed_date
                continue

        line_date, title = _split_todo_line_date_prefix(cleaned)
        if title:
            items.append({
                'title': title,
                'todo_date': line_date or default_date,
                'note': default_note,
                'remind_time': default_remind_time,
            })

    deduped = []
    seen = set()
    for item in items:
        key = (item.get('todo_date', ''), item.get('title', ''), item.get('note', ''))
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)
    return deduped


def create_line_todo(fields, event):
    """覆寫新增：不回傳 ID；仍把完整日期存入 Firestore。"""
    title = (fields.get('title') or '').strip()
    todo_date = _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    note = (fields.get('note') or '').strip()
    remind_time = (fields.get('remind_time') or '').strip()

    if not title:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：請填「事項」。\n\n範例：\n#新增代辦\n日期: 明天\n厝米排版 土地現廣稿\n拍水哥爸爸土地'}
    if not todo_date:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：日期格式看不懂，請用 5/29、今天、明天。'}

    target_id, target_type = _line_todo_target_from_event(event)
    source = event.get('source') or {}
    sender_display_name = get_line_sender_display_name(event)
    now = now_taipei().isoformat()

    doc_ref = db.collection(LINE_TODO_COLLECTION).document()
    doc_ref.set({
        'title': title,
        'todo_date': todo_date,
        'note': note,
        'remind_time': remind_time,
        'status': 'open',
        'line_target_id': target_id,
        'line_target_type': target_type,
        'line_group_id': source.get('groupId', ''),
        'line_room_id': source.get('roomId', ''),
        'line_user_id': source.get('userId', ''),
        'sender_display_name': sender_display_name,
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': sender_display_name or 'LINE Bot',
        'reminder_sent_dates': [],
        'tomorrow_reminder_sent_dates': [],
    })

    return {
        'handled': True,
        'ok': True,
        'reply_text': f"已新增代辦\n日期：{_todo_display_md(todo_date)}\n事項：{title}\n\n查詢請回：#今日代辦",
        'parsed_tag': '新增代辦',
    }


def create_line_todos_bulk(fields_list, event):
    """覆寫批次新增摘要：不顯示 ID；同批不同日期時自動分組。"""
    if not fields_list:
        return {
            'handled': True,
            'ok': False,
            'reply_text': (
                '未新增：請在 #新增代辦 下一行開始貼代辦事項。\n\n'
                '範例一：同一天\n'
                '#新增代辦\n'
                '日期: 明天\n'
                '厝米排版 土地現廣稿\n'
                '拍水哥爸爸土地\n\n'
                '範例二：不同天\n'
                '#新增代辦\n'
                '5/29 厝米排版 土地現廣稿\n'
                '5/30 拍水哥爸爸土地'
            ),
            'parsed_tag': '新增代辦',
        }

    ok_items = []
    failed_msgs = []

    for fields in fields_list:
        result = create_line_todo(fields, event)
        if result.get('ok'):
            ok_items.append({
                'todo_date': fields.get('todo_date', ''),
                'title': fields.get('title', ''),
            })
        else:
            failed_msgs.append(result.get('reply_text', '新增失敗'))

    if ok_items:
        grouped = {}
        for item in ok_items:
            d = item.get('todo_date', '') or now_taipei().strftime('%Y-%m-%d')
            grouped.setdefault(d, []).append(item.get('title', ''))

        lines = [f'已新增 {len(ok_items)} 筆代辦']
        if len(grouped) == 1:
            only_date = next(iter(grouped.keys()))
            lines.append(f'日期：{_todo_display_md(only_date)}')
            lines.append('')
            for idx, title in enumerate(grouped[only_date], 1):
                lines.append(f'{idx}. {title}')
        else:
            lines.append('')
            running = 1
            for d in sorted(grouped.keys()):
                lines.append(f'【{_todo_display_md(d)}】')
                for title in grouped[d]:
                    lines.append(f'{running}. {title}')
                    running += 1
                lines.append('')
            if lines and lines[-1] == '':
                lines.pop()

        lines.append('')
        lines.append('查詢請回：#今日代辦')
        lines.append('完成可回：#完成代辦 1')
        reply = '\n'.join(lines)
        if failed_msgs:
            reply += '\n\n有部分未新增：\n' + '\n'.join(failed_msgs[:3])
        return {'handled': True, 'ok': True, 'reply_text': reply[:5000], 'parsed_tag': '新增代辦'}

    return {'handled': True, 'ok': False, 'reply_text': '\n'.join(failed_msgs)[:5000], 'parsed_tag': '新增代辦'}


def _format_line_todo_sections(overdue_items, today_items, title, today_label='今天'):
    """覆寫清單顯示：完全不顯示 ID，只靠畫面序號完成。"""
    if not overdue_items and not today_items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    idx = 1

    if overdue_items:
        lines.append('')
        lines.append('【尚未完成】')
        for doc in overdue_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            old_date = _todo_display_md((data.get('todo_date') or '').strip())
            lines.append(f"{idx}. {old_date}｜{data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if today_items:
        lines.append('')
        lines.append(f'【{today_label}要做】')
        for doc in today_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    lines.append('一次完成多筆：#完成代辦 1 3')
    lines.append('清除請回：#清除代辦 1')
    return '\n'.join(lines)[:5000]


def _format_line_todo_list(items, title):
    """保險覆寫舊版清單格式：不顯示 ID。"""
    if not items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    for idx, item in enumerate(items, 1):
        data = item.to_dict() or {}
        note = (data.get('note') or '').strip()
        line = f"{idx}. {data.get('title', '')}"
        if note:
            line += f"\n   備註: {note}"
        lines.append(line)
    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    return '\n'.join(lines)[:5000]


def _find_line_todo(todo_key: str, target_id=''):
    """覆寫搜尋錯誤訊息：避免提示 ID，優先建議用序號完成。"""
    key = (todo_key or '').strip()
    if not key:
        return None, '請提供代辦序號或事項關鍵字。'

    # 仍保留隱藏能力：若你剛好知道完整文件 ID，程式仍可處理，但 LINE 畫面不顯示。
    direct = db.collection(LINE_TODO_COLLECTION).document(key).get()
    if direct.exists:
        data = direct.to_dict() or {}
        if target_id and data.get('line_target_id') != target_id:
            return None, '這筆代辦不是在目前這個 LINE 對話建立的。'
        return direct, ''

    matches = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if target_id and data.get('line_target_id') != target_id:
            continue
        title = data.get('title', '')
        if doc.id.startswith(key) or key in title:
            matches.append(doc)

    if len(matches) == 1:
        return matches[0], ''
    if len(matches) > 1:
        preview = '\n'.join([f"- {(d.to_dict() or {}).get('title','')}" for d in matches[:8]])
        return None, '找到多筆相似代辦，請先輸入 #今日代辦，再用序號完成：\n' + preview
    return None, '找不到這筆未完成代辦，請先輸入 #今日代辦 查看序號。'


def send_tomorrow_line_todo_reminders():
    """
    晚上 23:00 使用：提醒明天要做的代辦。
    注意：這不會影響隔天早上 08:00 的今日提醒，兩者用不同欄位記錄。
    """
    tomorrow_date = (now_taipei().date() + timedelta(days=1)).strftime('%Y-%m-%d')
    grouped = {}

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        todo_date = (data.get('todo_date') or '').strip()
        if todo_date != tomorrow_date:
            continue
        sent_dates = data.get('tomorrow_reminder_sent_dates') or []
        if tomorrow_date in sent_dates:
            continue
        target_id = data.get('line_target_id', '')
        if not target_id:
            continue
        grouped.setdefault(target_id, []).append(doc)

    sent_count = 0
    failed = []

    for target_id, items in grouped.items():
        tomorrow_items = _sort_line_todo_docs(items)
        text = _format_line_todo_sections(
            [],
            tomorrow_items,
            f'明天 {_todo_display_md(tomorrow_date)} 要做的事情',
            today_label='明天',
        )
        ok, msg = push_line_text(target_id, text)
        if ok:
            sent_count += 1
            for doc in items:
                doc.reference.update({
                    'tomorrow_reminder_sent_dates': firestore.ArrayUnion([tomorrow_date]),
                    'last_tomorrow_reminded_at': now_taipei().isoformat(),
                })
        else:
            failed.append({'target_id': target_id, 'error': msg})

    return {'date': tomorrow_date, 'target_count': len(grouped), 'sent_count': sent_count, 'failed': failed}


@app.route('/line/todos/remind-tomorrow', methods=['GET', 'POST'])
def line_todos_remind_tomorrow():
    # Render Cron Job / UptimeRobot 建議呼叫：/line/todos/remind-tomorrow?key=你的密鑰
    secret = os.environ.get('TODO_REMINDER_SECRET', '').strip()
    key = request.args.get('key', '').strip() or request.form.get('key', '').strip()
    if secret and key != secret:
        return {'ok': False, 'message': 'Invalid key'}, 403
    result = send_tomorrow_line_todo_reminders()
    return {'ok': True, 'result': result}, 200

# ========= LINE Bot 代辦事項 Patch v5 End =========



# ========= LINE Bot 代辦事項 Patch v6：用 LINE 指令設定提醒時間 =========
# 貼在 v5 代辦事項 Patch 的最底部即可。
# 重點：
# 1. LINE 指令只負責「儲存提醒設定」。
# 2. 真正定時發送由 /line/todos/reminder-check 搭配 Render Cron / UptimeRobot 每 5~10 分鐘呼叫。
# 3. 每個 LINE 對話（群組 / 聊天室 / 個人）可以有自己的提醒時間。

LINE_TODO_SETTINGS_COLLECTION = os.environ.get('LINE_TODO_SETTINGS_COLLECTION', 'line_todo_settings')


LINE_TODO_DEFAULT_REMINDER_SETTINGS = {
    'today_enabled': True,
    'today_reminder_time': '08:00',
    'tomorrow_enabled': True,
    'tomorrow_reminder_time': '23:00',
}


def _line_todo_setting_doc_id(target_id: str) -> str:
    """避免 LINE target_id 內有特殊字元，統一用 sha1 當 Firestore 文件 ID。"""
    raw = (target_id or '').strip()
    if not raw:
        raw = 'unknown'
    return hashlib.sha1(raw.encode('utf-8')).hexdigest()


def _normalize_line_todo_time(value: str):
    """
    支援：
    - 08:00 / 8:00 / 0800
    - 8點 / 8點30 / 8時30分
    - 早上8點 / 上午8點
    - 晚上11點 / 下午11點 -> 23:00
    回傳 HH:MM；看不懂回傳空字串。
    """
    raw = (value or '').strip()
    if not raw:
        return ''

    raw = raw.translate(str.maketrans('０１２３４５６７８９：', '0123456789:'))
    raw = re.sub(r'\s+', '', raw)

    is_pm = any(x in raw for x in ['下午', '晚上', '晚間', '傍晚'])
    is_am = any(x in raw for x in ['上午', '早上', '清晨'])
    raw2 = raw
    for word in ['上午', '早上', '清晨', '下午', '晚上', '晚間', '傍晚', '中午']:
        raw2 = raw2.replace(word, '')

    hour = None
    minute = 0

    m = re.search(r'^(\d{1,2}):(\d{1,2})$', raw2)
    if m:
        hour = int(m.group(1))
        minute = int(m.group(2))
    else:
        m = re.search(r'^(\d{3,4})$', raw2)
        if m:
            digits = m.group(1)
            hour = int(digits[:-2])
            minute = int(digits[-2:])
        else:
            m = re.search(r'(\d{1,2})(?:點|時)(?:(\d{1,2})(?:分)?)?$', raw2)
            if m:
                hour = int(m.group(1))
                minute = int(m.group(2) or 0)
            else:
                m = re.search(r'^(\d{1,2})$', raw2)
                if m:
                    hour = int(m.group(1))
                    minute = 0

    if hour is None:
        return ''

    if is_pm and hour < 12:
        hour += 12
    if is_am and hour == 12:
        hour = 0

    if hour < 0 or hour > 23 or minute < 0 or minute > 59:
        return ''

    return f'{hour:02d}:{minute:02d}'


def _time_to_minutes(hhmm: str, default='00:00'):
    t = _normalize_line_todo_time(hhmm) or default
    h, m = t.split(':')
    return int(h) * 60 + int(m)


def _get_line_todo_reminder_settings(target_id: str, target_type: str = ''):
    settings = dict(LINE_TODO_DEFAULT_REMINDER_SETTINGS)
    if not target_id:
        return settings

    doc = db.collection(LINE_TODO_SETTINGS_COLLECTION).document(_line_todo_setting_doc_id(target_id)).get()
    if doc.exists:
        data = doc.to_dict() or {}
        settings.update({k: v for k, v in data.items() if v is not None})

    settings['line_target_id'] = target_id
    if target_type:
        settings['line_target_type'] = target_type
    return settings


def _save_line_todo_reminder_settings(target_id: str, target_type: str, updates: dict, event=None):
    if not target_id:
        return False, '找不到目前 LINE 對話 ID，無法設定提醒。'

    now = now_taipei().isoformat()
    payload = {
        'line_target_id': target_id,
        'line_target_type': target_type,
        'updated_at': now,
        'updated_by_id': 'line_bot',
        'updated_by_name': get_line_sender_display_name(event) if event else 'LINE Bot',
    }
    payload.update(updates)
    db.collection(LINE_TODO_SETTINGS_COLLECTION).document(_line_todo_setting_doc_id(target_id)).set(payload, merge=True)
    return True, ''


def _parse_line_todo_reminder_setting_command(raw_text: str):
    """解析 #設定代辦提醒 / #設定提醒。"""
    text = (raw_text or '').strip()
    if text.startswith('#設定代辦提醒'):
        body = text.replace('#設定代辦提醒', '', 1).strip()
    elif text.startswith('#設定提醒'):
        body = text.replace('#設定提醒', '', 1).strip()
    else:
        body = text

    lines = [ln.strip() for ln in body.splitlines() if ln.strip()]
    joined = ' '.join(lines)

    updates = {}
    errors = []

    key_aliases = {
        '今日': 'today_reminder_time',
        '今天': 'today_reminder_time',
        '今日提醒': 'today_reminder_time',
        '今天提醒': 'today_reminder_time',
        '今日代辦': 'today_reminder_time',
        '明日': 'tomorrow_reminder_time',
        '明天': 'tomorrow_reminder_time',
        '明日提醒': 'tomorrow_reminder_time',
        '明天提醒': 'tomorrow_reminder_time',
        '明天代辦': 'tomorrow_reminder_time',
    }

    # 多行 key: value 格式
    for line in lines:
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.+)$', line)
        if not m:
            continue
        key = re.sub(r'\s+', '', m.group(1).strip())
        value = m.group(2).strip()
        field = key_aliases.get(key)
        if field:
            t = _normalize_line_todo_time(value)
            if t:
                updates[field] = t
            else:
                errors.append(f'{key} 的時間看不懂：{value}')

    # 一行自然語句格式，例如：#設定提醒 今日 08:00 明天 23:00
    patterns = [
        (r'(今日|今天|今日代辦|今天代辦|今日提醒|今天提醒)\D{0,10}((?:上午|早上|下午|晚上|晚間|傍晚|清晨)?\s*\d{1,2}(?::\d{1,2}|點\d{0,2}|時\d{0,2}分?)?)', 'today_reminder_time'),
        (r'(明日|明天|明日代辦|明天代辦|明日提醒|明天提醒)\D{0,10}((?:上午|早上|下午|晚上|晚間|傍晚|清晨)?\s*\d{1,2}(?::\d{1,2}|點\d{0,2}|時\d{0,2}分?)?)', 'tomorrow_reminder_time'),
    ]
    for pattern, field in patterns:
        m = re.search(pattern, joined)
        if m and field not in updates:
            t = _normalize_line_todo_time(m.group(2))
            if t:
                updates[field] = t

    return updates, errors


def set_line_todo_reminder_settings_from_command(raw_text: str, event):
    target_id, target_type = _line_todo_target_from_event(event)
    updates, errors = _parse_line_todo_reminder_setting_command(raw_text)

    if not updates:
        msg = (
            '請告訴我要設定的提醒時間。\n\n'
            '範例：\n'
            '#設定代辦提醒\n'
            '今日提醒: 08:00\n'
            '明天提醒: 23:00\n\n'
            '也可以：#設定提醒 今日 08:00 明天 23:00'
        )
        if errors:
            msg += '\n\n' + '\n'.join(errors[:3])
        return {'handled': True, 'ok': False, 'reply_text': msg, 'parsed_tag': '設定代辦提醒'}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '設定代辦提醒'}

    current = _get_line_todo_reminder_settings(target_id, target_type)
    lines = ['已更新代辦提醒設定']
    lines.append(f"今日代辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}")
    lines.append(f"明天代辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}")
    lines.append('')
    lines.append('查詢設定：#代辦提醒設定')
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines), 'parsed_tag': '設定代辦提醒'}


def query_line_todo_reminder_settings(event):
    target_id, target_type = _line_todo_target_from_event(event)
    current = _get_line_todo_reminder_settings(target_id, target_type)
    lines = [
        '目前代辦提醒設定',
        f"今日代辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}",
        f"明天代辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}",
        '',
        '修改範例：',
        '#設定代辦提醒',
        '今日提醒: 08:00',
        '明天提醒: 23:00',
        '',
        '關閉範例：#關閉代辦提醒 明天',
    ]
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines), 'parsed_tag': '代辦提醒設定'}


def switch_line_todo_reminder(raw_text: str, event, enabled: bool):
    target_id, target_type = _line_todo_target_from_event(event)
    body = raw_text.replace('#關閉代辦提醒', '', 1).replace('#開啟代辦提醒', '', 1).strip()
    body = body or '全部'

    updates = {}
    if any(x in body for x in ['今日', '今天']):
        updates['today_enabled'] = enabled
    if any(x in body for x in ['明日', '明天']):
        updates['tomorrow_enabled'] = enabled
    if not updates or '全部' in body or '全開' in body or '全關' in body:
        updates = {'today_enabled': enabled, 'tomorrow_enabled': enabled}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '代辦提醒設定'}

    current = _get_line_todo_reminder_settings(target_id, target_type)
    action = '開啟' if enabled else '關閉'
    lines = [
        f'已{action}代辦提醒',
        f"今日代辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}",
        f"明天代辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}",
    ]
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines), 'parsed_tag': '代辦提醒設定'}


def process_line_todo_reminder_settings_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if not raw_text.startswith('#'):
        return {'handled': False}

    if raw_text.startswith('#設定代辦提醒') or raw_text.startswith('#設定提醒'):
        result = set_line_todo_reminder_settings_from_command(raw_text, event)
        save_line_log({'tag': result.get('parsed_tag', '設定代辦提醒'), 'action': 'line_todo_reminder_setting_update', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#代辦提醒設定') or raw_text.startswith('#查詢代辦提醒') or raw_text.startswith('#提醒設定'):
        result = query_line_todo_reminder_settings(event)
        save_line_log({'tag': result.get('parsed_tag', '代辦提醒設定'), 'action': 'line_todo_reminder_setting_query', 'fields': {}, 'raw_text': raw_text}, event, 'success', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#關閉代辦提醒'):
        result = switch_line_todo_reminder(raw_text, event, enabled=False)
        save_line_log({'tag': result.get('parsed_tag', '代辦提醒設定'), 'action': 'line_todo_reminder_setting_disable', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#開啟代辦提醒'):
        result = switch_line_todo_reminder(raw_text, event, enabled=True)
        save_line_log({'tag': result.get('parsed_tag', '代辦提醒設定'), 'action': 'line_todo_reminder_setting_enable', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    return {'handled': False}


# 先攔截「提醒設定」指令，其它代辦指令與原本客需 / 委託 / 開發指令照舊。
_process_line_todo_message_event_before_reminder_settings = process_line_todo_message_event


def process_line_todo_message_event(event):
    setting_result = process_line_todo_reminder_settings_message_event(event)
    if setting_result.get('handled'):
        return setting_result
    return _process_line_todo_message_event_before_reminder_settings(event)


def _collect_line_todo_targets_with_open_items():
    targets = {}
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        target_id = data.get('line_target_id', '')
        if not target_id:
            continue
        targets[target_id] = data.get('line_target_type', '')
    return targets


def send_due_line_todo_reminders_by_settings():
    """
    建議由 Cron 每 5~10 分鐘呼叫一次。
    會依各 LINE 對話儲存的提醒設定判斷是否該發：
    - 今日提醒：todo_date <= 今天，包含尚未完成
    - 明天提醒：todo_date == 明天
    每一種提醒每天只發一次。
    """
    now_dt = now_taipei()
    today = now_dt.strftime('%Y-%m-%d')
    tomorrow = (now_dt.date() + timedelta(days=1)).strftime('%Y-%m-%d')
    current_minutes = now_dt.hour * 60 + now_dt.minute

    targets = _collect_line_todo_targets_with_open_items()
    sent_count = 0
    failed = []
    checked_targets = 0

    for target_id, target_type in targets.items():
        checked_targets += 1
        settings = _get_line_todo_reminder_settings(target_id, target_type)

        # 今日提醒：尚未完成 + 今日要做
        if settings.get('today_enabled', True):
            today_time = settings.get('today_reminder_time', '08:00')
            if current_minutes >= _time_to_minutes(today_time, default='08:00'):
                items = []
                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if not todo_date or todo_date > today:
                        continue
                    sent_dates = data.get('reminder_sent_dates') or []
                    if today in sent_dates:
                        continue
                    items.append(doc)

                if items:
                    overdue_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) < today])
                    today_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) == today])
                    text = _format_line_todo_sections(
                        overdue_items,
                        today_items,
                        f'今日代辦提醒 {_todo_display_md(today)}',
                        today_label='今天',
                    )
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in items:
                            doc.reference.update({
                                'reminder_sent_dates': firestore.ArrayUnion([today]),
                                'last_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'today', 'error': msg})

        # 明天提醒：明天要做
        if settings.get('tomorrow_enabled', True):
            tomorrow_time = settings.get('tomorrow_reminder_time', '23:00')
            if current_minutes >= _time_to_minutes(tomorrow_time, default='23:00'):
                items = []
                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if todo_date != tomorrow:
                        continue
                    sent_dates = data.get('tomorrow_reminder_sent_dates') or []
                    if tomorrow in sent_dates:
                        continue
                    items.append(doc)

                if items:
                    tomorrow_items = _sort_line_todo_docs(items)
                    text = _format_line_todo_sections(
                        [],
                        tomorrow_items,
                        f'明天 {_todo_display_md(tomorrow)} 要做的事情',
                        today_label='明天',
                    )
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in items:
                            doc.reference.update({
                                'tomorrow_reminder_sent_dates': firestore.ArrayUnion([tomorrow]),
                                'last_tomorrow_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'tomorrow', 'error': msg})

    return {
        'now': now_dt.strftime('%Y-%m-%d %H:%M:%S'),
        'checked_targets': checked_targets,
        'sent_count': sent_count,
        'failed': failed,
    }


@app.route('/line/todos/reminder-check', methods=['GET', 'POST'])
def line_todos_reminder_check():
    # Render Cron / UptimeRobot 建議每 5~10 分鐘呼叫：/line/todos/reminder-check?key=你的密鑰
    secret = os.environ.get('TODO_REMINDER_SECRET', '').strip()
    key = request.args.get('key', '').strip() or request.form.get('key', '').strip()
    if secret and key != secret:
        return {'ok': False, 'message': 'Invalid key'}, 403
    result = send_due_line_todo_reminders_by_settings()
    return {'ok': True, 'result': result}, 200

# ========= LINE Bot 代辦事項 Patch v6 End =========

# ========= LINE Bot 代辦事項 Patch v7：固定每日提醒時間 + 自訂開頭 =========
# 貼在 v6 代辦事項 Patch 的最底部即可。
# 功能：
# 1. 固定每天「今日代辦提醒」時間，例如 08:00。
# 2. 固定每天「明天代辦提醒」時間，例如 23:00。
# 3. 可用 LINE 指令設定早上開頭 / 晚上開頭。
# 4. 提醒內容格式：開頭文字 -> 空一行 -> 代辦事項。

from datetime import timedelta

# 覆寫 v6 預設設定：加入開頭文字欄位。
LINE_TODO_DEFAULT_REMINDER_SETTINGS = {
    'today_enabled': True,
    'today_reminder_time': '08:00',
    'tomorrow_enabled': True,
    'tomorrow_reminder_time': '23:00',
    'today_opening_text': '各位厝米的夥伴早安 ☀️\n今天的代辦事項如下：',
    'tomorrow_opening_text': '各位厝米的夥伴晚安 🌙\n先看一下明天要完成的事情：',
}


def _line_todo_clean_opening_text(value: str) -> str:
    """整理提醒開頭文字。"""
    text = (value or '').strip()
    if text in ('無', '不用', '不要', '不要開頭', '關閉', '清空', 'none', 'None'):
        return ''
    # 避免太長塞爆 LINE，最多保留 600 字。
    return text[:600]


def _line_todo_preview_opening(value: str) -> str:
    text = (value or '').strip()
    if not text:
        return '未設定'
    one_line = ' / '.join([ln.strip() for ln in text.splitlines() if ln.strip()])
    return one_line[:120]


def _line_todo_add_opening(opening_text: str, todo_text: str) -> str:
    opening = (opening_text or '').strip()
    body = (todo_text or '').strip()
    if not opening:
        return body[:5000]
    return (opening + '\n\n' + body)[:5000]


def _parse_line_todo_reminder_setting_command(raw_text: str):
    """覆寫 v6：解析提醒時間 + 提醒開頭。"""
    text = (raw_text or '').strip()
    if text.startswith('#設定代辦提醒'):
        body = text.replace('#設定代辦提醒', '', 1).strip()
    elif text.startswith('#設定提醒'):
        body = text.replace('#設定提醒', '', 1).strip()
    else:
        body = text

    lines = [ln.rstrip() for ln in body.splitlines() if ln.strip()]
    joined = ' '.join([ln.strip() for ln in lines])

    updates = {}
    errors = []

    time_key_aliases = {
        '今日': 'today_reminder_time',
        '今天': 'today_reminder_time',
        '早上': 'today_reminder_time',
        '早安': 'today_reminder_time',
        '今日提醒': 'today_reminder_time',
        '今天提醒': 'today_reminder_time',
        '早上提醒': 'today_reminder_time',
        '早安提醒': 'today_reminder_time',
        '今日代辦': 'today_reminder_time',
        '今天代辦': 'today_reminder_time',
        '明日': 'tomorrow_reminder_time',
        '明天': 'tomorrow_reminder_time',
        '晚上': 'tomorrow_reminder_time',
        '晚安': 'tomorrow_reminder_time',
        '明日提醒': 'tomorrow_reminder_time',
        '明天提醒': 'tomorrow_reminder_time',
        '晚上提醒': 'tomorrow_reminder_time',
        '晚安提醒': 'tomorrow_reminder_time',
        '明天代辦': 'tomorrow_reminder_time',
    }

    opening_key_aliases = {
        '今日開頭': 'today_opening_text',
        '今天開頭': 'today_opening_text',
        '早上開頭': 'today_opening_text',
        '早安開頭': 'today_opening_text',
        '今日提醒開頭': 'today_opening_text',
        '今天提醒開頭': 'today_opening_text',
        '早上提醒開頭': 'today_opening_text',
        '今日代辦開頭': 'today_opening_text',
        '明天開頭': 'tomorrow_opening_text',
        '明日開頭': 'tomorrow_opening_text',
        '晚上開頭': 'tomorrow_opening_text',
        '晚安開頭': 'tomorrow_opening_text',
        '明天提醒開頭': 'tomorrow_opening_text',
        '明日提醒開頭': 'tomorrow_opening_text',
        '晚上提醒開頭': 'tomorrow_opening_text',
        '明天代辦開頭': 'tomorrow_opening_text',
    }

    # 多行 key: value 格式。
    for line in lines:
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.*)$', line)
        if not m:
            continue
        key = re.sub(r'\s+', '', m.group(1).strip())
        value = (m.group(2) or '').strip()

        time_field = time_key_aliases.get(key)
        if time_field:
            t = _normalize_line_todo_time(value)
            if t:
                updates[time_field] = t
            else:
                errors.append(f'{key} 的時間看不懂：{value}')
            continue

        opening_field = opening_key_aliases.get(key)
        if opening_field:
            updates[opening_field] = _line_todo_clean_opening_text(value)
            continue

    # 一行自然語句格式，例如：#設定提醒 今日 08:00 明天 23:00
    patterns = [
        (r'(今日|今天|早上|早安|今日代辦|今天代辦|今日提醒|今天提醒|早上提醒|早安提醒)\D{0,10}((?:上午|早上|下午|晚上|晚間|傍晚|清晨)?\s*\d{1,2}(?::\d{1,2}|點\d{0,2}|時\d{0,2}分?)?)', 'today_reminder_time'),
        (r'(明日|明天|晚上|晚安|明日代辦|明天代辦|明日提醒|明天提醒|晚上提醒|晚安提醒)\D{0,10}((?:上午|早上|下午|晚上|晚間|傍晚|清晨)?\s*\d{1,2}(?::\d{1,2}|點\d{0,2}|時\d{0,2}分?)?)', 'tomorrow_reminder_time'),
    ]
    for pattern, field in patterns:
        m = re.search(pattern, joined)
        if m and field not in updates:
            t = _normalize_line_todo_time(m.group(2))
            if t:
                updates[field] = t

    return updates, errors


def _set_line_todo_opening_from_command(raw_text: str, event, target_field: str):
    """處理 #設定今日開頭 / #設定明天開頭。"""
    target_id, target_type = _line_todo_target_from_event(event)

    command_aliases = [
        '#設定今日開頭', '#設定今天開頭', '#設定早安開頭', '#設定早上開頭',
        '#設定明天開頭', '#設定明日開頭', '#設定晚安開頭', '#設定晚上開頭',
    ]
    body = (raw_text or '').strip()
    for cmd in command_aliases:
        if body.startswith(cmd):
            body = body.replace(cmd, '', 1).strip()
            break

    opening = _line_todo_clean_opening_text(body)
    if not opening and body not in ('無', '不用', '不要', '不要開頭', '關閉', '清空'):
        label = '今日提醒開頭' if target_field == 'today_opening_text' else '明天提醒開頭'
        example = '各位厝米的夥伴早安\n今天的代辦事項如下：' if target_field == 'today_opening_text' else '各位厝米的夥伴晚安\n先看一下明天要完成的事情：'
        return {
            'handled': True,
            'ok': False,
            'reply_text': f'請輸入要設定的{label}。\n\n範例：\n#設定{"今日" if target_field == "today_opening_text" else "明天"}開頭\n{example}',
            'parsed_tag': '設定代辦提醒開頭',
        }

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, {target_field: opening}, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '設定代辦提醒開頭'}

    label = '今日提醒開頭' if target_field == 'today_opening_text' else '明天提醒開頭'
    msg = f'已更新{label}：\n{opening or "未設定"}'
    return {'handled': True, 'ok': True, 'reply_text': msg[:5000], 'parsed_tag': '設定代辦提醒開頭'}


def set_line_todo_reminder_settings_from_command(raw_text: str, event):
    """覆寫 v6：設定時間時也可以同時設定開頭。"""
    target_id, target_type = _line_todo_target_from_event(event)
    updates, errors = _parse_line_todo_reminder_setting_command(raw_text)

    if not updates:
        msg = (
            '請告訴我要設定的固定每日提醒時間或開頭。\n\n'
            '範例：\n'
            '#設定代辦提醒\n'
            '今日提醒: 08:00\n'
            '今日開頭: 各位厝米的夥伴早安 ☀️\n'
            '明天提醒: 23:00\n'
            '明天開頭: 各位厝米的夥伴晚安 🌙\n\n'
            '也可以單獨設定：\n'
            '#設定今日開頭\n'
            '各位厝米的夥伴早安\n'
            '今天的代辦事項如下：'
        )
        if errors:
            msg += '\n\n' + '\n'.join(errors[:3])
        return {'handled': True, 'ok': False, 'reply_text': msg[:5000], 'parsed_tag': '設定代辦提醒'}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '設定代辦提醒'}

    current = _get_line_todo_reminder_settings(target_id, target_type)
    lines = ['已更新固定每日代辦提醒設定']
    lines.append(f"今日代辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}")
    lines.append(f"今日開頭：{_line_todo_preview_opening(current.get('today_opening_text', ''))}")
    lines.append(f"明天代辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}")
    lines.append(f"明天開頭：{_line_todo_preview_opening(current.get('tomorrow_opening_text', ''))}")
    lines.append('')
    lines.append('查詢設定：#代辦提醒設定')
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '設定代辦提醒'}


def query_line_todo_reminder_settings(event):
    """覆寫 v6：查詢設定時顯示開頭。"""
    target_id, target_type = _line_todo_target_from_event(event)
    current = _get_line_todo_reminder_settings(target_id, target_type)
    lines = [
        '目前固定每日代辦提醒設定',
        f"今日代辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}",
        f"今日開頭：{_line_todo_preview_opening(current.get('today_opening_text', ''))}",
        f"明天代辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}",
        f"明天開頭：{_line_todo_preview_opening(current.get('tomorrow_opening_text', ''))}",
        '',
        '修改範例：',
        '#設定代辦提醒',
        '今日提醒: 08:00',
        '今日開頭: 各位厝米的夥伴早安 ☀️',
        '明天提醒: 23:00',
        '明天開頭: 各位厝米的夥伴晚安 🌙',
        '',
        '單獨修改開頭：#設定今日開頭',
        '關閉範例：#關閉代辦提醒 明天',
    ]
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '代辦提醒設定'}


def switch_line_todo_reminder(raw_text: str, event, enabled: bool):
    """覆寫 v6：開關後也顯示開頭摘要。"""
    target_id, target_type = _line_todo_target_from_event(event)
    body = raw_text.replace('#關閉代辦提醒', '', 1).replace('#開啟代辦提醒', '', 1).strip()
    body = body or '全部'

    updates = {}
    if any(x in body for x in ['今日', '今天', '早上', '早安']):
        updates['today_enabled'] = enabled
    if any(x in body for x in ['明日', '明天', '晚上', '晚安']):
        updates['tomorrow_enabled'] = enabled
    if not updates or '全部' in body or '全開' in body or '全關' in body:
        updates = {'today_enabled': enabled, 'tomorrow_enabled': enabled}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '代辦提醒設定'}

    current = _get_line_todo_reminder_settings(target_id, target_type)
    action = '開啟' if enabled else '關閉'
    lines = [
        f'已{action}代辦提醒',
        f"今日代辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}",
        f"今日開頭：{_line_todo_preview_opening(current.get('today_opening_text', ''))}",
        f"明天代辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}",
        f"明天開頭：{_line_todo_preview_opening(current.get('tomorrow_opening_text', ''))}",
    ]
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '代辦提醒設定'}


def process_line_todo_reminder_settings_message_event(event):
    """覆寫 v6：新增 #設定今日開頭 / #設定明天開頭。"""
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if not raw_text.startswith('#'):
        return {'handled': False}

    today_opening_commands = ('#設定今日開頭', '#設定今天開頭', '#設定早安開頭', '#設定早上開頭')
    tomorrow_opening_commands = ('#設定明天開頭', '#設定明日開頭', '#設定晚安開頭', '#設定晚上開頭')

    if raw_text.startswith(today_opening_commands):
        result = _set_line_todo_opening_from_command(raw_text, event, 'today_opening_text')
        save_line_log({'tag': result.get('parsed_tag', '設定代辦提醒開頭'), 'action': 'line_todo_today_opening_update', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith(tomorrow_opening_commands):
        result = _set_line_todo_opening_from_command(raw_text, event, 'tomorrow_opening_text')
        save_line_log({'tag': result.get('parsed_tag', '設定代辦提醒開頭'), 'action': 'line_todo_tomorrow_opening_update', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#設定代辦提醒') or raw_text.startswith('#設定提醒'):
        result = set_line_todo_reminder_settings_from_command(raw_text, event)
        save_line_log({'tag': result.get('parsed_tag', '設定代辦提醒'), 'action': 'line_todo_reminder_setting_update', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#代辦提醒設定') or raw_text.startswith('#查詢代辦提醒') or raw_text.startswith('#提醒設定'):
        result = query_line_todo_reminder_settings(event)
        save_line_log({'tag': result.get('parsed_tag', '代辦提醒設定'), 'action': 'line_todo_reminder_setting_query', 'fields': {}, 'raw_text': raw_text}, event, 'success', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#關閉代辦提醒'):
        result = switch_line_todo_reminder(raw_text, event, enabled=False)
        save_line_log({'tag': result.get('parsed_tag', '代辦提醒設定'), 'action': 'line_todo_reminder_setting_disable', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    if raw_text.startswith('#開啟代辦提醒'):
        result = switch_line_todo_reminder(raw_text, event, enabled=True)
        save_line_log({'tag': result.get('parsed_tag', '代辦提醒設定'), 'action': 'line_todo_reminder_setting_enable', 'fields': {}, 'raw_text': raw_text}, event, 'success' if result.get('ok') else 'failed', note=result.get('reply_text', ''), sender_display_name=get_line_sender_display_name(event))
        return result

    return {'handled': False}


def send_due_line_todo_reminders_by_settings():
    """
    覆寫 v6：依固定每日時間提醒，並在代辦清單前加上自訂開頭。
    建議 Cron 每 5~10 分鐘呼叫 /line/todos/reminder-check?key=你的密鑰。
    """
    now_dt = now_taipei()
    today = now_dt.strftime('%Y-%m-%d')
    tomorrow = (now_dt.date() + timedelta(days=1)).strftime('%Y-%m-%d')
    current_minutes = now_dt.hour * 60 + now_dt.minute

    targets = _collect_line_todo_targets_with_open_items()
    sent_count = 0
    failed = []
    checked_targets = 0

    for target_id, target_type in targets.items():
        checked_targets += 1
        settings = _get_line_todo_reminder_settings(target_id, target_type)

        # 今日提醒：尚未完成 + 今日要做。
        if settings.get('today_enabled', True):
            today_time = settings.get('today_reminder_time', '08:00')
            if current_minutes >= _time_to_minutes(today_time, default='08:00'):
                items = []
                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if not todo_date or todo_date > today:
                        continue
                    sent_dates = data.get('reminder_sent_dates') or []
                    if today in sent_dates:
                        continue
                    items.append(doc)

                if items:
                    overdue_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) < today])
                    today_items = _sort_line_todo_docs([d for d in items if _todo_date_value(d) == today])
                    body = _format_line_todo_sections(
                        overdue_items,
                        today_items,
                        f'{_todo_display_md(today)} 今日代辦',
                        today_label='今天',
                    )
                    text = _line_todo_add_opening(settings.get('today_opening_text', ''), body)
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in items:
                            doc.reference.update({
                                'reminder_sent_dates': firestore.ArrayUnion([today]),
                                'last_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'today', 'error': msg})

        # 明天提醒：明天要做。
        if settings.get('tomorrow_enabled', True):
            tomorrow_time = settings.get('tomorrow_reminder_time', '23:00')
            if current_minutes >= _time_to_minutes(tomorrow_time, default='23:00'):
                items = []
                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if todo_date != tomorrow:
                        continue
                    sent_dates = data.get('tomorrow_reminder_sent_dates') or []
                    if tomorrow in sent_dates:
                        continue
                    items.append(doc)

                if items:
                    tomorrow_items = _sort_line_todo_docs(items)
                    body = _format_line_todo_sections(
                        [],
                        tomorrow_items,
                        f'{_todo_display_md(tomorrow)} 明天代辦',
                        today_label='明天',
                    )
                    text = _line_todo_add_opening(settings.get('tomorrow_opening_text', ''), body)
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in items:
                            doc.reference.update({
                                'tomorrow_reminder_sent_dates': firestore.ArrayUnion([tomorrow]),
                                'last_tomorrow_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'tomorrow', 'error': msg})

    return {
        'now': now_dt.strftime('%Y-%m-%d %H:%M:%S'),
        'checked_targets': checked_targets,
        'sent_count': sent_count,
        'failed': failed,
    }

# ========= LINE Bot 代辦事項 Patch v7 End =========


# ========= LINE Bot 代辦事項 Patch v8：防止重複新增 + 清除既有重複代辦 =========
# 貼在 v7 patch 的最底部即可。
# 目的：
# 1. 同一個 LINE 訊息重送時，不會重複新增代辦。
# 2. 同一個 LINE 對話、同一天、同一個事項，已經有未完成代辦時，不再新增第二筆。
# 3. 清單顯示時會自動去重，避免同一事項出現兩次。
# 4. 可用 #清除重複代辦 一次整理目前 LINE 對話中的重複資料。

LINE_TODO_PROCESSED_COLLECTION = os.environ.get('LINE_TODO_PROCESSED_COLLECTION', 'line_todo_processed_messages')


def _line_todo_normalize_title_for_dedupe(value: str) -> str:
    """把代辦標題正規化，避免空白不同造成重複判斷失敗。"""
    text = (value or '').strip()
    text = re.sub(r'\s+', ' ', text)
    return text.lower()


def _line_todo_dedupe_key(target_id: str, todo_date: str, title: str) -> str:
    return f"{target_id}|{todo_date}|{_line_todo_normalize_title_for_dedupe(title)}"


def _line_todo_doc_dedupe_key(doc):
    data = doc.to_dict() or {}
    return _line_todo_dedupe_key(
        data.get('line_target_id', ''),
        (data.get('todo_date') or '').strip(),
        data.get('title', ''),
    )


def _find_existing_open_line_todo(target_id: str, todo_date: str, title: str):
    """找同一對話、同一天、同事項的未完成代辦。"""
    key = _line_todo_dedupe_key(target_id, todo_date, title)
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if data.get('line_target_id', '') != target_id:
            continue
        if (data.get('todo_date') or '').strip() != todo_date:
            continue
        existing_key = data.get('dedupe_key') or _line_todo_dedupe_key(
            data.get('line_target_id', ''),
            (data.get('todo_date') or '').strip(),
            data.get('title', ''),
        )
        if existing_key == key:
            return doc
    return None


def _line_todo_dedupe_docs_for_display(items):
    """顯示清單前去重；保留最早建立的那筆。"""
    sorted_items = sorted(
        list(items or []),
        key=lambda d: ((d.to_dict() or {}).get('created_at') or '', d.id)
    )
    seen = set()
    result = []
    for doc in sorted_items:
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        key = data.get('dedupe_key') or _line_todo_doc_dedupe_key(doc)
        if key in seen:
            continue
        seen.add(key)
        result.append(doc)
    return result


# 覆寫新增單筆：加入「同訊息防重送」與「同日期同事項防重複」。
def create_line_todo(fields, event):
    title = (fields.get('title') or '').strip()
    todo_date = _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    note = (fields.get('note') or '').strip()
    remind_time = (fields.get('remind_time') or '').strip()

    if not title:
        return {
            'handled': True,
            'ok': False,
            'reply_text': '未新增：請填「事項」。\n\n範例：\n#新增代辦\n日期: 明天\n厝米排版 土地現廣稿\n拍水哥爸爸土地',
            'parsed_tag': '新增代辦',
        }
    if not todo_date:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：日期格式看不懂，請用 5/29、今天、明天。', 'parsed_tag': '新增代辦'}

    target_id, target_type = _line_todo_target_from_event(event)
    duplicate_doc = _find_existing_open_line_todo(target_id, todo_date, title)
    if duplicate_doc:
        return {
            'handled': True,
            'ok': False,
            'duplicate': True,
            'reply_text': f'已存在，不重複新增：{_todo_display_md(todo_date)}｜{title}',
            'parsed_tag': '新增代辦',
        }

    source = event.get('source') or {}
    message = event.get('message') or {}
    sender_display_name = get_line_sender_display_name(event)
    now = now_taipei().isoformat()
    dedupe_key = _line_todo_dedupe_key(target_id, todo_date, title)

    doc_ref = db.collection(LINE_TODO_COLLECTION).document()
    doc_ref.set({
        'title': title,
        'todo_date': todo_date,
        'note': note,
        'remind_time': remind_time,
        'status': 'open',
        'dedupe_key': dedupe_key,
        'line_message_id': message.get('id', ''),
        'line_target_id': target_id,
        'line_target_type': target_type,
        'line_group_id': source.get('groupId', ''),
        'line_room_id': source.get('roomId', ''),
        'line_user_id': source.get('userId', ''),
        'sender_display_name': sender_display_name,
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': sender_display_name or 'LINE Bot',
        'reminder_sent_dates': [],
        'tomorrow_reminder_sent_dates': [],
    })

    return {
        'handled': True,
        'ok': True,
        'todo_date': todo_date,
        'title': title,
        'reply_text': f"已新增代辦\n日期：{_todo_display_md(todo_date)}\n事項：{title}\n\n查詢請回：#今日代辦",
        'parsed_tag': '新增代辦',
    }


# 覆寫批次新增：若同批或資料庫已存在相同代辦，就跳過不新增。
def create_line_todos_bulk(fields_list, event):
    if not fields_list:
        return {
            'handled': True,
            'ok': False,
            'reply_text': (
                '未新增：請在 #新增代辦 下一行開始貼代辦事項。\n\n'
                '範例一：同一天\n'
                '#新增代辦\n'
                '日期: 明天\n'
                '厝米排版 土地現廣稿\n'
                '拍水哥爸爸土地\n\n'
                '範例二：不同天\n'
                '#新增代辦\n'
                '5/29 厝米排版 土地現廣稿\n'
                '5/30 拍水哥爸爸土地'
            ),
            'parsed_tag': '新增代辦',
        }

    ok_items = []
    skipped = []
    failed_msgs = []
    seen_in_this_message = set()
    target_id, _target_type = _line_todo_target_from_event(event)

    for fields in fields_list:
        title = (fields.get('title') or '').strip()
        todo_date = _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
        local_key = _line_todo_dedupe_key(target_id, todo_date, title)
        if title and todo_date and local_key in seen_in_this_message:
            skipped.append(f'{_todo_display_md(todo_date)}｜{title}')
            continue
        if title and todo_date:
            seen_in_this_message.add(local_key)

        result = create_line_todo(fields, event)
        if result.get('ok'):
            ok_items.append({
                'todo_date': result.get('todo_date') or todo_date,
                'title': result.get('title') or title,
            })
        elif result.get('duplicate'):
            skipped.append(f'{_todo_display_md(todo_date)}｜{title}')
        else:
            failed_msgs.append(result.get('reply_text', '新增失敗'))

    lines = []
    if ok_items:
        grouped = {}
        for item in ok_items:
            d = item.get('todo_date', '') or now_taipei().strftime('%Y-%m-%d')
            grouped.setdefault(d, []).append(item.get('title', ''))

        lines.append(f'已新增 {len(ok_items)} 筆代辦')
        if len(grouped) == 1:
            only_date = next(iter(grouped.keys()))
            lines.append(f'日期：{_todo_display_md(only_date)}')
            lines.append('')
            for idx, title in enumerate(grouped[only_date], 1):
                lines.append(f'{idx}. {title}')
        else:
            lines.append('')
            running = 1
            for d in sorted(grouped.keys()):
                lines.append(f'【{_todo_display_md(d)}】')
                for title in grouped[d]:
                    lines.append(f'{running}. {title}')
                    running += 1
                lines.append('')
            if lines and lines[-1] == '':
                lines.pop()

    if skipped:
        if lines:
            lines.append('')
        lines.append(f'已跳過 {len(skipped)} 筆重複代辦')
        for item in skipped[:8]:
            lines.append(f'- {item}')
        if len(skipped) > 8:
            lines.append(f'...另有 {len(skipped) - 8} 筆')

    if failed_msgs:
        if lines:
            lines.append('')
        lines.append('有部分未新增：')
        lines.extend(failed_msgs[:3])

    if lines:
        lines.append('')
        lines.append('查詢請回：#今日代辦')
        lines.append('完成可回：#完成代辦 1')
        return {'handled': True, 'ok': bool(ok_items), 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '新增代辦'}

    return {'handled': True, 'ok': False, 'reply_text': '沒有新增任何代辦。', 'parsed_tag': '新增代辦'}


# 覆寫清單格式：顯示前去重，避免舊資料中已有重複時畫面重複。
def _format_line_todo_sections(overdue_items, today_items, title, today_label='今天'):
    overdue_items = _line_todo_dedupe_docs_for_display(overdue_items)
    today_items = _line_todo_dedupe_docs_for_display(today_items)

    if not overdue_items and not today_items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    idx = 1

    if overdue_items:
        lines.append('')
        lines.append('【尚未完成】')
        for doc in overdue_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            date_text = _todo_display_md(data.get('todo_date', ''))
            lines.append(f"{idx}. {date_text}｜{data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if today_items:
        lines.append('')
        lines.append(f'【{today_label}要做】')
        for doc in today_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    lines.append('一次完成多筆：#完成代辦 1 3')
    lines.append('清除請回：#清除代辦 1')
    return '\n'.join(lines)[:5000]


def _cleanup_duplicate_line_todos_for_target(event):
    """刪除目前 LINE 對話中重複的未完成代辦，保留最早建立的一筆。"""
    target_id, _target_type = _line_todo_target_from_event(event)
    docs = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if data.get('line_target_id') != target_id:
            continue
        docs.append(doc)

    docs = sorted(docs, key=lambda d: ((d.to_dict() or {}).get('created_at') or '', d.id))
    seen = {}
    deleted_titles = []
    updated_keys = 0

    for doc in docs:
        data = doc.to_dict() or {}
        key = data.get('dedupe_key') or _line_todo_doc_dedupe_key(doc)
        # 順手補上舊資料沒有的 dedupe_key，之後更好判斷。
        if not data.get('dedupe_key'):
            try:
                doc.reference.update({'dedupe_key': key})
                updated_keys += 1
            except Exception:
                pass

        if key in seen:
            title = data.get('title', '')
            date_text = _todo_display_md(data.get('todo_date', ''))
            doc.reference.delete()
            deleted_titles.append(f'{date_text}｜{title}')
        else:
            seen[key] = doc.id

    if not deleted_titles:
        msg = '目前沒有找到重複的未完成代辦。'
        if updated_keys:
            msg += f'\n已順手整理 {updated_keys} 筆代辦索引。'
        return {'handled': True, 'ok': True, 'reply_text': msg, 'parsed_tag': '清除重複代辦'}

    lines = [f'已清除 {len(deleted_titles)} 筆重複代辦']
    for item in deleted_titles[:12]:
        lines.append(f'- {item}')
    if len(deleted_titles) > 12:
        lines.append(f'...另有 {len(deleted_titles) - 12} 筆')
    lines.append('')
    lines.append('查詢請回：#今日代辦')
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '清除重複代辦'}


# 先攔截「清除重複代辦」與「同一 LINE 訊息重送」。
_process_line_todo_message_event_before_duplicate_guard = process_line_todo_message_event


def process_line_todo_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if not raw_text.startswith('#'):
        return {'handled': False}

    if raw_text.startswith('#清除重複代辦') or raw_text.startswith('#整理代辦'):
        result = _cleanup_duplicate_line_todos_for_target(event)
        save_line_log(
            {'tag': result.get('parsed_tag', '清除重複代辦'), 'action': 'line_todo_cleanup_duplicates', 'fields': {}, 'raw_text': raw_text},
            event,
            'success' if result.get('ok') else 'failed',
            note=result.get('reply_text', ''),
            sender_display_name=get_line_sender_display_name(event),
        )
        return result

    # LINE webhook 有時會 retry，同一 message_id 若已處理過，就不要再新增一次。
    if raw_text.startswith('#新增代辦'):
        message_id = message.get('id', '')
        if message_id:
            processed_ref = db.collection(LINE_TODO_PROCESSED_COLLECTION).document(str(message_id))
            processed_doc = processed_ref.get()
            if processed_doc.exists:
                return {
                    'handled': True,
                    'ok': True,
                    'reply_text': '這則新增代辦已經處理過，為避免重複新增，本次不再新增。\n\n查詢請回：#今日代辦',
                    'parsed_tag': '新增代辦',
                }

            result = _process_line_todo_message_event_before_duplicate_guard(event)
            try:
                processed_ref.set({
                    'message_id': message_id,
                    'raw_text': raw_text[:2000],
                    'ok': bool(result.get('ok')),
                    'created_at': now_taipei().isoformat(),
                }, merge=True)
            except Exception:
                pass
            return result

    return _process_line_todo_message_event_before_duplicate_guard(event)

# ========= LINE Bot 代辦事項 Patch v8 End =========


# ========= LINE Bot 代辦事項 Patch v9：台灣時間確認 + 設定後重置提醒標記 + 診斷指令 =========
# 使用方式：貼在 v8 patch 的最底部。
# 目的：
# 1. 明確顯示「目前台灣時間」與提醒時間判斷。
# 2. 設定提醒時間後，會清除該對話今天/明天的已提醒標記，避免你改時間後因為已提醒過而不再跳出。
# 3. 提供 #測試代辦時間 / #檢查代辦時間 指令，確認設定是否真的存到目前 LINE 對話。
# 4. /line/todos/reminder-check 回傳結果也會顯示 taipei_now 與 timezone。

from datetime import timezone


def _line_todo_taipei_now_text():
    dt = now_taipei()
    return dt.strftime('%m/%d %H:%M')


def _line_todo_current_minutes_taipei():
    dt = now_taipei()
    return dt.hour * 60 + dt.minute


def _line_todo_time_status_text(hhmm: str, default: str):
    now_minutes = _line_todo_current_minutes_taipei()
    target_minutes = _time_to_minutes(hhmm, default=default)
    return '已到提醒時間' if now_minutes >= target_minutes else '尚未到提醒時間'


def _line_todo_reset_sent_markers_for_target(target_id: str, reset_today=False, reset_tomorrow=False):
    """設定提醒時間後，清除目前對話的已提醒標記，避免改時間後被舊標記擋住。"""
    if not target_id:
        return 0

    now_dt = now_taipei()
    today = now_dt.strftime('%Y-%m-%d')
    tomorrow = (now_dt.date() + timedelta(days=1)).strftime('%Y-%m-%d')
    count = 0

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if data.get('line_target_id') != target_id:
            continue

        updates = {}
        if reset_today:
            updates['reminder_sent_dates'] = firestore.ArrayRemove([today])
        if reset_tomorrow:
            updates['tomorrow_reminder_sent_dates'] = firestore.ArrayRemove([tomorrow])

        if updates:
            try:
                doc.reference.update(updates)
                count += 1
            except Exception as e:
                print('⚠️ 重置代辦提醒標記失敗：', e)

    return count


def _line_todo_count_open_items_for_target(target_id: str):
    now_dt = now_taipei()
    today = now_dt.strftime('%Y-%m-%d')
    tomorrow = (now_dt.date() + timedelta(days=1)).strftime('%Y-%m-%d')
    overdue_or_today = 0
    tomorrow_count = 0

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if data.get('line_target_id') != target_id:
            continue
        todo_date = (data.get('todo_date') or '').strip()
        if todo_date and todo_date <= today:
            overdue_or_today += 1
        if todo_date == tomorrow:
            tomorrow_count += 1

    return overdue_or_today, tomorrow_count


def _line_todo_settings_debug_text(event):
    target_id, target_type = _line_todo_target_from_event(event)
    settings = _get_line_todo_reminder_settings(target_id, target_type)
    today_count, tomorrow_count = _line_todo_count_open_items_for_target(target_id)

    today_time = settings.get('today_reminder_time', '08:00')
    tomorrow_time = settings.get('tomorrow_reminder_time', '23:00')

    lines = [
        '目前固定每日代辦提醒設定',
        f'現在台灣時間：{_line_todo_taipei_now_text()}',
        '時區：Asia/Taipei',
        '',
        f"今日代辦提醒：{'開啟' if settings.get('today_enabled', True) else '關閉'}｜{today_time}｜{_line_todo_time_status_text(today_time, '08:00')}",
        f"今日開頭：{_line_todo_preview_opening(settings.get('today_opening_text', ''))}",
        f'今日＋尚未完成筆數：{today_count}',
        '',
        f"明天代辦提醒：{'開啟' if settings.get('tomorrow_enabled', True) else '關閉'}｜{tomorrow_time}｜{_line_todo_time_status_text(tomorrow_time, '23:00')}",
        f"明天開頭：{_line_todo_preview_opening(settings.get('tomorrow_opening_text', ''))}",
        f'明天代辦筆數：{tomorrow_count}',
        '',
        '提醒是用台灣時間判斷；如果到了時間仍沒跳，請確認排程器有每 5 分鐘呼叫：',
        '/line/todos/reminder-check?key=你的密鑰',
    ]
    return '\n'.join(lines)[:5000]


# 覆寫查詢設定：直接顯示台灣時間與是否已到提醒時間。
def query_line_todo_reminder_settings(event):
    return {
        'handled': True,
        'ok': True,
        'reply_text': _line_todo_settings_debug_text(event),
        'parsed_tag': '代辦提醒設定',
    }


# 覆寫設定提醒：設定時間後重置已提醒標記，並回覆目前台灣時間與判斷狀態。
def set_line_todo_reminder_settings_from_command(raw_text: str, event):
    target_id, target_type = _line_todo_target_from_event(event)
    updates, errors = _parse_line_todo_reminder_setting_command(raw_text)

    if not updates:
        msg = (
            '請告訴我要設定的固定每日提醒時間或開頭。\n\n'
            '範例：\n'
            '#設定代辦提醒\n'
            '今日提醒: 08:00\n'
            '今日開頭: 各位厝米的夥伴早安 ☀️\n'
            '明天提醒: 23:00\n'
            '明天開頭: 各位厝米的夥伴晚安 🌙'
        )
        if errors:
            msg += '\n\n' + '\n'.join(errors[:3])
        return {'handled': True, 'ok': False, 'reply_text': msg[:5000], 'parsed_tag': '設定代辦提醒'}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '設定代辦提醒'}

    reset_today = 'today_reminder_time' in updates
    reset_tomorrow = 'tomorrow_reminder_time' in updates
    reset_count = _line_todo_reset_sent_markers_for_target(target_id, reset_today=reset_today, reset_tomorrow=reset_tomorrow)

    lines = ['已更新固定每日代辦提醒設定']
    if reset_count:
        lines.append(f'已重新開放 {reset_count} 筆代辦的提醒判斷')
    lines.append('')
    lines.append(_line_todo_settings_debug_text(event))
    lines.append('')
    lines.append('如果現在已經超過設定時間，等排程器下一次呼叫 reminder-check 就會推播。')
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '設定代辦提醒'}


# 新增診斷指令：#測試代辦時間 / #檢查代辦時間 / #確認代辦時間。
_process_line_todo_reminder_settings_message_event_before_v9_time_debug = process_line_todo_reminder_settings_message_event


def process_line_todo_reminder_settings_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if raw_text.startswith(('#測試代辦時間', '#檢查代辦時間', '#確認代辦時間', '#代辦時間')):
        result = query_line_todo_reminder_settings(event)
        save_line_log(
            {'tag': '代辦提醒時間診斷', 'action': 'line_todo_reminder_time_debug', 'fields': {}, 'raw_text': raw_text},
            event,
            'success',
            note=result.get('reply_text', ''),
            sender_display_name=get_line_sender_display_name(event),
        )
        return result

    return _process_line_todo_reminder_settings_message_event_before_v9_time_debug(event)


# 覆寫 reminder-check route：回傳台灣時間，方便你用瀏覽器確認是不是用台灣時間。
# 注意：如果你原本已經有同 endpoint route，Flask 不允許重複註冊 endpoint。
# 所以這裡不重複註冊 route，只覆寫 view_functions 裡的 endpoint 函式。
def line_todos_reminder_check_v9():
    secret = os.environ.get('TODO_REMINDER_SECRET', '').strip()
    key = request.args.get('key', '').strip() or request.form.get('key', '').strip()
    if secret and key != secret:
        return {
            'ok': False,
            'message': 'Invalid key',
            'taipei_now': now_taipei().strftime('%Y-%m-%d %H:%M:%S'),
            'timezone': 'Asia/Taipei',
        }, 403

    result = send_due_line_todo_reminders_by_settings()
    result['taipei_now'] = now_taipei().strftime('%Y-%m-%d %H:%M:%S')
    result['timezone'] = 'Asia/Taipei'
    return {'ok': True, 'result': result}, 200


try:
    app.view_functions['line_todos_reminder_check'] = line_todos_reminder_check_v9
except Exception as e:
    print('⚠️ 套用 reminder-check v9 view 覆寫失敗：', e)

# ========= LINE Bot 代辦事項 Patch v9 End =========


# ========= LINE Bot 待辦事項 Patch v10：用字修正（代辦 → 待辦） =========
# 使用方式：貼在目前 v9 patch 的最底部。
# 目的：
# 1. LINE 顯示文字統一改成「待辦 / 待辦事項」。
# 2. 指令支援新寫法：#新增待辦、#今日待辦、#完成待辦、#清除待辦、#設定待辦提醒。
# 3. 舊指令仍相容：#新增代辦、#今日代辦、#完成代辦、#清除代辦、#設定代辦提醒。
# 4. Firestore collection / 既有資料不改，避免資料搬移風險。


def _line_todo_wording_fix_text(value):
    """只修正使用者看得到的文字，不動資料庫欄位。"""
    if not isinstance(value, str):
        return value
    return value.replace('代辦事項', '待辦事項').replace('代辦', '待辦')


def _line_todo_normalize_command_wording(value):
    """讓使用者輸入「待辦」時，內部仍可走原本 v9 的「代辦」判斷。"""
    if not isinstance(value, str):
        return value
    return value.replace('待辦事項', '代辦事項').replace('待辦', '代辦')


# 1) LINE push 推播文字修正，例如早上 / 晚上固定提醒。
try:
    _push_line_text_before_v10_wording = push_line_text

    def push_line_text(to_id: str, text_message: str):
        return _push_line_text_before_v10_wording(to_id, _line_todo_wording_fix_text(text_message))
except Exception as e:
    print('⚠️ 套用 push_line_text 待辦用字修正失敗：', e)


# 2) reply_line_text 保險修正：如果有直接呼叫 reply_line_text，也會顯示待辦。
try:
    _reply_line_text_before_v10_wording = reply_line_text

    def reply_line_text(reply_token: str, text_message: str):
        return _reply_line_text_before_v10_wording(reply_token, _line_todo_wording_fix_text(text_message))
except Exception as e:
    print('⚠️ 套用 reply_line_text 待辦用字修正失敗：', e)


# 3) 使用者輸入 #新增待辦 / #今日待辦 時，轉成舊版可辨識的 #新增代辦；回覆再改回待辦。
try:
    _process_line_message_event_before_v10_wording = process_line_message_event

    def process_line_message_event(event):
        patched_event = event
        try:
            message = (event or {}).get('message') or {}
            if message.get('type') == 'text':
                raw_text = message.get('text') or ''
                normalized_text = _line_todo_normalize_command_wording(raw_text)
                if normalized_text != raw_text:
                    # 複製 event，避免影響原始 webhook 內容太多。
                    patched_event = dict(event)
                    patched_message = dict(message)
                    patched_message['text'] = normalized_text
                    patched_event['message'] = patched_message
        except Exception as e:
            print('⚠️ 待辦指令文字正規化失敗：', e)
            patched_event = event

        result = _process_line_message_event_before_v10_wording(patched_event)
        try:
            if isinstance(result, dict) and 'reply_text' in result:
                result = dict(result)
                result['reply_text'] = _line_todo_wording_fix_text(result.get('reply_text', ''))
                if 'parsed_tag' in result:
                    result['parsed_tag'] = _line_todo_wording_fix_text(result.get('parsed_tag', ''))
        except Exception as e:
            print('⚠️ 待辦回覆文字修正失敗：', e)
        return result
except Exception as e:
    print('⚠️ 套用 process_line_message_event 待辦用字修正失敗：', e)


# 4) reminder-check 的 JSON 回傳也修正顯示文字，方便瀏覽器測試時看到「待辦」。
try:
    _line_todos_reminder_check_before_v10_wording = app.view_functions.get('line_todos_reminder_check')

    def line_todos_reminder_check_v10_wording():
        response = _line_todos_reminder_check_before_v10_wording()
        try:
            # Flask view 可能回傳 (dict, status_code) 或 dict。
            if isinstance(response, tuple) and response and isinstance(response[0], dict):
                data = response[0]
                status = response[1] if len(response) > 1 else 200
                data_text = _line_todo_wording_fix_text(json.dumps(data, ensure_ascii=False))
                return json.loads(data_text), status
            if isinstance(response, dict):
                data_text = _line_todo_wording_fix_text(json.dumps(response, ensure_ascii=False))
                return json.loads(data_text)
        except Exception as e:
            print('⚠️ reminder-check 待辦用字修正失敗：', e)
        return response

    if _line_todos_reminder_check_before_v10_wording:
        app.view_functions['line_todos_reminder_check'] = line_todos_reminder_check_v10_wording
except Exception as e:
    print('⚠️ 套用 reminder-check 待辦用字修正失敗：', e)

# ========= LINE Bot 待辦事項 Patch v10 End =========
# ========= LINE Bot 待辦事項 Patch v11：新增「未來待辦」區塊 =========
# 使用方式：貼在目前 v10 patch 的最底部。
# 功能：
# 1. #今日待辦 會多一區「未來待辦」，預設顯示未來 7 天。
# 2. 早上今日提醒也會一起帶出「未來待辦」。
# 3. 新增指令：#未來待辦，可查未來 7 天；#未來待辦 14 可查未來 14 天。
# 4. 完成 / 清除仍可用畫面序號；如果未來待辦顯示成第 5 筆，就可用 #完成待辦 5。
# 5. Firestore 資料欄位不搬移；只新增 future_reminder_sent_dates 用於避免同一天重複提醒。

LINE_TODO_FUTURE_DAYS_DEFAULT = int(os.environ.get('LINE_TODO_FUTURE_DAYS_DEFAULT', '7') or '7')
LINE_TODO_FUTURE_DAYS_MAX = int(os.environ.get('LINE_TODO_FUTURE_DAYS_MAX', '30') or '30')


def _line_todo_parse_days_from_text(raw_text: str, default_days=None):
    """從 #未來代辦 7 或 天數: 14 取出查詢天數。"""
    default_days = default_days or LINE_TODO_FUTURE_DAYS_DEFAULT
    text = raw_text or ''
    m = re.search(r'(?:天數|未來|查詢天數)\s*[:：]?\s*(\d{1,2})', text)
    if not m:
        # 例如：#未來代辦 14
        parts = re.split(r'\s+', text.strip())
        if len(parts) >= 2 and re.fullmatch(r'\d{1,2}', parts[1]):
            m_value = parts[1]
        else:
            m_value = ''
    else:
        m_value = m.group(1)

    try:
        days = int(m_value) if m_value else int(default_days)
    except Exception:
        days = int(default_days)
    return max(1, min(LINE_TODO_FUTURE_DAYS_MAX, days))


def _get_future_line_todos(start_date='', target_id='', days=None):
    """取得 start_date 之後、days 天內的未完成待辦。"""
    days = days or LINE_TODO_FUTURE_DAYS_DEFAULT
    start_date = start_date or now_taipei().strftime('%Y-%m-%d')
    try:
        start_dt = datetime.strptime(start_date, '%Y-%m-%d').date()
    except Exception:
        start_dt = now_taipei().date()
        start_date = start_dt.strftime('%Y-%m-%d')

    end_date = (start_dt + timedelta(days=int(days))).strftime('%Y-%m-%d')
    result = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = _todo_date_value(doc)
        if start_date < d <= end_date:
            result.append(doc)
    return _sort_line_todo_docs(result)


def _line_todo_group_docs_by_date(items):
    grouped = {}
    for doc in _sort_line_todo_docs(items or []):
        data = doc.to_dict() or {}
        d = (data.get('todo_date') or '').strip()
        if not d:
            continue
        grouped.setdefault(d, []).append(doc)
    return grouped


def _format_line_todo_sections(overdue_items, today_items, title, today_label='今天', future_items=None, future_days=None):
    """覆寫清單格式：新增「未來待辦」區塊。"""
    overdue_items = _line_todo_dedupe_docs_for_display(overdue_items or [])
    today_items = _line_todo_dedupe_docs_for_display(today_items or [])
    future_items = _line_todo_dedupe_docs_for_display(future_items or [])
    future_days = future_days or LINE_TODO_FUTURE_DAYS_DEFAULT

    if not overdue_items and not today_items and not future_items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    idx = 1

    if overdue_items:
        lines.append('')
        lines.append('【尚未完成】')
        for doc in overdue_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            date_text = _todo_display_md(data.get('todo_date', ''))
            lines.append(f"{idx}. {date_text}｜{data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if today_items:
        lines.append('')
        lines.append(f'【{today_label}要做】')
        for doc in today_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if future_items:
        lines.append('')
        lines.append(f'【未來待辦｜未來 {future_days} 天】')
        grouped = _line_todo_group_docs_by_date(future_items)
        for d in sorted(grouped.keys()):
            lines.append(f'〔{_todo_display_md(d)}〕')
            for doc in grouped[d]:
                data = doc.to_dict() or {}
                note = (data.get('note') or '').strip()
                lines.append(f"{idx}. {data.get('title', '')}")
                if note:
                    lines.append(f'   備註: {note}')
                idx += 1

    lines.append('')
    lines.append('完成請回：#完成代辦 1')
    lines.append('一次完成多筆：#完成代辦 1 3')
    lines.append('清除請回：#清除代辦 1')
    lines.append('查未來請回：#未來代辦')
    return '\n'.join(lines)[:5000]


def _get_display_line_todos(todo_date='', target_id='', include_future=True, future_days=None):
    """畫面顯示順序：尚未完成 → 指定日期 → 未來待辦。序號完成會依這個順序。"""
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    overdue_items = _get_overdue_line_todos(query_date, target_id=target_id)
    today_items = _get_open_line_todos(query_date, target_id=target_id, include_overdue=False)
    items = overdue_items + today_items

    # 只有查今天時，才把未來待辦接在同一份清單後面，避免查指定日期時序號混亂。
    if include_future and query_date == now_taipei().strftime('%Y-%m-%d'):
        items += _get_future_line_todos(start_date=query_date, target_id=target_id, days=future_days or LINE_TODO_FUTURE_DAYS_DEFAULT)
    return items


def _find_line_todo_v2(todo_key: str, target_id='', todo_date=''):
    """覆寫序號完成邏輯：支援完成「未來待辦」區塊的序號。"""
    key = (todo_key or '').strip()
    if not key:
        return None, '請提供待辦序號、ID 或事項關鍵字。'

    query_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')

    if re.fullmatch(r'\d+', key):
        items = _get_display_line_todos(todo_date=query_date, target_id=target_id, include_future=True)
        idx = int(key) - 1
        if 0 <= idx < len(items):
            return items[idx], ''
        return None, f'找不到第 {key} 筆待辦，請先輸入 #今日待辦 確認序號。'

    return _find_line_todo(key, target_id=target_id)


def query_line_todos(fields, event, force_today=False):
    """覆寫查詢：今天清單會附上未來待辦。"""
    target_id, _ = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    if not todo_date:
        todo_date = now_taipei().strftime('%Y-%m-%d')

    overdue_items = _get_overdue_line_todos(todo_date=todo_date, target_id=target_id)
    today_items = _get_open_line_todos(todo_date=todo_date, target_id=target_id, include_overdue=False)
    future_items = []
    future_days = LINE_TODO_FUTURE_DAYS_DEFAULT

    if todo_date == now_taipei().strftime('%Y-%m-%d'):
        future_items = _get_future_line_todos(start_date=todo_date, target_id=target_id, days=future_days)
        title = f'{_todo_display_md(todo_date)} 代辦清單'
        today_label = '今天'
    else:
        title = f'{_todo_display_md(todo_date)} 代辦清單'
        today_label = _todo_display_md(todo_date)

    return {
        'handled': True,
        'ok': True,
        'reply_text': _format_line_todo_sections(overdue_items, today_items, title, today_label=today_label, future_items=future_items, future_days=future_days),
        'parsed_tag': '查詢代辦',
    }


def query_future_line_todos(event, days=None):
    """查詢未來待辦，不含今天、不含逾期。"""
    target_id, _ = _line_todo_target_from_event(event)
    days = days or LINE_TODO_FUTURE_DAYS_DEFAULT
    today = now_taipei().strftime('%Y-%m-%d')
    future_items = _get_future_line_todos(start_date=today, target_id=target_id, days=days)

    if not future_items:
        return {
            'handled': True,
            'ok': True,
            'reply_text': f'未來 {days} 天目前沒有未完成待辦。',
            'parsed_tag': '未來代辦',
        }

    lines = [f'未來 {days} 天待辦']
    idx = 1
    grouped = _line_todo_group_docs_by_date(future_items)
    for d in sorted(grouped.keys()):
        lines.append('')
        lines.append(f'【{_todo_display_md(d)}】')
        for doc in grouped[d]:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    lines.append('')
    lines.append('完成請先回 #今日代辦 看完整序號，或用事項關鍵字完成。')
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '未來代辦'}


_process_line_todo_message_event_before_v11_future = process_line_todo_message_event


def process_line_todo_message_event(event):
    """新增 #未來代辦；其它指令照原本 v10 流程。"""
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    if not raw_text.startswith('#'):
        return {'handled': False}

    if raw_text.startswith(('#未來代辦', '#近期代辦', '#未來清單')):
        days = _line_todo_parse_days_from_text(raw_text, default_days=LINE_TODO_FUTURE_DAYS_DEFAULT)
        result = query_future_line_todos(event, days=days)
        save_line_log(
            {'tag': result.get('parsed_tag', '未來代辦'), 'action': 'line_todo_future_query', 'fields': {'days': days}, 'raw_text': raw_text},
            event,
            'success' if result.get('ok') else 'failed',
            note=result.get('reply_text', ''),
            sender_display_name=get_line_sender_display_name(event),
        )
        return result

    return _process_line_todo_message_event_before_v11_future(event)


def send_due_line_todo_reminders_by_settings():
    """
    覆寫每日提醒：今日提醒加入「未來待辦」。
    - 今日提醒：尚未完成 + 今天要做 + 未來 7 天。
    - 明天提醒：維持只提醒明天要做。
    - 未來待辦使用 future_reminder_sent_dates 防止同一天重複推播。
    """
    now_dt = now_taipei()
    today = now_dt.strftime('%Y-%m-%d')
    tomorrow = (now_dt.date() + timedelta(days=1)).strftime('%Y-%m-%d')
    current_minutes = now_dt.hour * 60 + now_dt.minute

    targets = _collect_line_todo_targets_with_open_items()
    sent_count = 0
    failed = []
    checked_targets = 0
    future_days = LINE_TODO_FUTURE_DAYS_DEFAULT

    for target_id, target_type in targets.items():
        checked_targets += 1
        settings = _get_line_todo_reminder_settings(target_id, target_type)

        # 今日提醒：尚未完成 + 今日要做 + 未來待辦。
        if settings.get('today_enabled', True):
            today_time = settings.get('today_reminder_time', '08:00')
            if current_minutes >= _time_to_minutes(today_time, default='08:00'):
                due_items = []
                future_items_for_push = []

                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if not todo_date:
                        continue

                    # 尚未完成 + 今天要做：每天各待辦提醒一次。
                    if todo_date <= today:
                        sent_dates = data.get('reminder_sent_dates') or []
                        if today not in sent_dates:
                            due_items.append(doc)
                        continue

                    # 未來待辦：每天提醒清單可以帶一次，不影響該待辦到當天時的正式提醒。
                    try:
                        end_date = (now_dt.date() + timedelta(days=future_days)).strftime('%Y-%m-%d')
                    except Exception:
                        end_date = today
                    if today < todo_date <= end_date:
                        sent_dates = data.get('future_reminder_sent_dates') or []
                        if today not in sent_dates:
                            future_items_for_push.append(doc)

                if due_items or future_items_for_push:
                    overdue_items = _sort_line_todo_docs([d for d in due_items if _todo_date_value(d) < today])
                    today_items = _sort_line_todo_docs([d for d in due_items if _todo_date_value(d) == today])
                    future_items = _sort_line_todo_docs(future_items_for_push)
                    body = _format_line_todo_sections(
                        overdue_items,
                        today_items,
                        f'{_todo_display_md(today)} 今日代辦',
                        today_label='今天',
                        future_items=future_items,
                        future_days=future_days,
                    )
                    text = _line_todo_add_opening(settings.get('today_opening_text', ''), body)
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in due_items:
                            doc.reference.update({
                                'reminder_sent_dates': firestore.ArrayUnion([today]),
                                'last_reminded_at': now_taipei().isoformat(),
                            })
                        for doc in future_items_for_push:
                            doc.reference.update({
                                'future_reminder_sent_dates': firestore.ArrayUnion([today]),
                                'last_future_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'today_with_future', 'error': msg})

        # 明天提醒：維持明天要做。
        if settings.get('tomorrow_enabled', True):
            tomorrow_time = settings.get('tomorrow_reminder_time', '23:00')
            if current_minutes >= _time_to_minutes(tomorrow_time, default='23:00'):
                items = []
                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if todo_date != tomorrow:
                        continue
                    sent_dates = data.get('tomorrow_reminder_sent_dates') or []
                    if tomorrow in sent_dates:
                        continue
                    items.append(doc)

                if items:
                    tomorrow_items = _sort_line_todo_docs(items)
                    body = _format_line_todo_sections(
                        [],
                        tomorrow_items,
                        f'{_todo_display_md(tomorrow)} 明天代辦',
                        today_label='明天',
                    )
                    text = _line_todo_add_opening(settings.get('tomorrow_opening_text', ''), body)
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in items:
                            doc.reference.update({
                                'tomorrow_reminder_sent_dates': firestore.ArrayUnion([tomorrow]),
                                'last_tomorrow_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'tomorrow', 'error': msg})

    return {
        'now': now_dt.strftime('%Y-%m-%d %H:%M:%S'),
        'checked_targets': checked_targets,
        'sent_count': sent_count,
        'future_days': future_days,
        'failed': failed,
    }

# ========= LINE Bot 待辦事項 Patch v11 End =========


# ========= LINE Bot 待辦事項 Patch v12：未來待辦「提前幾天提醒」設定 =========
# 使用方式：貼在目前 v11 patch 的最底部。
# 功能：
# 1. 可用 LINE 指令設定未來待辦提前幾天開始出現在每日提醒。
#    例如：未來提醒: 3天前，代表未來 3 天內的待辦會在早上提醒中顯示。
# 2. 支援 #設定未來待辦提醒 3天前。
# 3. 支援 #關閉待辦提醒 未來 / #開啟待辦提醒 未來。
# 4. #待辦提醒設定 / #測試待辦時間 會顯示未來提醒設定與筆數。
# 5. Firestore 新增欄位：future_enabled、future_reminder_days。

LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT = int(
    os.environ.get(
        'LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT',
        os.environ.get('LINE_TODO_FUTURE_DAYS_DEFAULT', '7')
    ) or '7'
)
LINE_TODO_FUTURE_REMINDER_DAYS_MAX = int(
    os.environ.get(
        'LINE_TODO_FUTURE_REMINDER_DAYS_MAX',
        os.environ.get('LINE_TODO_FUTURE_DAYS_MAX', '30')
    ) or '30'
)


def _line_todo_clamp_future_days(value, default=None):
    default = default or LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT
    try:
        days = int(str(value).strip())
    except Exception:
        days = int(default)
    return max(1, min(LINE_TODO_FUTURE_REMINDER_DAYS_MAX, days))


def _line_todo_get_future_reminder_days(settings=None):
    settings = settings or {}
    return _line_todo_clamp_future_days(
        settings.get('future_reminder_days', LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT),
        default=LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT,
    )


def _line_todo_is_future_enabled(settings=None):
    settings = settings or {}
    return bool(settings.get('future_enabled', True))


def _line_todo_extract_days_number(text: str):
    text = str(text or '').strip()
    m = re.search(r'(\d{1,2})\s*(?:天前|天|日|days?|d)?', text, flags=re.IGNORECASE)
    if not m:
        return None
    try:
        return int(m.group(1))
    except Exception:
        return None


def _line_todo_parse_future_reminder_updates(raw_text: str):
    """解析未來待辦提醒設定，例如：未來提醒: 3天前 / 未來提醒: 關閉 / #設定未來待辦提醒 3天前。"""
    text = (raw_text or '').strip()
    normalized = _line_todo_normalize_command_wording(text) if '_line_todo_normalize_command_wording' in globals() else text
    lines = [ln.strip() for ln in normalized.splitlines() if ln.strip()]
    joined = ' '.join(lines)

    updates = {}
    errors = []

    def apply_value(label, value):
        v = str(value or '').strip()
        compact = re.sub(r'\s+', '', v).lower()
        if any(x in compact for x in ['關閉', '停用', '取消', '不要', 'off', 'false', 'no']):
            updates['future_enabled'] = False
            return
        if any(x in compact for x in ['開啟', '啟用', '恢復', 'on', 'true', 'yes']):
            updates['future_enabled'] = True
            # 如果同時有數字，例如「開啟 3天前」，也一起更新天數。
        days = _line_todo_extract_days_number(v)
        if days is not None:
            if days <= 0:
                updates['future_enabled'] = False
            else:
                updates['future_enabled'] = True
                updates['future_reminder_days'] = _line_todo_clamp_future_days(days)
            return
        if label:
            errors.append(f'{label} 的天數看不懂：{value}')

    future_key_words = ['未來提醒', '未來代辦提醒', '未來代辦', '提前提醒', '提前天數', '幾天前提醒', '提前幾天']

    # 多行 key: value 格式。
    for line in lines:
        m = re.match(r'^([^:：]+)\s*[:：]\s*(.+)$', line)
        if not m:
            continue
        key = re.sub(r'\s+', '', m.group(1).strip())
        value = m.group(2).strip()
        if any(k in key for k in future_key_words):
            apply_value(key, value)

    # 一行格式，例如：#設定待辦提醒 未來提醒 3天前 / #設定未來待辦提醒 3天前。
    if any(k in joined for k in future_key_words) or normalized.startswith(('#設定未來代辦提醒', '#設定未來提醒', '#設定提前提醒')):
        # 優先找「3天前」這種清楚格式。
        m = re.search(r'(\d{1,2})\s*天前', joined)
        if not m:
            m = re.search(r'(?:未來(?:代辦)?提醒|未來代辦|未來提醒|提前(?:提醒|天數|幾天)?|幾天前提醒)\D{0,12}(\d{1,2})', joined)
        if m:
            days = int(m.group(1))
            if days <= 0:
                updates['future_enabled'] = False
            else:
                updates['future_enabled'] = True
                updates['future_reminder_days'] = _line_todo_clamp_future_days(days)
        elif re.search(r'(?:未來|提前).*(關閉|停用|取消|不要|off|false|no)', joined, flags=re.IGNORECASE):
            updates['future_enabled'] = False
        elif re.search(r'(?:未來|提前).*(開啟|啟用|恢復|on|true|yes)', joined, flags=re.IGNORECASE):
            updates['future_enabled'] = True

    return updates, errors


_parse_line_todo_reminder_setting_command_before_v12_future_days = _parse_line_todo_reminder_setting_command


def _parse_line_todo_reminder_setting_command(raw_text: str):
    updates, errors = _parse_line_todo_reminder_setting_command_before_v12_future_days(raw_text)
    future_updates, future_errors = _line_todo_parse_future_reminder_updates(raw_text)
    updates.update(future_updates)
    errors.extend(future_errors)
    return updates, errors


def _line_todo_days_until_text(date_str: str):
    try:
        target_date = datetime.strptime(date_str, '%Y-%m-%d').date()
        diff = (target_date - now_taipei().date()).days
        if diff == 1:
            return '明天'
        if diff == 2:
            return '後天'
        if diff > 0:
            return f'{diff}天後'
    except Exception:
        pass
    return ''


def _format_line_todo_sections(overdue_items, today_items, title, today_label='今天', future_items=None, future_days=None):
    """覆寫清單格式：未來待辦改成「提前 N 天提醒」。"""
    overdue_items = _line_todo_dedupe_docs_for_display(overdue_items or [])
    today_items = _line_todo_dedupe_docs_for_display(today_items or [])
    future_items = _line_todo_dedupe_docs_for_display(future_items or [])
    future_days = _line_todo_clamp_future_days(future_days or LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT)

    if not overdue_items and not today_items and not future_items:
        return f'{title}\n目前沒有未完成代辦。'

    lines = [title]
    idx = 1

    if overdue_items:
        lines.append('')
        lines.append('【尚未完成】')
        for doc in overdue_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            date_text = _todo_display_md(data.get('todo_date', ''))
            lines.append(f"{idx}. {date_text}｜{data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if today_items:
        lines.append('')
        lines.append(f'【{today_label}要做】')
        for doc in today_items:
            data = doc.to_dict() or {}
            note = (data.get('note') or '').strip()
            lines.append(f"{idx}. {data.get('title', '')}")
            if note:
                lines.append(f'   備註: {note}')
            idx += 1

    if future_items:
        lines.append('')
        lines.append(f'【未來待辦｜提前 {future_days} 天提醒】')
        grouped = _line_todo_group_docs_by_date(future_items)
        for d in sorted(grouped.keys()):
            days_text = _line_todo_days_until_text(d)
            if days_text:
                lines.append(f'〔{_todo_display_md(d)}｜{days_text}〕')
            else:
                lines.append(f'〔{_todo_display_md(d)}〕')
            for doc in grouped[d]:
                data = doc.to_dict() or {}
                note = (data.get('note') or '').strip()
                lines.append(f"{idx}. {data.get('title', '')}")
                if note:
                    lines.append(f'   備註: {note}')
                idx += 1

    lines.append('')
    lines.append('完成請回：#完成待辦 1')
    lines.append('一次完成多筆：#完成待辦 1 3')
    lines.append('清除請回：#清除待辦 1')
    lines.append('查未來請回：#未來待辦')
    return '\n'.join(lines)[:5000]


def _line_todo_count_future_items_for_target(target_id: str, days=None):
    days = _line_todo_clamp_future_days(days or LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT)
    today_date = now_taipei().date()
    today = today_date.strftime('%Y-%m-%d')
    end_date = (today_date + timedelta(days=days)).strftime('%Y-%m-%d')
    count = 0
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if data.get('line_target_id') != target_id:
            continue
        todo_date = (data.get('todo_date') or '').strip()
        if today < todo_date <= end_date:
            count += 1
    return count


def _line_todo_reset_future_sent_markers_for_target(target_id: str):
    """設定未來提醒天數 / 開關後，清除今天的未來提醒標記，讓新設定可重新判斷。"""
    if not target_id:
        return 0
    today = now_taipei().strftime('%Y-%m-%d')
    count = 0
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if data.get('line_target_id') != target_id:
            continue
        try:
            doc.reference.update({'future_reminder_sent_dates': firestore.ArrayRemove([today])})
            count += 1
        except Exception as e:
            print('⚠️ 重置未來待辦提醒標記失敗：', e)
    return count


def _line_todo_settings_debug_text(event):
    target_id, target_type = _line_todo_target_from_event(event)
    settings = _get_line_todo_reminder_settings(target_id, target_type)
    today_count, tomorrow_count = _line_todo_count_open_items_for_target(target_id)

    today_time = settings.get('today_reminder_time', '08:00')
    tomorrow_time = settings.get('tomorrow_reminder_time', '23:00')
    future_days = _line_todo_get_future_reminder_days(settings)
    future_count = _line_todo_count_future_items_for_target(target_id, days=future_days)

    lines = [
        '目前固定每日待辦提醒設定',
        f'現在台灣時間：{_line_todo_taipei_now_text()}',
        '時區：Asia/Taipei',
        '',
        f"今日待辦提醒：{'開啟' if settings.get('today_enabled', True) else '關閉'}｜{today_time}｜{_line_todo_time_status_text(today_time, '08:00')}",
        f"今日開頭：{_line_todo_preview_opening(settings.get('today_opening_text', ''))}",
        f'今日＋尚未完成筆數：{today_count}',
        '',
        f"未來待辦提醒：{'開啟' if _line_todo_is_future_enabled(settings) else '關閉'}｜提前 {future_days} 天開始提醒",
        f'未來 {future_days} 天待辦筆數：{future_count}',
        '',
        f"明天待辦提醒：{'開啟' if settings.get('tomorrow_enabled', True) else '關閉'}｜{tomorrow_time}｜{_line_todo_time_status_text(tomorrow_time, '23:00')}",
        f"明天開頭：{_line_todo_preview_opening(settings.get('tomorrow_opening_text', ''))}",
        f'明天待辦筆數：{tomorrow_count}',
        '',
        '提醒是用台灣時間判斷；排程器仍需要每 5 分鐘呼叫：',
        '/line/todos/reminder-check?key=你的密鑰',
    ]
    return '\n'.join(lines)[:5000]


def query_line_todo_reminder_settings(event):
    return {
        'handled': True,
        'ok': True,
        'reply_text': _line_todo_settings_debug_text(event),
        'parsed_tag': '待辦提醒設定',
    }


def set_line_todo_reminder_settings_from_command(raw_text: str, event):
    target_id, target_type = _line_todo_target_from_event(event)
    updates, errors = _parse_line_todo_reminder_setting_command(raw_text)

    if not updates:
        msg = (
            '請告訴我要設定的固定每日提醒時間、開頭或未來待辦提前天數。\n\n'
            '範例：\n'
            '#設定待辦提醒\n'
            '今日提醒: 08:00\n'
            '今日開頭: 各位厝米的夥伴早安 ☀️\n'
            '未來提醒: 3天前\n'
            '明天提醒: 23:00\n'
            '明天開頭: 各位厝米的夥伴晚安 🌙\n\n'
            '也可以：#設定未來待辦提醒 3天前'
        )
        if errors:
            msg += '\n\n' + '\n'.join(errors[:3])
        return {'handled': True, 'ok': False, 'reply_text': msg[:5000], 'parsed_tag': '設定待辦提醒'}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '設定待辦提醒'}

    reset_today = 'today_reminder_time' in updates
    reset_tomorrow = 'tomorrow_reminder_time' in updates
    reset_future = any(k in updates for k in ['future_enabled', 'future_reminder_days'])
    reset_count = _line_todo_reset_sent_markers_for_target(target_id, reset_today=reset_today, reset_tomorrow=reset_tomorrow)
    if reset_future:
        reset_count += _line_todo_reset_future_sent_markers_for_target(target_id)

    lines = ['已更新固定每日待辦提醒設定']
    if reset_count:
        lines.append(f'已重新開放 {reset_count} 筆待辦的提醒判斷')
    lines.append('')
    lines.append(_line_todo_settings_debug_text(event))
    lines.append('')
    lines.append('未來提醒說明：例如設定 3天前，代表未來 3 天內的待辦會出現在每日提醒的【未來待辦】區塊。')
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines)[:5000], 'parsed_tag': '設定待辦提醒'}


_switch_line_todo_reminder_before_v12_future_days = switch_line_todo_reminder


def switch_line_todo_reminder(raw_text: str, event, enabled: bool):
    target_id, target_type = _line_todo_target_from_event(event)
    normalized = _line_todo_normalize_command_wording(raw_text) if '_line_todo_normalize_command_wording' in globals() else raw_text
    body = normalized.replace('#關閉代辦提醒', '', 1).replace('#開啟代辦提醒', '', 1).strip()
    body = body or '全部'

    updates = {}
    if any(x in body for x in ['今日', '今天']):
        updates['today_enabled'] = enabled
    if any(x in body for x in ['明日', '明天']):
        updates['tomorrow_enabled'] = enabled
    if any(x in body for x in ['未來', '提前']):
        updates['future_enabled'] = enabled
    if not updates or '全部' in body or '全開' in body or '全關' in body:
        updates = {'today_enabled': enabled, 'tomorrow_enabled': enabled, 'future_enabled': enabled}

    ok, err = _save_line_todo_reminder_settings(target_id, target_type, updates, event=event)
    if not ok:
        return {'handled': True, 'ok': False, 'reply_text': err, 'parsed_tag': '待辦提醒設定'}

    if 'future_enabled' in updates:
        _line_todo_reset_future_sent_markers_for_target(target_id)

    current = _get_line_todo_reminder_settings(target_id, target_type)
    action = '開啟' if enabled else '關閉'
    future_days = _line_todo_get_future_reminder_days(current)
    lines = [
        f'已{action}待辦提醒',
        f"今日待辦提醒：{'開啟' if current.get('today_enabled', True) else '關閉'}｜{current.get('today_reminder_time', '08:00')}",
        f"未來待辦提醒：{'開啟' if _line_todo_is_future_enabled(current) else '關閉'}｜提前 {future_days} 天開始提醒",
        f"明天待辦提醒：{'開啟' if current.get('tomorrow_enabled', True) else '關閉'}｜{current.get('tomorrow_reminder_time', '23:00')}",
    ]
    return {'handled': True, 'ok': True, 'reply_text': '\n'.join(lines), 'parsed_tag': '待辦提醒設定'}


_process_line_todo_reminder_settings_message_event_before_v12_future_days = process_line_todo_reminder_settings_message_event


def process_line_todo_reminder_settings_message_event(event):
    message = event.get('message') or {}
    if message.get('type') != 'text':
        return {'handled': False}

    raw_text = (message.get('text') or '').strip()
    normalized = _line_todo_normalize_command_wording(raw_text) if '_line_todo_normalize_command_wording' in globals() else raw_text

    if normalized.startswith(('#設定未來代辦提醒', '#設定未來提醒', '#設定提前提醒')):
        result = set_line_todo_reminder_settings_from_command(normalized, event)
        save_line_log(
            {'tag': result.get('parsed_tag', '設定待辦提醒'), 'action': 'line_todo_future_reminder_setting_update', 'fields': {}, 'raw_text': raw_text},
            event,
            'success' if result.get('ok') else 'failed',
            note=result.get('reply_text', ''),
            sender_display_name=get_line_sender_display_name(event),
        )
        return result

    if normalized.startswith(('#關閉未來代辦提醒', '#開啟未來代辦提醒')):
        enabled = normalized.startswith('#開啟')
        fake_text = ('#開啟代辦提醒 未來' if enabled else '#關閉代辦提醒 未來')
        result = switch_line_todo_reminder(fake_text, event, enabled=enabled)
        save_line_log(
            {'tag': result.get('parsed_tag', '待辦提醒設定'), 'action': 'line_todo_future_reminder_switch', 'fields': {'enabled': enabled}, 'raw_text': raw_text},
            event,
            'success' if result.get('ok') else 'failed',
            note=result.get('reply_text', ''),
            sender_display_name=get_line_sender_display_name(event),
        )
        return result

    return _process_line_todo_reminder_settings_message_event_before_v12_future_days(event)


def _get_display_line_todos(todo_date='', target_id='', include_future=True, future_days=None):
    """畫面顯示順序：尚未完成 → 指定日期 → 未來待辦。序號完成會依這個順序。"""
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    overdue_items = _get_overdue_line_todos(query_date, target_id=target_id)
    today_items = _get_open_line_todos(query_date, target_id=target_id, include_overdue=False)
    items = overdue_items + today_items

    if include_future and query_date == now_taipei().strftime('%Y-%m-%d'):
        settings = _get_line_todo_reminder_settings(target_id, '')
        if _line_todo_is_future_enabled(settings):
            days = future_days or _line_todo_get_future_reminder_days(settings)
            items += _get_future_line_todos(start_date=query_date, target_id=target_id, days=days)
    return items


def query_line_todos(fields, event, force_today=False):
    """覆寫查詢：今天清單會依設定附上未來待辦。"""
    target_id, target_type = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    if not todo_date:
        todo_date = now_taipei().strftime('%Y-%m-%d')

    overdue_items = _get_overdue_line_todos(todo_date=todo_date, target_id=target_id)
    today_items = _get_open_line_todos(todo_date=todo_date, target_id=target_id, include_overdue=False)
    future_items = []
    settings = _get_line_todo_reminder_settings(target_id, target_type)
    future_days = _line_todo_get_future_reminder_days(settings)

    if todo_date == now_taipei().strftime('%Y-%m-%d'):
        if _line_todo_is_future_enabled(settings):
            future_items = _get_future_line_todos(start_date=todo_date, target_id=target_id, days=future_days)
        title = f'{_todo_display_md(todo_date)} 待辦清單'
        today_label = '今天'
    else:
        title = f'{_todo_display_md(todo_date)} 待辦清單'
        today_label = _todo_display_md(todo_date)

    return {
        'handled': True,
        'ok': True,
        'reply_text': _format_line_todo_sections(overdue_items, today_items, title, today_label=today_label, future_items=future_items, future_days=future_days),
        'parsed_tag': '查詢待辦',
    }


def send_due_line_todo_reminders_by_settings():
    """
    覆寫每日提醒：今日提醒加入「未來待辦」，且可依設定控制提前幾天開始提醒。
    - 今日提醒：尚未完成 + 今天要做 + 未來 N 天。
    - 明天提醒：維持只提醒明天要做。
    """
    now_dt = now_taipei()
    today = now_dt.strftime('%Y-%m-%d')
    tomorrow = (now_dt.date() + timedelta(days=1)).strftime('%Y-%m-%d')
    current_minutes = now_dt.hour * 60 + now_dt.minute

    targets = _collect_line_todo_targets_with_open_items()
    sent_count = 0
    failed = []
    checked_targets = 0

    for target_id, target_type in targets.items():
        checked_targets += 1
        settings = _get_line_todo_reminder_settings(target_id, target_type)
        future_days = _line_todo_get_future_reminder_days(settings)

        # 今日提醒：尚未完成 + 今日要做 + 未來待辦。
        if settings.get('today_enabled', True):
            today_time = settings.get('today_reminder_time', '08:00')
            if current_minutes >= _time_to_minutes(today_time, default='08:00'):
                due_items = []
                future_items_for_push = []

                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if not todo_date:
                        continue

                    # 尚未完成 + 今天要做：每天各待辦提醒一次。
                    if todo_date <= today:
                        sent_dates = data.get('reminder_sent_dates') or []
                        if today not in sent_dates:
                            due_items.append(doc)
                        continue

                    # 未來待辦：依設定提前 N 天開始顯示；每天最多提醒一次。
                    if _line_todo_is_future_enabled(settings):
                        end_date = (now_dt.date() + timedelta(days=future_days)).strftime('%Y-%m-%d')
                        if today < todo_date <= end_date:
                            sent_dates = data.get('future_reminder_sent_dates') or []
                            if today not in sent_dates:
                                future_items_for_push.append(doc)

                if due_items or future_items_for_push:
                    overdue_items = _sort_line_todo_docs([d for d in due_items if _todo_date_value(d) < today])
                    today_items = _sort_line_todo_docs([d for d in due_items if _todo_date_value(d) == today])
                    future_items = _sort_line_todo_docs(future_items_for_push)
                    body = _format_line_todo_sections(
                        overdue_items,
                        today_items,
                        f'{_todo_display_md(today)} 今日待辦',
                        today_label='今天',
                        future_items=future_items,
                        future_days=future_days,
                    )
                    text = _line_todo_add_opening(settings.get('today_opening_text', ''), body)
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in due_items:
                            doc.reference.update({
                                'reminder_sent_dates': firestore.ArrayUnion([today]),
                                'last_reminded_at': now_taipei().isoformat(),
                            })
                        for doc in future_items_for_push:
                            doc.reference.update({
                                'future_reminder_sent_dates': firestore.ArrayUnion([today]),
                                'last_future_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'today_with_future', 'error': msg})

        # 明天提醒：維持明天要做。
        if settings.get('tomorrow_enabled', True):
            tomorrow_time = settings.get('tomorrow_reminder_time', '23:00')
            if current_minutes >= _time_to_minutes(tomorrow_time, default='23:00'):
                items = []
                for doc in db.collection(LINE_TODO_COLLECTION).stream():
                    data = doc.to_dict() or {}
                    if data.get('status', 'open') != 'open':
                        continue
                    if data.get('line_target_id') != target_id:
                        continue
                    todo_date = (data.get('todo_date') or '').strip()
                    if todo_date != tomorrow:
                        continue
                    sent_dates = data.get('tomorrow_reminder_sent_dates') or []
                    if tomorrow in sent_dates:
                        continue
                    items.append(doc)

                if items:
                    tomorrow_items = _sort_line_todo_docs(items)
                    body = _format_line_todo_sections(
                        [],
                        tomorrow_items,
                        f'{_todo_display_md(tomorrow)} 明天待辦',
                        today_label='明天',
                    )
                    text = _line_todo_add_opening(settings.get('tomorrow_opening_text', ''), body)
                    ok, msg = push_line_text(target_id, text)
                    if ok:
                        sent_count += 1
                        for doc in items:
                            doc.reference.update({
                                'tomorrow_reminder_sent_dates': firestore.ArrayUnion([tomorrow]),
                                'last_tomorrow_reminded_at': now_taipei().isoformat(),
                            })
                    else:
                        failed.append({'target_id': target_id, 'type': 'tomorrow', 'error': msg})

    return {
        'now': now_dt.strftime('%Y-%m-%d %H:%M:%S'),
        'timezone': 'Asia/Taipei',
        'checked_targets': checked_targets,
        'sent_count': sent_count,
        'future_reminder_days_default': LINE_TODO_FUTURE_REMINDER_DAYS_DEFAULT,
        'failed': failed,
    }

# ========= LINE Bot 待辦事項 Patch v12 End =========

# =============================================================================
# 行事曆後台 + LINE 行程卡片 Flex Message Patch v1
# 加入位置：請放在所有既有 route / LINE 待辦 patch 之後，app.run 之前。
# 功能：
# 1. CRM 後台新增 /calendar 行事曆，每 30 分鐘一格。
# 2. CRM 後台新增 /line-card-settings，可控制 LINE 行程卡片顯示內容與顏色。
# 3. LINE 指令支援：#今日行程、#明日行程、#本週行程、#新增行程。
# 4. LINE 回覆改支援 Flex Message / Carousel 卡片，不再只能純文字。
# =============================================================================

from datetime import timedelta
from urllib.parse import quote_plus

CALENDAR_EVENT_COLLECTION = "calendar_events"
LINE_CARD_SETTINGS_COLLECTION = "line_card_settings"

CALENDAR_CATEGORY_OPTIONS = [
    "帶看",
    "回電",
    "開發",
    "簽約",
    "拍照",
    "收服務費",
    "待辦",
    "其他",
]

CALENDAR_CATEGORY_COLOR_MAP = {
    "帶看": "#C9874A",
    "回電": "#E0A800",
    "開發": "#5E9F45",
    "簽約": "#7B61FF",
    "拍照": "#8A63D2",
    "收服務費": "#B36B00",
    "待辦": "#6C757D",
    "其他": "#8B8B8B",
}

DEFAULT_LINE_CARD_SETTINGS = {
    "title_today": "您今天的行程如下：",
    "title_tomorrow": "您明天的行程如下：",
    "title_week": "本週行程如下：",
    "show_phone": True,
    "show_location": True,
    "show_note": True,
    "show_customer": True,
    "show_quick_actions": True,
    "primary_color": "#C9874A",
    "button_color": "#C9874A",
    "calendar_start_hour": 8,
    "calendar_end_hour": 22,
    "line_only_enabled_events": False,
    "quick_actions": ["新增行程", "今日行程", "本週行程", "客戶查詢", "設定提醒"],
    "notify_target_id": "",
}


def calendar_safe_date(value=None):
    """回傳 YYYY-MM-DD；無效時用台北今天。"""
    if value:
        try:
            return datetime.strptime(str(value).strip(), "%Y-%m-%d").strftime("%Y-%m-%d")
        except Exception:
            pass
    return now_taipei().strftime("%Y-%m-%d")


def format_calendar_date_label(date_text):
    """把 2026-06-21 顯示成 6/21(日)，後台與 LINE 卡片共用。"""
    if not date_text:
        return ""
    try:
        dt = datetime.strptime(str(date_text)[:10], "%Y-%m-%d")
        week_map = ["一", "二", "三", "四", "五", "六", "日"]
        return f"{dt.month}/{dt.day}({week_map[dt.weekday()]})"
    except Exception:
        return str(date_text)


def calendar_safe_time(value=None, default="09:00"):
    """回傳 HH:MM；無效時用 default。"""
    raw = (value or "").strip()
    try:
        return datetime.strptime(raw, "%H:%M").strftime("%H:%M")
    except Exception:
        return default


def calendar_time_to_minutes(hhmm: str):
    try:
        h, m = [int(x) for x in str(hhmm).split(":", 1)]
        return h * 60 + m
    except Exception:
        return 0


def build_30_min_slots(start_hour=None, end_hour=None):
    settings = get_line_card_settings()
    if start_hour is None:
        start_hour = int(settings.get("calendar_start_hour", 8) or 8)
    if end_hour is None:
        end_hour = int(settings.get("calendar_end_hour", 22) or 22)

    start_hour = max(0, min(23, start_hour))
    end_hour = max(start_hour, min(23, end_hour))

    slots = []
    for hour in range(start_hour, end_hour + 1):
        slots.append(f"{hour:02d}:00")
        if hour != end_hour:
            slots.append(f"{hour:02d}:30")
    return slots


def next_30_min_time(hhmm: str):
    minutes = calendar_time_to_minutes(hhmm) + 30
    minutes = min(minutes, 23 * 60 + 59)
    return f"{minutes // 60:02d}:{minutes % 60:02d}"


def get_line_card_settings():
    settings = dict(DEFAULT_LINE_CARD_SETTINGS)
    try:
        doc = db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").get()
        if doc.exists:
            data = doc.to_dict() or {}
            settings.update(data)
    except Exception as e:
        print("⚠️ 讀取 LINE 卡片設定失敗：", e)
    return settings


def save_line_card_settings_from_form(form):
    quick_actions_raw = (form.get("quick_actions", "") or "").strip()
    quick_actions = [x.strip() for x in re.split(r"[，,、\n]+", quick_actions_raw) if x.strip()]
    if not quick_actions:
        quick_actions = DEFAULT_LINE_CARD_SETTINGS["quick_actions"]

    def _int_field(name, default):
        try:
            return int(form.get(name, default))
        except Exception:
            return default

    start_hour = max(0, min(23, _int_field("calendar_start_hour", 8)))
    end_hour = max(start_hour, min(23, _int_field("calendar_end_hour", 22)))

    updates = {
        "title_today": (form.get("title_today", "") or "").strip() or DEFAULT_LINE_CARD_SETTINGS["title_today"],
        "title_tomorrow": (form.get("title_tomorrow", "") or "").strip() or DEFAULT_LINE_CARD_SETTINGS["title_tomorrow"],
        "title_week": (form.get("title_week", "") or "").strip() or DEFAULT_LINE_CARD_SETTINGS["title_week"],
        "show_phone": form.get("show_phone") == "on",
        "show_location": form.get("show_location") == "on",
        "show_note": form.get("show_note") == "on",
        "show_customer": form.get("show_customer") == "on",
        "show_quick_actions": form.get("show_quick_actions") == "on",
        "line_only_enabled_events": form.get("line_only_enabled_events") == "on",
        "primary_color": (form.get("primary_color", "") or "").strip() or DEFAULT_LINE_CARD_SETTINGS["primary_color"],
        "button_color": (form.get("button_color", "") or "").strip() or DEFAULT_LINE_CARD_SETTINGS["button_color"],
        "calendar_start_hour": start_hour,
        "calendar_end_hour": end_hour,
        "quick_actions": quick_actions,
        "notify_target_id": (form.get("notify_target_id", "") or "").strip(),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }
    db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set(updates, merge=True)
    return updates


def doc_to_calendar_event(doc):
    data = doc_to_dict(doc)
    data["start_time"] = calendar_safe_time(data.get("start_time"), "09:00")
    data["end_time"] = calendar_safe_time(data.get("end_time"), next_30_min_time(data["start_time"]))
    data["event_date"] = calendar_safe_date(data.get("event_date"))
    data["event_date_label"] = format_calendar_date_label(data.get("event_date"))
    data["category"] = data.get("category") or "其他"
    data["custom_category"] = data.get("custom_category") or ""
    data["display_category"] = data.get("display_category") or (data["custom_category"] if data["category"] == "其他" and data["custom_category"] else data["category"])
    data["visibility"] = data.get("visibility") or "public"
    data["category_color"] = data.get("category_color") or CALENDAR_CATEGORY_COLOR_MAP.get(data["category"], "#8B8B8B")
    return data


def fetch_calendar_events(start_date: str, end_date: str = None, line_only: bool = False):
    """讀取日期區間行程。end_date 未傳時只讀單日。"""
    start_date = calendar_safe_date(start_date)
    end_date = calendar_safe_date(end_date or start_date)
    events = []

    try:
        if start_date == end_date:
            docs = db.collection(CALENDAR_EVENT_COLLECTION).where("event_date", "==", start_date).stream()
        else:
            # Firestore 需要索引時，後台會提示建立；若無索引，下面 except 會走備援。
            docs = db.collection(CALENDAR_EVENT_COLLECTION).where("event_date", ">=", start_date).where("event_date", "<=", end_date).stream()
        for d in docs:
            item = doc_to_calendar_event(d)
            if line_only and not item.get("line_enabled", True):
                continue
            # 後台登入時：個人行程只給本人看，公開行程所有人可看；舊資料無 owner/visibility 時保留顯示。
            if not line_only:
                vis = item.get("visibility") or "public"
                owner = item.get("owner_user_id") or ""
                uid = session.get("user_id", "") if request else ""
                if vis == "personal" and owner and uid and owner != uid:
                    continue
            events.append(item)
    except Exception as e:
        print("⚠️ Firestore 日期區間查詢失敗，改用備援讀取：", e)
        try:
            for d in db.collection(CALENDAR_EVENT_COLLECTION).stream():
                item = doc_to_calendar_event(d)
                if start_date <= item.get("event_date", "") <= end_date:
                    if line_only and not item.get("line_enabled", True):
                        continue
                    if not line_only:
                        vis = item.get("visibility") or "public"
                        owner = item.get("owner_user_id") or ""
                        uid = session.get("user_id", "") if request else ""
                        if vis == "personal" and owner and uid and owner != uid:
                            continue
                    events.append(item)
        except Exception as e2:
            print("❌ 讀取行事曆失敗：", e2)

    events.sort(key=lambda x: (x.get("event_date", ""), x.get("start_time", ""), x.get("created_at", "")))
    return events




def calc_calendar_slot_span(start_time: str, end_time: str):
    """計算一筆行程要橫跨幾個 30 分鐘格。"""
    start_minutes = calendar_time_to_minutes(start_time)
    end_minutes = calendar_time_to_minutes(end_time)
    if end_minutes <= start_minutes:
        end_minutes = start_minutes + 30
    span = (end_minutes - start_minutes + 29) // 30
    return max(1, span)


def build_calendar_slot_cells(events, slots):
    """建立行事曆表格資料，讓長行程可以用 rowspan 橫跨多個 30 分鐘格。"""
    cells = {slot: {"event": None, "span": 1, "skip": False, "extra_events": []} for slot in slots}
    slot_index = {slot: idx for idx, slot in enumerate(slots)}

    # 同一個起始時間可能有多筆；第一筆用 rowspan，其餘顯示在同一張卡下方，避免資料消失。
    grouped = {}
    outside_events = []
    for event in events:
        start = (event.get("start_time") or "").strip()
        if start in slot_index:
            grouped.setdefault(start, []).append(event)
        else:
            outside_events.append(event)

    for start, group in grouped.items():
        main_event = group[0]
        span = calc_calendar_slot_span(main_event.get("start_time"), main_event.get("end_time"))
        start_idx = slot_index[start]
        max_span = max(1, min(span, len(slots) - start_idx))
        cells[start] = {"event": main_event, "span": max_span, "skip": False, "extra_events": group[1:]}

        for offset in range(1, max_span):
            idx = start_idx + offset
            if idx < len(slots):
                covered_slot = slots[idx]
                # 如果該時段本來也有行程，保留到起始卡片的 extra_events，不讓它消失。
                if cells.get(covered_slot, {}).get("event"):
                    cells[start]["extra_events"].append(cells[covered_slot]["event"])
                cells[covered_slot] = {"event": None, "span": 1, "skip": True, "extra_events": []}

    if outside_events and slots:
        first_slot = slots[0]
        if cells[first_slot].get("event"):
            cells[first_slot].setdefault("extra_events", []).extend(outside_events)
        else:
            cells[first_slot] = {"event": outside_events[0], "span": 1, "skip": False, "extra_events": outside_events[1:]}

    return cells

def build_calendar_event_payload(form, existing=None):
    existing = existing or {}
    event_date = calendar_safe_date(form.get("event_date") or existing.get("event_date"))
    start_time = calendar_safe_time(form.get("start_time") or existing.get("start_time"), "09:00")
    end_time = calendar_safe_time(form.get("end_time") or existing.get("end_time"), next_30_min_time(start_time))
    if calendar_time_to_minutes(end_time) <= calendar_time_to_minutes(start_time):
        end_time = next_30_min_time(start_time)

    category = (form.get("category", "") or existing.get("category") or "其他").strip()
    if category not in CALENDAR_CATEGORY_OPTIONS:
        category = "其他"

    custom_category = (form.get("custom_category", "") or existing.get("custom_category", "") or "").strip()
    display_category = custom_category if category == "其他" and custom_category else category
    visibility = (form.get("visibility", "") or existing.get("visibility") or "personal").strip()
    if visibility not in ("personal", "public"):
        visibility = "personal"

    payload = {
        "title": (form.get("title", "") or "").strip(),
        "event_date": event_date,
        "start_time": start_time,
        "end_time": end_time,
        "category": category,
        "custom_category": custom_category,
        "display_category": display_category,
        "visibility": visibility,
        "owner_user_id": existing.get("owner_user_id") or session.get("user_id", ""),
        "owner_user_name": existing.get("owner_user_name") or session.get("user_name", ""),
        "category_color": CALENDAR_CATEGORY_COLOR_MAP.get(category, "#8B8B8B"),
        "related_type": (form.get("related_type", "") or "").strip(),
        "related_id": (form.get("related_id", "") or "").strip(),
        "customer_name": (form.get("customer_name", "") or "").strip(),
        "phone": (form.get("phone", "") or "").strip(),
        "location": (form.get("location", "") or "").strip(),
        "note": (form.get("note", "") or "").strip(),
        "line_enabled": form.get("line_enabled") == "on",
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }
    if not payload["title"]:
        name = payload.get("customer_name") or payload.get("location") or category
        payload["title"] = f"{name} {category}".strip()
    return payload


def calendar_prev_next_dates(selected_date):
    dt = datetime.strptime(calendar_safe_date(selected_date), "%Y-%m-%d").date()
    return {
        "prev_date": (dt - timedelta(days=1)).strftime("%Y-%m-%d"),
        "next_date": (dt + timedelta(days=1)).strftime("%Y-%m-%d"),
        "today": now_taipei().strftime("%Y-%m-%d"),
    }


CALENDAR_PAGE_TEMPLATE = r'''
<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>行事曆｜Team M.E CRM</title>
  <style>
    body { margin:0; background:#f7f3ed; color:#3f3028; font-family:-apple-system,BlinkMacSystemFont,"Microsoft JhengHei","Noto Sans TC",Arial,sans-serif; }
    .topbar { position:sticky; top:0; z-index:10; background:#fffaf3; border-bottom:1px solid #eadfd0; padding:14px 22px; display:flex; gap:12px; align-items:center; justify-content:space-between; }
    .brand { font-weight:800; color:#7A4E2D; letter-spacing:.5px; }
    .nav a, .btn { display:inline-block; text-decoration:none; border:1px solid #d8b28a; border-radius:12px; padding:8px 12px; color:#7A4E2D; background:#fff; font-size:14px; }
    .btn.primary { background:#C9874A; color:#fff; border-color:#C9874A; }
    .btn.danger { color:#a33; border-color:#e4b6b6; }
    .wrap { max-width:1180px; margin:24px auto; padding:0 18px; }
    .panel { background:#fff; border:1px solid #eadfd0; border-radius:20px; box-shadow:0 10px 30px rgba(120,80,40,.08); overflow:hidden; }
    .panel-head { padding:18px 20px; border-bottom:1px solid #f0e6d9; display:flex; flex-wrap:wrap; gap:12px; align-items:center; justify-content:space-between; }
    .date-control { display:flex; gap:8px; align-items:center; flex-wrap:wrap; }
    input[type=date], select, input[type=text], input[type=time], textarea { border:1px solid #dec9b2; border-radius:10px; padding:9px 10px; font-size:15px; background:#fff; box-sizing:border-box; }
    .schedule { display:grid; grid-template-columns:110px 1fr; }
    .slot-time { padding:16px 18px; border-bottom:1px solid #f3eadf; background:#fffaf6; color:#8b6b4f; font-weight:700; }
    .slot-events { padding:10px 14px; border-bottom:1px solid #f3eadf; min-height:52px; }
    .empty { color:#c7b7a8; font-size:14px; padding:8px 0; }
    .event-card { border:1px solid #eadfd0; border-left:7px solid var(--cat,#C9874A); border-radius:14px; padding:12px 14px; margin:4px 0 8px; background:#fffdf9; display:flex; justify-content:space-between; gap:14px; align-items:flex-start; }
    .event-title { font-size:17px; font-weight:800; margin-bottom:6px; color:#3d2d23; }
    .event-meta { font-size:14px; color:#6f6258; line-height:1.75; }
    .pill { display:inline-block; padding:3px 9px; border-radius:999px; background:#f4e3d2; color:#7A4E2D; font-size:13px; font-weight:700; margin-left:8px; }
    .actions { display:flex; gap:7px; flex-wrap:wrap; justify-content:flex-end; min-width:150px; }
    .msg { margin:0 0 14px; padding:12px 14px; border-radius:12px; background:#fff6e8; border:1px solid #f0d7b5; color:#7A4E2D; }
    .preview-grid { display:grid; grid-template-columns:repeat(auto-fit,minmax(220px,1fr)); gap:14px; margin-top:18px; }
    .preview-card { background:#fff; border:1px solid #eadfd0; border-radius:20px; padding:16px; box-shadow:0 8px 20px rgba(120,80,40,.08); }
    .preview-time { color:#C9874A; font-weight:900; }
    @media (max-width:720px){ .schedule{grid-template-columns:80px 1fr;} .event-card{display:block;} .actions{justify-content:flex-start; margin-top:10px;} }
  </style>
</head>
<body>
  <div class="topbar">
    <div class="brand">Team M.E CRM｜行事曆</div>
    <div class="nav">
      <a href="{{ url_for('buyers') }}">客需</a>
      <a href="{{ url_for('sellers') }}">委託</a>
      <a href="{{ url_for('calendar_page') }}">行事曆</a>
      <a href="{{ url_for('line_card_settings') }}">LINE卡片設定</a>
    </div>
  </div>

  <div class="wrap">
    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}
        {% for category, message in messages %}<div class="msg">{{ message }}</div>{% endfor %}
      {% endif %}
    {% endwith %}

    <div class="panel">
      <div class="panel-head">
        <div>
          <h2 style="margin:0 0 6px;">{{ selected_date }} 行程</h2>
          <div style="color:#8b6b4f;font-size:14px;">每 30 分鐘一格，可新增帶看、回電、開發、簽約與提醒。</div>
        </div>
        <div class="date-control">
          <a class="btn" href="{{ url_for('calendar_page', date=prev_date) }}">前一天</a>
          <form method="get" action="{{ url_for('calendar_page') }}" style="display:flex;gap:8px;align-items:center;">
            <input type="date" name="date" value="{{ selected_date }}">
            <button class="btn" type="submit">切換日期</button>
          </form>
          <a class="btn" href="{{ url_for('calendar_page', date=today) }}">今天</a>
          <a class="btn" href="{{ url_for('calendar_page', date=next_date) }}">後一天</a>
          <a class="btn primary" href="{{ url_for('calendar_new', date=selected_date) }}">＋新增行程</a>
        </div>
      </div>

      <div class="schedule">
        {% for slot in slots %}
          <div class="slot-time">{{ slot }}</div>
          <div class="slot-events">
            {% set items = event_map.get(slot, []) %}
            {% if not items %}
              <div class="empty">空</div>
            {% endif %}
            {% for e in items %}
              <div class="event-card" style="--cat: {{ e.category_color }};">
                <div>
                  <div class="event-title">{{ e.title }} <span class="pill">{{ e.category }}</span></div>
                  <div class="event-meta">
                    {{ e.start_time }} - {{ e.end_time }}<br>
                    {% if e.customer_name %}客戶：{{ e.customer_name }}{% if e.phone %}｜{{ e.phone }}{% endif %}<br>{% endif %}
                    {% if e.location %}地點：{{ e.location }}<br>{% endif %}
                    {% if e.note %}備註：{{ e.note }}{% endif %}
                  </div>
                </div>
                <div class="actions">
                  <a class="btn" href="{{ url_for('calendar_edit', event_id=e.id) }}">編輯</a>
                  <form method="post" action="{{ url_for('calendar_delete', event_id=e.id) }}" onsubmit="return confirm('確定刪除此行程？');">
                    <button class="btn danger" type="submit">刪除</button>
                  </form>
                </div>
              </div>
            {% endfor %}
          </div>
        {% endfor %}
      </div>
    </div>

    <div class="preview-grid">
      <div class="preview-card">
        <div class="preview-time">LINE 指令</div>
        <div style="line-height:1.9;margin-top:8px;">
          #今日行程<br>
          #明日行程<br>
          #本週行程<br>
          #新增行程<br>
          日期: 今天<br>
          時間: 10:00<br>
          類型: 帶看<br>
          標題: 童先生看農舍
        </div>
      </div>
      <div class="preview-card">
        <div class="preview-time">LINE 卡片</div>
        <div style="line-height:1.9;margin-top:8px;">
          會依照行程數量自動產生左右滑動卡片。<br>
          可在「LINE卡片設定」控制電話、地點、備註是否顯示。
        </div>
      </div>
    </div>
  </div>
</body>
</html>
'''


CALENDAR_FORM_TEMPLATE = r'''
<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{{ '編輯' if event else '新增' }}行程｜Team M.E CRM</title>
  <style>
    body { margin:0; background:#f7f3ed; color:#3f3028; font-family:-apple-system,BlinkMacSystemFont,"Microsoft JhengHei","Noto Sans TC",Arial,sans-serif; }
    .topbar { background:#fffaf3; border-bottom:1px solid #eadfd0; padding:14px 22px; display:flex; justify-content:space-between; align-items:center; }
    .brand { font-weight:800; color:#7A4E2D; }
    .nav a, .btn { display:inline-block; text-decoration:none; border:1px solid #d8b28a; border-radius:12px; padding:9px 13px; color:#7A4E2D; background:#fff; font-size:14px; }
    .btn.primary { background:#C9874A; color:#fff; border-color:#C9874A; }
    .wrap { max-width:760px; margin:24px auto; padding:0 18px; }
    .panel { background:#fff; border:1px solid #eadfd0; border-radius:22px; padding:24px; box-shadow:0 10px 30px rgba(120,80,40,.08); }
    .grid { display:grid; grid-template-columns:1fr 1fr; gap:16px; }
    label { display:block; font-weight:800; margin:0 0 7px; color:#5a4032; }
    input, select, textarea { width:100%; border:1px solid #dec9b2; border-radius:12px; padding:10px 12px; font-size:15px; background:#fff; box-sizing:border-box; }
    textarea { min-height:110px; resize:vertical; }
    .full { grid-column:1 / -1; }
    .hint { color:#8b6b4f; font-size:13px; margin-top:5px; }
    .actions { margin-top:22px; display:flex; gap:10px; justify-content:flex-end; flex-wrap:wrap; }
    .msg { margin:0 0 14px; padding:12px 14px; border-radius:12px; background:#fff6e8; border:1px solid #f0d7b5; color:#7A4E2D; }
    @media (max-width:720px){ .grid{grid-template-columns:1fr;} }
  </style>
</head>
<body>
  <div class="topbar">
    <div class="brand">Team M.E CRM｜{{ '編輯' if event else '新增' }}行程</div>
    <div class="nav"><a href="{{ url_for('calendar_page', date=(event.event_date if event else selected_date)) }}">返回行事曆</a></div>
  </div>
  <div class="wrap">
    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}{% for category, message in messages %}<div class="msg">{{ message }}</div>{% endfor %}{% endif %}
    {% endwith %}
    <form class="panel" method="post">
      <h2 style="margin-top:0;">{{ '編輯' if event else '新增' }}行程</h2>
      <div class="grid">
        <div>
          <label>日期</label>
          <input type="date" name="event_date" value="{{ event.event_date if event else selected_date }}" required>
        </div>
        <div>
          <label>類型</label>
          <select name="category">
            {% for c in category_options %}<option value="{{ c }}" {% if event and event.category == c %}selected{% endif %}>{{ c }}</option>{% endfor %}
          </select>
        </div>
        <div>
          <label>開始時間</label>
          <select name="start_time">
            {% for s in slots %}<option value="{{ s }}" {% if (event and event.start_time == s) or ((not event) and default_start == s) %}selected{% endif %}>{{ s }}</option>{% endfor %}
          </select>
        </div>
        <div>
          <label>結束時間</label>
          <select name="end_time">
            {% for s in slots %}<option value="{{ s }}" {% if (event and event.end_time == s) or ((not event) and default_end == s) %}selected{% endif %}>{{ s }}</option>{% endfor %}
          </select>
        </div>
        <div class="full">
          <label>標題</label>
          <input type="text" name="title" value="{{ event.title if event else '' }}" placeholder="例如：童先生 看農舍 / 李太太 回電追蹤">
        </div>
        <div>
          <label>客戶姓名</label>
          <input type="text" name="customer_name" value="{{ event.customer_name if event else '' }}">
        </div>
        <div>
          <label>電話</label>
          <input type="text" name="phone" value="{{ event.phone if event else '' }}">
        </div>
        <div>
          <label>關聯類型</label>
          <select name="related_type">
            {% set rt = event.related_type if event else '' %}
            <option value="" {% if not rt %}selected{% endif %}>不關聯</option>
            <option value="buyer" {% if rt == 'buyer' %}selected{% endif %}>買方 / 客需</option>
            <option value="seller" {% if rt == 'seller' %}selected{% endif %}>賣方 / 委託</option>
            <option value="development" {% if rt == 'development' %}selected{% endif %}>開發</option>
            <option value="todo" {% if rt == 'todo' %}selected{% endif %}>待辦</option>
          </select>
        </div>
        <div>
          <label>關聯ID</label>
          <input type="text" name="related_id" value="{{ event.related_id if event else '' }}" placeholder="可先空白">
        </div>
        <div class="full">
          <label>地點</label>
          <input type="text" name="location" value="{{ event.location if event else '' }}" placeholder="例如：清水、梧棲交界 / 沙鹿中山路">
        </div>
        <div class="full">
          <label>備註</label>
          <textarea name="note" placeholder="客戶需求、下一步、注意事項">{{ event.note if event else '' }}</textarea>
        </div>
        <div class="full">
          <label style="display:flex;align-items:center;gap:8px;"><input type="checkbox" name="line_enabled" style="width:auto;" {% if (event and event.line_enabled) or not event %}checked{% endif %}> LINE 查詢行程時顯示這筆</label>
          <div class="hint">若關閉，且 LINE 卡片設定選擇「只顯示已啟用」，這筆就不會出現在 LINE 卡片。</div>
        </div>
      </div>
      <div class="actions">
        <a class="btn" href="{{ url_for('calendar_page', date=(event.event_date if event else selected_date)) }}">取消</a>
        <button class="btn primary" type="submit">儲存</button>
      </div>
    </form>
  </div>
</body>
</html>
'''


LINE_CARD_SETTINGS_TEMPLATE = r'''
<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>LINE卡片設定｜Team M.E CRM</title>
  <style>
    body { margin:0; background:#f7f3ed; color:#3f3028; font-family:-apple-system,BlinkMacSystemFont,"Microsoft JhengHei","Noto Sans TC",Arial,sans-serif; }
    .topbar { background:#fffaf3; border-bottom:1px solid #eadfd0; padding:14px 22px; display:flex; justify-content:space-between; align-items:center; }
    .brand { font-weight:800; color:#7A4E2D; }
    .nav a, .btn { display:inline-block; text-decoration:none; border:1px solid #d8b28a; border-radius:12px; padding:9px 13px; color:#7A4E2D; background:#fff; font-size:14px; }
    .btn.primary { background:#C9874A; color:#fff; border-color:#C9874A; }
    .wrap { max-width:850px; margin:24px auto; padding:0 18px; }
    .panel { background:#fff; border:1px solid #eadfd0; border-radius:22px; padding:24px; box-shadow:0 10px 30px rgba(120,80,40,.08); }
    .grid { display:grid; grid-template-columns:1fr 1fr; gap:16px; }
    .full { grid-column:1 / -1; }
    label { display:block; font-weight:800; margin:0 0 7px; color:#5a4032; }
    input, textarea { width:100%; border:1px solid #dec9b2; border-radius:12px; padding:10px 12px; font-size:15px; background:#fff; box-sizing:border-box; }
    textarea { min-height:105px; resize:vertical; }
    .check { display:flex; gap:8px; align-items:center; margin:8px 0; font-weight:700; }
    .check input { width:auto; }
    .hint { color:#8b6b4f; font-size:13px; margin-top:5px; line-height:1.7; }
    .actions { margin-top:22px; display:flex; gap:10px; justify-content:flex-end; }
    .msg { margin:0 0 14px; padding:12px 14px; border-radius:12px; background:#fff6e8; border:1px solid #f0d7b5; color:#7A4E2D; }
    @media (max-width:720px){ .grid{grid-template-columns:1fr;} }
  </style>
</head>
<body>
  <div class="topbar">
    <div class="brand">Team M.E CRM｜LINE卡片設定</div>
    <div class="nav"><a href="{{ url_for('calendar_page') }}">返回行事曆</a></div>
  </div>
  <div class="wrap">
    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}{% for category, message in messages %}<div class="msg">{{ message }}</div>{% endfor %}{% endif %}
    {% endwith %}
    <form class="panel" method="post">
      <h2 style="margin-top:0;">LINE 行程卡片設定</h2>
      <div class="grid">
        <div class="full"><label>今日行程標題</label><input name="title_today" value="{{ settings.title_today }}"></div>
        <div class="full"><label>明日行程標題</label><input name="title_tomorrow" value="{{ settings.title_tomorrow }}"></div>
        <div class="full"><label>本週行程標題</label><input name="title_week" value="{{ settings.title_week }}"></div>
        <div><label>主色</label><input name="primary_color" value="{{ settings.primary_color }}" placeholder="#C9874A"></div>
        <div><label>按鈕色</label><input name="button_color" value="{{ settings.button_color }}" placeholder="#C9874A"></div>
        <div><label>行事曆開始小時</label><input name="calendar_start_hour" value="{{ settings.calendar_start_hour }}" placeholder="8"></div>
        <div><label>行事曆結束小時</label><input name="calendar_end_hour" value="{{ settings.calendar_end_hour }}" placeholder="22"></div>
        <div class="full">
          <label>LINE 卡片顯示內容</label>
          <label class="check"><input type="checkbox" name="show_customer" {% if settings.show_customer %}checked{% endif %}> 顯示客戶姓名</label>
          <label class="check"><input type="checkbox" name="show_phone" {% if settings.show_phone %}checked{% endif %}> 顯示電話</label>
          <label class="check"><input type="checkbox" name="show_location" {% if settings.show_location %}checked{% endif %}> 顯示地點</label>
          <label class="check"><input type="checkbox" name="show_note" {% if settings.show_note %}checked{% endif %}> 顯示備註</label>
          <label class="check"><input type="checkbox" name="show_quick_actions" {% if settings.show_quick_actions %}checked{% endif %}> 顯示快速操作</label>
          <label class="check"><input type="checkbox" name="line_only_enabled_events" {% if settings.line_only_enabled_events %}checked{% endif %}> LINE 只顯示有勾選「LINE查詢顯示」的行程</label>
        </div>
        <div class="full">
          <label>快速操作按鈕文字</label>
          <textarea name="quick_actions">{{ quick_actions_text }}</textarea>
          <div class="hint">一行一個或用逗號分隔，例如：新增行程、今日行程、本週行程、客戶查詢、設定提醒。</div>
        </div>
      </div>
      <div class="actions">
        <a class="btn" href="{{ url_for('calendar_page') }}">取消</a>
        <button class="btn primary" type="submit">儲存設定</button>
      </div>
    </form>
  </div>
</body>
</html>
'''


@app.route("/calendar")
@login_required
def calendar_page():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    selected_date_label = format_calendar_date_label(selected_date)
    events = fetch_calendar_events(selected_date)
    slots = build_30_min_slots()
    slot_cells = build_calendar_slot_cells(events, slots)

    # 舊版 calendar.html 若仍使用 event_map，也保留傳入，避免模板尚未覆蓋時出錯。
    event_map = {}
    for e in events:
        event_map.setdefault(e.get("start_time"), []).append(e)

    dates = calendar_prev_next_dates(selected_date)
    return render_template(
        "calendar.html",
        selected_date=selected_date,
        selected_date_label=selected_date_label,
        slots=slots,
        slot_cells=slot_cells,
        event_map=event_map,
        events=events,
        category_options=CALENDAR_CATEGORY_OPTIONS,
        **dates,
    )


@app.route("/calendar/new", methods=["GET", "POST"])
@login_required
def calendar_new():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    default_start = calendar_safe_time(request.args.get("start", ""), "09:00")
    default_end = next_30_min_time(default_start)

    if request.method == "POST":
        payload = build_calendar_event_payload(request.form)
        payload.update({
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id", ""),
            "created_by_name": session.get("user_name", ""),
        })
        db.collection(CALENDAR_EVENT_COLLECTION).add(payload)
        flash("已新增行程", "success")
        return redirect(url_for("calendar_page", date=payload["event_date"]))

    return render_template_string(
        CALENDAR_FORM_TEMPLATE,
        event=None,
        selected_date=selected_date,
        default_start=default_start,
        default_end=default_end,
        slots=build_30_min_slots(),
        category_options=CALENDAR_CATEGORY_OPTIONS,
    )


@app.route("/calendar/<event_id>/edit", methods=["GET", "POST"])
@login_required
def calendar_edit(event_id):
    doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆行程", "danger")
        return redirect(url_for("calendar_page"))

    event = doc_to_calendar_event(doc)

    if request.method == "POST":
        payload = build_calendar_event_payload(request.form, existing=event)
        doc_ref.update(payload)
        flash("已更新行程", "success")
        return redirect(url_for("calendar_page", date=payload["event_date"]))

    return render_template_string(
        CALENDAR_FORM_TEMPLATE,
        event=event,
        selected_date=event.get("event_date"),
        default_start=event.get("start_time", "09:00"),
        default_end=event.get("end_time", "09:30"),
        slots=build_30_min_slots(),
        category_options=CALENDAR_CATEGORY_OPTIONS,
    )


@app.route("/calendar/<event_id>/delete", methods=["POST"])
@login_required
def calendar_delete(event_id):
    doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id)
    doc = doc_ref.get()
    selected_date = now_taipei().strftime("%Y-%m-%d")
    if doc.exists:
        selected_date = (doc.to_dict() or {}).get("event_date", selected_date)
        doc_ref.delete()
        flash("已刪除行程", "info")
    else:
        flash("找不到這筆行程", "danger")
    return redirect(url_for("calendar_page", date=selected_date))


@app.route("/line-card-settings", methods=["GET", "POST"])
@login_required
def line_card_settings():
    if request.method == "POST":
        save_line_card_settings_from_form(request.form)
        flash("LINE 卡片設定已更新", "success")
        return redirect(url_for("line_card_settings"))

    settings = get_line_card_settings()
    return render_template_string(
        LINE_CARD_SETTINGS_TEMPLATE,
        settings=settings,
        quick_actions_text="\n".join(settings.get("quick_actions") or []),
    )


# 給外部或測試用：JSON 版今日 / 指定日期行程。
@app.route("/calendar/api/events")
@login_required
def calendar_api_events():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    events = fetch_calendar_events(selected_date)
    return {"ok": True, "date": selected_date, "events": events}, 200


# =============================================================================
# LINE Flex Message：回覆函式與卡片產生器
# =============================================================================

def reply_line_messages(reply_token: str, messages: list):
    if not LINE_CHANNEL_ACCESS_TOKEN or not reply_token:
        return None
    payload = {
        "replyToken": reply_token,
        "messages": messages,
    }
    try:
        import requests
        res = requests.post(
            "https://api.line.me/v2/bot/message/reply",
            headers=line_api_headers(),
            json=payload,
            timeout=8,
        )
        print("LINE reply status:", res.status_code, res.text[:300])
        data = {}
        try:
            data = res.json() if res.text else {}
        except Exception:
            data = {}
        return {
            "status_code": res.status_code,
            "data": data,
            "sent_messages": (data or {}).get("sentMessages", []),
        }
    except Exception as e:
        print("⚠️ LINE flex reply 發生錯誤：", e)
        return None


def reply_line_flex(reply_token: str, alt_text: str, contents: dict, quick_reply_items=None):
    message = {
        "type": "flex",
        "altText": (alt_text or "CRM 行程卡片")[:400],
        "contents": contents,
    }
    if quick_reply_items:
        message["quickReply"] = {"items": quick_reply_items[:13]}
    return reply_line_messages(reply_token, [message])


def line_truncate(text, max_len=80):
    text = str(text or "").strip()
    if len(text) <= max_len:
        return text
    return text[:max_len - 1] + "…"


def flex_text(text, size="sm", color="#333333", weight=None, wrap=True, margin=None):
    item = {"type": "text", "text": str(text or "-"), "size": size, "color": color, "wrap": wrap}
    if weight:
        item["weight"] = weight
    if margin:
        item["margin"] = margin
    return item


def flex_info_row(label, value):
    return {
        "type": "box",
        "layout": "baseline",
        "spacing": "sm",
        "contents": [
            {"type": "text", "text": label, "size": "xs", "color": "#999999", "flex": 2},
            {"type": "text", "text": line_truncate(value, 90) or "-", "size": "sm", "color": "#333333", "flex": 5, "wrap": True},
        ],
    }




def get_app_public_base_url():
    """取得 LINE 卡片按鈕要用的公開網址。
    優先使用 Render / ngrok 設定的 APP_BASE_URL；沒有設定時，後台預覽可用目前 request 的網址。
    """
    base = os.environ.get("APP_BASE_URL", "").strip().rstrip("/")
    if base:
        return base
    try:
        return request.url_root.strip().rstrip("/")
    except Exception:
        return ""


def build_app_url(path: str) -> str:
    base = get_app_public_base_url()
    if not base or not path:
        return ""
    path = str(path)
    if not path.startswith("/"):
        path = "/" + path
    return base + path


def calendar_event_edit_path(event):
    event_id = (event or {}).get("id", "")
    if not event_id:
        return ""
    return f"/calendar/{event_id}/edit"


def calendar_related_edit_path_and_label(event):
    """回傳行程關聯的後台編輯網址與按鈕文字。
    若有 related_type + related_id：直接進入對應編輯頁。
    若只有電話：先帶到客需搜尋頁，避免沒有 ID 時無法直接判斷是哪一筆。
    """
    event = event or {}
    related_type = (event.get("related_type") or "").strip()
    related_id = (event.get("related_id") or "").strip()
    phone = (event.get("phone") or "").strip()

    if related_type in ("buyer", "buyers", "客需", "買方"):
        if related_id:
            return f"/buyers/{related_id}/edit", "編輯客需"
        if phone:
            return f"/buyers?q={quote_plus(phone)}", "搜尋客需"

    if related_type in ("seller", "sellers", "委託", "賣方", "屋主"):
        if related_id:
            return f"/sellers/{related_id}/edit", "編輯委託"
        if phone:
            return f"/sellers?q={quote_plus(phone)}", "搜尋委託"

    if related_type in ("development", "developments", "開發"):
        if related_id:
            return f"/developments/{related_id}/edit", "編輯開發"
        if phone:
            return f"/developments?q={quote_plus(phone)}", "搜尋開發"

    # 沒有設定關聯類型時，用電話先帶到客需搜尋，避免錯連。
    if phone:
        return f"/buyers?q={quote_plus(phone)}", "搜尋客需"

    return "", ""


def calendar_related_button_label(event):
    return calendar_related_edit_path_and_label(event)[1]


def calendar_related_edit_url(event):
    path, _label = calendar_related_edit_path_and_label(event)
    return build_app_url(path)


def calendar_event_edit_url(event):
    return build_app_url(calendar_event_edit_path(event))

def build_calendar_event_bubble(event, settings=None):
    settings = settings or get_line_card_settings()
    category = event.get("display_category") or event.get("category") or "行程"
    category_color = event.get("category_color") or CALENDAR_CATEGORY_COLOR_MAP.get(category, settings.get("primary_color", "#C9874A"))
    title = event.get("title") or f"{category}行程"
    raw_date_text = event.get("event_date", "") or ""
    date_text = format_calendar_date_label(raw_date_text) or "-"
    time_text = f"{date_text} {event.get('start_time', '')} - {event.get('end_time', '')}"

    body_contents = [
        {
            "type": "box",
            "layout": "horizontal",
            "contents": [
                {"type": "text", "text": f"📅 {date_text}", "size": "sm", "color": category_color, "weight": "bold", "flex": 4},
                {"type": "text", "text": category, "size": "xs", "color": "#3F3028", "align": "center", "gravity": "center", "flex": 2, "weight": "bold"},
            ],
        },
        flex_text(f"🕒 {event.get('start_time', '')} - {event.get('end_time', '')}", size="sm", color="#666666", weight="bold", margin="sm"),
        flex_text(line_truncate(title, 48), size="lg", color="#222222", weight="bold", margin="md"),
        {"type": "separator", "margin": "md"},
        {"type": "box", "layout": "vertical", "spacing": "sm", "margin": "md", "contents": []},
    ]

    info_box = body_contents[-1]
    if settings.get("show_customer", True):
        info_box["contents"].append(flex_info_row("客戶", event.get("customer_name") or "-"))
    if settings.get("show_phone", True):
        info_box["contents"].append(flex_info_row("電話", event.get("phone") or "-"))
    if settings.get("show_location", True):
        info_box["contents"].append(flex_info_row("地點", event.get("location") or "-"))
    if settings.get("show_note", True):
        info_box["contents"].append(flex_info_row("備註", event.get("note") or "-"))

    if not info_box["contents"]:
        info_box["contents"].append(flex_text(time_text, size="sm", color="#666666"))

    footer_contents = []

    event_edit_url = calendar_event_edit_url(event)
    if event_edit_url:
        footer_contents.append({
            "type": "button",
            "style": "primary",
            "color": settings.get("button_color", "#C9874A"),
            "height": "sm",
            "action": {
                "type": "uri",
                "label": "編輯行程",
                "uri": event_edit_url,
            },
        })

    related_url = calendar_related_edit_url(event)
    related_label = calendar_related_button_label(event)
    if related_url and related_label:
        footer_contents.append({
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "uri",
                "label": related_label,
                "uri": related_url,
            },
        })

    if event.get("location"):
        footer_contents.append({
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "uri",
                "label": "查看地圖",
                "uri": "https://www.google.com/maps/search/?api=1&query=" + quote_plus(event.get("location", "")),
            },
        })

    if event.get("phone"):
        footer_contents.append({
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "uri",
                "label": "撥打電話",
                "uri": "tel:" + re.sub(r"[^0-9+]", "", event.get("phone", "")),
            },
        })

    if not footer_contents:
        footer_contents.append({
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "message",
                "label": "本週行程",
                "text": "#本週行程",
            },
        })

    return {
        "type": "bubble",
        "size": "mega",
        "body": {
            "type": "box",
            "layout": "vertical",
            "spacing": "md",
            "contents": body_contents,
        },
        "footer": {
            "type": "box",
            "layout": "vertical",
            "spacing": "sm",
            "contents": footer_contents[:4],
        },
        "styles": {
            "footer": {"separator": True},
        },
    }


def build_calendar_empty_bubble(title, date_text, settings=None):
    settings = settings or get_line_card_settings()
    return {
        "type": "bubble",
        "size": "mega",
        "body": {
            "type": "box",
            "layout": "vertical",
            "spacing": "md",
            "contents": [
                flex_text("📅 " + title, size="lg", color="#222222", weight="bold"),
                flex_text(date_text, size="sm", color="#8b6b4f"),
                {"type": "separator", "margin": "md"},
                flex_text("目前沒有行程，可以到後台新增，或直接用 #新增行程 建立。", size="sm", color="#666666", margin="md"),
            ],
        },
        "footer": {
            "type": "box",
            "layout": "vertical",
            "spacing": "sm",
            "contents": [
                {
                    "type": "button",
                    "style": "primary",
                    "color": settings.get("button_color", "#C9874A"),
                    "action": {"type": "message", "label": "新增行程格式", "text": "#新增行程"},
                }
            ],
        },
    }


def build_calendar_carousel(events, title="行程", date_text="", settings=None):
    settings = settings or get_line_card_settings()
    bubbles = []
    if not events:
        bubbles.append(build_calendar_empty_bubble(title, date_text, settings))
    else:
        for event in events[:12]:
            bubbles.append(build_calendar_event_bubble(event, settings))

    if len(bubbles) == 1:
        return bubbles[0]
    return {"type": "carousel", "contents": bubbles}


def build_calendar_quick_reply(settings=None):
    settings = settings or get_line_card_settings()
    if not settings.get("show_quick_actions", True):
        return []

    label_to_text = {
        "新增行程": "#新增行程",
        "今日行程": "#今日行程",
        "今天行程": "#今日行程",
        "明日行程": "#明日行程",
        "明天行程": "#明日行程",
        "本週行程": "#本週行程",
        "客戶查詢": "#查詢紀錄",
        "設定提醒": "#設定待辦提醒",
    }

    items = []
    for label in settings.get("quick_actions") or []:
        text = label_to_text.get(label, f"#{label}")
        items.append({
            "type": "action",
            "action": {
                "type": "message",
                "label": str(label)[:20],
                "text": text,
            },
        })
    return items[:13]


def build_calendar_reply_for_range(start_date, end_date=None, mode="today"):
    settings = get_line_card_settings()
    line_only = bool(settings.get("line_only_enabled_events", False))
    events = fetch_calendar_events(start_date, end_date or start_date, line_only=line_only)

    if mode == "week":
        title = settings.get("title_week") or DEFAULT_LINE_CARD_SETTINGS["title_week"]
        date_text = f"{format_calendar_date_label(start_date)} ~ {format_calendar_date_label(end_date)}"
    elif mode == "tomorrow":
        title = settings.get("title_tomorrow") or DEFAULT_LINE_CARD_SETTINGS["title_tomorrow"]
        date_text = format_calendar_date_label(start_date)
    else:
        title = settings.get("title_today") or DEFAULT_LINE_CARD_SETTINGS["title_today"]
        date_text = format_calendar_date_label(start_date)

    flex = build_calendar_carousel(events, title=title, date_text=date_text, settings=settings)
    alt = f"{title} {len(events)} 筆"
    return {
        "handled": True,
        "ok": True,
        "reply_text": alt,
        "reply_flex": flex,
        "reply_quick_reply": build_calendar_quick_reply(settings),
        "parsed_tag": "行事曆",
    }


def parse_line_calendar_create_fields(text):
    fields = {}
    lines = [ln.strip() for ln in (text or "").splitlines() if ln.strip()]
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.+)$", line)
        if not m:
            # 沒寫 key 的行，先當備註追加。
            fields["note"] = (fields.get("note", "") + "\n" + line).strip()
            continue
        key = normalize_line_key(m.group(1))
        raw_key = (m.group(1) or "").strip().replace(" ", "")
        value = m.group(2).strip()
        if raw_key in ("日期", "行程日期"):
            fields["event_date_raw"] = value
        elif raw_key in ("時間", "開始時間", "時段"):
            fields["time_raw"] = value
        elif raw_key in ("結束時間",):
            fields["end_time"] = value
        elif raw_key in ("類型", "分類"):
            fields["category"] = value
        elif raw_key in ("標題", "行程", "事項"):
            fields["title"] = value
        elif raw_key in ("客戶", "姓名", "客戶姓名"):
            fields["customer_name"] = value
        elif raw_key in ("電話", "手機"):
            fields["phone"] = value
        elif raw_key in ("地點", "地址", "位置"):
            fields["location"] = value
        elif raw_key in ("備註", "內容", "說明"):
            fields["note"] = value
        elif key in ("record_id", "related_id"):
            fields["related_id"] = value
        else:
            fields[raw_key] = value
    return fields


def parse_calendar_date_word(raw):
    raw = (raw or "").strip()
    today = now_taipei().date()
    if raw in ("", "今天", "今日"):
        return today.strftime("%Y-%m-%d")
    if raw in ("明天", "明日"):
        return (today + timedelta(days=1)).strftime("%Y-%m-%d")
    if raw in ("後天",):
        return (today + timedelta(days=2)).strftime("%Y-%m-%d")
    return calendar_safe_date(raw)


def parse_calendar_time_range(raw_time, raw_end=None):
    text = (raw_time or "").strip()
    # 支援 10:00-10:30 / 1000-1030 / 10點 / 10點30
    found = re.findall(r"(\d{1,2})(?:[:：點](\d{1,2}))?", text)
    if found:
        h1, m1 = found[0]
        start = f"{int(h1):02d}:{int(m1 or 0):02d}"
        if len(found) >= 2:
            h2, m2 = found[1]
            end = f"{int(h2):02d}:{int(m2 or 0):02d}"
        else:
            end = calendar_safe_time(raw_end, next_30_min_time(start))
        return calendar_safe_time(start, "09:00"), calendar_safe_time(end, next_30_min_time(start))
    start = calendar_safe_time(text, "09:00")
    end = calendar_safe_time(raw_end, next_30_min_time(start))
    return start, end


def process_line_calendar_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    raw_text = (message.get("text") or "").strip()
    normalized = raw_text.replace("＃", "#").strip()
    today = now_taipei().date()

    if normalized in ("#今日行程", "#今天行程", "今日行程", "今天行程"):
        d = today.strftime("%Y-%m-%d")
        return build_calendar_reply_for_range(d, mode="today")

    if normalized in ("#明日行程", "#明天行程", "明日行程", "明天行程"):
        d = (today + timedelta(days=1)).strftime("%Y-%m-%d")
        return build_calendar_reply_for_range(d, mode="tomorrow")

    if normalized in ("#本週行程", "本週行程"):
        # 以今天起算未來 7 天，符合房仲使用情境。
        start = today.strftime("%Y-%m-%d")
        end = (today + timedelta(days=6)).strftime("%Y-%m-%d")
        return build_calendar_reply_for_range(start, end, mode="week")

    if normalized.startswith("#新增行程"):
        fields = parse_line_calendar_create_fields(normalized)
        if len(normalized.splitlines()) == 1:
            example = (
                "新增行程格式：\n"
                "#新增行程\n"
                "日期: 今天\n"
                "時間: 10:00-10:30\n"
                "類型: 帶看\n"
                "標題: 童先生看農舍\n"
                "客戶: 童先生\n"
                "電話: 0921-123-456\n"
                "地點: 清水、梧棲交界\n"
                "備註: 退休夫妻，想看農舍、有空地"
            )
            return {"handled": True, "ok": True, "reply_text": example, "parsed_tag": "新增行程格式"}

        event_date = parse_calendar_date_word(fields.get("event_date_raw", "今天"))
        start_time, end_time = parse_calendar_time_range(fields.get("time_raw", "09:00"), fields.get("end_time", ""))
        category = fields.get("category", "其他")
        if category not in CALENDAR_CATEGORY_OPTIONS:
            category = "其他"

        title = fields.get("title", "").strip()
        if not title:
            title = f"{fields.get('customer_name', '')} {category}".strip() or category

        payload = {
            "title": title,
            "event_date": event_date,
            "start_time": start_time,
            "end_time": end_time,
            "category": category,
            "category_color": CALENDAR_CATEGORY_COLOR_MAP.get(category, "#8B8B8B"),
            "related_type": fields.get("related_type", ""),
            "related_id": fields.get("related_id", ""),
            "customer_name": fields.get("customer_name", ""),
            "phone": fields.get("phone", ""),
            "location": fields.get("location", ""),
            "note": fields.get("note", ""),
            "line_enabled": True,
            "created_at": now_taipei().isoformat(),
            "created_by_id": "line_bot",
            "created_by_name": "LINE Bot",
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": "line_bot",
            "updated_by_name": "LINE Bot",
        }
        doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document()
        doc_ref.set(payload)
        payload["id"] = doc_ref.id
        event_item = dict(payload)
        settings = get_line_card_settings()
        flex = build_calendar_event_bubble(event_item, settings)
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已新增行程：{title}（{format_calendar_date_label(event_date)} {start_time}）",
            "reply_flex": flex,
            "reply_quick_reply": build_calendar_quick_reply(settings),
            "parsed_tag": "新增行程",
        }

    return {"handled": False}


# 將行事曆 LINE 指令掛到現有 process_line_message_event 前面；其他客需/委託/待辦照原本流程。
try:
    _process_line_message_event_before_calendar_patch = process_line_message_event

    def process_line_message_event(event):
        calendar_result = process_line_calendar_message_event(event)
        if calendar_result.get("handled"):
            return calendar_result
        return _process_line_message_event_before_calendar_patch(event)

    print("✅ 已啟用 LINE 行事曆指令 Patch")
except Exception as e:
    print("⚠️ 啟用 LINE 行事曆指令 Patch 失敗：", e)


# 覆寫既有 line_webhook view：支援 result['reply_flex'] / result['reply_messages']，仍保留原文字回覆。
def line_webhook_with_flex_calendar():
    raw_body = request.get_data(cache=False, as_text=False)
    signature = request.headers.get("x-line-signature", "")

    if not verify_line_signature(raw_body, signature):
        return "Invalid signature", 400

    try:
        payload = json.loads(raw_body.decode("utf-8"))
    except Exception as e:
        print("⚠️ LINE webhook JSON 解析失敗：", e)
        return "Bad Request", 400

    events = payload.get("events", [])
    for event in events:
        try:
            result = process_line_message_event(event)
            if not result or not result.get("handled"):
                continue

            reply_result = None
            reply_token = event.get("replyToken")
            if not reply_token:
                continue

            if result.get("reply_messages"):
                reply_result = reply_line_messages(reply_token, result.get("reply_messages") or [])
            elif result.get("reply_flex"):
                reply_result = reply_line_flex(
                    reply_token,
                    result.get("reply_text", "CRM 行程卡片"),
                    result.get("reply_flex"),
                    quick_reply_items=result.get("reply_quick_reply"),
                )
            elif result.get("reply_text"):
                reply_result = reply_line_text(
                    reply_token,
                    result["reply_text"] if result.get("ok") else result["reply_text"],
                )

            if result.get("ok") and result.get("target_type") and result.get("target_id") and reply_result:
                for sent in reply_result.get("sent_messages", []):
                    sent_id = str(sent.get("id", "")).strip()
                    if sent_id:
                        save_line_message_link(
                            sent_id,
                            result["target_type"],
                            result["target_id"],
                            tag=result.get("parsed_tag", ""),
                            action="bot_reply",
                            customer_name=result.get("customer_name", ""),
                            phone=result.get("phone", ""),
                            source_event=event,
                        )
        except Exception as e:
            print("⚠️ 處理 LINE event 發生錯誤：", e)

    return "OK", 200


try:
    app.view_functions["line_webhook"] = line_webhook_with_flex_calendar
    print("✅ line_webhook 已升級：支援 Flex Message / 行事曆卡片")
except Exception as e:
    print("⚠️ 覆寫 line_webhook 失敗：", e)

# =============================================================================
# 行事曆後台 + LINE 行程卡片 Flex Message Patch v1 End
# =============================================================================





# ========= DEBUG：檢查目前 Flask 有註冊哪些網址 =========
@app.route("/debug/routes", strict_slashes=False)
def debug_routes():
    rules = []
    for rule in app.url_map.iter_rules():
        methods = ",".join(sorted(rule.methods - {"HEAD", "OPTIONS"}))
        rules.append(f"{methods:10s} {rule.rule}  ->  {rule.endpoint}")
    return "<pre>" + "\n".join(sorted(rules)) + "</pre>"

@app.route("/debug/line-preview-check", strict_slashes=False)
def debug_line_preview_check():
    ok = any(rule.rule == "/line-card-preview" for rule in app.url_map.iter_rules())
    rules = [rule.rule for rule in app.url_map.iter_rules() if "line" in rule.rule or "calendar" in rule.rule]
    return {
        "line_card_preview_registered": ok,
        "calendar_line_routes": sorted(rules),
        "message": "如果這裡顯示 true，就可以開 /line-card-preview?date=2026-06-22"
    }

# =============================================================================
# LINE 卡片預覽頁：不用真的連 LINE，也可以先看 Flex Message 長相
# =============================================================================
@app.route("/line-card-preview", strict_slashes=False)
@app.route("/line_card_preview", strict_slashes=False)
@login_required
def line_card_preview():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    selected_date_label = format_calendar_date_label(selected_date)
    mode = (request.args.get("mode", "day") or "day").strip()

    if mode == "week":
        start_dt = datetime.strptime(selected_date, "%Y-%m-%d").date()
        end_date = (start_dt + timedelta(days=6)).strftime("%Y-%m-%d")
        result = build_calendar_reply_for_range(selected_date, end_date, mode="week")
        events = fetch_calendar_events(selected_date, end_date, line_only=False)
        page_title = f"本週行程卡片預覽｜{format_calendar_date_label(selected_date)} ~ {format_calendar_date_label(end_date)}"
        back_query = f"?date={selected_date}&mode=day"
    else:
        end_date = selected_date
        result = build_calendar_reply_for_range(selected_date, selected_date, mode="today")
        events = fetch_calendar_events(selected_date, selected_date, line_only=False)
        page_title = f"LINE 行程卡片預覽｜{selected_date_label}"
        back_query = f"?date={selected_date}&mode=week"

    settings = get_line_card_settings()
    flex_json_text = json.dumps(result.get("reply_flex", {}), ensure_ascii=False, indent=2)
    quick_reply_text = json.dumps(result.get("reply_quick_reply", []), ensure_ascii=False, indent=2)

    return render_template_string("""
{% extends "base.html" %}
{% block content %}
<style>
  .preview-wrap {display:grid;grid-template-columns:minmax(320px, 430px) 1fr;gap:20px;align-items:start;}
  .phone-frame {background:#dbeafe;border-radius:32px;padding:18px;box-shadow:0 14px 40px rgba(0,0,0,.12);border:1px solid rgba(255,255,255,.8);}
  .phone-top {text-align:center;color:#6b7280;font-size:12px;margin-bottom:12px;}
  .line-area {background:linear-gradient(180deg,#cfe3ff,#eef6ff);border-radius:24px;padding:16px;min-height:640px;overflow:hidden;}
  .msg-row {display:flex;gap:10px;align-items:flex-start;margin-bottom:12px;}
  .bot-icon {width:38px;height:38px;border-radius:50%;background:#7A4E2D;color:white;display:flex;align-items:center;justify-content:center;font-size:11px;flex:0 0 auto;}
  .bot-bubble {background:white;border-radius:18px;padding:10px 14px;font-weight:700;box-shadow:0 4px 16px rgba(0,0,0,.08);}
  .carousel-preview {display:flex;gap:12px;overflow-x:auto;padding:8px 2px 16px;scroll-snap-type:x mandatory;}
  .event-card {background:white;border-radius:18px;width:310px;min-width:310px;padding:16px;box-shadow:0 8px 22px rgba(0,0,0,.13);scroll-snap-align:start;}
  .event-head {display:flex;justify-content:space-between;gap:8px;align-items:center;margin-bottom:10px;}
  .time-pill {font-weight:800;font-size:15px;}
  .cat-pill {border-radius:999px;padding:5px 12px;font-size:13px;font-weight:700;background:#fff4df;border:1px solid #f1c27d;}
  .event-title {font-size:20px;font-weight:900;margin:8px 0 12px;line-height:1.25;}
  .event-line {font-size:14px;margin:8px 0;color:#374151;display:flex;gap:8px;}
  .event-line .label {color:#9ca3af;width:44px;flex:0 0 auto;}
  .event-buttons {display:grid;grid-template-columns:1fr 1fr;gap:8px;margin-top:14px;}
  .mock-btn {border:1px solid #d1d5db;border-radius:10px;text-align:center;padding:8px;font-size:14px;font-weight:700;background:#fff;}
  .mock-btn.primary {background:{{ settings.button_color or '#C9874A' }};color:white;border-color:{{ settings.button_color or '#C9874A' }};}
  .json-box {font-family:Consolas,Menlo,monospace;font-size:12px;white-space:pre;max-height:520px;overflow:auto;background:#111827;color:#e5e7eb;border-radius:12px;padding:14px;}
  @media (max-width: 992px){.preview-wrap{grid-template-columns:1fr}.phone-frame{max-width:460px;margin:auto}.line-area{min-height:560px}}
</style>

<div class="d-flex justify-content-between align-items-center mb-3 flex-wrap gap-2">
  <div>
    <h3 class="mb-1">{{ page_title }}</h3>
    <div class="text-muted small">這裡是本機預覽，不需要部署，也不會真的發送到 LINE。</div>
  </div>
  <div class="d-flex gap-2 flex-wrap">
    <form method="get" action="{{ url_for('line_card_preview') }}" class="d-flex gap-2">
      <input type="date" name="date" value="{{ selected_date }}" class="form-control">
      <input type="hidden" name="mode" value="{{ mode }}">
      <button class="btn btn-outline-secondary" type="submit">切換日期</button>
    </form>
    <a class="btn btn-outline-primary" href="{{ url_for('line_card_preview') }}{{ back_query }}">{{ '看單日' if mode == 'week' else '看本週' }}</a>
    <a class="btn btn-primary" href="{{ url_for('calendar_page', date=selected_date) }}">回行事曆</a>
  </div>
</div>

<div class="preview-wrap">
  <div class="phone-frame">
    <div class="phone-top">LINE Bot 預覽</div>
    <div class="line-area">
      <div class="msg-row">
        <div class="bot-icon">Team<br>ME</div>
        <div class="bot-bubble">📅 {{ result.reply_text }}</div>
      </div>

      <div class="carousel-preview">
        {% if events %}
          {% for e in events[:12] %}
            <div class="event-card">
              <div class="event-head">
                <div class="time-pill" style="color:{{ e.category_color or settings.primary_color or '#C9874A' }};">📅 {{ e.event_date_label or e.event_date }}</div>
                <div class="cat-pill">{{ e.category or '行程' }}</div>
              </div>
              <div class="event-line" style="font-weight:700;color:#6b7280;"><span class="label">時間</span><span>{{ e.start_time }} - {{ e.end_time }}</span></div>
              <div class="event-title">{{ e.title or '未命名行程' }}</div>
              {% if settings.show_customer %}<div class="event-line"><span class="label">客戶</span><span>{{ e.customer_name or '-' }}</span></div>{% endif %}
              {% if settings.show_phone %}<div class="event-line"><span class="label">電話</span><span>{{ e.phone or '-' }}</span></div>{% endif %}
              {% if settings.show_location %}<div class="event-line"><span class="label">地點</span><span>{{ e.location or '-' }}</span></div>{% endif %}
              {% if settings.show_note %}<div class="event-line"><span class="label">備註</span><span>{{ e.note or '-' }}</span></div>{% endif %}
              <div class="event-buttons">
                <div class="mock-btn primary">編輯行程</div>
                {% set related_label = calendar_related_button_label(e) %}
                {% if related_label %}<div class="mock-btn">{{ related_label }}</div>{% endif %}
                {% if e.location %}<div class="mock-btn">查看地圖</div>{% endif %}
                {% if e.phone %}<div class="mock-btn">撥打電話</div>{% endif %}
              </div>
            </div>
          {% endfor %}
        {% else %}
          <div class="event-card">
            <div class="event-title">目前沒有行程</div>
            <div class="event-line"><span>可以先到行事曆新增一筆行程，再回來預覽。</span></div>
            <div class="event-buttons"><div class="mock-btn primary">新增行程格式</div></div>
          </div>
        {% endif %}
      </div>
    </div>
  </div>

  <div>
    <div class="card mb-3">
      <div class="card-header d-flex justify-content-between align-items-center">
        <strong>Flex Message JSON</strong>
        <span class="badge bg-secondary">可貼到 LINE Flex Simulator</span>
      </div>
      <div class="card-body">
        <div class="json-box">{{ flex_json_text }}</div>
      </div>
    </div>
    <div class="card">
      <div class="card-header"><strong>Quick Reply JSON</strong></div>
      <div class="card-body"><div class="json-box" style="max-height:220px;">{{ quick_reply_text }}</div></div>
    </div>
  </div>
</div>
{% endblock %}
    """,
    selected_date=selected_date,
    selected_date_label=selected_date_label,
    mode=mode,
    end_date=end_date,
    result=result,
    events=events,
    settings=settings,
    flex_json_text=flex_json_text,
    quick_reply_text=quick_reply_text,
    page_title=page_title,
    back_query=back_query,
    calendar_related_button_label=calendar_related_button_label,
    )


# =============================================================================
# 後台 BUG 修正 + 群組推播 + 待辦事項後台 + 卡片加入行事曆 Patch v20260621B
# =============================================================================
from urllib.parse import urlencode as _urlencode


def _crm_keep_status(form_value, existing_value, fallback=""):
    value = (form_value or "").strip()
    if value:
        return value
    return (existing_value or fallback or "").strip()


def _crm_line_notify_target_id():
    target = (os.environ.get("LINE_NOTIFY_TARGET_ID") or "").strip()
    if target:
        return target
    try:
        settings = get_line_card_settings()
        return (settings.get("notify_target_id") or settings.get("line_notify_target_id") or "").strip()
    except Exception:
        return ""


def line_push_messages(target_id: str, messages: list):
    if not LINE_CHANNEL_ACCESS_TOKEN:
        return {"ok": False, "error": "LINE_CHANNEL_ACCESS_TOKEN 未設定"}
    if not target_id:
        return {"ok": False, "error": "LINE_NOTIFY_TARGET_ID 未設定"}
    try:
        import requests
        res = requests.post(
            "https://api.line.me/v2/bot/message/push",
            headers=line_api_headers(),
            json={"to": target_id, "messages": messages[:5]},
            timeout=8,
        )
        print("LINE push status:", res.status_code, res.text[:300])
        return {"ok": 200 <= res.status_code < 300, "status_code": res.status_code, "text": res.text[:500]}
    except Exception as e:
        print("⚠️ LINE push 發生錯誤：", e)
        return {"ok": False, "error": str(e)}


def _crm_public_url_for(path, **params):
    if params:
        qs = _urlencode({k: v for k, v in params.items() if v is not None and str(v) != ""})
        path = path + (("&" if "?" in path else "?") + qs if qs else "")
    return build_app_url(path)


def _record_calendar_params(record_type: str, record_id: str, data: dict):
    record_type = record_type or ""
    name = data.get("name") or data.get("customer_name") or ""
    phone = data.get("phone") or ""
    location = data.get("address") or data.get("preferred_areas") or data.get("registered_address") or ""
    if record_type == "buyer":
        title = f"{name} 客需追蹤".strip()
        note = data.get("note") or data.get("requirement_must") or data.get("preferred_areas") or ""
        category = "回電"
    elif record_type == "seller":
        title = f"{name} 委託追蹤".strip()
        note = data.get("note") or data.get("reason") or ""
        category = "回電"
    else:
        title = f"{name or data.get('address') or '開發'} 開發追蹤".strip()
        note = data.get("note") or data.get("next_action") or ""
        category = "開發"
    return {
        "related_type": record_type,
        "related_id": record_id,
        "title": title,
        "category": category,
        "customer_name": name,
        "phone": phone,
        "location": location,
        "note": note,
        "visibility": data.get("visibility") or "public",
        "owner_line_user_id": data.get("owner_line_user_id") or "",
    }


def _record_calendar_url(record_type: str, record_id: str, data: dict):
    params = _record_calendar_params(record_type, record_id, data)
    return _crm_public_url_for("/calendar/new", **params)


def _record_edit_url(record_type: str, record_id: str):
    if record_type == "buyer":
        return _crm_public_url_for(f"/buyers/{record_id}/edit")
    if record_type == "seller":
        return _crm_public_url_for(f"/sellers/{record_id}/edit")
    if record_type == "development":
        return _crm_public_url_for(f"/developments/{record_id}/edit")
    return build_app_url("/")


def _record_detail_url(record_type: str, record_id: str):
    if record_type == "buyer":
        return _crm_public_url_for(f"/buyers/{record_id}")
    if record_type == "seller":
        return _crm_public_url_for(f"/sellers/{record_id}")
    if record_type == "development":
        return _crm_public_url_for(f"/developments/{record_id}")
    return build_app_url("/")


def _build_record_flex_bubble(record_type: str, record_id: str, data: dict, title_prefix="CRM 資料"):
    label_map = {"buyer": "客需", "seller": "委託", "development": "開發"}
    label = label_map.get(record_type, "CRM")
    name = data.get("name") or "未填姓名"
    phone = data.get("phone") or "-"
    source = data.get("source") or "-"
    status = data.get("current_stage") or data.get("stage") or "-"
    next_action = data.get("next_action") or "-"
    area = data.get("preferred_areas") or data.get("address") or data.get("registered_address") or "-"
    note = data.get("note") or data.get("requirement_must") or "-"

    btns = [
        {"type": "button", "style": "primary", "height": "sm", "color": "#C9874A", "action": {"type": "uri", "label": f"編輯{label}", "uri": _record_edit_url(record_type, record_id)}},
        {"type": "button", "style": "secondary", "height": "sm", "action": {"type": "uri", "label": "加入行事曆", "uri": _record_calendar_url(record_type, record_id, data)}},
        {"type": "button", "style": "secondary", "height": "sm", "action": {"type": "uri", "label": "查看後台", "uri": _record_detail_url(record_type, record_id)}},
    ]
    # 在群組點這個按鈕會丟出可被機器人解析的格式，方便群組快速回覆/追蹤。
    if record_type == "buyer":
        msg_text = f"#買方追蹤\n客戶ID: {record_id}\n內容: "
        msg_label = "回覆追蹤"
    elif record_type == "seller":
        msg_text = f"#賣方追蹤\n客戶ID: {record_id}\n內容: "
        msg_label = "回覆追蹤"
    else:
        msg_text = f"#開發追蹤\nID: {record_id}\n內容: "
        msg_label = "回覆開發"
    # LINE 的 message action 會「直接送出」，無法讓同仁繼續打字；
    # 改用 postback + inputOption=openKeyboard + fillInText，點「回覆追蹤」後會開啟鍵盤，
    # 文字先帶入輸入框，同仁補上內容再送出，送出後會被原本 #買方追蹤/#賣方追蹤/#開發追蹤 parser 寫回後台。
    btns.append({
        "type": "button",
        "style": "secondary",
        "height": "sm",
        "action": {
            "type": "postback",
            "label": msg_label,
            "data": f"action=followup_input&record_type={record_type}&record_id={record_id}",
            "inputOption": "openKeyboard",
            "fillInText": msg_text,
        },
    })

    return {
        "type": "bubble",
        "size": "mega",
        "body": {"type": "box", "layout": "vertical", "spacing": "md", "contents": [
            {"type": "text", "text": f"{title_prefix}｜{label}", "size": "xs", "color": "#C9874A", "weight": "bold"},
            {"type": "text", "text": line_truncate(name, 45), "size": "lg", "weight": "bold", "wrap": True, "color": "#222222"},
            {"type": "separator", "margin": "md"},
            {"type": "box", "layout": "vertical", "spacing": "sm", "margin": "md", "contents": [
                flex_info_row("電話", phone),
                flex_info_row("來源", source),
                flex_info_row("狀態", status),
                flex_info_row("下一步", next_action),
                flex_info_row("區域/地址", area),
                flex_info_row("備註", note),
            ]},
        ]},
        "footer": {"type": "box", "layout": "vertical", "spacing": "sm", "contents": btns[:4]},
        "styles": {"footer": {"separator": True}},
    }


def _push_record_to_group(record_type: str, record_id: str, title_prefix="CRM 資料"):
    coll = {"buyer": "buyers", "seller": "sellers", "development": "developments"}.get(record_type)
    if not coll:
        return {"ok": False, "error": "record_type 不正確"}
    snap = db.collection(coll).document(record_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到資料"}
    data = snap.to_dict() or {}
    bubble = _build_record_flex_bubble(record_type, record_id, data, title_prefix=title_prefix)
    return line_push_messages(_crm_line_notify_target_id(), [{"type": "flex", "altText": f"{title_prefix}：{data.get('name','')}", "contents": bubble}])


def _push_calendar_event_to_group(event_id: str, title_prefix="行程資料"):
    snap = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到行程"}
    event = doc_to_calendar_event(snap)
    bubble = build_calendar_event_bubble(event)
    return line_push_messages(_crm_line_notify_target_id(), [{"type": "flex", "altText": f"{title_prefix}：{event.get('title','')}", "contents": bubble}])


@app.route("/buyers/<buyer_id>/send-to-line", methods=["POST"])
@login_required
def buyer_send_to_line(buyer_id):
    res = _push_record_to_group("buyer", buyer_id, title_prefix="後台傳送")
    flash("已傳送到 LINE 群組" if res.get("ok") else f"傳送失敗：{res.get('error') or res.get('text')}", "success" if res.get("ok") else "danger")
    return redirect(request.referrer or url_for("buyer_detail", buyer_id=buyer_id))


@app.route("/sellers/<seller_id>/send-to-line", methods=["POST"])
@login_required
def seller_send_to_line(seller_id):
    res = _push_record_to_group("seller", seller_id, title_prefix="後台傳送")
    flash("已傳送到 LINE 群組" if res.get("ok") else f"傳送失敗：{res.get('error') or res.get('text')}", "success" if res.get("ok") else "danger")
    return redirect(request.referrer or url_for("seller_detail", seller_id=seller_id))


@app.route("/developments/<development_id>/send-to-line", methods=["POST"])
@login_required
def development_send_to_line(development_id):
    res = _push_record_to_group("development", development_id, title_prefix="後台傳送")
    flash("已傳送到 LINE 群組" if res.get("ok") else f"傳送失敗：{res.get('error') or res.get('text')}", "success" if res.get("ok") else "danger")
    return redirect(request.referrer or url_for("development_detail", development_id=development_id))


@app.route("/calendar/<event_id>/send-to-line", methods=["POST"])
@login_required
def calendar_send_to_line(event_id):
    res = _push_calendar_event_to_group(event_id, title_prefix="後台傳送行程")
    flash("已傳送行程到 LINE 群組" if res.get("ok") else f"傳送失敗：{res.get('error') or res.get('text')}", "success" if res.get("ok") else "danger")
    return redirect(request.referrer or url_for("calendar_page"))


@app.after_request
def _backend_update_notify_group(response):
    # 後台編輯 / 新增追蹤後，自動推送到群組。沒有設定 target id 時會安靜略過。
    try:
        if request.method != "POST" or response.status_code not in (200, 302, 303):
            return response
        if not _crm_line_notify_target_id():
            return response
        endpoint = request.endpoint or ""
        path = request.path or ""
        m = None
        if endpoint in ("buyer_edit", "add_buyer_followup"):
            m = re.search(r"/buyers/([^/]+)", path)
            if m:
                _push_record_to_group("buyer", m.group(1), title_prefix="後台更新")
        elif endpoint in ("seller_edit", "add_seller_followup"):
            m = re.search(r"/sellers/([^/]+)", path)
            if m:
                _push_record_to_group("seller", m.group(1), title_prefix="後台更新")
        elif endpoint in ("development_edit", "add_development_followup"):
            m = re.search(r"/developments/([^/]+)", path)
            if m:
                _push_record_to_group("development", m.group(1), title_prefix="後台更新")
        elif endpoint == "calendar_edit":
            m = re.search(r"/calendar/([^/]+)/edit", path)
            if m:
                _push_calendar_event_to_group(m.group(1), title_prefix="行程更新")
    except Exception as e:
        print("⚠️ 後台更新推播失敗：", e)
    return response


# 讓 /calendar/new 可以吃客需/委託/開發卡片帶入的 query string。
def calendar_new_prefill():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    default_start = calendar_safe_time(request.args.get("start", ""), "09:00")
    default_end = calendar_safe_time(request.args.get("end", ""), next_30_min_time(default_start))

    if request.method == "POST":
        payload = build_calendar_event_payload(request.form)
        payload.update({
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id", ""),
            "created_by_name": session.get("user_name", ""),
        })
        doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).add(payload)[1]
        flash("已新增行程", "success")
        if _crm_line_notify_target_id():
            try:
                _push_calendar_event_to_group(doc_ref.id, title_prefix="新增行程")
            except Exception as e:
                print("⚠️ 新增行程推播失敗：", e)
        return redirect(url_for("calendar_page", date=payload["event_date"]))

    prefill = {}
    for key in ("title", "category", "custom_category", "related_type", "related_id", "customer_name", "phone", "location", "note", "visibility"):
        value = (request.args.get(key) or "").strip()
        if value:
            prefill[key] = value
    if prefill:
        prefill["event_date"] = selected_date
        prefill["start_time"] = default_start
        prefill["end_time"] = default_end
        prefill["line_enabled"] = True
    return render_template(
        "calendar_form.html",
        event=prefill if prefill else None,
        selected_date=selected_date,
        default_start=default_start,
        default_end=default_end,
        slots=build_30_min_slots(),
        category_options=CALENDAR_CATEGORY_OPTIONS,
    )

# 覆蓋原本 calendar_new endpoint，但保留原本 URL rule。
app.view_functions["calendar_new"] = login_required(calendar_new_prefill)


# 修正 development edit：表單沒有送 current_stage 時保留原本最後狀態，並補齊 template 變數名稱。
def development_edit_fixed(development_id):
    doc_ref = db.collection("developments").document(development_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆開發", "danger")
        return redirect(url_for("developments"))
    development = doc_to_dict(doc)
    if request.method == "POST":
        form = request.form
        current_stage = normalize_development_status(_crm_keep_status(form.get("current_stage") or form.get("stage"), development.get("current_stage") or development.get("stage"), "待聯繫"))
        next_action = normalize_development_next_action(_crm_keep_status(form.get("next_action"), development.get("next_action"), ""))
        next_action_date = (form.get("next_action_date", "") or development.get("next_action_date", "") or "").strip()
        crm_vis = crm_record_visibility_payload_from_form(form, development)
        updated = {
            **crm_vis,
            "name": form.get("name", "").strip() or development.get("name") or "未填姓名",
            "phone": form.get("phone", "").strip(),
            "source": form.get("source", "").strip(),
            "address": form.get("address", "").strip(),
            "registered_address": form.get("registered_address", "").strip(),
            "url": form.get("url", "").strip(),
            "current_stage": current_stage,
            "stage": current_stage,
            "next_action": next_action,
            "next_action_date": next_action_date,
            "note": form.get("note", "").strip(),
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": session.get("user_id"),
            "updated_by_name": session.get("user_name"),
        }
        doc_ref.update(updated)
        flash("已更新開發資料", "success")
        return redirect(url_for("development_detail", development_id=development_id))
    return render_template(
        "development_edit.html",
        development=development,
        status_options=DEVELOPMENT_STATUS_OPTIONS,
        next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )

app.view_functions["development_edit"] = login_required(development_edit_fixed)


# 記住開發列表的篩選狀態，從詳細頁返回時不會跳回預設。
@app.before_request
def _remember_developments_list_url():
    try:
        if request.endpoint == "developments" and request.method == "GET":
            session["developments_return_url"] = request.full_path
    except Exception:
        pass


# 待辦事項後台。
def _todo_doc_to_dict(doc):
    data = doc.to_dict() or {}
    data["id"] = doc.id
    return data


@app.route("/todos")
@login_required
def todos_page():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    show_done = request.args.get("show_done", "") == "1"
    docs = db.collection(LINE_TODO_COLLECTION).where("todo_date", "==", selected_date).stream()
    items = [_todo_doc_to_dict(d) for d in docs]
    if not show_done:
        items = [x for x in items if x.get("status", "open") != "done"]
    items.sort(key=lambda x: (x.get("status", "open") == "done", x.get("created_at", "")))
    return render_template("todos.html", items=items, selected_date=selected_date, selected_date_label=format_calendar_date_label(selected_date), show_done=show_done)


@app.route("/todos/new", methods=["POST"])
@login_required
def todos_new():
    title = (request.form.get("title") or "").strip()
    todo_date = calendar_safe_date(request.form.get("todo_date") or "")
    note = (request.form.get("note") or "").strip()
    if not title:
        flash("請輸入待辦事項", "warning")
        return redirect(url_for("todos_page", date=todo_date))
    db.collection(LINE_TODO_COLLECTION).add({
        "title": title,
        "content": note,
        "todo_date": todo_date,
        "status": "open",
        "source": "後台",
        "created_at": now_taipei().isoformat(),
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
    })
    flash("已新增待辦事項", "success")
    return redirect(url_for("todos_page", date=todo_date))


@app.route("/todos/<todo_id>/done", methods=["POST"])
@login_required
def todos_done(todo_id):
    ref = db.collection(LINE_TODO_COLLECTION).document(todo_id)
    snap = ref.get()
    date = now_taipei().strftime("%Y-%m-%d")
    if snap.exists:
        date = (snap.to_dict() or {}).get("todo_date") or date
        ref.update({"status": "done", "completed_at": now_taipei().isoformat(), "completed_by_name": session.get("user_name", "")})
        flash("已完成待辦", "success")
    return redirect(url_for("todos_page", date=date))


@app.route("/todos/<todo_id>/delete", methods=["POST"])
@login_required
def todos_delete(todo_id):
    ref = db.collection(LINE_TODO_COLLECTION).document(todo_id)
    snap = ref.get()
    date = now_taipei().strftime("%Y-%m-%d")
    if snap.exists:
        date = (snap.to_dict() or {}).get("todo_date") or date
        ref.delete()
        flash("已刪除待辦", "info")
    return redirect(url_for("todos_page", date=date))

# =============================================================================
# 後台 BUG 修正 + 新增功能 Patch End
# =============================================================================


# =============================================================================
# 追蹤回覆輸入修正 + 開發傳群組穩定修正 Patch v20260621C
# =============================================================================

# 讓 LINE postback「回覆追蹤」按鈕只負責打開鍵盤/預填文字，不需要另外回覆。
def process_line_postback_event(event):
    try:
        postback = event.get("postback") or {}
        data = postback.get("data", "")
        if "followup_input" in data:
            return {"handled": True, "ok": True, "reply_text": ""}
    except Exception:
        pass
    return {"handled": False}


# 如果原本 webhook 有處理 event 但沒處理 postback，這裡只補一個安全 helper；
# 主要輸入寫回仍然是使用者補完 #買方追蹤/#賣方追蹤/#開發追蹤 後送出的文字。


# 重新覆蓋開發傳群組，避免開發頁按鈕因 referrer 或資料欄位缺漏出錯。
def development_send_to_line_fixed(development_id):
    try:
        res = _push_record_to_group("development", development_id, title_prefix="後台傳送")
        if res.get("ok"):
            flash("已傳送到 LINE 群組", "success")
        else:
            flash(f"傳送失敗：{res.get('error') or res.get('text') or res}", "danger")
    except Exception as e:
        print("⚠️ 開發傳到群組失敗：", e)
        flash(f"傳送失敗：{e}", "danger")
    return redirect(request.referrer or url_for("development_detail", development_id=development_id))

try:
    app.view_functions["development_send_to_line"] = login_required(development_send_to_line_fixed)
except Exception as e:
    print("⚠️ 套用 development_send_to_line_fixed 失敗：", e)

# =============================================================================
# 追蹤回覆輸入修正 + 開發傳群組穩定修正 Patch End
# =============================================================================


# =============================================================================
# 同電話多物件：屋主/物件分組 Patch v20260621D
# - 電話用來判斷「同一個屋主」
# - 地址用來判斷「同一個物件」
# - 同電話不同地址時，不再覆蓋原本開發資料，改成新增同屋主底下的新物件
# =============================================================================

def development_owner_key_from_phone(phone: str) -> str:
    return normalize_phone(phone or "")


def development_address_key(address: str) -> str:
    return re.sub(r"\s+", "", (address or "").strip())


def attach_development_owner_groups(items):
    """替開發清單補上同屋主分組資訊，不需要先搬移資料庫。"""
    groups = {}
    for item in items:
        key = development_owner_key_from_phone(item.get("phone", ""))
        if not key:
            continue
        groups.setdefault(key, []).append(item)

    for key, group in groups.items():
        # 同屋主物件依建立時間排序，方便標出第幾筆物件
        group.sort(key=lambda x: x.get("created_at") or "")
        for idx, item in enumerate(group, start=1):
            item["owner_key"] = key
            item["property_group"] = key
            item["owner_group_count"] = len(group)
            item["owner_group_index"] = idx
            item["owner_group_label"] = f"同屋主 {len(group)} 筆物件" if len(group) > 1 else "單一物件"
            item["same_owner_properties"] = [
                {
                    "id": x.get("id", ""),
                    "address": x.get("address") or x.get("registered_address") or "未填地址",
                    "current_stage": x.get("current_stage") or x.get("stage") or "-",
                    "next_action": x.get("next_action") or "-",
                    "source": x.get("source") or "-",
                }
                for x in group
            ]
    return items


def get_development_same_owner_items(phone: str, exclude_id: str = ""):
    key = development_owner_key_from_phone(phone)
    if not key:
        return []
    result = []
    for doc in db.collection("developments").stream():
        item = doc_to_dict(doc)
        item_key = item.get("owner_key") or development_owner_key_from_phone(item.get("phone", ""))
        if item_key == key:
            result.append(item)
    result.sort(key=lambda x: (x.get("created_at") or "", x.get("address") or ""))
    attach_development_owner_groups(result)
    return result


def find_same_development_property(phone: str = "", address: str = "", registered_address: str = ""):
    """同電話 + 同地址/戶籍地址才視為同一個物件。"""
    phone_key = development_owner_key_from_phone(phone)
    addr_keys = {development_address_key(address), development_address_key(registered_address)}
    addr_keys = {x for x in addr_keys if x}
    if not phone_key or not addr_keys:
        return None

    for doc in db.collection("developments").stream():
        item = doc.to_dict() or {}
        item_phone_key = item.get("owner_key") or development_owner_key_from_phone(item.get("phone", ""))
        if item_phone_key != phone_key:
            continue
        item_addr_keys = {
            development_address_key(item.get("address", "")),
            development_address_key(item.get("registered_address", "")),
        }
        item_addr_keys = {x for x in item_addr_keys if x}
        if addr_keys & item_addr_keys:
            return doc
    return None


def enrich_development_owner_payload(payload: dict):
    phone_key = development_owner_key_from_phone(payload.get("phone", ""))
    if phone_key:
        payload["owner_key"] = phone_key
        payload["property_group"] = phone_key
        payload["owner_name"] = payload.get("name", "")
        payload["owner_phone"] = payload.get("phone", "")
    return payload


# 覆蓋開發列表：加入同屋主分組欄位。
def developments_grouped_view():
    q = request.args.get("q", "").strip()
    current_stage = request.args.get("current_stage", "").strip()
    next_action = request.args.get("next_action", "").strip()
    source = request.args.get("source", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")
    show_done = request.args.get("show_done", "").strip()

    docs = db.collection("developments").stream()
    all_items = [doc_to_dict(d) for d in docs]
    total_count = len(all_items)
    source_options = sorted({(x.get("source") or "").strip() for x in all_items if (x.get("source") or "").strip()})

    # 所有資料先補分組，篩選後也保留完整 group count
    attach_development_owner_groups(all_items)
    items = list(all_items)

    if q:
        q_key = development_owner_key_from_phone(q)
        items = [
            x for x in items
            if q in (x.get("name") or "")
            or q in (x.get("phone") or "")
            or q in (x.get("address") or "")
            or q in (x.get("registered_address") or "")
            or (q_key and q_key == (x.get("owner_key") or development_owner_key_from_phone(x.get("phone", ""))))
        ]

    if current_stage:
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") == current_stage]
    if next_action:
        items = [x for x in items if (x.get("next_action") or "") == next_action]
    if source:
        items = [x for x in items if (x.get("source") or "") == source]
    if show_done != "1":
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") not in DEVELOPMENT_HIDDEN_BY_DEFAULT]

    if sort_by == "created_at_asc":
        items.sort(key=lambda x: x.get("created_at") or "")
    elif sort_by == "created_at_desc":
        items.sort(key=lambda x: x.get("created_at") or "", reverse=True)
    elif sort_by == "name_asc":
        items.sort(key=lambda x: (x.get("name") or ""))
    elif sort_by == "name_desc":
        items.sort(key=lambda x: (x.get("name") or ""), reverse=True)
    else:
        items.sort(key=lambda x: x.get("created_at") or "", reverse=True)

    return render_template(
        "developments.html",
        developments=items,
        q=q,
        current_stage=current_stage,
        next_action=next_action,
        source=source,
        source_options=source_options,
        show_done=show_done,
        sort_by=sort_by,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
        total_count=total_count,
        filtered_count=len(items),
        label_docx_enabled=(next_action == "寄開發信"),
        label_docx_count=len([x for x in items if (x.get("registered_address") or "").strip()]),
    )

app.view_functions["developments"] = login_required(developments_grouped_view)


# 覆蓋新增開發：新增 owner_key；同電話同地址提示並更新，反之新增成同屋主多物件。
def developments_new_owner_grouped():
    form = request.form
    phone = form.get("phone", "").strip()
    address = form.get("address", "").strip()
    registered_address = form.get("registered_address", "").strip()
    existing_doc = find_same_development_property(phone=phone, address=address, registered_address=registered_address)

    current_stage = normalize_development_status(form.get("current_stage", "").strip() or form.get("stage", "").strip() or "待聯繫")
    next_action = normalize_development_next_action(form.get("next_action", "").strip())
    next_action_date = form.get("next_action_date", "").strip()
    _manual_url = form.get("url", "").strip()
    nav_url = _make_google_nav_url(registered_address or address)
    crm_vis = crm_record_visibility_payload_from_form(form)

    data = {
        **crm_vis,
        "name": form.get("name", "").strip() or "未填姓名",
        "phone": phone,
        "source": infer_development_source(form.get("source", "").strip(), _manual_url),
        "address": address,
        "registered_address": registered_address,
        "registered_address_google_maps_url": nav_url,
        "url": _manual_url,
        "current_stage": current_stage,
        "stage": current_stage,
        "next_action": next_action,
        "next_action_date": next_action_date,
        "note": form.get("note", "").strip(),
        "record_date": now_taipei().strftime("%Y-%m-%d"),
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id"),
        "updated_by_name": session.get("user_name"),
    }
    enrich_development_owner_payload(data)

    if not (data["name"] or data["phone"] or data["address"] or data["url"]):
        flash("至少請填姓名、電話、地址、網址其中一項", "danger")
        return redirect(url_for("developments"))

    if existing_doc:
        doc_ref = db.collection("developments").document(existing_doc.id)
        doc_ref.update({k: v for k, v in data.items() if v not in (None, "")})
        doc_id = existing_doc.id
        flash("已找到同電話同地址物件，已更新原資料", "success")
    else:
        data.update({
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id"),
            "created_by_name": session.get("user_name"),
        })
        doc_ref = db.collection("developments").document()
        doc_ref.set(data)
        doc_id = doc_ref.id
        if phone and len(get_development_same_owner_items(phone)) > 1:
            flash("已新增同屋主的新物件", "success")
        else:
            flash("已新增開發", "success")

    if data["note"]:
        db.collection("development_followups").add({
            "development_id": doc_id,
            "contact_time": now_taipei().strftime("%Y-%m-%d %H:%M"),
            "channel": "手動新增",
            "current_stage": current_stage,
            "stage": current_stage,
            "next_action": next_action,
            "next_action_date": next_action_date,
            "registered_address": data["registered_address"],
            "content": data["note"],
            "next_contact_date": next_action_date,
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id"),
            "created_by_name": session.get("user_name"),
            "sender_display_name": session.get("user_name"),
        })
    return redirect(url_for("developments", q=phone or address))

app.view_functions["developments_new"] = login_required(developments_new_owner_grouped)


# 覆蓋 LINE 新增開發：同電話不同地址新增，不再擋住；同電話同地址才更新。
def create_development(fields, event):
    phone = (fields.get('phone') or '').strip()
    name = (fields.get('name') or '').strip() or '未填姓名'
    url = (fields.get('url') or '').strip()
    source_value = infer_development_source(fields.get('source', ''), url)
    address = (fields.get('address') or '').strip()
    registered_address = (fields.get('registered_address') or '').strip()
    nav_url = _make_google_nav_url(registered_address or address)

    exact_doc = find_same_development_property(phone=phone, address=address, registered_address=registered_address)
    if not exact_doc and not phone and address:
        exact_doc = find_development_record(address=address)

    labels = build_development_labels(fields.get('labels'))
    content_text = (fields.get('content') or '').strip() or registered_address or address or url or '新增開發'
    note_content = build_line_summary(content_text, event)
    current_stage = normalize_development_status((fields.get('current_stage') or '').strip() or (fields.get('stage') or '').strip() or '待聯繫')

    payload = {
        'name': name,
        'phone': phone,
        'source': source_value,
        'url': url,
        'address': address,
        'registered_address': registered_address,
        'registered_address_google_maps_url': nav_url,
        'current_stage': current_stage,
        'stage': current_stage,
        'next_action': normalize_development_next_action((fields.get('next_action') or '').strip()),
        'next_action_date': (fields.get('next_action_date') or '').strip() or (fields.get('next_contact_date') or '').strip(),
        'record_date': (fields.get('record_date') or '').strip() or now_taipei().strftime('%Y-%m-%d'),
        'labels': labels,
        'updated_at': now_taipei().isoformat(),
        'updated_by_id': 'line_bot',
        'updated_by_name': 'LINE Bot',
        'sender_display_name': get_line_sender_display_name(event) or '',
    }
    enrich_development_owner_payload(payload)

    if exact_doc:
        doc_ref = db.collection('developments').document(exact_doc.id)
        update_customer_note_and_labels(target_type='development', doc_ref=doc_ref, content=note_content, labels=labels, stage=payload['stage'], source=source_value, event=event, registered_address=registered_address)
        clean_updates = {k: v for k, v in payload.items() if v not in ('', None)}
        doc_ref.update(clean_updates)
        add_customer_followup(target_type='development', customer_id=exact_doc.id, content=note_content, next_action=payload.get('next_action', ''), next_contact_date=payload.get('next_action_date', ''), labels=labels, line_event=event)
        updated_doc = doc_ref.get().to_dict() or {}
        reply_text = f"已更新同屋主同物件：{updated_doc.get('name', '')}（{updated_doc.get('phone', '-') or '-'}）"
        return {'handled': True, 'ok': True, 'reply_text': reply_text[:5000], 'target_type': 'development', 'target_id': exact_doc.id, 'customer_name': updated_doc.get('name', ''), 'phone': updated_doc.get('phone', ''), 'parsed_tag': '新增開發'}

    now = now_taipei().isoformat()
    payload.update({'created_at': now, 'created_by_id': 'line_bot', 'created_by_name': 'LINE Bot', 'note': append_note_block('', note_content, build_line_operator_label(event))})
    doc_ref = db.collection('developments').document()
    doc_ref.set(payload)
    add_customer_followup(target_type='development', customer_id=doc_ref.id, content=note_content, next_action=payload.get('next_action', ''), next_contact_date=payload.get('next_action_date', ''), labels=labels, line_event=event)

    group_count = len(get_development_same_owner_items(phone)) if phone else 1
    reply_text = f"已新增開發物件：{name}（{phone or '-'}）"
    if group_count > 1:
        reply_text += f"\n同屋主目前共有 {group_count} 筆物件"
    return {'handled': True, 'ok': True, 'reply_text': reply_text[:5000], 'target_type': 'development', 'target_id': doc_ref.id, 'customer_name': name, 'phone': phone, 'parsed_tag': '新增開發'}


# 詳細頁補同屋主物件清單。
def development_detail_owner_grouped(development_id):
    doc = db.collection("developments").document(development_id).get()
    if not doc.exists:
        flash("找不到這筆開發", "danger")
        return redirect(url_for("developments"))

    development = doc_to_dict(doc)
    same_owner_items = get_development_same_owner_items(development.get("phone", ""))
    followups_ref = db.collection("development_followups").where("development_id", "==", development_id)
    followups = [doc_to_dict(f) for f in followups_ref.stream()]
    followups.sort(key=lambda x: x.get("contact_time", ""), reverse=True)
    return render_template(
        "development_detail.html",
        development=development,
        same_owner_items=same_owner_items,
        followups=followups,
        status_options=DEVELOPMENT_STATUS_OPTIONS,
        next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )

app.view_functions["development_detail"] = login_required(development_detail_owner_grouped)


@app.route("/developments/owner/<owner_key>")
@login_required
def development_owner_group(owner_key):
    owner_key = development_owner_key_from_phone(owner_key)
    items = []
    owner_name = "同屋主"
    owner_phone = owner_key
    for doc in db.collection("developments").stream():
        item = doc_to_dict(doc)
        item_key = item.get("owner_key") or development_owner_key_from_phone(item.get("phone", ""))
        if item_key == owner_key:
            items.append(item)
            owner_name = item.get("name") or owner_name
            owner_phone = item.get("phone") or owner_phone
    items.sort(key=lambda x: (x.get("created_at") or "", x.get("address") or ""))
    attach_development_owner_groups(items)
    return render_template(
        "development_owner_group.html",
        owner_key=owner_key,
        owner_name=owner_name,
        owner_phone=owner_phone,
        items=items,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )

# =============================================================================
# 同電話多物件：屋主/物件分組 Patch End
# =============================================================================


# =============================================================================
# LINE 回覆追蹤：預填輸入框 + 送出後寫回後台 Patch v20260621E
# - 修正舊版 message action 會直接把空白範本送出
# - 改成 postback + openKeyboard + fillInText：點按鈕只預填，不直接送出
# - 使用者補上「內容」後送出，仍由 #買方追蹤/#賣方追蹤/#開發追蹤 寫回後台
# =============================================================================

def _followup_prefill_text(record_type: str, record_id: str) -> str:
    if record_type == "buyer":
        return f"#買方追蹤\n客戶ID: {record_id}\n內容: "
    if record_type == "seller":
        return f"#賣方追蹤\n客戶ID: {record_id}\n內容: "
    if record_type == "development":
        return f"#開發追蹤\nID: {record_id}\n內容: "
    return f"#查詢紀錄\nID: {record_id}\n"


def _followup_postback_action(record_type: str, record_id: str, label: str = "回覆追蹤"):
    # LINE postback action 支援 inputOption=openKeyboard + fillInText。
    # 注意：部分 LINE 電腦版可能不支援預填，手機 LINE 測試最準。
    return {
        "type": "postback",
        "label": label,
        "data": f"action=followup_input&record_type={record_type}&record_id={record_id}",
        "inputOption": "openKeyboard",
        "fillInText": _followup_prefill_text(record_type, record_id),
    }


def process_line_postback_event(event):
    try:
        postback = event.get("postback") or {}
        data = postback.get("data", "") or ""
        if "action=followup_input" in data:
            # 只打開鍵盤 / 預填文字，不需要回覆任何訊息。
            return {"handled": True, "ok": True, "reply_text": ""}
    except Exception as e:
        print("⚠️ process_line_postback_event 發生錯誤：", e)
    return {"handled": False}


# 重新覆蓋群組卡片：保留原欄位與按鈕，但「回覆追蹤」必定使用 postback/fillInText。
def _build_record_flex_bubble(record_type: str, record_id: str, data: dict, title_prefix="CRM 資料"):
    label_map = {"buyer": "客需", "seller": "委託", "development": "開發"}
    label = label_map.get(record_type, "CRM")
    name = data.get("name") or "未填姓名"
    phone = data.get("phone") or "-"
    source = data.get("source") or "-"
    status = data.get("current_stage") or data.get("stage") or "-"
    next_action = data.get("next_action") or "-"
    area = data.get("preferred_areas") or data.get("address") or data.get("registered_address") or "-"
    note = data.get("note") or data.get("requirement_must") or "-"

    if record_type == "development":
        reply_label = "回覆開發"
    else:
        reply_label = "回覆追蹤"

    btns = [
        {
            "type": "button",
            "style": "primary",
            "height": "sm",
            "color": "#C9874A",
            "action": {
                "type": "uri",
                "label": f"編輯{label}",
                "uri": _record_edit_url(record_type, record_id),
            },
        },
        {
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "uri",
                "label": "加入行事曆",
                "uri": _record_calendar_url(record_type, record_id, data),
            },
        },
        {
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "uri",
                "label": "查看後台",
                "uri": _record_detail_url(record_type, record_id),
            },
        },
        {
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": _followup_postback_action(record_type, record_id, reply_label),
        },
    ]

    return {
        "type": "bubble",
        "size": "mega",
        "body": {
            "type": "box",
            "layout": "vertical",
            "spacing": "md",
            "contents": [
                {"type": "text", "text": f"{title_prefix}｜{label}", "size": "xs", "color": "#C9874A", "weight": "bold"},
                {"type": "text", "text": line_truncate(name, 45), "size": "lg", "weight": "bold", "wrap": True, "color": "#222222"},
                {"type": "separator", "margin": "md"},
                {
                    "type": "box",
                    "layout": "vertical",
                    "spacing": "sm",
                    "margin": "md",
                    "contents": [
                        flex_info_row("電話", phone),
                        flex_info_row("來源", source),
                        flex_info_row("狀態", status),
                        flex_info_row("下一步", next_action),
                        flex_info_row("區域/地址", area),
                        flex_info_row("備註", note),
                    ],
                },
            ],
        },
        "footer": {"type": "box", "layout": "vertical", "spacing": "sm", "contents": btns},
        "styles": {"footer": {"separator": True}},
    }


# 重新覆寫 webhook：先處理 postback，再處理文字訊息。
def line_webhook_with_postback_followup():
    raw_body = request.get_data(cache=False, as_text=False)
    signature = request.headers.get("x-line-signature", "")

    if not verify_line_signature(raw_body, signature):
        return "Invalid signature", 400

    try:
        payload = json.loads(raw_body.decode("utf-8"))
    except Exception as e:
        print("⚠️ LINE webhook JSON 解析失敗：", e)
        return "Bad Request", 400

    events = payload.get("events", [])
    for event in events:
        try:
            if event.get("type") == "postback":
                result = process_line_postback_event(event)
            else:
                result = process_line_message_event(event)

            if not result or not result.get("handled"):
                continue

            reply_result = None
            reply_token = event.get("replyToken")
            if not reply_token:
                continue

            if result.get("reply_messages"):
                reply_result = reply_line_messages(reply_token, result.get("reply_messages") or [])
            elif result.get("reply_flex"):
                reply_result = reply_line_flex(
                    reply_token,
                    result.get("reply_text", "CRM 卡片"),
                    result.get("reply_flex"),
                    quick_reply_items=result.get("reply_quick_reply"),
                )
            elif result.get("reply_text"):
                reply_result = reply_line_text(
                    reply_token,
                    result["reply_text"] if result.get("ok") else result["reply_text"],
                )

            if result.get("ok") and result.get("target_type") and result.get("target_id") and reply_result:
                for sent in reply_result.get("sent_messages", []):
                    sent_id = str(sent.get("id", "")).strip()
                    if sent_id:
                        save_line_message_link(
                            sent_id,
                            result["target_type"],
                            result["target_id"],
                            tag=result.get("parsed_tag", ""),
                            action="bot_reply",
                            customer_name=result.get("customer_name", ""),
                            phone=result.get("phone", ""),
                            source_event=event,
                        )
        except Exception as e:
            print("⚠️ 處理 LINE event 發生錯誤：", e)

    return "OK", 200

try:
    app.view_functions["line_webhook"] = line_webhook_with_postback_followup
    print("✅ LINE 回覆追蹤已修正：postback + fillInText 預填輸入框")
except Exception as e:
    print("⚠️ LINE 回覆追蹤修正套用失敗：", e)

# =============================================================================
# LINE 回覆追蹤 Patch End
# =============================================================================


# =============================================================================
# 待辦事項後台卡片式 + 勾選完成即時更新 Patch v20260621D
# =============================================================================

def _todo_item_sort_key(item):
    status = item.get("status", "open")
    return (
        1 if status == "done" else 0,
        item.get("todo_date", ""),
        item.get("created_at", ""),
        item.get("id", ""),
    )


def todos_page_card_view():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    show_done = request.args.get("show_done", "") == "1"

    docs = db.collection(LINE_TODO_COLLECTION).where("todo_date", "==", selected_date).stream()
    all_items = [_todo_doc_to_dict(d) for d in docs]
    all_items.sort(key=_todo_item_sort_key)

    open_count = len([x for x in all_items if x.get("status", "open") != "done"])
    done_count = len([x for x in all_items if x.get("status", "open") == "done"])

    items = all_items if show_done else [x for x in all_items if x.get("status", "open") != "done"]

    return render_template(
        "todos.html",
        items=items,
        selected_date=selected_date,
        selected_date_label=format_calendar_date_label(selected_date),
        show_done=show_done,
        open_count=open_count,
        done_count=done_count,
        total_count=len(all_items),
    )


try:
    app.view_functions["todos_page"] = login_required(todos_page_card_view)
    print("✅ 待辦事項後台已升級：卡片式 + 勾選完成")
except Exception as e:
    print("⚠️ 待辦事項後台升級失敗：", e)


@app.route("/todos/<todo_id>/toggle", methods=["POST"])
@login_required
def todos_toggle(todo_id):
    ref = db.collection(LINE_TODO_COLLECTION).document(todo_id)
    snap = ref.get()

    fallback_date = calendar_safe_date(request.form.get("date") or "")
    show_done = request.form.get("show_done", "") == "1"
    next_url = request.form.get("next") or ""
    desired_status = (request.form.get("status") or "").strip()

    if snap.exists:
        data = snap.to_dict() or {}
        fallback_date = data.get("todo_date") or fallback_date

        if desired_status == "done":
            ref.update({
                "status": "done",
                "completed_at": now_taipei().isoformat(),
                "done_at": now_taipei().isoformat(),
                "completed_by_id": session.get("user_id", ""),
                "completed_by_name": session.get("user_name", ""),
                "done_by_id": session.get("user_id", ""),
                "done_by_name": session.get("user_name", ""),
                "updated_at": now_taipei().isoformat(),
            })
            flash("已完成待辦事項", "success")
        else:
            ref.update({
                "status": "open",
                "completed_at": "",
                "done_at": "",
                "completed_by_id": "",
                "completed_by_name": "",
                "done_by_id": "",
                "done_by_name": "",
                "updated_at": now_taipei().isoformat(),
            })
            flash("已改回待處理", "info")

    if next_url.startswith("/todos"):
        return redirect(next_url)
    return redirect(url_for("todos_page", date=fallback_date, show_done="1" if show_done else ""))

# =============================================================================
# 待辦事項後台卡片式 Patch End
# =============================================================================


# =============================================================================
# LINE 待辦事項卡片版 + 點擊勾選直接更新後台 Patch v20260621L
# - #今日待辦 / #待辦 / #查詢待辦 會回覆同一張 Flex 卡片
# - 每筆待辦的「☐」按鈕是 LINE postback，點擊後直接更新 Firestore status=done
# - 點擊完成後會立即回覆更新後的待辦卡片
# - LINE 沒有真正的 checkbox 元件，這裡用 postback button 模擬勾選效果
# =============================================================================

def _line_todo_flex_safe_text(value, fallback='-'):
    value = str(value or '').strip()
    return value if value else fallback


def _line_todo_flex_truncate(value, max_len=80):
    value = _line_todo_flex_safe_text(value, '')
    if len(value) <= max_len:
        return value
    return value[:max_len - 1] + '…'


def _line_todo_doc_payload(doc):
    data = doc.to_dict() or {}
    data['id'] = doc.id
    return data


def _line_todo_complete_postback_action(todo_id, todo_date=''):
    return {
        'type': 'postback',
        'label': '☐',
        'data': f'action=todo_done&todo_id={todo_id}&date={todo_date or ""}',
        'displayText': '完成待辦',
    }


def _line_todo_done_text_row(title='已完成'):
    return {
        'type': 'box',
        'layout': 'horizontal',
        'spacing': 'sm',
        'contents': [
            {'type': 'text', 'text': '☑', 'size': 'sm', 'color': '#16A34A', 'flex': 1, 'align': 'center'},
            {'type': 'text', 'text': _line_todo_flex_truncate(title, 68), 'size': 'sm', 'color': '#777777', 'flex': 8, 'wrap': True, 'decoration': 'line-through'},
        ],
    }


def _line_todo_item_box(doc, todo_date='', section_label=''):
    data = _line_todo_doc_payload(doc)
    title = data.get('title') or data.get('content') or '未命名待辦'
    note = (data.get('note') or '').strip()
    date_label = _todo_display_md(data.get('todo_date') or todo_date or '') if '_todo_display_md' in globals() else (data.get('todo_date') or todo_date or '')
    created_by = data.get('created_by_name') or data.get('sender_display_name') or ''

    detail_lines = []
    if date_label and section_label not in ('今天', '今日'):
        detail_lines.append(date_label)
    if note:
        detail_lines.append(_line_todo_flex_truncate(note, 55))
    if created_by:
        detail_lines.append(f'建立：{created_by}')
    detail_text = '｜'.join(detail_lines)

    text_contents = [
        {'type': 'text', 'text': _line_todo_flex_truncate(title, 68), 'size': 'sm', 'color': '#222222', 'wrap': True, 'weight': 'bold'},
    ]
    if detail_text:
        text_contents.append({'type': 'text', 'text': _line_todo_flex_truncate(detail_text, 92), 'size': 'xs', 'color': '#888888', 'wrap': True, 'margin': 'xs'})

    return {
        'type': 'box',
        'layout': 'horizontal',
        'spacing': 'sm',
        'paddingTop': 'sm',
        'paddingBottom': 'sm',
        'contents': [
            {
                'type': 'button',
                'style': 'secondary',
                'height': 'sm',
                'flex': 2,
                'action': _line_todo_complete_postback_action(data.get('id', ''), data.get('todo_date') or todo_date),
            },
            {
                'type': 'box',
                'layout': 'vertical',
                'flex': 8,
                'contents': text_contents,
            },
        ],
    }


def _line_todo_section_box(section_title, docs, todo_date='', empty_text='目前沒有待辦', max_items=20):
    docs = list(docs or [])
    contents = [
        {'type': 'text', 'text': section_title, 'size': 'xs', 'color': '#C9874A', 'weight': 'bold'},
    ]

    if not docs:
        contents.append({'type': 'text', 'text': empty_text, 'size': 'sm', 'color': '#999999', 'margin': 'sm'})
    else:
        for doc in docs[:max_items]:
            contents.append(_line_todo_item_box(doc, todo_date=todo_date, section_label=section_title))
        if len(docs) > max_items:
            contents.append({
                'type': 'text',
                'text': f'還有 {len(docs) - max_items} 筆，請到後台查看。',
                'size': 'xs',
                'color': '#999999',
                'margin': 'sm',
                'wrap': True,
            })

    return {
        'type': 'box',
        'layout': 'vertical',
        'spacing': 'xs',
        'margin': 'md',
        'contents': contents,
    }


def _line_todo_get_card_docs(todo_date, target_id=''):
    todo_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    overdue_items = []
    today_items = []
    future_items = []

    try:
        overdue_items = _get_overdue_line_todos(todo_date=todo_date, target_id=target_id)
    except Exception as e:
        print('⚠️ 取得逾期待辦失敗：', e)

    try:
        today_items = _get_open_line_todos(todo_date=todo_date, target_id=target_id, include_overdue=False)
    except Exception as e:
        print('⚠️ 取得當日待辦失敗：', e)

    try:
        # 若目前版本有未來待辦功能，卡片也一起顯示未來待辦。
        if '_get_future_line_todos' in globals():
            future_days = int(globals().get('LINE_TODO_FUTURE_DAYS_DEFAULT', 7) or 7)
            future_items = _get_future_line_todos(start_date=todo_date, target_id=target_id, days=future_days)
    except Exception as e:
        print('⚠️ 取得未來待辦失敗：', e)
        future_items = []

    return overdue_items, today_items, future_items


def build_line_todo_flex_card(todo_date='', target_id='', title_prefix='待辦事項'):
    todo_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    display_date = _todo_display_md(todo_date) if '_todo_display_md' in globals() else todo_date
    overdue_items, today_items, future_items = _line_todo_get_card_docs(todo_date, target_id=target_id)
    total_open = len(overdue_items) + len(today_items) + len(future_items)

    body_contents = [
        {
            'type': 'box',
            'layout': 'horizontal',
            'contents': [
                {'type': 'text', 'text': f'📌 {title_prefix}', 'size': 'lg', 'weight': 'bold', 'color': '#222222', 'flex': 5},
                {'type': 'text', 'text': f'{display_date}', 'size': 'sm', 'color': '#C9874A', 'align': 'end', 'gravity': 'center', 'flex': 3, 'weight': 'bold'},
            ],
        },
        {'type': 'text', 'text': f'未完成 {total_open} 筆｜點左側 ☐ 可直接完成並更新後台', 'size': 'xs', 'color': '#888888', 'wrap': True, 'margin': 'sm'},
        {'type': 'separator', 'margin': 'md'},
    ]

    if not overdue_items and not today_items and not future_items:
        body_contents.append({
            'type': 'box',
            'layout': 'vertical',
            'margin': 'lg',
            'contents': [
                {'type': 'text', 'text': '目前沒有未完成待辦。', 'size': 'sm', 'color': '#666666', 'wrap': True},
            ],
        })
    else:
        if overdue_items:
            body_contents.append(_line_todo_section_box('尚未完成', overdue_items, todo_date=todo_date, empty_text='', max_items=20))
        if today_items:
            body_contents.append(_line_todo_section_box('今天待辦', today_items, todo_date=todo_date, empty_text='', max_items=20))
        if future_items:
            body_contents.append(_line_todo_section_box('未來待辦', future_items, todo_date=todo_date, empty_text='', max_items=10))

    footer_contents = []
    todos_url = build_app_url(f'/todos?date={todo_date}') if 'build_app_url' in globals() else ''
    if todos_url:
        footer_contents.append({
            'type': 'button',
            'style': 'primary',
            'height': 'sm',
            'color': '#C9874A',
            'action': {'type': 'uri', 'label': '查看後台', 'uri': todos_url},
        })
    footer_contents.append({
        'type': 'button',
        'style': 'secondary',
        'height': 'sm',
        'action': {
            'type': 'postback',
            'label': '新增待辦',
            'data': 'action=todo_new_input',
            'inputOption': 'openKeyboard',
            'fillInText': '#新增待辦\n日期: 今天\n事項: ',
        },
    })

    return {
        'type': 'bubble',
        'size': 'mega',
        'body': {
            'type': 'box',
            'layout': 'vertical',
            'spacing': 'sm',
            'contents': body_contents,
        },
        'footer': {
            'type': 'box',
            'layout': 'vertical',
            'spacing': 'sm',
            'contents': footer_contents,
        },
        'styles': {'footer': {'separator': True}},
    }


# 覆寫查詢待辦：文字指令回覆 Flex 卡片。
def query_line_todos(fields, event, force_today=False):
    target_id, _target_type = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    if not todo_date:
        todo_date = now_taipei().strftime('%Y-%m-%d')
    flex = build_line_todo_flex_card(todo_date=todo_date, target_id=target_id, title_prefix='待辦事項')
    return {
        'handled': True,
        'ok': True,
        'reply_text': f'{_todo_display_md(todo_date) if "_todo_display_md" in globals() else todo_date} 待辦事項',
        'reply_flex': flex,
        'parsed_tag': '查詢待辦',
    }


def _line_todo_mark_done_from_postback(todo_id, event=None):
    todo_id = (todo_id or '').strip()
    if not todo_id:
        return False, '缺少待辦 ID', ''

    ref = db.collection(LINE_TODO_COLLECTION).document(todo_id)
    snap = ref.get()
    if not snap.exists:
        return False, '找不到這筆待辦，可能已刪除。', ''

    data = snap.to_dict() or {}
    sender = get_line_sender_display_name(event) if event else ''
    ref.update({
        'status': 'done',
        'completed_at': now_taipei().isoformat(),
        'done_at': now_taipei().isoformat(),
        'completed_by_id': ((event or {}).get('source') or {}).get('userId', ''),
        'completed_by_name': sender or 'LINE',
        'done_by_id': ((event or {}).get('source') or {}).get('userId', ''),
        'done_by_name': sender or 'LINE',
        'updated_at': now_taipei().isoformat(),
    })
    return True, f'已完成：{data.get("title", "待辦")}', data.get('todo_date') or now_taipei().strftime('%Y-%m-%d')


# 擴充 postback：處理待辦勾選完成。
try:
    _process_line_postback_event_before_todo_card = process_line_postback_event
except Exception:
    _process_line_postback_event_before_todo_card = None


def process_line_postback_event(event):
    try:
        from urllib.parse import parse_qs
        postback = event.get('postback') or {}
        raw_data = postback.get('data', '') or ''
        params = {k: (v[0] if isinstance(v, list) and v else '') for k, v in parse_qs(raw_data).items()}
        action = params.get('action') or ''

        if action == 'todo_new_input':
            # 按「新增待辦」只負責打開鍵盤並預填文字；
            # 使用者補上事項後按送出，才會由 #新增待辦 文字流程寫入後台。
            return {
                'handled': True,
                'ok': True,
                'reply_text': '',
                'parsed_tag': '新增待辦輸入',
                'silent': True,
            }

        if action == 'todo_done':
            todo_id = params.get('todo_id') or ''
            ok, msg, todo_date = _line_todo_mark_done_from_postback(todo_id, event=event)
            target_id, _target_type = _line_todo_target_from_event(event)
            query_date = _parse_line_todo_date(params.get('date') or todo_date or '') or now_taipei().strftime('%Y-%m-%d')
            flex = build_line_todo_flex_card(todo_date=query_date, target_id=target_id, title_prefix='待辦事項已更新' if ok else '待辦事項')
            return {
                'handled': True,
                'ok': ok,
                'reply_text': msg,
                'reply_flex': flex,
                'parsed_tag': '完成待辦',
            }
    except Exception as e:
        print('⚠️ LINE 待辦 postback 處理失敗：', e)
        return {'handled': True, 'ok': False, 'reply_text': f'待辦更新失敗：{e}'}

    if _process_line_postback_event_before_todo_card:
        return _process_line_postback_event_before_todo_card(event)
    return {'handled': False}


def push_line_flex(to_id: str, alt_text: str, contents: dict):
    if not LINE_CHANNEL_ACCESS_TOKEN or not to_id:
        return False, 'LINE_CHANNEL_ACCESS_TOKEN 或 to_id 為空'
    payload = {
        'to': to_id,
        'messages': [{
            'type': 'flex',
            'altText': (alt_text or '待辦事項')[:400],
            'contents': contents,
        }],
    }
    try:
        import requests
        res = requests.post(
            'https://api.line.me/v2/bot/message/push',
            headers=line_api_headers(),
            json=payload,
            timeout=8,
        )
        print('LINE push flex status:', res.status_code, res.text[:300])
        return res.status_code in (200, 202), res.text[:300]
    except Exception as e:
        print('⚠️ LINE flex push 發生錯誤：', e)
        return False, str(e)


# 每日提醒也改成推送 Flex 卡片；提醒後仍會寫入 reminder_sent_dates 避免重複推播。
def send_today_line_todo_reminders():
    today = now_taipei().strftime('%Y-%m-%d')
    grouped = {}

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        todo_date = (data.get('todo_date') or '').strip()
        if not todo_date or todo_date > today:
            continue
        sent_dates = data.get('reminder_sent_dates') or []
        if today in sent_dates:
            continue
        target_id = data.get('line_target_id', '')
        if not target_id:
            continue
        grouped.setdefault(target_id, []).append(doc)

    sent_count = 0
    failed = []
    for target_id, docs in grouped.items():
        flex = build_line_todo_flex_card(todo_date=today, target_id=target_id, title_prefix='今日待辦提醒')
        ok, msg = push_line_flex(target_id, f'{_todo_display_md(today) if "_todo_display_md" in globals() else today} 今日待辦提醒', flex)
        if ok:
            sent_count += 1
            for doc in docs:
                try:
                    doc.reference.update({
                        'reminder_sent_dates': firestore.ArrayUnion([today]),
                        'last_reminded_at': now_taipei().isoformat(),
                    })
                except Exception as e:
                    print('⚠️ 更新 reminder_sent_dates 失敗：', e)
        else:
            failed.append({'target_id': target_id, 'error': msg})
    return {'date': today, 'target_count': len(grouped), 'sent_count': sent_count, 'failed': failed}

print('✅ LINE 待辦事項已升級：同一張卡片 + 點擊勾選直接更新後台')
# =============================================================================
# LINE 待辦事項卡片版 Patch End
# =============================================================================



# =============================================================================
# LINE 待辦事項查詢修正 Patch v20260621M
# - LINE #今日待辦 / #待辦 會抓得到後台 /todos 原本新增的待辦
# - 後台新增的待辦通常沒有 line_target_id，視為「共用待辦」，所有群組可查詢
# - LINE 群組自己新增的待辦只顯示在該群組，避免別群資料混在一起
# - 顯示規則：指定日期以前未完成 + 指定日期當天未完成；不顯示未來待辦
# - 點 LINE 卡片左側 ☐ 一樣直接更新 Firestore 後台 status=done
# =============================================================================


def _line_todo_visible_for_target(data: dict, target_id: str = '') -> bool:
    """判斷一筆待辦是否應該在目前 LINE 對話中顯示。

    規則：
    1. 後台新增的待辦通常沒有 line_target_id，視為共用待辦，LINE 查詢要看得到。
    2. LINE 群組 / 個人對話新增的待辦有 line_target_id，只顯示在同一個 target。
    """
    item_target = (data.get('line_target_id') or '').strip()
    target_id = (target_id or '').strip()

    # 後台待辦 / 舊資料：沒有 target，視為共用待辦
    if not item_target:
        return True

    # 沒有目標限制時也可顯示
    if not target_id:
        return True

    return item_target == target_id


def _is_open_todo_doc(doc, target_id=''):
    """覆寫：LINE 查詢要包含後台共用待辦。"""
    data = doc.to_dict() or {}
    if data.get('status', 'open') != 'open':
        return False
    if not _line_todo_visible_for_target(data, target_id=target_id):
        return False
    if not (data.get('todo_date') or '').strip():
        return False
    return True


def _line_todo_doc_note(data: dict) -> str:
    """兼容後台欄位：後台舊資料可能把備註存在 content，不是 note。"""
    return (data.get('note') or data.get('content') or '').strip()


def _get_open_line_todos(todo_date='', target_id='', include_overdue=False):
    """指定日期待辦。include_overdue=True 時包含指定日期以前。"""
    query_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    result = []

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = _todo_date_value(doc)
        if include_overdue:
            if d <= query_date:
                result.append(doc)
        else:
            if d == query_date:
                result.append(doc)

    return _sort_line_todo_docs(result)


def _get_overdue_line_todos(todo_date='', target_id=''):
    """指定日期以前尚未完成的待辦；會持續顯示直到完成或刪除。"""
    query_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    result = []

    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = _todo_date_value(doc)
        if d < query_date:
            result.append(doc)

    return _sort_line_todo_docs(result)


def _get_display_line_todos(todo_date='', target_id=''):
    """LINE 畫面使用：逾期未完成在前，當日待辦在後。"""
    query_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    return _get_overdue_line_todos(query_date, target_id=target_id) + _get_open_line_todos(query_date, target_id=target_id, include_overdue=False)


def _line_todo_item_box(doc, todo_date='', section_label=''):
    """覆寫卡片列：備註相容 note/content，日期顯示 M/D。"""
    data = _line_todo_doc_payload(doc)
    title = data.get('title') or data.get('content') or '未命名待辦'
    note = _line_todo_doc_note(data)
    date_label = _todo_display_md(data.get('todo_date') or todo_date or '') if '_todo_display_md' in globals() else (data.get('todo_date') or todo_date or '')
    created_by = data.get('created_by_name') or data.get('sender_display_name') or ''
    source = data.get('source') or ''

    detail_lines = []
    if date_label and section_label not in ('今天', '今日', '今天待辦'):
        detail_lines.append(date_label)
    if note:
        detail_lines.append(_line_todo_flex_truncate(note, 55))
    if created_by:
        detail_lines.append(f'建立：{created_by}')
    if source == '後台' and not (data.get('line_target_id') or '').strip():
        detail_lines.append('共用')
    detail_text = '｜'.join(detail_lines)

    text_contents = [
        {'type': 'text', 'text': _line_todo_flex_truncate(title, 68), 'size': 'sm', 'color': '#222222', 'wrap': True, 'weight': 'bold'},
    ]
    if detail_text:
        text_contents.append({'type': 'text', 'text': _line_todo_flex_truncate(detail_text, 92), 'size': 'xs', 'color': '#888888', 'wrap': True, 'margin': 'xs'})

    return {
        'type': 'box',
        'layout': 'horizontal',
        'spacing': 'sm',
        'paddingTop': 'sm',
        'paddingBottom': 'sm',
        'contents': [
            {
                'type': 'button',
                'style': 'secondary',
                'height': 'sm',
                'flex': 2,
                'action': _line_todo_complete_postback_action(data.get('id', ''), data.get('todo_date') or todo_date),
            },
            {
                'type': 'box',
                'layout': 'vertical',
                'flex': 8,
                'contents': text_contents,
            },
        ],
    }


def _line_todo_get_card_docs(todo_date, target_id=''):
    """覆寫：只顯示逾期未完成 + 當日未完成，不顯示未來待辦。"""
    todo_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    overdue_items = _get_overdue_line_todos(todo_date=todo_date, target_id=target_id)
    today_items = _get_open_line_todos(todo_date=todo_date, target_id=target_id, include_overdue=False)
    future_items = []
    return overdue_items, today_items, future_items


def build_line_todo_flex_card(todo_date='', target_id='', title_prefix='待辦事項'):
    """覆寫 LINE 待辦卡片：包含後台共用待辦 + 未完成延續顯示。"""
    todo_date = _parse_line_todo_date(todo_date or '') or now_taipei().strftime('%Y-%m-%d')
    display_date = _todo_display_md(todo_date) if '_todo_display_md' in globals() else todo_date
    overdue_items, today_items, future_items = _line_todo_get_card_docs(todo_date, target_id=target_id)
    total_open = len(overdue_items) + len(today_items)

    body_contents = [
        {
            'type': 'box',
            'layout': 'horizontal',
            'contents': [
                {'type': 'text', 'text': f'📌 {title_prefix}', 'size': 'lg', 'weight': 'bold', 'color': '#222222', 'flex': 5},
                {'type': 'text', 'text': f'{display_date}', 'size': 'sm', 'color': '#C9874A', 'align': 'end', 'gravity': 'center', 'flex': 3, 'weight': 'bold'},
            ],
        },
        {'type': 'text', 'text': f'未完成 {total_open} 筆｜包含今日與過去尚未完成｜點左側 ☐ 可完成並更新後台', 'size': 'xs', 'color': '#888888', 'wrap': True, 'margin': 'sm'},
        {'type': 'separator', 'margin': 'md'},
    ]

    if not overdue_items and not today_items:
        body_contents.append({
            'type': 'box',
            'layout': 'vertical',
            'margin': 'lg',
            'contents': [
                {'type': 'text', 'text': '目前沒有未完成待辦。', 'size': 'sm', 'color': '#666666', 'wrap': True},
            ],
        })
    else:
        if overdue_items:
            body_contents.append(_line_todo_section_box('尚未完成', overdue_items, todo_date=todo_date, empty_text='', max_items=20))
        if today_items:
            body_contents.append(_line_todo_section_box('今天待辦', today_items, todo_date=todo_date, empty_text='', max_items=20))

    footer_contents = []
    todos_url = build_app_url(f'/todos?date={todo_date}') if 'build_app_url' in globals() else ''
    if todos_url:
        footer_contents.append({
            'type': 'button',
            'style': 'primary',
            'height': 'sm',
            'color': '#C9874A',
            'action': {'type': 'uri', 'label': '查看後台', 'uri': todos_url},
        })
    footer_contents.append({
        'type': 'button',
        'style': 'secondary',
        'height': 'sm',
        'action': {
            'type': 'postback',
            'label': '新增待辦',
            'data': 'action=todo_new_input',
            'inputOption': 'openKeyboard',
            'fillInText': '#新增待辦\n日期: 今天\n事項: ',
        },
    })

    return {
        'type': 'bubble',
        'size': 'mega',
        'body': {
            'type': 'box',
            'layout': 'vertical',
            'spacing': 'sm',
            'contents': body_contents,
        },
        'footer': {
            'type': 'box',
            'layout': 'vertical',
            'spacing': 'sm',
            'contents': footer_contents,
        },
        'styles': {'footer': {'separator': True}},
    }


def query_line_todos(fields, event, force_today=False):
    """覆寫 LINE 查詢：#今日待辦 會顯示今日 + 過去未完成，並包含後台共用待辦。"""
    target_id, _target_type = _line_todo_target_from_event(event)
    todo_date = now_taipei().strftime('%Y-%m-%d') if force_today else _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    if not todo_date:
        todo_date = now_taipei().strftime('%Y-%m-%d')
    flex = build_line_todo_flex_card(todo_date=todo_date, target_id=target_id, title_prefix='待辦事項')
    return {
        'handled': True,
        'ok': True,
        'reply_text': f'{_todo_display_md(todo_date) if "_todo_display_md" in globals() else todo_date} 待辦事項',
        'reply_flex': flex,
        'parsed_tag': '查詢待辦',
    }


def _line_todo_mark_done_from_postback(todo_id, event=None):
    """點 LINE 卡片 ☐ 後直接完成後台資料。"""
    todo_id = (todo_id or '').strip()
    if not todo_id:
        return False, '缺少待辦 ID', ''

    ref = db.collection(LINE_TODO_COLLECTION).document(todo_id)
    snap = ref.get()
    if not snap.exists:
        return False, '找不到這筆待辦，可能已刪除。', ''

    data = snap.to_dict() or {}
    sender = get_line_sender_display_name(event) if event else ''
    now_iso = now_taipei().isoformat()
    ref.update({
        'status': 'done',
        'completed_at': now_iso,
        'done_at': now_iso,
        'completed_by_id': ((event or {}).get('source') or {}).get('userId', ''),
        'completed_by_name': sender or 'LINE',
        'done_by_id': ((event or {}).get('source') or {}).get('userId', ''),
        'done_by_name': sender or 'LINE',
        'updated_at': now_iso,
    })
    return True, f'已完成：{data.get("title", "待辦")}', data.get('todo_date') or now_taipei().strftime('%Y-%m-%d')


# 後台新增待辦時，兩個欄位都存，避免 LINE 讀不到備註。
def todos_new_line_compatible():
    title = (request.form.get('title') or '').strip()
    todo_date = calendar_safe_date(request.form.get('todo_date') or '')
    note = (request.form.get('note') or '').strip()
    if not title:
        flash('請輸入待辦事項', 'warning')
        return redirect(url_for('todos_page', date=todo_date))
    db.collection(LINE_TODO_COLLECTION).add({
        'title': title,
        'content': note,
        'note': note,
        'todo_date': todo_date,
        'status': 'open',
        'source': '後台',
        # line_target_id 留空：代表共用待辦，LINE 群組查詢要能看得到。
        'line_target_id': '',
        'line_target_type': 'backend_shared',
        'created_at': now_taipei().isoformat(),
        'created_by_id': session.get('user_id'),
        'created_by_name': session.get('user_name'),
    })
    flash('已新增待辦事項', 'success')
    return redirect(url_for('todos_page', date=todo_date))

try:
    app.view_functions['todos_new'] = login_required(todos_new_line_compatible)
    print('✅ LINE 待辦查詢已修正：包含後台共用待辦 + 未完成延續顯示')
except Exception as e:
    print('⚠️ LINE 待辦查詢修正套用失敗：', e)
# =============================================================================
# LINE 待辦事項查詢修正 Patch End
# =============================================================================



# =============================================================================
# 行事曆分類客製化 Patch v1
# 功能：
# 1. 在 /line-card-settings 裡設定「行事曆行程分類」。
# 2. /calendar/new 與 /calendar/<id>/edit 的類型下拉選單改吃後台設定。
# 3. LINE #新增行程 的「類型」也改吃後台設定。
# 4. 自訂分類沒有指定顏色時，預設使用 LINE 卡片設定的主色。
# =============================================================================

DEFAULT_CALENDAR_CATEGORIES = [
    "帶看",
    "回電",
    "開發",
    "簽約",
    "拍照",
    "收服務費",
    "待辦",
    "其他",
]

try:
    DEFAULT_LINE_CARD_SETTINGS.setdefault("calendar_categories", list(DEFAULT_CALENDAR_CATEGORIES))
except Exception:
    pass


def _parse_calendar_category_list(value):
    """把後台輸入的分類文字轉成乾淨清單。支援一行一個、逗號、頓號。"""
    if isinstance(value, list):
        raw_items = value
    else:
        raw_items = re.split(r"[，,、\n]+", str(value or ""))

    result = []
    seen = set()
    for item in raw_items:
        label = str(item or "").strip()
        if not label:
            continue
        if label in seen:
            continue
        seen.add(label)
        result.append(label)

    if not result:
        result = list(DEFAULT_CALENDAR_CATEGORIES)

    if "其他" not in result:
        result.append("其他")

    return result


def _sync_calendar_category_global(category_list):
    """同步舊程式會使用的 CALENDAR_CATEGORY_OPTIONS，避免舊函式讀到舊清單。"""
    try:
        categories = _parse_calendar_category_list(category_list)
        globals()["CALENDAR_CATEGORY_OPTIONS"][:] = categories
        return categories
    except Exception:
        globals()["CALENDAR_CATEGORY_OPTIONS"] = _parse_calendar_category_list(category_list)
        return globals()["CALENDAR_CATEGORY_OPTIONS"]


_original_get_line_card_settings_for_categories = get_line_card_settings


def get_line_card_settings():
    settings = _original_get_line_card_settings_for_categories()
    categories = _parse_calendar_category_list(settings.get("calendar_categories") or DEFAULT_CALENDAR_CATEGORIES)
    settings["calendar_categories"] = categories
    _sync_calendar_category_global(categories)
    return settings


_original_save_line_card_settings_from_form_for_categories = save_line_card_settings_from_form


def save_line_card_settings_from_form(form):
    updates = _original_save_line_card_settings_from_form_for_categories(form)
    categories = _parse_calendar_category_list(form.get("calendar_categories") or updates.get("calendar_categories") or DEFAULT_CALENDAR_CATEGORIES)
    updates["calendar_categories"] = categories
    updates["updated_at"] = now_taipei().isoformat()
    updates["updated_by_id"] = session.get("user_id", "")
    updates["updated_by_name"] = session.get("user_name", "")
    db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set({
        "calendar_categories": categories,
        "updated_at": updates["updated_at"],
        "updated_by_id": updates["updated_by_id"],
        "updated_by_name": updates["updated_by_name"],
    }, merge=True)
    _sync_calendar_category_global(categories)
    return updates


def get_calendar_category_options():
    settings = get_line_card_settings()
    return _parse_calendar_category_list(settings.get("calendar_categories") or DEFAULT_CALENDAR_CATEGORIES)


def get_calendar_category_color(category, settings=None):
    settings = settings or get_line_card_settings()
    category = (category or "其他").strip()
    return CALENDAR_CATEGORY_COLOR_MAP.get(category) or settings.get("primary_color") or "#C9874A"


# 讓 LINE 卡片設定頁增加「行事曆行程分類」輸入欄位。
_CATEGORY_SETTINGS_BLOCK = """
        <div class="full">
          <label>行事曆行程分類</label>
          <textarea name="calendar_categories">{{ calendar_categories_text }}</textarea>
          <div class="hint">
            一行一個分類，會同步用在：後台新增行程、編輯行程、LINE #新增行程。<br>
            例如：帶看、回電、開發、簽約、拍照、收服務費、跑銀行、會議、私人行程、其他。<br>
            建議保留「其他」，若沒填系統會自動補上。
          </div>
        </div>
"""

try:
    LINE_CARD_SETTINGS_TEMPLATE_CUSTOM_CATEGORIES = LINE_CARD_SETTINGS_TEMPLATE.replace(
        '        <div class="full">\n          <label>快速操作按鈕文字</label>',
        _CATEGORY_SETTINGS_BLOCK + '        <div class="full">\n          <label>快速操作按鈕文字</label>'
    )
except Exception:
    LINE_CARD_SETTINGS_TEMPLATE_CUSTOM_CATEGORIES = LINE_CARD_SETTINGS_TEMPLATE


def line_card_settings_custom_categories():
    if request.method == "POST":
        save_line_card_settings_from_form(request.form)
        flash("LINE 卡片設定已更新", "success")
        return redirect(url_for("line_card_settings"))

    settings = get_line_card_settings()
    return render_template_string(
        LINE_CARD_SETTINGS_TEMPLATE_CUSTOM_CATEGORIES,
        settings=settings,
        quick_actions_text="\n".join(settings.get("quick_actions") or []),
        calendar_categories_text="\n".join(get_calendar_category_options()),
    )


try:
    app.view_functions["line_card_settings"] = login_required(line_card_settings_custom_categories)
except Exception as e:
    print("⚠️ 套用行事曆分類設定頁失敗：", e)


def doc_to_calendar_event(doc):
    data = doc_to_dict(doc)
    data["start_time"] = calendar_safe_time(data.get("start_time"), "09:00")
    data["end_time"] = calendar_safe_time(data.get("end_time"), next_30_min_time(data["start_time"]))
    data["event_date"] = calendar_safe_date(data.get("event_date"))
    data["event_date_label"] = format_calendar_date_label(data.get("event_date"))
    data["category"] = (data.get("category") or "其他").strip()
    data["custom_category"] = data.get("custom_category") or ""
    data["display_category"] = data.get("display_category") or (
        data["custom_category"] if data["category"] == "其他" and data["custom_category"] else data["category"]
    )
    data["visibility"] = data.get("visibility") or "public"
    data["category_color"] = data.get("category_color") or get_calendar_category_color(data["category"])
    return data


def build_calendar_event_payload(form, existing=None):
    existing = existing or {}
    event_date = calendar_safe_date(form.get("event_date") or existing.get("event_date"))
    start_time = calendar_safe_time(form.get("start_time") or existing.get("start_time"), "09:00")
    end_time = calendar_safe_time(form.get("end_time") or existing.get("end_time"), next_30_min_time(start_time))
    if calendar_time_to_minutes(end_time) <= calendar_time_to_minutes(start_time):
        end_time = next_30_min_time(start_time)

    category_options = get_calendar_category_options()
    category = (form.get("category", "") or existing.get("category") or "其他").strip()
    if category not in category_options:
        category = "其他"

    custom_category = (form.get("custom_category", "") or existing.get("custom_category", "") or "").strip()
    display_category = custom_category if category == "其他" and custom_category else category

    visibility = (form.get("visibility", "") or existing.get("visibility") or "personal").strip()
    if visibility not in ("personal", "public"):
        visibility = "personal"

    payload = {
        "title": (form.get("title", "") or "").strip(),
        "event_date": event_date,
        "start_time": start_time,
        "end_time": end_time,
        "category": category,
        "custom_category": custom_category,
        "display_category": display_category,
        "visibility": visibility,
        "owner_user_id": existing.get("owner_user_id") or session.get("user_id", ""),
        "owner_user_name": existing.get("owner_user_name") or session.get("user_name", ""),
        "category_color": get_calendar_category_color(category),
        "related_type": (form.get("related_type", "") or "").strip(),
        "related_id": (form.get("related_id", "") or "").strip(),
        "customer_name": (form.get("customer_name", "") or "").strip(),
        "phone": (form.get("phone", "") or "").strip(),
        "location": (form.get("location", "") or "").strip(),
        "note": (form.get("note", "") or "").strip(),
        "line_enabled": form.get("line_enabled") == "on",
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }
    if not payload["title"]:
        name = payload.get("customer_name") or payload.get("location") or category
        payload["title"] = f"{name} {category}".strip()
    return payload


def calendar_page_custom_categories():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    selected_date_label = format_calendar_date_label(selected_date)
    events = fetch_calendar_events(selected_date)
    slots = build_30_min_slots()
    slot_cells = build_calendar_slot_cells(events, slots)

    event_map = {}
    for e in events:
        event_map.setdefault(e.get("start_time"), []).append(e)

    dates = calendar_prev_next_dates(selected_date)
    return render_template(
        "calendar.html",
        selected_date=selected_date,
        selected_date_label=selected_date_label,
        slots=slots,
        slot_cells=slot_cells,
        event_map=event_map,
        events=events,
        category_options=get_calendar_category_options(),
        **dates,
    )


def calendar_edit_custom_categories(event_id):
    doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆行程", "danger")
        return redirect(url_for("calendar_page"))

    event = doc_to_calendar_event(doc)
    if request.method == "POST":
        payload = build_calendar_event_payload(request.form, existing=event)
        doc_ref.update(payload)
        flash("已更新行程", "success")
        return redirect(url_for("calendar_page", date=payload["event_date"]))

    return render_template(
        "calendar_form.html",
        event=event,
        selected_date=event.get("event_date"),
        default_start=event.get("start_time", "09:00"),
        default_end=event.get("end_time", "09:30"),
        slots=build_30_min_slots(),
        category_options=get_calendar_category_options(),
    )


def calendar_new_prefill_custom_categories():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    default_start = calendar_safe_time(request.args.get("start", ""), "09:00")
    default_end = calendar_safe_time(request.args.get("end", ""), next_30_min_time(default_start))

    if request.method == "POST":
        payload = build_calendar_event_payload(request.form)
        payload.update({
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id", ""),
            "created_by_name": session.get("user_name", ""),
        })
        doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).add(payload)[1]
        flash("已新增行程", "success")
        if _crm_line_notify_target_id():
            try:
                _push_calendar_event_to_group(doc_ref.id, title_prefix="新增行程")
            except Exception as e:
                print("⚠️ 新增行程推播失敗：", e)
        return redirect(url_for("calendar_page", date=payload["event_date"]))

    prefill = {}
    for key in ("title", "category", "custom_category", "related_type", "related_id", "customer_name", "phone", "location", "note", "visibility"):
        value = (request.args.get(key) or "").strip()
        if value:
            prefill[key] = value

    if prefill:
        prefill["event_date"] = selected_date
        prefill["start_time"] = default_start
        prefill["end_time"] = default_end
        prefill["line_enabled"] = True

    return render_template(
        "calendar_form.html",
        event=prefill if prefill else None,
        selected_date=selected_date,
        default_start=default_start,
        default_end=default_end,
        slots=build_30_min_slots(),
        category_options=get_calendar_category_options(),
    )


# LINE #新增行程 也吃後台自訂行程分類。
def process_line_calendar_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    raw_text = (message.get("text") or "").strip()
    normalized = raw_text.replace("＃", "#").strip()
    today = now_taipei().date()

    if normalized in ("#今日行程", "#今天行程", "今日行程", "今天行程"):
        return build_calendar_reply_for_range(today.strftime("%Y-%m-%d"), mode="today")

    if normalized in ("#明日行程", "#明天行程", "明日行程", "明天行程"):
        return build_calendar_reply_for_range((today + timedelta(days=1)).strftime("%Y-%m-%d"), mode="tomorrow")

    if normalized in ("#本週行程", "本週行程"):
        start = today.strftime("%Y-%m-%d")
        end = (today + timedelta(days=6)).strftime("%Y-%m-%d")
        return build_calendar_reply_for_range(start, end, mode="week")

    if normalized.startswith("#新增行程"):
        fields = parse_line_calendar_create_fields(normalized)
        category_options = get_calendar_category_options()
        if len(normalized.splitlines()) == 1:
            sample_category = category_options[0] if category_options else "帶看"
            example = (
                "新增行程格式：\n"
                "#新增行程\n"
                "日期: 今天\n"
                "時間: 10:00-10:30\n"
                f"類型: {sample_category}\n"
                "標題: 童先生看農舍\n"
                "客戶: 童先生\n"
                "電話: 0921-123-456\n"
                "地點: 清水、梧棲交界\n"
                "備註: 退休夫妻，想看農舍、有空地\n\n"
                "目前可用類型：" + "、".join(category_options)
            )
            return {"handled": True, "ok": True, "reply_text": example, "parsed_tag": "新增行程格式"}

        event_date = parse_calendar_date_word(fields.get("event_date_raw", "今天"))
        start_time, end_time = parse_calendar_time_range(fields.get("time_raw", "09:00"), fields.get("end_time", ""))
        category = (fields.get("category") or "其他").strip()
        if category not in category_options:
            category = "其他"

        title = fields.get("title", "").strip()
        if not title:
            title = f"{fields.get('customer_name', '')} {category}".strip() or category

        payload = {
            "title": title,
            "event_date": event_date,
            "start_time": start_time,
            "end_time": end_time,
            "category": category,
            "display_category": category,
            "category_color": get_calendar_category_color(category),
            "related_type": fields.get("related_type", ""),
            "related_id": fields.get("related_id", ""),
            "customer_name": fields.get("customer_name", ""),
            "phone": fields.get("phone", ""),
            "location": fields.get("location", ""),
            "note": fields.get("note", ""),
            "visibility": "public",
            "line_enabled": True,
            "created_at": now_taipei().isoformat(),
            "created_by_id": "line_bot",
            "created_by_name": "LINE Bot",
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": "line_bot",
            "updated_by_name": "LINE Bot",
        }
        doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document()
        doc_ref.set(payload)
        payload["id"] = doc_ref.id
        settings = get_line_card_settings()
        flex = build_calendar_event_bubble(dict(payload), settings)
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已新增行程：{title}（{format_calendar_date_label(event_date)} {start_time}）",
            "reply_flex": flex,
            "reply_quick_reply": build_calendar_quick_reply(settings),
            "parsed_tag": "新增行程",
        }

    return {"handled": False}


try:
    _sync_calendar_category_global(get_calendar_category_options())
    app.view_functions["calendar_page"] = login_required(calendar_page_custom_categories)
    app.view_functions["calendar_new"] = login_required(calendar_new_prefill_custom_categories)
    app.view_functions["calendar_edit"] = login_required(calendar_edit_custom_categories)
    print("✅ 行事曆分類客製化已啟用：可在 /line-card-settings 設定行程分類")
except Exception as e:
    print("⚠️ 行事曆分類客製化套用失敗：", e)
# =============================================================================
# 行事曆分類客製化 Patch End
# =============================================================================


# =============================================================================
# LINE 設定中心 + 多群組權限 + 指令控管 Patch v20260622A
# - /line-card-settings 改成「設定中心」，需管理員密碼
# - 可設定多個 LINE 群組 / 房間，每個群組可接收不同類型訊息
# - 每個群組可使用的指令不同
# - 未設定的群組 / 私訊使用者不能使用指令
# =============================================================================

LINE_RECEIVE_TYPE_OPTIONS = [
    ("buyer", "客需卡片 / 客需更新"),
    ("seller", "委託卡片 / 委託更新"),
    ("development", "開發卡片 / 開發更新"),
    ("calendar", "行事曆 / 行程卡片"),
    ("todo", "待辦事項"),
]

LINE_COMMAND_TYPE_OPTIONS = [
    ("calendar", "行事曆指令：#今日行程、#明日行程、#本週行程、#新增行程"),
    ("todo", "待辦指令：#今日待辦、#待辦、#新增待辦"),
    ("buyer", "客需指令：#新增客需、#買方追蹤、#帶看、#成交"),
    ("seller", "委託指令：#新增委託、#賣方追蹤、#查詢委託到期"),
    ("development", "開發指令：#新增開發、#開發追蹤、開發自由格式"),
    ("query", "查詢指令：#查詢紀錄"),
    ("followup", "卡片回覆追蹤 / 勾選待辦 / postback 操作"),
]

LINE_SETTINGS_ADMIN_SESSION_KEY = "line_settings_admin_ok"


def _settings_admin_password_ok(password: str) -> bool:
    password = str(password or "")
    if not password:
        return False

    # 1. Firestore 內儲存的雜湊密碼優先。
    try:
        settings = get_line_card_settings()
        password_hash = settings.get("settings_admin_password_hash", "") or ""
        if password_hash and check_password_hash(password_hash, password):
            return True
    except Exception:
        pass

    # 2. 環境變數可作為備援 / 忘記密碼時覆蓋。
    env_password = (os.environ.get("LINE_SETTINGS_ADMIN_PASSWORD") or "").strip()
    if env_password and hmac.compare_digest(env_password, password):
        return True

    # 3. 第一次安裝預設密碼：123456。登入後請立即到設定中心修改。
    if not env_password:
        return hmac.compare_digest("123456", password)

    return False


def _normalize_permission_list(values):
    if not values:
        return []
    if isinstance(values, str):
        values = re.split(r"[，,、\s]+", values)
    out = []
    seen = set()
    for item in values:
        item = str(item or "").strip()
        if item and item not in seen:
            seen.add(item)
            out.append(item)
    return out


def _normalize_line_group_config(group):
    group = dict(group or {})
    return {
        "enabled": bool(group.get("enabled", True)),
        "name": (group.get("name") or "未命名群組").strip(),
        "target_id": (group.get("target_id") or group.get("group_id") or group.get("room_id") or "").strip(),
        "receive_types": _normalize_permission_list(group.get("receive_types") or group.get("receives") or []),
        "command_types": _normalize_permission_list(group.get("command_types") or group.get("commands") or []),
        "note": (group.get("note") or "").strip(),
    }


def get_line_group_settings():
    """讀取設定中心的群組清單。舊版 notify_target_id 會自動轉成一個預設群組。"""
    settings = get_line_card_settings()
    groups = []

    for g in settings.get("line_groups") or []:
        ng = _normalize_line_group_config(g)
        if ng.get("target_id"):
            groups.append(ng)

    # 舊版單一群組相容：有 notify_target_id 但還沒設定 line_groups 時，視為預設群組。
    old_target = (settings.get("notify_target_id") or settings.get("line_notify_target_id") or os.environ.get("LINE_NOTIFY_TARGET_ID") or "").strip()
    if old_target and not any(g.get("target_id") == old_target for g in groups):
        groups.append({
            "enabled": True,
            "name": "預設群組",
            "target_id": old_target,
            "receive_types": [x[0] for x in LINE_RECEIVE_TYPE_OPTIONS],
            "command_types": [x[0] for x in LINE_COMMAND_TYPE_OPTIONS],
            "note": "由舊版 LINE_NOTIFY_TARGET_ID / notify_target_id 自動帶入",
        })

    return groups


def get_enabled_line_groups():
    return [g for g in get_line_group_settings() if g.get("enabled") and g.get("target_id")]


def line_event_source_kind_and_id(event):
    source = (event or {}).get("source") or {}
    if source.get("groupId"):
        return "group", source.get("groupId")
    if source.get("roomId"):
        return "room", source.get("roomId")
    if source.get("userId"):
        return "user", source.get("userId")
    return "unknown", ""


def find_line_group_by_target_id(target_id: str):
    target_id = (target_id or "").strip()
    if not target_id:
        return None
    for g in get_enabled_line_groups():
        if (g.get("target_id") or "").strip() == target_id:
            return g
    return None


def line_group_allows_receive(group, receive_type: str) -> bool:
    if not group or not group.get("enabled"):
        return False
    allowed = set(group.get("receive_types") or [])
    return "all" in allowed or receive_type in allowed


def line_group_allows_command(group, command_type: str) -> bool:
    if not group or not group.get("enabled"):
        return False
    allowed = set(group.get("command_types") or [])
    return "all" in allowed or command_type in allowed


def detect_line_command_type(text: str, event=None) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    first = text.splitlines()[0].strip().replace(" ", "")
    first_no_hash = first[1:] if first.startswith("#") else first

    if first_no_hash in ("綁定", "群組ID", "查詢群組ID", "取得群組ID", "GroupID", "groupid"):
        return "system_group_id"

    if first_no_hash in ("今日行程", "今天行程", "明日行程", "明天行程", "本週行程", "新增行程"):
        return "calendar"
    if first_no_hash in ("今日待辦", "待辦", "新增待辦"):
        return "todo"
    if first_no_hash in ("新增客需", "買方追蹤", "帶看", "成交", "客戶分類"):
        return "buyer"
    if first_no_hash in ("新增委託", "賣方追蹤", "查詢委託到期", "委託"):
        return "seller"
    if first_no_hash in ("新增開發", "開發追蹤"):
        return "development"
    if first_no_hash in ("查詢紀錄",):
        return "query"

    # 開發自由格式通常沒有 #，給開發群組使用。
    if not first.startswith("#") and ("地址" in text or "屋主" in text or "自售" in text):
        return "development"

    return "unknown"


def detect_line_postback_command_type(event) -> str:
    try:
        data = ((event or {}).get("postback") or {}).get("data", "") or ""
        if "todo" in data:
            return "todo"
        if "followup" in data:
            return "followup"
        if "calendar" in data:
            return "calendar"
    except Exception:
        pass
    return "followup"


def line_access_gate(event):
    """檢查群組是否被授權，以及該群組是否能使用此指令。"""
    kind, target_id = line_event_source_kind_and_id(event)

    # 允許任何地方詢問群組 ID，方便設定。
    if (event.get("message") or {}).get("type") == "text":
        text = (event.get("message") or {}).get("text", "")
        if detect_line_command_type(text) == "system_group_id":
            return True, "system_group_id", None

    group = find_line_group_by_target_id(target_id)
    if not group:
        if kind == "user":
            return False, "未授權：此官方帳號不開放私人指令。請在已設定的 LINE 群組中使用。", None
        if target_id:
            return False, f"未授權：此群組尚未在後台設定。\n群組ID：{target_id}\n請到「設定 → LINE群組權限」新增此群組。", None
        return False, "未授權：無法辨識 LINE 來源。", None

    if event.get("type") == "postback":
        cmd_type = detect_line_postback_command_type(event)
    else:
        text = (event.get("message") or {}).get("text", "")
        cmd_type = detect_line_command_type(text)

    # unknown 文字不主動放行，避免陌生人加了官方帳號後亂試。
    if cmd_type == "unknown":
        return False, "這個群組沒有可辨識的指令。請使用已開放的 #指令。", group

    if cmd_type == "system_group_id":
        return True, cmd_type, group

    if not line_group_allows_command(group, cmd_type):
        return False, f"此群組未開放「{cmd_type}」類指令，請到後台設定中心調整群組權限。", group

    return True, cmd_type, group


def line_get_push_targets(receive_type: str):
    targets = []
    for g in get_enabled_line_groups():
        if line_group_allows_receive(g, receive_type):
            targets.append(g)
    return targets


def line_push_messages_to_allowed_groups(receive_type: str, messages: list):
    targets = line_get_push_targets(receive_type)
    if not targets:
        return {"ok": False, "error": f"沒有設定可接收「{receive_type}」的 LINE 群組"}

    results = []
    ok_count = 0
    for g in targets:
        res = line_push_messages(g.get("target_id"), messages)
        res["group_name"] = g.get("name")
        res["target_id"] = g.get("target_id")
        results.append(res)
        if res.get("ok"):
            ok_count += 1
    return {"ok": ok_count > 0, "sent_count": ok_count, "total_count": len(targets), "results": results}


# 覆寫舊版單一群組：依群組設定發送到所有允許接收該類型的群組。
def _push_record_to_group(record_type: str, record_id: str, title_prefix="CRM 資料"):
    coll = {"buyer": "buyers", "seller": "sellers", "development": "developments"}.get(record_type)
    if not coll:
        return {"ok": False, "error": "record_type 不正確"}
    snap = db.collection(coll).document(record_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到資料"}
    data = snap.to_dict() or {}
    bubble = _build_record_flex_bubble(record_type, record_id, data, title_prefix=title_prefix)
    return line_push_messages_to_allowed_groups(record_type, [{"type": "flex", "altText": f"{title_prefix}：{data.get('name','')}", "contents": bubble}])


def _push_calendar_event_to_group(event_id: str, title_prefix="行程資料"):
    snap = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到行程"}
    event = doc_to_calendar_event(snap)
    bubble = build_calendar_event_bubble(event)
    return line_push_messages_to_allowed_groups("calendar", [{"type": "flex", "altText": f"{title_prefix}：{event.get('title','')}", "contents": bubble}])


def _line_push_result_flash_message(res):
    if not res.get("ok"):
        return f"傳送失敗：{res.get('error') or res.get('text') or res}", "danger"
    if res.get("total_count"):
        return f"已傳送到 {res.get('sent_count', 0)}/{res.get('total_count', 0)} 個 LINE 群組", "success"
    return "已傳送到 LINE 群組", "success"


# 讓舊有 send-to-line route 的 flash 顯示多群組結果。
def buyer_send_to_line_multi_group(buyer_id):
    res = _push_record_to_group("buyer", buyer_id, title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("buyer_detail", buyer_id=buyer_id))


def seller_send_to_line_multi_group(seller_id):
    res = _push_record_to_group("seller", seller_id, title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("seller_detail", seller_id=seller_id))


def development_send_to_line_multi_group(development_id):
    res = _push_record_to_group("development", development_id, title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("development_detail", development_id=development_id))


def calendar_send_to_line_multi_group(event_id):
    res = _push_calendar_event_to_group(event_id, title_prefix="後台傳送行程")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("calendar_page"))


try:
    app.view_functions["buyer_send_to_line"] = login_required(buyer_send_to_line_multi_group)
    app.view_functions["seller_send_to_line"] = login_required(seller_send_to_line_multi_group)
    app.view_functions["development_send_to_line"] = login_required(development_send_to_line_multi_group)
    app.view_functions["calendar_send_to_line"] = login_required(calendar_send_to_line_multi_group)
except Exception as e:
    print("⚠️ 套用多群組發送 view 失敗：", e)


LINE_SETTINGS_CENTER_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8">
  <title>設定中心｜厝米 Team M.E</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>
    body{background:#f6f3ef;}
    .setting-card{background:#fff;border:1px solid #eadbca;border-radius:18px;box-shadow:0 6px 22px rgba(120,80,40,.08);}
    .section-title{font-weight:800;color:#7a4e2d;}
    .hint{font-size:.86rem;color:#777;}
    .sticky-side{position:sticky;top:20px;}
    textarea{min-height:140px;}
    .group-box{border:1px solid #eadbca;border-radius:14px;padding:14px;background:#fffaf5;}
    .code{font-family:Consolas,monospace;font-size:.85rem;}
  </style>
</head>
<body>
<nav class="navbar navbar-expand-lg navbar-dark bg-dark mb-4">
  <div class="container-fluid">
    <a class="navbar-brand" href="{{ url_for('buyers') }}">厝米 Team M.E</a>
    <div class="d-flex ms-auto flex-wrap gap-2">
      <a href="{{ url_for('buyers') }}" class="btn btn-sm btn-outline-light">客需</a>
      <a href="{{ url_for('sellers') }}" class="btn btn-sm btn-outline-light">委託</a>
      <a href="{{ url_for('developments') }}" class="btn btn-sm btn-outline-light">開發</a>
      <a href="{{ url_for('calendar_page') }}" class="btn btn-sm btn-outline-light">行事曆</a>
      <a href="{{ url_for('todos_page') }}" class="btn btn-sm btn-outline-light">待辦事項</a>
      <a href="{{ url_for('logout') }}" class="btn btn-sm btn-warning">登出</a>
    </div>
  </div>
</nav>

<div class="container mb-5">
  {% with messages = get_flashed_messages(with_categories=true) %}
    {% if messages %}
      {% for category, msg in messages %}
        <div class="alert alert-{{ category }} alert-dismissible fade show" role="alert">
          {{ msg }}
          <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
        </div>
      {% endfor %}
    {% endif %}
  {% endwith %}

  <div class="d-flex justify-content-between align-items-center mb-3 flex-wrap gap-2">
    <div>
      <h2 class="mb-1">設定中心</h2>
      <div class="text-muted small">LINE 群組、權限、卡片樣式、行事曆分類都在這裡管理。</div>
    </div>
    <form method="post" action="{{ url_for('line_settings_admin_logout') }}">
      <button class="btn btn-outline-danger btn-sm">離開管理模式</button>
    </form>
  </div>

  <form method="post">
    <div class="row g-4">
      <div class="col-lg-3">
        <div class="setting-card p-3 sticky-side">
          <div class="section-title mb-2">設定區塊</div>
          <div class="list-group small">
            <a class="list-group-item list-group-item-action" href="#security">管理員密碼</a>
            <a class="list-group-item list-group-item-action" href="#groups">LINE 群組權限</a>
            <a class="list-group-item list-group-item-action" href="#card">LINE 卡片樣式</a>
            <a class="list-group-item list-group-item-action" href="#calendar">行事曆設定</a>
            <a class="list-group-item list-group-item-action" href="#help">指令說明</a>
          </div>
          <button class="btn btn-primary w-100 mt-3" type="submit">儲存設定</button>
        </div>
      </div>

      <div class="col-lg-9">
        <div id="security" class="setting-card p-4 mb-4">
          <h4 class="section-title">管理員密碼</h4>
          <div class="hint mb-3">進入此設定頁需要輸入管理員密碼。第一次預設密碼是 <span class="code">123456</span>，建議立即修改；也可以用環境變數 <span class="code">LINE_SETTINGS_ADMIN_PASSWORD</span> 強制覆蓋。</div>
          <div class="row g-3">
            <div class="col-md-6">
              <label class="form-label">新管理員密碼</label>
              <input type="password" name="new_settings_admin_password" class="form-control" placeholder="留空則不修改">
            </div>
            <div class="col-md-6">
              <label class="form-label">再次輸入新密碼</label>
              <input type="password" name="new_settings_admin_password_confirm" class="form-control" placeholder="留空則不修改">
            </div>
          </div>
        </div>

        <div id="groups" class="setting-card p-4 mb-4">
          <h4 class="section-title">LINE 群組權限</h4>
          <div class="hint mb-3">
            只有列在這裡且啟用的群組 / 房間可以使用指令。私人聊天預設不能使用。<br>
            群組 ID 可以在群組輸入 <span class="code">#綁定</span> 取得。
          </div>

          <input type="hidden" name="group_count" value="{{ group_rows|length }}">
          {% for g in group_rows %}
            {% set idx = loop.index0 %}
            <div class="group-box mb-3">
              <div class="d-flex justify-content-between align-items-center mb-2">
                <strong>群組 {{ loop.index }}</strong>
                <div class="form-check form-switch">
                  <input class="form-check-input" type="checkbox" name="group_{{ idx }}_enabled" {% if g.enabled %}checked{% endif %}>
                  <label class="form-check-label">啟用</label>
                </div>
              </div>
              <div class="row g-2 mb-3">
                <div class="col-md-4">
                  <label class="form-label">群組名稱</label>
                  <input class="form-control" name="group_{{ idx }}_name" value="{{ g.name }}" placeholder="例如：小秘書群">
                </div>
                <div class="col-md-8">
                  <label class="form-label">群組 / 房間 ID</label>
                  <input class="form-control code" name="group_{{ idx }}_target_id" value="{{ g.target_id }}" placeholder="Cxxxxxxxx 或 Rxxxxxxxx">
                </div>
              </div>

              <div class="mb-2 fw-bold small">這個群組可以接收哪些後台推播？</div>
              <div class="row g-2 mb-3">
                {% for key, label in receive_options %}
                  <div class="col-md-6">
                    <label class="form-check">
                      <input class="form-check-input" type="checkbox" name="group_{{ idx }}_receive_types" value="{{ key }}" {% if key in g.receive_types %}checked{% endif %}>
                      <span class="form-check-label">{{ label }}</span>
                    </label>
                  </div>
                {% endfor %}
              </div>

              <div class="mb-2 fw-bold small">這個群組可以下哪些 LINE 指令？</div>
              <div class="row g-2 mb-3">
                {% for key, label in command_options %}
                  <div class="col-md-6">
                    <label class="form-check">
                      <input class="form-check-input" type="checkbox" name="group_{{ idx }}_command_types" value="{{ key }}" {% if key in g.command_types %}checked{% endif %}>
                      <span class="form-check-label">{{ label }}</span>
                    </label>
                  </div>
                {% endfor %}
              </div>

              <label class="form-label">備註</label>
              <input class="form-control" name="group_{{ idx }}_note" value="{{ g.note }}" placeholder="例如：店內公告用、業務群、測試群">
            </div>
          {% endfor %}
        </div>

        <div id="card" class="setting-card p-4 mb-4">
          <h4 class="section-title">LINE 卡片樣式</h4>
          <div class="row g-3">
            <div class="col-md-4"><label class="form-label">今天行程標題</label><input name="title_today" class="form-control" value="{{ settings.title_today }}"></div>
            <div class="col-md-4"><label class="form-label">明日行程標題</label><input name="title_tomorrow" class="form-control" value="{{ settings.title_tomorrow }}"></div>
            <div class="col-md-4"><label class="form-label">本週行程標題</label><input name="title_week" class="form-control" value="{{ settings.title_week }}"></div>
            <div class="col-md-3"><label class="form-label">主色</label><input type="color" name="primary_color" class="form-control form-control-color" value="{{ settings.primary_color or '#C9874A' }}"></div>
            <div class="col-md-3"><label class="form-label">按鈕色</label><input type="color" name="button_color" class="form-control form-control-color" value="{{ settings.button_color or '#C9874A' }}"></div>
            <div class="col-md-6">
              <label class="form-label">快速操作按鈕文字</label>
              <textarea name="quick_actions" class="form-control">{{ quick_actions_text }}</textarea>
            </div>
            <div class="col-12">
              <div class="row g-2">
                {% for key, label in [('show_phone','顯示電話'),('show_location','顯示地點'),('show_note','顯示備註'),('show_customer','顯示客戶'),('show_quick_actions','顯示快速操作'),('line_only_enabled_events','只顯示勾選 LINE 的行程')] %}
                  <div class="col-md-4"><label class="form-check"><input class="form-check-input" type="checkbox" name="{{ key }}" {% if settings.get(key) %}checked{% endif %}> {{ label }}</label></div>
                {% endfor %}
              </div>
            </div>
          </div>
        </div>

        <div id="calendar" class="setting-card p-4 mb-4">
          <h4 class="section-title">行事曆設定</h4>
          <div class="row g-3">
            <div class="col-md-3"><label class="form-label">開始小時</label><input type="number" name="calendar_start_hour" class="form-control" value="{{ settings.calendar_start_hour or 8 }}" min="0" max="23"></div>
            <div class="col-md-3"><label class="form-label">結束小時</label><input type="number" name="calendar_end_hour" class="form-control" value="{{ settings.calendar_end_hour or 22 }}" min="0" max="23"></div>
            <div class="col-md-6"></div>
            <div class="col-12">
              <label class="form-label">行事曆行程分類</label>
              <textarea name="calendar_categories" class="form-control">{{ calendar_categories_text }}</textarea>
              <div class="hint mt-1">一行一個分類，會同步用在後台新增 / 編輯行程與 LINE #新增行程。</div>
            </div>
          </div>
        </div>

        <div id="help" class="setting-card p-4 mb-4">
          <h4 class="section-title">指令說明</h4>
          <div class="row g-3 small">
            <div class="col-md-6"><strong>查群組 ID</strong><br><span class="code">#綁定</span></div>
            <div class="col-md-6"><strong>行事曆</strong><br><span class="code">#今日行程 / #明日行程 / #本週行程 / #新增行程</span></div>
            <div class="col-md-6"><strong>待辦</strong><br><span class="code">#今日待辦 / #待辦 / #新增待辦</span></div>
            <div class="col-md-6"><strong>客需 / 委託 / 開發</strong><br><span class="code">#新增客需 / #新增委託 / #新增開發</span></div>
          </div>
        </div>

        <button class="btn btn-primary btn-lg w-100" type="submit">儲存全部設定</button>
      </div>
    </div>
  </form>
</div>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""


LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8">
  <title>設定中心管理員驗證</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>body{background:#f6f3ef}.box{max-width:420px;margin:80px auto;background:#fff;border-radius:18px;padding:28px;box-shadow:0 8px 28px rgba(120,80,40,.12)}</style>
</head>
<body>
  <div class="box">
    <h3 class="mb-2">設定中心</h3>
    <p class="text-muted small">此頁可管理 LINE 群組權限與指令，請輸入管理員密碼。</p>
    {% with messages = get_flashed_messages(with_categories=true) %}
      {% if messages %}
        {% for category, msg in messages %}
          <div class="alert alert-{{ category }} py-2">{{ msg }}</div>
        {% endfor %}
      {% endif %}
    {% endwith %}
    <form method="post">
      <input type="hidden" name="admin_login" value="1">
      <div class="mb-3">
        <label class="form-label">管理員密碼</label>
        <input type="password" name="admin_password" class="form-control" autofocus required>
      </div>
      <button class="btn btn-primary w-100">進入設定中心</button>
      <a href="{{ url_for('buyers') }}" class="btn btn-link w-100 mt-2">回後台</a>
    </form>
  </div>
</body>
</html>
"""


def parse_line_groups_from_form(form):
    try:
        count = int(form.get("group_count", 0))
    except Exception:
        count = 0
    count = max(0, min(20, count))
    groups = []
    for i in range(count):
        target_id = (form.get(f"group_{i}_target_id") or "").strip()
        name = (form.get(f"group_{i}_name") or "").strip()
        # 空白列不儲存。
        if not target_id and not name:
            continue
        groups.append({
            "enabled": form.get(f"group_{i}_enabled") == "on",
            "name": name or f"群組 {i+1}",
            "target_id": target_id,
            "receive_types": form.getlist(f"group_{i}_receive_types"),
            "command_types": form.getlist(f"group_{i}_command_types"),
            "note": (form.get(f"group_{i}_note") or "").strip(),
        })
    return groups


def save_line_settings_center_from_form(form):
    # 保留既有設定儲存邏輯：卡片樣式、行事曆分類。
    updates = save_line_card_settings_from_form(form)

    groups = parse_line_groups_from_form(form)

    # 若使用者沒有填任何群組，但有舊版 notify_target_id，就保留相容。
    if not groups:
        old_target = (form.get("notify_target_id") or updates.get("notify_target_id") or os.environ.get("LINE_NOTIFY_TARGET_ID") or "").strip()
        if old_target:
            groups = [{
                "enabled": True,
                "name": "預設群組",
                "target_id": old_target,
                "receive_types": [x[0] for x in LINE_RECEIVE_TYPE_OPTIONS],
                "command_types": [x[0] for x in LINE_COMMAND_TYPE_OPTIONS],
                "note": "舊版單一群組相容",
            }]

    extra = {
        "line_groups": groups,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }

    new_pw = (form.get("new_settings_admin_password") or "").strip()
    new_pw2 = (form.get("new_settings_admin_password_confirm") or "").strip()
    if new_pw or new_pw2:
        if new_pw != new_pw2:
            raise ValueError("兩次輸入的新管理員密碼不一致")
        if len(new_pw) < 4:
            raise ValueError("管理員密碼至少 4 碼")
        extra["settings_admin_password_hash"] = generate_password_hash(new_pw)

    db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set(extra, merge=True)
    return {**updates, **extra}


def line_card_settings_center():
    if request.method == "POST" and request.form.get("admin_login") == "1":
        if _settings_admin_password_ok(request.form.get("admin_password", "")):
            session[LINE_SETTINGS_ADMIN_SESSION_KEY] = True
            flash("已進入設定中心", "success")
            return redirect(url_for("line_card_settings"))
        flash("管理員密碼錯誤", "danger")
        return render_template_string(LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE)

    if not session.get(LINE_SETTINGS_ADMIN_SESSION_KEY):
        return render_template_string(LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE)

    if request.method == "POST":
        try:
            save_line_settings_center_from_form(request.form)
            flash("設定已更新", "success")
        except Exception as e:
            flash(f"設定儲存失敗：{e}", "danger")
        return redirect(url_for("line_card_settings"))

    settings = get_line_card_settings()
    groups = get_line_group_settings()

    # 顯示固定 6 列；不足補空白列，方便新增群組。
    group_rows = list(groups)
    while len(group_rows) < 6:
        group_rows.append({
            "enabled": False,
            "name": "",
            "target_id": "",
            "receive_types": [],
            "command_types": [],
            "note": "",
        })

    return render_template_string(
        LINE_SETTINGS_CENTER_TEMPLATE,
        settings=settings,
        group_rows=group_rows,
        receive_options=LINE_RECEIVE_TYPE_OPTIONS,
        command_options=LINE_COMMAND_TYPE_OPTIONS,
        quick_actions_text="\n".join(settings.get("quick_actions") or []),
        calendar_categories_text="\n".join(get_calendar_category_options()),
    )


@app.route("/line-card-settings/logout-admin", methods=["POST"], endpoint="line_settings_admin_logout")
@login_required
def line_settings_admin_logout():
    session.pop(LINE_SETTINGS_ADMIN_SESSION_KEY, None)
    flash("已離開設定中心管理模式", "info")
    return redirect(url_for("buyers"))


try:
    app.view_functions["line_card_settings"] = login_required(line_card_settings_center)
    print("✅ 設定中心已啟用：/line-card-settings 需管理員密碼，支援多群組權限")
except Exception as e:
    print("⚠️ 設定中心套用失敗：", e)


# 覆寫 LINE webhook：加入群組/指令權限控管。
def line_webhook_with_group_acl():
    raw_body = request.get_data(cache=False, as_text=False)
    signature = request.headers.get("x-line-signature", "")

    if not verify_line_signature(raw_body, signature):
        return "Invalid signature", 400

    try:
        payload = json.loads(raw_body.decode("utf-8"))
    except Exception as e:
        print("⚠️ LINE webhook JSON 解析失敗：", e)
        return "Bad Request", 400

    events = payload.get("events", [])
    for event in events:
        try:
            reply_token = event.get("replyToken")

            # 綁定/ID查詢：任何來源都可以用，方便建立設定。
            if (event.get("message") or {}).get("type") == "text":
                text = (event.get("message") or {}).get("text", "")
                if detect_line_command_type(text) == "system_group_id":
                    kind, target_id = line_event_source_kind_and_id(event)
                    msg = f"來源類型：{kind}\nID：{target_id or '-'}\n請複製這個 ID 到後台「設定 → LINE 群組權限」。"
                    if reply_token:
                        reply_line_text(reply_token, msg)
                    continue

            allowed, reason_or_cmd, group = line_access_gate(event)
            if not allowed:
                if reply_token:
                    reply_line_text(reply_token, reason_or_cmd)
                continue

            if event.get("type") == "postback":
                result = process_line_postback_event(event)
            else:
                result = process_line_message_event(event)

            if not result or not result.get("handled"):
                continue

            reply_result = None
            if not reply_token:
                continue

            if result.get("reply_messages"):
                reply_result = reply_line_messages(reply_token, result.get("reply_messages") or [])
            elif result.get("reply_flex"):
                reply_result = reply_line_flex(
                    reply_token,
                    result.get("reply_text", "CRM 卡片"),
                    result.get("reply_flex"),
                    quick_reply_items=result.get("reply_quick_reply"),
                )
            elif result.get("reply_text"):
                reply_result = reply_line_text(
                    reply_token,
                    result["reply_text"] if result.get("ok") else result["reply_text"],
                )

            if result.get("ok") and result.get("target_type") and result.get("target_id") and reply_result:
                for sent in reply_result.get("sent_messages", []):
                    sent_id = str(sent.get("id", "")).strip()
                    if sent_id:
                        save_line_message_link(
                            sent_id,
                            result["target_type"],
                            result["target_id"],
                            tag=result.get("parsed_tag", ""),
                            action="bot_reply",
                            customer_name=result.get("customer_name", ""),
                            phone=result.get("phone", ""),
                            source_event=event,
                        )
        except Exception as e:
            print("⚠️ 處理 LINE event 發生錯誤：", e)

    return "OK", 200


try:
    app.view_functions["line_webhook"] = line_webhook_with_group_acl
    print("✅ LINE 群組權限控管已啟用：未設定群組不可使用指令")
except Exception as e:
    print("⚠️ LINE 群組權限控管套用失敗：", e)

# =============================================================================
# LINE 設定中心 + 多群組權限 + 指令控管 Patch End
# =============================================================================



# =============================================================================
# 行事曆：個人 / 公開分流 + LINE 個人帳號綁定 Patch v20260622B
# - 公開行程：傳送到設定中心中可接收 calendar 的群組
# - 個人行程：綁定 LINE 個人 userId，單獨 push 給該人
# - LINE 私訊指令：只有設定中心已綁定的 LINE 個人帳號可使用
# - 後台行事曆：公開 / 個人可分開篩選顯示
# =============================================================================

LINE_PERSONAL_USER_COUNT_DEFAULT = 8


def _normalize_line_personal_user_config(user):
    user = dict(user or {})
    return {
        "enabled": bool(user.get("enabled", True)),
        "name": (user.get("name") or user.get("display_name") or "未命名使用者").strip(),
        "user_id": (user.get("user_id") or user.get("line_user_id") or "").strip(),
        "receive_types": _normalize_permission_list(user.get("receive_types") or user.get("receives") or ["calendar"]),
        "command_types": _normalize_permission_list(user.get("command_types") or user.get("commands") or ["calendar", "todo", "followup"]),
        "note": (user.get("note") or "").strip(),
    }


def get_line_personal_users(include_disabled=False):
    settings = get_line_card_settings()
    users = []
    for u in settings.get("line_personal_users") or []:
        nu = _normalize_line_personal_user_config(u)
        if nu.get("user_id") and (include_disabled or nu.get("enabled")):
            users.append(nu)
    return users


def find_line_personal_user_by_user_id(user_id: str):
    user_id = (user_id or "").strip()
    if not user_id:
        return None
    for u in get_line_personal_users(include_disabled=False):
        if (u.get("user_id") or "").strip() == user_id:
            return u
    return None


def line_personal_user_allows_receive(user, receive_type: str) -> bool:
    if not user or not user.get("enabled"):
        return False
    allowed = set(user.get("receive_types") or [])
    return "all" in allowed or receive_type in allowed


def line_personal_user_allows_command(user, command_type: str) -> bool:
    if not user or not user.get("enabled"):
        return False
    allowed = set(user.get("command_types") or [])
    return "all" in allowed or command_type in allowed


def calendar_event_visibility_text(event):
    return "個人行程" if (event.get("visibility") or "") == "personal" else "公開行程"


def calendar_event_target_line_user_id(event):
    return (event.get("owner_line_user_id") or event.get("line_user_id") or event.get("target_user_id") or "").strip()


def calendar_event_target_line_name(event):
    return (event.get("owner_line_name") or event.get("line_user_name") or event.get("target_user_name") or "").strip()


# 覆寫 doc_to_calendar_event：補個人 LINE 綁定欄位與可見性顯示。
def doc_to_calendar_event(doc):
    data = doc_to_dict(doc)
    data["start_time"] = calendar_safe_time(data.get("start_time"), "09:00")
    data["end_time"] = calendar_safe_time(data.get("end_time"), next_30_min_time(data["start_time"]))
    data["event_date"] = calendar_safe_date(data.get("event_date"))
    data["event_date_label"] = format_calendar_date_label(data.get("event_date"))
    data["category"] = (data.get("category") or "其他").strip()
    data["custom_category"] = data.get("custom_category") or ""
    data["display_category"] = data.get("display_category") or (
        data["custom_category"] if data["category"] == "其他" and data["custom_category"] else data["category"]
    )
    data["visibility"] = (data.get("visibility") or "public").strip()
    if data["visibility"] not in ("public", "personal"):
        data["visibility"] = "public"
    data["owner_line_user_id"] = calendar_event_target_line_user_id(data)
    data["owner_line_name"] = calendar_event_target_line_name(data)
    data["visibility_text"] = calendar_event_visibility_text(data)
    data["category_color"] = data.get("category_color") or get_calendar_category_color(data["category"])
    return data


# 讀取日期區間，支援後台篩選與 LINE 來源篩選。
def fetch_calendar_events(start_date: str, end_date: str = None, line_only: bool = False, calendar_view: str = "all", line_source_kind: str = "", line_source_id: str = ""):
    start_date = calendar_safe_date(start_date)
    end_date = calendar_safe_date(end_date or start_date)
    events = []

    try:
        if start_date == end_date:
            docs = db.collection(CALENDAR_EVENT_COLLECTION).where("event_date", "==", start_date).stream()
        else:
            docs = db.collection(CALENDAR_EVENT_COLLECTION).where("event_date", ">=", start_date).where("event_date", "<=", end_date).stream()
        docs = list(docs)
    except Exception as e:
        print("⚠️ Firestore 日期區間查詢失敗，改用備援讀取：", e)
        docs = list(db.collection(CALENDAR_EVENT_COLLECTION).stream())

    current_user_id = ""
    try:
        current_user_id = session.get("user_id", "") if request else ""
    except Exception:
        current_user_id = ""

    for d in docs:
        try:
            item = doc_to_calendar_event(d)
        except Exception:
            continue
        if not (start_date <= item.get("event_date", "") <= end_date):
            continue
        if line_only and not item.get("line_enabled", True):
            continue

        vis = item.get("visibility") or "public"
        owner_uid = item.get("owner_user_id") or ""
        owner_line_uid = calendar_event_target_line_user_id(item)

        if line_only:
            # LINE 群組 / 房間只看公開行程；私人使用者看公開行程 + 綁定給自己的個人行程。
            if line_source_kind in ("group", "room"):
                if vis != "public":
                    continue
            elif line_source_kind == "user":
                if vis == "personal" and owner_line_uid != line_source_id:
                    continue
                # vis == public 仍顯示，讓個人也能看到公開行程。
            else:
                if vis != "public":
                    continue
        else:
            # 後台：公開都可見；個人只給建立者看。若 calendar_view 指定，做分流。
            if calendar_view == "public" and vis != "public":
                continue
            if calendar_view == "personal" and vis != "personal":
                continue
            if vis == "personal" and owner_uid and current_user_id and owner_uid != current_user_id:
                continue

        events.append(item)

    events.sort(key=lambda x: (x.get("event_date", ""), x.get("start_time", ""), x.get("created_at", "")))
    return events


def build_calendar_event_payload(form, existing=None):
    existing = existing or {}
    event_date = calendar_safe_date(form.get("event_date") or existing.get("event_date"))
    start_time = calendar_safe_time(form.get("start_time") or existing.get("start_time"), "09:00")
    end_time = calendar_safe_time(form.get("end_time") or existing.get("end_time"), next_30_min_time(start_time))
    if calendar_time_to_minutes(end_time) <= calendar_time_to_minutes(start_time):
        end_time = next_30_min_time(start_time)

    category_options = get_calendar_category_options()
    category = (form.get("category", "") or existing.get("category") or "其他").strip()
    if category not in category_options:
        category = "其他"

    custom_category = (form.get("custom_category", "") or existing.get("custom_category", "") or "").strip()
    display_category = custom_category if category == "其他" and custom_category else category

    visibility = (form.get("visibility", "") or existing.get("visibility") or "personal").strip()
    if visibility not in ("personal", "public"):
        visibility = "personal"

    owner_line_user_id = (form.get("owner_line_user_id") or existing.get("owner_line_user_id") or "").strip()
    owner_line_name = ""
    for u in get_line_personal_users(include_disabled=True):
        if u.get("user_id") == owner_line_user_id:
            owner_line_name = u.get("name") or ""
            break
    if not owner_line_name:
        owner_line_name = (form.get("owner_line_name") or existing.get("owner_line_name") or "").strip()

    payload = {
        "title": (form.get("title", "") or "").strip(),
        "event_date": event_date,
        "start_time": start_time,
        "end_time": end_time,
        "category": category,
        "custom_category": custom_category,
        "display_category": display_category,
        "visibility": visibility,
        "owner_user_id": existing.get("owner_user_id") or session.get("user_id", ""),
        "owner_user_name": existing.get("owner_user_name") or session.get("user_name", ""),
        "owner_line_user_id": owner_line_user_id if visibility == "personal" else "",
        "owner_line_name": owner_line_name if visibility == "personal" else "",
        "category_color": get_calendar_category_color(category),
        "related_type": (form.get("related_type", "") or "").strip(),
        "related_id": (form.get("related_id", "") or "").strip(),
        "customer_name": (form.get("customer_name", "") or "").strip(),
        "phone": (form.get("phone", "") or "").strip(),
        "location": (form.get("location", "") or "").strip(),
        "note": (form.get("note", "") or "").strip(),
        "line_enabled": form.get("line_enabled") == "on",
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }
    if not payload["title"]:
        name = payload.get("customer_name") or payload.get("location") or category
        payload["title"] = f"{name} {category}".strip()
    return payload


def calendar_page_personal_public():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    selected_date_label = format_calendar_date_label(selected_date)
    calendar_view = (request.args.get("view") or "all").strip()
    if calendar_view not in ("all", "public", "personal"):
        calendar_view = "all"

    events = fetch_calendar_events(selected_date, calendar_view=calendar_view)
    public_events = [e for e in events if (e.get("visibility") or "public") == "public"]
    personal_events = [e for e in events if (e.get("visibility") or "public") == "personal"]

    slots = build_30_min_slots()
    slot_cells = build_calendar_slot_cells(events, slots)
    event_map = {}
    for e in events:
        event_map.setdefault(e.get("start_time"), []).append(e)

    dates = calendar_prev_next_dates(selected_date)
    return render_template(
        "calendar.html",
        selected_date=selected_date,
        selected_date_label=selected_date_label,
        calendar_view=calendar_view,
        slots=slots,
        slot_cells=slot_cells,
        event_map=event_map,
        events=events,
        public_events=public_events,
        personal_events=personal_events,
        category_options=get_calendar_category_options(),
        **dates,
    )


def calendar_new_personal_public():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    default_start = calendar_safe_time(request.args.get("start", ""), "09:00")
    default_end = calendar_safe_time(request.args.get("end", ""), next_30_min_time(default_start))

    if request.method == "POST":
        payload = build_calendar_event_payload(request.form)
        payload.update({
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id", ""),
            "created_by_name": session.get("user_name", ""),
        })
        doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).add(payload)[1]
        if request.form.get("send_line_now") == "on":
            try:
                _push_calendar_event_to_group(doc_ref.id, title_prefix="新增行程")
            except Exception as e:
                print("⚠️ 新增行程推播失敗：", e)
        flash("已新增行程", "success")
        return redirect(url_for("calendar_page", date=payload["event_date"], view=payload.get("visibility") or "all"))

    prefill = {}
    for key in ("title", "category", "custom_category", "related_type", "related_id", "customer_name", "phone", "location", "note", "visibility", "owner_line_user_id"):
        value = (request.args.get(key) or "").strip()
        if value:
            prefill[key] = value
    if prefill:
        prefill["event_date"] = selected_date
        prefill["start_time"] = default_start
        prefill["end_time"] = default_end
        prefill["line_enabled"] = True

    return render_template(
        "calendar_form.html",
        event=prefill if prefill else None,
        selected_date=selected_date,
        default_start=default_start,
        default_end=default_end,
        slots=build_30_min_slots(),
        category_options=get_calendar_category_options(),
        line_user_options=get_line_personal_users(include_disabled=True),
    )


def calendar_edit_personal_public(event_id):
    doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id)
    doc = doc_ref.get()
    if not doc.exists:
        flash("找不到這筆行程", "danger")
        return redirect(url_for("calendar_page"))

    event = doc_to_calendar_event(doc)
    if request.method == "POST":
        payload = build_calendar_event_payload(request.form, existing=event)
        doc_ref.update(payload)
        if request.form.get("send_line_now") == "on":
            try:
                _push_calendar_event_to_group(event_id, title_prefix="更新行程")
            except Exception as e:
                print("⚠️ 更新行程推播失敗：", e)
        flash("已更新行程", "success")
        return redirect(url_for("calendar_page", date=payload["event_date"], view=payload.get("visibility") or "all"))

    return render_template(
        "calendar_form.html",
        event=event,
        selected_date=event.get("event_date"),
        default_start=event.get("start_time", "09:00"),
        default_end=event.get("end_time", "09:30"),
        slots=build_30_min_slots(),
        category_options=get_calendar_category_options(),
        line_user_options=get_line_personal_users(include_disabled=True),
    )


# LINE 行程卡片：顯示個人/公開與綁定對象。
_original_build_calendar_event_bubble_before_personal_public = build_calendar_event_bubble


def build_calendar_event_bubble(event, settings=None):
    bubble = _original_build_calendar_event_bubble_before_personal_public(event, settings)
    try:
        vis_text = calendar_event_visibility_text(event)
        line_name = calendar_event_target_line_name(event)
        body = bubble.get("body", {}).get("contents", [])
        # 加在標題與資訊間：LINE 手機上可清楚看到公開 / 個人。
        insert_at = 3 if len(body) >= 3 else len(body)
        extra_text = f"{vis_text}" + (f"｜{line_name}" if line_name and event.get("visibility") == "personal" else "")
        body.insert(insert_at, flex_text(extra_text, size="xs", color="#8b6b4f", weight="bold", margin="sm"))
    except Exception:
        pass
    return bubble


# LINE 查詢行程：群組只看公開；個人已綁定帳號可看公開 + 自己的個人行程。
def build_calendar_reply_for_range(start_date, end_date=None, mode="today", event=None):
    settings = get_line_card_settings()
    line_only = bool(settings.get("line_only_enabled_events", False))
    kind = ""
    target_id = ""
    if event:
        kind, target_id = line_event_source_kind_and_id(event)
    events = fetch_calendar_events(
        start_date,
        end_date or start_date,
        line_only=line_only,
        line_source_kind=kind,
        line_source_id=target_id,
    )

    if mode == "week":
        title = settings.get("title_week") or DEFAULT_LINE_CARD_SETTINGS["title_week"]
        date_text = f"{format_calendar_date_label(start_date)} ~ {format_calendar_date_label(end_date)}"
    elif mode == "tomorrow":
        title = settings.get("title_tomorrow") or DEFAULT_LINE_CARD_SETTINGS["title_tomorrow"]
        date_text = format_calendar_date_label(start_date)
    else:
        title = settings.get("title_today") or DEFAULT_LINE_CARD_SETTINGS["title_today"]
        date_text = format_calendar_date_label(start_date)

    flex = build_calendar_carousel(events, title=title, date_text=date_text, settings=settings)
    alt = f"{title} {len(events)} 筆"
    return {
        "handled": True,
        "ok": True,
        "reply_text": alt,
        "reply_flex": flex,
        "reply_quick_reply": build_calendar_quick_reply(settings),
        "parsed_tag": "行事曆",
    }


def process_line_calendar_message_event(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}

    raw_text = (message.get("text") or "").strip()
    normalized = raw_text.replace("＃", "#").strip()
    today = now_taipei().date()

    if normalized in ("#今日行程", "#今天行程", "今日行程", "今天行程"):
        return build_calendar_reply_for_range(today.strftime("%Y-%m-%d"), mode="today", event=event)

    if normalized in ("#明日行程", "#明天行程", "明日行程", "明天行程"):
        return build_calendar_reply_for_range((today + timedelta(days=1)).strftime("%Y-%m-%d"), mode="tomorrow", event=event)

    if normalized in ("#本週行程", "本週行程"):
        start = today.strftime("%Y-%m-%d")
        end = (today + timedelta(days=6)).strftime("%Y-%m-%d")
        return build_calendar_reply_for_range(start, end, mode="week", event=event)

    if normalized.startswith("#新增行程"):
        fields = parse_line_calendar_create_fields(normalized)
        category_options = get_calendar_category_options()
        if len(normalized.splitlines()) == 1:
            sample_category = category_options[0] if category_options else "帶看"
            example = (
                "新增行程格式：\n"
                "#新增行程\n"
                "日期: 今天\n"
                "時間: 10:00-10:30\n"
                f"類型: {sample_category}\n"
                "可見性: 公開\n"
                "標題: 童先生看農舍\n"
                "客戶: 童先生\n"
                "電話: 0921-123-456\n"
                "地點: 清水、梧棲交界\n"
                "備註: 退休夫妻，想看農舍、有空地\n\n"
                "若要新增個人行程，請改成：可見性: 個人"
            )
            return {"handled": True, "ok": True, "reply_text": example, "parsed_tag": "新增行程格式"}

        event_date = parse_calendar_date_word(fields.get("event_date_raw", "今天"))
        start_time, end_time = parse_calendar_time_range(fields.get("time_raw", "09:00"), fields.get("end_time", ""))
        category = (fields.get("category") or "其他").strip()
        if category not in category_options:
            category = "其他"

        title = fields.get("title", "").strip()
        if not title:
            title = f"{fields.get('customer_name', '')} {category}".strip() or category

        kind, target_id = line_event_source_kind_and_id(event)
        visibility_raw = (fields.get("可見性") or fields.get("visibility") or fields.get("行程可見性") or "公開").strip()
        visibility = "personal" if visibility_raw in ("個人", "私人", "personal") else "public"
        owner_line_user_id = ""
        owner_line_name = ""
        if visibility == "personal" and kind == "user":
            user_cfg = find_line_personal_user_by_user_id(target_id)
            owner_line_user_id = target_id
            owner_line_name = (user_cfg or {}).get("name") or get_line_sender_display_name(event) or ""

        payload = {
            "title": title,
            "event_date": event_date,
            "start_time": start_time,
            "end_time": end_time,
            "category": category,
            "display_category": category,
            "category_color": get_calendar_category_color(category),
            "related_type": fields.get("related_type", ""),
            "related_id": fields.get("related_id", ""),
            "customer_name": fields.get("customer_name", ""),
            "phone": fields.get("phone", ""),
            "location": fields.get("location", ""),
            "note": fields.get("note", ""),
            "visibility": visibility,
            "owner_line_user_id": owner_line_user_id,
            "owner_line_name": owner_line_name,
            "line_enabled": True,
            "created_at": now_taipei().isoformat(),
            "created_by_id": "line_bot",
            "created_by_name": get_line_sender_display_name(event) or "LINE Bot",
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": "line_bot",
            "updated_by_name": get_line_sender_display_name(event) or "LINE Bot",
        }
        doc_ref = db.collection(CALENDAR_EVENT_COLLECTION).document()
        doc_ref.set(payload)
        payload["id"] = doc_ref.id
        settings = get_line_card_settings()
        flex = build_calendar_event_bubble(dict(payload), settings)
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"已新增行程：{title}（{format_calendar_date_label(event_date)} {start_time}）",
            "reply_flex": flex,
            "reply_quick_reply": build_calendar_quick_reply(settings),
            "parsed_tag": "新增行程",
        }

    return {"handled": False}


# 個人行程推送：個人行程送到綁定的 LINE userId；公開行程送到可接收 calendar 的群組。
def _push_calendar_event_to_group(event_id: str, title_prefix="行程資料"):
    snap = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到行程"}
    event = doc_to_calendar_event(snap)
    bubble = build_calendar_event_bubble(event)
    message = {"type": "flex", "altText": f"{title_prefix}：{event.get('title','')}", "contents": bubble}

    if (event.get("visibility") or "public") == "personal":
        user_id = calendar_event_target_line_user_id(event)
        if not user_id:
            return {"ok": False, "error": "此個人行程尚未綁定 LINE 個人帳號，無法單獨傳送"}
        user_cfg = find_line_personal_user_by_user_id(user_id)
        if user_cfg and not line_personal_user_allows_receive(user_cfg, "calendar"):
            return {"ok": False, "error": "此 LINE 個人帳號未開放接收行事曆訊息"}
        res = line_push_messages(user_id, [message])
        res["personal_target_name"] = calendar_event_target_line_name(event)
        return res

    return line_push_messages_to_allowed_groups("calendar", [message])


# 私訊權限：只有設定中心已綁定的 LINE 個人帳號可以使用私人指令；群組仍照原本群組權限。
def line_access_gate(event):
    kind, target_id = line_event_source_kind_and_id(event)

    if (event.get("message") or {}).get("type") == "text":
        text = (event.get("message") or {}).get("text", "")
        if detect_line_command_type(text) == "system_group_id":
            return True, "system_group_id", None

    if event.get("type") == "postback":
        cmd_type = detect_line_postback_command_type(event)
    else:
        text = (event.get("message") or {}).get("text", "")
        cmd_type = detect_line_command_type(text)

    if kind == "user":
        user_cfg = find_line_personal_user_by_user_id(target_id)
        if not user_cfg:
            return False, "未授權：此 LINE 個人帳號尚未在設定中心綁定，無法使用私人指令。請先輸入 #綁定 取得個人 ID 後交給管理員設定。", None
        if cmd_type == "unknown":
            return False, "這個指令無法辨識，請使用已開放的 #指令。", user_cfg
        if cmd_type == "system_group_id":
            return True, cmd_type, user_cfg
        if not line_personal_user_allows_command(user_cfg, cmd_type):
            return False, f"你的個人帳號未開放「{cmd_type}」類指令，請到設定中心調整個人權限。", user_cfg
        return True, cmd_type, user_cfg

    group = find_line_group_by_target_id(target_id)
    if not group:
        if target_id:
            return False, f"未授權：此群組尚未在後台設定。\n群組ID：{target_id}\n請到「設定 → LINE群組權限」新增此群組。", None
        return False, "未授權：無法辨識 LINE 來源。", None

    if cmd_type == "unknown":
        return False, "這個群組沒有可辨識的指令。請使用已開放的 #指令。", group
    if cmd_type == "system_group_id":
        return True, cmd_type, group
    if not line_group_allows_command(group, cmd_type):
        return False, f"此群組未開放「{cmd_type}」類指令，請到後台設定中心調整群組權限。", group
    return True, cmd_type, group


# 取得/綁定 ID 指令文案：同一指令可查群組 ID 或個人 userId。
def _line_source_id_reply_text(event):
    kind, target_id = line_event_source_kind_and_id(event)
    label = {"group": "群組ID", "room": "聊天室ID", "user": "個人LINE ID"}.get(kind, "來源ID")
    extra = ""
    if kind == "user":
        extra = "\n用途：可貼到設定中心的「LINE 個人帳號綁定」，用來接收個人行程。"
    else:
        extra = "\n用途：可貼到設定中心的「LINE 群組權限」。"
    return f"來源類型：{kind}\n{label}：{target_id or '-'}{extra}"


# 覆寫 webhook 一小段，讓 #綁定 私訊時文案顯示「個人LINE ID」。
def line_webhook_with_group_acl():
    raw_body = request.get_data(cache=False, as_text=False)
    signature = request.headers.get("x-line-signature", "")

    if not verify_line_signature(raw_body, signature):
        return "Invalid signature", 400

    try:
        payload = json.loads(raw_body.decode("utf-8"))
    except Exception as e:
        print("⚠️ LINE webhook JSON 解析失敗：", e)
        return "Bad Request", 400

    events = payload.get("events", [])
    for event in events:
        try:
            reply_token = event.get("replyToken")

            if (event.get("message") or {}).get("type") == "text":
                text = (event.get("message") or {}).get("text", "")
                if detect_line_command_type(text) == "system_group_id":
                    if reply_token:
                        reply_line_text(reply_token, _line_source_id_reply_text(event))
                    continue

            allowed, reason_or_cmd, source_cfg = line_access_gate(event)
            if not allowed:
                if reply_token:
                    reply_line_text(reply_token, reason_or_cmd)
                continue

            if event.get("type") == "postback":
                result = process_line_postback_event(event)
            else:
                result = process_line_message_event(event)

            if not result or not result.get("handled"):
                continue

            reply_result = None
            if not reply_token:
                continue

            if result.get("reply_messages"):
                reply_result = reply_line_messages(reply_token, result.get("reply_messages") or [])
            elif result.get("reply_flex"):
                reply_result = reply_line_flex(
                    reply_token,
                    result.get("reply_text", "CRM 卡片"),
                    result.get("reply_flex"),
                    quick_reply_items=result.get("reply_quick_reply"),
                )
            elif result.get("reply_text"):
                reply_result = reply_line_text(reply_token, result["reply_text"] if result.get("ok") else result["reply_text"])

            if result.get("ok") and result.get("target_type") and result.get("target_id") and reply_result:
                for sent in reply_result.get("sent_messages", []):
                    sent_id = str(sent.get("id", "")).strip()
                    if sent_id:
                        save_line_message_link(
                            sent_id,
                            result["target_type"],
                            result["target_id"],
                            tag=result.get("parsed_tag", ""),
                            action="bot_reply",
                            customer_name=result.get("customer_name", ""),
                            phone=result.get("phone", ""),
                            source_event=event,
                        )
        except Exception as e:
            print("⚠️ 處理 LINE event 發生錯誤：", e)

    return "OK", 200


# 設定中心：加入 LINE 個人帳號綁定區塊。
def parse_line_personal_users_from_form(form):
    try:
        count = int(form.get("personal_user_count", 0))
    except Exception:
        count = 0
    count = max(0, min(30, count))
    users = []
    for i in range(count):
        user_id = (form.get(f"personal_{i}_user_id") or "").strip()
        name = (form.get(f"personal_{i}_name") or "").strip()
        if not user_id and not name:
            continue
        users.append({
            "enabled": form.get(f"personal_{i}_enabled") == "on",
            "name": name or f"個人 {i+1}",
            "user_id": user_id,
            "receive_types": form.getlist(f"personal_{i}_receive_types"),
            "command_types": form.getlist(f"personal_{i}_command_types"),
            "note": (form.get(f"personal_{i}_note") or "").strip(),
        })
    return users


def save_line_settings_center_from_form(form):
    updates = save_line_card_settings_from_form(form)
    groups = parse_line_groups_from_form(form)
    personal_users = parse_line_personal_users_from_form(form)

    if not groups:
        old_target = (form.get("notify_target_id") or updates.get("notify_target_id") or os.environ.get("LINE_NOTIFY_TARGET_ID") or "").strip()
        if old_target:
            groups = [{
                "enabled": True,
                "name": "預設群組",
                "target_id": old_target,
                "receive_types": [x[0] for x in LINE_RECEIVE_TYPE_OPTIONS],
                "command_types": [x[0] for x in LINE_COMMAND_TYPE_OPTIONS],
                "note": "舊版單一群組相容",
            }]

    extra = {
        "line_groups": groups,
        "line_personal_users": personal_users,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }

    new_pw = (form.get("new_settings_admin_password") or "").strip()
    new_pw2 = (form.get("new_settings_admin_password_confirm") or "").strip()
    if new_pw or new_pw2:
        if new_pw != new_pw2:
            raise ValueError("兩次輸入的新管理員密碼不一致")
        if len(new_pw) < 4:
            raise ValueError("管理員密碼至少 4 碼")
        extra["settings_admin_password_hash"] = generate_password_hash(new_pw)

    db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set(extra, merge=True)
    return {**updates, **extra}


LINE_SETTINGS_CENTER_TEMPLATE_V2 = """
<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8">
  <title>設定中心｜厝米 Team M.E</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>
    body{background:#f6f3ef}.setting-card{background:#fff;border:1px solid #eadbca;border-radius:18px;box-shadow:0 8px 24px rgba(120,80,40,.08)}
    .group-box,.user-box{border:1px solid #eadbca;border-radius:14px;padding:14px;background:#fffaf5}.code{font-family:Consolas,monospace}.hint{font-size:.88rem;color:#8b6b4f;line-height:1.7}.sticky-nav{position:sticky;top:16px}
  </style>
</head>
<body>
<div class="container-fluid py-4">
  <div class="d-flex justify-content-between align-items-center mb-3 flex-wrap gap-2">
    <div><h2 class="mb-1">設定中心</h2><div class="text-muted small">管理 LINE 群組、個人帳號、指令權限與行事曆卡片。</div></div>
    <div class="d-flex gap-2">
      <a href="{{ url_for('buyers') }}" class="btn btn-outline-secondary btn-sm">回後台</a>
      <form method="post" action="{{ url_for('line_settings_admin_logout') }}"><button class="btn btn-outline-danger btn-sm">離開管理模式</button></form>
    </div>
  </div>
  {% with messages = get_flashed_messages(with_categories=true) %}{% if messages %}{% for category, msg in messages %}<div class="alert alert-{{ category }} py-2">{{ msg }}</div>{% endfor %}{% endif %}{% endwith %}
  <form method="post">
    <div class="row g-3">
      <div class="col-lg-3">
        <div class="setting-card p-3 sticky-nav">
          <div class="list-group small">
            <a class="list-group-item list-group-item-action" href="#security">管理員密碼</a>
            <a class="list-group-item list-group-item-action" href="#groups">LINE 群組權限</a>
            <a class="list-group-item list-group-item-action" href="#personal">LINE 個人帳號綁定</a>
            <a class="list-group-item list-group-item-action" href="#card">LINE 卡片樣式</a>
            <a class="list-group-item list-group-item-action" href="#calendar">行事曆設定</a>
            <a class="list-group-item list-group-item-action" href="#help">指令說明</a>
          </div>
        </div>
      </div>
      <div class="col-lg-9">
        <div id="security" class="setting-card p-4 mb-4">
          <h5>管理員密碼</h5>
          <div class="hint mb-3">第一次預設密碼是 <span class="code">123456</span>，建議修改；忘記時可用環境變數 <span class="code">LINE_SETTINGS_ADMIN_PASSWORD</span> 覆蓋。</div>
          <div class="row g-2">
            <div class="col-md-6"><label class="form-label">新密碼</label><input type="password" name="new_settings_admin_password" class="form-control" placeholder="留空則不修改"></div>
            <div class="col-md-6"><label class="form-label">再次輸入</label><input type="password" name="new_settings_admin_password_confirm" class="form-control" placeholder="留空則不修改"></div>
          </div>
        </div>
        <div id="groups" class="setting-card p-4 mb-4">
          <h5>LINE 群組權限</h5>
          <div class="hint mb-3">群組輸入 <span class="code">#綁定</span> 可取得/綁定 ID。公開行程會傳送到有勾選「行事曆」接收權限的群組。</div>
          <input type="hidden" name="group_count" value="{{ group_rows|length }}">
          {% for g in group_rows %}{% set idx = loop.index0 %}
          <div class="group-box mb-3">
            <div class="row g-2 align-items-end">
              <div class="col-md-1"><label class="form-label">啟用</label><input class="form-check-input d-block" type="checkbox" name="group_{{ idx }}_enabled" {% if g.enabled %}checked{% endif %}></div>
              <div class="col-md-3"><label class="form-label">群組名稱</label><input class="form-control" name="group_{{ idx }}_name" value="{{ g.name }}" placeholder="例如：業務群"></div>
              <div class="col-md-5"><label class="form-label">群組 / 房間 ID</label><input class="form-control code" name="group_{{ idx }}_target_id" value="{{ g.target_id }}" placeholder="Cxxxxxxxx 或 Rxxxxxxxx"></div>
              <div class="col-md-3"><label class="form-label">備註</label><input class="form-control" name="group_{{ idx }}_note" value="{{ g.note }}"></div>
            </div>
            <div class="row mt-2">
              <div class="col-md-6"><div class="fw-bold small mb-1">可接收訊息</div>{% for key,label in receive_options %}<label class="form-check form-check-inline small"><input class="form-check-input" type="checkbox" name="group_{{ idx }}_receive_types" value="{{ key }}" {% if key in g.receive_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div>
              <div class="col-md-6"><div class="fw-bold small mb-1">可使用指令</div>{% for key,label in command_options %}<label class="form-check d-block small"><input class="form-check-input" type="checkbox" name="group_{{ idx }}_command_types" value="{{ key }}" {% if key in g.command_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div>
            </div>
          </div>
          {% endfor %}
        </div>
        <div id="personal" class="setting-card p-4 mb-4">
          <h5>LINE 個人帳號綁定</h5>
          <div class="hint mb-3">個人私訊官方帳號輸入 <span class="code">#綁定</span>，會回覆「個人LINE ID」。把該 ID 貼到這裡，個人行程就可以單獨推送給他。</div>
          <input type="hidden" name="personal_user_count" value="{{ personal_user_rows|length }}">
          {% for u in personal_user_rows %}{% set idx = loop.index0 %}
          <div class="user-box mb-3">
            <div class="row g-2 align-items-end">
              <div class="col-md-1"><label class="form-label">啟用</label><input class="form-check-input d-block" type="checkbox" name="personal_{{ idx }}_enabled" {% if u.enabled %}checked{% endif %}></div>
              <div class="col-md-3"><label class="form-label">姓名 / 暱稱</label><input class="form-control" name="personal_{{ idx }}_name" value="{{ u.name }}" placeholder="例如：曜昀"></div>
              <div class="col-md-5"><label class="form-label">LINE 個人 userId</label><input class="form-control code" name="personal_{{ idx }}_user_id" value="{{ u.user_id }}" placeholder="Uxxxxxxxx"></div>
              <div class="col-md-3"><label class="form-label">備註</label><input class="form-control" name="personal_{{ idx }}_note" value="{{ u.note }}"></div>
            </div>
            <div class="row mt-2">
              <div class="col-md-6"><div class="fw-bold small mb-1">可接收訊息</div>{% for key,label in receive_options %}<label class="form-check form-check-inline small"><input class="form-check-input" type="checkbox" name="personal_{{ idx }}_receive_types" value="{{ key }}" {% if key in u.receive_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div>
              <div class="col-md-6"><div class="fw-bold small mb-1">可使用指令</div>{% for key,label in command_options %}<label class="form-check d-block small"><input class="form-check-input" type="checkbox" name="personal_{{ idx }}_command_types" value="{{ key }}" {% if key in u.command_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div>
            </div>
          </div>
          {% endfor %}
        </div>
        <div id="card" class="setting-card p-4 mb-4">
          <h5>LINE 卡片樣式</h5>
          <div class="row g-2">
            <div class="col-12"><label class="form-label">今日行程標題</label><input class="form-control" name="title_today" value="{{ settings.title_today }}"></div>
            <div class="col-12"><label class="form-label">明日行程標題</label><input class="form-control" name="title_tomorrow" value="{{ settings.title_tomorrow }}"></div>
            <div class="col-12"><label class="form-label">本週行程標題</label><input class="form-control" name="title_week" value="{{ settings.title_week }}"></div>
            <div class="col-md-6"><label class="form-label">主色</label><input class="form-control" name="primary_color" value="{{ settings.primary_color }}"></div>
            <div class="col-md-6"><label class="form-label">按鈕色</label><input class="form-control" name="button_color" value="{{ settings.button_color }}"></div>
            <div class="col-12">
              <label class="form-label">LINE 卡片顯示內容</label><br>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_customer" {% if settings.show_customer %}checked{% endif %}> 客戶</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_phone" {% if settings.show_phone %}checked{% endif %}> 電話</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_location" {% if settings.show_location %}checked{% endif %}> 地點</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_note" {% if settings.show_note %}checked{% endif %}> 備註</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_quick_actions" {% if settings.show_quick_actions %}checked{% endif %}> 快速操作</label>
              <label class="form-check d-block mt-2"><input class="form-check-input" type="checkbox" name="line_only_enabled_events" {% if settings.line_only_enabled_events %}checked{% endif %}> LINE 只顯示有勾「LINE查詢顯示」的行程</label>
            </div>
            <div class="col-12"><label class="form-label">快速操作按鈕文字</label><textarea class="form-control" rows="4" name="quick_actions">{{ quick_actions_text }}</textarea></div>
          </div>
        </div>
        <div id="calendar" class="setting-card p-4 mb-4">
          <h5>行事曆設定</h5>
          <div class="row g-2">
            <div class="col-md-6"><label class="form-label">行事曆開始小時</label><input class="form-control" name="calendar_start_hour" value="{{ settings.calendar_start_hour }}"></div>
            <div class="col-md-6"><label class="form-label">行事曆結束小時</label><input class="form-control" name="calendar_end_hour" value="{{ settings.calendar_end_hour }}"></div>
            <div class="col-12"><label class="form-label">行事曆行程分類</label><textarea class="form-control" rows="5" name="calendar_categories">{{ calendar_categories_text }}</textarea><div class="hint">一行一個分類，例如：帶看、回電、開發、簽約、拍照、私人行程、其他。</div></div>
          </div>
        </div>
        <div id="help" class="setting-card p-4 mb-4"><h5>指令說明</h5><div class="row g-2 small"><div class="col-md-6"><strong>取得/綁定 ID</strong><br><span class="code">#綁定</span><br>群組中回覆群組ID；私訊中回覆個人LINE ID。</div><div class="col-md-6"><strong>行程</strong><br><span class="code">#今日行程 / #新增行程</span><br>群組只看公開行程；個人可看公開 + 自己的個人行程。</div></div></div>
        <button class="btn btn-primary btn-lg w-100" type="submit">儲存全部設定</button>
      </div>
    </div>
  </form>
</div>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""


def line_card_settings_center():
    if request.method == "POST" and request.form.get("admin_login") == "1":
        if _settings_admin_password_ok(request.form.get("admin_password", "")):
            session[LINE_SETTINGS_ADMIN_SESSION_KEY] = True
            flash("已進入設定中心", "success")
            return redirect(url_for("line_card_settings"))
        flash("管理員密碼錯誤", "danger")
        return render_template_string(LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE)

    if not session.get(LINE_SETTINGS_ADMIN_SESSION_KEY):
        return render_template_string(LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE)

    if request.method == "POST":
        try:
            save_line_settings_center_from_form(request.form)
            flash("設定已更新", "success")
        except Exception as e:
            flash(f"設定儲存失敗：{e}", "danger")
        return redirect(url_for("line_card_settings"))

    settings = get_line_card_settings()
    group_rows = list(get_line_group_settings())
    while len(group_rows) < 6:
        group_rows.append({"enabled": False, "name": "", "target_id": "", "receive_types": [], "command_types": [], "note": ""})

    personal_user_rows = list(get_line_personal_users(include_disabled=True))
    while len(personal_user_rows) < LINE_PERSONAL_USER_COUNT_DEFAULT:
        personal_user_rows.append({"enabled": False, "name": "", "user_id": "", "receive_types": ["calendar"], "command_types": ["calendar", "todo", "followup"], "note": ""})

    return render_template_string(
        LINE_SETTINGS_CENTER_TEMPLATE_V2,
        settings=settings,
        group_rows=group_rows,
        personal_user_rows=personal_user_rows,
        receive_options=LINE_RECEIVE_TYPE_OPTIONS,
        command_options=LINE_COMMAND_TYPE_OPTIONS,
        quick_actions_text="\n".join(settings.get("quick_actions") or []),
        calendar_categories_text="\n".join(get_calendar_category_options()),
    )


try:
    app.view_functions["calendar_page"] = login_required(calendar_page_personal_public)
    app.view_functions["calendar_new"] = login_required(calendar_new_personal_public)
    app.view_functions["calendar_edit"] = login_required(calendar_edit_personal_public)
    app.view_functions["line_card_settings"] = login_required(line_card_settings_center)
    app.view_functions["line_webhook"] = line_webhook_with_group_acl
    print("✅ 行事曆個人/公開分流已啟用：個人行程可綁定 LINE 個人帳號並單獨推送")
except Exception as e:
    print("⚠️ 行事曆個人/公開分流套用失敗：", e)

# =============================================================================
# 行事曆：個人 / 公開分流 + LINE 個人帳號綁定 Patch End
# =============================================================================


# =============================================================================
# 個人後台：點擊右上角使用者名稱綁定 LINE 個人帳號 Patch v20260622C
# - 不再需要到設定中心新增個人帳號；每位後台使用者可在自己的個人後台綁定 LINE userId
# - 綁定後會同步寫入 users 文件與 line_card_settings.default.line_personal_users
# - 行事曆個人行程可直接選擇已綁定的 LINE 個人帳號並單獨推送
# =============================================================================

DEFAULT_PERSONAL_RECEIVE_TYPES = ["calendar"]
DEFAULT_PERSONAL_COMMAND_TYPES = ["calendar", "todo", "followup"]


def _current_crm_user_doc():
    user_id = session.get("user_id", "")
    if not user_id:
        return None, {}
    try:
        doc_ref = db.collection("users").document(user_id)
        snap = doc_ref.get()
        if snap.exists:
            data = snap.to_dict() or {}
        else:
            data = {}
        return doc_ref, data
    except Exception as e:
        print("⚠️ 讀取個人使用者資料失敗：", e)
        return None, {}


def _find_personal_user_raw_by_crm_user_id(settings, crm_user_id):
    crm_user_id = (crm_user_id or "").strip()
    if not crm_user_id:
        return None
    for item in settings.get("line_personal_users") or []:
        if (item.get("crm_user_id") or "").strip() == crm_user_id:
            return item
    return None


def get_current_user_line_binding():
    """取得目前登入者的 LINE 綁定資料，優先讀 users 文件，沒有就從設定中心舊資料找。"""
    user_id = session.get("user_id", "")
    user_name = session.get("user_name", "") or "未命名使用者"
    user_email = session.get("user_email", "") or ""

    _, user_doc = _current_crm_user_doc()
    settings = get_line_card_settings()
    raw = _find_personal_user_raw_by_crm_user_id(settings, user_id) or {}

    line_user_id = (
        user_doc.get("line_user_id")
        or user_doc.get("line_personal_user_id")
        or raw.get("user_id")
        or raw.get("line_user_id")
        or ""
    ).strip()

    return {
        "crm_user_id": user_id,
        "crm_user_name": user_name,
        "crm_user_email": user_email,
        "enabled": bool(user_doc.get("line_bind_enabled", raw.get("enabled", True))),
        "name": (user_doc.get("line_display_name") or raw.get("name") or user_name or "未命名使用者").strip(),
        "user_id": line_user_id,
        "receive_types": _normalize_permission_list(user_doc.get("line_receive_types") or raw.get("receive_types") or DEFAULT_PERSONAL_RECEIVE_TYPES),
        "command_types": _normalize_permission_list(user_doc.get("line_command_types") or raw.get("command_types") or DEFAULT_PERSONAL_COMMAND_TYPES),
        "note": (user_doc.get("line_bind_note") or raw.get("note") or "").strip(),
        "updated_at": user_doc.get("line_bound_at") or raw.get("updated_at") or "",
    }


def upsert_current_user_line_binding(form):
    user_id = session.get("user_id", "")
    user_name = session.get("user_name", "") or "未命名使用者"
    user_email = session.get("user_email", "") or ""
    if not user_id:
        raise ValueError("請先登入")

    line_user_id = (form.get("line_user_id") or "").strip()
    display_name = (form.get("line_display_name") or user_name).strip()
    enabled = form.get("enabled") == "on"
    receive_types = _normalize_permission_list(form.getlist("receive_types") or DEFAULT_PERSONAL_RECEIVE_TYPES)
    command_types = _normalize_permission_list(form.getlist("command_types") or DEFAULT_PERSONAL_COMMAND_TYPES)
    note = (form.get("note") or "").strip()

    # LINE userId 通常是 U 開頭；不強制擋掉，避免測試環境或未來格式調整。
    if line_user_id and len(line_user_id) < 10:
        raise ValueError("LINE 個人 ID 看起來太短，請確認是否貼到完整 ID")

    now = now_taipei().isoformat()

    doc_ref, _ = _current_crm_user_doc()
    if doc_ref:
        doc_ref.set({
            "line_user_id": line_user_id,
            "line_personal_user_id": line_user_id,
            "line_display_name": display_name,
            "line_bind_enabled": enabled,
            "line_receive_types": receive_types,
            "line_command_types": command_types,
            "line_bind_note": note,
            "line_bound_at": now,
            "updated_at": now,
            "updated_by_id": user_id,
            "updated_by_name": user_name,
        }, merge=True)

    settings = get_line_card_settings()
    rows = []
    updated = False
    for item in settings.get("line_personal_users") or []:
        old_crm_user_id = (item.get("crm_user_id") or "").strip()
        old_user_id = (item.get("user_id") or item.get("line_user_id") or "").strip()
        # 同一位 CRM 使用者，或同一個 LINE userId，都視為同一筆綁定。
        if (old_crm_user_id and old_crm_user_id == user_id) or (line_user_id and old_user_id == line_user_id):
            item = dict(item)
            item.update({
                "enabled": enabled,
                "name": display_name or user_name,
                "user_id": line_user_id,
                "line_user_id": line_user_id,
                "crm_user_id": user_id,
                "crm_user_email": user_email,
                "receive_types": receive_types,
                "command_types": command_types,
                "note": note,
                "updated_at": now,
            })
            rows.append(item)
            updated = True
        else:
            rows.append(item)

    if not updated and line_user_id:
        rows.append({
            "enabled": enabled,
            "name": display_name or user_name,
            "user_id": line_user_id,
            "line_user_id": line_user_id,
            "crm_user_id": user_id,
            "crm_user_email": user_email,
            "receive_types": receive_types,
            "command_types": command_types,
            "note": note,
            "updated_at": now,
        })

    db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set({
        "line_personal_users": rows,
        "updated_at": now,
        "updated_by_id": user_id,
        "updated_by_name": user_name,
    }, merge=True)

    return get_current_user_line_binding()


PERSONAL_PROFILE_TEMPLATE = """
{% extends "base.html" %}
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3 flex-wrap gap-2">
  <div>
    <h3 class="mb-1">個人後台</h3>
    <div class="text-muted small">設定你的 LINE 個人帳號，個人行程會單獨推送到這個帳號。</div>
  </div>
  <a href="{{ url_for('calendar_page') }}" class="btn btn-outline-secondary btn-sm">回行事曆</a>
</div>

<div class="row g-4">
  <div class="col-lg-5">
    <div class="card shadow-sm border-0 h-100">
      <div class="card-body">
        <h5 class="card-title mb-3">後台使用者資料</h5>
        <p class="mb-1"><strong>姓名：</strong>{{ binding.crm_user_name or '-' }}</p>
        <p class="mb-1"><strong>Email：</strong>{{ binding.crm_user_email or '-' }}</p>
        <p class="mb-1"><strong>綁定狀態：</strong>
          {% if binding.user_id and binding.enabled %}
            <span class="badge bg-success">已啟用</span>
          {% elif binding.user_id %}
            <span class="badge bg-secondary">已綁定但停用</span>
          {% else %}
            <span class="badge bg-warning text-dark">尚未綁定</span>
          {% endif %}
        </p>
        {% if binding.updated_at %}
          <p class="small text-muted mt-2 mb-0">最後更新：{{ binding.updated_at[:16] }}</p>
        {% endif %}
        <hr>
        <div class="alert alert-info small mb-0">
          <div class="fw-bold mb-1">如何取得 LINE 個人 ID？</div>
          <ol class="mb-0 ps-3">
            <li>用你的個人 LINE 私訊官方帳號。</li>
            <li>輸入 <span class="font-monospace">#綁定</span>。</li>
            <li>Bot 會回覆「個人LINE ID」，複製後貼到右側欄位。</li>
          </ol>
        </div>
      </div>
    </div>
  </div>

  <div class="col-lg-7">
    <div class="card shadow-sm border-0">
      <div class="card-body">
        <h5 class="card-title mb-3">LINE 個人帳號綁定</h5>
        <form method="post">
          <div class="form-check form-switch mb-3">
            <input class="form-check-input" type="checkbox" name="enabled" id="enabled" {% if binding.enabled %}checked{% endif %}>
            <label class="form-check-label" for="enabled">啟用此 LINE 個人帳號</label>
          </div>

          <div class="mb-3">
            <label class="form-label">顯示名稱</label>
            <input type="text" class="form-control" name="line_display_name" value="{{ binding.name }}" placeholder="例：黃曜昀">
            <div class="form-text">這個名稱會出現在行事曆個人行程的指派選單。</div>
          </div>

          <div class="mb-3">
            <label class="form-label">LINE 個人 ID</label>
            <input type="text" class="form-control font-monospace" name="line_user_id" value="{{ binding.user_id }}" placeholder="Uxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx">
          </div>

          <div class="row">
            <div class="col-md-6 mb-3">
              <label class="form-label">可接收訊息</label>
              <div class="border rounded p-3 bg-light">
                {% for key, label in receive_options %}
                  <div class="form-check">
                    <input class="form-check-input" type="checkbox" name="receive_types" value="{{ key }}" id="receive_{{ key }}"
                      {% if key in binding.receive_types %}checked{% endif %}>
                    <label class="form-check-label" for="receive_{{ key }}">{{ label }}</label>
                  </div>
                {% endfor %}
              </div>
            </div>
            <div class="col-md-6 mb-3">
              <label class="form-label">可使用指令</label>
              <div class="border rounded p-3 bg-light">
                {% for key, label in command_options %}
                  <div class="form-check">
                    <input class="form-check-input" type="checkbox" name="command_types" value="{{ key }}" id="command_{{ key }}"
                      {% if key in binding.command_types %}checked{% endif %}>
                    <label class="form-check-label" for="command_{{ key }}">{{ label }}</label>
                  </div>
                {% endfor %}
              </div>
            </div>
          </div>

          <div class="mb-3">
            <label class="form-label">備註</label>
            <textarea class="form-control" name="note" rows="2" placeholder="例：曜昀個人帳號">{{ binding.note }}</textarea>
          </div>

          <button type="submit" class="btn btn-primary">儲存 LINE 綁定</button>
          <a href="{{ url_for('calendar_new') }}" class="btn btn-outline-primary">新增個人行程</a>
        </form>
      </div>
    </div>
  </div>
</div>
{% endblock %}
"""


@app.route("/me", methods=["GET", "POST"], endpoint="personal_profile")
@login_required
def personal_profile():
    if request.method == "POST":
        try:
            upsert_current_user_line_binding(request.form)
            flash("LINE 個人帳號綁定已更新", "success")
        except Exception as e:
            flash(f"LINE 綁定更新失敗：{e}", "danger")
        return redirect(url_for("personal_profile"))

    binding = get_current_user_line_binding()
    return render_template_string(
        PERSONAL_PROFILE_TEMPLATE,
        binding=binding,
        receive_options=LINE_RECEIVE_TYPE_OPTIONS,
        command_options=LINE_COMMAND_TYPE_OPTIONS,
    )


# 覆寫設定中心：個人帳號綁定移到右上角使用者名稱的「個人後台」，設定中心保留群組與權限設定。
_original_line_card_settings_center_before_profile = line_card_settings_center

def line_card_settings_center_with_profile_hint():
    response = _original_line_card_settings_center_before_profile()
    return response

try:
    app.view_functions["line_card_settings"] = login_required(line_card_settings_center_with_profile_hint)
    print("✅ 個人後台 LINE 綁定已啟用：點右上角使用者名稱可設定個人 LINE 帳號")
except Exception as e:
    print("⚠️ 個人後台 LINE 綁定套用失敗：", e)

# =============================================================================
# 個人後台 LINE 綁定 Patch End
# =============================================================================


# =============================================================================
# CRM 客需 / 委託 / 開發 個人 / 公開 Patch Start
# =============================================================================

CRM_VISIBILITY_OPTIONS = [("public", "公開資料"), ("personal", "個人資料")]


def crm_record_visibility_payload_from_form(form, existing=None):
    """客需/委託/開發共用：公開或個人，個人可綁 LINE 個人帳號。"""
    existing = existing or {}
    visibility = (form.get("visibility") or existing.get("visibility") or "public").strip()
    if visibility not in ("public", "personal"):
        visibility = "public"

    owner_line_user_id = (form.get("owner_line_user_id") or existing.get("owner_line_user_id") or "").strip()
    owner_line_name = ""
    try:
        for u in get_line_personal_users(include_disabled=True):
            if u.get("user_id") == owner_line_user_id:
                owner_line_name = u.get("name") or ""
                break
    except Exception:
        pass
    if not owner_line_name:
        owner_line_name = (form.get("owner_line_name") or existing.get("owner_line_name") or "").strip()

    return {
        "visibility": visibility,
        "owner_user_id": existing.get("owner_user_id") or session.get("user_id", ""),
        "owner_user_name": existing.get("owner_user_name") or session.get("user_name", ""),
        "owner_line_user_id": owner_line_user_id if visibility == "personal" else "",
        "owner_line_name": owner_line_name if visibility == "personal" else "",
    }


def crm_record_visibility_text(data: dict):
    return "個人資料" if (data or {}).get("visibility") == "personal" else "公開資料"


def crm_record_target_line_user_id(data: dict):
    return ((data or {}).get("owner_line_user_id") or "").strip()


def crm_record_target_line_name(data: dict):
    uid = crm_record_target_line_user_id(data)
    if uid:
        try:
            u = find_line_personal_user_by_user_id(uid)
            if u and u.get("name"):
                return u.get("name")
        except Exception:
            pass
    return (data or {}).get("owner_line_name") or (data or {}).get("owner_user_name") or ""


@app.context_processor
def inject_crm_visibility_options():
    try:
        users = get_line_personal_users(include_disabled=True)
    except Exception:
        users = []
    return {
        "crm_visibility_options": CRM_VISIBILITY_OPTIONS,
        "line_user_options": users,
    }


# LINE 群組卡片：補上公開/個人與對象顯示。
_original_build_record_flex_bubble_before_visibility = _build_record_flex_bubble

def _build_record_flex_bubble(record_type: str, record_id: str, data: dict, title_prefix="CRM 資料"):
    bubble = _original_build_record_flex_bubble_before_visibility(record_type, record_id, data, title_prefix=title_prefix)
    try:
        vis_text = crm_record_visibility_text(data)
        target_name = crm_record_target_line_name(data)
        extra = vis_text + (f"｜{target_name}" if data.get("visibility") == "personal" and target_name else "")
        body = bubble.get("body", {}).get("contents", [])
        insert_at = 2 if len(body) >= 2 else len(body)
        body.insert(insert_at, flex_text(extra, size="xs", color="#8b6b4f", weight="bold", margin="sm"))
    except Exception as e:
        print("⚠️ CRM 卡片補上公開/個人標示失敗：", e)
    return bubble


# 加入行事曆時，把客需/委託/開發的公開/個人設定一起帶進行程。
def _record_calendar_params(record_type: str, record_id: str, data: dict):
    record_type = record_type or ""
    name = data.get("name") or data.get("customer_name") or ""
    phone = data.get("phone") or ""
    location = data.get("address") or data.get("preferred_areas") or data.get("registered_address") or ""
    if record_type == "buyer":
        title = f"{name} 客需追蹤".strip()
        note = data.get("note") or data.get("requirement_must") or data.get("preferred_areas") or ""
        category = "回電"
    elif record_type == "seller":
        title = f"{name} 委託追蹤".strip()
        note = data.get("note") or data.get("reason") or ""
        category = "回電"
    else:
        title = f"{name or data.get('address') or '開發'} 開發追蹤".strip()
        note = data.get("note") or data.get("next_action") or ""
        category = "開發"
    return {
        "related_type": record_type,
        "related_id": record_id,
        "title": title,
        "category": category,
        "customer_name": name,
        "phone": phone,
        "location": location,
        "note": note,
        "visibility": data.get("visibility") or "public",
        "owner_line_user_id": data.get("owner_line_user_id") or "",
    }


# 傳送客需 / 委託 / 開發：公開傳群組，個人傳該 LINE 個人帳號。
def _push_record_to_group(record_type: str, record_id: str, title_prefix="CRM 資料"):
    coll = {"buyer": "buyers", "seller": "sellers", "development": "developments"}.get(record_type)
    if not coll:
        return {"ok": False, "error": "record_type 不正確"}
    snap = db.collection(coll).document(record_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到資料"}
    data = snap.to_dict() or {}
    bubble = _build_record_flex_bubble(record_type, record_id, data, title_prefix=title_prefix)
    message = {"type": "flex", "altText": f"{title_prefix}：{data.get('name','')}", "contents": bubble}

    if (data.get("visibility") or "public") == "personal":
        user_id = crm_record_target_line_user_id(data)
        if not user_id:
            return {"ok": False, "error": "此資料設定為個人，但尚未綁定 LINE 個人帳號"}
        try:
            user_cfg = find_line_personal_user_by_user_id(user_id)
            if user_cfg and not line_personal_user_allows_receive(user_cfg, record_type):
                return {"ok": False, "error": f"{crm_record_target_line_name(data) or '此帳號'} 未開放接收「{record_type}」訊息"}
        except Exception:
            pass
        res = line_push_messages(user_id, [message])
        res["personal_target_name"] = crm_record_target_line_name(data)
        return res

    return line_push_messages_to_allowed_groups(record_type, [message])


# 重新套用送出按鈕 view，確保使用新的公開/個人邏輯。
def buyer_send_to_line_multi_group(buyer_id):
    res = _push_record_to_group("buyer", buyer_id, title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("buyer_detail", buyer_id=buyer_id))


def seller_send_to_line_multi_group(seller_id):
    res = _push_record_to_group("seller", seller_id, title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("seller_detail", seller_id=seller_id))


def development_send_to_line_multi_group(development_id):
    res = _push_record_to_group("development", development_id, title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("development_detail", development_id=development_id))

try:
    app.view_functions["buyer_send_to_line"] = login_required(buyer_send_to_line_multi_group)
    app.view_functions["seller_send_to_line"] = login_required(seller_send_to_line_multi_group)
    app.view_functions["development_send_to_line"] = login_required(development_send_to_line_multi_group)
    print("✅ CRM 客需/委託/開發公開/個人傳送已啟用")
except Exception as e:
    print("⚠️ CRM 公開/個人傳送套用失敗：", e)

# =============================================================================
# CRM 客需 / 委託 / 開發 個人 / 公開 Patch End
# =============================================================================



# =============================================================================
# 設定中心：誰可以看什麼 / 全客製化權限 Patch v20260623A
# - 群組 / 個人可分別設定：可接收、可使用指令、可查詢資料類型、可見範圍
# - LINE 群組查詢預設只看公開資料；可在後台設定成店長/管理員權限看全部
# - LINE 個人查詢預設看公開 + 自己綁定的個人資料；也可客製為只看自己 / 看全部
# =============================================================================

LINE_VIEW_TYPE_OPTIONS = [
    ("buyer", "客需資料"),
    ("seller", "委託資料"),
    ("development", "開發資料"),
    ("calendar", "行事曆資料"),
    ("todo", "待辦事項"),
]

LINE_VISIBILITY_SCOPE_OPTIONS = [
    ("public_only", "只看公開資料"),
    ("public_and_own", "公開資料 + 自己的個人資料"),
    ("own_only", "只看自己的個人資料"),
    ("all", "全部資料（公開 + 所有人個人資料）"),
]

_LINE_PERMISSION_CONTEXT = {}


def _permission_normalize_scope(value, default="public_only"):
    value = (value or "").strip()
    valid = {x[0] for x in LINE_VISIBILITY_SCOPE_OPTIONS}
    return value if value in valid else default


def _permission_record_type_from_collection(collection_name: str):
    mapping = {
        "buyers": "buyer",
        "sellers": "seller",
        "developments": "development",
        CALENDAR_EVENT_COLLECTION if 'CALENDAR_EVENT_COLLECTION' in globals() else 'calendar_events': "calendar",
        "todos": "todo",
        "line_todos": "todo",
    }
    return mapping.get(collection_name, "")


def _permission_normalize_view_types(values, default=None):
    vals = _normalize_permission_list(values)
    if vals:
        return vals
    return list(default or [])


# 重新覆寫群組設定 normalize：加上 view_types / visibility_scope。
def _normalize_line_group_config(group):
    group = dict(group or {})
    return {
        "enabled": bool(group.get("enabled", True)),
        "name": (group.get("name") or "未命名群組").strip(),
        "target_id": (group.get("target_id") or group.get("group_id") or group.get("room_id") or "").strip(),
        "receive_types": _normalize_permission_list(group.get("receive_types") or group.get("receives") or []),
        "command_types": _normalize_permission_list(group.get("command_types") or group.get("commands") or []),
        "view_types": _permission_normalize_view_types(group.get("view_types") or group.get("visible_types") or [], default=[x[0] for x in LINE_VIEW_TYPE_OPTIONS]),
        "visibility_scope": _permission_normalize_scope(group.get("visibility_scope"), default="public_only"),
        "note": (group.get("note") or "").strip(),
    }


# 重新覆寫個人帳號 normalize：加上 view_types / visibility_scope。
def _normalize_line_personal_user_config(user):
    user = dict(user or {})
    return {
        "enabled": bool(user.get("enabled", True)),
        "name": (user.get("name") or "未命名使用者").strip(),
        "user_id": (user.get("user_id") or user.get("line_user_id") or "").strip(),
        "receive_types": _normalize_permission_list(user.get("receive_types") or user.get("receives") or []),
        "command_types": _normalize_permission_list(user.get("command_types") or user.get("commands") or []),
        "view_types": _permission_normalize_view_types(user.get("view_types") or user.get("visible_types") or [], default=[x[0] for x in LINE_VIEW_TYPE_OPTIONS]),
        "visibility_scope": _permission_normalize_scope(user.get("visibility_scope"), default="public_and_own"),
        "note": (user.get("note") or "").strip(),
    }


def permission_config_can_view(source_cfg, data_type: str, data: dict, source_kind: str = "", source_id: str = "") -> bool:
    """依設定中心權限判斷 LINE 來源是否可看某筆資料。"""
    data_type = (data_type or "").strip()
    if not source_cfg:
        return False

    view_types = set(source_cfg.get("view_types") or [])
    if "all" not in view_types and data_type and data_type not in view_types:
        return False

    visibility = (data or {}).get("visibility") or "public"
    if visibility not in ("public", "personal"):
        visibility = "public"

    scope = _permission_normalize_scope(source_cfg.get("visibility_scope"), default="public_only")
    owner_line_id = (
        (data or {}).get("owner_line_user_id")
        or (data or {}).get("line_user_id")
        or (data or {}).get("target_user_id")
        or ""
    ).strip()

    # 舊資料如果沒有 visibility，視為公開資料，避免舊資料突然完全查不到。
    if visibility == "public":
        return scope in ("public_only", "public_and_own", "all")

    # 個人資料
    if scope == "all":
        return True
    if scope in ("public_and_own", "own_only"):
        return bool(source_kind == "user" and owner_line_id and owner_line_id == source_id)
    return False


def permission_current_line_source():
    ctx = dict(_LINE_PERMISSION_CONTEXT or {})
    if ctx.get("source_cfg"):
        return ctx
    return {"kind": "", "target_id": "", "source_cfg": None}


def permission_filter_doc_snapshot(record_type: str, doc_snapshot):
    ctx = permission_current_line_source()
    if not ctx.get("source_cfg"):
        return True
    try:
        data = doc_snapshot.to_dict() or {}
    except Exception:
        data = {}
    return permission_config_can_view(ctx.get("source_cfg"), record_type, data, ctx.get("kind"), ctx.get("target_id"))


# 讓 LINE 處理期間帶入目前來源權限；後台操作不受影響。
try:
    _process_line_message_event_before_permission_matrix = process_line_message_event

    def process_line_message_event(event):
        global _LINE_PERMISSION_CONTEXT
        old_ctx = dict(_LINE_PERMISSION_CONTEXT or {})
        kind, target_id = line_event_source_kind_and_id(event)
        cfg = None
        if kind == "user":
            cfg = find_line_personal_user_by_user_id(target_id)
        elif kind in ("group", "room"):
            cfg = find_line_group_by_target_id(target_id)
        _LINE_PERMISSION_CONTEXT = {"kind": kind, "target_id": target_id, "source_cfg": cfg}
        try:
            return _process_line_message_event_before_permission_matrix(event)
        finally:
            _LINE_PERMISSION_CONTEXT = old_ctx

    print("✅ LINE 查詢權限矩陣：已掛入 process_line_message_event")
except Exception as e:
    print("⚠️ LINE 查詢權限矩陣 process_line_message_event 掛入失敗：", e)


# 過濾電話搜尋結果，避免群組 / 個人查到不該看的私有資料。
try:
    _find_records_by_phone_before_permission_matrix = find_records_by_phone

    def find_records_by_phone(collection_name: str, phone: str):
        docs = _find_records_by_phone_before_permission_matrix(collection_name, phone)
        record_type = _permission_record_type_from_collection(collection_name)
        ctx = permission_current_line_source()
        if not ctx.get("source_cfg") or not record_type:
            return docs
        return [d for d in docs if permission_filter_doc_snapshot(record_type, d)]

    print("✅ LINE 查詢權限矩陣：已過濾電話搜尋")
except Exception as e:
    print("⚠️ LINE 查詢權限矩陣 find_records_by_phone 掛入失敗：", e)


# 過濾單筆查詢結果。
try:
    _find_customer_record_before_permission_matrix = find_customer_record

    def find_customer_record(target_type: str, record_id: str = "", phone: str = "", name: str = "", address: str = ""):
        try:
            doc = _find_customer_record_before_permission_matrix(target_type, record_id, phone, name, address)
        except TypeError:
            doc = _find_customer_record_before_permission_matrix(target_type, record_id, phone, name)
        if not doc:
            return doc
        record_type = "development" if target_type in ("development", "developments") else target_type
        ctx = permission_current_line_source()
        if ctx.get("source_cfg") and record_type in ("buyer", "seller", "development"):
            if not permission_filter_doc_snapshot(record_type, doc):
                return None
        return doc

    print("✅ LINE 查詢權限矩陣：已過濾單筆查詢")
except Exception as e:
    print("⚠️ LINE 查詢權限矩陣 find_customer_record 掛入失敗：", e)


# 行事曆查詢依權限矩陣過濾，不再固定「群組只能公開、個人只能自己」。
try:
    _fetch_calendar_events_before_permission_matrix = fetch_calendar_events

    def fetch_calendar_events(start_date: str, end_date: str = None, line_only: bool = False, calendar_view: str = "all", line_source_kind: str = "", line_source_id: str = ""):
        start_date = calendar_safe_date(start_date)
        end_date = calendar_safe_date(end_date or start_date)
        if not line_source_kind and permission_current_line_source().get("kind"):
            ctx = permission_current_line_source()
            line_source_kind = ctx.get("kind") or ""
            line_source_id = ctx.get("target_id") or ""

        # 後台照舊。
        if not line_source_kind:
            return _fetch_calendar_events_before_permission_matrix(start_date, end_date, line_only=line_only, calendar_view=calendar_view)

        try:
            if start_date == end_date:
                docs = list(db.collection(CALENDAR_EVENT_COLLECTION).where("event_date", "==", start_date).stream())
            else:
                docs = list(db.collection(CALENDAR_EVENT_COLLECTION).where("event_date", ">=", start_date).where("event_date", "<=", end_date).stream())
        except Exception as e:
            print("⚠️ 權限矩陣行事曆查詢失敗，改用全表掃描：", e)
            docs = list(db.collection(CALENDAR_EVENT_COLLECTION).stream())

        source_cfg = None
        if line_source_kind == "user":
            source_cfg = find_line_personal_user_by_user_id(line_source_id)
        elif line_source_kind in ("group", "room"):
            source_cfg = find_line_group_by_target_id(line_source_id)

        events = []
        for d in docs:
            try:
                item = doc_to_calendar_event(d)
            except Exception:
                continue
            if not (start_date <= item.get("event_date", "") <= end_date):
                continue
            if line_only and not item.get("line_enabled", True):
                continue
            if not permission_config_can_view(source_cfg, "calendar", item, line_source_kind, line_source_id):
                continue
            events.append(item)
        events.sort(key=lambda x: (x.get("event_date", ""), x.get("start_time", ""), x.get("created_at", "")))
        return events

    print("✅ LINE 查詢權限矩陣：行事曆查詢已套用")
except Exception as e:
    print("⚠️ LINE 查詢權限矩陣 fetch_calendar_events 掛入失敗：", e)


# line_access_gate 加一道「資料可查詢類型」檢查：有指令權限但沒有可看資料，也不放行查詢。
try:
    _line_access_gate_before_permission_matrix = line_access_gate

    def line_access_gate(event):
        allowed, reason_or_cmd, source_cfg = _line_access_gate_before_permission_matrix(event)
        if not allowed:
            return allowed, reason_or_cmd, source_cfg

        cmd_type = reason_or_cmd
        if cmd_type in ("system_group_id", "followup"):
            return allowed, reason_or_cmd, source_cfg

        view_map = {
            "calendar": "calendar",
            "todo": "todo",
            "buyer": "buyer",
            "seller": "seller",
            "development": "development",
        }
        view_type = view_map.get(cmd_type)
        if view_type and source_cfg:
            view_types = set(source_cfg.get("view_types") or [])
            if "all" not in view_types and view_type not in view_types:
                return False, f"此來源未開放查看「{view_type}」資料，請到設定中心調整『可查詢資料』。", source_cfg
        return allowed, reason_or_cmd, source_cfg

    print("✅ LINE 查詢權限矩陣：line_access_gate 已加上可查詢資料檢查")
except Exception as e:
    print("⚠️ LINE 查詢權限矩陣 line_access_gate 掛入失敗：", e)


def parse_line_groups_from_form(form):
    try:
        count = int(form.get("group_count", 0))
    except Exception:
        count = 0
    count = max(0, min(30, count))
    groups = []
    for i in range(count):
        target_id = (form.get(f"group_{i}_target_id") or "").strip()
        name = (form.get(f"group_{i}_name") or "").strip()
        if not target_id and not name:
            continue
        groups.append({
            "enabled": form.get(f"group_{i}_enabled") == "on",
            "name": name or f"群組 {i+1}",
            "target_id": target_id,
            "receive_types": form.getlist(f"group_{i}_receive_types"),
            "command_types": form.getlist(f"group_{i}_command_types"),
            "view_types": form.getlist(f"group_{i}_view_types"),
            "visibility_scope": _permission_normalize_scope(form.get(f"group_{i}_visibility_scope"), default="public_only"),
            "note": (form.get(f"group_{i}_note") or "").strip(),
        })
    return groups


def parse_line_personal_users_from_form(form):
    try:
        count = int(form.get("personal_user_count", 0))
    except Exception:
        count = 0
    count = max(0, min(50, count))
    users = []
    for i in range(count):
        user_id = (form.get(f"personal_{i}_user_id") or "").strip()
        name = (form.get(f"personal_{i}_name") or "").strip()
        if not user_id and not name:
            continue
        users.append({
            "enabled": form.get(f"personal_{i}_enabled") == "on",
            "name": name or f"個人 {i+1}",
            "user_id": user_id,
            "receive_types": form.getlist(f"personal_{i}_receive_types"),
            "command_types": form.getlist(f"personal_{i}_command_types"),
            "view_types": form.getlist(f"personal_{i}_view_types"),
            "visibility_scope": _permission_normalize_scope(form.get(f"personal_{i}_visibility_scope"), default="public_and_own"),
            "note": (form.get(f"personal_{i}_note") or "").strip(),
        })
    return users


def save_line_settings_center_from_form(form):
    updates = save_line_card_settings_from_form(form)
    groups = parse_line_groups_from_form(form)
    personal_users = parse_line_personal_users_from_form(form)

    if not groups:
        old_target = (form.get("notify_target_id") or updates.get("notify_target_id") or os.environ.get("LINE_NOTIFY_TARGET_ID") or "").strip()
        if old_target:
            groups = [{
                "enabled": True,
                "name": "預設群組",
                "target_id": old_target,
                "receive_types": [x[0] for x in LINE_RECEIVE_TYPE_OPTIONS],
                "command_types": [x[0] for x in LINE_COMMAND_TYPE_OPTIONS],
                "view_types": [x[0] for x in LINE_VIEW_TYPE_OPTIONS],
                "visibility_scope": "public_only",
                "note": "舊版單一群組相容",
            }]

    extra = {
        "line_groups": groups,
        "line_personal_users": personal_users,
        "updated_at": now_taipei().isoformat(),
        "updated_by_id": session.get("user_id", ""),
        "updated_by_name": session.get("user_name", ""),
    }

    new_pw = (form.get("new_settings_admin_password") or "").strip()
    new_pw2 = (form.get("new_settings_admin_password_confirm") or "").strip()
    if new_pw or new_pw2:
        if new_pw != new_pw2:
            raise ValueError("兩次輸入的新管理員密碼不一致")
        if len(new_pw) < 4:
            raise ValueError("管理員密碼至少 4 碼")
        extra["settings_admin_password_hash"] = generate_password_hash(new_pw)

    db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set(extra, merge=True)
    return {**updates, **extra}


LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX = """
<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8">
  <title>設定中心｜厝米 Team M.E</title>
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/css/bootstrap.min.css" rel="stylesheet">
  <style>
    body{background:#f6f3ef}.setting-card{background:#fff;border:1px solid #eadbca;border-radius:18px;box-shadow:0 8px 24px rgba(120,80,40,.08)}
    .group-box,.user-box{border:1px solid #eadbca;border-radius:14px;padding:14px;background:#fffaf5}.code{font-family:Consolas,monospace}.hint{font-size:.88rem;color:#8b6b4f;line-height:1.7}.sticky-nav{position:sticky;top:16px}.small-check label{margin-right:.65rem;margin-bottom:.25rem}.scope-select{max-width:360px}.badge-soft{background:#f8eadc;color:#7a4e2d;border:1px solid #eadbca}
  </style>
</head>
<body>
<div class="container-fluid py-4">
  <div class="d-flex justify-content-between align-items-center mb-3 flex-wrap gap-2">
    <div><h2 class="mb-1">設定中心</h2><div class="text-muted small">自訂誰可以接收、誰可以下指令、誰可以看哪些資料。</div></div>
    <div class="d-flex gap-2"><a href="{{ url_for('buyers') }}" class="btn btn-outline-secondary btn-sm">回後台</a><form method="post" action="{{ url_for('line_settings_admin_logout') }}"><button class="btn btn-outline-danger btn-sm">離開管理模式</button></form></div>
  </div>
  {% with messages = get_flashed_messages(with_categories=true) %}{% if messages %}{% for category, msg in messages %}<div class="alert alert-{{ category }} py-2">{{ msg }}</div>{% endfor %}{% endif %}{% endwith %}
  <form method="post">
    <div class="row g-3">
      <div class="col-lg-3"><div class="setting-card p-3 sticky-nav"><div class="list-group small">
        <a class="list-group-item list-group-item-action" href="#security">管理員密碼</a>
        <a class="list-group-item list-group-item-action" href="#groups">LINE 群組權限</a>
        <a class="list-group-item list-group-item-action" href="#personal">LINE 個人權限</a>
        <a class="list-group-item list-group-item-action" href="#card">LINE 卡片樣式</a>
        <a class="list-group-item list-group-item-action" href="#calendar">行事曆設定</a>
        <a class="list-group-item list-group-item-action" href="#rules">權限規則</a>
      </div></div></div>
      <div class="col-lg-9">
        <div id="security" class="setting-card p-4 mb-4">
          <h5>管理員密碼</h5><div class="hint mb-3">第一次預設密碼是 <span class="code">123456</span>。修改後，下次進入設定中心需輸入新密碼。</div>
          <div class="row g-2"><div class="col-md-6"><label class="form-label">新密碼</label><input type="password" name="new_settings_admin_password" class="form-control" placeholder="留空則不修改"></div><div class="col-md-6"><label class="form-label">再次輸入</label><input type="password" name="new_settings_admin_password_confirm" class="form-control" placeholder="留空則不修改"></div></div>
        </div>

        <div id="groups" class="setting-card p-4 mb-4">
          <h5>LINE 群組權限</h5>
          <div class="hint mb-3">群組輸入 <span class="code">#綁定</span> 可取得群組 ID。每個群組可以獨立設定：接收什麼、能下什麼指令、能查什麼資料、能看公開或全部。</div>
          <input type="hidden" name="group_count" value="{{ group_rows|length }}">
          {% for g in group_rows %}{% set idx = loop.index0 %}
          <div class="group-box mb-3">
            <div class="row g-2 align-items-end">
              <div class="col-md-1"><label class="form-label">啟用</label><input class="form-check-input d-block" type="checkbox" name="group_{{ idx }}_enabled" {% if g.enabled %}checked{% endif %}></div>
              <div class="col-md-3"><label class="form-label">群組名稱</label><input class="form-control" name="group_{{ idx }}_name" value="{{ g.name }}" placeholder="例如：業務群"></div>
              <div class="col-md-5"><label class="form-label">群組 / 房間 ID</label><input class="form-control code" name="group_{{ idx }}_target_id" value="{{ g.target_id }}" placeholder="Cxxxxxxxx 或 Rxxxxxxxx"></div>
              <div class="col-md-3"><label class="form-label">備註</label><input class="form-control" name="group_{{ idx }}_note" value="{{ g.note }}"></div>
            </div>
            <div class="row mt-3 g-3">
              <div class="col-md-6"><div class="fw-bold small mb-1">1. 可接收推播</div><div class="small-check">{% for key,label in receive_options %}<label class="form-check form-check-inline small"><input class="form-check-input" type="checkbox" name="group_{{ idx }}_receive_types" value="{{ key }}" {% if key in g.receive_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div></div>
              <div class="col-md-6"><div class="fw-bold small mb-1">2. 可使用指令</div>{% for key,label in command_options %}<label class="form-check d-block small"><input class="form-check-input" type="checkbox" name="group_{{ idx }}_command_types" value="{{ key }}" {% if key in g.command_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div>
              <div class="col-md-6"><div class="fw-bold small mb-1">3. 可查詢資料</div><div class="small-check">{% for key,label in view_options %}<label class="form-check form-check-inline small"><input class="form-check-input" type="checkbox" name="group_{{ idx }}_view_types" value="{{ key }}" {% if key in g.view_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div></div>
              <div class="col-md-6"><div class="fw-bold small mb-1">4. 可見範圍</div><select class="form-select scope-select" name="group_{{ idx }}_visibility_scope">{% for key,label in visibility_scope_options %}<option value="{{ key }}" {% if g.visibility_scope == key %}selected{% endif %}>{{ label }}</option>{% endfor %}</select><div class="hint mt-1">一般群組建議用「只看公開資料」；店長群可改「全部資料」。</div></div>
            </div>
          </div>
          {% endfor %}
        </div>

        <div id="personal" class="setting-card p-4 mb-4">
          <h5>LINE 個人權限</h5>
          <div class="hint mb-3">個人私訊官方帳號輸入 <span class="code">#綁定</span>，會回覆個人 LINE ID。把 ID 貼在這裡後，就可以設定這個人能看什麼。</div>
          <input type="hidden" name="personal_user_count" value="{{ personal_user_rows|length }}">
          {% for u in personal_user_rows %}{% set idx = loop.index0 %}
          <div class="user-box mb-3">
            <div class="row g-2 align-items-end">
              <div class="col-md-1"><label class="form-label">啟用</label><input class="form-check-input d-block" type="checkbox" name="personal_{{ idx }}_enabled" {% if u.enabled %}checked{% endif %}></div>
              <div class="col-md-3"><label class="form-label">姓名 / 暱稱</label><input class="form-control" name="personal_{{ idx }}_name" value="{{ u.name }}" placeholder="例如：曜昀"></div>
              <div class="col-md-5"><label class="form-label">LINE 個人 userId</label><input class="form-control code" name="personal_{{ idx }}_user_id" value="{{ u.user_id }}" placeholder="Uxxxxxxxx"></div>
              <div class="col-md-3"><label class="form-label">備註</label><input class="form-control" name="personal_{{ idx }}_note" value="{{ u.note }}"></div>
            </div>
            <div class="row mt-3 g-3">
              <div class="col-md-6"><div class="fw-bold small mb-1">1. 可接收推播</div><div class="small-check">{% for key,label in receive_options %}<label class="form-check form-check-inline small"><input class="form-check-input" type="checkbox" name="personal_{{ idx }}_receive_types" value="{{ key }}" {% if key in u.receive_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div></div>
              <div class="col-md-6"><div class="fw-bold small mb-1">2. 可使用指令</div>{% for key,label in command_options %}<label class="form-check d-block small"><input class="form-check-input" type="checkbox" name="personal_{{ idx }}_command_types" value="{{ key }}" {% if key in u.command_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div>
              <div class="col-md-6"><div class="fw-bold small mb-1">3. 可查詢資料</div><div class="small-check">{% for key,label in view_options %}<label class="form-check form-check-inline small"><input class="form-check-input" type="checkbox" name="personal_{{ idx }}_view_types" value="{{ key }}" {% if key in u.view_types %}checked{% endif %}> {{ label }}</label>{% endfor %}</div></div>
              <div class="col-md-6"><div class="fw-bold small mb-1">4. 可見範圍</div><select class="form-select scope-select" name="personal_{{ idx }}_visibility_scope">{% for key,label in visibility_scope_options %}<option value="{{ key }}" {% if u.visibility_scope == key %}selected{% endif %}>{{ label }}</option>{% endfor %}</select><div class="hint mt-1">一般個人建議用「公開資料 + 自己的個人資料」。</div></div>
            </div>
          </div>
          {% endfor %}
        </div>

        <div id="card" class="setting-card p-4 mb-4">
          <h5>LINE 卡片樣式</h5>
          <div class="row g-2">
            <div class="col-12"><label class="form-label">今日行程標題</label><input class="form-control" name="title_today" value="{{ settings.title_today }}"></div>
            <div class="col-12"><label class="form-label">明日行程標題</label><input class="form-control" name="title_tomorrow" value="{{ settings.title_tomorrow }}"></div>
            <div class="col-12"><label class="form-label">本週行程標題</label><input class="form-control" name="title_week" value="{{ settings.title_week }}"></div>
            <div class="col-md-6"><label class="form-label">主色</label><input class="form-control" name="primary_color" value="{{ settings.primary_color }}"></div>
            <div class="col-md-6"><label class="form-label">按鈕色</label><input class="form-control" name="button_color" value="{{ settings.button_color }}"></div>
            <div class="col-12"><label class="form-label">LINE 卡片顯示內容</label><br>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_customer" {% if settings.show_customer %}checked{% endif %}> 客戶</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_phone" {% if settings.show_phone %}checked{% endif %}> 電話</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_location" {% if settings.show_location %}checked{% endif %}> 地點</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_note" {% if settings.show_note %}checked{% endif %}> 備註</label>
              <label class="form-check form-check-inline"><input class="form-check-input" type="checkbox" name="show_quick_actions" {% if settings.show_quick_actions %}checked{% endif %}> 快速操作</label>
              <label class="form-check d-block mt-2"><input class="form-check-input" type="checkbox" name="line_only_enabled_events" {% if settings.line_only_enabled_events %}checked{% endif %}> LINE 只顯示有勾「LINE查詢顯示」的行程</label>
            </div>
            <div class="col-12"><label class="form-label">快速操作按鈕文字</label><textarea class="form-control" rows="4" name="quick_actions">{{ quick_actions_text }}</textarea></div>
          </div>
        </div>

        <div id="calendar" class="setting-card p-4 mb-4">
          <h5>行事曆設定</h5>
          <div class="row g-2">
            <div class="col-md-6"><label class="form-label">行事曆開始小時</label><input class="form-control" name="calendar_start_hour" value="{{ settings.calendar_start_hour }}"></div>
            <div class="col-md-6"><label class="form-label">行事曆結束小時</label><input class="form-control" name="calendar_end_hour" value="{{ settings.calendar_end_hour }}"></div>
            <div class="col-12"><label class="form-label">行事曆行程分類</label><textarea class="form-control" rows="5" name="calendar_categories">{{ calendar_categories_text }}</textarea><div class="hint">一行一個分類，例如：帶看、回電、開發、簽約、拍照、私人行程、其他。</div></div>
          </div>
        </div>

        <div id="rules" class="setting-card p-4 mb-4">
          <h5>權限規則說明</h5>
          <div class="row g-3 small">
            <div class="col-md-6"><span class="badge badge-soft mb-2">群組</span><br>群組通常設定「只看公開資料」。如果設定為「全部資料」，該群組查詢時會看得到所有人的個人資料，請只給店長或管理者群組。</div>
            <div class="col-md-6"><span class="badge badge-soft mb-2">個人</span><br>個人建議設定「公開資料 + 自己的個人資料」。這樣他私訊查詢時，看得到公開資料和綁定給自己的資料，但看不到別人的個人資料。</div>
            <div class="col-md-6"><span class="badge badge-soft mb-2">可使用指令</span><br>控制能不能下 #今日行程、#新增客需、#查詢紀錄 等指令。</div>
            <div class="col-md-6"><span class="badge badge-soft mb-2">可查詢資料</span><br>控制指令通過後，實際能不能看到客需、委託、開發、行事曆、待辦。</div>
          </div>
        </div>
        <button class="btn btn-primary btn-lg w-100" type="submit">儲存全部設定</button>
      </div>
    </div>
  </form>
</div>
<script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
"""


def line_card_settings_center():
    if request.method == "POST" and request.form.get("admin_login") == "1":
        if _settings_admin_password_ok(request.form.get("admin_password", "")):
            session[LINE_SETTINGS_ADMIN_SESSION_KEY] = True
            flash("已進入設定中心", "success")
            return redirect(url_for("line_card_settings"))
        flash("管理員密碼錯誤", "danger")
        return render_template_string(LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE)

    if not session.get(LINE_SETTINGS_ADMIN_SESSION_KEY):
        return render_template_string(LINE_SETTINGS_ADMIN_LOGIN_TEMPLATE)

    if request.method == "POST":
        try:
            save_line_settings_center_from_form(request.form)
            flash("設定已更新", "success")
        except Exception as e:
            flash(f"設定儲存失敗：{e}", "danger")
        return redirect(url_for("line_card_settings"))

    settings = get_line_card_settings()
    group_rows = list(get_line_group_settings())
    while len(group_rows) < 8:
        group_rows.append({"enabled": False, "name": "", "target_id": "", "receive_types": [], "command_types": [], "view_types": ["buyer", "seller", "development", "calendar", "todo"], "visibility_scope": "public_only", "note": ""})

    personal_user_rows = list(get_line_personal_users(include_disabled=True))
    while len(personal_user_rows) < max(10, LINE_PERSONAL_USER_COUNT_DEFAULT if 'LINE_PERSONAL_USER_COUNT_DEFAULT' in globals() else 8):
        personal_user_rows.append({"enabled": False, "name": "", "user_id": "", "receive_types": ["calendar"], "command_types": ["calendar", "todo", "followup"], "view_types": ["buyer", "seller", "development", "calendar", "todo"], "visibility_scope": "public_and_own", "note": ""})

    return render_template_string(
        LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX,
        settings=settings,
        group_rows=group_rows,
        personal_user_rows=personal_user_rows,
        receive_options=LINE_RECEIVE_TYPE_OPTIONS,
        command_options=LINE_COMMAND_TYPE_OPTIONS,
        view_options=LINE_VIEW_TYPE_OPTIONS,
        visibility_scope_options=LINE_VISIBILITY_SCOPE_OPTIONS,
        quick_actions_text="\n".join(settings.get("quick_actions") or []),
        calendar_categories_text="\n".join(get_calendar_category_options()),
    )

try:
    app.view_functions["line_card_settings"] = login_required(line_card_settings_center)
    print("✅ 設定中心權限矩陣已啟用：可客製誰可以看什麼")
except Exception as e:
    print("⚠️ 設定中心權限矩陣套用失敗：", e)

# =============================================================================
# 設定中心：誰可以看什麼 / 全客製化權限 Patch End
# =============================================================================



# =============================================================================
# LINE 個人 / 群組精準權限 + 可選傳送對象 + 待辦公開/個人 Patch v20260623B
# - LINE 私訊新增客需/委託/開發：自動個人資料
# - LINE 群組新增客需/委託/開發：自動公開資料
# - 後台傳 LINE：可選傳給「目前登入者個人」或「目前登入者所在且已設定的群組」
# - 待辦事項：也支援公開 / 個人，LINE 查詢與勾選依權限過濾
# =============================================================================


def _line_source_kind_id_safe(event=None):
    try:
        return line_event_source_kind_and_id(event or {})
    except Exception:
        source = ((event or {}).get("source") or {})
        if source.get("groupId"):
            return "group", source.get("groupId")
        if source.get("roomId"):
            return "room", source.get("roomId")
        return "user", source.get("userId", "")


def _line_visibility_payload_from_event(event=None):
    """LINE 建立資料時：群組/聊天室=公開，個人私訊=個人。"""
    kind, target_id = _line_source_kind_id_safe(event)
    payload = {
        "line_created_source_kind": kind,
        "line_created_source_id": target_id,
        "updated_at": now_taipei().isoformat(),
    }
    if kind == "user":
        try:
            user_cfg = find_line_personal_user_by_user_id(target_id) or {}
        except Exception:
            user_cfg = {}
        name = (user_cfg.get("name") or get_line_sender_display_name(event or {}) or "").strip()
        payload.update({
            "visibility": "personal",
            "owner_line_user_id": target_id,
            "owner_line_name": name,
        })
    else:
        payload.update({
            "visibility": "public",
            "owner_line_user_id": "",
            "owner_line_name": "",
        })
    return payload


def _apply_line_visibility_to_created_result(result, event=None):
    if not result or not result.get("ok"):
        return result
    target_type = (result.get("target_type") or "").strip()
    target_id = (result.get("target_id") or "").strip()
    coll = {"buyer": "buyers", "seller": "sellers", "development": "developments"}.get(target_type)
    if not coll or not target_id:
        return result
    try:
        payload = _line_visibility_payload_from_event(event)
        db.collection(coll).document(target_id).set(payload, merge=True)
        result["visibility"] = payload.get("visibility")
        result["owner_line_user_id"] = payload.get("owner_line_user_id")
        result["owner_line_name"] = payload.get("owner_line_name")
        if payload.get("visibility") == "personal":
            result["reply_text"] = (result.get("reply_text") or "已寫入") + "\n已設為：個人資料"
        else:
            result["reply_text"] = (result.get("reply_text") or "已寫入") + "\n已設為：公開資料"
    except Exception as e:
        print("⚠️ 自動設定 LINE 新增資料公開/個人失敗：", e)
    return result


try:
    _create_buyer_need_before_visibility_auto = create_buyer_need
    def create_buyer_need(fields, event):
        result = _create_buyer_need_before_visibility_auto(fields, event)
        return _apply_line_visibility_to_created_result(result, event)
except Exception as e:
    print("⚠️ 掛入客需 LINE 自動公開/個人失敗：", e)

try:
    _create_seller_listing_before_visibility_auto = create_seller_listing
    def create_seller_listing(fields, event):
        result = _create_seller_listing_before_visibility_auto(fields, event)
        return _apply_line_visibility_to_created_result(result, event)
except Exception as e:
    print("⚠️ 掛入委託 LINE 自動公開/個人失敗：", e)

try:
    _create_development_before_visibility_auto = create_development
    def create_development(fields, event):
        result = _create_development_before_visibility_auto(fields, event)
        return _apply_line_visibility_to_created_result(result, event)
except Exception as e:
    print("⚠️ 掛入開發 LINE 自動公開/個人失敗：", e)


# ---------- 後台：可選擇傳送到哪個群組或個人 ----------

def _current_bound_line_user_id():
    try:
        binding = get_current_user_line_binding()
        return (binding.get("user_id") or "").strip(), (binding.get("name") or binding.get("crm_user_name") or session.get("user_name") or "個人").strip()
    except Exception:
        return "", (session.get("user_name") or "個人").strip()


def _line_group_contains_user(group_id: str, line_user_id: str) -> bool:
    """只列出官方帳號所在、且目前登入者也在的群組。"""
    group_id = (group_id or "").strip()
    line_user_id = (line_user_id or "").strip()
    if not group_id or not line_user_id:
        return False

    cache_key = f"member:{group_id}:{line_user_id}"
    cache = app.config.setdefault("LINE_GROUP_MEMBER_CACHE", {})
    if cache_key in cache:
        return bool(cache[cache_key])

    # 先用 webhook log 判斷：此 user 曾在此群組傳訊息給 bot，就代表兩者同群。
    try:
        logs = db.collection("line_logs").where("line_group_id", "==", group_id).where("line_user_id", "==", line_user_id).limit(1).stream()
        if list(logs):
            cache[cache_key] = True
            return True
    except Exception:
        pass

    # 再用 LINE API 判斷 member profile。官方帳號不在該群或 user 不在，通常會失敗。
    if LINE_CHANNEL_ACCESS_TOKEN:
        try:
            import requests
            res = requests.get(
                f"https://api.line.me/v2/bot/group/{group_id}/member/{line_user_id}",
                headers=line_api_headers(),
                timeout=5,
            )
            ok = (res.status_code == 200)
            cache[cache_key] = ok
            return ok
        except Exception as e:
            print("⚠️ 檢查 LINE 群組成員失敗：", e)

    cache[cache_key] = False
    return False


def _source_config_for_delivery_value(value: str):
    value = (value or "").strip()
    if value.startswith("user:"):
        uid = value.split(":", 1)[1].strip()
        return "user", uid, find_line_personal_user_by_user_id(uid)
    if value.startswith("group:"):
        gid = value.split(":", 1)[1].strip()
        return "group", gid, find_line_group_by_target_id(gid)
    return "", "", None


def line_delivery_options(record_type: str):
    """Jinja 使用：列出目前登入者可傳送的個人/群組。"""
    record_type = (record_type or "").strip()
    options = []
    line_user_id, line_name = _current_bound_line_user_id()

    if line_user_id:
        try:
            personal_cfg = find_line_personal_user_by_user_id(line_user_id)
            if personal_cfg and line_personal_user_allows_receive(personal_cfg, record_type):
                options.append({"value": f"user:{line_user_id}", "label": f"傳給個人｜{personal_cfg.get('name') or line_name}"})
            elif record_type in ("buyer", "seller", "development", "calendar", "todo"):
                # 即使未在接收設定勾選，仍顯示個人選項，但送出時會再次檢查。
                options.append({"value": f"user:{line_user_id}", "label": f"傳給個人｜{line_name}"})
        except Exception:
            options.append({"value": f"user:{line_user_id}", "label": f"傳給個人｜{line_name}"})

    for g in get_enabled_line_groups():
        try:
            if not line_group_allows_receive(g, record_type):
                continue
            gid = (g.get("target_id") or "").strip()
            if not gid:
                continue
            if line_user_id and not _line_group_contains_user(gid, line_user_id):
                continue
            options.append({"value": f"group:{gid}", "label": f"傳到群組｜{g.get('name') or gid[-6:]}"})
        except Exception:
            continue
    return options


@app.context_processor
def inject_line_delivery_options_patch():
    return {"line_delivery_options": line_delivery_options}


def _line_push_result_flash_message_selected(res):
    if not res or not res.get("ok"):
        return f"傳送失敗：{(res or {}).get('error') or (res or {}).get('text') or res}", "danger"
    label = (res.get("target_label") or res.get("group_name") or res.get("personal_target_name") or "LINE").strip()
    return f"已傳送到：{label}", "success"


def _build_record_message_for_delivery(record_type, record_id, title_prefix="後台傳送"):
    coll = {"buyer": "buyers", "seller": "sellers", "development": "developments"}.get(record_type)
    if not coll:
        return None, None, {"ok": False, "error": "record_type 不正確"}
    snap = db.collection(coll).document(record_id).get()
    if not snap.exists:
        return None, None, {"ok": False, "error": "找不到資料"}
    data = snap.to_dict() or {}
    bubble = _build_record_flex_bubble(record_type, record_id, data, title_prefix=title_prefix)
    msg = {"type": "flex", "altText": f"{title_prefix}：{data.get('name','')}", "contents": bubble}
    return data, msg, None


def _push_record_to_selected_destination(record_type, record_id, destination, title_prefix="後台傳送"):
    data, msg, err = _build_record_message_for_delivery(record_type, record_id, title_prefix=title_prefix)
    if err:
        return err
    kind, target_id, cfg = _source_config_for_delivery_value(destination)
    if not kind or not target_id:
        # 沒有選擇時保留舊邏輯：依公開/個人自動傳。
        return _push_record_to_group(record_type, record_id, title_prefix=title_prefix)

    # 目前登入者只能傳到自己個人，或自己也在的群組。
    my_line_id, my_name = _current_bound_line_user_id()
    if kind == "user" and my_line_id and target_id != my_line_id:
        return {"ok": False, "error": "只能傳送到目前登入者綁定的個人 LINE。"}
    if kind == "group" and my_line_id and not _line_group_contains_user(target_id, my_line_id):
        return {"ok": False, "error": "你不是這個 LINE 群組成員，或官方帳號不在該群組，不能傳送。"}

    if cfg:
        if kind == "user":
            if not line_personal_user_allows_receive(cfg, record_type):
                return {"ok": False, "error": "此個人帳號未開放接收這類訊息。"}
        else:
            if not line_group_allows_receive(cfg, record_type):
                return {"ok": False, "error": "此群組未開放接收這類訊息。"}
        if not permission_config_can_view(cfg, record_type, data, kind, target_id):
            return {"ok": False, "error": "依權限設定，此對象不可查看這筆資料。"}

    res = line_push_messages(target_id, [msg])
    res["target_label"] = (cfg or {}).get("name") or ("個人" if kind == "user" else "群組")
    return res


def _push_calendar_to_selected_destination(event_id, destination, title_prefix="後台傳送行程"):
    snap = db.collection(CALENDAR_EVENT_COLLECTION).document(event_id).get()
    if not snap.exists:
        return {"ok": False, "error": "找不到行程"}
    event = doc_to_calendar_event(snap)
    bubble = build_calendar_event_bubble(event)
    msg = {"type": "flex", "altText": f"{title_prefix}：{event.get('title','')}", "contents": bubble}
    kind, target_id, cfg = _source_config_for_delivery_value(destination)
    if not kind or not target_id:
        return _push_calendar_event_to_group(event_id, title_prefix=title_prefix)

    my_line_id, my_name = _current_bound_line_user_id()
    if kind == "user" and my_line_id and target_id != my_line_id:
        return {"ok": False, "error": "只能傳送到目前登入者綁定的個人 LINE。"}
    if kind == "group" and my_line_id and not _line_group_contains_user(target_id, my_line_id):
        return {"ok": False, "error": "你不是這個 LINE 群組成員，或官方帳號不在該群組，不能傳送。"}

    if cfg:
        if kind == "user" and not line_personal_user_allows_receive(cfg, "calendar"):
            return {"ok": False, "error": "此個人帳號未開放接收行事曆。"}
        if kind == "group" and not line_group_allows_receive(cfg, "calendar"):
            return {"ok": False, "error": "此群組未開放接收行事曆。"}
        if not permission_config_can_view(cfg, "calendar", event, kind, target_id):
            return {"ok": False, "error": "依權限設定，此對象不可查看這筆行程。"}

    res = line_push_messages(target_id, [msg])
    res["target_label"] = (cfg or {}).get("name") or ("個人" if kind == "user" else "群組")
    return res


def buyer_send_to_line_selected(buyer_id):
    res = _push_record_to_selected_destination("buyer", buyer_id, request.form.get("line_destination", ""), title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message_selected(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("buyer_detail", buyer_id=buyer_id))


def seller_send_to_line_selected(seller_id):
    res = _push_record_to_selected_destination("seller", seller_id, request.form.get("line_destination", ""), title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message_selected(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("seller_detail", seller_id=seller_id))


def development_send_to_line_selected(development_id):
    res = _push_record_to_selected_destination("development", development_id, request.form.get("line_destination", ""), title_prefix="後台傳送")
    msg, cat = _line_push_result_flash_message_selected(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("development_detail", development_id=development_id))


def calendar_send_to_line_selected(event_id):
    res = _push_calendar_to_selected_destination(event_id, request.form.get("line_destination", ""), title_prefix="後台傳送行程")
    msg, cat = _line_push_result_flash_message_selected(res)
    flash(msg, cat)
    return redirect(request.referrer or url_for("calendar_page"))

try:
    app.view_functions["buyer_send_to_line"] = login_required(buyer_send_to_line_selected)
    app.view_functions["seller_send_to_line"] = login_required(seller_send_to_line_selected)
    app.view_functions["development_send_to_line"] = login_required(development_send_to_line_selected)
    app.view_functions["calendar_send_to_line"] = login_required(calendar_send_to_line_selected)
    print("✅ LINE 後台傳送：已改成可選群組/個人，且群組需包含目前使用者與官方帳號")
except Exception as e:
    print("⚠️ LINE 可選傳送對象套用失敗：", e)


# ---------- 待辦事項：公開 / 個人 ----------

def _todo_visibility_payload_from_event(event=None):
    return _line_visibility_payload_from_event(event)


def _todo_normalize_visibility(data: dict):
    data = dict(data or {})
    visibility = (data.get("visibility") or "").strip()
    if visibility not in ("public", "personal"):
        # 舊資料兼容：個人 userId 建立的視為個人；後台共用/群組視為公開。
        if (data.get("line_target_type") == "user") or (str(data.get("line_target_id") or "").startswith("U")):
            visibility = "personal"
        else:
            visibility = "public"
    data["visibility"] = visibility
    if visibility == "personal":
        data["owner_line_user_id"] = (data.get("owner_line_user_id") or data.get("line_user_id") or (data.get("line_target_id") if str(data.get("line_target_id") or "").startswith("U") else "") or "").strip()
        data["owner_line_name"] = (data.get("owner_line_name") or data.get("sender_display_name") or "").strip()
    else:
        data["owner_line_user_id"] = data.get("owner_line_user_id") or ""
        data["owner_line_name"] = data.get("owner_line_name") or ""
    return data


def _todo_source_cfg_for_target(target_id: str):
    target_id = (target_id or "").strip()
    if target_id.startswith("U"):
        return "user", find_line_personal_user_by_user_id(target_id)
    if target_id.startswith("C") or target_id.startswith("R"):
        return "group", find_line_group_by_target_id(target_id)
    ctx = permission_current_line_source()
    return ctx.get("kind", ""), ctx.get("source_cfg")


def _todo_doc_visible_for_target(doc, target_id=""):
    try:
        raw = doc.to_dict() or {}
    except Exception:
        raw = {}
    data = _todo_normalize_visibility(raw)
    # 後台無 target 時顯示全部，後台頁面再用篩選控制。
    if not target_id:
        ctx = permission_current_line_source()
        if not ctx.get("source_cfg"):
            return True
        return permission_config_can_view(ctx.get("source_cfg"), "todo", data, ctx.get("kind"), ctx.get("target_id"))
    kind, cfg = _todo_source_cfg_for_target(target_id)
    if not cfg:
        # 舊相容：沒有設定的來源，只看完全綁在自己 target 的資料。
        return (data.get("line_target_id") or "") == target_id
    return permission_config_can_view(cfg, "todo", data, kind, target_id)


def _is_open_todo_doc(doc, target_id=''):
    data = doc.to_dict() or {}
    if data.get('status', 'open') != 'open':
        return False
    if not (data.get('todo_date') or '').strip():
        return False
    if target_id and not _todo_doc_visible_for_target(doc, target_id=target_id):
        return False
    if not target_id and not _todo_doc_visible_for_target(doc, target_id=''):
        return False
    return True


def _get_open_line_todos(todo_date='', target_id='', include_overdue=False):
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    result = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = (doc.to_dict() or {}).get('todo_date', '')
        if include_overdue:
            if d <= query_date:
                result.append(doc)
        else:
            if d == query_date:
                result.append(doc)
    return _sort_line_todo_docs(result)


def _get_overdue_line_todos(todo_date='', target_id=''):
    query_date = todo_date or now_taipei().strftime('%Y-%m-%d')
    result = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        if not _is_open_todo_doc(doc, target_id=target_id):
            continue
        d = (doc.to_dict() or {}).get('todo_date', '')
        if d < query_date:
            result.append(doc)
    return _sort_line_todo_docs(result)


def _find_line_todo(todo_key: str, target_id=''):
    key = (todo_key or '').strip()
    if not key:
        return None, '請提供代辦 ID 或事項關鍵字。'
    direct = db.collection(LINE_TODO_COLLECTION).document(key).get()
    if direct.exists:
        if target_id and not _todo_doc_visible_for_target(direct, target_id=target_id):
            return None, '依權限設定，這筆待辦不能在目前 LINE 來源查看或完成。'
        return direct, ''
    matches = []
    for doc in db.collection(LINE_TODO_COLLECTION).stream():
        data = doc.to_dict() or {}
        if data.get('status', 'open') != 'open':
            continue
        if target_id and not _todo_doc_visible_for_target(doc, target_id=target_id):
            continue
        title = data.get('title', '') or data.get('content', '')
        if doc.id.startswith(key) or key in title:
            matches.append(doc)
    if len(matches) == 1:
        return matches[0], ''
    if len(matches) > 1:
        preview = '\n'.join([f"- [{d.id[:6]}] {(d.to_dict() or {}).get('title','')}" for d in matches[:8]])
        return None, '找到多筆待辦，請用 ID 完成：\n' + preview
    return None, '找不到這筆未完成待辦，請先輸入 #今日待辦 查看 ID。'


try:
    _create_line_todo_before_visibility_patch = create_line_todo
except Exception:
    _create_line_todo_before_visibility_patch = None


def create_line_todo(fields, event):
    title = (fields.get('title') or '').strip()
    todo_date = _parse_line_todo_date(fields.get('todo_date') or fields.get('todo_date_raw') or '')
    note = (fields.get('note') or '').strip()
    if not title:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：請填「事項」。\n\n範例：\n#新增待辦\n日期: 明天\n事項: 打給王小姐確認貸款資料'}
    if not todo_date:
        return {'handled': True, 'ok': False, 'reply_text': '未新增：日期格式看不懂，請用 2026-05-29、5/29、今天、明天。'}

    target_id, target_type = _line_todo_target_from_event(event)
    source = event.get('source') or {}
    sender_display_name = get_line_sender_display_name(event)
    now = now_taipei().isoformat()
    vis_payload = _todo_visibility_payload_from_event(event)

    doc_ref = db.collection(LINE_TODO_COLLECTION).document()
    doc_ref.set({
        'title': title,
        'content': note,
        'todo_date': todo_date,
        'note': note,
        'status': 'open',
        'visibility': vis_payload.get('visibility'),
        'owner_line_user_id': vis_payload.get('owner_line_user_id', ''),
        'owner_line_name': vis_payload.get('owner_line_name', ''),
        'line_target_id': target_id,
        'line_target_type': target_type,
        'line_group_id': source.get('groupId', ''),
        'line_room_id': source.get('roomId', ''),
        'line_user_id': source.get('userId', ''),
        'sender_display_name': sender_display_name,
        'created_at': now,
        'created_by_id': 'line_bot',
        'created_by_name': sender_display_name or 'LINE Bot',
        'reminder_sent_dates': [],
    })
    vis_text = '個人待辦' if vis_payload.get('visibility') == 'personal' else '公開待辦'
    return {'handled': True, 'ok': True, 'reply_text': f"已新增{vis_text}：{title}\n日期：{todo_date}\nID：{doc_ref.id[:6]}", 'parsed_tag': '新增代辦'}


# 後台新增待辦：可選公開 / 個人。
def todos_new_visibility_compatible():
    title = (request.form.get('title') or '').strip()
    todo_date = calendar_safe_date(request.form.get('todo_date') or '')
    note = (request.form.get('note') or '').strip()
    visibility = (request.form.get('visibility') or 'public').strip()
    if visibility not in ('public', 'personal'):
        visibility = 'public'
    owner_line_user_id = (request.form.get('owner_line_user_id') or '').strip()
    owner_line_name = ''
    if visibility == 'personal':
        if not owner_line_user_id:
            owner_line_user_id, owner_line_name = _current_bound_line_user_id()
        else:
            cfg = find_line_personal_user_by_user_id(owner_line_user_id) or {}
            owner_line_name = cfg.get('name') or ''
    if not title:
        flash('請輸入待辦事項', 'warning')
        return redirect(url_for('todos_page', date=todo_date))
    db.collection(LINE_TODO_COLLECTION).add({
        'title': title,
        'content': note,
        'note': note,
        'todo_date': todo_date,
        'status': 'open',
        'visibility': visibility,
        'owner_line_user_id': owner_line_user_id if visibility == 'personal' else '',
        'owner_line_name': owner_line_name if visibility == 'personal' else '',
        'source': '後台',
        'line_target_id': owner_line_user_id if visibility == 'personal' else '',
        'line_target_type': 'user' if visibility == 'personal' else 'backend_shared',
        'created_at': now_taipei().isoformat(),
        'created_by_id': session.get('user_id'),
        'created_by_name': session.get('user_name'),
    })
    flash('已新增待辦事項', 'success')
    return redirect(url_for('todos_page', date=todo_date))

try:
    app.view_functions['todos_new'] = login_required(todos_new_visibility_compatible)
    print('✅ 待辦事項公開/個人設定已啟用')
except Exception as e:
    print('⚠️ 待辦事項公開/個人設定套用失敗：', e)

# =============================================================================
# LINE 個人 / 群組精準權限 + 可選傳送對象 + 待辦公開/個人 Patch End
# =============================================================================


# =============================================================================
# 後台個人資料隔離 Patch v20260623B
# - 後台登入者只能看到：公開資料 + 自己的個人資料
# - 其他人的個人客需 / 委託 / 開發 / 行事曆 / 待辦不會出現在列表、詳細頁、同屋主頁
# - 個人資料如果綁定 LINE userId，會以目前登入者的個人後台 LINE 綁定為準
# =============================================================================


def _backend_current_line_user_ids():
    """取得目前登入者已綁定的 LINE 個人 userId，供後台個人資料過濾使用。"""
    ids = set()
    try:
        binding = get_current_user_line_binding()
        if (binding.get("user_id") or "").strip():
            ids.add((binding.get("user_id") or "").strip())
    except Exception:
        pass

    try:
        _, user_doc = _current_crm_user_doc()
        for key in ("line_user_id", "line_personal_user_id"):
            val = (user_doc.get(key) or "").strip()
            if val:
                ids.add(val)
    except Exception:
        pass
    return ids


def _backend_normalize_record_visibility(data: dict):
    data = dict(data or {})
    visibility = (data.get("visibility") or "").strip()
    if visibility not in ("public", "personal"):
        # 舊資料預設公開，避免既有資料突然消失；有 U 開頭 target 的舊待辦視為個人。
        if str(data.get("line_target_id") or "").startswith("U") or (data.get("line_target_type") == "user"):
            visibility = "personal"
        else:
            visibility = "public"
    data["visibility"] = visibility
    return data


def backend_can_view_personal_record(data: dict) -> bool:
    """
    後台頁面用：
    - 公開資料：所有登入者可看
    - 個人資料：只能該登入者看
      1) 若有 owner_line_user_id，以登入者個人後台綁定 LINE ID 比對
      2) 若沒有 owner_line_user_id，才以 owner_user_id / created_by_id 比對
    """
    data = _backend_normalize_record_visibility(data)
    if data.get("visibility") != "personal":
        return True

    uid = (session.get("user_id") or "").strip()
    if not uid:
        return False

    owner_line_id = (
        data.get("owner_line_user_id")
        or data.get("line_user_id")
        or (data.get("line_target_id") if str(data.get("line_target_id") or "").startswith("U") else "")
        or ""
    ).strip()

    if owner_line_id:
        return owner_line_id in _backend_current_line_user_ids()

    owner_user_id = (data.get("owner_user_id") or "").strip()
    created_by_id = (data.get("created_by_id") or "").strip()
    return bool((owner_user_id and owner_user_id == uid) or (created_by_id and created_by_id == uid))


def _backend_visible_items(items):
    return [x for x in (items or []) if backend_can_view_personal_record(x)]


def _backend_can_view_doc(collection_name: str, doc_id: str) -> bool:
    if not doc_id:
        return False
    try:
        snap = db.collection(collection_name).document(doc_id).get()
        if not snap.exists:
            return False
        return backend_can_view_personal_record(snap.to_dict() or {})
    except Exception as e:
        print("⚠️ 後台權限檢查失敗：", collection_name, doc_id, e)
        return False


def _backend_forbidden_redirect(label="資料", fallback_endpoint="buyers"):
    flash(f"這筆{label}是其他使用者的個人資料，無法查看或操作。", "warning")
    try:
        return redirect(request.referrer or url_for(fallback_endpoint))
    except Exception:
        return redirect(url_for("buyers"))


# ---------- 客需列表：只顯示公開 + 自己個人 ----------
def buyers_backend_visible():
    q = request.args.get("q", "").strip()
    level = request.args.get("level", "").strip()
    intent_type = request.args.get("intent_type", "").strip()
    stage = request.args.get("stage", "").strip()
    source = request.args.get("source", "").strip()
    label = request.args.get("label", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")

    docs = db.collection("buyers").stream()
    visible_buyers = _backend_visible_items([doc_to_dict(d) for d in docs])

    source_options = sorted({(b.get("source") or "").strip() for b in visible_buyers if (b.get("source") or "").strip()})
    label_options = build_label_options(visible_buyers)
    buyers_list = list(visible_buyers)

    if q:
        buyers_list = [b for b in buyers_list if q in (b.get("name") or "") or q in (b.get("phone") or "")]
    if level:
        buyers_list = [b for b in buyers_list if b.get("level") == level]
    if intent_type:
        buyers_list = [b for b in buyers_list if b.get("intent_type") == intent_type]
    if stage:
        buyers_list = [b for b in buyers_list if (b.get("stage") or "") == stage]
    if source:
        buyers_list = [b for b in buyers_list if (b.get("source") or "") == source]
    if label:
        buyers_list = [b for b in buyers_list if label in ensure_list(b.get("labels"))]

    if sort_by == "created_at_asc":
        buyers_list.sort(key=lambda b: b.get("created_at") or "")
    elif sort_by == "created_at_desc":
        buyers_list.sort(key=lambda b: b.get("created_at") or "", reverse=True)
    elif sort_by == "name_asc":
        buyers_list.sort(key=lambda b: (b.get("name") or ""))
    elif sort_by == "name_desc":
        buyers_list.sort(key=lambda b: (b.get("name") or ""), reverse=True)

    buyers_list = attach_latest_followup(buyers_list, "buyer_followups", "buyer_id")
    return render_template(
        "buyers.html",
        buyers=buyers_list,
        q=q,
        level=level,
        intent_type=intent_type,
        stage=stage,
        source=source,
        source_options=source_options,
        label=label,
        label_options=label_options,
        sort_by=sort_by,
        buyer_stage_options=BUYER_STAGE_OPTIONS,
        total_count=len(visible_buyers),
        filtered_count=len(buyers_list),
    )


# ---------- 委託列表：只顯示公開 + 自己個人 ----------
def sellers_backend_visible():
    q = request.args.get("q", "").strip()
    level = request.args.get("level", "").strip()
    stage = request.args.get("stage", "").strip()
    source = request.args.get("source", "").strip()
    label = request.args.get("label", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")

    docs = db.collection("sellers").stream()
    visible_sellers = _backend_visible_items([doc_to_dict(d) for d in docs])

    source_options = sorted({(s.get("source") or "").strip() for s in visible_sellers if (s.get("source") or "").strip()})
    label_options = build_label_options(visible_sellers)
    sellers_list = list(visible_sellers)

    if q:
        sellers_list = [s for s in sellers_list if q in (s.get("name") or "") or q in (s.get("phone") or "")]
    if level:
        sellers_list = [s for s in sellers_list if s.get("level") == level]
    if stage:
        sellers_list = [s for s in sellers_list if (s.get("stage") or "") == stage]
    if source:
        sellers_list = [s for s in sellers_list if (s.get("source") or "") == source]
    if label:
        sellers_list = [s for s in sellers_list if label in ensure_list(s.get("labels"))]

    if sort_by == "created_at_asc":
        sellers_list.sort(key=lambda s: s.get("created_at") or "")
    elif sort_by == "created_at_desc":
        sellers_list.sort(key=lambda s: s.get("created_at") or "", reverse=True)
    elif sort_by == "name_asc":
        sellers_list.sort(key=lambda s: (s.get("name") or ""))
    elif sort_by == "name_desc":
        sellers_list.sort(key=lambda s: (s.get("name") or ""), reverse=True)

    sellers_list = attach_latest_followup(sellers_list, "seller_followups", "seller_id")
    return render_template(
        "sellers.html",
        sellers=sellers_list,
        q=q,
        level=level,
        stage=stage,
        source=source,
        source_options=source_options,
        label=label,
        label_options=label_options,
        sort_by=sort_by,
        seller_stage_options=SELLER_STAGE_OPTIONS,
        total_count=len(visible_sellers),
        filtered_count=len(sellers_list),
    )


# ---------- 開發列表：同屋主分組也只用可見資料計算 ----------
def developments_backend_visible_grouped():
    q = request.args.get("q", "").strip()
    current_stage = request.args.get("current_stage", "").strip()
    next_action = request.args.get("next_action", "").strip()
    source = request.args.get("source", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")
    show_done = request.args.get("show_done", "").strip()

    docs = db.collection("developments").stream()
    visible_items = _backend_visible_items([doc_to_dict(d) for d in docs])
    total_count = len(visible_items)
    source_options = sorted({(x.get("source") or "").strip() for x in visible_items if (x.get("source") or "").strip()})

    attach_development_owner_groups(visible_items)
    items = list(visible_items)

    if q:
        q_key = development_owner_key_from_phone(q)
        items = [
            x for x in items
            if q in (x.get("name") or "")
            or q in (x.get("phone") or "")
            or q in (x.get("address") or "")
            or q in (x.get("registered_address") or "")
            or (q_key and q_key == (x.get("owner_key") or development_owner_key_from_phone(x.get("phone", ""))))
        ]
    if current_stage:
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") == current_stage]
    if next_action:
        items = [x for x in items if (x.get("next_action") or "") == next_action]
    if source:
        items = [x for x in items if (x.get("source") or "") == source]
    if show_done != "1":
        items = [x for x in items if (x.get("current_stage") or x.get("stage") or "") not in DEVELOPMENT_HIDDEN_BY_DEFAULT]

    if sort_by == "created_at_asc":
        items.sort(key=lambda x: x.get("created_at") or "")
    elif sort_by == "created_at_desc":
        items.sort(key=lambda x: x.get("created_at") or "", reverse=True)
    elif sort_by == "name_asc":
        items.sort(key=lambda x: (x.get("name") or ""))
    elif sort_by == "name_desc":
        items.sort(key=lambda x: (x.get("name") or ""), reverse=True)
    else:
        items.sort(key=lambda x: x.get("created_at") or "", reverse=True)

    return render_template(
        "developments.html",
        developments=items,
        q=q,
        current_stage=current_stage,
        next_action=next_action,
        source=source,
        source_options=source_options,
        show_done=show_done,
        sort_by=sort_by,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
        total_count=total_count,
        filtered_count=len(items),
        label_docx_enabled=(next_action == "寄開發信"),
        label_docx_count=len([x for x in items if (x.get("registered_address") or "").strip()]),
    )


# 同屋主物件清單：只列出目前登入者可看的物件。
try:
    _get_development_same_owner_items_before_backend_isolation = get_development_same_owner_items
except Exception:
    _get_development_same_owner_items_before_backend_isolation = None


def get_development_same_owner_items(phone: str, exclude_id: str = ""):
    key = development_owner_key_from_phone(phone)
    if not key:
        return []
    result = []
    for doc in db.collection("developments").stream():
        item = doc_to_dict(doc)
        if not backend_can_view_personal_record(item):
            continue
        item_key = item.get("owner_key") or development_owner_key_from_phone(item.get("phone", ""))
        if item_key == key and (not exclude_id or item.get("id") != exclude_id):
            result.append(item)
    result.sort(key=lambda x: (x.get("created_at") or "", x.get("address") or ""))
    attach_development_owner_groups(result)
    return result


# 開發同屋主獨立頁：只列出目前登入者可看的物件。
def development_owner_group_backend_visible(owner_key):
    owner_key = development_owner_key_from_phone(owner_key)
    items = []
    owner_name = "同屋主"
    owner_phone = owner_key
    for doc in db.collection("developments").stream():
        item = doc_to_dict(doc)
        if not backend_can_view_personal_record(item):
            continue
        item_key = item.get("owner_key") or development_owner_key_from_phone(item.get("phone", ""))
        if item_key == owner_key:
            items.append(item)
            owner_name = item.get("name") or owner_name
            owner_phone = item.get("phone") or owner_phone
    items.sort(key=lambda x: (x.get("created_at") or "", x.get("address") or ""))
    attach_development_owner_groups(items)
    return render_template(
        "development_owner_group.html",
        owner_key=owner_key,
        owner_name=owner_name,
        owner_phone=owner_phone,
        items=items,
        development_current_stage_options=DEVELOPMENT_STATUS_OPTIONS,
        development_next_action_options=DEVELOPMENT_NEXT_ACTION_OPTIONS,
    )


# 行事曆：後台個人行程依「登入者 LINE 綁定 / owner_user_id」過濾。
try:
    _fetch_calendar_events_before_backend_isolation = fetch_calendar_events
except Exception:
    _fetch_calendar_events_before_backend_isolation = None


def fetch_calendar_events(start_date: str, end_date: str = None, line_only: bool = False, calendar_view: str = "all", line_source_kind: str = "", line_source_id: str = ""):
    events = _fetch_calendar_events_before_backend_isolation(start_date, end_date, line_only=line_only, calendar_view=calendar_view, line_source_kind=line_source_kind, line_source_id=line_source_id)
    # LINE 查詢仍走原本權限矩陣；只加強後台查詢。
    if line_only or line_source_kind or line_source_id:
        return events
    filtered = []
    for item in events:
        if not backend_can_view_personal_record(item):
            continue
        if calendar_view == "public" and (item.get("visibility") or "public") != "public":
            continue
        if calendar_view == "personal" and (item.get("visibility") or "public") != "personal":
            continue
        filtered.append(item)
    return filtered


# 待辦後台：只顯示公開待辦 + 自己個人待辦。
def todos_page_backend_visible():
    selected_date = calendar_safe_date(request.args.get("date", ""))
    show_done = request.args.get("show_done", "") == "1"

    docs = db.collection(LINE_TODO_COLLECTION).where("todo_date", "==", selected_date).stream()
    all_items = [_todo_doc_to_dict(d) for d in docs]
    all_items = _backend_visible_items(all_items)
    all_items.sort(key=_todo_item_sort_key)

    open_count = len([x for x in all_items if x.get("status", "open") != "done"])
    done_count = len([x for x in all_items if x.get("status", "open") == "done"])
    items = all_items if show_done else [x for x in all_items if x.get("status", "open") != "done"]

    return render_template(
        "todos.html",
        items=items,
        selected_date=selected_date,
        selected_date_label=format_calendar_date_label(selected_date),
        show_done=show_done,
        open_count=open_count,
        done_count=done_count,
        total_count=len(all_items),
    )


def _guard_backend_record(collection_name, record_id, label, fallback_endpoint):
    if not _backend_can_view_doc(collection_name, record_id):
        return _backend_forbidden_redirect(label=label, fallback_endpoint=fallback_endpoint)
    return None


def _wrap_backend_access(endpoint, collection_name, id_param, label, fallback_endpoint):
    old_func = app.view_functions.get(endpoint)
    if not old_func:
        return
    def wrapped(*args, **kwargs):
        rid = kwargs.get(id_param)
        if session.get("user_id"):
            blocked = _guard_backend_record(collection_name, rid, label, fallback_endpoint)
            if blocked:
                return blocked
        return old_func(*args, **kwargs)
    wrapped.__name__ = getattr(old_func, "__name__", endpoint) + "_backend_isolated"
    app.view_functions[endpoint] = wrapped


# 套用列表頁。
try:
    app.view_functions["buyers"] = login_required(buyers_backend_visible)
    app.view_functions["sellers"] = login_required(sellers_backend_visible)
    app.view_functions["developments"] = login_required(developments_backend_visible_grouped)
    app.view_functions["development_owner_group"] = login_required(development_owner_group_backend_visible)
    app.view_functions["todos_page"] = login_required(todos_page_backend_visible)
except Exception as e:
    print("⚠️ 後台列表個人資料隔離套用失敗：", e)


# 套用詳細 / 編輯 / 刪除 / 追蹤 / 傳送等操作保護。
for _endpoint, _collection, _id_param, _label, _fallback in [
    ("buyer_detail", "buyers", "buyer_id", "客需", "buyers"),
    ("buyer_edit", "buyers", "buyer_id", "客需", "buyers"),
    ("buyer_delete", "buyers", "buyer_id", "客需", "buyers"),
    ("buyer_quick_stage", "buyers", "buyer_id", "客需", "buyers"),
    ("add_buyer_followup", "buyers", "buyer_id", "客需", "buyers"),
    ("buyer_followup_edit", "buyers", "buyer_id", "客需", "buyers"),
    ("buyer_followup_delete", "buyers", "buyer_id", "客需", "buyers"),
    ("buyer_send_to_line", "buyers", "buyer_id", "客需", "buyers"),

    ("seller_detail", "sellers", "seller_id", "委託", "sellers"),
    ("seller_edit", "sellers", "seller_id", "委託", "sellers"),
    ("seller_delete", "sellers", "seller_id", "委託", "sellers"),
    ("seller_quick_stage", "sellers", "seller_id", "委託", "sellers"),
    ("add_seller_followup", "sellers", "seller_id", "委託", "sellers"),
    ("seller_followup_edit", "sellers", "seller_id", "委託", "sellers"),
    ("seller_followup_delete", "sellers", "seller_id", "委託", "sellers"),
    ("seller_send_to_line", "sellers", "seller_id", "委託", "sellers"),

    ("development_detail", "developments", "development_id", "開發", "developments"),
    ("development_edit", "developments", "development_id", "開發", "developments"),
    ("development_delete", "developments", "development_id", "開發", "developments"),
    ("development_quick_flow", "developments", "development_id", "開發", "developments"),
    ("add_development_followup", "developments", "development_id", "開發", "developments"),
    ("development_followup_edit", "developments", "development_id", "開發", "developments"),
    ("development_followup_delete", "developments", "development_id", "開發", "developments"),
    ("development_send_to_line", "developments", "development_id", "開發", "developments"),

    ("calendar_edit", CALENDAR_EVENT_COLLECTION if 'CALENDAR_EVENT_COLLECTION' in globals() else 'calendar_events', "event_id", "行程", "calendar_page"),
    ("calendar_delete", CALENDAR_EVENT_COLLECTION if 'CALENDAR_EVENT_COLLECTION' in globals() else 'calendar_events', "event_id", "行程", "calendar_page"),
    ("calendar_send_to_line", CALENDAR_EVENT_COLLECTION if 'CALENDAR_EVENT_COLLECTION' in globals() else 'calendar_events', "event_id", "行程", "calendar_page"),
    ("todos_toggle", LINE_TODO_COLLECTION if 'LINE_TODO_COLLECTION' in globals() else 'line_todos', "todo_id", "待辦", "todos_page"),
    ("todos_delete", LINE_TODO_COLLECTION if 'LINE_TODO_COLLECTION' in globals() else 'line_todos', "todo_id", "待辦", "todos_page"),
]:
    try:
        _wrap_backend_access(_endpoint, _collection, _id_param, _label, _fallback)
    except Exception as e:
        print("⚠️ 後台操作個人資料隔離套用失敗：", _endpoint, e)

print("✅ 後台個人資料隔離已啟用：登入者只能看到公開資料 + 自己的個人資料")

# =============================================================================
# 後台個人資料隔離 Patch End
# =============================================================================


# =============================================================================
# Gemini AI 物件推薦 + LINE #推薦物件 指令 Patch v20260624
# - LINE 輸入 #推薦物件，可用電話 / 客戶ID 查詢客需並回覆物件卡片
# - 客需卡片下方新增「推薦物件」按鈕，點擊後直接回覆符合條件的物件卡片
# - 物件網址會顯示在每張推薦卡片的「查看物件」按鈕
# - 物件資料來源：Firestore properties；出售 sale、出租 rent 分開讀取
# =============================================================================

GEMINI_DEFAULT_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")
AI_PROPERTY_MAX_CANDIDATES = int(os.environ.get("AI_PROPERTY_MAX_CANDIDATES", "25") or 25)
AI_PROPERTY_DEFAULT_TOP_N = int(os.environ.get("AI_PROPERTY_DEFAULT_TOP_N", "5") or 5)


def _ai_safe_float(value, default=None):
    if value is None:
        return default
    try:
        text = str(value).replace(",", "").replace("萬", "").replace("元", "").strip()
        m = re.search(r"-?\d+(?:\.\d+)?", text)
        if not m:
            return default
        return float(m.group(0))
    except Exception:
        return default


def _ai_safe_int(value, default=None):
    num = _ai_safe_float(value, None)
    if num is None:
        return default
    try:
        return int(round(num))
    except Exception:
        return default


def _ai_norm(text):
    return str(text or "").strip()


def _ai_compact_text(text, max_len=240):
    text = re.sub(r"\s+", " ", str(text or "").strip())
    if len(text) <= max_len:
        return text
    return text[:max_len - 1] + "…"




def _ai_should_show_risk(text):
    """只有 Gemini 真的有判斷出具體風險時才顯示。"""
    text = re.sub(r"\s+", " ", str(text or "").strip())
    if not text:
        return False
    bad_phrases = [
        "尚未由 Gemini",
        "尚未設定 Gemini",
        "規則粗篩",
        "建議再人工確認",
        "需現場確認細節",
        "可再人工確認",
        "無",
        "-",
        "—",
        "無明顯風險",
        "目前沒有",
    ]
    return not any(p in text for p in bad_phrases)

try:
    app.jinja_env.globals["_ai_should_show_risk"] = _ai_should_show_risk
except Exception:
    pass

def _ai_json_loads(text):
    if not text:
        return {}
    text = str(text).strip()
    text = re.sub(r"^```json", "", text, flags=re.I).strip()
    text = re.sub(r"^```", "", text).strip()
    text = re.sub(r"```$", "", text).strip()
    try:
        return json.loads(text)
    except Exception:
        pass
    m = re.search(r"\{.*\}", text, re.S)
    if m:
        try:
            return json.loads(m.group(0))
        except Exception:
            return {}
    return {}


def _gemini_generate_json(prompt: str):
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("尚未設定 GEMINI_API_KEY")
    try:
        from google import genai
        client = genai.Client(api_key=api_key)
        res = client.models.generate_content(
            model=os.environ.get("GEMINI_MODEL", GEMINI_DEFAULT_MODEL),
            contents=prompt,
        )
        return _ai_json_loads(getattr(res, "text", "") or "")
    except Exception as e:
        raise RuntimeError(f"Gemini 呼叫失敗：{e}")


def _ai_split_keywords(text):
    text = str(text or "")
    parts = re.split(r"[，,、/\n\s]+", text)
    return [p.strip() for p in parts if p.strip()]


def _ai_parse_room_min(value):
    text = str(value or "")
    if not text:
        return None
    m = re.search(r"(\d+)\s*房", text)
    if m:
        return int(m.group(1))
    m = re.search(r"\d+", text)
    return int(m.group(0)) if m else None


def _ai_parse_age_max_from_text(text):
    text = str(text or "")
    # 例如：不要超過30年、30年內、屋齡30以下
    m = re.search(r"(?:屋齡)?\s*(?:不要超過|不超過|低於|小於|以下|內)?\s*(\d+)\s*年", text)
    if m and any(k in text for k in ["屋齡", "年內", "以下", "不超過", "不要超過", "低於", "小於", "新"]):
        return int(m.group(1))
    return None


def _ai_parse_ping_min_from_text(text, keyword_group=None):
    """
    從口語條件抓坪數下限。
    例如：主建物坪30坪以上、主建30坪以上、室內25坪以上、使用坪數20坪以上、建物40坪以上。
    """
    text = str(text or "")
    if not text.strip():
        return None
    keyword_group = keyword_group or ["坪"]
    # 關鍵字在數字前：主建物坪30坪以上 / 室內 25 坪以上
    for kw in keyword_group:
        pattern = rf"{re.escape(kw)}[^0-9]{{0,8}}(\d+(?:\.\d+)?)\s*坪?\s*(?:以上|起|up|UP|至少|不低於|大於|超過)?"
        m = re.search(pattern, text)
        if m:
            return float(m.group(1))
    # 數字在關鍵字前：30坪以上主建物 / 25坪以上室內
    for kw in keyword_group:
        pattern = rf"(\d+(?:\.\d+)?)\s*坪\s*(?:以上|起|至少|不低於|大於|超過)?[^，,。；;\n]{{0,12}}{re.escape(kw)}"
        m = re.search(pattern, text)
        if m:
            return float(m.group(1))
    return None


def _ai_parse_budget_max_buy_wan(value):
    text = str(value or "")
    num = _ai_safe_float(text, None)
    if num is None:
        return None
    # 如果輸入像 15000000，轉成萬；如果本來就是 1500萬，就維持 1500。
    if num > 100000:
        return round(num / 10000, 2)
    return num


def _ai_buyer_need_to_parsed(buyer: dict, extra_text: str = ""):
    buyer = buyer or {}
    raw_all = " ".join([
        _ai_norm(buyer.get("name")),
        _ai_norm(buyer.get("intent_type")),
        _ai_norm(buyer.get("budget_max")),
        _ai_norm(buyer.get("rent_max")),
        _ai_norm(buyer.get("preferred_areas")),
        _ai_norm(buyer.get("property_type")),
        _ai_norm(buyer.get("room_range")),
        _ai_norm(buyer.get("car_need")),
        _ai_norm(buyer.get("requirement_must")),
        _ai_norm(buyer.get("requirement_nice")),
        _ai_norm(buyer.get("note")),
        _ai_norm(extra_text),
    ])
    intent = (buyer.get("intent_type") or "").strip()
    if intent in ("rent", "租", "租屋", "承租"):
        intent_type = "rent"
    else:
        intent_type = "buy"

    budget_max = _ai_parse_budget_max_buy_wan(buyer.get("budget_max") or buyer.get("budget") or "")
    rent_max = _ai_safe_float(buyer.get("rent_max") or buyer.get("rent") or "", None)
    areas = _ai_split_keywords(buyer.get("preferred_areas") or buyer.get("area") or "")
    property_types = _ai_split_keywords(buyer.get("property_type") or "")
    room_min = _ai_parse_room_min(buyer.get("room_range") or "")
    age_max = _ai_parse_age_max_from_text(raw_all)
    main_building_area_min = _ai_parse_ping_min_from_text(raw_all, ["主建物坪", "主建物", "主建", "室內坪", "室內", "使用坪數", "使用坪"])
    building_area_min = _ai_parse_ping_min_from_text(raw_all, ["建物坪", "建物面積", "建坪", "權狀坪", "登記坪數", "權狀面積"])
    need_parking = None
    if any(k in raw_all for k in ["車位", "停車", "平車", "機械", "雙車", "車庫"]):
        need_parking = True
    if any(k in raw_all for k in ["不用車位", "無車位", "不要車位"]):
        need_parking = False

    keywords = []
    for k in ["有空地", "大地坪", "前院", "後院", "庭院", "可放東西", "雙車", "平車", "電梯", "低樓層", "屋況佳", "新一點", "近商圈", "生活機能", "可寵", "可開伙"]:
        if k in raw_all:
            keywords.append(k)
    keywords.extend(_ai_split_keywords(buyer.get("requirement_must") or ""))
    keywords.extend(_ai_split_keywords(buyer.get("requirement_nice") or ""))
    keywords.extend(_ai_split_keywords(extra_text or ""))

    return {
        "intent_type": intent_type,
        "budget_max_buy_wan": budget_max,
        "rent_max": rent_max,
        "areas": areas,
        "property_types": property_types,
        "room_min": room_min,
        "age_max": age_max,
        "main_building_area_min": main_building_area_min,
        "building_area_min": building_area_min,
        "need_parking": need_parking,
        "must_have": _ai_split_keywords(buyer.get("requirement_must") or ""),
        "nice_to_have": _ai_split_keywords(buyer.get("requirement_nice") or ""),
        "exclude": [],
        "search_keywords": dedupe_keep_order([k for k in keywords if k]),
        "summary": _ai_compact_text(raw_all, 180),
    }


def _ai_try_parse_buyer_need_with_gemini(buyer: dict, extra_text: str = ""):
    parsed_fallback = _ai_buyer_need_to_parsed(buyer, extra_text=extra_text)
    if not os.environ.get("GEMINI_API_KEY", "").strip():
        return parsed_fallback
    raw_text = "\n".join([
        f"姓名：{buyer.get('name','')}",
        f"需求類型：{buyer.get('intent_type','')}",
        f"預算：{buyer.get('budget_max') or buyer.get('rent_max') or ''}",
        f"區域：{buyer.get('preferred_areas','')}",
        f"產品類型：{buyer.get('property_type','')}",
        f"房數：{buyer.get('room_range','')}",
        f"車位：{buyer.get('car_need','')}",
        f"必要條件：{buyer.get('requirement_must','')}",
        f"加分條件：{buyer.get('requirement_nice','')}",
        f"備註：{buyer.get('note','')}",
        f"補充：{extra_text or ''}",
    ]).strip()
    prompt = f"""
你是台中海線房仲 CRM 的客需解析助手。
請把以下客需整理成 JSON。只輸出 JSON，不要有其他文字。

客需：
{raw_text}

輸出格式：
{{
  "intent_type": "buy 或 rent",
  "budget_max_buy_wan": 數字或 null,
  "rent_max": 數字或 null,
  "areas": ["清水區", "梧棲區"],
  "property_types": ["透天", "別墅", "電梯大樓", "華廈", "土地", "農舍", "店面", "廠房", "公寓"],
  "room_min": 數字或 null,
  "age_max": 數字或 null,
  "main_building_area_min": 主建物坪數下限，數字或 null,
  "building_area_min": 建物/權狀坪數下限，數字或 null,
  "need_parking": true/false/null,
  "must_have": ["必要條件"],
  "nice_to_have": ["加分條件"],
  "exclude": ["排除條件"],
  "search_keywords": ["用於物件描述比對的關鍵字"],
  "summary": "一句話需求整理"
}}
""".strip()
    try:
        parsed = _gemini_generate_json(prompt)
        if not parsed:
            return parsed_fallback
        # 保留 fallback 中有抓到但 Gemini 漏掉的資訊。
        for k, v in parsed_fallback.items():
            if parsed.get(k) in (None, "", [], {}):
                parsed[k] = v
        return parsed
    except Exception as e:
        print("⚠️ Gemini 解析客需失敗，改用規則解析：", e)
        return parsed_fallback


def _property_field(prop: dict, *keys, default=""):
    for k in keys:
        if k in prop and prop.get(k) not in (None, ""):
            return prop.get(k)
    raw = prop.get("raw") or {}
    if isinstance(raw, dict):
        for k in keys:
            if k in raw and raw.get(k) not in (None, ""):
                return raw.get(k)
    return default


def _property_normalized(prop: dict):
    prop = prop or {}
    title = _ai_norm(_property_field(prop, "title", "rakuya_物件名稱", "物件名稱", "案名"))
    area = _ai_norm(_property_field(prop, "area", "rakuya_行政區", "行政區"))
    address = _ai_norm(_property_field(prop, "address", "rakuya_地址", "地址"))
    ptype = _ai_norm(_property_field(prop, "property_type", "rakuya_現況類型", "現況類型", "產品類型"))
    desc = _ai_norm(_property_field(prop, "description", "rakuya_特色描述", "特色描述", "描述", "searchable_text"))
    url = _ai_norm(_property_field(prop, "url", "rakuya_刊登來源網址", "網址", "物件網址", "source_url"))
    price_wan = _ai_safe_float(_property_field(prop, "price_wan", "rakuya_總價_萬", "總價_萬", "總價"), None)
    rent_price = _ai_safe_float(_property_field(prop, "rent_price", "rakuya_租金_元", "租金", "租金_元"), None)
    rooms = _ai_safe_float(_property_field(prop, "rooms", "rakuya_房", "房"), None)
    halls = _ai_safe_float(_property_field(prop, "halls", "rakuya_廳", "廳"), None)
    baths = _ai_safe_float(_property_field(prop, "baths", "rakuya_衛", "衛"), None)
    age = _ai_safe_float(_property_field(prop, "age", "rakuya_屋齡", "屋齡"), None)
    main_building_area = _ai_safe_float(_property_field(
        prop,
        "main_building_area", "rakuya_主建物", "raw_主建物坪", "主建物", "主建物坪",
        "rakuya_使用坪數", "使用坪數", "室內坪"
    ), None)
    building_area = _ai_safe_float(_property_field(
        prop,
        "building_area", "rakuya_建物登記", "rakuya_權狀面積", "raw_建物面積",
        "raw_登記坪數", "建物面積", "權狀面積", "登記坪數"
    ), None)
    parking = _ai_norm(_property_field(prop, "parking", "車位", "rakuya_車位"))
    full_text = " ".join([
        title, area, address, ptype, desc, parking, _ai_norm(prop.get("searchable_text")),
        f"主建物{main_building_area}坪" if main_building_area is not None else "",
        f"建物{building_area}坪" if building_area is not None else ""
    ])
    return {
        "id": prop.get("id") or "",
        "deal_type": prop.get("deal_type") or "sale",
        "title": title or "未命名物件",
        "area": area,
        "address": address,
        "property_type": ptype,
        "description": desc,
        "url": url,
        "price_wan": price_wan,
        "rent_price": rent_price,
        "rooms": rooms,
        "halls": halls,
        "baths": baths,
        "age": age,
        "main_building_area": main_building_area,
        "building_area": building_area,
        "parking": parking,
        "searchable_text": full_text,
        "raw": prop,
    }


def _fetch_active_properties(deal_type="sale", max_docs=2500):
    deal_type = "rent" if deal_type == "rent" else "sale"
    docs = []
    try:
        docs = list(db.collection("properties").where("deal_type", "==", deal_type).limit(max_docs).stream())
    except Exception:
        docs = list(db.collection("properties").limit(max_docs).stream())
    items = []
    for d in docs:
        data = d.to_dict() or {}
        data["id"] = d.id
        if (data.get("deal_type") or deal_type) != deal_type:
            continue
        if data.get("active") is False:
            continue
        items.append(_property_normalized(data))
    return items


def _hard_filter_properties_for_buyer(parsed_need: dict, top_limit=AI_PROPERTY_MAX_CANDIDATES):
    deal_type = "rent" if (parsed_need.get("intent_type") == "rent") else "sale"
    props = _fetch_active_properties(deal_type=deal_type)
    areas = [str(a).replace("台中市", "").strip() for a in (parsed_need.get("areas") or []) if str(a).strip()]
    ptypes = [str(t).strip() for t in (parsed_need.get("property_types") or []) if str(t).strip()]
    keywords = [str(k).strip() for k in (parsed_need.get("search_keywords") or []) if str(k).strip()]
    budget_max = _ai_safe_float(parsed_need.get("budget_max_buy_wan"), None)
    rent_max = _ai_safe_float(parsed_need.get("rent_max"), None)
    room_min = _ai_safe_float(parsed_need.get("room_min"), None)
    age_max = _ai_safe_float(parsed_need.get("age_max"), None)
    main_building_area_min = _ai_safe_float(parsed_need.get("main_building_area_min"), None)
    building_area_min = _ai_safe_float(parsed_need.get("building_area_min"), None)
    need_parking = parsed_need.get("need_parking")

    scored = []
    for p in props:
        full_text = p.get("searchable_text") or ""
        score = 0
        reasons = []

        # 價格 / 租金是硬條件；若資料沒有價格，不直接排除但不加分。
        if deal_type == "sale" and budget_max:
            if p.get("price_wan") is not None:
                if p["price_wan"] <= budget_max:
                    score += 25
                    reasons.append("總價符合預算")
                else:
                    continue
        if deal_type == "rent" and rent_max:
            if p.get("rent_price") is not None:
                if p["rent_price"] <= rent_max:
                    score += 25
                    reasons.append("租金符合預算")
                else:
                    continue

        if areas:
            if any(a and (a in p.get("area", "") or a in p.get("address", "") or a in full_text) for a in areas):
                score += 22
                reasons.append("區域符合")
            else:
                continue

        if ptypes:
            if any(t and (t in p.get("property_type", "") or t in p.get("title", "") or t in full_text) for t in ptypes):
                score += 18
                reasons.append("產品類型符合")
            else:
                # 類型不完全符合仍可作候選，但分數較低，給 AI 判斷接近性。
                score += 3

        if room_min:
            if p.get("rooms") is not None and p["rooms"] >= room_min:
                score += 8
                reasons.append("房數符合")
            elif p.get("rooms") is not None:
                score -= 4

        if age_max:
            if p.get("age") is not None and p["age"] <= age_max:
                score += 10
                reasons.append("屋齡符合")
            elif p.get("age") is not None:
                score -= 8

        # 坪數條件是硬條件：使用者指定主建物/室內/使用坪數時，沒有達標就排除。
        if main_building_area_min:
            if p.get("main_building_area") is not None and p["main_building_area"] >= main_building_area_min:
                score += 18
                reasons.append(f"主建物坪數符合 {main_building_area_min:g} 坪以上")
            else:
                continue

        if building_area_min:
            if p.get("building_area") is not None and p["building_area"] >= building_area_min:
                score += 12
                reasons.append(f"建物坪數符合 {building_area_min:g} 坪以上")
            else:
                continue

        if need_parking is True:
            if any(k in full_text for k in ["車位", "平車", "雙車", "停車", "車庫", "前院"]):
                score += 9
                reasons.append("可能符合停車需求")
            else:
                score -= 3

        matched_keywords = [k for k in keywords if k and k in full_text]
        if matched_keywords:
            score += min(12, len(matched_keywords) * 3)
            reasons.append("關鍵字符合：" + "、".join(matched_keywords[:4]))

        # 基本資料完整度。
        if p.get("url"):
            score += 2
        if p.get("description"):
            score += 2

        item = dict(p)
        item["rule_score"] = score
        item["basic_reasons"] = dedupe_keep_order(reasons)
        scored.append(item)

    scored.sort(key=lambda x: x.get("rule_score", 0), reverse=True)
    return scored[:max(1, int(top_limit or AI_PROPERTY_MAX_CANDIDATES))]


def _rank_properties_with_gemini(parsed_need: dict, candidates: list, top_n=AI_PROPERTY_DEFAULT_TOP_N):
    top_n = max(1, min(10, int(top_n or AI_PROPERTY_DEFAULT_TOP_N)))
    if not candidates:
        return []
    if not os.environ.get("GEMINI_API_KEY", "").strip():
        fallback = []
        for c in candidates[:top_n]:
            item = dict(c)
            item["ai_score"] = min(100, max(0, int(c.get("rule_score", 0))))
            item["fit_level"] = "可介紹"
            item["reason"] = "、".join(c.get("basic_reasons") or []) or "依條件粗篩後符合度較高。"
            item["risk"] = ""
            item["talking_point"] = "可先確認客戶是否接受此物件條件，再安排進一步介紹。"
            fallback.append(item)
        return fallback

    compact = []
    for idx, c in enumerate(candidates, start=1):
        compact.append({
            "idx": idx,
            "title": c.get("title"),
            "area": c.get("area"),
            "address": c.get("address"),
            "property_type": c.get("property_type"),
            "price_wan": c.get("price_wan"),
            "rent_price": c.get("rent_price"),
            "rooms": c.get("rooms"),
            "age": c.get("age"),
            "main_building_area": c.get("main_building_area"),
            "building_area": c.get("building_area"),
            "parking": c.get("parking"),
            "description": _ai_compact_text(c.get("description"), 260),
            "url": c.get("url"),
            "rule_score": c.get("rule_score"),
            "basic_reasons": c.get("basic_reasons"),
        })
    prompt = f"""
你是台中海線房仲 AI 物件推薦助手。
請根據客需，從候選物件挑出最適合的前 {top_n} 筆。
請只輸出 JSON，不要有其他文字。

客需條件：
{json.dumps(parsed_need, ensure_ascii=False, indent=2)}

候選物件：
{json.dumps(compact, ensure_ascii=False, indent=2)}

輸出格式：
{{
  "recommendations": [
    {{
      "idx": 候選物件idx,
      "ai_score": 0到100,
      "fit_level": "很適合 / 可介紹 / 勉強可看 / 不建議",
      "reason": "推薦原因，房仲口吻，具體說明符合哪些需求",
      "risk": "可能不符合或需要確認的地方",
      "talking_point": "可以怎麼跟客戶介紹這間"
    }}
  ]
}}
""".strip()
    try:
        data = _gemini_generate_json(prompt)
        recs = data.get("recommendations") or []
        by_idx = {i + 1: c for i, c in enumerate(candidates)}
        results = []
        for rec in recs:
            idx = _ai_safe_int(rec.get("idx"), None)
            base = by_idx.get(idx)
            if not base:
                continue
            item = dict(base)
            item["ai_score"] = _ai_safe_int(rec.get("ai_score"), base.get("rule_score") or 0)
            item["fit_level"] = rec.get("fit_level") or "可介紹"
            item["reason"] = rec.get("reason") or "依條件判斷符合度較高。"
            item["risk"] = rec.get("risk") or ""
            item["talking_point"] = rec.get("talking_point") or "可先提供給客戶參考。"
            results.append(item)
        results.sort(key=lambda x: x.get("ai_score") or 0, reverse=True)
        return results[:top_n] if results else _rank_properties_with_gemini_no_api(candidates, top_n)
    except Exception as e:
        print("⚠️ Gemini 排名失敗，改用規則排序：", e)
        return _rank_properties_with_gemini_no_api(candidates, top_n)


def _rank_properties_with_gemini_no_api(candidates, top_n=AI_PROPERTY_DEFAULT_TOP_N):
    out = []
    for c in (candidates or [])[:top_n]:
        item = dict(c)
        item["ai_score"] = min(100, max(0, int(c.get("rule_score", 0))))
        item["fit_level"] = "可介紹"
        item["reason"] = "、".join(c.get("basic_reasons") or []) or "依系統條件比對，符合度較高。"
        item["risk"] = "尚未由 Gemini 深度判斷，建議再人工確認屋況與細節。"
        item["talking_point"] = "可先用價格、區域、類型符合做切入介紹。"
        out.append(item)
    return out


def recommend_properties_for_buyer_data(buyer: dict, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = ""):
    parsed = _ai_try_parse_buyer_need_with_gemini(buyer, extra_text=extra_text)
    candidates = _hard_filter_properties_for_buyer(parsed, top_limit=AI_PROPERTY_MAX_CANDIDATES)
    ranked = _rank_properties_with_gemini(parsed, candidates, top_n=top_n)
    return {"parsed_need": parsed, "recommendations": ranked}


def _ai_property_price_text(item):
    if (item.get("deal_type") == "rent") or item.get("rent_price"):
        rent = _ai_safe_int(item.get("rent_price"), None)
        return f"租金 {rent:,} 元" if rent else "租金未填"
    price = item.get("price_wan")
    if price is None:
        return "總價未填"
    try:
        if float(price).is_integer():
            return f"{int(price):,} 萬"
    except Exception:
        pass
    return f"{price} 萬"


def _ai_valid_uri(url):
    url = str(url or "").strip()
    return url if url.startswith("http://") or url.startswith("https://") else ""


def build_property_recommend_bubble(item, idx=1):
    score = item.get("ai_score") or item.get("rule_score") or 0
    title = item.get("title") or "未命名物件"
    url = _ai_valid_uri(item.get("url"))
    info_contents = [
        flex_info_row("價格", _ai_property_price_text(item)),
        flex_info_row("區域", item.get("area") or "-"),
        flex_info_row("類型", item.get("property_type") or "-"),
        flex_info_row("格局", f"{_ai_safe_int(item.get('rooms'), '-') }房" if item.get("rooms") is not None else "-"),
        flex_info_row("屋齡", f"{item.get('age')}年" if item.get("age") is not None else "-"),
        flex_info_row("主建物", f"{item.get('main_building_area')}坪" if item.get("main_building_area") is not None else "-"),
    ]
    body = [
        {"type": "text", "text": f"推薦物件 #{idx}", "size": "xs", "color": "#C9874A", "weight": "bold"},
        {"type": "text", "text": line_truncate(title, 48), "size": "lg", "weight": "bold", "wrap": True, "color": "#222222"},
        {"type": "text", "text": f"{score}分｜{item.get('fit_level') or '可介紹'}", "size": "sm", "color": "#8b6b4f", "weight": "bold", "margin": "sm"},
        {"type": "separator", "margin": "md"},
        {"type": "box", "layout": "vertical", "spacing": "sm", "margin": "md", "contents": info_contents},
        {"type": "separator", "margin": "md"},
        {"type": "text", "text": "推薦原因", "size": "xs", "color": "#999999", "margin": "md"},
        {"type": "text", "text": _ai_compact_text(item.get("reason"), 150) or "符合條件，可先介紹。", "size": "sm", "color": "#333333", "wrap": True},
    ]
    if _ai_should_show_risk(item.get("risk")):
        body.extend([
            {"type": "text", "text": "注意", "size": "xs", "color": "#B00020", "margin": "md"},
            {"type": "text", "text": _ai_compact_text(item.get("risk"), 120), "size": "sm", "color": "#B00020", "wrap": True},
        ])

    footer = []
    if url:
        footer.append({
            "type": "button",
            "style": "primary",
            "height": "sm",
            "color": "#C9874A",
            "action": {"type": "uri", "label": "查看物件", "uri": url},
        })
    if item.get("talking_point"):
        footer.append({
            "type": "button",
            "style": "secondary",
            "height": "sm",
            "action": {
                "type": "postback",
                "label": "介紹話術",
                "data": f"action=property_talking_point&idx={idx}",
                "inputOption": "openKeyboard",
                "fillInText": _ai_compact_text(item.get("talking_point"), 300),
            },
        })
    return {
        "type": "bubble",
        "size": "mega",
        "body": {"type": "box", "layout": "vertical", "spacing": "md", "contents": body},
        "footer": {"type": "box", "layout": "vertical", "spacing": "sm", "contents": footer} if footer else {"type": "box", "layout": "vertical", "contents": []},
        "styles": {"footer": {"separator": True}},
    }


def build_property_recommend_flex(buyer: dict, recommendations: list, parsed_need: dict = None):
    buyer = buyer or {}
    recommendations = recommendations or []
    if not recommendations:
        return {
            "type": "bubble",
            "size": "mega",
            "body": {"type": "box", "layout": "vertical", "spacing": "md", "contents": [
                {"type": "text", "text": "AI推薦物件", "size": "xs", "color": "#C9874A", "weight": "bold"},
                {"type": "text", "text": buyer.get("name") or "客需", "size": "lg", "weight": "bold", "wrap": True},
                {"type": "separator", "margin": "md"},
                {"type": "text", "text": "目前沒有找到符合條件的物件。可以放寬預算、區域、產品類型或屋齡條件後再查詢。", "size": "sm", "wrap": True, "color": "#333333", "margin": "md"},
            ]},
        }
    bubbles = []
    summary = (parsed_need or {}).get("summary") or ""
    head_body = [
        {"type": "text", "text": "AI推薦物件", "size": "xs", "color": "#C9874A", "weight": "bold"},
        {"type": "text", "text": buyer.get("name") or "客需", "size": "lg", "weight": "bold", "wrap": True},
        {"type": "separator", "margin": "md"},
        flex_info_row("電話", buyer.get("phone") or "-"),
        flex_info_row("需求", _ai_compact_text(summary or buyer.get("note") or "-", 120)),
        flex_info_row("筆數", f"找到 {len(recommendations)} 筆推薦"),
    ]
    bubbles.append({
        "type": "bubble",
        "size": "mega",
        "body": {"type": "box", "layout": "vertical", "spacing": "md", "contents": head_body},
    })
    for i, item in enumerate(recommendations[:10], start=1):
        bubbles.append(build_property_recommend_bubble(item, i))
    return {"type": "carousel", "contents": bubbles[:11]}


def _parse_recommend_command_fields(text: str):
    lines = [ln.strip() for ln in str(text or "").splitlines() if ln.strip()]
    fields = {}
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line)
        if not m:
            continue
        key = normalize_line_key(m.group(1))
        value = m.group(2).strip()
        fields[key] = value
    return fields


def _find_buyer_doc_for_recommend(fields: dict):
    record_id = (fields.get("record_id") or fields.get("buyer_id") or fields.get("id") or "").strip()
    phone = (fields.get("phone") or "").strip()
    name = (fields.get("name") or "").strip()
    if record_id:
        doc = find_customer_record("buyer", record_id=record_id)
        if doc:
            return doc, ""
        return None, "找不到這筆客需，請確認客戶ID。"
    if phone or name:
        doc = find_customer_record("buyer", phone=phone, name=name)
        if doc:
            return doc, ""
        return None, "找不到唯一客需，請改用電話或客戶ID。"
    return None, "請輸入電話或客戶ID。\n\n格式：\n#推薦物件\n電話: 0928xxxxxx\n筆數: 5"


def _line_event_can_view_buyer(event, buyer_data: dict):
    try:
        kind, target_id = line_event_source_kind_and_id(event)
        cfg = None
        if kind == "user":
            cfg = find_line_personal_user_by_user_id(target_id)
        elif kind in ("group", "room"):
            cfg = find_line_group_by_target_id(target_id)
        if cfg and "permission_config_can_view" in globals():
            return permission_config_can_view(cfg, "buyer", buyer_data, kind, target_id)
    except Exception:
        pass
    # 沒有權限設定時，保守處理：群組只看公開；個人只能看公開或自己的。
    visibility = (buyer_data or {}).get("visibility") or "public"
    if visibility != "personal":
        return True
    try:
        kind, target_id = line_event_source_kind_and_id(event)
        return bool(kind == "user" and target_id and target_id == (buyer_data or {}).get("owner_line_user_id"))
    except Exception:
        return False


def _make_recommendation_result_for_buyer_doc(buyer_doc, event=None, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text=""):
    if not buyer_doc or not getattr(buyer_doc, "exists", False):
        return {"handled": True, "ok": False, "reply_text": "找不到客需資料。"}
    buyer = buyer_doc.to_dict() or {}
    buyer["id"] = buyer_doc.id
    if event is not None and not _line_event_can_view_buyer(event, buyer):
        return {"handled": True, "ok": False, "reply_text": "你沒有權限查詢這筆個人客需的推薦物件。"}
    try:
        result = recommend_properties_for_buyer_data(buyer, top_n=top_n, extra_text=extra_text)
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"{buyer.get('name','客需')} 推薦物件",
            "reply_flex": flex,
            "target_type": "buyer",
            "target_id": buyer_doc.id,
            "customer_name": buyer.get("name", ""),
            "phone": buyer.get("phone", ""),
            "parsed_tag": "推薦物件",
        }
    except Exception as e:
        return {"handled": True, "ok": False, "reply_text": f"AI推薦物件失敗：{e}"}


def process_line_recommend_properties_event(event):
    message = event.get("message") or {}
    text = (message.get("text") or "").strip()
    if not text:
        return {"handled": False}
    first = text.splitlines()[0].strip().replace(" ", "")
    if first not in ("#推薦物件", "#AI推薦物件", "#物件推薦"):
        return {"handled": False}
    fields = _parse_recommend_command_fields(text)
    top_n = parse_int_limit(fields.get("limit") or fields.get("筆數") or 5, default=5, max_value=10)
    extra_text = fields.get("補充") or fields.get("content") or fields.get("note") or ""
    buyer_doc, err = _find_buyer_doc_for_recommend(fields)
    if err:
        return {"handled": True, "ok": False, "reply_text": err}
    return _make_recommendation_result_for_buyer_doc(buyer_doc, event=event, top_n=top_n, extra_text=extra_text)


def _ai_recommend_properties_postback_action(buyer_id: str, label="推薦物件"):
    return {
        "type": "postback",
        "label": label,
        "data": f"action=ai_recommend_properties&buyer_id={buyer_id}",
    }


# 客需 LINE 卡片下方新增「推薦物件」按鈕。
try:
    _build_record_flex_bubble_before_ai_recommend = _build_record_flex_bubble

    def _build_record_flex_bubble(record_type: str, record_id: str, data: dict, title_prefix="CRM 資料"):
        bubble = _build_record_flex_bubble_before_ai_recommend(record_type, record_id, data, title_prefix=title_prefix)
        if record_type != "buyer" or not record_id:
            return bubble
        try:
            footer = bubble.setdefault("footer", {"type": "box", "layout": "vertical", "spacing": "sm", "contents": []})
            contents = footer.setdefault("contents", [])
            # 避免重複插入。
            for btn in contents:
                act = (btn or {}).get("action") or {}
                if "ai_recommend_properties" in str(act.get("data", "")):
                    return bubble
            contents.append({
                "type": "button",
                "style": "secondary",
                "height": "sm",
                "action": _ai_recommend_properties_postback_action(record_id, label="推薦物件"),
            })
        except Exception as e:
            print("⚠️ 客需卡片加入推薦物件按鈕失敗：", e)
        return bubble

    print("✅ 客需 LINE 卡片已加入『推薦物件』按鈕")
except Exception as e:
    print("⚠️ 客需推薦物件按鈕 patch 失敗：", e)


# LINE 權限 gate：讓 #推薦物件 被視為 buyer 類指令；postback 也走 buyer 權限。
try:
    _detect_line_command_type_before_ai_recommend = detect_line_command_type

    def detect_line_command_type(text: str, event=None) -> str:
        first = (text or "").strip().splitlines()[0].strip().replace(" ", "") if (text or "").strip() else ""
        first_no_hash = first[1:] if first.startswith("#") else first
        if first_no_hash in ("推薦物件", "AI推薦物件", "物件推薦"):
            return "buyer"
        return _detect_line_command_type_before_ai_recommend(text, event=event)

    print("✅ LINE 指令已加入 #推薦物件")
except Exception as e:
    print("⚠️ #推薦物件 指令 patch 失敗：", e)

try:
    _detect_line_postback_command_type_before_ai_recommend = detect_line_postback_command_type

    def detect_line_postback_command_type(event) -> str:
        try:
            data = ((event or {}).get("postback") or {}).get("data", "") or ""
            if "action=ai_recommend_properties" in data:
                return "buyer"
        except Exception:
            pass
        return _detect_line_postback_command_type_before_ai_recommend(event)

    print("✅ LINE postback 權限已加入推薦物件")
except Exception as e:
    print("⚠️ 推薦物件 postback 權限 patch 失敗：", e)


# LINE 文字訊息：攔截 #推薦物件；新增客需成功時也改回卡片，卡片下方有推薦物件按鈕。
try:
    _process_line_message_event_before_ai_recommend = process_line_message_event

    def process_line_message_event(event):
        msg = event.get("message") or {}
        if msg.get("type") == "text":
            rec_result = process_line_recommend_properties_event(event)
            if rec_result.get("handled"):
                return rec_result

        result = _process_line_message_event_before_ai_recommend(event)
        try:
            # 只把新增客需/委託/開發這類成功結果改成卡片；查詢紀錄維持文字。
            if (
                result and result.get("ok") and result.get("target_type") in ("buyer", "seller", "development")
                and not result.get("reply_flex") and not result.get("reply_messages")
                and result.get("parsed_tag") in ("新增客需", "新增委託", "新增開發")
            ):
                coll = {"buyer": "buyers", "seller": "sellers", "development": "developments"}.get(result.get("target_type"))
                snap = db.collection(coll).document(result.get("target_id")).get()
                if snap.exists:
                    data = snap.to_dict() or {}
                    result["reply_flex"] = _build_record_flex_bubble(result.get("target_type"), result.get("target_id"), data, title_prefix=result.get("parsed_tag") or "CRM 資料")
                    result["reply_text"] = f"{result.get('parsed_tag') or 'CRM'}：{data.get('name','')}"
        except Exception as e:
            print("⚠️ 新增資料回覆改卡片失敗：", e)
        return result

    print("✅ process_line_message_event 已支援 #推薦物件 與新增客需卡片回覆")
except Exception as e:
    print("⚠️ process_line_message_event 推薦物件 patch 失敗：", e)


# LINE postback：點客需卡片的「推薦物件」時，直接回覆推薦物件 carousel。
try:
    _process_line_postback_event_before_ai_recommend = process_line_postback_event

    def process_line_postback_event(event):
        try:
            data = ((event or {}).get("postback") or {}).get("data", "") or ""
            if "action=ai_recommend_properties" in data:
                params = {}
                for part in data.split("&"):
                    if "=" in part:
                        k, v = part.split("=", 1)
                        params[k] = v
                buyer_id = (params.get("buyer_id") or params.get("record_id") or "").strip()
                if not buyer_id:
                    return {"handled": True, "ok": False, "reply_text": "缺少客需ID，無法推薦物件。"}
                buyer_doc = db.collection("buyers").document(buyer_id).get()
                return _make_recommendation_result_for_buyer_doc(buyer_doc, event=event, top_n=AI_PROPERTY_DEFAULT_TOP_N)
        except Exception as e:
            return {"handled": True, "ok": False, "reply_text": f"推薦物件失敗：{e}"}
        return _process_line_postback_event_before_ai_recommend(event)

    print("✅ LINE postback 已支援點擊『推薦物件』")
except Exception as e:
    print("⚠️ 推薦物件 postback 處理 patch 失敗：", e)


# 後台客需詳細頁也可直接看 AI 推薦；沒有模板也可運作。
@app.route("/buyers/<buyer_id>/ai-recommend", strict_slashes=False)
@login_required
def buyer_ai_recommend_page(buyer_id):
    snap = db.collection("buyers").document(buyer_id).get()
    if not snap.exists:
        flash("找不到這筆客需", "danger")
        return redirect(url_for("buyers"))
    buyer = snap.to_dict() or {}
    buyer["id"] = snap.id
    try:
        if "backend_can_view_personal_record" in globals() and not backend_can_view_personal_record(buyer):
            flash("你沒有權限查看這筆個人客需", "danger")
            return redirect(url_for("buyers"))
    except Exception:
        pass
    try:
        result = recommend_properties_for_buyer_data(buyer, top_n=10)
    except Exception as e:
        flash(f"AI推薦失敗：{e}", "danger")
        return redirect(url_for("buyer_detail", buyer_id=buyer_id))
    html = """
    {% extends "base.html" %}
    {% block content %}
    <div class="d-flex justify-content-between align-items-center mb-3">
      <div>
        <h3 class="mb-1">AI推薦物件</h3>
        <div class="text-muted small">客戶：{{ buyer.name or '-' }}｜電話：{{ buyer.phone or '-' }}</div>
      </div>
      <a class="btn btn-secondary" href="{{ url_for('buyer_detail', buyer_id=buyer.id) }}">回客需詳細</a>
    </div>
    <div class="card mb-3"><div class="card-header fw-bold">AI解析客需</div><div class="card-body"><pre class="mb-0" style="white-space:pre-wrap;">{{ parsed_need | tojson(indent=2, ensure_ascii=False) }}</pre></div></div>
    <div class="card"><div class="card-header fw-bold">推薦物件</div><div class="card-body">
      {% if recommendations %}
        {% for item in recommendations %}
          <div class="border rounded p-3 mb-3 bg-white">
            <div class="d-flex justify-content-between align-items-start gap-3">
              <div><h5 class="mb-1">{{ loop.index }}. {{ item.title }}</h5><div class="text-muted small">{{ item.area or '-' }}｜{{ item.property_type or '-' }}｜{% if item.price_wan %}{{ item.price_wan }}萬{% elif item.rent_price %}租金 {{ item.rent_price }}元{% else %}價格未填{% endif %}{% if item.age %}｜屋齡 {{ item.age }}年{% endif %}</div></div>
              <span class="badge bg-success fs-6">{{ item.ai_score or item.rule_score }} 分</span>
            </div>
            <div class="mt-2"><strong>適合度：</strong>{{ item.fit_level or '-' }}</div>
            <div class="mt-2"><strong>推薦原因：</strong>{{ item.reason or '-' }}</div>
            {% if _ai_should_show_risk(item.risk) %}<div class="mt-2 text-danger"><strong>注意：</strong>{{ item.risk }}</div>{% endif %}
            <div class="mt-2"><strong>話術：</strong>{{ item.talking_point or '-' }}</div>
            {% if item.url %}<a class="btn btn-sm btn-outline-primary mt-2" target="_blank" href="{{ item.url }}">查看物件</a>{% endif %}
          </div>
        {% endfor %}
      {% else %}
        <div class="text-muted">目前沒有找到符合條件的物件。</div>
      {% endif %}
    </div></div>
    {% endblock %}
    """
    return render_template_string(html, buyer=buyer, parsed_need=result.get("parsed_need") or {}, recommendations=result.get("recommendations") or [])

print("✅ Gemini AI 物件推薦 Patch 已載入：#推薦物件 / 客需卡片推薦按鈕 / 後台AI推薦頁")
# =============================================================================
# Gemini AI 物件推薦 + LINE #推薦物件 指令 Patch End
# =============================================================================


# =============================================================================
# AI 推薦物件效能優化 Patch v20260624_FAST
# - Firestore properties 加入短時間記憶體快取，避免每次推薦都重新讀 1000+ 筆
# - Gemini 推薦結果寫入 ai_recommendation_cache
# - 後台 AI推薦物件頁改成：有快取立即顯示；無快取先建立背景任務，不阻塞後台
# - LINE #推薦物件 / 卡片推薦按鈕：有快取立即回卡片；無快取先回覆「AI推薦中」，完成後再推送卡片
# =============================================================================

AI_PROPERTY_CACHE_TTL_SECONDS = int(os.environ.get("AI_PROPERTY_CACHE_TTL_SECONDS", "600") or 600)
AI_RECOMMEND_BACKGROUND_ENABLED = os.environ.get("AI_RECOMMEND_BACKGROUND_ENABLED", "1").strip() not in ("0", "false", "False", "no", "NO")
AI_RECOMMEND_CACHE_COLLECTION = os.environ.get("AI_RECOMMEND_CACHE_COLLECTION", "ai_recommendation_cache")

try:
    import threading as _ai_threading
    import traceback as _ai_traceback
except Exception:
    _ai_threading = None
    _ai_traceback = None

_AI_PROPERTY_MEMORY_CACHE = globals().get("_AI_PROPERTY_MEMORY_CACHE") or {}
_AI_RUNNING_RECOMMEND_JOBS = globals().get("_AI_RUNNING_RECOMMEND_JOBS") or set()

try:
    _fetch_active_properties_before_fast_patch = _fetch_active_properties
except Exception:
    _fetch_active_properties_before_fast_patch = None


def _ai_now_ts():
    try:
        return time.time()
    except Exception:
        return 0


def _ai_property_import_token(deal_type: str):
    """取得目前 sale/rent 物件資料版本。用 import_logs 判斷，找不到就用 properties 數量備援。"""
    deal_type = "rent" if deal_type == "rent" else "sale"
    try:
        logs = []
        for d in db.collection("property_import_logs").where("deal_type", "==", deal_type).stream():
            item = d.to_dict() or {}
            item["id"] = d.id
            logs.append(item)
        if logs:
            logs.sort(key=lambda x: x.get("created_at") or x.get("uploaded_at") or x.get("imported_at") or "", reverse=True)
            latest = logs[0]
            return "|".join([
                deal_type,
                str(latest.get("id") or ""),
                str(latest.get("created_at") or latest.get("uploaded_at") or latest.get("imported_at") or ""),
                str(latest.get("inserted_count") or latest.get("total_count") or ""),
                str(latest.get("file_name") or ""),
            ])
    except Exception as e:
        print("⚠️ 讀取 property_import_logs 失敗，改用 properties 備援：", e)

    try:
        # 備援：最多只計算到 3000 筆，避免過慢。
        count = 0
        latest_imported_at = ""
        for d in db.collection("properties").where("deal_type", "==", deal_type).limit(3000).stream():
            data = d.to_dict() or {}
            if data.get("active") is False:
                continue
            count += 1
            ts = data.get("imported_at") or data.get("created_at") or ""
            if ts > latest_imported_at:
                latest_imported_at = ts
        return f"{deal_type}|fallback|{count}|{latest_imported_at}"
    except Exception:
        return f"{deal_type}|unknown"


def clear_ai_property_memory_cache(deal_type: str = ""):
    """物件 CSV 匯入後可呼叫這個，讓下一次推薦重新讀 Firestore。"""
    try:
        if not deal_type:
            _AI_PROPERTY_MEMORY_CACHE.clear()
        else:
            deal_type = "rent" if deal_type == "rent" else "sale"
            for k in list(_AI_PROPERTY_MEMORY_CACHE.keys()):
                if str(k).startswith(deal_type + "|"):
                    _AI_PROPERTY_MEMORY_CACHE.pop(k, None)
        print("✅ AI 物件記憶體快取已清除", deal_type or "all")
    except Exception as e:
        print("⚠️ 清除 AI 物件快取失敗：", e)


def _fetch_active_properties(deal_type="sale", max_docs=2500):
    """覆寫原本讀物件資料的方法：加入 TTL 快取 + 物件版本 token。"""
    deal_type = "rent" if deal_type == "rent" else "sale"
    token = _ai_property_import_token(deal_type)
    key = f"{deal_type}|{max_docs}|{token}"
    now_ts = _ai_now_ts()
    cached = _AI_PROPERTY_MEMORY_CACHE.get(key)
    if cached and (now_ts - cached.get("ts", 0) <= AI_PROPERTY_CACHE_TTL_SECONDS):
        return deepcopy(cached.get("items") or [])

    # 清掉同類型舊版本快取，避免記憶體一直累積。
    for old_key in list(_AI_PROPERTY_MEMORY_CACHE.keys()):
        if str(old_key).startswith(deal_type + "|") and old_key != key:
            _AI_PROPERTY_MEMORY_CACHE.pop(old_key, None)

    if _fetch_active_properties_before_fast_patch:
        items = _fetch_active_properties_before_fast_patch(deal_type=deal_type, max_docs=max_docs)
    else:
        items = []
        docs = db.collection("properties").where("deal_type", "==", deal_type).limit(max_docs).stream()
        for d in docs:
            data = d.to_dict() or {}
            data["id"] = d.id
            if data.get("active") is False:
                continue
            items.append(_property_normalized(data))

    _AI_PROPERTY_MEMORY_CACHE[key] = {"ts": now_ts, "items": deepcopy(items), "token": token}
    return items


def _ai_buyer_cache_signature(buyer: dict, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = ""):
    buyer = buyer or {}
    intent = (buyer.get("intent_type") or "buy").strip()
    deal_type = "rent" if intent in ("rent", "租", "租屋", "承租") else "sale"
    prop_token = _ai_property_import_token(deal_type)
    important = {
        "buyer_id": buyer.get("id") or "",
        "updated_at": buyer.get("updated_at") or buyer.get("created_at") or "",
        "name": buyer.get("name") or "",
        "phone": normalize_phone(buyer.get("phone") or ""),
        "intent_type": buyer.get("intent_type") or "",
        "budget_max": buyer.get("budget_max") or "",
        "rent_max": buyer.get("rent_max") or "",
        "preferred_areas": buyer.get("preferred_areas") or "",
        "property_type": buyer.get("property_type") or "",
        "room_range": buyer.get("room_range") or "",
        "car_need": buyer.get("car_need") or "",
        "requirement_must": buyer.get("requirement_must") or "",
        "requirement_nice": buyer.get("requirement_nice") or "",
        "note": buyer.get("note") or "",
        "extra_text": extra_text or "",
        "top_n": int(top_n or AI_PROPERTY_DEFAULT_TOP_N),
        "gemini_model": os.environ.get("GEMINI_MODEL", GEMINI_DEFAULT_MODEL),
        "property_token": prop_token,
        "deal_type": deal_type,
    }
    raw = json.dumps(important, ensure_ascii=False, sort_keys=True)
    doc_id = hashlib.sha256(raw.encode("utf-8")).hexdigest()
    return doc_id, important


def _get_ai_recommend_cache(buyer: dict, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = ""):
    doc_id, sig = _ai_buyer_cache_signature(buyer, top_n=top_n, extra_text=extra_text)
    try:
        doc = db.collection(AI_RECOMMEND_CACHE_COLLECTION).document(doc_id).get()
        if doc.exists:
            data = doc.to_dict() or {}
            if data.get("signature") == sig and data.get("status") == "success":
                return data
    except Exception as e:
        print("⚠️ 讀取 AI 推薦快取失敗：", e)
    return None


def _save_ai_recommend_cache(buyer: dict, result: dict, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = ""):
    doc_id, sig = _ai_buyer_cache_signature(buyer, top_n=top_n, extra_text=extra_text)
    try:
        db.collection(AI_RECOMMEND_CACHE_COLLECTION).document(doc_id).set({
            "status": "success",
            "buyer_id": buyer.get("id") or "",
            "buyer_name": buyer.get("name") or "",
            "buyer_phone": buyer.get("phone") or "",
            "signature": sig,
            "parsed_need": result.get("parsed_need") or {},
            "recommendations": result.get("recommendations") or [],
            "created_at": now_taipei().isoformat(),
            "updated_at": now_taipei().isoformat(),
        }, merge=True)
    except Exception as e:
        print("⚠️ 寫入 AI 推薦快取失敗：", e)


def _mark_ai_recommend_job_status(buyer: dict, status: str, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = "", error: str = ""):
    doc_id, sig = _ai_buyer_cache_signature(buyer, top_n=top_n, extra_text=extra_text)
    try:
        db.collection(AI_RECOMMEND_CACHE_COLLECTION).document(doc_id).set({
            "status": status,
            "buyer_id": buyer.get("id") or "",
            "buyer_name": buyer.get("name") or "",
            "buyer_phone": buyer.get("phone") or "",
            "signature": sig,
            "error": error,
            "updated_at": now_taipei().isoformat(),
        }, merge=True)
    except Exception as e:
        print("⚠️ 更新 AI 推薦任務狀態失敗：", e)


try:
    _recommend_properties_for_buyer_data_before_fast_patch = recommend_properties_for_buyer_data
except Exception:
    _recommend_properties_for_buyer_data_before_fast_patch = None


def recommend_properties_for_buyer_data(buyer: dict, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = "", force_refresh: bool = False):
    """覆寫推薦主流程：先看 Firestore 快取；需要重算才跑 Gemini。"""
    buyer = buyer or {}
    top_n = max(1, min(10, int(top_n or AI_PROPERTY_DEFAULT_TOP_N)))
    if not force_refresh:
        cached = _get_ai_recommend_cache(buyer, top_n=top_n, extra_text=extra_text)
        if cached:
            return {
                "parsed_need": cached.get("parsed_need") or {},
                "recommendations": cached.get("recommendations") or [],
                "cache_hit": True,
                "cached_at": cached.get("created_at") or "",
            }

    parsed = _ai_try_parse_buyer_need_with_gemini(buyer, extra_text=extra_text)
    candidates = _hard_filter_properties_for_buyer(parsed, top_limit=AI_PROPERTY_MAX_CANDIDATES)
    ranked = _rank_properties_with_gemini(parsed, candidates, top_n=top_n)
    result = {"parsed_need": parsed, "recommendations": ranked, "cache_hit": False}
    _save_ai_recommend_cache(buyer, result, top_n=top_n, extra_text=extra_text)
    return result


def _line_event_target_id(event):
    source = (event or {}).get("source") or {}
    return source.get("groupId") or source.get("roomId") or source.get("userId") or ""


def _push_ai_recommendation_to_line(target_id: str, buyer: dict, result: dict):
    if not target_id:
        return False, "missing target_id"
    try:
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        messages = [{"type": "flex", "altText": f"推薦物件：{buyer.get('name') or ''}", "contents": flex}]
        if "line_push_messages" in globals():
            return line_push_messages(target_id, messages)
        if "push_line_flex" in globals():
            return push_line_flex(target_id, f"推薦物件：{buyer.get('name') or ''}", flex)
        if "push_line_text" in globals():
            return push_line_text(target_id, f"推薦物件完成：{buyer.get('name') or ''}")
    except Exception as e:
        print("⚠️ 推送 AI 推薦卡片失敗：", e)
        return False, str(e)
    return False, "no push function"


def _start_ai_recommend_background_job(buyer: dict, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text: str = "", target_id: str = ""):
    """建立背景推薦任務。完成後如果有 target_id，會把卡片推回 LINE。"""
    if not AI_RECOMMEND_BACKGROUND_ENABLED or not _ai_threading:
        return False
    buyer = dict(buyer or {})
    top_n = max(1, min(10, int(top_n or AI_PROPERTY_DEFAULT_TOP_N)))
    doc_id, _sig = _ai_buyer_cache_signature(buyer, top_n=top_n, extra_text=extra_text)
    if doc_id in _AI_RUNNING_RECOMMEND_JOBS:
        return True

    def _runner():
        _AI_RUNNING_RECOMMEND_JOBS.add(doc_id)
        try:
            _mark_ai_recommend_job_status(buyer, "running", top_n=top_n, extra_text=extra_text)
            result = recommend_properties_for_buyer_data(buyer, top_n=top_n, extra_text=extra_text, force_refresh=True)
            if target_id:
                _push_ai_recommendation_to_line(target_id, buyer, result)
        except Exception as e:
            print("❌ AI推薦背景任務失敗：", e)
            if _ai_traceback:
                print(_ai_traceback.format_exc())
            _mark_ai_recommend_job_status(buyer, "failed", top_n=top_n, extra_text=extra_text, error=str(e))
            if target_id and "push_line_text" in globals():
                try:
                    push_line_text(target_id, f"AI推薦物件失敗：{e}")
                except Exception:
                    pass
        finally:
            _AI_RUNNING_RECOMMEND_JOBS.discard(doc_id)

    t = _ai_threading.Thread(target=_runner, daemon=True)
    t.start()
    return True


try:
    _make_recommendation_result_for_buyer_doc_before_fast_patch = _make_recommendation_result_for_buyer_doc
except Exception:
    _make_recommendation_result_for_buyer_doc_before_fast_patch = None


def _make_recommendation_result_for_buyer_doc(buyer_doc, event=None, top_n=AI_PROPERTY_DEFAULT_TOP_N, extra_text=""):
    if not buyer_doc or not getattr(buyer_doc, "exists", False):
        return {"handled": True, "ok": False, "reply_text": "找不到這筆客需，無法推薦物件。"}
    buyer = buyer_doc.to_dict() or {}
    buyer["id"] = buyer_doc.id
    top_n = max(1, min(10, int(top_n or AI_PROPERTY_DEFAULT_TOP_N)))

    cached = _get_ai_recommend_cache(buyer, top_n=top_n, extra_text=extra_text)
    if cached:
        result = {"parsed_need": cached.get("parsed_need") or {}, "recommendations": cached.get("recommendations") or [], "cache_hit": True}
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"推薦物件：{buyer.get('name','')}",
            "reply_flex": flex,
            "target_type": "buyer",
            "target_id": buyer_doc.id,
            "customer_name": buyer.get("name", ""),
            "phone": buyer.get("phone", ""),
            "parsed_tag": "推薦物件",
        }

    target_id = _line_event_target_id(event) if event else ""
    if event and AI_RECOMMEND_BACKGROUND_ENABLED:
        started = _start_ai_recommend_background_job(buyer, top_n=top_n, extra_text=extra_text, target_id=target_id)
        if started:
            return {
                "handled": True,
                "ok": True,
                "reply_text": f"已收到，正在幫 {buyer.get('name','這位客戶')} AI推薦物件。完成後會自動把物件卡片傳回來。",
                "target_type": "buyer",
                "target_id": buyer_doc.id,
                "customer_name": buyer.get("name", ""),
                "phone": buyer.get("phone", ""),
                "parsed_tag": "推薦物件",
            }

    # 如果背景任務不可用，就保留同步模式。
    try:
        result = recommend_properties_for_buyer_data(buyer, top_n=top_n, extra_text=extra_text, force_refresh=False)
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {
            "handled": True,
            "ok": True,
            "reply_text": f"推薦物件：{buyer.get('name','')}",
            "reply_flex": flex,
            "target_type": "buyer",
            "target_id": buyer_doc.id,
            "customer_name": buyer.get("name", ""),
            "phone": buyer.get("phone", ""),
            "parsed_tag": "推薦物件",
        }
    except Exception as e:
        return {"handled": True, "ok": False, "reply_text": f"AI推薦物件失敗：{e}"}


def _render_ai_recommend_html(buyer, parsed_need=None, recommendations=None, running=False, cache_hit=False, error=""):
    html = """
    {% extends "base.html" %}
    {% block content %}
    <div class="d-flex justify-content-between align-items-center mb-3">
      <div>
        <h3 class="mb-1">AI推薦物件</h3>
        <div class="text-muted small">客戶：{{ buyer.name or '-' }}｜電話：{{ buyer.phone or '-' }}</div>
      </div>
      <div class="d-flex gap-2">
        <a class="btn btn-outline-primary" href="{{ url_for('buyer_ai_recommend_page', buyer_id=buyer.id) }}?refresh=1">重新AI推薦</a>
        <a class="btn btn-secondary" href="{{ url_for('buyer_detail', buyer_id=buyer.id) }}">回客需詳細</a>
      </div>
    </div>

    {% if running %}
      <meta http-equiv="refresh" content="5">
      <div class="card border-warning mb-3">
        <div class="card-body">
          <h5 class="mb-2">AI推薦中...</h5>
          <div class="text-muted">系統已在背景處理，這個頁面會每 5 秒自動更新。你也可以先回後台做其他事情。</div>
        </div>
      </div>
    {% endif %}

    {% if error %}
      <div class="alert alert-danger">{{ error }}</div>
    {% endif %}

    {% if cache_hit %}
      <div class="alert alert-success small">已使用上次推薦快取，所以開啟速度會比較快。若客需或物件資料有變動，可按「重新AI推薦」。</div>
    {% endif %}

    {% if parsed_need %}
      <div class="card mb-3"><div class="card-header fw-bold">AI解析客需</div><div class="card-body"><pre class="mb-0" style="white-space:pre-wrap;">{{ parsed_need | tojson(indent=2, ensure_ascii=False) }}</pre></div></div>
    {% endif %}

    <div class="card"><div class="card-header fw-bold">推薦物件</div><div class="card-body">
      {% if recommendations %}
        {% for item in recommendations %}
          <div class="border rounded p-3 mb-3 bg-white">
            <div class="d-flex justify-content-between align-items-start gap-3">
              <div><h5 class="mb-1">{{ loop.index }}. {{ item.title }}</h5><div class="text-muted small">{{ item.area or '-' }}｜{{ item.property_type or '-' }}｜{% if item.price_wan %}{{ item.price_wan }}萬{% elif item.rent_price %}租金 {{ item.rent_price }}元{% else %}價格未填{% endif %}{% if item.age %}｜屋齡 {{ item.age }}年{% endif %}</div></div>
              <span class="badge bg-success fs-6">{{ item.ai_score or item.rule_score }} 分</span>
            </div>
            <div class="mt-2"><strong>適合度：</strong>{{ item.fit_level or '-' }}</div>
            <div class="mt-2"><strong>推薦原因：</strong>{{ item.reason or '-' }}</div>
            {% if _ai_should_show_risk(item.risk) %}<div class="mt-2 text-danger"><strong>注意：</strong>{{ item.risk }}</div>{% endif %}
            <div class="mt-2"><strong>話術：</strong>{{ item.talking_point or '-' }}</div>
            {% if item.url %}<a class="btn btn-sm btn-outline-primary mt-2" target="_blank" href="{{ item.url }}">查看物件</a>{% endif %}
          </div>
        {% endfor %}
      {% elif not running %}
        <div class="text-muted">目前沒有找到符合條件的物件。</div>
      {% endif %}
    </div></div>
    {% endblock %}
    """
    return render_template_string(html, buyer=buyer, parsed_need=parsed_need or {}, recommendations=recommendations or [], running=running, cache_hit=cache_hit, error=error)


def _buyer_ai_recommend_page_fast(buyer_id):
    snap = db.collection("buyers").document(buyer_id).get()
    if not snap.exists:
        flash("找不到這筆客需", "danger")
        return redirect(url_for("buyers"))
    buyer = snap.to_dict() or {}
    buyer["id"] = snap.id
    try:
        if "backend_can_view_personal_record" in globals() and not backend_can_view_personal_record(buyer):
            flash("你沒有權限查看這筆個人客需", "danger")
            return redirect(url_for("buyers"))
    except Exception:
        pass

    force_refresh = (request.args.get("refresh") or "").strip() in ("1", "true", "yes")
    top_n = parse_int_limit(request.args.get("top_n") or 10, default=10, max_value=10)

    if not force_refresh:
        cached = _get_ai_recommend_cache(buyer, top_n=top_n, extra_text="")
        if cached:
            return _render_ai_recommend_html(
                buyer,
                parsed_need=cached.get("parsed_need") or {},
                recommendations=cached.get("recommendations") or [],
                cache_hit=True,
            )

    # 背景處理，不阻塞後台。
    if AI_RECOMMEND_BACKGROUND_ENABLED:
        _start_ai_recommend_background_job(buyer, top_n=top_n, extra_text="", target_id="")
        return _render_ai_recommend_html(buyer, running=True)

    # 背景關閉時才同步跑。
    try:
        result = recommend_properties_for_buyer_data(buyer, top_n=top_n, force_refresh=force_refresh)
        return _render_ai_recommend_html(buyer, parsed_need=result.get("parsed_need") or {}, recommendations=result.get("recommendations") or [], cache_hit=bool(result.get("cache_hit")))
    except Exception as e:
        return _render_ai_recommend_html(buyer, error=f"AI推薦失敗：{e}")

try:
    app.view_functions["buyer_ai_recommend_page"] = _buyer_ai_recommend_page_fast
    print("✅ 後台 AI推薦物件頁已改成快取 + 背景任務，不再長時間卡住後台")
except Exception as e:
    print("⚠️ 覆寫後台 AI推薦頁失敗：", e)

print("✅ AI 推薦效能優化已啟用：properties快取 / 推薦快取 / 背景任務")
# =============================================================================
# AI 推薦物件效能優化 Patch End
# =============================================================================


# =============================================================================
# AI推薦 V3：風險顯示邏輯 / 條件找物件 / 設定中心開關 Patch
# =============================================================================

def ai_recommend_feature_enabled():
    """後台設定中心可控制 AI推薦是否啟用；預設啟用。"""
    env_val = os.environ.get("AI_RECOMMEND_ENABLED", "").strip().lower()
    if env_val in ("0", "false", "no", "off"):
        return False
    try:
        settings = get_line_card_settings() if "get_line_card_settings" in globals() else {}
        if "ai_recommend_enabled" in settings:
            return bool(settings.get("ai_recommend_enabled"))
    except Exception:
        pass
    return True

try:
    _save_line_settings_center_from_form_before_ai_toggle = save_line_settings_center_from_form

    def save_line_settings_center_from_form(form):
        result = _save_line_settings_center_from_form_before_ai_toggle(form)
        extra = {
            "ai_recommend_enabled": form.get("ai_recommend_enabled") == "on",
            "ai_free_search_enabled": form.get("ai_free_search_enabled") == "on",
            "updated_at": now_taipei().isoformat(),
            "updated_by_id": session.get("user_id", ""),
            "updated_by_name": session.get("user_name", ""),
        }
        db.collection(LINE_CARD_SETTINGS_COLLECTION).document("default").set(extra, merge=True)
        result.update(extra)
        return result

    _ai_sidebar_link = '<a class="list-group-item list-group-item-action" href="#ai-recommend">AI推薦設定</a>'
    if _ai_sidebar_link not in LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX:
        LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX = LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX.replace(
            '<a class="list-group-item list-group-item-action" href="#rules">權限規則</a>',
            _ai_sidebar_link + '\n        <a class="list-group-item list-group-item-action" href="#rules">權限規則</a>'
        )

    _ai_settings_card = """
        <div id="ai-recommend" class="setting-card p-4 mb-4">
          <h5>AI推薦設定</h5>
          <div class="hint mb-3">控制 Gemini AI推薦物件與條件找物件功能是否開啟。關閉後，LINE 指令與後台 AI推薦都會停止執行，避免誤用額度。</div>
          <label class="form-check mb-2">
            <input class="form-check-input" type="checkbox" name="ai_recommend_enabled" {% if settings.ai_recommend_enabled is not defined or settings.ai_recommend_enabled %}checked{% endif %}>
            啟用 AI推薦物件功能
          </label>
          <label class="form-check mb-2">
            <input class="form-check-input" type="checkbox" name="ai_free_search_enabled" {% if settings.ai_free_search_enabled is not defined or settings.ai_free_search_enabled %}checked{% endif %}>
            啟用 LINE 條件找物件指令（#找物件 / #搜尋物件）
          </label>
          <div class="hint mt-2">
            建議開啟方式：平常啟用，需要控管 Gemini 額度時可先關閉。<br>
            指令範例：<span class="code">#找物件 清水 18000內 2房 平車</span>
          </div>
        </div>
"""
    if 'id="ai-recommend"' not in LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX:
        LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX = LINE_SETTINGS_CENTER_TEMPLATE_PERMISSION_MATRIX.replace(
            '        <div id="rules" class="setting-card p-4 mb-4">',
            _ai_settings_card + '\n        <div id="rules" class="setting-card p-4 mb-4">'
        )

    try:
        app.jinja_env.globals["ai_recommend_feature_enabled"] = ai_recommend_feature_enabled
    except Exception:
        pass
    print("✅ 設定中心已加入 AI推薦開關")
except Exception as e:
    print("⚠️ AI推薦開關套用失敗：", e)


def _ai_free_search_feature_enabled():
    if not ai_recommend_feature_enabled():
        return False
    try:
        settings = get_line_card_settings() if "get_line_card_settings" in globals() else {}
        if "ai_free_search_enabled" in settings:
            return bool(settings.get("ai_free_search_enabled"))
    except Exception:
        pass
    return True


def _parse_ai_free_property_search_text(text: str):
    text = (text or "").strip()
    if not text:
        return None
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    first = lines[0].strip()
    first_no_space = first.replace(" ", "")
    valid = ("#找物件", "#搜尋物件", "#AI找物件", "#查物件", "#找公司物件")
    if not any(first_no_space.startswith(v) for v in valid):
        return None

    inline = first
    for v in valid:
        if inline.replace(" ", "").startswith(v):
            inline = re.sub(r"^\s*" + re.escape(v) + r"\s*", "", inline, flags=re.I)
            break

    fields = {}
    free_parts = []
    if inline.strip():
        free_parts.append(inline.strip())
    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line)
        if m:
            key = normalize_line_key(m.group(1)) if "normalize_line_key" in globals() else m.group(1).strip()
            fields[key] = (m.group(2) or "").strip()
        else:
            free_parts.append(line)

    all_text = "\n".join(free_parts + [f"{k}: {v}" for k, v in fields.items() if v]).strip()
    if not all_text:
        all_text = "請依一般客需條件推薦物件"

    all_text_for_type = all_text + " " + " ".join(str(v) for v in fields.values())
    is_rent = any(k in all_text_for_type for k in ["租", "月租", "租金", "出租", "承租"])

    budget = fields.get("budget") or fields.get("預算") or fields.get("price") or ""
    rent_max = fields.get("rent_max") or fields.get("租金") or ""
    if is_rent and not rent_max:
        rent_max = budget

    buyer = {
        "id": "free_search_" + hashlib.sha1(all_text.encode("utf-8")).hexdigest()[:16],
        "name": "條件搜尋",
        "phone": "",
        "intent_type": "rent" if is_rent else "buy",
        "budget_max": "" if is_rent else budget,
        "rent_max": rent_max if is_rent else "",
        "preferred_areas": fields.get("preferred_areas") or fields.get("區域") or _ai_extract_areas_from_text(all_text),
        "property_type": fields.get("property_type") or fields.get("產品類型") or fields.get("類型") or _ai_extract_property_types_from_text(all_text),
        "room_range": fields.get("room_range") or fields.get("房數") or "",
        "car_need": fields.get("car_need") or fields.get("車位") or "",
        "requirement_must": fields.get("必要條件") or fields.get("must_have") or "",
        "requirement_nice": fields.get("加分條件") or fields.get("nice_to_have") or "",
        "note": all_text,
        "source": "LINE條件搜尋",
        "visibility": "public",
    }
    top_n = parse_int_limit(fields.get("limit") or fields.get("筆數") or 5, default=5, max_value=10) if "parse_int_limit" in globals() else 5
    return buyer, all_text, top_n


def process_line_free_property_search_event(event):
    msg = (event or {}).get("message") or {}
    if msg.get("type") != "text":
        return {"handled": False}
    parsed = _parse_ai_free_property_search_text(msg.get("text") or "")
    if not parsed:
        return {"handled": False}
    if not _ai_free_search_feature_enabled():
        return {"handled": True, "ok": False, "reply_text": "AI條件找物件功能目前尚未啟用，請到設定中心開啟。"}

    buyer, extra_text, top_n = parsed
    cached = _get_ai_recommend_cache(buyer, top_n=top_n, extra_text=extra_text) if "_get_ai_recommend_cache" in globals() else None
    if cached:
        result = {"parsed_need": cached.get("parsed_need") or {}, "recommendations": cached.get("recommendations") or [], "cache_hit": True}
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {"handled": True, "ok": True, "reply_text": "條件搜尋推薦物件", "reply_flex": flex, "parsed_tag": "找物件"}

    target_id = _line_event_target_id(event) if "_line_event_target_id" in globals() else ""
    if AI_RECOMMEND_BACKGROUND_ENABLED and "_start_ai_recommend_background_job" in globals():
        started = _start_ai_recommend_background_job(buyer, top_n=top_n, extra_text=extra_text, target_id=target_id)
        if started:
            return {"handled": True, "ok": True, "reply_text": "已收到，正在依照你輸入的條件 AI 搜尋公司物件。完成後會自動把推薦卡片傳回來。", "parsed_tag": "找物件"}

    try:
        result = recommend_properties_for_buyer_data(buyer, top_n=top_n, extra_text=extra_text, force_refresh=False)
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {"handled": True, "ok": True, "reply_text": "條件搜尋推薦物件", "reply_flex": flex, "parsed_tag": "找物件"}
    except Exception as e:
        return {"handled": True, "ok": False, "reply_text": f"AI找物件失敗：{e}"}


try:
    _process_line_message_event_before_ai_free_search = process_line_message_event

    def process_line_message_event(event):
        msg = (event or {}).get("message") or {}
        if msg.get("type") == "text":
            text = (msg.get("text") or "").strip()
            first = text.splitlines()[0].strip().replace(" ", "") if text else ""
            if first.startswith(("#推薦物件", "#AI推薦物件", "#物件推薦")) and not ai_recommend_feature_enabled():
                return {"handled": True, "ok": False, "reply_text": "AI推薦物件功能目前尚未啟用，請到設定中心開啟。"}
            free_result = process_line_free_property_search_event(event)
            if free_result.get("handled"):
                return free_result
        return _process_line_message_event_before_ai_free_search(event)

    print("✅ LINE 已支援 #找物件 / #搜尋物件 條件搜尋")
except Exception as e:
    print("⚠️ #找物件 指令套用失敗：", e)

try:
    _detect_line_command_type_before_ai_free_search = detect_line_command_type

    def detect_line_command_type(text: str, event=None) -> str:
        first = (text or "").strip().splitlines()[0].strip().replace(" ", "") if (text or "").strip() else ""
        first_no_hash = first[1:] if first.startswith("#") else first
        if first_no_hash in ("找物件", "搜尋物件", "AI找物件", "查物件", "找公司物件"):
            return "buyer"
        return _detect_line_command_type_before_ai_free_search(text, event=event)

    print("✅ LINE 指令權限已加入 #找物件")
except Exception as e:
    print("⚠️ #找物件 權限套用失敗：", e)

try:
    _buyer_ai_recommend_page_before_ai_toggle = app.view_functions.get("buyer_ai_recommend_page")

    def buyer_ai_recommend_page_ai_toggle_wrapper(buyer_id):
        if not ai_recommend_feature_enabled():
            flash("AI推薦物件功能目前尚未啟用，請到設定中心開啟。", "warning")
            return redirect(url_for("buyer_detail", buyer_id=buyer_id))
        return _buyer_ai_recommend_page_before_ai_toggle(buyer_id)

    if _buyer_ai_recommend_page_before_ai_toggle:
        app.view_functions["buyer_ai_recommend_page"] = login_required(buyer_ai_recommend_page_ai_toggle_wrapper)
        print("✅ 後台 AI推薦頁已加入開關控制")
except Exception as e:
    print("⚠️ 後台 AI推薦頁開關控制套用失敗：", e)

print("✅ AI推薦 V3 已載入：風險有內容才顯示、保留介紹話術、支援 #找物件、設定中心可開關")
# =============================================================================
# AI推薦 V3 Patch End
# =============================================================================


# =============================================================================
# AI推薦 V4：群組/個人權限 + 物件搜尋明確區分買賣/租賃
# - 設定中心的群組 / 個人「可使用指令」新增 ai_recommend
# - #推薦物件、#找物件、#搜尋物件、推薦物件 postback 都改走 ai_recommend 權限
# - 未授權的群組/個人不能使用 AI推薦與條件找物件
# - #找物件 必須能判斷買賣或租賃，避免混用出售/出租物件庫
# =============================================================================

AI_RECOMMEND_COMMAND_KEY = "ai_recommend"

try:
    # 讓設定中心的群組 / 個人可勾選 AI推薦權限。
    if "LINE_COMMAND_TYPE_OPTIONS" in globals():
        if not any(k == AI_RECOMMEND_COMMAND_KEY for k, _ in LINE_COMMAND_TYPE_OPTIONS):
            LINE_COMMAND_TYPE_OPTIONS.append((
                AI_RECOMMEND_COMMAND_KEY,
                "AI推薦 / 找物件指令：#推薦物件、#找物件、#搜尋物件、#AI找物件"
            ))
    print("✅ AI推薦 V4：設定中心已加入群組/個人 AI推薦權限選項")
except Exception as e:
    print("⚠️ AI推薦 V4 加入權限選項失敗：", e)


def _line_source_ai_recommend_allowed(event):
    """檢查目前 LINE 來源是否有 AI推薦/找物件權限。"""
    try:
        kind, target_id = line_event_source_kind_and_id(event)
    except Exception:
        kind, target_id = "unknown", ""

    cfg = None
    label = ""
    if kind in ("group", "room"):
        try:
            cfg = find_line_group_by_target_id(target_id)
        except Exception:
            cfg = None
        label = "此群組"
    elif kind == "user":
        try:
            cfg = find_line_personal_user_by_user_id(target_id)
        except Exception:
            cfg = None
        label = "此個人帳號"
    else:
        label = "此 LINE 來源"

    if not cfg:
        return False, f"未授權：{label}尚未在後台設定中心綁定，無法使用 AI推薦 / 找物件。"

    allowed = set(cfg.get("command_types") or [])
    if "all" in allowed or AI_RECOMMEND_COMMAND_KEY in allowed:
        return True, ""

    return False, f"未授權：{label}尚未開放 AI推薦 / 找物件權限，請到設定中心勾選「AI推薦 / 找物件指令」。"


try:
    _detect_line_command_type_before_ai_v4 = detect_line_command_type

    def detect_line_command_type(text: str, event=None) -> str:
        first = (text or "").strip().splitlines()[0].strip().replace(" ", "") if (text or "").strip() else ""
        first_no_hash = first[1:] if first.startswith("#") else first
        if first_no_hash in (
            "推薦物件", "AI推薦物件", "物件推薦",
            "找物件", "搜尋物件", "AI找物件", "查物件", "找公司物件"
        ):
            return AI_RECOMMEND_COMMAND_KEY
        return _detect_line_command_type_before_ai_v4(text, event=event)

    print("✅ AI推薦 V4：LINE 文字指令已改走 ai_recommend 權限")
except Exception as e:
    print("⚠️ AI推薦 V4 文字權限 patch 失敗：", e)


try:
    _detect_line_postback_command_type_before_ai_v4 = detect_line_postback_command_type

    def detect_line_postback_command_type(event) -> str:
        try:
            data = ((event or {}).get("postback") or {}).get("data", "") or ""
            if "action=ai_recommend_properties" in data or "property_talking_point" in data:
                return AI_RECOMMEND_COMMAND_KEY
        except Exception:
            pass
        return _detect_line_postback_command_type_before_ai_v4(event)

    print("✅ AI推薦 V4：推薦物件 postback 已改走 ai_recommend 權限")
except Exception as e:
    print("⚠️ AI推薦 V4 postback 權限 patch 失敗：", e)


def _detect_property_search_deal_type(text: str, fields: dict = None):
    """回傳 buy / rent / 空字串。租賃關鍵字優先，避免「月租1.8萬」被誤判買賣。"""
    fields = fields or {}
    values_text = " ".join(str(v or "") for v in fields.values())
    source_text = f"{text or ''} {values_text}".strip().lower()
    source_text_no_space = source_text.replace(" ", "")

    # 指定欄位優先。
    explicit = (
        fields.get("intent_type_raw") or fields.get("deal_type_raw") or
        fields.get("intent_type") or fields.get("deal_type") or fields.get("需求類型") or ""
    )
    explicit = str(explicit).strip().lower()

    rent_words = ["租賃", "租屋", "出租", "承租", "租金", "月租", "租"]
    buy_words = ["買賣", "買屋", "買房", "購屋", "出售", "售屋", "總價", "買"]

    if any(w in explicit for w in rent_words):
        return "rent"
    if any(w in explicit for w in buy_words):
        return "buy"

    # 全文判斷：租賃優先。
    if any(w in source_text_no_space for w in rent_words):
        return "rent"
    if any(w in source_text_no_space for w in buy_words):
        return "buy"

    # 有明確「萬」通常是買賣，但租賃文字已在上面優先排除。
    if re.search(r"\d+(?:\.\d+)?\s*萬", source_text):
        return "buy"

    return ""


def _ai_extract_areas_from_text(text: str):
    text = str(text or "")
    area_words = ["清水", "清水區", "梧棲", "梧棲區", "沙鹿", "沙鹿區", "龍井", "龍井區", "大甲", "大甲區", "大安", "大安區", "外埔", "外埔區", "后里", "后里區", "神岡", "神岡區", "西屯", "西屯區", "北屯", "北屯區", "南屯", "南屯區"]
    found = []
    for a in area_words:
        if a in text:
            name = a if a.endswith("區") else a + "區"
            found.append(name)
    return " ".join(dedupe_keep_order(found))


def _ai_extract_property_types_from_text(text: str):
    text = str(text or "")
    types = ["透天", "別墅", "大樓", "電梯大樓", "華廈", "公寓", "套房", "店面", "土地", "農地", "農舍", "廠房", "倉庫"]
    found = []
    for t in types:
        if t in text:
            found.append(t)
    return " ".join(dedupe_keep_order(found))


def _parse_ai_free_property_search_text(text: str):
    """解析 #找物件；V5 支援主建物坪/室內坪/使用坪數硬篩。"""
    text = (text or "").strip()
    if not text:
        return None

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if not lines:
        return None

    first = lines[0].strip()
    first_no_space = first.replace(" ", "")
    valid = ("#找物件", "#搜尋物件", "#AI找物件", "#查物件", "#找公司物件")
    if not any(first_no_space.startswith(v) for v in valid):
        return None

    inline = first
    for v in valid:
        if inline.replace(" ", "").startswith(v):
            inline = re.sub(r"^\s*" + re.escape(v) + r"\s*", "", inline, flags=re.I)
            break

    fields = {}
    free_parts = []
    if inline.strip():
        free_parts.append(inline.strip())

    for line in lines[1:]:
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line)
        if m:
            key = normalize_line_key(m.group(1)) if "normalize_line_key" in globals() else m.group(1).strip()
            fields[key] = (m.group(2) or "").strip()
        else:
            free_parts.append(line)

    all_text = "\n".join(free_parts + [f"{k}: {v}" for k, v in fields.items() if v]).strip()
    if not all_text:
        all_text = "請依一般客需條件推薦物件"

    deal_type = _detect_property_search_deal_type(all_text, fields)
    if not deal_type:
        return {
            "error": (
                "請先指定要找「買賣」或「租賃」物件，避免系統混用出售/出租資料。\n\n"
                "範例一：#找物件 租賃 清水 18000內 2房 平車\n"
                "範例二：#找物件 買賣 梧棲 1500萬內 透天 3房\n\n"
                "也可以用多行：\n"
                "#找物件\n需求類型: 租賃\n預算: 18000\n區域: 清水\n產品類型: 大樓\n房數: 2房"
            )
        }

    budget = fields.get("budget") or fields.get("price") or fields.get("預算") or ""
    rent_max = fields.get("rent_max") or fields.get("租金") or ""
    # 支援口語：租 2.5萬 / 25000內 / 月租18000內
    if not budget and deal_type == "buy":
        m = re.search(r"(\d+(?:\.\d+)?)\s*萬\s*(?:內|以下|以內)?", all_text)
        if m:
            budget = m.group(1)
    if deal_type == "rent" and not rent_max:
        m = re.search(r"(?:租金|月租|租)?\s*(\d+(?:\.\d+)?)\s*萬\s*(?:內|以下|以內)?", all_text)
        if m:
            rent_max = str(float(m.group(1)) * 10000)
        else:
            m = re.search(r"(?:租金|月租|租)?\s*(\d{4,6})\s*(?:元|內|以下|以內)?", all_text)
            if m:
                rent_max = m.group(1)
        if not rent_max:
            rent_max = budget

    buyer = {
        "id": f"free_search_{deal_type}_" + hashlib.sha1((deal_type + "\n" + all_text).encode("utf-8")).hexdigest()[:16],
        "name": "租賃條件搜尋" if deal_type == "rent" else "買賣條件搜尋",
        "phone": "",
        "intent_type": "rent" if deal_type == "rent" else "buy",
        "budget_max": "" if deal_type == "rent" else budget,
        "rent_max": rent_max if deal_type == "rent" else "",
        "preferred_areas": fields.get("preferred_areas") or fields.get("區域") or "",
        "property_type": fields.get("property_type") or fields.get("產品類型") or fields.get("類型") or "",
        "room_range": fields.get("room_range") or fields.get("房數") or "",
        "car_need": fields.get("car_need") or fields.get("車位") or "",
        "requirement_must": fields.get("必要條件") or fields.get("must_have") or "",
        "requirement_nice": fields.get("加分條件") or fields.get("nice_to_have") or "",
        "note": all_text,
        "source": "LINE條件搜尋",
        "visibility": "public",
    }
    top_n = parse_int_limit(fields.get("limit") or fields.get("筆數") or 5, default=5, max_value=10) if "parse_int_limit" in globals() else 5
    return buyer, all_text, top_n


try:
    _process_line_recommend_properties_event_before_ai_v4 = process_line_recommend_properties_event

    def process_line_recommend_properties_event(event):
        msg = (event or {}).get("message") or {}
        text = (msg.get("text") or "").strip()
        first = text.splitlines()[0].strip().replace(" ", "") if text else ""
        if first not in ("#推薦物件", "#AI推薦物件", "#物件推薦"):
            return {"handled": False}
        ok, reason = _line_source_ai_recommend_allowed(event)
        if not ok:
            return {"handled": True, "ok": False, "reply_text": reason}
        return _process_line_recommend_properties_event_before_ai_v4(event)

    print("✅ AI推薦 V4：#推薦物件 已套用群組/個人權限")
except Exception as e:
    print("⚠️ AI推薦 V4 #推薦物件 權限 patch 失敗：", e)


try:
    _process_line_free_property_search_event_before_ai_v4 = process_line_free_property_search_event
except Exception:
    _process_line_free_property_search_event_before_ai_v4 = None


def process_line_free_property_search_event(event):
    msg = (event or {}).get("message") or {}
    if msg.get("type") != "text":
        return {"handled": False}

    parsed = _parse_ai_free_property_search_text(msg.get("text") or "")
    if not parsed:
        return {"handled": False}

    if isinstance(parsed, dict) and parsed.get("error"):
        return {"handled": True, "ok": False, "reply_text": parsed.get("error")}

    ok, reason = _line_source_ai_recommend_allowed(event)
    if not ok:
        return {"handled": True, "ok": False, "reply_text": reason}

    if not _ai_free_search_feature_enabled():
        return {"handled": True, "ok": False, "reply_text": "AI條件找物件功能目前尚未啟用，請到設定中心開啟。"}

    buyer, extra_text, top_n = parsed
    cached = _get_ai_recommend_cache(buyer, top_n=top_n, extra_text=extra_text) if "_get_ai_recommend_cache" in globals() else None
    if cached:
        result = {"parsed_need": cached.get("parsed_need") or {}, "recommendations": cached.get("recommendations") or [], "cache_hit": True}
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {"handled": True, "ok": True, "reply_text": "條件搜尋推薦物件", "reply_flex": flex, "parsed_tag": "找物件"}

    target_id = _line_event_target_id(event) if "_line_event_target_id" in globals() else ""
    if AI_RECOMMEND_BACKGROUND_ENABLED and "_start_ai_recommend_background_job" in globals():
        started = _start_ai_recommend_background_job(buyer, top_n=top_n, extra_text=extra_text, target_id=target_id)
        if started:
            return {"handled": True, "ok": True, "reply_text": "已收到，正在依照你輸入的買賣/租賃條件 AI 搜尋公司物件。完成後會自動把推薦卡片傳回來。", "parsed_tag": "找物件"}

    try:
        result = recommend_properties_for_buyer_data(buyer, top_n=top_n, extra_text=extra_text, force_refresh=False)
        flex = build_property_recommend_flex(buyer, result.get("recommendations") or [], parsed_need=result.get("parsed_need") or {})
        return {"handled": True, "ok": True, "reply_text": "條件搜尋推薦物件", "reply_flex": flex, "parsed_tag": "找物件"}
    except Exception as e:
        return {"handled": True, "ok": False, "reply_text": f"AI找物件失敗：{e}"}


try:
    # 再包一次 process_line_message_event，確保 V4 的解析與權限最後生效。
    _process_line_message_event_before_ai_v4 = process_line_message_event

    def process_line_message_event(event):
        msg = (event or {}).get("message") or {}
        if msg.get("type") == "text":
            text = (msg.get("text") or "").strip()
            first = text.splitlines()[0].strip().replace(" ", "") if text else ""
            if first.startswith(("#推薦物件", "#AI推薦物件", "#物件推薦")):
                rec_result = process_line_recommend_properties_event(event)
                if rec_result.get("handled"):
                    return rec_result
            free_result = process_line_free_property_search_event(event)
            if free_result.get("handled"):
                return free_result
        return _process_line_message_event_before_ai_v4(event)

    print("✅ AI推薦 V4：LINE 訊息處理已套用最後版權限與買賣/租賃判斷")
except Exception as e:
    print("⚠️ AI推薦 V4 訊息處理 patch 失敗：", e)


try:
    _process_line_postback_event_before_ai_v4 = process_line_postback_event

    def process_line_postback_event(event):
        try:
            data = ((event or {}).get("postback") or {}).get("data", "") or ""
            if "action=ai_recommend_properties" in data:
                ok, reason = _line_source_ai_recommend_allowed(event)
                if not ok:
                    return {"handled": True, "ok": False, "reply_text": reason}
        except Exception:
            pass
        return _process_line_postback_event_before_ai_v4(event)

    print("✅ AI推薦 V4：推薦物件卡片按鈕已套用群組/個人權限")
except Exception as e:
    print("⚠️ AI推薦 V4 postback 處理 patch 失敗：", e)


print("✅ AI推薦 V4 已載入：AI推薦依群組/個人權限控管，#找物件 需明確指定買賣或租賃")
# =============================================================================
# AI推薦 V4 End
# =============================================================================

# =============================================================================
# 委託：屋主回報管理 Patch v20260624_OWNER_REPORT
# - 委託詳細頁新增屋主回報管理
# - 可手動新增回報紀錄
# - Gemini 依委託資料、追蹤紀錄、備註、歷史回報生成屋主回報草稿
# - 下一次回報日自動建立待辦事項
# - 委託列表顯示上次回報 / 需回報狀態
# =============================================================================

OWNER_REPORT_COLLECTION = os.environ.get("OWNER_REPORT_COLLECTION", "seller_owner_reports")
OWNER_REPORT_DRAFT_COLLECTION = os.environ.get("OWNER_REPORT_DRAFT_COLLECTION", "seller_owner_report_drafts")

OWNER_REPORT_TYPE_OPTIONS = [
    ("general", "一般進度回報"),
    ("price_negotiation", "議價回報"),
    ("market_feedback", "市場反應說明"),
    ("showing_feedback", "帶看後回報"),
    ("low_inquiry", "詢問量不足回報"),
    ("contract_expiring", "委託到期前維繫"),
]


def _owner_report_today():
    return now_taipei().strftime("%Y-%m-%d")


def _owner_report_label(value):
    mapping = dict(OWNER_REPORT_TYPE_OPTIONS)
    return mapping.get(value or "", value or "一般回報")


def _owner_report_safe_int(value, default=0):
    try:
        text = str(value or "").replace(",", "")
        m = re.search(r"\d+", text)
        return int(m.group(0)) if m else default
    except Exception:
        return default


def _owner_report_doc_to_dict(doc):
    data = doc.to_dict() or {}
    data["id"] = doc.id
    data["report_type_label"] = _owner_report_label(data.get("report_type"))
    return data


def _get_seller_or_redirect(seller_id):
    snap = db.collection("sellers").document(seller_id).get()
    if not snap.exists:
        return None, redirect(url_for("sellers"))
    seller = snap.to_dict() or {}
    seller["id"] = snap.id
    try:
        if "backend_can_view_personal_record" in globals() and not backend_can_view_personal_record(seller):
            flash("你沒有權限查看這筆個人委託", "danger")
            return None, redirect(url_for("sellers"))
    except Exception:
        pass
    return seller, None


def fetch_seller_owner_reports(seller_id, limit=50):
    reports = []
    try:
        docs = db.collection(OWNER_REPORT_COLLECTION).where("seller_id", "==", seller_id).stream()
        reports = [_owner_report_doc_to_dict(d) for d in docs]
    except Exception as e:
        print("⚠️ 讀取屋主回報失敗：", e)
        reports = []
    reports.sort(key=lambda x: (x.get("report_date") or x.get("created_at") or ""), reverse=True)
    return reports[:limit]


def attach_owner_report_info_to_sellers(sellers_list):
    reports_by_seller = {}
    try:
        for d in db.collection(OWNER_REPORT_COLLECTION).stream():
            item = _owner_report_doc_to_dict(d)
            sid = item.get("seller_id")
            if sid:
                reports_by_seller.setdefault(sid, []).append(item)
    except Exception as e:
        print("⚠️ 委託列表讀取屋主回報狀態失敗：", e)

    today = _owner_report_today()
    for s in sellers_list:
        sid = s.get("id")
        reports = reports_by_seller.get(sid, [])
        reports.sort(key=lambda x: (x.get("report_date") or x.get("created_at") or ""), reverse=True)
        last = reports[0] if reports else None
        s["last_owner_report"] = last
        if not last:
            s["owner_report_status"] = "尚未回報"
            s["owner_report_status_class"] = "secondary"
            s["owner_report_due_text"] = "尚未設定"
            continue
        next_date = last.get("next_report_date") or ""
        s["owner_report_due_text"] = next_date or "未設定"
        if next_date and next_date <= today:
            s["owner_report_status"] = "需回報"
            s["owner_report_status_class"] = "danger"
        elif (last.get("owner_response") or "").find("待回覆") >= 0:
            s["owner_report_status"] = "屋主待回覆"
            s["owner_report_status_class"] = "warning"
        else:
            s["owner_report_status"] = "已回報"
            s["owner_report_status_class"] = "success"
    return sellers_list


def _owner_report_format_seller_basic(seller):
    return "\n".join([
        f"屋主/委託人：{seller.get('name','')}",
        f"電話：{seller.get('phone','')}",
        f"物件地址：{seller.get('address','')}",
        f"產品類型：{seller.get('property_type','')}",
        f"目前狀態：{seller.get('stage','')}",
        f"開價/期望價：{seller.get('expected_price','')}",
        f"底價：{seller.get('min_price','')}",
        f"委託到期日：{seller.get('contract_end_date','')}",
        f"出售/出租原因：{seller.get('reason','')}",
        f"內部備註：{seller.get('note','')}",
    ]).strip()


def _fetch_seller_followups_for_report(seller_id, limit=20):
    rows = []
    try:
        for d in db.collection("seller_followups").where("seller_id", "==", seller_id).stream():
            item = d.to_dict() or {}
            item["id"] = d.id
            rows.append(item)
    except Exception:
        rows = []
    rows.sort(key=lambda x: (x.get("contact_time") or x.get("created_at") or ""), reverse=True)
    return rows[:limit]


def _owner_report_gemini_json(prompt):
    if "_gemini_generate_json" in globals():
        return _gemini_generate_json(prompt)
    api_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not api_key:
        raise RuntimeError("尚未設定 GEMINI_API_KEY")
    from google import genai
    client = genai.Client(api_key=api_key)
    res = client.models.generate_content(
        model=os.environ.get("GEMINI_MODEL", "gemini-2.5-flash"),
        contents=prompt,
    )
    text = getattr(res, "text", "") or ""
    return json.loads(re.search(r"\{.*\}", text, re.S).group(0))


def generate_owner_report_draft_with_ai(seller, report_type="general", period_start="", period_end="", extra_note=""):
    followups = _fetch_seller_followups_for_report(seller.get("id"), limit=25)
    reports = fetch_seller_owner_reports(seller.get("id"), limit=5)
    followup_text = []
    for f in followups:
        followup_text.append(
            f"- {f.get('contact_time') or f.get('created_at') or ''}｜{f.get('channel') or ''}｜內容：{f.get('content') or ''}｜下一步：{f.get('next_action') or ''}"
        )
    report_text = []
    for r in reports:
        report_text.append(
            f"- {r.get('report_date') or ''}｜{r.get('report_type_label') or ''}｜內容：{r.get('report_message') or r.get('progress_summary') or ''}｜屋主回應：{r.get('owner_response') or ''}"
        )

    prompt = f"""
你是台中海線房仲的屋主回報助理。請根據委託資料、追蹤紀錄、歷史屋主回報，生成一份可以傳給屋主的回報草稿。

請注意：
1. 語氣要專業、口語、讓屋主覺得有在積極處理。
2. 不要編造沒有資料支持的詢問量或帶看量。
3. 如果資料不足，請用「目前紀錄顯示」或「目前回報重點」表達，不要硬說很多詢問。
4. 如果回報目的是議價，要用客戶反應與市場回饋委婉建議，不要太強硬。
5. 請只輸出 JSON，不要輸出其他文字。

回報目的：{_owner_report_label(report_type)}
回報區間：{period_start or '-'} ～ {period_end or _owner_report_today()}
補充需求：{extra_note or '-'}

委託資料：
{_owner_report_format_seller_basic(seller)}

近期追蹤紀錄：
{chr(10).join(followup_text) if followup_text else '目前沒有追蹤紀錄'}

歷史屋主回報：
{chr(10).join(report_text) if report_text else '目前沒有歷史屋主回報'}

請輸出格式：
{{
  "progress_summary": "本次進度摘要",
  "customer_feedback": "客戶反應整理",
  "owner_response": "如果沒有屋主回應就留空字串",
  "next_action": "下一步建議",
  "ai_suggestion": "給房仲看的策略建議，例如是否議價、是否補曝光、是否確認屋主想法",
  "report_message": "可以直接傳給屋主的完整回報文字，繁體中文，口語但專業"
}}
""".strip()

    try:
        data = _owner_report_gemini_json(prompt) or {}
    except Exception as e:
        print("⚠️ Gemini 生成屋主回報失敗，改用規則草稿：", e)
        data = {}

    if not data:
        recent = followups[0] if followups else {}
        latest_content = recent.get("content") or seller.get("note") or "目前持續曝光與追蹤物件進度。"
        data = {
            "progress_summary": latest_content,
            "customer_feedback": "目前客戶反應仍需持續彙整。",
            "owner_response": "",
            "next_action": "持續追蹤詢問與帶看狀況，並定期回報屋主。",
            "ai_suggestion": "資料不足，建議先補充詢問量、帶看量與客戶反應，再進一步判斷是否需要議價。",
            "report_message": f"您好，跟您回報一下目前物件的進度。\n\n目前我們這邊有持續追蹤與整理物件狀況，近期紀錄重點為：{latest_content}\n\n接下來我會持續觀察詢問與帶看反應，有新的客戶回饋或進一步進度，我會再跟您回報。",
        }
    for key in ("progress_summary", "customer_feedback", "owner_response", "next_action", "ai_suggestion", "report_message"):
        data[key] = str(data.get(key) or "").strip()
    return data


def create_owner_report_todo(seller, report, report_id=""):
    next_date = (report.get("next_report_date") or "").strip()
    if not next_date:
        return ""
    try:
        collection_name = LINE_TODO_COLLECTION if "LINE_TODO_COLLECTION" in globals() else "line_todos"
    except Exception:
        collection_name = "line_todos"
    visibility = report.get("visibility") or seller.get("visibility") or "public"
    owner_line_user_id = report.get("owner_line_user_id") or seller.get("owner_line_user_id") or ""
    owner_line_name = report.get("owner_line_name") or seller.get("owner_line_name") or ""
    data = {
        "title": f"屋主回報：{seller.get('name') or '委託'}",
        "content": f"物件：{seller.get('address') or '-'}\n下一步：{report.get('next_action') or '-'}",
        "note": f"關聯委託：{seller.get('name') or ''}｜{seller.get('phone') or ''}",
        "todo_date": next_date,
        "status": "open",
        "visibility": visibility if visibility in ("public", "personal") else "public",
        "owner_line_user_id": owner_line_user_id if visibility == "personal" else "",
        "owner_line_name": owner_line_name if visibility == "personal" else "",
        "line_target_id": owner_line_user_id if visibility == "personal" else "",
        "line_target_type": "user" if visibility == "personal" else "backend_shared",
        "source": "屋主回報管理",
        "related_type": "seller_owner_report",
        "seller_id": seller.get("id"),
        "seller_name": seller.get("name"),
        "seller_phone": seller.get("phone"),
        "owner_report_id": report_id,
        "created_at": now_taipei().isoformat(),
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
    }
    ref = db.collection(collection_name).document()
    ref.set(data)
    return ref.id


def _owner_report_form_payload(seller, form):
    report_date = (form.get("report_date") or _owner_report_today()).strip()
    report = {
        "seller_id": seller.get("id"),
        "seller_name": seller.get("name") or "",
        "seller_phone": seller.get("phone") or "",
        "seller_address": seller.get("address") or "",
        "report_date": report_date,
        "report_period_start": (form.get("report_period_start") or "").strip(),
        "report_period_end": (form.get("report_period_end") or "").strip(),
        "report_type": (form.get("report_type") or "general").strip(),
        "inquiry_count": _owner_report_safe_int(form.get("inquiry_count"), 0),
        "showing_count": _owner_report_safe_int(form.get("showing_count"), 0),
        "customer_feedback": (form.get("customer_feedback") or "").strip(),
        "progress_summary": (form.get("progress_summary") or "").strip(),
        "owner_response": (form.get("owner_response") or "").strip(),
        "next_action": (form.get("next_action") or "").strip(),
        "next_report_date": (form.get("next_report_date") or "").strip(),
        "ai_suggestion": (form.get("ai_suggestion") or "").strip(),
        "report_message": (form.get("report_message") or "").strip(),
        "sent_to_owner": bool(form.get("sent_to_owner")),
        "sent_at": now_taipei().isoformat() if form.get("sent_to_owner") else "",
        "visibility": seller.get("visibility") or "public",
        "owner_line_user_id": seller.get("owner_line_user_id") or "",
        "owner_line_name": seller.get("owner_line_name") or "",
        "created_at": now_taipei().isoformat(),
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
    }
    return report


@app.route("/sellers/<seller_id>/owner-report/new", methods=["POST"])
@login_required
def seller_owner_report_create(seller_id):
    seller, resp = _get_seller_or_redirect(seller_id)
    if resp:
        return resp
    report = _owner_report_form_payload(seller, request.form)
    if not report.get("report_message") and not report.get("progress_summary") and not report.get("customer_feedback"):
        flash("請輸入回報內容，或先用 AI 生成屋主回報草稿", "warning")
        return redirect(url_for("seller_detail", seller_id=seller_id))
    ref = db.collection(OWNER_REPORT_COLLECTION).document()
    ref.set(report)
    todo_id = create_owner_report_todo(seller, report, ref.id) if report.get("next_report_date") else ""
    if todo_id:
        ref.update({"next_report_todo_id": todo_id})
    # 同步寫入一筆委託追蹤，讓整體紀錄串在一起。
    try:
        db.collection("seller_followups").add({
            "seller_id": seller_id,
            "contact_time": f"{report.get('report_date')} 09:00",
            "channel": "屋主回報",
            "content": report.get("report_message") or report.get("progress_summary") or "已完成屋主回報",
            "next_action": report.get("next_action") or "",
            "next_contact_date": report.get("next_report_date") or "",
            "created_at": now_taipei().isoformat(),
            "created_by_id": session.get("user_id"),
            "created_by_name": session.get("user_name"),
        })
    except Exception as e:
        print("⚠️ 屋主回報同步追蹤紀錄失敗：", e)
    flash("已新增屋主回報" + ("，並建立下一次回報待辦" if todo_id else ""), "success")
    return redirect(url_for("seller_detail", seller_id=seller_id))


@app.route("/sellers/<seller_id>/owner-report/ai", methods=["POST"])
@login_required
def seller_owner_report_ai_generate(seller_id):
    seller, resp = _get_seller_or_redirect(seller_id)
    if resp:
        return resp
    report_type = (request.form.get("report_type") or "general").strip()
    period_start = (request.form.get("report_period_start") or "").strip()
    period_end = (request.form.get("report_period_end") or _owner_report_today()).strip()
    extra_note = (request.form.get("extra_note") or "").strip()
    draft = generate_owner_report_draft_with_ai(seller, report_type=report_type, period_start=period_start, period_end=period_end, extra_note=extra_note)
    draft.update({
        "seller_id": seller_id,
        "report_type": report_type,
        "report_period_start": period_start,
        "report_period_end": period_end,
        "report_date": _owner_report_today(),
        "extra_note": extra_note,
        "created_at": now_taipei().isoformat(),
        "created_by_id": session.get("user_id"),
        "created_by_name": session.get("user_name"),
    })
    ref = db.collection(OWNER_REPORT_DRAFT_COLLECTION).document()
    ref.set(draft)
    flash("AI 已生成屋主回報草稿，請確認後儲存", "success")
    return redirect(url_for("seller_detail", seller_id=seller_id, owner_report_draft_id=ref.id))


@app.route("/sellers/<seller_id>/owner-report/<report_id>/delete", methods=["POST"])
@login_required
def seller_owner_report_delete(seller_id, report_id):
    seller, resp = _get_seller_or_redirect(seller_id)
    if resp:
        return resp
    ref = db.collection(OWNER_REPORT_COLLECTION).document(report_id)
    snap = ref.get()
    if snap.exists:
        ref.delete()
        flash("已刪除屋主回報", "info")
    return redirect(url_for("seller_detail", seller_id=seller_id))


def seller_detail_with_owner_reports(seller_id):
    seller, resp = _get_seller_or_redirect(seller_id)
    if resp:
        flash("找不到這位委託或無權限查看", "danger")
        return resp
    followups_ref = db.collection("seller_followups").where("seller_id", "==", seller_id)
    followups = [doc_to_dict(f) for f in followups_ref.stream()]
    followups.sort(key=lambda x: x.get("contact_time", ""), reverse=True)
    owner_reports = fetch_seller_owner_reports(seller_id, limit=30)
    draft = {}
    draft_id = (request.args.get("owner_report_draft_id") or "").strip()
    if draft_id:
        try:
            ds = db.collection(OWNER_REPORT_DRAFT_COLLECTION).document(draft_id).get()
            if ds.exists:
                d = ds.to_dict() or {}
                if d.get("seller_id") == seller_id:
                    draft = d
                    draft["id"] = draft_id
        except Exception:
            draft = {}
    return render_template(
        "seller_detail.html",
        seller=seller,
        followups=followups,
        owner_reports=owner_reports,
        ai_owner_report_draft=draft,
        owner_report_type_options=OWNER_REPORT_TYPE_OPTIONS,
        today_date=_owner_report_today(),
    )


def sellers_with_owner_reports():
    q = request.args.get("q", "").strip()
    level = request.args.get("level", "").strip()
    stage = request.args.get("stage", "").strip()
    source = request.args.get("source", "").strip()
    label = request.args.get("label", "").strip()
    sort_by = request.args.get("sort_by", "created_at_desc")

    docs = db.collection("sellers").stream()
    all_sellers = [doc_to_dict(d) for d in docs]
    try:
        visible_sellers = _backend_visible_items(all_sellers) if "_backend_visible_items" in globals() else all_sellers
    except Exception:
        visible_sellers = all_sellers

    source_options = sorted({(s.get("source") or "").strip() for s in visible_sellers if (s.get("source") or "").strip()})
    label_options = build_label_options(visible_sellers)
    sellers_list = list(visible_sellers)

    if q:
        sellers_list = [s for s in sellers_list if q in (s.get("name") or "") or q in (s.get("phone") or "") or q in (s.get("address") or "")]
    if level:
        sellers_list = [s for s in sellers_list if s.get("level") == level]
    if stage:
        sellers_list = [s for s in sellers_list if (s.get("stage") or "") == stage]
    if source:
        sellers_list = [s for s in sellers_list if (s.get("source") or "") == source]
    if label:
        sellers_list = [s for s in sellers_list if label in ensure_list(s.get("labels"))]

    if sort_by == "created_at_asc":
        sellers_list.sort(key=lambda s: s.get("created_at") or "")
    elif sort_by == "created_at_desc":
        sellers_list.sort(key=lambda s: s.get("created_at") or "", reverse=True)
    elif sort_by == "name_asc":
        sellers_list.sort(key=lambda s: (s.get("name") or ""))
    elif sort_by == "name_desc":
        sellers_list.sort(key=lambda s: (s.get("name") or ""), reverse=True)

    sellers_list = attach_latest_followup(sellers_list, "seller_followups", "seller_id")
    sellers_list = attach_owner_report_info_to_sellers(sellers_list)
    return render_template(
        "sellers.html",
        sellers=sellers_list,
        q=q,
        level=level,
        stage=stage,
        source=source,
        source_options=source_options,
        label=label,
        label_options=label_options,
        sort_by=sort_by,
        seller_stage_options=SELLER_STAGE_OPTIONS,
        total_count=len(visible_sellers),
        filtered_count=len(sellers_list),
    )

try:
    app.view_functions["seller_detail"] = login_required(seller_detail_with_owner_reports)
    app.view_functions["sellers"] = login_required(sellers_with_owner_reports)
    print("✅ 屋主回報管理已啟用：委託詳細 / AI生成回報 / 下次回報待辦 / 委託列表回報狀態")
except Exception as e:
    print("⚠️ 屋主回報管理套用失敗：", e)

# =============================================================================
# 委託：屋主回報管理 Patch End
# =============================================================================


if __name__ == "__main__":
    print("✅ FULL_READY_20260621 已載入")
    print("✅ /line-card-preview 已註冊：", any(rule.rule == "/line-card-preview" for rule in app.url_map.iter_rules()))
    print("✅ 可用 /debug/routes 檢查目前所有 route")
    app.run(debug=True)



# =============================================================================
# LINE URI 長度修正 Patch
# - LINE Flex URI action 的 uri 最長 1000 字元。
# - 原本「加入行事曆」把 note / location 等長內容塞進 query string，
#   委託或客需備註很長時會導致傳送群組失敗：size must be between 0 and 1000。
# - 這裡覆寫 _record_calendar_url：只帶必要參數，長文字截短，避免超過 LINE 限制。
# =============================================================================

def _crm_short_text_for_uri(value, max_len=80):
    value = str(value or '').replace('\r', ' ').replace('\n', ' ').strip()
    value = re.sub(r'\s+', ' ', value)
    return value[:max_len]


def _record_calendar_url(record_type: str, record_id: str, data: dict):
    """產生 LINE Flex 可用的短網址，避免 URI 超過 1000 字元。"""
    record_type = (record_type or '').strip()
    record_id = (record_id or '').strip()
    data = data or {}

    name = data.get('name') or data.get('customer_name') or ''
    phone = data.get('phone') or ''

    if record_type == 'buyer':
        title = f"{name} 客需追蹤".strip() or '客需追蹤'
        category = '回電'
        location = data.get('preferred_areas') or ''
    elif record_type == 'seller':
        title = f"{name} 委託追蹤".strip() or '委託追蹤'
        category = '回電'
        location = data.get('address') or ''
    elif record_type == 'development':
        title = f"{name or data.get('address') or '開發'} 開發追蹤".strip()
        category = '開發'
        location = data.get('address') or data.get('registered_address') or ''
    else:
        title = 'CRM 追蹤'
        category = '回電'
        location = ''

    # note 容易非常長，是造成 LINE 傳送失敗的主因。
    # 行事曆表單可用 related_type / related_id 回到原資料查看完整內容，因此這裡不塞完整 note。
    short_note = f"來源：{record_type}｜ID：{record_id}"

    params = {
        'related_type': record_type,
        'related_id': record_id,
        'title': _crm_short_text_for_uri(title, 50),
        'category': category,
        'customer_name': _crm_short_text_for_uri(name, 30),
        'phone': re.sub(r'[^0-9+]', '', str(phone or ''))[:20],
        'location': _crm_short_text_for_uri(location, 60),
        'note': _crm_short_text_for_uri(short_note, 80),
        'visibility': data.get('visibility') or 'public',
        'owner_line_user_id': _crm_short_text_for_uri(data.get('owner_line_user_id') or '', 80),
    }
    url = _crm_public_url_for('/calendar/new', **{k: v for k, v in params.items() if v})

    # 最後防呆：若 APP_BASE_URL 很長或特殊資料導致仍超過，降級到最短網址。
    if len(url) > 950:
        url = _crm_public_url_for('/calendar/new', related_type=record_type, related_id=record_id)
    return url

print('✅ LINE URI 長度修正已載入：加入行事曆按鈕不再塞入超長備註')
# =============================================================================
# LINE URI 長度修正 Patch End
# =============================================================================

# =============================================================================
# 委託：案件輸入表 + AI 強銷文案 Patch v20260624_CASE_FORM_AI
# - 委託詳細頁可進入「案件表 / AI強銷」
# - 可填屋主資料、坪數資料、帶看方式、備註與原始群組文案
# - Gemini 生成：強銷標題、五點特色、公司群組文案、刊登描述
# - 一鍵產生 A4 可列印 PDF 案件輸入表
# - LINE 指令：#補委託資料 / #生成強銷 / #生成案件表
# =============================================================================

CASE_DETAIL_FIELD_KEYS = [
    "property_title", "community_name", "case_address", "case_price", "layout",
    "total_ping", "main_ping", "attached_ping", "public_ping", "land_ping", "parking_ping",
    "floor", "floor_total", "building_age", "facing", "showing_method", "case_note",
    "life_note", "property_highlight_note", "target_customer_note", "raw_group_text", "source_url",
    "ai_sales_title", "ai_selling_points", "ai_group_copy", "ai_listing_description", "ai_feature_note",
]

CASE_BASIC_LABELS = {
    "property_title": "物件標題",
    "community_name": "社區名稱",
    "case_address": "完整地址",
    "case_price": "開價/售價",
    "layout": "格局",
    "total_ping": "總建坪",
    "main_ping": "主建坪",
    "attached_ping": "附屬坪",
    "public_ping": "公設坪",
    "land_ping": "地坪",
    "parking_ping": "車位坪",
    "floor": "所在樓層",
    "floor_total": "樓高/總樓層",
    "building_age": "屋齡",
    "facing": "朝向/座向",
    "showing_method": "帶看方式",
    "case_note": "備註",
    "life_note": "生活機能補充",
    "property_highlight_note": "物件亮點補充",
    "target_customer_note": "適合客群",
    "source_url": "參考網址",
}


def _case_clean_text(value, max_len=None):
    text = str(value or "").replace("\r", "").strip()
    text = re.sub(r"[ \t]+", " ", text)
    if max_len and len(text) > max_len:
        return text[:max_len - 1] + "…"
    return text


def _case_normalize_fullwidth(text):
    text = str(text or "")
    table = str.maketrans("０１２３４５６７８９：，．／－（）", "0123456789:,./-()")
    return text.translate(table)


def _case_first_match(text, patterns, default=""):
    text = _case_normalize_fullwidth(text)
    for pat in patterns:
        m = re.search(pat, text, flags=re.I | re.M)
        if m:
            val = (m.group(1) or "").strip()
            val = re.split(r"\n|_{3,}|-{3,}|—{3,}", val)[0].strip()
            return val
    return default


def _case_parse_raw_listing_text(raw_text: str):
    """從公司群組貼文 / 補充資料中先用規則拆欄位，AI 生成前可先儲存。"""
    raw = str(raw_text or "").strip()
    text = _case_normalize_fullwidth(raw)
    data = {"raw_group_text": raw}
    data["case_address"] = _case_first_match(text, [r"(?:地址|物件地址)\s*[:：]\s*(.+)"])
    data["case_price"] = _case_first_match(text, [r"(?:開價|售價|總價|價格)\s*[:：]\s*(.+)"])
    data["layout"] = _case_first_match(text, [r"(?:格局|房廳衛)\s*[:：]\s*(.+)"])
    data["total_ping"] = _case_first_match(text, [r"(?:總建坪|總坪|建坪|建物面積)\s*[:：]\s*(.+)"])
    data["main_ping"] = _case_first_match(text, [r"(?:主建坪|主建物|主建)\s*[:：]\s*(.+)"])
    data["attached_ping"] = _case_first_match(text, [r"(?:附屬坪|附屬)\s*[:：]\s*(.*)"])
    data["public_ping"] = _case_first_match(text, [r"(?:公設坪|公設)\s*[:：]\s*(.*)"])
    data["land_ping"] = _case_first_match(text, [r"(?:土地坪|地坪|土地面積)\s*[:：]\s*(.+)"])
    data["parking_ping"] = _case_first_match(text, [r"(?:車位坪|車位)\s*[:：]\s*(.+)"])
    data["floor_total"] = _case_first_match(text, [r"(?:樓高|總樓層)\s*[:：]\s*(.+)"])
    data["floor"] = _case_first_match(text, [r"(?:樓層|所在樓層)\s*[:：]\s*(.+)"])
    data["building_age"] = _case_first_match(text, [r"(?:屋齡)\s*[:：]\s*(.+)"])
    data["facing"] = _case_first_match(text, [r"(?:朝向|座向)\s*[:：]\s*(.*)"])
    data["showing_method"] = _case_first_match(text, [r"(?:帶看方式|看屋方式)\s*[:：]\s*(.+)"])
    data["case_note"] = _case_first_match(text, [r"(?:備註|特色備註)\s*[:：]\s*(.*)"])
    url_match = re.search(r"https?://[^\s]+", raw)
    if url_match:
        data["source_url"] = url_match.group(0).strip()

    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]
    candidates = []
    for ln in lines[:8]:
        if re.search(r"[:：]", ln):
            continue
        if re.fullmatch(r"[_\-—=\s]+", ln):
            continue
        if re.search(r"新接|再麻煩|學長|學姐|謝謝|其餘資料後補", ln):
            continue
        candidates.append(ln)
    if candidates:
        data["property_title"] = candidates[0][:80]

    bullet_lines = []
    for ln in lines:
        if re.match(r"^\d+[\.、．)]", _case_normalize_fullwidth(ln)):
            bullet_lines.append(re.sub(r"^\d+[\.、．)]\s*", "", ln).strip())
    if bullet_lines:
        data["property_highlight_note"] = "\n".join(bullet_lines[:8])

    return {k: v for k, v in data.items() if str(v or "").strip()}


def _case_merge_seller_and_case_data(seller: dict, incoming=None):
    incoming = incoming or {}
    data = {}
    data["property_title"] = seller.get("property_title") or seller.get("ai_sales_title") or ""
    data["community_name"] = seller.get("community_name") or ""
    data["case_address"] = seller.get("case_address") or seller.get("address") or ""
    data["case_price"] = seller.get("case_price") or seller.get("expected_price") or ""
    data["layout"] = seller.get("layout") or ""
    data["total_ping"] = seller.get("total_ping") or seller.get("building_area") or ""
    data["main_ping"] = seller.get("main_ping") or ""
    data["attached_ping"] = seller.get("attached_ping") or ""
    data["public_ping"] = seller.get("public_ping") or ""
    data["land_ping"] = seller.get("land_ping") or ""
    data["parking_ping"] = seller.get("parking_ping") or ""
    data["floor"] = seller.get("floor") or ""
    data["floor_total"] = seller.get("floor_total") or seller.get("total_floors") or ""
    data["building_age"] = seller.get("building_age") or seller.get("age") or ""
    data["facing"] = seller.get("facing") or ""
    data["showing_method"] = seller.get("showing_method") or ""
    data["case_note"] = seller.get("case_note") or seller.get("note") or ""
    data["life_note"] = seller.get("life_note") or ""
    data["property_highlight_note"] = seller.get("property_highlight_note") or ""
    data["target_customer_note"] = seller.get("target_customer_note") or ""
    data["raw_group_text"] = seller.get("raw_group_text") or ""
    data["source_url"] = seller.get("source_url") or ""
    data["ai_sales_title"] = seller.get("ai_sales_title") or ""
    data["ai_selling_points"] = seller.get("ai_selling_points") or []
    data["ai_group_copy"] = seller.get("ai_group_copy") or ""
    data["ai_listing_description"] = seller.get("ai_listing_description") or ""
    data["ai_feature_note"] = seller.get("ai_feature_note") or ""
    for k, v in (incoming or {}).items():
        if k in CASE_DETAIL_FIELD_KEYS and str(v or "").strip():
            data[k] = v
    return data


def _case_seller_payload_from_form(form):
    payload = {}
    for key in CASE_DETAIL_FIELD_KEYS:
        if key == "ai_selling_points":
            continue
        val = form.get(key, "") if hasattr(form, "get") else ""
        if val is not None:
            payload[key] = _case_clean_text(val)
    return payload


def _case_format_seller_context(seller: dict, case_data: dict):
    lines = []
    base = {
        "屋主姓名": seller.get("name"),
        "電話": seller.get("phone"),
        "委託類型": seller.get("deal_type"),
        "產品類型": seller.get("property_type"),
        "開價": seller.get("expected_price"),
        "底價": seller.get("min_price"),
        "委託到期日": seller.get("contract_end_date"),
        "內部備註": seller.get("note"),
    }
    for label, value in base.items():
        if value:
            lines.append(f"{label}：{value}")
    for key, label in CASE_BASIC_LABELS.items():
        value = case_data.get(key)
        if value:
            lines.append(f"{label}：{value}")
    if case_data.get("raw_group_text"):
        lines.append("原始群組資料：")
        lines.append(_case_clean_text(case_data.get("raw_group_text"), 2500))
    return "\n".join(lines)


def _case_ai_generate_sales_copy(seller: dict, case_data: dict):
    context = _case_format_seller_context(seller, case_data)
    prompt = f"""
你是台中海線房仲的委託物件強銷文案助手。
請根據資料產生 JSON，不要輸出 JSON 以外的文字。

物件資料：
{context}

請輸出格式：
{{
  "ai_sales_title": "最多 24 字的強銷標題",
  "ai_selling_points": ["五點特色，每點 18 到 45 字，共 5 點"],
  "ai_group_copy": "公司 LINE 群組可直接貼上的完整文案，包含標題、基本資料、五大強銷亮點與帶看方式",
  "ai_listing_description": "可放到樂屋/591的較完整描述，約 120-220 字",
  "ai_feature_note": "案件輸入表特色備註，約 50-100 字"
}}

生成規則：
1. 五點特色必須至少涵蓋：生活機能、物件本身亮點、物件優勢判斷、適合客群、成交切入點。
2. 不要把總坪、主建、地坪、附屬坪數當成五點特色的主要賣點；坪數只放在基本資料區。
3. 不要亂編未提供或未查證的地標、學校、商圈。如果資料只有地址，請用保守說法如「在地生活圈」「主要道路動線」「日常生活需求」。
4. 語氣要像房仲公司群組可直接轉傳，清楚、有賣點、不浮誇。
5. 若屋齡較高，不要用負面字眼；可說「適合依買方喜好重新整理」。
6. 如果帶看方式有提供，群組文案要保留。
""".strip()
    try:
        if "_gemini_generate_json" not in globals():
            raise RuntimeError("Gemini helper 不存在")
        data = _gemini_generate_json(prompt)
        if not isinstance(data, dict):
            raise RuntimeError("Gemini 回傳格式不是 JSON object")
        points = data.get("ai_selling_points") or data.get("selling_points") or []
        if isinstance(points, str):
            points = [p.strip() for p in re.split(r"\n+", points) if p.strip()]
        points = [str(p).strip() for p in points if str(p).strip()][:5]
        if len(points) < 5:
            raise RuntimeError("Gemini 回傳特色不足 5 點")
        data["ai_selling_points"] = points
        data["ai_sales_title"] = _case_clean_text(data.get("ai_sales_title") or data.get("title") or "", 40)
        data["ai_group_copy"] = _case_clean_text(data.get("ai_group_copy") or "", 4000)
        data["ai_listing_description"] = _case_clean_text(data.get("ai_listing_description") or "", 800)
        data["ai_feature_note"] = _case_clean_text(data.get("ai_feature_note") or "", 300)
        return data
    except Exception as e:
        print("⚠️ Gemini 生成強銷文案失敗，改用規則版：", e)
        return _case_rule_generate_sales_copy(seller, case_data)


def _case_extract_district(address):
    m = re.search(r"([\u4e00-\u9fff]{2,4}區)", str(address or ""))
    return m.group(1) if m else "在地"


def _case_rule_generate_sales_copy(seller: dict, case_data: dict):
    address = case_data.get("case_address") or seller.get("address") or ""
    district = _case_extract_district(address)
    ptype = seller.get("property_type") or case_data.get("property_type") or "物件"
    price = case_data.get("case_price") or seller.get("expected_price") or ""
    layout = case_data.get("layout") or "多房格局"
    showing = case_data.get("showing_method") or "請提前預約"
    note = case_data.get("property_highlight_note") or case_data.get("case_note") or seller.get("note") or ""

    title_core = case_data.get("property_title") or f"{district}｜低總價{layout}{ptype}"
    title = _case_clean_text(title_core, 30) or f"{district}優質委託物件"
    points = [
        f"位於{district}生活圈，適合想找在地生活機能與通勤動線的買方。",
        f"{layout}空間好運用，適合家庭成員多、需要書房或工作室的客群。",
        f"{ptype}使用彈性高，生活規劃自由度比一般大樓更有空間。",
        f"{price}的總價帶好切入，適合首購、換屋或在地自住客評估。" if price else "總價帶具討論空間，適合有明確自住需求的客戶評估。",
        "屋況可依買方喜好重新整理，適合想打造自己風格住家的客戶。",
    ]
    if "全新" in note or "新" in note:
        points[4] = "屋況條件佳，買方後續整理成本較低，入住規劃更省心。"
    if "朝南" in note or "朝南" in str(case_data.get("facing")):
        points[1] = "朝南條件加分，採光與居住舒適度更容易吸引自住型買方。"
    if "學" in note or "國小" in note or "國中" in note:
        points[0] = "鄰近學區與在地生活圈，對有小孩接送需求的家庭更有吸引力。"

    basic_lines = [
        title, "",
        f"地址：{address or '-'}",
        "——————————————",
        f"開價：{price or seller.get('expected_price') or '-'}",
        f"格局：{layout or '-'}",
        f"總建坪：{case_data.get('total_ping') or '-'}",
        f"主建坪：{case_data.get('main_ping') or '-'}",
        f"附屬：{case_data.get('attached_ping') or '-'}",
        f"地坪：{case_data.get('land_ping') or '-'}",
        f"樓高：{case_data.get('floor_total') or '-'}",
        f"屋齡：{case_data.get('building_age') or '-'}",
        f"朝向：{case_data.get('facing') or '-'}",
        f"帶看方式：{showing or '-'}",
        "——————————————",
        "五大強銷亮點：",
    ]
    for i, p in enumerate(points, 1):
        basic_lines.append(f"{i}. {p}")
    basic_lines.extend(["", f"帶看方式：{showing or '請提前預約'}", "再麻煩學長姐多多介紹，謝謝！"])
    return {
        "ai_sales_title": title,
        "ai_selling_points": points,
        "ai_group_copy": "\n".join(basic_lines),
        "ai_listing_description": " ".join(points),
        "ai_feature_note": "；".join(points[:3]),
    }


def _case_apply_ai_to_seller(seller_id: str, case_payload: dict, ai_payload: dict = None):
    updates = dict(case_payload or {})
    ai_payload = ai_payload or {}
    for k in ["ai_sales_title", "ai_group_copy", "ai_listing_description", "ai_feature_note"]:
        if k in ai_payload:
            updates[k] = ai_payload.get(k)
    if ai_payload.get("ai_selling_points"):
        updates["ai_selling_points"] = ai_payload.get("ai_selling_points")[:5]
    updates["case_form_updated_at"] = now_taipei().isoformat()
    updates["updated_at"] = now_taipei().isoformat()
    updates["updated_by_id"] = session.get("user_id") or "line_bot"
    updates["updated_by_name"] = session.get("user_name") or "LINE Bot"
    db.collection("sellers").document(seller_id).set(updates, merge=True)
    return updates


def _case_missing_fields(case_data: dict):
    required = ["case_address", "case_price", "layout", "total_ping", "main_ping", "land_ping", "floor_total", "building_age", "showing_method"]
    return [CASE_BASIC_LABELS.get(k, k) for k in required if not str(case_data.get(k) or "").strip()]


CASE_TOOLS_HTML = r'''
{% extends "base.html" %}
{% block content %}
<div class="d-flex justify-content-between align-items-center mb-3">
  <div>
    <h3 class="mb-1">案件輸入表 / AI強銷文案</h3>
    <div class="text-muted small">委託：{{ seller.name or '-' }}｜{{ seller.phone or '-' }}</div>
  </div>
  <div>
    <a class="btn btn-outline-primary" target="_blank" href="{{ url_for('seller_case_form_pdf', seller_id=seller.id) }}">下載案件輸入表 PDF</a>
    <a class="btn btn-secondary" href="{{ url_for('seller_detail', seller_id=seller.id) }}">回委託詳細</a>
  </div>
</div>

{% if missing_fields %}
<div class="alert alert-warning small">
  <strong>資料尚可補強：</strong>{{ missing_fields|join('、') }}。沒有資料的欄位 PDF 會先留白或顯示待確認。
</div>
{% endif %}

<div class="row g-4">
  <div class="col-lg-7">
    <form method="post" class="card">
      <div class="card-header fw-bold">案件資料</div>
      <div class="card-body">
        <div class="mb-3">
          <label class="form-label">貼上公司群組原始文案 / 補充資料</label>
          <textarea name="raw_group_text" class="form-control" rows="8" placeholder="可直接貼公司群組整段物件資料，系統會先自動拆欄位，再交給 AI 生成標題與五點強銷。">{{ case_data.raw_group_text or '' }}</textarea>
          <div class="form-text">送出時會自動解析：地址、售價、格局、坪數、屋齡、帶看方式、網址與條列特色。</div>
        </div>

        <div class="row g-2">
          {% for key, label in field_labels.items() %}
            {% if key not in ['life_note','property_highlight_note','target_customer_note','source_url'] %}
              <div class="col-md-6">
                <label class="form-label">{{ label }}</label>
                <input type="text" name="{{ key }}" class="form-control" value="{{ case_data.get(key, '') }}">
              </div>
            {% endif %}
          {% endfor %}
        </div>

        <hr>
        <div class="row g-2">
          <div class="col-md-4">
            <label class="form-label">生活機能補充</label>
            <textarea name="life_note" class="form-control" rows="4" placeholder="例：中山路生活圈、近市場、近學校、近主要道路">{{ case_data.life_note or '' }}</textarea>
          </div>
          <div class="col-md-4">
            <label class="form-label">物件亮點補充</label>
            <textarea name="property_highlight_note" class="form-control" rows="4" placeholder="例：朝南、間間套房、全新、屋主自住、可整理">{{ case_data.property_highlight_note or '' }}</textarea>
          </div>
          <div class="col-md-4">
            <label class="form-label">適合客群</label>
            <textarea name="target_customer_note" class="form-control" rows="4" placeholder="例：在地換屋、大家庭、首購、自住整理型客戶">{{ case_data.target_customer_note or '' }}</textarea>
          </div>
          <div class="col-12">
            <label class="form-label">參考網址</label>
            <input type="text" name="source_url" class="form-control" value="{{ case_data.source_url or '' }}">
          </div>
        </div>
      </div>
      <div class="card-footer d-flex gap-2 flex-wrap">
        <button class="btn btn-primary" name="action" value="generate_ai" type="submit">AI整理並生成強銷文案</button>
        <button class="btn btn-outline-secondary" name="action" value="save_only" type="submit">只儲存資料</button>
        <a class="btn btn-outline-primary" target="_blank" href="{{ url_for('seller_case_form_pdf', seller_id=seller.id) }}">下載案件輸入表 PDF</a>
      </div>
    </form>
  </div>

  <div class="col-lg-5">
    <div class="card mb-3">
      <div class="card-header fw-bold">AI 強銷標題與五點特色</div>
      <div class="card-body">
        <label class="form-label">強銷標題</label>
        <input class="form-control mb-3" readonly value="{{ case_data.ai_sales_title or '' }}">
        <label class="form-label">五點特色</label>
        {% if case_data.ai_selling_points %}
          <ol class="mb-0">
            {% for p in case_data.ai_selling_points %}
              <li class="mb-2">{{ p }}</li>
            {% endfor %}
          </ol>
        {% else %}
          <div class="text-muted">尚未生成。左側填資料後按「AI整理並生成強銷文案」。</div>
        {% endif %}
      </div>
    </div>

    <div class="card mb-3">
      <div class="card-header fw-bold">公司群組文案</div>
      <div class="card-body">
        <textarea id="aiGroupCopy" class="form-control" rows="16" readonly>{{ case_data.ai_group_copy or '' }}</textarea>
        <button class="btn btn-sm btn-outline-secondary mt-2" onclick="navigator.clipboard.writeText(document.getElementById('aiGroupCopy').value); alert('已複製群組文案');">複製群組文案</button>
      </div>
    </div>

    <div class="card">
      <div class="card-header fw-bold">案件輸入表特色備註</div>
      <div class="card-body">
        <textarea class="form-control" rows="5" readonly>{{ case_data.ai_feature_note or '' }}</textarea>
      </div>
    </div>
  </div>
</div>
{% endblock %}
'''


@app.route("/sellers/<seller_id>/case-tools", methods=["GET", "POST"])
@login_required
def seller_case_tools(seller_id):
    doc_ref = db.collection("sellers").document(seller_id)
    snap = doc_ref.get()
    if not snap.exists:
        flash("找不到這筆委託", "danger")
        return redirect(url_for("sellers"))
    seller = doc_to_dict(snap)
    if request.method == "POST":
        form_payload = _case_seller_payload_from_form(request.form)
        raw_text = form_payload.get("raw_group_text") or ""
        parsed_payload = _case_parse_raw_listing_text(raw_text) if raw_text else {}
        merged_payload = dict(parsed_payload)
        for k, v in form_payload.items():
            if str(v or "").strip():
                merged_payload[k] = v
        case_data = _case_merge_seller_and_case_data(seller, merged_payload)
        action = request.form.get("action", "save_only")
        if action == "generate_ai":
            ai_payload = _case_ai_generate_sales_copy(seller, case_data)
            _case_apply_ai_to_seller(seller_id, case_data, ai_payload)
            flash("已整理案件資料，並生成 AI 強銷文案", "success")
        else:
            _case_apply_ai_to_seller(seller_id, case_data, {})
            flash("已儲存案件資料", "success")
        return redirect(url_for("seller_case_tools", seller_id=seller_id))
    case_data = _case_merge_seller_and_case_data(seller)
    missing_fields = _case_missing_fields(case_data)
    return render_template_string(CASE_TOOLS_HTML, seller=seller, case_data=case_data, field_labels=CASE_BASIC_LABELS, missing_fields=missing_fields)


def _case_pdf_value(case_data, key, default=""):
    value = case_data.get(key)
    if isinstance(value, list):
        return "、".join(map(str, value))
    return str(value or default or "")


def _case_generate_pdf_bytes(seller: dict, case_data: dict):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.pdfbase import pdfmetrics
    except Exception as e:
        raise RuntimeError("缺少 reportlab，請先安裝：pip install reportlab") from e
    buffer = BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=10*mm, leftMargin=10*mm, topMargin=10*mm, bottomMargin=10*mm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="CJKTitle", parent=styles["Title"], fontName="STSong-Light", fontSize=18, leading=22, alignment=1))
    styles.add(ParagraphStyle(name="CJK", parent=styles["Normal"], fontName="STSong-Light", fontSize=9, leading=13))
    styles.add(ParagraphStyle(name="CJKSmall", parent=styles["Normal"], fontName="STSong-Light", fontSize=8, leading=11))
    def P(text, small=False):
        return Paragraph(str(text or "").replace("\n", "<br/>"), styles["CJKSmall" if small else "CJK"])
    story = [Paragraph("太平洋房屋 幸福房仲團隊 - 案件輸入表", styles["CJKTitle"]), Spacer(1, 5*mm)]
    top_rows = [
        [P("委託類別"), P("☑ 出售  □ 出租" if (seller.get("deal_type") or "sale") != "rent" else "□ 出售  ☑ 出租"), P("案件編號"), P(seller.get("id") or "")],
        [P("屋主姓名"), P(seller.get("name") or ""), P("電話"), P(seller.get("phone") or "")],
        [P("委託到期日"), P(seller.get("contract_end_date") or ""), P("帶看方式"), P(_case_pdf_value(case_data, "showing_method"))],
    ]
    table_style = TableStyle([("FONTNAME", (0,0), (-1,-1), "STSong-Light"), ("FONTSIZE", (0,0), (-1,-1), 9), ("GRID", (0,0), (-1,-1), 0.4, colors.black), ("VALIGN", (0,0), (-1,-1), "MIDDLE"), ("BACKGROUND", (0,0), (0,-1), colors.whitesmoke), ("BACKGROUND", (2,0), (2,-1), colors.whitesmoke)])
    t = Table(top_rows, colWidths=[28*mm, 67*mm, 28*mm, 67*mm]); t.setStyle(table_style); story.append(t); story.append(Spacer(1,4*mm))
    basic_rows = [
        [P("物件名稱"), P(_case_pdf_value(case_data, "property_title") or _case_pdf_value(case_data, "ai_sales_title")), P("社區名稱"), P(_case_pdf_value(case_data, "community_name"))],
        [P("物件地址"), P(_case_pdf_value(case_data, "case_address") or seller.get("address") or ""), P("產品類型"), P(seller.get("property_type") or "")],
        [P("開價/售價"), P(_case_pdf_value(case_data, "case_price") or seller.get("expected_price") or ""), P("底價"), P(seller.get("min_price") or "")],
        [P("格局"), P(_case_pdf_value(case_data, "layout")), P("樓層 / 樓高"), P(f"{_case_pdf_value(case_data, 'floor') or '-'} / {_case_pdf_value(case_data, 'floor_total') or '-'}")],
        [P("總建坪"), P(_case_pdf_value(case_data, "total_ping")), P("主建坪"), P(_case_pdf_value(case_data, "main_ping"))],
        [P("附屬坪"), P(_case_pdf_value(case_data, "attached_ping")), P("公設坪"), P(_case_pdf_value(case_data, "public_ping"))],
        [P("地坪"), P(_case_pdf_value(case_data, "land_ping")), P("車位坪"), P(_case_pdf_value(case_data, "parking_ping"))],
        [P("屋齡"), P(_case_pdf_value(case_data, "building_age")), P("朝向"), P(_case_pdf_value(case_data, "facing"))],
    ]
    t = Table(basic_rows, colWidths=[28*mm, 67*mm, 28*mm, 67*mm]); t.setStyle(table_style); story.append(Paragraph("1. 基本資料", styles["CJK"])); story.append(t); story.append(Spacer(1,4*mm))
    points = case_data.get("ai_selling_points") or []
    points_text = "<br/>".join([f"{i}. {p}" for i, p in enumerate(points[:5], 1)]) if points else (_case_pdf_value(case_data, "ai_feature_note") or _case_pdf_value(case_data, "case_note"))
    desc_rows = [[P("特色備註"), P(points_text or "")], [P("生活機能補充"), P(_case_pdf_value(case_data, "life_note"))], [P("物件亮點補充"), P(_case_pdf_value(case_data, "property_highlight_note"))], [P("適合客群"), P(_case_pdf_value(case_data, "target_customer_note"))], [P("產權特別注意事項"), P("")], [P("合約日 / 租約日"), P("合約日：______年____月____日  至  ______年____月____日")]]
    t = Table(desc_rows, colWidths=[38*mm, 152*mm]); t.setStyle(TableStyle([("FONTNAME", (0,0), (-1,-1), "STSong-Light"), ("FONTSIZE", (0,0), (-1,-1), 9), ("GRID", (0,0), (-1,-1), 0.4, colors.black), ("VALIGN", (0,0), (-1,-1), "TOP"), ("BACKGROUND", (0,0), (0,-1), colors.whitesmoke)])); story.append(Paragraph("2. 學區 / 環境 / 特色備註", styles["CJK"])); story.append(t); story.append(Spacer(1,3*mm)); story.append(Paragraph("PS. 未提供或待確認資料請列印後手寫補上。AI 生成內容仍建議人工確認。", styles["CJKSmall"]))
    doc.build(story)
    buffer.seek(0)
    return buffer


@app.route("/sellers/<seller_id>/case-form.pdf")
@login_required
def seller_case_form_pdf(seller_id):
    snap = db.collection("sellers").document(seller_id).get()
    if not snap.exists:
        flash("找不到這筆委託", "danger")
        return redirect(url_for("sellers"))
    seller = doc_to_dict(snap)
    case_data = _case_merge_seller_and_case_data(seller)
    try:
        buf = _case_generate_pdf_bytes(seller, case_data)
    except Exception as e:
        flash(f"產生 PDF 失敗：{e}", "danger")
        return redirect(url_for("seller_case_tools", seller_id=seller_id))
    filename = f"案件輸入表_{seller.get('name') or seller_id}.pdf"
    return send_file(buf, mimetype="application/pdf", as_attachment=True, download_name=filename)


def _case_find_seller_for_line(fields, raw_text=""):
    record_id = fields.get("record_id") or fields.get("seller_id") or fields.get("委託ID") or ""
    phone = fields.get("phone") or fields.get("電話") or ""
    name = fields.get("name") or fields.get("姓名") or ""
    if record_id:
        doc = db.collection("sellers").document(str(record_id).strip()).get()
        if doc.exists:
            return doc
    if phone:
        doc = find_customer_record("seller", phone=phone)
        if doc:
            return doc
    if name:
        doc = find_customer_record("seller", name=name)
        if doc:
            return doc
    parsed = _case_parse_raw_listing_text(raw_text or "")
    addr = parsed.get("case_address") or ""
    if addr:
        key = re.sub(r"\s+", "", addr)[-12:]
        if key:
            for d in db.collection("sellers").stream():
                data = d.to_dict() or {}
                saddr = re.sub(r"\s+", "", (data.get("address") or data.get("case_address") or ""))
                if key and key in saddr:
                    return d
    return None


def _case_parse_simple_fields_from_text(text):
    fields = {}
    for line in str(text or "").splitlines():
        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line.strip())
        if not m:
            continue
        key = normalize_line_key(m.group(1)) if "normalize_line_key" in globals() else m.group(1).strip()
        val = m.group(2).strip()
        fields[key] = val
        fields[m.group(1).strip()] = val
    return fields


def _case_guess_property_type_from_case_text(case_data, fields=None):
    """依案名 / 原文 / 地址簡單推測產品類型，避免 #新增委託 模板完全空白。"""
    fields = fields or {}
    existing = fields.get("property_type") or fields.get("產品類型") or ""
    if str(existing or "").strip():
        return str(existing).strip()
    blob = " ".join([
        str(case_data.get("property_title") or ""),
        str(case_data.get("raw_group_text") or ""),
        str(case_data.get("case_note") or ""),
        str(case_data.get("property_highlight_note") or ""),
    ])
    mapping = [
        ("別墅", ["別墅", "美墅", "墅"]),
        ("透天", ["透天", "透店", "店住"]),
        ("電梯大樓", ["大樓", "三房平車", "兩房平車", "社區", "高樓"]),
        ("華廈", ["華廈"]),
        ("土地", ["土地", "建地", "農地"]),
        ("店面", ["店面", "店鋪"]),
        ("廠房", ["廠房", "工業"]),
        ("公寓", ["公寓"]),
    ]
    for label, keys in mapping:
        if any(k in blob for k in keys):
            return label
    return ""


def _case_short_line_text(value, max_len=80):
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if len(text) <= max_len:
        return text
    return text[:max_len-1] + "…"


def _case_build_new_seller_fill_text(fields=None, case_data=None):
    """LINE 卡片的「新增委託」按鈕使用：打開鍵盤並預填 #新增委託 模板。"""
    fields = fields or {}
    case_data = case_data or {}
    address = case_data.get("case_address") or fields.get("address") or fields.get("地址") or ""
    price = case_data.get("case_price") or fields.get("price") or fields.get("開價") or ""
    ptype = _case_guess_property_type_from_case_text(case_data, fields)
    title = case_data.get("property_title") or case_data.get("ai_sales_title") or ""
    # fillInText 不要太長，避免 LINE postback action 超過限制。
    lines = [
        "#新增委託",
        "姓名:",
        "電話:",
        "客戶來源:",
        "委託類型: 賣",
        f"地址:{_case_short_line_text(address, 60)}" if address else "地址:",
        f"產品類型:{_case_short_line_text(ptype, 20)}" if ptype else "產品類型:",
        f"開價:{_case_short_line_text(price, 24)}" if price else "開價:",
        "底價:",
        "委託到期日:",
        f"備註:{_case_short_line_text(title, 45)}" if title else "備註:",
    ]
    fill = "\n".join(lines)
    return fill[:300]


def _case_build_missing_seller_flex(ai_payload=None, case_data=None, fields=None):
    """已生成 AI 文案但找不到委託資料時，回一張卡片讓使用者一鍵開啟 #新增委託 模板。"""
    ai_payload = ai_payload or {}
    case_data = case_data or {}
    fields = fields or {}
    title = ai_payload.get("ai_sales_title") or case_data.get("property_title") or "AI強銷文案已生成"
    address = case_data.get("case_address") or "未填地址"
    price = case_data.get("case_price") or "未填開價"
    fill_text = _case_build_new_seller_fill_text(fields, {**case_data, **{"ai_sales_title": ai_payload.get("ai_sales_title", "")}})
    return {
        "type": "bubble",
        "size": "mega",
        "body": {
            "type": "box",
            "layout": "vertical",
            "spacing": "md",
            "contents": [
                {"type": "text", "text": "AI文案已生成", "size": "xs", "color": "#C9874A", "weight": "bold"},
                {"type": "text", "text": _case_short_line_text(title, 52), "size": "lg", "weight": "bold", "wrap": True, "color": "#222222"},
                {"type": "separator", "margin": "md"},
                {"type": "box", "layout": "baseline", "spacing": "sm", "margin": "md", "contents": [
                    {"type": "text", "text": "狀態", "size": "sm", "color": "#999999", "flex": 2},
                    {"type": "text", "text": "尚未找到對應委託資料", "size": "sm", "color": "#B00020", "wrap": True, "flex": 5},
                ]},
                {"type": "box", "layout": "baseline", "spacing": "sm", "contents": [
                    {"type": "text", "text": "地址", "size": "sm", "color": "#999999", "flex": 2},
                    {"type": "text", "text": _case_short_line_text(address, 90), "size": "sm", "color": "#333333", "wrap": True, "flex": 5},
                ]},
                {"type": "box", "layout": "baseline", "spacing": "sm", "contents": [
                    {"type": "text", "text": "開價", "size": "sm", "color": "#999999", "flex": 2},
                    {"type": "text", "text": _case_short_line_text(price, 50), "size": "sm", "color": "#333333", "wrap": True, "flex": 5},
                ]},
                {"type": "text", "text": "請先新增委託人資料，之後這份強銷文案與案件資料才會寫入後台。", "size": "sm", "color": "#555555", "wrap": True, "margin": "md"},
            ],
        },
        "footer": {
            "type": "box",
            "layout": "vertical",
            "spacing": "sm",
            "contents": [
                {
                    "type": "button",
                    "style": "primary",
                    "height": "sm",
                    "color": "#C9874A",
                    "action": {
                        "type": "postback",
                        "label": "新增委託",
                        "data": "action=new_seller_from_ai_copy",
                        "inputOption": "openKeyboard",
                        "fillInText": fill_text,
                    },
                },
                {
                    "type": "button",
                    "style": "secondary",
                    "height": "sm",
                    "action": {
                        "type": "postback",
                        "label": "補委託資料模板",
                        "data": "action=case_fill_seller_detail",
                        "inputOption": "openKeyboard",
                        "fillInText": "#補委託資料\n電話:\n\n請貼上物件完整資料",
                    },
                },
            ],
        },
        "styles": {"footer": {"separator": True}},
    }


def _case_reply_ai_copy_with_new_seller_card(ai_payload, case_data, fields=None):
    """LINE 回覆：第一則給 AI 文案，第二則給新增委託卡片。"""
    copy_text = (ai_payload or {}).get("ai_group_copy") or "已生成強銷文案。"
    flex = _case_build_missing_seller_flex(ai_payload, case_data, fields)
    return {
        "handled": True,
        "ok": True,
        "reply_messages": [
            {"type": "text", "text": _case_clean_text(copy_text, 4500)},
            {"type": "flex", "altText": "AI文案已生成，請新增委託資料", "contents": flex},
        ],
        "reply_text": _case_clean_text(copy_text, 4500),
        "parsed_tag": "生成強銷",
    }


def _case_process_line_commands(event):
    message = event.get("message") or {}
    if message.get("type") != "text":
        return {"handled": False}
    text = (message.get("text") or "").strip()
    if not text.startswith(("#補委託資料", "#生成強銷", "#生成群組文案", "#生成案件表")):
        return {"handled": False}
    fields = _case_parse_simple_fields_from_text(text)
    raw_text = re.sub(r"^#(補委託資料|生成強銷|生成群組文案|生成案件表).*\n?", "", text, count=1).strip()
    seller_doc = _case_find_seller_for_line(fields, raw_text)
    if text.startswith("#生成強銷") or text.startswith("#生成群組文案"):
        parsed_payload = _case_parse_raw_listing_text(raw_text)
        if seller_doc:
            seller = doc_to_dict(seller_doc)
            case_data = _case_merge_seller_and_case_data(seller, parsed_payload)
            ai_payload = _case_ai_generate_sales_copy(seller, case_data)
            _case_apply_ai_to_seller(seller_doc.id, case_data, ai_payload)
            return {"handled": True, "ok": True, "reply_text": (ai_payload.get("ai_group_copy") or "已生成強銷文案。")[:4500], "target_type": "seller", "target_id": seller_doc.id, "customer_name": seller.get("name", ""), "phone": seller.get("phone", ""), "parsed_tag": "生成強銷"}
        pseudo_seller = {"name": fields.get("name", ""), "phone": fields.get("phone", ""), "property_type": fields.get("property_type", "")}
        case_data = _case_merge_seller_and_case_data(pseudo_seller, parsed_payload)
        ai_payload = _case_ai_generate_sales_copy(pseudo_seller, case_data)
        # 找不到對應委託時，仍先生成文案，並回一張「新增委託」卡片讓使用者直接 key in 屋主資料。
        return _case_reply_ai_copy_with_new_seller_card(ai_payload, case_data, fields)
    if text.startswith("#生成案件表"):
        if not seller_doc:
            return {"handled": True, "ok": False, "reply_text": "請提供電話或委託ID，才能產生案件輸入表。"}
        seller = doc_to_dict(seller_doc)
        try:
            url = _crm_public_url_for(f"/sellers/{seller_doc.id}/case-form.pdf") if "_crm_public_url_for" in globals() else url_for("seller_case_form_pdf", seller_id=seller_doc.id, _external=True)
        except Exception:
            url = f"/sellers/{seller_doc.id}/case-form.pdf"
        return {"handled": True, "ok": True, "reply_text": f"案件輸入表下載連結：\n{url}\n若尚未登入後台，請先登入後再下載。", "target_type": "seller", "target_id": seller_doc.id, "customer_name": seller.get("name", ""), "phone": seller.get("phone", ""), "parsed_tag": "生成案件表"}
    if not seller_doc:
        parsed_payload = _case_parse_raw_listing_text(raw_text)
        pseudo_seller = {"name": fields.get("name", ""), "phone": fields.get("phone", ""), "property_type": fields.get("property_type", "")}
        case_data = _case_merge_seller_and_case_data(pseudo_seller, parsed_payload)
        ai_payload = _case_ai_generate_sales_copy(pseudo_seller, case_data)
        # 補資料時若尚未建立委託，先生成文案，再提供新增委託模板。
        result = _case_reply_ai_copy_with_new_seller_card(ai_payload, case_data, fields)
        result["parsed_tag"] = "補委託資料"
        return result
    seller = doc_to_dict(seller_doc)
    parsed_payload = _case_parse_raw_listing_text(raw_text)
    case_data = _case_merge_seller_and_case_data(seller, parsed_payload)
    ai_payload = _case_ai_generate_sales_copy(seller, case_data)
    _case_apply_ai_to_seller(seller_doc.id, case_data, ai_payload)
    return {"handled": True, "ok": True, "reply_text": ("已補委託資料並生成強銷文案：\n\n" + (ai_payload.get("ai_group_copy") or ""))[:4500], "target_type": "seller", "target_id": seller_doc.id, "customer_name": seller.get("name", ""), "phone": seller.get("phone", ""), "parsed_tag": "補委託資料"}


try:
    _process_line_message_event_before_case_form_ai = process_line_message_event
    def process_line_message_event(event):
        result = _case_process_line_commands(event)
        if result.get("handled"):
            return result
        return _process_line_message_event_before_case_form_ai(event)
    print("✅ 委託案件輸入表/AI強銷：LINE 指令已啟用 #補委託資料 / #生成強銷 / #生成案件表（找不到委託時會顯示新增委託卡片）")
except Exception as e:
    print("⚠️ 委託案件輸入表/AI強銷 LINE 指令掛入失敗：", e)

try:
    app.jinja_env.globals["seller_case_tools_enabled"] = True
except Exception:
    pass

print("✅ 委託案件輸入表 / AI 強銷文案已啟用：/sellers/<id>/case-tools")
# =============================================================================
# 委託：案件輸入表 + AI 強銷文案 Patch End
# =============================================================================


# =============================================================================
# CASE_FORM_AI V3：修正 LINE 群組「#生成強銷 / #補委託資料 / #生成案件表」被權限閘門判定 unknown
# 原因：line_access_gate 會先呼叫 detect_line_command_type；如果指令未被辨識，還沒進入
# _case_process_line_commands 就會回覆「這個群組沒有可辨識的指令」。
# 這裡把案件表 / AI 強銷指令歸類為 seller 指令，沿用委託權限即可。
# =============================================================================
try:
    CASE_FORM_AI_LINE_COMMANDS = (
        "補委託資料", "生成強銷", "生成群組文案", "生成案件表",
        "AI強銷", "AI文案", "案件表", "產生案件表"
    )

    _detect_line_command_type_before_case_form_ai_v3 = detect_line_command_type

    def detect_line_command_type(text: str, event=None) -> str:
        raw = (text or "").strip().replace("＃", "#")
        if raw:
            first = raw.splitlines()[0].strip().replace(" ", "")
            first_no_hash = first[1:] if first.startswith("#") else first
            if first_no_hash in CASE_FORM_AI_LINE_COMMANDS:
                return "seller"
        return _detect_line_command_type_before_case_form_ai_v3(text, event=event)

    # 設定中心顯示文字同步補上，實際權限仍走 seller。
    try:
        if "LINE_COMMAND_TYPE_OPTIONS" in globals():
            for i, (key, label) in enumerate(LINE_COMMAND_TYPE_OPTIONS):
                if key == "seller" and "#生成強銷" not in label:
                    LINE_COMMAND_TYPE_OPTIONS[i] = (
                        key,
                        label + "、#補委託資料、#生成強銷、#生成群組文案、#生成案件表"
                    )
                    break
    except Exception:
        pass

    print("✅ CASE_FORM_AI V3：LINE 權限閘門已辨識 #生成強銷 / #補委託資料 / #生成案件表")
except Exception as e:
    print("⚠️ CASE_FORM_AI V3 指令辨識修正失敗：", e)


# =============================================================================
# CASE_FORM_AI V4：設定中心新增「AI文案 / 案件表」獨立權限勾選
# - V3 先把 #生成強銷 歸到 seller 指令；但設定頁看起來沒有獨立勾選項目
# - V4 新增 case_ai 指令類型，讓群組 / 個人權限可單獨勾選
# - 同時保留相容：如果舊資料已勾 seller，也仍可使用 case_ai，避免突然失效
# =============================================================================
try:
    CASE_FORM_AI_COMMAND_TYPE = "case_ai"
    CASE_FORM_AI_LINE_COMMANDS = (
        "補委託資料", "生成強銷", "生成群組文案", "生成案件表",
        "AI強銷", "AI文案", "案件表", "產生案件表", "物件文案", "強銷文案"
    )

    # 1) 設定中心可使用指令：新增獨立勾選項目
    try:
        if "LINE_COMMAND_TYPE_OPTIONS" in globals():
            # 先把 V3 加在 seller label 後面的長文字清乾淨，避免重複顯示
            cleaned = []
            for key, label in LINE_COMMAND_TYPE_OPTIONS:
                if key == "seller":
                    label = label.replace("、#補委託資料、#生成強銷、#生成群組文案、#生成案件表", "")
                if key != CASE_FORM_AI_COMMAND_TYPE:
                    cleaned.append((key, label))
            # 放在 seller 後面比較直覺
            inserted = False
            new_options = []
            for key, label in cleaned:
                new_options.append((key, label))
                if key == "seller":
                    new_options.append((CASE_FORM_AI_COMMAND_TYPE, "AI文案 / 案件表：#補委託資料、#生成強銷、#生成群組文案、#生成案件表"))
                    inserted = True
            if not inserted:
                new_options.append((CASE_FORM_AI_COMMAND_TYPE, "AI文案 / 案件表：#補委託資料、#生成強銷、#生成群組文案、#生成案件表"))
            LINE_COMMAND_TYPE_OPTIONS[:] = new_options
    except Exception as e:
        print("⚠️ CASE_FORM_AI V4 設定中心指令選項更新失敗：", e)

    # 2) 指令辨識：#生成強銷 等歸類為 case_ai，不再混在 seller
    _detect_line_command_type_before_case_form_ai_v4 = detect_line_command_type
    def detect_line_command_type(text: str, event=None) -> str:
        raw = (text or "").strip().replace("＃", "#")
        if raw:
            first = raw.splitlines()[0].strip().replace(" ", "")
            first_no_hash = first[1:] if first.startswith("#") else first
            if first_no_hash in CASE_FORM_AI_LINE_COMMANDS:
                return CASE_FORM_AI_COMMAND_TYPE
        return _detect_line_command_type_before_case_form_ai_v4(text, event=event)

    # 3) 指令權限相容：有勾 case_ai 可用；舊版只勾 seller 也可用
    _line_group_allows_command_before_case_form_ai_v4 = line_group_allows_command
    def line_group_allows_command(group, command_type: str) -> bool:
        if command_type == CASE_FORM_AI_COMMAND_TYPE:
            allowed = set((group or {}).get("command_types") or [])
            return "all" in allowed or CASE_FORM_AI_COMMAND_TYPE in allowed or "seller" in allowed
        return _line_group_allows_command_before_case_form_ai_v4(group, command_type)

    _line_personal_user_allows_command_before_case_form_ai_v4 = line_personal_user_allows_command
    def line_personal_user_allows_command(user, command_type: str) -> bool:
        if command_type == CASE_FORM_AI_COMMAND_TYPE:
            allowed = set((user or {}).get("command_types") or [])
            return "all" in allowed or CASE_FORM_AI_COMMAND_TYPE in allowed or "seller" in allowed
        return _line_personal_user_allows_command_before_case_form_ai_v4(user, command_type)

    # 4) 可查詢資料也要看 seller 權限，避免能生成委託文案但不能看委託資料
    _line_access_gate_before_case_form_ai_v4 = line_access_gate
    def line_access_gate(event):
        allowed, reason_or_cmd, source_cfg = _line_access_gate_before_case_form_ai_v4(event)
        if not allowed:
            return allowed, reason_or_cmd, source_cfg
        if reason_or_cmd == CASE_FORM_AI_COMMAND_TYPE and source_cfg:
            view_types = set(source_cfg.get("view_types") or [])
            if "all" not in view_types and "seller" not in view_types:
                return False, "此來源已開放 AI文案 / 案件表指令，但未開放查看『委託』資料，請到設定中心勾選可查詢資料：委託。", source_cfg
        return allowed, reason_or_cmd, source_cfg

    print("✅ CASE_FORM_AI V4：設定中心已新增 AI文案 / 案件表獨立權限勾選 case_ai")
except Exception as e:
    print("⚠️ CASE_FORM_AI V4 權限勾選修正失敗：", e)


# =============================================================================
# LINE CRM Patch 20260705：
# 1) #新增客需 / #買方追蹤 等欄位支援多行內容，備註/內容換行也會完整寫入。
# 2) LINE 訊息第一行沒有 # 時視為一般聊天，不回覆「沒有可辨識的指令」。
#    注意：此段放在檔案最後，目的是覆蓋前面多個歷史版本的同名函式。
# =============================================================================
def _teamme_first_nonempty_line(text: str) -> str:
    for ln in (text or "").splitlines():
        if ln.strip():
            return ln.strip()
    return ""


def _teamme_parse_key_value_multiline(body_lines):
    """
    解析 LINE #指令的欄位，支援：
      備註: 第一行
      第二行
      第三行
    會解析成 fields['content'] = '第一行\n第二行\n第三行'

    規則：
    - 遇到「欄位: 值」就開新欄位
    - 沒有冒號的行，接到上一個欄位後面
    - 空白行也會保留在上一個欄位中，避免使用者分段備註消失
    """
    fields = {}
    current_key = None

    for raw in body_lines:
        raw = raw.rstrip("\r")
        line = raw.strip()

        m = re.match(r"^([^:：]+)\s*[:：]\s*(.*)$", line) if line else None
        if m:
            key = normalize_line_key(m.group(1))
            value = (m.group(2) or "").strip()
            current_key = key
            if key == "labels":
                fields[key] = parse_label_csv(value)
            else:
                fields[key] = value
            continue

        # 沒有「欄位:」的行，視為上一個欄位的續行。
        if current_key:
            if current_key == "labels":
                old = ensure_list(fields.get(current_key))
                more = parse_label_csv(line)
                fields[current_key] = dedupe_keep_order(old + more)
            else:
                old = fields.get(current_key, "")
                if old:
                    fields[current_key] = old + "\n" + line
                else:
                    fields[current_key] = line

    return fields


def parse_line_formatted_message(text: str):
    raw_text = text or ""
    lines = [ln.rstrip() for ln in raw_text.splitlines()]
    nonempty = [ln.strip() for ln in lines if ln.strip()]
    if not nonempty:
        return None

    first = nonempty[0].replace("＃", "#")
    if not first.startswith("#"):
        return None

    tag = first.lstrip("#").strip()
    tag_map = {
        "新增客需": "create_buyer_need",
        "新增委託": "create_seller_listing",
        "新增開發": "create_development",
        "新增開發批次": "create_development_batch",
        "開發追蹤": "development_followup",
        "買方追蹤": "buyer_followup",
        "賣方追蹤": "seller_followup",
        "客戶分類": "classify",
        "查詢紀錄": "query_records",
        "查詢委託到期": "query_contract_end",
        "帶看": "buyer_followup",
        "成交": "buyer_followup",
        "委託": "seller_followup",
        "紀錄": "generic_note",
    }
    action = tag_map.get(tag)
    if not action:
        return None

    # 取第一個 #標題 後面的所有原始行，保留換行。
    started = False
    body_lines = []
    for ln in lines:
        stripped = ln.strip().replace("＃", "#")
        if not started:
            if stripped == first:
                started = True
            continue
        body_lines.append(ln)
    raw_body = "\n".join(body_lines).strip("\n")

    if action == "create_development":
        fields = _strict_parse_development_fields(raw_body)
        return {"tag": tag, "action": action, "fields": fields, "raw_text": raw_text, "raw_body": raw_body}

    if action == "create_development_batch":
        items = _strict_parse_development_batch(raw_body)
        if not items:
            return None
        return {"tag": tag, "action": action, "fields": {}, "raw_text": raw_text, "raw_body": raw_body}

    fields = _teamme_parse_key_value_multiline(body_lines)

    if tag in ("買方追蹤", "帶看", "成交") and not fields.get("target_type"):
        fields["target_type"] = "buyer"
    if tag in ("賣方追蹤", "委託") and not fields.get("target_type"):
        fields["target_type"] = "seller"
    if tag == "開發追蹤" and not fields.get("target_type"):
        fields["target_type"] = "development"

    fields["target_type"] = normalize_target_type(fields.get("target_type", "")) or fields.get("target_type", "")
    fields["intent_type"] = normalize_intent_type(fields.get("intent_type_raw", ""), fields)
    fields["deal_type"] = normalize_deal_type(fields.get("deal_type_raw", ""))
    fields["limit"] = parse_int_limit(fields.get("limit", 10), default=10, max_value=30)

    if action == "create_buyer_need":
        if not (fields.get("name") and fields.get("phone")):
            return None
    elif action == "create_seller_listing":
        if not (fields.get("name") and fields.get("phone")):
            return None
    elif action in ("buyer_followup", "seller_followup", "classify", "query_records", "query_contract_end", "development_followup"):
        if not (fields.get("record_id") or fields.get("phone") or fields.get("name") or fields.get("address")):
            return None

    return {"tag": tag, "action": action, "fields": fields, "raw_text": raw_text, "raw_body": raw_body}


# 沒有文字就不要呼叫 LINE reply API，避免用空訊息硬回覆。
_teamme_reply_line_text_before_no_empty = reply_line_text
def reply_line_text(reply_token: str, text_message: str):
    if not (text_message or "").strip():
        return None
    return _teamme_reply_line_text_before_no_empty(reply_token, text_message)


_teamme_line_access_gate_before_ignore_non_hash = line_access_gate
def line_access_gate(event):
    if (event.get("message") or {}).get("type") == "text":
        text = ((event.get("message") or {}).get("text") or "").replace("＃", "#")
        first = _teamme_first_nonempty_line(text)
        # 沒有 # 開頭：當成群組一般聊天，不回覆、不處理。
        if first and not first.startswith("#"):
            return False, "", None
    return _teamme_line_access_gate_before_ignore_non_hash(event)

print("✅ LINE CRM Patch 20260705：多行備註保存 + 非 # 訊息靜默忽略 已啟用")
