# -*- coding: utf-8 -*-
"""
Team M.E｜案件輸入表 Word 版（PDF 完全同版型）

這版不是重新畫 Word 表格，而是：
1. 把原本「案件輸入表-使用中.pdf」轉成同版型背景圖
2. 在 Word 裡把背景圖鋪滿 A4
3. 用可編輯的文字框把謄本解析資料填到對應位置

好處：Word 打開後版面會跟 PDF 幾乎一模一樣，不會再因 Word 表格重排跑版。
注意：底圖本身不可編輯，但填入的文字框可在 Word 中點選修改。

安裝：
    pip install python-docx
"""
from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any
import html
import json
import re

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.oxml import parse_xml
    from docx.oxml.ns import qn
except Exception as e:  # pragma: no cover
    raise RuntimeError("缺少 python-docx，請先安裝：pip install python-docx") from e

A4_W_PT = 595.0
A4_H_PT = 842.0
FONT_NAME = "Microsoft JhengHei"


def clean(v: Any) -> str:
    if v is None:
        return ""
    return str(v).strip()


def _module_dir() -> Path:
    return Path(__file__).resolve().parent


def find_template_png(path: str | Path | None = None) -> Path:
    candidates = []
    if path:
        candidates.append(Path(path))
    root = _module_dir()
    candidates.extend([
        root / "assets" / "案件輸入表-使用中.png",
        root / "案件輸入表-使用中.png",
        Path.cwd() / "assets" / "案件輸入表-使用中.png",
    ])
    for p in candidates:
        if p.exists():
            return p
    raise FileNotFoundError("找不到 Word 背景圖：assets/案件輸入表-使用中.png")


def nsdecls_extra():
    return ('xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:v="urn:schemas-microsoft-com:vml" '
            'xmlns:o="urn:schemas-microsoft-com:office:office"')


def _xml_text_run(text: str, size: float = 8, bold: bool = False, color: str = "000000") -> str:
    b = '<w:b/>' if bold else ''
    return (
        f'<w:r><w:rPr>{b}'
        f'<w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:eastAsia="{FONT_NAME}"/>'
        f'<w:sz w:val="{int(size*2)}"/><w:color w:val="{color}"/>'
        f'</w:rPr><w:t xml:space="preserve">{html.escape(text)}</w:t></w:r>'
    )


def add_textbox(paragraph, text: str, x: float, y_top: float, w: float, h: float, size: float = 8,
                name: str = "box", bold: bool = False, color: str = "000000", max_lines: int = 1,
                line_height_pt: float | None = None):
    """在 Word 頁面上以 PDF 點數座標加入可編輯文字框。x/y_top 都是 pt。"""
    text = clean(text)
    if not text:
        return
    line_height_pt = line_height_pt or (size + 1.8)

    # 簡易中文換行：依文字框寬度估算每行可容納字數
    def units(s: str) -> float:
        return sum(1.0 if ord(ch) > 127 else 0.55 for ch in s)

    max_units = max(1, int(w / max(size * 0.82, 1)))
    lines = []
    for raw in text.splitlines():
        raw = raw.strip()
        cur = ""
        for ch in raw:
            if units(cur + ch) > max_units:
                if cur:
                    lines.append(cur)
                cur = ch
            else:
                cur += ch
        if cur:
            lines.append(cur)
    lines = lines[:max_lines]
    paras = []
    for line in lines:
        paras.append(
            '<w:p><w:pPr><w:spacing w:before="0" w:after="0" w:lineRule="exact" '
            f'w:line="{int(line_height_pt*20)}"/></w:pPr>' +
            _xml_text_run(line, size=size, bold=bold, color=color) +
            '</w:p>'
        )
    xml = (
        f'<w:pict {nsdecls_extra()}>'
        f'<v:shape id="{html.escape(name)}" type="#_x0000_t202" '
        f'style="position:absolute;margin-left:{x:.2f}pt;margin-top:{y_top:.2f}pt;'
        f'width:{w:.2f}pt;height:{h:.2f}pt;z-index:251659264;'
        f'mso-position-horizontal-relative:page;mso-position-vertical-relative:page" '
        f'stroked="f" fillcolor="none">'
        f'<v:textbox inset="0,0,0,0"><w:txbxContent>{"".join(paras)}</w:txbxContent></v:textbox>'
        f'</v:shape></w:pict>'
    )
    paragraph.add_run()._r.append(parse_xml(xml))


def add_check(paragraph, x: float, y_top: float, name: str = "check"):
    # 用文字框放大勾，避免 Word/LibreOffice 的符號字型問題。實務開 Word 可再點選移動。
    add_textbox(paragraph, "✓", x-1, y_top-5, 16, 16, size=12, name=name, bold=True, max_lines=1, line_height_pt=13)


TEXT_POS = {
    # 客戶資料
    "owner_name": (45, 121, 88, 10, 7.8),
    "owner_mobile": (418, 147, 88, 10, 7.5),
    "owner_address": (55, 173, 420, 12, 7.2),

    # 1. 基本資料
    "property_title": (88, 213, 125, 11, 7.4),
    "community_name": (500, 213, 72, 11, 7.4),
    "case_address": (80, 331, 410, 12, 7.3),
    "floor_total": (92, 349, 32, 10, 7.2),
    "floor": (92, 365, 42, 10, 7.2),
    "layout": (238, 365, 188, 10, 7.4),
    "completed_year": (130, 381, 24, 10, 7.2),
    "completed_month": (174, 381, 18, 10, 7.2),
    "completed_day": (211, 381, 18, 10, 7.2),
    "building_age": (302, 381, 24, 10, 7.2),
    "facing": (510, 290, 30, 10, 7.2),

    # 3. 面積金額
    "total_ping": (98, 578, 45, 10, 7.1),
    "main_ping": (203, 578, 45, 10, 7.1),
    "attached_ping": (293, 578, 45, 10, 7.1),
    "public_ping": (383, 578, 45, 10, 7.1),
    "parking_ping": (486, 578, 45, 10, 7.1),
    "land_ping": (98, 604, 45, 10, 7.1),
    "base_land_ping": (209, 604, 45, 10, 7.1),
    "land_share_ping": (362, 604, 45, 10, 7.1),
    "case_price": (66, 629, 58, 10, 7.2),
    "rent_price": (208, 629, 58, 10, 7.2),
    "deposit": (330, 629, 55, 10, 7.2),
    "deposit_months": (501, 629, 45, 10, 7.2),

    # 5. 特色備註
    "feature_note": (135, 746, 420, 13, 6.6),
    "special_note": (135, 804, 420, 13, 6.4),
}

CHECK_POS = {
    "deal_sale": (86, 50),
    "deal_rent": (125, 50),
    "mandate_exclusive": (264, 50),
    "mandate_general": (314, 50),
    "source_deed": (302, 97),
    "type_toutian": (189, 231),
    "type_villa": (236, 231),
    "type_store": (329, 231),
    "type_factory": (421, 231),
    "status_empty": (84, 266),
    "status_self_use": (134, 266),
    "status_rented": (184, 266),
    "structure_rc": (247, 441),
    "use_residential": (132, 542),
    "use_store": (190, 542),
    "use_parking": (302, 542),
    "use_factory": (377, 542),
    "use_commercial": (461, 542),
    "use_office": (530, 542),
}


def infer_checks(case_data: dict[str, Any], seller: dict[str, Any] | None = None) -> dict[str, bool]:
    seller = seller or {}
    checks: dict[str, bool] = {}
    deal = clean(seller.get("deal_type") or case_data.get("deal_type") or "sale").lower()
    checks["deal_rent"] = deal in {"rent", "出租", "租"}
    checks["deal_sale"] = not checks["deal_rent"]
    checks["source_deed"] = bool(case_data.get("deed_parsed_note") or case_data.get("deed_raw_text"))

    ptype = clean(seller.get("property_type") or case_data.get("property_type"))
    main_use = clean(case_data.get("deed_main_use"))
    floor_total = clean(case_data.get("floor_total"))
    material = clean(case_data.get("deed_main_material"))

    if "透天" in ptype or ("住宅" in main_use and floor_total.isdigit() and int(floor_total) <= 5):
        checks["type_toutian"] = True
    if "別墅" in ptype:
        checks["type_villa"] = True
    if "店" in ptype or "店" in main_use:
        checks["type_store"] = True
    if "廠" in ptype or "工業" in main_use:
        checks["type_factory"] = True

    if "自用" in clean(seller.get("occupancy_status")):
        checks["status_self_use"] = True
    elif "出租" in clean(seller.get("occupancy_status")):
        checks["status_rented"] = True
    elif "空" in clean(seller.get("occupancy_status")):
        checks["status_empty"] = True

    if "鋼筋混凝土" in material or "RC" in material.upper():
        checks["structure_rc"] = True

    if "住宅" in main_use or "住家" in main_use:
        checks["use_residential"] = True
    if "車庫" in main_use or "停車" in main_use:
        checks["use_parking"] = True
    if "店" in main_use:
        checks["use_store"] = True
    if "工業" in main_use or "廠" in main_use:
        checks["use_factory"] = True
    if "商業" in main_use:
        checks["use_commercial"] = True
    if "辦公" in main_use:
        checks["use_office"] = True
    return checks


def build_fields(case_data: dict[str, Any], seller: dict[str, Any] | None = None) -> dict[str, str]:
    seller = seller or {}
    fields: dict[str, str] = {}

    def put(key: str, *values):
        for v in values:
            v = clean(v)
            if v:
                fields[key] = v
                return

    put("owner_name", seller.get("name"))
    put("owner_mobile", seller.get("phone"))
    put("owner_address", seller.get("contact_address"), seller.get("address"), case_data.get("case_address"))
    put("property_title", case_data.get("property_title"), case_data.get("ai_sales_title"))
    put("community_name", case_data.get("community_name"))
    put("case_address", case_data.get("case_address"), seller.get("address"))
    put("floor_total", case_data.get("floor_total"))
    put("floor", case_data.get("floor"))
    put("layout", case_data.get("layout"))
    put("completed_year", case_data.get("completed_minguo_year"))
    put("completed_month", case_data.get("completed_month"))
    put("completed_day", case_data.get("completed_day"))
    put("building_age", case_data.get("building_age"))
    put("facing", case_data.get("facing"))
    put("total_ping", case_data.get("total_ping"))
    put("main_ping", case_data.get("main_ping"))
    put("attached_ping", case_data.get("attached_ping"))
    put("public_ping", case_data.get("public_ping"))
    put("parking_ping", case_data.get("parking_ping"))
    put("land_ping", case_data.get("land_ping"))
    put("base_land_ping", case_data.get("land_ping"))
    put("land_share_ping", case_data.get("land_ping"))
    put("case_price", case_data.get("case_price"), seller.get("expected_price"))
    put("rent_price", case_data.get("rent_price"))
    put("deposit", case_data.get("deposit"))
    put("deposit_months", case_data.get("deposit_months"))

    feature_parts = []
    for k in ("ai_feature_note", "property_highlight_note", "life_note", "target_customer_note"):
        if clean(case_data.get(k)):
            feature_parts.append(clean(case_data.get(k)))
    feature_text = " ".join(feature_parts) or clean(case_data.get("deed_parsed_note"))
    put("feature_note", feature_text[:110])

    raw_special = clean(case_data.get("deed_mortgage_note") or case_data.get("case_note"))
    special = raw_special
    # 表格這一欄只有一條線，過長會壓到合約日；這裡自動濃縮成可放入表格的一句話。
    if raw_special:
        parts = []
        if "禁止處分" in raw_special:
            parts.append("有禁止處分登記，須確認塗銷/移轉")
        if "抵押權" in raw_special or "最高限額" in raw_special:
            creditors = []
            for name in ["玉山", "中租", "台新", "國泰", "中信", "聯邦", "土地銀行", "合作金庫"]:
                if name in raw_special:
                    creditors.append(name)
            if creditors:
                parts.append("有最高限額抵押權：" + "、".join(dict.fromkeys(creditors)))
            else:
                parts.append("有他項權利/抵押權須確認")
        if parts:
            special = "；".join(parts)
    put("special_note", special[:120])
    return fields


def build_case_form_docx(case_data: dict[str, Any], seller: dict[str, Any] | None = None,
                         template_png: str | Path | None = None) -> Document:
    seller = seller or {}
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(0)
    sec.bottom_margin = Cm(0)
    sec.left_margin = Cm(0)
    sec.right_margin = Cm(0)

    # Normal style
    try:
        style = doc.styles["Normal"]
        style.font.name = FONT_NAME
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        style.font.size = Pt(8)
    except Exception:
        pass

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    # 背景：原 PDF 表格轉成的 PNG，鋪滿 A4，所以格式和 PDF 幾乎一模一樣。
    bg = find_template_png(template_png)
    run = p.add_run()
    run.add_picture(str(bg), width=Cm(21.0), height=Cm(29.7))

    fields = build_fields(case_data, seller)
    for key, value in fields.items():
        if key not in TEXT_POS:
            continue
        x, y, w, h, size = TEXT_POS[key]
        max_lines = 1
        if key == "feature_note":
            max_lines = 1
        elif key == "special_note":
            max_lines = 1
        add_textbox(p, value, x, y, w, h, size=size, name=f"tm_{key}", max_lines=max_lines)

    checks = infer_checks(case_data, seller)
    for key, enabled in checks.items():
        if enabled and key in CHECK_POS:
            x, y = CHECK_POS[key]
            add_check(p, x, y, name=f"tm_check_{key}")

    return doc


def build_case_form_docx_bytes(case_data: dict[str, Any], seller: dict[str, Any] | None = None,
                               template_png: str | Path | None = None) -> bytes:
    doc = build_case_form_docx(case_data, seller=seller, template_png=template_png)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def save_case_form_docx(case_data: dict[str, Any], output_path: str | Path,
                        seller: dict[str, Any] | None = None,
                        template_png: str | Path | None = None) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_bytes(build_case_form_docx_bytes(case_data, seller=seller, template_png=template_png))
    return output_path


def main(argv=None):
    import argparse
    parser = argparse.ArgumentParser(description="產生與 PDF 版面一致的 Word 案件輸入表")
    parser.add_argument("json_file", help="deed_case_form_tool.py 解析出的 JSON")
    parser.add_argument("output_docx")
    parser.add_argument("--template-png", default="")
    parser.add_argument("--seller-name", default="")
    parser.add_argument("--seller-phone", default="")
    parser.add_argument("--price", default="")
    parser.add_argument("--property-title", default="")
    parser.add_argument("--layout", default="")
    args = parser.parse_args(argv)

    parsed = json.loads(Path(args.json_file).read_text(encoding="utf-8"))
    case_data = parsed.get("case_data") or parsed
    if args.price:
        case_data["case_price"] = args.price
    if args.property_title:
        case_data["property_title"] = args.property_title
    if args.layout:
        case_data["layout"] = args.layout
    seller = {"name": args.seller_name, "phone": args.seller_phone, "deal_type": "sale", "property_type": "透天"}
    save_case_form_docx(case_data, args.output_docx, seller=seller, template_png=args.template_png or None)
    print(f"已產生：{args.output_docx}")


if __name__ == "__main__":
    main()
